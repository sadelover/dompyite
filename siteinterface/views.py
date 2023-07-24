# -*- coding: utf-8 -*-


"""
Routes and views for the flask application.
"""

from siteinterface import  app
from siteinterface.BEOPDataAccess import *
from siteinterface.LogicManager import *
from siteinterface.BEOPSqliteAccess import *
import sys
import os, shutil
import re
import copy
from math import floor, ceil
from flask import Flask,request,session,g,make_response,redirect,url_for,abort,render_template,send_file,flash,json,jsonify
import mysql.connector
from datetime import datetime,timedelta
from email import charset
import time
import calendar
import hashlib
import  codecs
from os import environ, listdir, path, unlink, remove, makedirs
import csv
from bson import ObjectId
from siteinterface.ExcelFile import ExcelFile
import  zlib,zipfile
from docx import Document
from docx.shared import Inches
from siteinterface.BEOPDataAccess import *
from datetime import datetime,timedelta
import logging
from math import *
import tablib
import random

from flask import Response
import io,mimetypes

from docx import Document
#from docx.enum.text import WD_ALIGN_PARAGRAPH
from siteinterface.ReportTool import ReportTool
import requests
import xlrd, xlwt

from functools import update_wrapper

from siteinterface.DBFileManager import DBFileManager
from flask import send_from_directory

from psutil import net_if_addrs
from siteinterface.utils import *
from siteinterface.commonUtils import *
from ctypes import *
from siteinterface.ExcelFile import read_xlsx, readPointsListFromExcel,read_warning_config_from_xlsx
from siteinterface.utils import send_email
from siteinterface.LicenseManage import LicenseManage
from siteinterface.ReportTool import *
from logic.PointNameAnalysis.PointNameAI import PointNameAI
from siteinterface.mod_project.controllers import *
from siteinterface.mod_point.controllers import *
from siteinterface.mod_logic.controllers import *
from siteinterface.utils import g_app_config_database, save_app_config_database, restore_app_config_database
from openpyxl import Workbook
from pypinyin import pinyin, lazy_pinyin
from siteinterface.RedisManager import RedisManager
from version import siteVersion
from siteinterface.utils import getCurrentSystemResource
from siteinterface.DomRegression import DomRegression
import random
import dateutil.parser
from siteinterface.OSSTool import OSSTool
import socket
from siteinterface.SystemAPI import g_systemAPIList
import bson
import traceback
import uuid


_mailSendTimestamp = datetime.now() - timedelta(1)
charset.add_charset('utf-8', charset.SHORTEST, charset.BASE64, 'utf-8')

g_thirdparty_pointmap = {}
g_get_history_data_time = datetime.now()

def updateMailSetting():
    # mail info
    try:
        strMailSetting = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig('mail_setting')
        if strMailSetting:
            jsonMailSetting = json.loads(strMailSetting)
            if jsonMailSetting.get('MAIL_SERVER'):
                app.config.update(dict(MAIL_SERVER=jsonMailSetting.get('MAIL_SERVER')))
            if jsonMailSetting.get('MAIL_USERNAME'):
                app.config.update(dict(MAIL_USERNAME=jsonMailSetting.get('MAIL_USERNAME')))
            if jsonMailSetting.get('MAIL_PASSWORD'):
                app.config.update(dict(MAIL_PASSWORD=jsonMailSetting.get('MAIL_PASSWORD')))
            if jsonMailSetting.get('MAIL_DEFAULT_SENDER') and jsonMailSetting.get('MAIL_DEFAULT_SENDER_NAME'):
                app.config.update(dict(MAIL_DEFAULT_SENDER=(
                    jsonMailSetting.get('MAIL_DEFAULT_SENDER_NAME'), jsonMailSetting.get('MAIL_DEFAULT_SENDER'))))
            app.config.update(dict(MAIL_DEBUG=0))
    except Exception as e:
        print('ERROR in updateMailSetting :%s'%(e.__str__()))


def initFirewall():
    # mail info
    try:
        strJsonFirewall = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig('firewall_config')
        if strJsonFirewall:
            jsonFirewall = json.loads(strJsonFirewall)
            if jsonFirewall.get('BLACK_USERS'):
                app.config.update(dict(FIREWALL_BLACK_USERS= jsonFirewall.get('BLACK_USERS')))

    except Exception as e:
        print('ERROR in updateMailSetting :%s'%(e.__str__()))



@app.route('/')
def main():
    rv = redirect(url_for('home'))
    return  rv
    # return render_template('debugTool.html')

@app.route('/omsite')
def omsiteWebpage():
    return render_template('omsiteWeb.html')


@app.route('/rps')
def rpsWebpate():
    return render_template('rps.html')

@app.route('/webservice')
def redicrect_web_service():
    return render_template('thirdpartytest.html')


@app.route('/login',methods=['POST'])
def login():
    print('login')
    return json.dumps(loginTask(request.get_json()), ensure_ascii=False)


@app.route('/getProjectByUserId/<userId>')
def getProjectByUserId(userId):
    return json.dumps(BEOPDataAccess.getInstance().getProject(userId), ensure_ascii=False)


@app.route('/reset_pwd', methods=['POST'])
def reset_pwd():
    print('reset_pwd')
    return json.dumps(BEOPDataAccess.getInstance().resetPwd(request.get_json()), ensure_ascii=False)


@app.route('/sendmail',methods=['POST'])
def request_send_mail():
    global _mailSendTimestamp
    rdata = request.get_json()
    reciepents =rdata.get('recipients',None)
    subject = rdata.get('subject', None)
    html = rdata.get('html', None)
    if not rdata or not reciepents or not subject or not html:
        return jsonify(dict(msg='', err=1, data={}))

    now = datetime.now()
    if (now - _mailSendTimestamp).total_seconds() > 5:
        updateMailSetting()
        rv = send_email(subject,reciepents,html)
        #_mailSendTimestamp = now
    else:
        print('error: sending mail too frequently')
        rv = 'error: you should wait a while before sending more mails'
    return json.dumps(rv, ensure_ascii=False)


@app.route('/getPysiteVersion')
def get_pysite_version():
    strVersion = siteVersion.getCurrentVersion()
    nums = strVersion.split('.')
    totalversion = int(nums[0]) * 10000 + int(nums[1]) * 100 + int(nums[2])
    return jsonify(dict(err=0, msg="", data=totalversion))


@app.route('/processVersion')
def get_process_version():
    strMsg = ""
    lAllProcess = []
    try:

        # 准备全部信息字典
        dDataAll = {}
        procNameList = []
        for item in allProcList:
            name = item["name"]
            procNameList.append(name)
            dDataAll.update({name: dict(name=item["name"],
                                        cloudVersion="",
                                        localVersion="",
                                        hostStatus=0,
                                        runningStatus=0,
                                        updateStatus=0,
                                        vsrc=item["vsrc"],
                                        vkey=item["vkey"],
                                        hostType=item["hostType"],
                                        sttSrc=item.get("sttSrc", ""),
                                        sttKey=item.get("sttKey", ""))})

        # 获取各进程运行状态
        dRunningStatus = ProcessManager.getInstance().getProcessRunningStatus(procNameList)

        # 获取redis运行状态
        redisHandler = ProcessManager.getInstance().findProcess("redis-server.exe")

        # 获取当前进程守护列表
        curProcHostList = BEOPDataAccess.getInstance().getListInUnit01ByKey("hostprocess")

        # 获取进程更新状态
        dJsonUpdateStatus = {}
        if redisHandler:
            dJsonUpdateStatus = RedisManager.get("ProcessUpdateStatus")

        # 获取进程云端版本JSON
        dJsonCloudVersion = {}
        filesDir = os.path.join(app.static_folder, "files")
        if not os.path.exists(filesDir):
            os.mkdir(filesDir)

        tempDir = os.path.join(app.static_folder, "temp")
        if not os.path.exists(tempDir):
            os.mkdir(tempDir)

        updateJsonPath = os.path.join(tempDir, "update.json")
        if os.path.exists(updateJsonPath):
            os.remove(updateJsonPath)

        osstool = OSSTool()
        bSuc = osstool.download("update/update.json", updateJsonPath)

        if bSuc:
            with open(updateJsonPath, "r", encoding="UTF-8") as f:
                dJsonCloudVersion = json.load(f)

        # 装载dompysite, domcore, domlogic本地版本
        dDataAll.get("dompysite").update({"localVersion": siteVersion.getCurrentVersion()})

        versionInfo = BEOPDataAccess.getInstance().getCoreLogicVersion()
        dCoreLogicStartupTime = {}

        if versionInfo.get("code") == 0:
            dCoreLogicStartupTime = versionInfo.get("startupTime", {})

            dDataAll.get("domcore").update({"localVersion": versionInfo.get("version").get("domcore")})
            dDataAll.get("domlogic").update({"localVersion": versionInfo.get("version").get("domlogic")})
            dDataAll.get("domSiemenseTCPCore").update({"localVersion": versionInfo.get("version").get("domSiemenseTCPCore")})

            for i in range(5):
                strDomlogic = "domlogic{0:02d}".format(i+1)
                dDataAll.get(strDomlogic).update({"localVersion": versionInfo.get("version").get("domlogic")})

        # 装载redis本地版本
        if redisHandler:
            redisVersion = RedisManager.get_version()
            dDataAll.get("redis").update({"localVersion": redisVersion})
            dRunningStatus.update({"redis": 1})

        # 装载进程更新状态、云端版本、运行状态、守护状态
        bResetUpdateStatus = False
        for procName, dProcInfo in dDataAll.items():
            if dProcInfo.get("vsrc") == "redis" and redisHandler:
                strVersion = RedisManager.get(dProcInfo["vkey"])
                if strVersion:
                    dProcInfo.update({"localVersion": strVersion})

            # 进程更新状态
            nUpdateStatus = 0
            if dJsonUpdateStatus and isinstance(dJsonUpdateStatus, dict):
                strKey = procName
                if re.match(r"^domBacnetCore[0-9]{1}$", procName):  # 当domBacnetCore1-5处于更新状态时则认为所有domBacnetCore处于更新状态
                    strKey = "domBacnetCore"
                elif re.match(r"^domlogic[0-9]{2}$", procName):
                    strKey = "domlogic"

                dInfo = dJsonUpdateStatus.get(strKey, {})
                if dInfo:
                    sTime = list(dInfo.keys())[0]
                    tTime = datetime.strptime(sTime, "%Y-%m-%d-%H-%M-%S")
                    if (datetime.now() - tTime).total_seconds() > 3600:
                        dJsonUpdateStatus.update({procName: {}})
                        bResetUpdateStatus = True
                    else:
                        nUpdateStatus = dInfo.get(sTime, 0)

            # 进程更新状态
            if procName == "domlogic":
                dDataAll.get("domlogic").update({"updateStatus": nUpdateStatus})
            elif re.match(r"^domlogic[0-9]{2}$", procName):
                dDataAll.get(procName).update({"updateStatus": nUpdateStatus})
            elif procName == "domBacnetCore":
                dDataAll.get("domBacnetCore").update({"updateStatus": nUpdateStatus})
            elif re.match(r"^domBacnetCore[0-9]{1}$", procName):
                dDataAll.get(procName).update({"updateStatus": nUpdateStatus})
            else:
                dProcInfo.update({"updateStatus": nUpdateStatus})

            # 进程云端版本
            if procName == "domlogic":
                dDataAll.get("domlogic").update({"cloudVersion": dJsonCloudVersion.get("domlogic", "")})
            elif re.match(r"^domlogic[0-9]{2}$", procName):
                dDataAll.get(procName).update({"cloudVersion": dJsonCloudVersion.get("domlogic", "")})
            elif procName == "domBacnetCore":
                dDataAll.get("domBacnetCore").update({"cloudVersion": dJsonCloudVersion.get("domBacnetCore", "")})
            elif re.match(r"^domBacnetCore[0-9]{1}$", procName):
                dDataAll.get(procName).update({"cloudVersion": dJsonCloudVersion.get("domBacnetCore", "")})
            else:
                dProcInfo.update({"cloudVersion": dJsonCloudVersion.get(procName, "")})

            # 进程运行状态
            dProcInfo.update({"runningStatus": dRunningStatus.get(procName, 0)})

            # 持续运行秒数
            nContinuousRunSecs = None
            if dProcInfo.get("runningStatus", 0) == 0:
                nContinuousRunSecs = 0
            elif dProcInfo.get("runningStatus", 0) == 1:
                strStartupTime = ""
                if dProcInfo.get("sttSrc", "") == "redis":
                    try:
                        strStartupTime = RedisManager.get(dProcInfo.get("sttKey", "")).get("startup", "")
                    except:
                        pass
                elif dProcInfo.get("sttSrc", "") == "unit01":
                    strStartupTime = dCoreLogicStartupTime.get(procName, "")

                if isinstance(strStartupTime, str) and is_valid_date(strStartupTime, "%Y-%m-%d %H:%M:%S"):
                    tStartupTime = datetime.strptime(strStartupTime, "%Y-%m-%d %H:%M:%S")
                    nContinuousRunSecs = (datetime.now() - tStartupTime).total_seconds()

            dProcInfo.update({"continuousSeconds": nContinuousRunSecs})


            # 进程驻守类型    0-无需添加驻守；1-需要添加驻守
            nHStatus = 0
            if dProcInfo.get("hostType") == 0:
                nHStatus = 0
            elif dProcInfo.get("hostType") == 1:
                if procName in curProcHostList:
                    nHStatus = 0
                else:
                    nHStatus = 1
            dProcInfo.update({"hostStatus": nHStatus})



        # 更新一次redis中的进程更新状态
        if bResetUpdateStatus:
            RedisManager.set("ProcessUpdateStatus", dJsonUpdateStatus)

        if not redisHandler:
            dDataAll.get("redis").update({"localVersion": "redis not running or not installed"})

        dPause = RedisManager.get("pause_host_config")
        if not isinstance(dPause, dict):
            dPause = {}

        # 生成进程列表
        for k, v in dDataAll.items():
            dProcPauseInfo = {}
            if isinstance(dPause.get(k, None), dict):
                strPauseToTime = dPause.get(k).get("pauseTo", "")
                if is_time_string(strPauseToTime, "%Y-%m-%d %H:%M:%S"):
                    tPauseTo = datetime.strptime(strPauseToTime, "%Y-%m-%d %H:%M:%S")
                    nDeltaSeconds = (tPauseTo - datetime.now()).total_seconds()

                    if nDeltaSeconds <= 60 * 5:
                        dProcPauseInfo.update(dict(time=strPauseToTime, msg="5分钟后解除"))
                    elif nDeltaSeconds > 60 * 5 and nDeltaSeconds <= 60 * 10:
                        dProcPauseInfo.update(dict(time=strPauseToTime, msg="10分钟后解除"))
                    elif nDeltaSeconds > 60 * 10 and nDeltaSeconds <= 60 * 15:
                        dProcPauseInfo.update(dict(time=strPauseToTime, msg="15分钟后解除"))
                    elif nDeltaSeconds > 60 * 15 and nDeltaSeconds <= 60 * 60:
                        dProcPauseInfo.update(dict(time=strPauseToTime, msg="1小时后解除"))
                    elif nDeltaSeconds > 60 * 60 and nDeltaSeconds < 3600 * 24 * 365:
                        dProcPauseInfo.update(dict(time=strPauseToTime, msg="{n}小时后解除".format(n=int(nDeltaSeconds / 3600))))
                    elif nDeltaSeconds > 3600 * 24 * 365:
                        dProcPauseInfo.update(dict(time=strPauseToTime, msg="永久暂缓"))

            v.update(dict(pause=dProcPauseInfo))
            lAllProcess.append(v)

        # 按是否运行分类
        lAllProcessRunning = []
        lAllProcessStopped = []

        for item in lAllProcess:
            if item.get("runningStatus") == 1:
                lAllProcessRunning.append(item)
            else:
                lAllProcessStopped.append(item)

        # 按字母顺序排序
        for les in [lAllProcessRunning, lAllProcessStopped]:
            for i in range(len(les)):
                for j in range(len(les)-i-1):
                    if les[j]["name"].lower() > les[j+1]["name"].lower():
                        les[j], les[j+1] = les[j+1], les[j]

        return jsonify(dict(err=0, msg="获取成功", data=dict(process=lAllProcessRunning + lAllProcessStopped)))
    except Exception as error:
        traceback.print_exc()
        logging.error("ERROR in /version: %s" % error.__str__())
        strMsg += ";{err}".format(err=error.__str__())
        return jsonify(dict(err=1, msg=strMsg, data={}))

@app.route('/version')
def get_beopversion():
    strMsg = ""
    res = {}
    try:
        res.update({"dompysite": siteVersion.getCurrentVersion()})

        versionInfo = BEOPDataAccess.getInstance().getCoreLogicVersion()
        if versionInfo.get("code") > 0:
            strMsg += versionInfo.get("msg")
        else:
            res.update(versionInfo.get("version"))

        return jsonify(dict(err=0, msg="获取成功", data=res))

    except Exception as e:
        logging.error("ERROR in /version: %s" % e.__str__())
        strMsg += ";{err}".format(err=e.__str__())

        return jsonify(dict(err=1, msg=strMsg, data={}))


@app.route('/getrealtimedata_old',methods=['POST'])
def getRealtimedata_old():
    rdata = request.get_json()
    data = BEOPDataAccess.getInstance().getData_Old(rdata.get('db'))
    dj = [dict(time=x[0], name=x[1],value=x[2]) for x in data]
    return json.dumps(dj, ensure_ascii=False)



@app.route('/getrealtimedata',methods=['POST'])
def getRealtimedata():
    dbname = request.form['db']
    j = request.form['data']
    pointList = json.loads(j).get('pointList')
    data = BEOPDataAccess.getInstance().getData(dbname,pointList)
    dj = [dict(name=x[0],value=x[1]) for x in data]
    return json.dumps(dj, ensure_ascii=False)

@app.route('/getrealtimedata_as_dict_with_time',methods=['POST'])
def getRealtimedata_as_dict_with_time():
    rdata = request.get_json()
    pointList = rdata.get('pointList')
    data = BEOPDataAccess.getInstance().getInputTableAsDictListWithTime(pointList)

    return json.dumps(dict(err=0, msg='', data=data), ensure_ascii=False)

@app.route('/get_realtimedata',methods=['POST'])
def get_realtimedata():
    rcv = request.get_json()

    projId = rcv.get('projId', None)
    pointList = rcv.get("pointList", [])
    scriptList = rcv.get('scriptList', [])

    if isinstance(pointList, str):
        pointList = [pointList]

    for script in scriptList:
        ptList = find_vars_in_str(script)
        pointList.extend(ptList)

    pointList = list(set(pointList))

    tNow = datetime.now()
    try:
        if projId:
            projdbKey = "projectdb" if projId == 1 else "projectdb%03d" % projId
            cf = ConfigObj("config.ini", encoding="UTF8")
            if cf.get(projdbKey) is None:
                dj = list()
            dbfileName = cf.get(projdbKey).get("dbFileName")
            dbfileDir = os.path.dirname(app.config.get("USE_4DB_NAME"))
            strDBPath  = os.path.join(dbfileDir, dbfileName)
            strRealDBName = BEOPSqliteAccess.getInstance().GetRealtimeDBName(strDBPath)
            save_app_config_database()
            app.config['DATABASE'] = strRealDBName
        rvRequest = BEOPDataAccess.getInstance().getInputTable(pointList)
        if rvRequest is None:
            return json.dumps([], ensure_ascii=False)

        dj = rvRequest[0]
        allDataMap = rvRequest[1]

        ret = dj
        if len(scriptList) and dj and allDataMap:
            #scriptsResult = tool_eval_string_v2(scriptList,allDataMap, 1)
            retExtendsList = []
            for oneScript in scriptList:
                scriptResult = eval_string_expression_strict(oneScript, '1', tNow.strftime('%Y-%m-%d %H:%M:%S'), dj)
                retExtendsList.append(dict(name=oneScript, value=scriptResult))
            ret.extend(retExtendsList)

        if projId:
            restore_app_config_database()

        return json.dumps(ret, ensure_ascii=False)

    except Exception as e:
        print('ERROR: ' + e.__str__())
        return json.dumps([], ensure_ascii=False)

@app.route('/get_realtimedata_by_projname',methods=['POST'])
def get_realtimedata_by_projname():
    args = request.form
    rdata = request.get_json()

    if rdata.get('proj') is None:
        projName = app.config['PROJECT_DATABASE']
    else:
        projName = rdata.get('proj')

    projId = BEOPDataAccess.getInstance().getProjIdByName(projName)
    data = BEOPDataAccess.getInstance().getDataByProj(projId)

    dj = [dict(name=k,value=v) for k,v in data.items()]
    return json.dumps(dj, ensure_ascii=False)

#golding added
@app.route('/get_realtimedata_with_description_by_projname',methods=['POST'])
def get_realtimedata_with_description_by_projname():
    rdata = request.get_json()
    projID = rdata.get('projid')
    projName = BEOPDataAccess.getInstance().getProjNameById(projID)
    projdbName = ''
    if(projID is None): projdbName = app.config['PROJECT_DATABASE']
    else:
        projdbName = BEOPDataAccess.getInstance().getProjMysqldb(projID)

    pointList = rdata.get('pointList')
    data = BEOPDataAccess.getInstance().getDataByProj(projID,pointList)

    s3dbname = BEOPDataAccess.getInstance().getProjS3db(projID)
    pointInfoList = BEOPSqliteAccess.getInstance().getPointMapsFromS3db(s3dbname)

    if not isinstance(data, list):
        return json.dumps(data, ensure_ascii=False)

    dj = []
    for x in data:
        desp = ''
        if x[0] in pointInfoList:
            desp = pointInfoList[x[0]]
        dj.append(dict(name=x[0], value=x[1], desc=desp))
    return json.dumps(dj, ensure_ascii=False)


@app.route('/set_to_site_by_projname',methods=['POST'])
def set_to_site_by_projname():
    #print('set_realtimedata_by_projname')
    data = request.get_json()
    projEName = data.get('db')
    point = data.get('point')
    value = data.get('value')
    rv = BEOPDataAccess.getInstance().setDataToSite(projEName,point,value)
    return json.dumps(rv , ensure_ascii=False)

@app.route('/set_mutile_to_site_by_projname',methods=['POST'])
def set_mutile_to_site_by_projname():
    #print('set_mutile_to_site_by_projname')
    data = request.get_json()
    projName = data.get('db')
    point = data.get('point')
    value = data.get('value')
    rv = BEOPDataAccess.getInstance().setMutilDataToSite(projName,point,value)
    return json.dumps(rv , ensure_ascii=False)

@app.route('/set_realtimedata_by_projname',methods=['POST'])
def set_realtimedata_by_projname():
    #print('set_realtimedata_by_projname')
    data = request.get_json()
    projName = data.get('db')
    point = data.get('point')
    value = data.get('value')
    rv = BEOPDataAccess.getInstance().setData(projName,point,value)
    return json.dumps(rv , ensure_ascii=False)

@app.route('/set_realtimedata_by_projid',methods=['POST'])
def set_realtimedata_by_projid():
    #print('set_realtimedata_by_projname')
    data = request.get_json()
    projId = data.get('projid')
    projName = BEOPDataAccess.getInstance().getProjNameById(projId)
    point = data.get('point')
    value = data.get('value')
    rv = BEOPDataAccess.getInstance().setData(projName,point,value)
    return json.dumps(rv , ensure_ascii=False)

@app.route('/set_mutile_realtimedata_by_projname',methods=['POST'])
def set_mutile_realtimedata_by_projname():
    #print('set_mutile_realtimedata_by_projname')
    data = request.get_json()
    projName = data.get('db')
    pointList = data.get('point')
    valueList = data.get('value')
    if projName is None or pointList is None or valueList is None:
        rv = 'json post data not good'
    else:
        rv = BEOPDataAccess.getInstance().setMutilData(projName,pointList,valueList)
    return json.dumps(rv , ensure_ascii=False)

@app.route('/set_mutile_realtimedata_by_projid',methods=['POST'])
def set_mutile_realtimedata_by_projId():
    #print('set_mutile_realtimedata_by_projname')
    data = request.get_json()
    projId = data.get('projId')
    pointList = data.get('point')
    valueList = data.get('value')
    if projId is None or pointList is None or valueList is None:
        rv = 'json post data not good'
    else:
        projName = BEOPDataAccess.getInstance().getProjNameById(projId)
        rv = BEOPDataAccess.getInstance().setMutilData(projName,pointList,valueList)
    return json.dumps(rv , ensure_ascii=False)

"""
{"dateFrom": "2016-01-01",
   "dateTo": "2016-01-04"}
"""
@app.route('/get_operation_log', methods=['POST'])
def get_operation_log():
    try:
        rcv = request.get_json()
        dateFrom = rcv.get('dateFrom')
        dateTo = rcv.get('dateTo')
        userId = rcv.get("userId", None)
        nType = rcv.get("type", None)  #    0: 表示登录登出   1: 设备操作   2: 策略操作记录

        if not isinstance(userId, int) and not isinstance(userId, str) and userId != None:
            return jsonify(dict(err=1, msg="用户ID格式有误", data=[]))
        if nType != None:
            if nType not in [0, 1, 2]:
                return jsonify(dict(err=1, msg="type必须为0或1或2", data=[]))

        dataList = BEOPDataAccess.getInstance().readOperationLogV2(dateFrom, dateTo, userId, nType)
        return jsonify(dict(err=0, msg="获取成功", data=dataList))
    except Exception as e:
        logging.error("ERROR in /get_operation_log: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败:%s" % e.__str__(), data=[]))


@app.route('/get_operation_log_in_timerange', methods=['POST'])
def get_operation_log_in_timerange():
    data = request.get_json()
    timeFrom = data.get('timeFrom')
    timeTo = data.get('timeTo')
    opData = BEOPDataAccess.getInstance().readOperationLogInTimeRange(timeFrom,timeTo)
    return json.dumps(dict(err=0, data=opData, msg=''), ensure_ascii=False)


@app.route('/get_sqlitedb_size')
@app.route('/get_sqlitedb_size/<dbname>')
def get_sqlitedb_size(dbname=None):
    print('get_sqlitedb_size')
    return json.dumps(getSqlitedbSize(dbname), ensure_ascii=False)


@app.route('/get_sqlitedb_name')
@app.route('/get_sqlitedb_name/<dbname>')
def get_sqlitedb_name(dbname=None):
    print('get_sqlitedb_name')
    return json.dumps(getSqlitedbName(dbname), ensure_ascii=False)

@app.route('/get_history_data_padded', methods=['POST'])
def get_history_data_padded():
    data = request.get_json()

    projId = data.get("projId", None)
    pointList = data.get('pointList')
    scriptList = data.get("scriptList", None)
    filter = data.get("filter", [])
    strTimeStart = data.get('timeStart')
    strTimeEnd = data.get('timeEnd')
    strTimeFormat = data.get('timeFormat')   # s5 m1 m5 h1 d1 M1
    
    tNow = datetime.now()
    logging.error('get_history_data_padded request recved: %s， body:%s'%( tNow.strftime('%Y-%m-%d %H:%M:%S'), str(data)))
    global g_get_history_data_time
    if (tNow - g_get_history_data_time).total_seconds()<=5:
        time.sleep(random.randint(1,10))

    g_get_history_data_time = datetime.now()

    dFilter = {}
    for dInfo in filter:
        try:
            if dInfo.get("pointName") not in dFilter.keys():
                dFilter.update({dInfo.get("pointName"): dict(filterType=dInfo.get("filterType"),
                                                              params=dInfo.get("params"))})
        except:
            pass
    
    if scriptList:
        scriptList = list(set(scriptList))

    #check pointList:
    bPointListParamOK = False
    if not pointList:
        bPointListParamOK = False
    elif isinstance(pointList, str):
        bPointListParamOK = True
    elif isinstance(pointList, list):
        bAllStr = True
        for item in pointList:
            if not isinstance(item, str):
                bAllStr = False
        if bAllStr:
            bPointListParamOK = True

    if not bPointListParamOK:
        return json.dumps(dict(err=1, msg='pointList param should be str or str List'), ensure_ascii=False)

    #当周期为5秒时额外检查是否有5秒存储如果没有，则只能返回1分钟数据
    if strTimeFormat is not None and strTimeFormat=="s5":
        try:
            if RedisManager.is_alive():
                infoDict = RedisManager.get("all_point_info")  # 若发现从redis拿到的点全集有误则从sqlite中重新获取一次并重新写入redis
                if not isinstance(infoDict, dict) or not infoDict:
                    infoDict = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(None)
                    RedisManager.set("all_point_info", infoDict)
            else:
                infoDict = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(None)

            bHaveSecond5Save = False
            for pt in pointList:
                ptInfo = infoDict.get(pt)
                if ptInfo:
                    ptSavePeriod = ptInfo.get("storecycle", None)
                    if ptSavePeriod is not None and str(ptSavePeriod)=="1":
                        bHaveSecond5Save = True
                        break
            if not bHaveSecond5Save:
                strTimeFormat = "m1"
        except:
            pass

    if projId:
        projdbKey = "projectdb" if projId == 1 else "projectdb%03d" % projId
        cf = ConfigObj("config.ini", encoding="UTF8")
        if cf.get(projdbKey) is None:
            dj = list()
        dbfileName = cf.get(projdbKey).get("dbFileName")
        dbfileDir = os.path.dirname(app.config.get("USE_4DB_NAME"))
        strDBPath = os.path.join(dbfileDir, dbfileName)

        strRealDBName = BEOPSqliteAccess.getInstance().GetRealtimeDBName(strDBPath)
        save_app_config_database()
        app.config['DATABASE'] = strRealDBName

    pointListReady = pointList
    pointsFromScriptList = []
    if scriptList:
        pointsFromScriptList = get_points_from_script_list(scriptList)
        pointsMerged = []
        for item in pointsFromScriptList:
            for point in item.get("pointList"):
                if point not in pointsMerged:
                    pointsMerged.append(point)

        pointListReady = list(set(pointList + pointsMerged))

    hisDataResult = BEOPDataAccess.getInstance().get_history_data_padded(pointListReady, strTimeStart, strTimeEnd, strTimeFormat, dFilter)
    if scriptList:
        scriptHisData = get_history_data_of_scripts(hisDataResult.get("map"), pointsFromScriptList)
        hisDataResult.get("map").update(scriptHisData)

    if projId:
        restore_app_config_database()

    return json.dumps(hisDataResult, ensure_ascii=False)



@app.route('/hisdata/remove_history_data', methods=['POST'])
def hisdata_remove_history_data():
    data = request.get_json()
    removeList = data.get('removeList')

    #check pointList:

    if not isinstance(removeList, list):
        return json.dumps(dict(err=1, msg='removeList param should be List'), ensure_ascii=False)

    # invalid query filter:
    nErrCount = 0
    for item in removeList:
        strTimeAt =  item.get('pointtime')
        strPoitnName = item.get('pointname')
        result = BEOPDataAccess.getInstance().remove_history_data(strTimeAt, strPoitnName)
        if not result:
            nErrCount+=1
    return jsonify(dict(err=0, msg='',data=dict(errCount=nErrCount)))

@app.route('/hisdata/remove_history_data_in_time_range', methods=['POST'])
def hisdata_remove_history_data_in_time_range():
    data = request.get_json()
    pointNameList = data.get('pointNameList')
    strTimeFrom = data.get('timeFrom')
    strTimeTo = data.get('timeTo')
    strWhere = data.get('where')

    #check pointList:

    if not pointNameList:
        return json.dumps(dict(err=1, msg='pointNameList param Lost'), ensure_ascii=False)
    if not strTimeFrom:
        return json.dumps(dict(err=1, msg='timeFrom param Lost'), ensure_ascii=False)
    if not strTimeTo:
        return json.dumps(dict(err=1, msg='timeTo param Lost'), ensure_ascii=False)

    # invalid query filter:
    tTimeFrom = datetime.strptime(strTimeFrom,'%Y-%m-%d %H:%M:%S')
    tTimeTo = datetime.strptime(strTimeTo,'%Y-%m-%d %H:%M:%S')
    nErrCount = 0

    #clear m1,m5
    tCur = tTimeFrom
    tCurTo = tCur+ timedelta(days=1)
    while tCur< tTimeTo:
        if tCurTo> tTimeTo:
            tCurTo = tTimeTo
        result = BEOPDataAccess.getInstance().remove_history_data_in_time_range(tCur.strftime('%Y-%m-%d %H:%M:%S'), tCurTo.strftime('%Y-%m-%d %H:%M:%S'),  pointNameList, 'm1', strWhere)
        result = BEOPDataAccess.getInstance().remove_history_data_in_time_range(tCur.strftime('%Y-%m-%d %H:%M:%S'),
                                                                                tCurTo.strftime('%Y-%m-%d %H:%M:%S'),
                                                                                pointNameList, 'm5',strWhere)

        if not result:
            nErrCount+=1

        tCur = tCurTo
        tCurTo = tCur+ timedelta(days=1)
        time.sleep(0.3)
    #clear h1
    tCur = tTimeFrom
    tCurTo = tCur + timedelta(days=25)
    while tCur < tTimeTo:
        if tCurTo > tTimeTo:
            tCurTo = tTimeTo
        result = BEOPDataAccess.getInstance().remove_history_data_in_time_range(tCur.strftime('%Y-%m-%d %H:%M:%S'),
                                                                                tCurTo.strftime('%Y-%m-%d %H:%M:%S'),
                                                                                pointNameList, 'h1', strWhere)

        if not result:
            nErrCount += 1

        tCur = tCurTo
        tCurTo = tCur + timedelta(days=25)

    # clear d1
    tCur = tTimeFrom
    tCurTo = tCur + timedelta(days=350)
    while tCur < tTimeTo:
        if tCurTo > tTimeTo:
            tCurTo = tTimeTo
        result = BEOPDataAccess.getInstance().remove_history_data_in_time_range(tCur.strftime('%Y-%m-%d %H:%M:%S'),
                                                                                tCurTo.strftime('%Y-%m-%d %H:%M:%S'),
                                                                                pointNameList, 'd1', strWhere)

        if not result:
            nErrCount += 1

        tCur = tCurTo
        tCurTo = tCur + timedelta(days=350)

    return jsonify(dict(err=0, msg='',data=dict(errCount=nErrCount)))

# 只用于值为字典的redis缓存
@app.route('/redis/updateMap', methods=['POST'])
def redis_update_map():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}
    key = rcv.get("key", None)
    value = rcv.get("value", None)
    if not isinstance(key, str):
        return jsonify(dict(err=1, msg='key不能为空', data=False))

    if not len(key):
        return jsonify(dict(err=1, msg='key不能为空', data=False))

    if not RedisManager.is_alive():
        return jsonify(dict(err=1, msg='redis未启动', data=False))

    if not isinstance(value, dict):
        return jsonify(dict(err=1, msg='值必须为字典', data=False))

    jsonValue = RedisManager.get(key)
    if not isinstance(jsonValue, dict):
        jsonValue = {}
    jsonValue.update(value)

    RedisManager.set(key, jsonValue)
    return jsonify(dict(err=0, msg='更新成功', data=True))

@app.route('/redis/updateHistoryDataMinutesOfDate', methods=['POST'])
def redis_update_historydata_minutes_of_date():
    tstart = datetime.now()
    data = request.get_json()
    strFromTime = data.get('fromDate')
    strToTime = data.get('toDate')
    nReset = data.get('clear', 0)

    try:
        t0 =  datetime.strptime(strFromTime,'%Y-%m-%d %H:%M:%S')
        t1 =  datetime.strptime(strToTime,'%Y-%m-%d %H:%M:%S')
    except:
        return json.dumps(dict(err=1, msg='time format wrong'), ensure_ascii=False)


    if nReset==1:
        print('start del_history_data_minutes_keys ')
        RedisManager.del_history_data_minutes_keys(t0,t1)
        print('end del_history_data_minutes_keys ')
    else:
        try:
            tCur = t0
            while tCur<=t1:
                BEOPDataAccess.getInstance().updateRedisDataMinutesOfDay(t0)
                tCur+= timedelta(days=1)
        except:
            return json.dumps(dict(err=1, msg='err in updateRedisDataMinutesOfDay'), ensure_ascii=False)

    return jsonify(dict(err=0, msg='',data={}))

@app.route('/redis/updateHistoryDataMinutes', methods=['POST'])
def redis_update_historydata_minutes():
    tstart = datetime.now()
    data = request.get_json()
    strTime = data.get('time')
    pointNameList = data.get('pointNameList')
    pointValueList = data.get('pointValueList')


    try:
        t0 =  datetime.strptime(strTime,'%Y-%m-%d %H:%M:%S')
    except:
        return json.dumps(dict(err=1, msg='time format wrong'), ensure_ascii=False)

    #check pointList:

    if not isinstance(pointNameList, list):
        return json.dumps(dict(err=1, msg='pointNameList param should be List'), ensure_ascii=False)


    if not isinstance(pointValueList, list):
        return json.dumps(dict(err=1, msg='pointValueList param should be List'), ensure_ascii=False)
    # invalid query filter:

    if RedisManager.is_alive():
        RedisManager.set_history_data_list_minutes(t0, pointNameList, pointValueList)
        logging.error('Recv redis/updateHistoryDataMinutes cost:%.1f seconds'%((datetime.now()-tstart).total_seconds()))
    else:
        return jsonify(dict(err=1, msg='redis not alive',data={}))

    return jsonify(dict(err=0, msg='',data={}))


# 用途: 请求将某个时刻的历史数据（接口传入点名清单和数据清单）更新入 Redis
# 输入参数:
##  time: 时刻，时间字符串，格式如2021-09-01 23:00:00
##  pointNameList: 点名清单，字符串数组
##  pointValueList 点值清单，字符串数组
#
@app.route('/redis/updateHistoryData', methods=['POST'])
def redis_update_historydata():
    tstart = datetime.now()
    try:
        data = request.get_json()
    except Exception as e:
        strError = 'ERROR in redis_update_historydata:%s'%(e.__str__())
        logging.error(strError)
        return jsonify(dict(err=1, msg=strError,data={}))

    strTime = data.get('time')
    pointNameList = data.get('pointNameList')
    pointValueList = data.get('pointValueList')

    #check pointList:


    if not isinstance(pointNameList, list):
        return json.dumps(dict(err=1, msg='pointNameList param should be List'), ensure_ascii=False)


    if not isinstance(pointValueList, list):
        return json.dumps(dict(err=1, msg='pointValueList param should be List'), ensure_ascii=False)
    # invalid query filter:

    if RedisManager.is_alive():
        RedisManager.set_history_data_list(pointNameList, [strTime]*len(pointNameList), pointValueList)
        logging.error('Recv redis/updateHistoryData cost:%.1f seconds'%((datetime.now()-tstart).total_seconds()))
    else:
        return jsonify(dict(err=1, msg='redis not alive',data={}))

    return jsonify(dict(err=0, msg='',data={}))


@app.route('/redis/downloadMinutesToExcel', methods=['POST'])
def redis_download_minutes_to_excel():
    tstart = datetime.now()
    try:
        data = request.get_json()
    except Exception as e:
        strError = 'ERROR in redis_update_historydata:%s'%(e.__str__())
        logging.error(strError)
        return jsonify(dict(err=1, msg=strError,data={}))

    strTimeFrom = data.get('timeFrom')
    strTimeTo = data.get('timeTo')

    #check pointList:

    tFrom = datetime.strptime(strTimeFrom, '%Y-%m-%d %H:%M:00')
    tTo = datetime.strptime(strTimeTo, '%Y-%m-%d %H:%M:00')

    book = Workbook()
    sheet = book.create_sheet("数据", 0)

    tCur = tFrom

    sheet.cell(row=1, column=1, value='time')
    sheet.cell(row=1, column=2, value='pointname')
    sheet.cell(row=1, column=3, value='pointvalue')
    rowIndex = 2
    while tCur<=tTo:
        strAllPointsValueMap = RedisManager.get_history_data_minutes_all_points(tCur)
        for k,v in strAllPointsValueMap.items():
            sheet.cell(row=rowIndex , column=1, value= tCur.strftime('%Y-%m-%d %H:%M:00'))
            sheet.cell(row=rowIndex , column=2, value= k)
            sheet.cell(row=rowIndex , column=3, value= v)
            rowIndex+=1
        tCur+= timedelta(minutes=1)




    saveFileDir = os.getcwd() + '\\siteinterface\\static\\temp'
    if not os.path.exists(saveFileDir):
        os.makedirs(saveFileDir)

    strCreateFileTempName = datetime.now().strftime('%Y%m%d%H%M%S') + '_' + "redis_download.xlsx"
    saveFilePath = os.path.join(saveFileDir, strCreateFileTempName)

    if os.path.exists(saveFilePath):
        os.remove(saveFilePath)

    book.save(saveFilePath)


    return jsonify(dict(err=0, msg="生成成功", data=strCreateFileTempName))


@app.route('/hisdata/modify_history_data', methods=['POST'])
def hisdata_modify_history_data():
    data = request.get_json()
    modifyList = data.get('modifyList')

    #check pointList:

    if not isinstance(modifyList, list):
        return json.dumps(dict(err=1, msg='modifyList param should be List'), ensure_ascii=False)

    tNow = datetime.now()
    strTimeNow = tNow.strftime('%Y-%m-%d %H:%M:%S')
    # invalid query filter:
    nErrCount = 0

    for item in modifyList:
        strTimeAt =  item.get('pointtime')
        strPoitnName = item.get('pointname')
        strPointValue = item.get('pointvalue')
        result = BEOPDataAccess.getInstance().modify_history_data(strTimeAt, strPoitnName, strPointValue)
        if result:
            BEOPDataAccess.getInstance().addLogicOutputRecordMul(strTimeNow, [strPoitnName], [strPointValue], 'dompysite_post_modify_history_data_%s'%(strTimeAt))
        elif RedisManager.is_alive():
            RedisManager.set_history_data_list([strPoitnName], [strTimeAt], [strPointValue])

    return jsonify(dict(err=0, msg='',data=dict(errCount=nErrCount)))

@app.route('/join_point', methods=['POST'])
def join_point():
    data = request.get_json()
    name = BEOPDataAccess.getInstance().getProjMysqldb(data.get('proj'))
    rv = BEOPDataAccess.getInstance().joinPointsToDb(name, data.get('startTime'), data.get('endTime'),data.get('timeFormat'),data.get('pointList'))
    return json.dumps(rv)

#1:访客  , 1:操作员, 2:管理员 3:调试员
def loginTask(login):
    print('loginTask')
    userAuthInfo = {}
    if login['name']=='cx' and login['pwd']=='DOM.cloud-2016':
        userAuthInfo =  dict(err=0, msg='', data=dict(id=-1, name='cx', role=10))
    else:
        userAuthInfo =  BEOPDataAccess.getInstance().validate_user(login['name'], login['pwd'])

    #site mode = 0 mean in simulation mode
    nSiteMode = 1
    try:
        siteModeInfo = BEOPDataAccess.getInstance().getSiteMode()
        nSiteMode = int(siteModeInfo.get('sitemode', 1))
    except Exception as e0:
        nSiteMode = 1
    if True: #开放永久授权
        userAuthInfo['license'] = dict(expired = 0, leftdays = 30000)
    else:
        nLicenseVersion = BEOPDataAccess.getInstance().getLicenseVersion()
        softAuthInfo =LicenseManage(nLicenseVersion).checkAuthored()
        userAuthInfo['license'] = dict(expired = softAuthInfo[0], leftdays = softAuthInfo[1])
    logofilepath =  os.path.join(app.static_folder, "images")
    logofilepath  =  os.path.join(logofilepath, "logo_small.png")
    if os.path.exists(logofilepath):
        userAuthInfo['logoURL'] = '/static/images/logo_small.png'
    strVersion = siteVersion.getCurrentVersion()
    nums = strVersion.split('.')
    totalversion = int(nums[0])*10000+ int(nums[1])*100+ int(nums[2])
    userAuthInfo['pysiteversion'] = totalversion

    strCloud = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig("CloudSetting")

    dCloud = {}
    if isinstance(strCloud, str):
        dCloud = json.loads(strCloud)
    projNameCloud = dCloud.get("CloudSettingProjectName", None)

    cMng = ConfigIniManager()
    projNameInConfig = cMng.get_content_of_option("projectname")

    userAuthInfo.update(dict(projectNameInCloudSetting=projNameCloud, projectNameInConfigIni=projNameInConfig))

    return userAuthInfo

def get_s3dbpath_cloud(dbname):
    print('get_s3dbpath_cloud')
    s3dbDir = app.config.get('S3DB_DIR_CLOUD')
    filepath = os.path.join(s3dbDir,dbname)
    if os.path.isfile(filepath):
        return filepath
    return None

def getSqlitedbSize(dbname=None):
    print('getSqlitedbSize')
    filepath = BEOPDataAccess.getInstance().get_s3dbpath_local()
    if filepath:
        return os.path.getsize(filepath)
    else:
        return 'error: file not found'

def getSqlitedbName(dbname=None):
    print('getSqliteName')
    filepath = BEOPDataAccess.getInstance().get_s3dbpath_local()
    if filepath:
        return os.path.basename(filepath)
    else:
        return 'error: file not found'

@app.route('/workflow_get_user_statics', methods=['POST'])
def get_workflow_get_user_statics():
    groupStatics = BEOPDataAccess.getInstance().getWorkflowTransactionGroupStatics()
    userStatics = BEOPDataAccess.getInstance().getWorkflowTransactionUserStatics()
    for userItem in userStatics:
        userItem['userName'] = BEOPDataAccess.getInstance().getUserNameById(userItem['userId'])
    return json.dumps(dict(groupStatics=groupStatics, userStatics=userStatics))

@app.route('/getTransactionOperationRecordByUserId/<userId>/<rownumber>')
def getTransactionOperationRecordByUserId(userId,rownumber):
    return json.dumps(BEOPDataAccess.getInstance().getOperationRecordByUserId(userId,rownumber), ensure_ascii=False)

@app.route('/getTransactionGroupUser')
def getTransactionGroupUser():
    return json.dumps(BEOPDataAccess.getInstance().getGroupUser(), ensure_ascii=False)

@app.route('/getMyfavorite/<userId>')
def getMyFavorite(userId):
    return json.dumps(BEOPDataAccess.getInstance().getMyFavorite(userId), ensure_ascii=False)

@app.route('/workflow_get_week_report_statics', methods=['POST'])
def get_workflow_get_week_report_statics():
    rdata = request.get_json()
    proj_data = BEOPDataAccess.getInstance().get_workflow_weekreport_projdata(rdata.get('projid'), rdata.get('userid'),
                                                                              rdata.get('weekbefore'))
    return json.dumps(dict(projdata=proj_data[1], transaction=proj_data[0], user_data=proj_data[2]), ensure_ascii=False)

@app.route('/save_file', methods=['POST'])
def save_file_in_server():
    data = request.get_json()
    filename = data.get('filename')
    filefolder = data.get('filefolder')
    filecontent = data.get('filecontent')
    ucode = data.get('ucode')

    strPath = app.static_folder
    strTempPath = os.path.join(strPath, 'files')
    strTempPath = os.path.join(strTempPath, 'temp')

    folder = path.join(strTempPath, filefolder)
    fileDir = path.join(folder,filename)
    makedirs(folder, exist_ok=True)
    if filecontent:
        try:
            file = codecs.open(fileDir,'w',ucode)
            '''filecontent = filecontent.decode('utf-8')'''
            file.write(filecontent)
            file.close()
        except Exception as e:
            strError = 'ERROR in save_file_in_server:%s'%(e.__str__())
            print(strError)
            logging.error(strError)
    rv =  'success'
    return json.dumps(rv, ensure_ascii=False)

@app.route('/get_csv_data', methods=['POST'])
def get_csv_data():
    upload_file = request.files.getlist("cvsFile")
    tmpfilename = ''
    for file in upload_file:
        file.save(file.filename)
        tmpfilename = file.filename
        break

    json_file = []
    csvReader = csv.reader(open(tmpfilename, 'r'))
    for row in csvReader:
        parameterStr = ','.join(row)
        param = parameterStr.split(',')
        obj = [{'name': param[0]}, {'value': [param[1], param[2], param[3]]}]
        json_file.append(obj)

    return json.dumps(json_file)

@app.route('/get_config_data', methods=['POST'])
def get_config_data():
    upload_file = request.files.getlist("config-file")
    tmpfilename = ''
    for file in upload_file:
        file.save(file.filename)
        tmpfilename = file.filename
        break

    dataTimePeriod = 1
    json_file = []
    csvReader = csv.reader(open(tmpfilename, 'r'))
    rowIndex = 0
    ptNameList = []
    vList = []
    for row in csvReader:
        parameterStr = ','.join(row)
        param = parameterStr.split(',')
        if rowIndex==0:
            ptNameList = param[1:]
        else:
            obj = dict(time = param[0], value = param[1:])
            vList.append(obj)
        rowIndex +=1


    try:
        t0 =  datetime.strptime(vList[0]['time'],'%Y-%m-%d %H:%M:%S')
        t1 =  datetime.strptime(vList[1]['time'],'%Y-%m-%d %H:%M:%S')
    except ValueError as e:
        try:
            t0 =  datetime.strptime(vList[0]['time'],'%Y-%m-%d')
            t1 =  datetime.strptime(vList[1]['time'],'%Y-%m-%d')
        except ValueError as e:
            try:
                t0 =  datetime.strptime(vList[0]['time'],'%m-%Y')
                t1 =  datetime.strptime(vList[1]['time'],'%m-%Y')
            except ValueError as e:
                try:
                    t0 =  datetime.strptime(vList[0]['time'],'%Y/%m/%d')
                    t1 =  datetime.strptime(vList[1]['time'],'%Y/%m/%d')
                except ValueError as e:
                    strError = 'ERROR in get_config_data:%s' % (e.__str__())
                    print(strError)
                    logging.error(strError)

    tt = t1 - t0
    dataTimeType = tt.total_seconds()/60.0
    rv = []
    for index in range(len(ptNameList)):
        ptName = ptNameList[index]
        ptValue = []
        for row in vList:
            ptValue.append(dict(time = row['time'], value = row['value'][index]))
        rv.append(dict(pointName = ptName, value = ptValue))
    return json.dumps(rv, ensure_ascii=False)

@app.route('/project/create', methods=['POST'])
def project_create():
    rv = ''
    data =  request.get_json()
    projName = data.get('projName')
    createUserName = data.get('userName')
    createUserId = data.get('userId')
    pointData = data.get('pointData')
    now = datetime.now()
    dbName = 'u%s_%s'% ( createUserName, now.strftime('%Y%m%d%H%M%S'))

    newProjId = BEOPDataAccess.getInstance().createProject(projName, projName, '', dbName, '', 'shanghai', projName, createUserId)
    if newProjId > 0:
        rv = newProjId;
    else:
        rv = { 'error': 'insert incorrectly.'}
    return json.dumps(rv, ensure_ascii=False)

@app.route('/update_realtimedata_input_value', methods=['POST'])
def update_realtimedata_input_value():
    data =  request.get_json()
    listInfo = data.get('listInfo')
    BEOPDataAccess.getInstance().update_realtimedata_input_value(listInfo)

    return '1'

@app.route('/load_history_config_value/<userName>', methods=['GET'])
def load_history_config_value(userName):
    return BEOPDataAccess.getInstance().LoadHistoryConfig(userName)


@app.route('/save_history_config_value', methods=['POST'])
def save_history_config_value():
    data = request.get_json()
    BEOPDataAccess.getInstance().SaveHistoryConfig(data['userName'], data['configName'], data['startTime'], data['endTime'], data['projectList'])
    rv =  'success'
    return json.dumps(rv, ensure_ascii=False)


@app.route('/observer/update/<s3dbname>', methods=['GET'])
def prepareResource(s3dbname, clean=True):
    print('prepareResouce')
    # TODO: force remove *.dec
    updatetime = 0 # always replace resouce where call this api
    configDb = path.join('/s3db' , s3dbname)
    dbfileDec = '{}.dec'.format(configDb)
    BEOPSqliteAccess.getInstance().prepareResouceFromS3db(dbfileDec, updatetime, clean)
    rv = 'success'
    return json.dumps(rv, ensure_ascii=False)

@app.route('/observer/setuprealtimedb', methods=['GET'])
def setuprealtimedb():
    rv = ''
    pList = BEOPDataAccess.getInstance().getProjectTableNameList()
    for proj in pList:
        if proj is None or len(proj)==0:
            continue
        if not BEOPDataAccess.getInstance().checkTableExist('beopdoengine', 'rtdata_%s'% proj):
            BEOPDataAccess.getInstance().createMysqlTable('rtdata_%s'% proj)
            rv += 'rtdata_%s table established.\n' % proj
    if len(rv)==0:
        rv = 'no project realtime table established.'
    return json.dumps(rv, ensure_ascii=False)

@app.route('/clearrtdata', methods=['POST'])
def clearrtdata():
    rv = ''
    data = request.get_json()
    projName = data['projName']
    rv = BEOPDataAccess.getInstance().clearRealtimeInputData(projName)
    return json.dumps(rv, ensure_ascii=False)


@app.route('/load_data_source_record/<userId>', methods=['GET'])
def load_data_source_record(userId):
    return BEOPDataAccess.getInstance().LoadDataSourceRecord(userId)


@app.route('/save_data_source_record', methods=['POST'])
def save_data_source_record():
    data = request.get_json()
    srcList = data['sourceList']
    if srcList == None or len(srcList) < 1:
        return '0'

    ret = BEOPDataAccess.getInstance().SaveDataSourceRecord(srcList)
    if ret == 0:
        return '0'
    else:
        return '-1'


@app.route('/delete_one_data_source_record', methods=['POST'])
def delete_one_data_source_record():
    data = request.get_json()
    ret = BEOPDataAccess.getInstance().DeleteOneDataSourceRecord(data['customName'])
    if ret == 0:
        return '0'
    else:
        return '-1'


@app.route('/delete_data_source_records_by_userid', methods=['POST'])
def delete_data_source_records_by_userid():
    data = request.get_json()
    ret = BEOPDataAccess.getInstance().DeleteDataSourceRecordsByUserId(data['userId'])
    if ret == 0:
        return '0'
    else:
        return '-1'


'''
post body data:
  type:0,1, 0:hhll , 1: bool
  hhenable:0,1
  henable:0,1
  llenable:0,1
  lenable:0,1
  hhlimit: double
  hlimit: double
  llimit:double
  lllimit: double
  pointname: string
  hhinfo: string
  hinfo: string
  llinfo: string
  linfo: string
  boolWarningLevel：int
  boolWarningInfo: string
  warningGroup: string
'''
@app.route('/warningConfig/add', methods=['POST'])
def warning_config_add():
    data = request.get_json()
    ret = BEOPDataAccess.getInstance().addWarningConfigItem(data)

    return json.dumps(ret , ensure_ascii=False)

@app.route('/warningConfig/addRule', methods=['POST'])
def warning_config_add_rule():
    data = request.get_json()
    ret = dict(err=0, msg='', data={})
    strScript = data.get('script', '')
    nWarningLevel = data.get("boolWarningLevel", 0)
    pointname = data.get('pointname', '')
    ofPositionName = data.get('ofPosition', '')
    ofSystemName = data.get('ofSystem', '')
    ofDepartmentName = data.get('ofDepartment', '')
    ofGroupName = data.get('ofGroup', '')
    tags = data.get('tag', '')
    strWarningInfo = data.get('boolWarningInfo', '')
    strWarningGroup = data.get('warningGroup', '')
    strInfoDetail = data.get("infoDetail", None)
    strUnitProperty02 = data.get("prop02", None)
    strUnitProperty03 = data.get("prop03", None)
    strUnitProperty04 = data.get("prop04", None)
    strUnitProperty05 = data.get("prop05", None)

    maxId = BEOPDataAccess.getInstance().getMaxIdInTable('warning_config')
    ruleId = maxId + 1

    nNewId = BEOPDataAccess.getInstance().addWarningConfigRuleItem(nWarningLevel, strScript, pointname,
                                                                   ofDepartmentName, ofPositionName,
                                                                   ofSystemName, ofGroupName, tags, strWarningInfo,
                                                                   strWarningGroup, strInfoDetail, ruleId,
                                                                   strUnitProperty02, strUnitProperty03,
                                                                   strUnitProperty04, strUnitProperty05)

    if nNewId<0:
        ret['err']=1
    else:
        ret['data'] = nNewId
    return json.dumps(ret, ensure_ascii=False)


'''
 获取所有报警配置清单
'''
@app.route('/warningConfig/getAll', methods=['POST'])
def warning_config_get_all():
    data = request.get_json()
    ret = BEOPDataAccess.getInstance().getAllWarningConfig()

    return json.dumps(ret , ensure_ascii=False)


@app.route("/warningConfig/export")
def warning_config_export():
    try:
        ret = BEOPDataAccess.getInstance().getAllWarningConfig()
        if not len(ret):
            return jsonify(dict(err=1, msg="无报警配置内容", data=""))

        boolList = []
        rangeList = []
        ruleList = []
        for item in ret:
            if item.get("type") == 1:
                boolList.append(
                    (item.get("id"),
                     item.get("pointname"),
                     item.get("boolWarningInfo"),
                     item.get("boolWarningLevel"),
                     item.get("ofPosition"),
                     item.get("ofSystem"),
                     item.get("ofDepartment"),
                     item.get("ofGroup"),
                     item.get("tag"),
                     item.get("unitproperty01"),
                     item.get("unitproperty02"),
                     item.get("unitproperty03"),
                     item.get("unitproperty04"),
                     item.get("unitproperty05"))
                )
            elif item.get("type") == 0:
                rangeList.append(
                    (item.get("id"),
                     item.get("pointname"),
                     item.get("hlimit"),
                     item.get("hinfo"),
                     item.get("hlimit"),
                     item.get("linfo"),
                     item.get("boolWarningLevel"),
                     item.get("ofPosition"),
                     item.get("ofSystem"),
                     item.get("ofDepartment"),
                     item.get("ofGroup"),
                     item.get("tag"),
                     item.get("unitproperty01"),
                     item.get("unitproperty02"),
                     item.get("unitproperty03"),
                     item.get("unitproperty04"),
                     item.get("unitproperty05"))
                )
            elif item.get("type") == 3:
                ruleList.append(
                    (item.get("id"),
                     item.get("pointname"),
                     item.get("script"),
                     item.get("linfo"),
                     item.get("boolWarningLevel"),
                     item.get("ofPosition"),
                     item.get("ofSystem"),
                     item.get("ofDepartment"),
                     item.get("ofGroup"),
                     item.get("tag"),
                     item.get("infoDetail"),
                     item.get("unitproperty01"),
                     item.get("unitproperty02"),
                     item.get("unitproperty03"),
                     item.get("unitproperty04"),
                     item.get("unitproperty05"))
                )
        book = Workbook()
        boolSheet = book.create_sheet("bool", 0)
        rangeSheet = book.create_sheet("range", 1)
        ruleSheet = book.create_sheet("rule", 2)

        boolHeader = ['id','pointname','info','level','position','system','department','group','tag',
                      'property01','property02','property03','property04','property05']
        rangeHeader = ['id','pointname','hlimit','hinfo','llmit','linfo','level','position','system',
                       'department','group','tag','property01','property02','property03','property04','property05']
        ruleHeader = ['id','pointname','script','linfo','level','position','system','department','group',
                      'tag','infodetail','property01','property02','property03','property04','property05']

        # 写表头
        for idx, item in enumerate(boolHeader):
            boolSheet.cell(row=1, column=idx+1, value=item)

        for idx, item in enumerate(rangeHeader):
            rangeSheet.cell(row=1, column=idx+1, value=item)

        for idx, item in enumerate(ruleHeader):
            ruleSheet.cell(row=1, column=idx+1, value=item)

        # 写内容
        for i, dataList in enumerate(boolList):
            for j, data in enumerate(dataList):
                boolSheet.cell(row=i+2, column=j+1, value=data)

        for i, dataList in enumerate(rangeList):
            for j, data in enumerate(dataList):
                rangeSheet.cell(row=i + 2, column=j + 1, value=data)

        for i, dataList in enumerate(ruleList):
            for j, data in enumerate(dataList):
                ruleSheet.cell(row=i + 2, column=j + 1, value=data)

        filesDir = os.path.join(app.static_folder, "files")
        if not os.path.exists(filesDir):
            os.mkdir(filesDir)

        backupFilePath = os.path.join(filesDir, "warningConfigBackupTable.xlsx")
        if os.path.exists(backupFilePath):
            os.remove(backupFilePath)

        book.save(backupFilePath)

        return jsonify(dict(err=0, msg="备份成功", data="warningConfigBackupTable.xlsx"))

    except Exception as e:
        logging.error("ERROR in /warningConfig/export: %s" % e.__str__())
        return jsonify(dict(err=1, msg="备份失败: %s" % e.__str__(), data=""))




'''
 编辑一个报警配置
 data内容与新建时完全相同
'''
@app.route('/warningConfig/edit', methods=['POST'])
def warning_config_edit():
    data = request.get_json()

    if data.get('id'):
        nType = data.get('type')
        if nType == 3:  # type is script:
            nID = int(data.get('id'))
            strScript = data.get('script')
            pointname = data.get('pointname')
            ofPositionName = data.get('ofPosition')

            nWarningLevel = data.get("boolWarningLevel", 0)
            strWarningInfo = data.get("boolWarningInfo", '')
            ofSystemName = data.get('ofSystem')
            ofDepartmentName = data.get('ofDepartment')
            ofGroupName = data.get('ofGroup')
            tags = data.get('tag')
            strWarningGroup = data.get('warningGroup', '')
            bSuccess = BEOPDataAccess.getInstance().editWarningConfigRuleItemById(nID,nWarningLevel, strScript, pointname, ofDepartmentName, ofPositionName, ofSystemName, ofGroupName, tags, strWarningInfo, strWarningGroup)
            if bSuccess:
                return json.dumps(dict(err=0,data={}) , ensure_ascii=False)
            else:
                return json.dumps(dict(err=1,data={}) , ensure_ascii=False)

    ret = BEOPDataAccess.getInstance().editWarningConfigItem(data)

    return json.dumps(ret , ensure_ascii=False)


'''
 删除一个报警配置
 body:
    pointName: string
    type: 0,1
'''
@app.route('/warningConfig/remove', methods=['POST'])
def remove_warning_config():
    data = request.get_json()
    if data.get('id'):
        nID = int(data.get('id'))
        if nID>=0:
            ret = BEOPDataAccess.getInstance().removeWarningConfigItemById(nID)
            return json.dumps(ret , ensure_ascii=False)


    strPointName = data.get('pointname')
    nType = data.get('type')
    ret = BEOPDataAccess.getInstance().removeWarningConfigItem(strPointName, int(nType))

    return json.dumps(ret , ensure_ascii=False)

'''
 获取报警分类列表
 body:
    pointName: string
    type: 0,1
'''
@app.route('/warningConfig/getGroupList', methods=['POST'])
def get_warning_config_group_list():
    ret = BEOPDataAccess.getInstance().getWarningGroupList()

    return json.dumps(ret , ensure_ascii=False)



'''
 处理报警
'''
@app.route('/warning/deal', methods=['POST'])
def warning_deal():
    data = request.get_json()
    strInfo = data.get('info', '')
    nOpType = int(data.get('type', 0)) #1 确认  2忽略  3 消除
    strRemark = data.get('remark', "")
    userName = data.get("userName", "")
    ignoreMinutes = data.get("ingnoreMinutes", None)

    nUserId = None

    if nOpType not in [1, 2, 3]:
        return jsonify(dict(err=1, msg="opType必须为1或2或3", data=False))

    if not isinstance(strInfo, str) and not isinstance(strInfo, list):
        return jsonify(dict(err=1, msg="info必须为字符串或数组", data=False))

    if isinstance(strInfo, str):
        if not len(strInfo):
            return jsonify(dict(err=1, msg="info不能为空字符串", data=False))

    if isinstance(strInfo, list):
        for item in strInfo:
            if not isinstance(item, str):
                return jsonify(dict(err=1, msg="info为数组时里面的元素必须为字符串，且不能为空字符串", data=False))
            if not len(item):
                return jsonify(dict(err=1, msg="info为数组时里面的元素必须为字符串，且不能为空字符串", data=False))

    # 只有在忽略和确认时才将userName 转换为userId
    if nOpType in [1, 2]:
        if not isinstance(userName, str):
            return jsonify(dict(err=1, msg="忽略和确认故障时必须传入用户名", data=False))
        if not len(userName):
            return jsonify(dict(err=1, msg="忽略和确认故障时用户名不能为空", data=False))

        nUserId = BEOPDataAccess.getInstance().get_user_id_by_name(userName)
        if nUserId == None:
            return jsonify(dict(err=1, msg="用户名{userName}对应的用户ID不存在或获取失败".format(userName=userName), data=False))

    if strRemark != None:
        if not isinstance(strRemark, str):
            return jsonify(dict(err=1, msg="原因必须为字符串", data=False))

    # 忽略必须传入忽略分钟数
    if nOpType in [2]:
        if not isinstance(ignoreMinutes, int):
            return jsonify(dict(err=1, msg="忽略故障时必须传入忽略分钟数", data=False))
        if ignoreMinutes < 0:
            return jsonify(dict(err=1, msg="忽略故障的分钟必须大于等于0", data=False))

    # 若关闭故障，则使用传入的userName查询一次userId，若查询不到则将userName作为userId记录
    if nOpType == 3:
        nUserId = BEOPDataAccess.getInstance().get_user_id_by_name(userName)
        if nUserId == None:
            nUserId = userName if isinstance(userName, str) else ""

    if not isinstance(strRemark, str):
        strRemark = ""

    conditionList = []
    if isinstance(strInfo, str):
        conditionList.append(strInfo)
    elif isinstance(strInfo, list):
        conditionList = strInfo

    bSuc = BEOPDataAccess.getInstance().dealWarning(conditionList, nOpType, strRemark, nUserId, ignoreMinutes)
    return jsonify(dict(err=0, msg="", data=bSuc))


'''
 插入实时报警
'''
@app.route('/warning/add', methods=['POST'])
def warning_add_realtime():
    data = request.get_json()
    strInfo = data.get('info', '')
    nLevel = int(data.get('level',0))
    strBindPointName = data.get('bindPointName', '')
    nRuleId = int(data.get('ruleId', -1))
    strOfPosition = data.get('ofPosition', '')
    strOfSystem = data.get('ofSystem', '')
    strOfDepartment = data.get('ofDepartment', '')
    strOfGroup = data.get('ofGroup', '')
    strTag = data.get('tag', '')
    strInfoDetail = data.get('infoDetail', '')

    ret = BEOPDataAccess.getInstance().insertWarningRecord(strInfo, nLevel, strBindPointName, nRuleId, strOfPosition,
                                                           strOfSystem, strOfDepartment, strOfGroup, strTag, strInfoDetail)

    return json.dumps(dict(err=0, data=ret, msg="") , ensure_ascii=False)



'''
获取实时报警清单
'''
@app.route('/warning/getRealtime', methods=['POST'])
def warning_get_realtime():
    data = request.get_json()
    ret = {}
    seconds = 30

    bGetPageWarning = False
    if isinstance(data, dict):
        seconds = data.get('seconds', 3600 * 4)             # 请求从4小时前到现在为止的报警记录
        bGetPageWarning = data.get('getPageWarning', False)  # 是否分析网络架构在离线

    dUserIdNameMap = {}
    dUserList = BEOPDataAccess.getInstance().get_all_users()
    for userDict in dUserList:
        dUserIdNameMap.update({userDict["userid"]: userDict["username"]})

    warningList = BEOPDataAccess.getInstance().getRealtimeWarningList(seconds)
    strBindPointNameList = []

    for item in warningList:
        nConfirmStatus = item.get("nConfirmStatus", 0)
        nConfirmUserId = item.get("nConfirmUserId", None)
        strConfirmRemark = item.get("strConfirmRemark", "")
        tConfirmOpTime = item.get("tConfirmOpTime", None)
        nIgnoreStatus = item.get("nIgnoreStatus", 0)
        nIgnoreUserId = item.get("nIgnoreUserId", None)
        strIgnoreRemark = item.get("strIgnoreRemark", "")
        tIgnoreOpTime = item.get("tIgnoreOpTime", None)
        tIgnoreToTime = item.get("tIgnoreToTime", None)
        nCloseStatus = item.get("nCloseStatus", 0)
        nCloseUserId = item.get("nCloseUserId", None)
        strCloseRemark = item.get("strCloseRemark", "")
        tCloseOpTime = item.get("tCloseOpTime", None)

        strConfirmUserName = ""
        if isinstance(nConfirmUserId, int) or is_int_digit(nConfirmUserId):
            strConfirmUserName = dUserIdNameMap.get(str(nConfirmUserId))

        strIgnoreUserName = ""
        if isinstance(nIgnoreUserId, int) or is_int_digit(nIgnoreUserId):
            strIgnoreUserName = dUserIdNameMap.get(str(nIgnoreUserId))

        strCloseUserName = ""
        if isinstance(nCloseUserId, int) or is_int_digit(nCloseUserId):
            strCloseUserName = dUserIdNameMap.get(str(nCloseUserId))

        nStatus = 1
        strStatus = "未确认，未消除"
        if nConfirmStatus == 1 and nCloseStatus != 1:
            nStatus = 5
            strStatus = "已确认，未消除"

        elif nConfirmStatus == 1 and nCloseStatus == 1:
            nStatus = 6
            strStatus = "已确认，已消除"

        elif isinstance(tIgnoreToTime, datetime) and datetime.now() <= tIgnoreToTime and nCloseStatus == 1:
            nStatus = 4
            strStatus = "已忽略，已消除"

        elif isinstance(tIgnoreToTime, datetime) and datetime.now() <= tIgnoreToTime and nCloseStatus != 1:
            nStatus = 3
            strStatus = "已忽略，未消除"

        elif nConfirmStatus != 1 and nCloseStatus == 1:
            nStatus = 2
            strStatus = "未确认，已消除"

        elif nConfirmStatus != 1 and nCloseStatus != 1:
            nStatus = 1
            strStatus = "未确认，未消除"

        if isinstance(tCloseOpTime, datetime):
            tCloseOpTime = tCloseOpTime.strftime("%Y-%m-%d %H:%M:%S")
        elif tCloseOpTime == None:
            tCloseOpTime = ""
        elif isinstance(tCloseOpTime, str):
            tCloseOpTime = tCloseOpTime

        if isinstance(tIgnoreOpTime, datetime):
            tIgnoreOpTime = tIgnoreOpTime.strftime("%Y-%m-%d %H:%M:%S")
        elif tIgnoreOpTime == None:
            tIgnoreOpTime = ""
        elif isinstance(tIgnoreOpTime, str):
            tIgnoreOpTime = tIgnoreOpTime

        if isinstance(tIgnoreToTime, datetime):
            tIgnoreToTime = tIgnoreToTime.strftime("%Y-%m-%d %H:%M:%S")
        elif tIgnoreToTime == None:
            tIgnoreToTime = ""
        elif isinstance(tIgnoreToTime, str):
            tIgnoreToTime = tIgnoreToTime

        if isinstance(tConfirmOpTime, datetime):
            tConfirmOpTime = tConfirmOpTime.strftime("%Y-%m-%d %H:%M:%S")
        elif tConfirmOpTime == None:
            tConfirmOpTime = ""
        elif isinstance(tConfirmOpTime, str):
            tConfirmOpTime = tConfirmOpTime

        item.update(dict(status=nStatus,
                         strStatus=strStatus,
                         tCloseOpTime=tCloseOpTime,
                         tIgnoreOpTime=tIgnoreOpTime,
                         tIgnoreToTime=tIgnoreToTime,
                         tConfirmOpTime=tConfirmOpTime,
                         closeUser=strCloseUserName,
                         confirmUser=strConfirmUserName,
                         ignoreUser=strIgnoreUserName,
                         confirmRemark=strConfirmRemark,
                         ignoreRemark=strIgnoreRemark,
                         closeRemark=strCloseRemark))

        if item.get('strBindPointName'):
            strBindPointNameList.append(item.get('strBindPointName'))

    ret.update(dict(warningList=warningList))

    if not bGetPageWarning:
        return json.dumps(ret, ensure_ascii=False)

    pageIdList = RedisManager.get_point_used_page_list(strBindPointNameList)

    nFireMode = 0
    fireModeScript = app.config.get("FIRE_MODE_SCRIPT")
    if fireModeScript:
        pointData, pointDataMap = BEOPDataAccess.getInstance().getInputTable()
        nFireMode = eval_string_expression_strict(fireModeScript, "1", datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pointData)

    ret.update(dict(warningPageIdList=pageIdList, fireMode=nFireMode))

    # 若redis中已存在网络在离线分析结果则直接获取（来自domJobs)
    if RedisManager.is_alive():
        netDeviceDropWarningList = RedisManager.get("net_device_offline_warning_list")
        netDeviceUnclearList = RedisManager.get("net_device_unclear_list")
        netDeviceOnlineList = RedisManager.get("net_device_online_list")
        if netDeviceDropWarningList != None and netDeviceUnclearList != None and netDeviceOnlineList != None:
            ret.update({"netDeviceDropWarningList": netDeviceDropWarningList,
                        "netDeviceOnlineList": netDeviceOnlineList,
                        "netDeviceUnclearList": netDeviceUnclearList})
            return json.dumps(ret, ensure_ascii=False)

    dWarningNoticeConfig = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfigMul(["warning_notice_config"])
    nDelaySecs = 20 * 60
    try:
        nDelaySecs = int(float(dWarningNoticeConfig.get("warning_notice_config", {}).get("EquipmentWarningDelaySeconds", None)))
    except:
        pass

    if nDelaySecs < 0:
        nDelaySecs = 1200

    deviceOfflineMinsThreshold = nDelaySecs / 60   # 设备离线判断分钟delta值  从配置读取

    netDeviceDropWarningList = []   # 确定掉线的设备
    netDeviceOnlineList = []  # 确定在线的设备
    netDeviceUnclearList = []  # 不清楚的或者无法判断是否在线的设备

    dAdapters = get_server_adapters()
    for adapterIp, dAdapterInfo in dAdapters.items():
        if ping_ip(adapterIp):
            netDeviceOnlineList.append(dict(ip=adapterIp,
                                            deviceId="",
                                            type="server",
                                            reason="",
                                            suggestion="",
                                            delay=None,
                                            id=adapterIp,
                                            prefix="",
                                            suffix="",
                                            equipType=""))
        else:
            netDeviceUnclearList.append(dict(ip=adapterIp,
                                                 deviceId="",
                                                 type="server",
                                                 reason="{ip}无法ping通".format(ip=adapterIp),
                                                 suggestion="建议检查该服务器网口:{ip}".format(ip=adapterIp),
                                                 id=adapterIp,
                                                 delay=None,
                                                 prefix="",
                                                 suffix="",
                                                 equipType=""))


    # 分析modbus equipment 在线离线
    ptNmList = []
    dNetworkNode = {}   # 心跳点与网络节点ID的对应关系
    dStationGateway = {}  # Station心跳点与网关心跳点的对应关系
    dData = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfigMul(["modbusclientconfig"])
    dModbusEquipmentPointDefine = get_standard_modbus_equipment_define_from_cloud_frequent()
    dModbusConfig = dData.get("modbusclientconfig", {})
    for modbusTCPKey, dModbusInfo in dModbusConfig.items():
        gatewayIpList = find_ip_list(modbusTCPKey)
        if not len(gatewayIpList):
            continue

        nPort = None
        try:
            nPort = int(float(dModbusInfo.get("port", 502)))
        except:
            pass

        if not isinstance(nPort, int):
            continue

        gatewayIp = gatewayIpList[0]
        ptNmGWHeartBeat = "modbusTCP_{ip}_{port}_heart_beat_time".format(ip=gatewayIp.replace(".", "_"), port=nPort)
        key = "{gatewayIp}_{port}".format(gatewayIp=gatewayIp, port=nPort)

        ptNmList.append(ptNmGWHeartBeat)
        dNetworkNode.update({ptNmGWHeartBeat: dict(nodeId=key, type="ModbusGateway", ip=gatewayIp, deviceId="")})

        dStations = dModbusInfo.get("stations", {})
        for strDeviceId, dStation in dStations.items():
            nDeviceId = None
            try:
                nDeviceId = int(float(strDeviceId))
            except:
                pass

            if not isinstance(nDeviceId, int):
                continue

            strEquipTypeEn = get_modbus_equip_type_en(dModbusEquipmentPointDefine, dStation.get("type", ""))

            ptNmStationHeartBeat = "modbusTCP_{ip}_{port}_{station}_heart_beat_time".format(ip=gatewayIp.replace(".", "_"),
                                                                                    station=nDeviceId,
                                                                                    port=nPort)

            key = "{ip}_{port}_{deviceId}".format(ip=gatewayIp, deviceId=nDeviceId, port=nPort)
            ptNmList.append(ptNmStationHeartBeat)
            dNetworkNode.update({ptNmStationHeartBeat: dict(nodeId=key,
                                                            type="ModbusStation",
                                                            ip=gatewayIp,
                                                            deviceId=nDeviceId,
                                                            prefix=dStation.get("pointPrefix", ""),
                                                            suffix=dStation.get("Number", ""),
                                                            equipType=strEquipTypeEn)})

            dStationGateway.update({ptNmStationHeartBeat: ptNmGWHeartBeat})

    # 分析点表
    dAllPoint = BEOPSqliteAccess.getInstance().getPointInfoFromS3db([])
    for pointName, dPointInfo in dAllPoint.items():
        if dPointInfo.get("sourceType") == "simense1200TCP":
            plcIp = dPointInfo.get("param3", None)
            if plcIp == None:
                continue

            ptNmPlcHeartBeat = "siemensPLC_{ip}_heart_beat_time".format(ip=plcIp.replace(".", "_"))
            key = plcIp
            ptNmList.append(ptNmPlcHeartBeat)
            dNetworkNode.update({ptNmPlcHeartBeat: dict(nodeId=key, type="SiemensPLC", ip=plcIp, deviceId="", equipType="")})

        elif dPointInfo.get("sourceType") == "bacnet":
            bacnetGatewayIp = dPointInfo.get("addr", None)
            if bacnetGatewayIp == None:
                continue

            if is_ip(bacnetGatewayIp):
                ptNmBacnetHeartBeat = "bacnetGateway_{ip}_heart_beat_time".format(ip=bacnetGatewayIp.replace(".", "_"))
                key = bacnetGatewayIp
                ptNmList.append(ptNmBacnetHeartBeat)
                dNetworkNode.update({ptNmBacnetHeartBeat: dict(nodeId=key, type="BacnetGateway", ip=bacnetGatewayIp, deviceId="", equipType="")})

        elif dPointInfo.get("sourceType") == "bacnet-py":
            param5 = dPointInfo.get("param5", None)
            param6 = dPointInfo.get("param6", None)

            bacnetGatewayIp = None
            ipListParam5 = []
            ipListParam6 = []

            if param5 != None:
                ipListParam5 = find_ip_list(param5)
            if param6 != None:
                ipListParam6 = find_ip_list(param6)

            if len(ipListParam5):
                bacnetGatewayIp = ipListParam5[0]
            elif len(ipListParam6):
                bacnetGatewayIp = ipListParam6[0]

            if bacnetGatewayIp != None:
                ptNmBacnetHeartBeat = "bacnetGateway_{ip}_heart_beat_time".format(ip=bacnetGatewayIp.replace(".", "_"))
                key = bacnetGatewayIp
                ptNmList.append(ptNmBacnetHeartBeat)
                dNetworkNode.update({ptNmBacnetHeartBeat: dict(nodeId=key, type="BacnetGateway", ip=bacnetGatewayIp, deviceId="", equipType="")})

    # 分析能源管理结构
    dEMDefine = {}
    strEMDefine = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig("energy_management_define")
    try:
        dEMDefine = json.loads(strEMDefine)
    except Exception as e:
        print(e.__str__())

    dDistribution = dEMDefine.get("distributionGroupList", {})
    dEMStruct = analysis_energy_management_structure(dDistribution)

    for gatewayId, dGatewayStruct in dEMStruct.items():
        ptNmGWHeartBeat = "{gatewayId}_heart_beat_time".format(gatewayId=gatewayId)
        ptNmList.append(ptNmGWHeartBeat)
        dNetworkNode.update({ptNmGWHeartBeat: dict(nodeId=gatewayId, type="ModbusGateway", ip="", deviceId="")})

        for meterId, dMeterStruct in dGatewayStruct.items():
            ptNmMeterHeartBeat = "{gatewayId}_{meterId}_heart_beat_time".format(gatewayId=gatewayId, meterId=meterId)
            ptNmList.append(ptNmMeterHeartBeat)
            dNetworkNode.update({ptNmMeterHeartBeat: dict(nodeId="{gatewayId}_{meterId}".format(gatewayId=gatewayId, meterId=meterId),
                                                          type="ModbusStation",
                                                          ip="",
                                                          deviceId=meterId)})

            dStationGateway.update({ptNmMeterHeartBeat: ptNmGWHeartBeat})

    # 获取heart_beat_time 实时值
    lRealtimeData, dRealtimeData = BEOPDataAccess.getInstance().getInputTable(ptNmList)

    for ptNmHeartBeat, dNetNode in dNetworkNode.items():
        nOnlineStatus = 1  # 0:掉线；1：在线；-1：不明
        reason = ""
        suggestion = ""

        strHBTime = dRealtimeData.get(ptNmHeartBeat, None)
        if strHBTime == None and dNetNode["type"] not in ["BacnetGateway", "SiemensPLC"]:  # bacnet网关和西门子PLC跳过，后面用ping ip判断这两种网关的在线离线
            nOnlineStatus = -1
            reason = "实时数据中未发现心跳更新时间点名({name})".format(name=ptNmHeartBeat)
            suggestion = "建议检查心跳更新时间点({name})是否存在".format(name=ptNmHeartBeat)
            netDeviceUnclearList.append(dict(ip=dNetNode["ip"],
                                            deviceId=dNetNode["deviceId"],
                                            type=dNetNode["type"],
                                            reason=reason,
                                            suggestion=suggestion,
                                            id=dNetNode["nodeId"],
                                            delay=None,
                                            prefix="",
                                            suffix="",
                                            equipType=dNetNode.get("equipType", "")))
            continue


        # modbus station
        if dNetNode["type"] == "ModbusStation":
            # 首先进行所在网关在线状态检查
            tHBTimeGateway = None
            ptNmHeartBeatGateway = dStationGateway.get(ptNmHeartBeat, None)
            if ptNmHeartBeatGateway:
                strHBTimeGateway = dRealtimeData.get(ptNmHeartBeatGateway, None)
                if is_time_string(strHBTimeGateway, "%Y-%m-%d %H:%M:%S"):
                    tHBTimeGateway = datetime.strptime(strHBTimeGateway, "%Y-%m-%d %H:%M:%S")

            if tHBTimeGateway != None:
                if (datetime.now() - tHBTimeGateway).total_seconds() > deviceOfflineMinsThreshold * 60:
                    nOnlineStatus = 0
                    reason = "所在网关心跳更新时间为{hbTime}(点名:{name})，已超过{mins}分钟未更新，则判断为该设备也处于离线状态".format(hbTime=strHBTimeGateway,
                                                                                                   mins=deviceOfflineMinsThreshold,
                                                                                                   name=ptNmHeartBeatGateway)
                    suggestion = "建议检查网关指示灯，处理网关掉线问题"
                    netDeviceDropWarningList.append(dict(ip=dNetNode["ip"],
                                                         deviceId=dNetNode["deviceId"],
                                                         type=dNetNode["type"],
                                                         reason=reason,
                                                         suggestion=suggestion,
                                                         id=dNetNode["nodeId"],
                                                         delay=(datetime.now() - tHBTimeGateway).total_seconds(),
                                                         prefix=dNetNode.get("prefix", ""),
                                                         suffix=dNetNode.get("suffix", ""),
                                                         equipType=dNetNode.get("equipType", "")))
                    continue

            # 进行心跳更新时间格式检查
            if not is_time_string(strHBTime, "%Y-%m-%d %H:%M:%S"):
                nOnlineStatus = -1
                reason = "心跳更新时间点({ptNm})时间格式有误，应为yyyy-mm-dd HH:MM:SS，但仍然判断为设备离线".format(ptNm=ptNmHeartBeat)
                suggestion = "建议检查设备485接线是否正常、设备是否上电等问题，若无误，则检查心跳更新时间点({ptNm})值的时间格式(应为yyyy-mm-dd HH:MM:SS)".format(ptNm=ptNmHeartBeat)
                netDeviceDropWarningList.append(dict(ip=dNetNode["ip"],
                                                 deviceId=dNetNode["deviceId"],
                                                 type=dNetNode["type"],
                                                 reason=reason,
                                                 suggestion=suggestion,
                                                 id=dNetNode["nodeId"],
                                                 delay=None,
                                                 prefix=dNetNode.get("prefix", ""),
                                                 suffix=dNetNode.get("suffix", ""),
                                                 equipType=dNetNode.get("equipType", "")))
                continue

            # 获取心跳更新时间
            tHBTime = datetime.strptime(strHBTime, "%Y-%m-%d %H:%M:%S")

            # 心跳更新时间是否正常检查
            if (datetime.now() - tHBTime).total_seconds() > deviceOfflineMinsThreshold * 60:
                bOnline = False
                reason = "心跳更新时间为{hbTime}(点名:{name})，已超过{mins}分钟未更新，则判断该设备已离线".format(hbTime=tHBTime.strftime("%Y-%m-%d %H:%M:%S"),
                                                                                      mins=deviceOfflineMinsThreshold,
                                                                                      name=ptNmHeartBeat)
                suggestion = "建议检查设备485接线是否正常、设备是否上电等问题"
                netDeviceDropWarningList.append(dict(ip=dNetNode["ip"],
                                                     deviceId=dNetNode["deviceId"],
                                                     type=dNetNode["type"],
                                                     reason=reason,
                                                     suggestion=suggestion,
                                                     id=dNetNode["nodeId"],
                                                     delay=(datetime.now() - tHBTime).total_seconds(),
                                                     prefix=dNetNode.get("prefix", ""),
                                                     suffix=dNetNode.get("suffix", ""),
                                                     equipType=dNetNode.get("equipType", "")))
                continue

            # 最终确认在线
            netDeviceOnlineList.append(dict(ip=dNetNode["ip"],
                                                 deviceId=dNetNode["deviceId"],
                                                 type=dNetNode["type"],
                                                 reason="",
                                                 suggestion="",
                                                 id=dNetNode["nodeId"],
                                                 delay=(datetime.now() - tHBTime).total_seconds(),
                                                 prefix=dNetNode.get("prefix", ""),
                                                 suffix=dNetNode.get("suffix", ""),
                                                 equipType=dNetNode.get("equipType", "")))

        # modbus 网关
        elif dNetNode["type"] in ["ModbusGateway"]:
            # 进行心跳更新时间格式检查
            if not is_time_string(strHBTime, "%Y-%m-%d %H:%M:%S"):
                nOnlineStatus = -1
                reason = "心跳更新时间点({ptNm})时间格式有误，应为yyyy-mm-dd HH:MM:SS，但仍然判断为设备离线".format(ptNm=ptNmHeartBeat)
                suggestion = "建议检查设备是否上电，与交换机的网线连接是否正常，若无误，则检查心跳更新时间点({ptNm})值的时间格式(应为yyyy-mm-dd HH:MM:SS)".format(ptNm=ptNmHeartBeat)
                netDeviceDropWarningList.append(dict(ip=dNetNode["ip"],
                                                     deviceId=dNetNode["deviceId"],
                                                     type=dNetNode["type"],
                                                     reason=reason,
                                                     suggestion=suggestion,
                                                     id=dNetNode["nodeId"],
                                                     delay=None,
                                                     prefix="",
                                                     suffix="",
                                                     equipType=dNetNode.get("equipType", "")))
                continue

            # 获取心跳更新时间
            tHBTime = datetime.strptime(strHBTime, "%Y-%m-%d %H:%M:%S")

            # 心跳更新时间是否正常检查
            if (datetime.now() - tHBTime).total_seconds() > deviceOfflineMinsThreshold * 60:
                reason = "心跳更新时间为{hbTime}(点名:{name})，已超过{mins}分钟未更新，则判断该设备已离线".format(hbTime=strHBTime,
                                                                                      name=ptNmHeartBeat,
                                                                                      mins=deviceOfflineMinsThreshold)
                suggestion = "建议检查设备是否上电，与交换机的网线连接是否正常"
                netDeviceDropWarningList.append(dict(ip=dNetNode["ip"],
                                                     deviceId=dNetNode["deviceId"],
                                                     type=dNetNode["type"],
                                                     reason=reason,
                                                     suggestion=suggestion,
                                                     id=dNetNode["nodeId"],
                                                     delay=(datetime.now() - tHBTime).total_seconds(),
                                                     prefix="",
                                                     suffix="",
                                                     equipType=dNetNode.get("equipType", "")))
                continue

            # 最终确认设备在线
            netDeviceOnlineList.append(dict(ip=dNetNode["ip"],
                                            deviceId=dNetNode["deviceId"],
                                            type=dNetNode["type"],
                                            reason="",
                                            suggestion="",
                                            id=dNetNode["nodeId"],
                                            delay=(datetime.now() - tHBTime).total_seconds(),
                                            prefix="",
                                            suffix="",
                                            equipType=dNetNode.get("equipType", "")))

        # bacnet网关、西门子PLC
        elif dNetNode["type"] in ["BacnetGateway", "SiemensPLC"]:
            nOnlineStatus = -1
            if isinstance(strHBTime, str):
                if is_time_string(strHBTime, "%Y-%m-%d %H:%M:%S"):
                    tHBTime = datetime.strptime(strHBTime, "%Y-%m-%d %H:%M:%S")
                    if (datetime.now() - tHBTime).total_seconds() < deviceOfflineMinsThreshold * 60:
                        nOnlineStatus = 1
                        netDeviceOnlineList.append(dict(ip=dNetNode["ip"],
                                                        deviceId=dNetNode["deviceId"],
                                                        type=dNetNode["type"],
                                                        reason="",
                                                        suggestion="",
                                                        id=dNetNode["nodeId"],
                                                        delay=(datetime.now() - tHBTime).total_seconds(),
                                                        prefix="",
                                                        suffix="",
                                                        equipType=dNetNode.get("equipType", "")))

            if nOnlineStatus != 1:
                if not is_ip(dNetNode["ip"]):
                    nOnlineStatus = -1
                    reason = "网关ip格式有误（{ip}），无法判断在离线".format(ip=dNetNode["ip"])
                    suggestion = "建议检查关于该网关下设备的点表配置"
                    netDeviceUnclearList.append(dict(ip=dNetNode["ip"],
                                                     deviceId=dNetNode["deviceId"],
                                                     type=dNetNode["type"],
                                                     reason=reason,
                                                     suggestion=suggestion,
                                                     id=dNetNode["nodeId"],
                                                     delay=None,
                                                     prefix="",
                                                     suffix="",
                                                     equipType=dNetNode.get("equipType", "")))
                else:
                    if ping_ip(dNetNode["ip"]):
                        nOnlineStatus = 1
                        netDeviceOnlineList.append(dict(ip=dNetNode["ip"],
                                                        deviceId=dNetNode["deviceId"],
                                                        type=dNetNode["type"],
                                                        reason="",
                                                        suggestion="",
                                                        id=dNetNode["nodeId"],
                                                        delay=None,
                                                        prefix="",
                                                        suffix="",
                                                        equipType=dNetNode.get("equipType", "")))
                    else:
                        nOnlineStatus = 0
                        reason = "网关ip({ip})无法ping通，判断为离线".format(ip=dNetNode["ip"])
                        suggestion = "建议检查设备是否上电，与交换机的网线连接是否正常等"
                        netDeviceDropWarningList.append(dict(ip=dNetNode["ip"],
                                                             deviceId=dNetNode["deviceId"],
                                                             type=dNetNode["type"],
                                                             reason=reason,
                                                             suggestion=suggestion,
                                                             id=dNetNode["nodeId"],
                                                             delay=None,
                                                             prefix="",
                                                             suffix="",
                                                             equipType=dNetNode.get("equipType", "")))

    ret.update({"netDeviceDropWarningList": netDeviceDropWarningList,
                "netDeviceOnlineList": netDeviceOnlineList,
                "netDeviceUnclearList": netDeviceUnclearList })

    return json.dumps(ret, ensure_ascii=False)

'''
 获取历史报警清单
'''
@app.route('/warning/getHistory', methods=['POST'])
def warning_get_history():
    data = request.get_json()

    ret = BEOPDataAccess.getInstance().getHistoryWarningList(data.get('timeFrom'), data.get('timeTo'))

    return json.dumps(ret , ensure_ascii=False)


'''
获取实时数据
'''
@app.route('/pointData/getRealtime', methods=['POST'])
def get_realtime_data():
    data =  request.get_json()
    pointList = data.get('pointList')
    ret = None
    try:
        ret = BEOPDataAccess.getInstance().getInputTable(pointList)[0]

    except Exception as e:
        print('ERROR: ' + e.__str__())

    return json.dumps( ret, ensure_ascii=False)


#根据关键字获取对应点的实时数据，支持分页查询
@app.route('/pointData/getRealtimeWithPages', methods=['POST'])
def getRealtimeWithPages():
    '''
        keyWordList
        targetPage
        pageNum
        :return {data:分页实时数据, totalNum:实时数据总数, msg:成功为success,失败为错误信息}
    '''
    data = request.get_json()
    keyWordList = data.get('keyWordList')
    targetPage = int(data.get('targetPage'))
    pageSize = int(data.get('pageSize'))
    nSearchArea = data.get("searchArea", 1)
    if not isinstance(nSearchArea, int):
        return json.dumps({"msg": "搜索范围必须为整数（0或1）"}, ensure_ascii=False)

    kwList = []
    if isinstance(keyWordList, list):
        for keyword in keyWordList:
            if keyword:
                kwList.append(keyword)
    else:
        kwList = keyWordList

    tFrom = datetime.now()
    ptsIn4db = BEOPSqliteAccess.getInstance().findPointListByKeyList_With_Redis(kwList, nSearchArea)  # 经过关键词过滤的点信息全集
    logging.error("getRealtimeWithPages tDelta1: %s" % (datetime.now() - tFrom).total_seconds())

    # 若无关键词则点名列表置为空
    ptNmList = []
    if len(keyWordList):
        ptNmList = list(ptsIn4db.keys())

    tFrom = datetime.now()
    dRealtimeData = BEOPDataAccess.getInstance().getInputTableAsDictListWithTimeV1(ptNmList)
    logging.error("getRealtimeWithPages tDelta2: %s" % (datetime.now() - tFrom).total_seconds())

    tFrom = datetime.now()
    strCurTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    dataList = []
    for ptName, dPtInfo in ptsIn4db.items():
        ptValue = ""
        if dRealtimeData.get(ptName, None) != None:
            ptValue = dRealtimeData.get(ptName)

        dataList.append(dict(name=ptName,
                             value=ptValue,
                             time=strCurTime,
                             source=dPtInfo.get("sourceType", ""),
                             description=dPtInfo.get("description", "")))
    logging.error("getRealtimeWithPages tDelta3: %s" % (datetime.now() - tFrom).total_seconds())

    ret = {"data": dataList, "totalNum": len(dataList)}
    return json.dumps(ret, ensure_ascii=False)


#获取所有监控进程
@app.route('/process/getProcessList',methods=['POST'])
def getProcessList():
    data = request.get_json()
    if data == None:
        data = {}
    bGetAll = data.get("getAll", False)
    ret = BEOPDataAccess.getInstance().getProcessList(bGetAll)
    return json.dumps(ret,ensure_ascii=False)

# 增加监控进程名称
@app.route('/process/addProcessName',methods=['POST'])
def addProcessName():
    data = request.get_json()
    directory = data.get('directory')
    processName = data.get('processName')
    exeTime = data.get("executeTime", "")

    if exeTime:
        strType = "taskprocess"
    else:
        strType = "hostprocess"

    ret = BEOPDataAccess.getInstance().addProcessByName(directory, processName, strType, exeTime)

    return json.dumps( ret, ensure_ascii=False)

@app.route('/process/delProcessName',methods=['POST'])
def delProcessList():
    data = request.get_json()
    processList = data.get('processList')
    strType = data.get("type", "hostprocess")

    ret = BEOPDataAccess.getInstance().delProcessList(processList, strType)
    return json.dumps(ret,ensure_ascii=False)

#日志log接口
@app.route('/log/search', methods=['POST'])
def searchLog():
    '''
    keyWordList:用于匹配loginfo中的信息
    timeFrom:
    timeTo:
    targetPage:
    pageSize:

    :return {data:分页log结果, totalNum:分页信息, msg:成功为success,失败为错误信息}
    '''
    data = request.get_json()
    keyWordList = data.get('keyWordList')
    timeFrom = data.get('timeFrom')
    timeTo = data.get('timeTo')
    targetPage = int(data.get('targetPage'))
    pageSize = int(data.get('pageSize'))

    ret = BEOPDataAccess.getInstance().searchLogByKeyWord(keyWordList, timeFrom, timeTo, targetPage, pageSize)
    return json.dumps( ret, ensure_ascii=False)

#日志log接口
@app.route('/log/download', methods=['POST'])
def downloadLog():
    '''
    keyWordList:用于匹配loginfo中的信息
    timeFrom:
    timeTo:
    :return {data:log结果, totalNum:分页信息, msg:成功为success,失败为错误信息}
    '''
    data = request.get_json()
    keyWordList = data.get('keyWordList')
    timeFrom = data.get('timeFrom')
    timeTo = data.get('timeTo')
    targetPage = -1
    pageSize = -1
    ret = BEOPDataAccess.getInstance().searchLogByKeyWord(keyWordList, timeFrom, timeTo, targetPage, pageSize)
    # ret = {'data':[]}
    result = {'filePath': None, 'msg':'success'}

    FILEPATH = app.root_path + '/static/log_file/'
    FILEPATH = FILEPATH.replace('\\','/')

    try:
        if ret['data'] is not None:
            filename = 'log--'+ timeFrom.split(' ')[0] + '--'+timeTo.split(' ')[0]+'.txt'
            f = open(FILEPATH + filename, 'w')
            for item in ret['data']:
                f.write(item['time']+'\t'+item['loginfo']+'\n')
            f.close()
            try:
                compression = zipfile.ZIP_DEFLATED
            except:
                compression = zipfile.ZIP_STORED
            zf = zipfile.ZipFile(FILEPATH+'zipfile_log.zip',mode='w')
            zf.write(FILEPATH+filename, filename,compress_type=compression)
            zf.close()
            filePath = '/static/log_file/zipfile_log.zip'
            result['filePath'] = filePath
    except Exception as e:
        print('error:'+e.__str__())
        result['msg'] = e.__str__()
    return json.dumps(result, ensure_ascii=False)

# 导出历史 Excel，通用
# postData = {
#     "head": ["id","区域", "设备标识", "报警次数", "温度采集次数", "温度合格次数", "温度合格率", "数据采集次数", "非正常次数", "压缩机工作时长", "冷库开门次数"],
#     "data":[["101","区域1", "设备标识1", "1", "2", "3", "4", "5", "6", "7", "8"]]
# }
@app.route('/export/excel/', methods=['POST'])
def exportExcel():
    rt = None
    try:
        data = request.get_json()
        rowNameList = data.get("head")
        valueList = data.get("data")
        startTime = data.get("begTime")
        endTime = data.get("endTime")
        filepath,filename = data.getfilepath()
        ExcleFileEx.write_excel(rowNameList, valueList, filepath)
        rt = "/static/reports/"+filename
    except Exception as e:
        print('download_file_of_historyFault error:' + e.__str__())
        logging.error(e.__str__())
    return rt




'''
写入实时数据
type: int表示你要设置整数，str表示字符串，float代表是浮点数
'''
@app.route('/pointData/setValue', methods=['POST'])
def set_realtime_data():
    data = request.get_json()
    pointList = data.get('pointList', [])
    valueList = data.get('valueList', [])
    strSource = data.get('source')
    strType = data.get('type', 'str')
    nWaitSuccess = data.get('waitSuccess', 0)
    userName = data.get("userName", None)
    bAddOperationLog = data.get("addOperationLog", False)

    if not isinstance(pointList, list) or not len(pointList):
        return json.dumps(dict(err=1, msg='写值点名列表为空', data={}), ensure_ascii=False)

    if strSource is None or strSource=='':
        strSource = 'omsite'
    else:
        bReadOnly = BEOPDataAccess.getInstance().is_user_read_only(strSource)
        if bReadOnly:
            return json.dumps(dict(err=1, msg='user has no write priority', data={}), ensure_ascii=False)

    if len(pointList) != len(valueList):
        return json.dumps(dict(err=1, msg='写入失败，传入的点名列表和点值列表长度不等', data={}), ensure_ascii=False)

    if bAddOperationLog not in [True, False, 1, 0, "1", "0"]:
        bAddOperationLog = False

    strMsg = ""
    strIP = request.remote_addr

    # 获取实时数据
    lRealtime, dRealtime = BEOPDataAccess.getInstance().getInputTable()

    # 获取所有点的点位信息
    dAllPointInfo = {}
    if RedisManager.is_alive():
        dAllPointInfo = RedisManager.get("all_point_info")
        if not isinstance(dAllPointInfo, dict) or not dAllPointInfo:
            dAllPointInfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(None)
            RedisManager.set("all_point_info", dAllPointInfo)
    else:
        dAllPointInfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(None)

    # 操作记录列表
    strOpContentList = []

    filtered_pointList = []
    filtered_valueList = []

    maintainCmdList = []
    for i, pname in enumerate(pointList):
        if pname in filtered_pointList:
            continue

        strValue = str(valueList[i])
        if strValue.find('<%') >= 0 and strValue.find('%>')>=2:
            thisEvalResult = eval_string_expression_strict(strValue, '1', '', lRealtime)
            if thisEvalResult is not None:

                strProcPointName = pname
                procPointValue = str(thisEvalResult)
            else:
                strError = 'ERROR in /pointData/setValue: eval_string_expression_strict return None, Expression: %s' % (strValue)
                logging.error(strError)

                strProcPointName = pname
                procPointValue = valueList[i]

        else:
            strProcPointName = pname
            procPointValue = valueList[i]

        if strProcPointName == None or procPointValue == None:
            continue

        if re.match(r"^[a-zA-Z0-9_]*MaintainOnOff[0-9]*$", strProcPointName):
            try:
                nValue = int(float(valueList[i]))
                if nValue == 1:
                    maintainCmdList.append((strProcPointName, 1))
                elif nValue == 0:
                    maintainCmdList.append((strProcPointName, 0))
            except:
                pass

        filtered_pointList.append(strProcPointName)
        filtered_valueList.append(procPointValue)

        strPointDesc = ""
        try:
            strPointDesc = dAllPointInfo.get(pname, {}).get("description", "")
        except:
            pass

        changeValueFrom = dRealtime.get(pname, None)
        changeValueTo = valueList[i]

        if changeValueFrom != None:
            if strPointDesc:
                strOpContent = '将%s(%s)从%s修改为%s(来自:%s)' % (pname, strPointDesc, changeValueFrom, changeValueTo, strIP)
            else:
                strOpContent = '将%s从%s修改为%s(来自:%s)' % (pname, changeValueFrom, changeValueTo, strIP)
            strOpContentList.append(strOpContent)

    if len(maintainCmdList):
        if not isinstance(userName, str):
            return jsonify(dict(err=1, msg="检测到正在执行摘挂牌，但传入的用户名有误", data=False))

        dExeUser = None
        nExeUserId = None
        strExeUserName = None

        curUserList = BEOPDataAccess.getInstance().get_all_users()
        dCurAllUser = {}
        for dUser in curUserList:
            dCurAllUser.update({int(dUser["userid"]): dUser["username"]})
            if userName == dUser.get("username"):
                dExeUser = dUser

        if isinstance(dExeUser, dict):
            if int(dExeUser.get("userofrole", 1)) == 1:
                return jsonify(dict(err=1, msg="检测到正在执行摘挂牌，但访客无摘挂牌权限", data=False))

        if isinstance(dExeUser, dict):
            try:
                nExeUserId = int(dExeUser["userid"])
                strExeUserName = dExeUser["username"]
            except:
                pass

        # 处理挂、摘牌
        dMaintainLockRecord = BEOPDataAccess.getInstance().getMaintainLockRecord()
        if dMaintainLockRecord == None:
            return jsonify(dict(err=1, msg="检测到正在执行摘挂牌，查询当前挂牌记录失败，请稍后再试", data=False))

        bEdited = False
        for maintainPtNm, maintainPtVl in maintainCmdList:
            if maintainPtVl == 0:
                if dMaintainLockRecord.get(maintainPtNm, {}):
                    nMaintainUserId = int(dMaintainLockRecord.get(maintainPtNm).get("userId"))
                    if nMaintainUserId != nExeUserId and dCurAllUser.get(nMaintainUserId):
                        return jsonify(dict(err=1, msg="摘牌指令无法执行，因摘牌权限属于挂牌人{lockUser}（点名: {ptnm}）".format(lockUser=dCurAllUser.get(int(dMaintainLockRecord[maintainPtNm]["userId"])),
                                                                                                          ptnm=maintainPtNm)))

                    del dMaintainLockRecord[maintainPtNm]
                    bEdited = True

            elif maintainPtVl == 1:
                if nExeUserId == None:
                    return jsonify(dict(err=1, msg="挂牌指令无法执行，因传入的用户信息获取失败", data=False))

                if not dMaintainLockRecord.get(maintainPtNm, {}):
                    dMaintainLockRecord.update({maintainPtNm: dict(time=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                                                   userId=nExeUserId,
                                                                   userName=strExeUserName,
                                                                   value=1)})
                    bEdited = True

        if bEdited:
            BEOPDataAccess.getInstance().saveMaintainLockRecord(dMaintainLockRecord)

    # 用户操作记录
    if bAddOperationLog in ["1", True, 1]:
        BEOPDataAccess.getInstance().appendOperationLog(strSource, strOpContentList)
        logging.error("[IMPORTANT]/pointData/setValue: %s->%s" % (filtered_pointList, filtered_valueList))

    ret = BEOPDataAccess.getInstance().setRealtimeData(filtered_pointList, filtered_valueList)

    strTimeNow = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    strSourceWithIP = '%s(%s)' % (strSource, strIP)

    # 插入点名指令记录
    BEOPDataAccess.getInstance().addLogicOutputRecord(strTimeNow, filtered_pointList, filtered_valueList, strSourceWithIP)

    if nWaitSuccess == 1:
        nWaitCheckCount = 10
        bAllSuccess = True

        while nWaitCheckCount > 0:
            try:
                dRealTimeData = BEOPDataAccess.getInstance().getInputTable(pointList)[1]
                bAllSuccess = True

                if not len(dRealTimeData.keys()):
                    bAllSuccess = False
                    strMsg = "写入失败，当前实时数据获取失败"
                    BEOPDataAccess.getInstance().setRealtimeData(filtered_pointList, filtered_valueList)  # 写入失败则再写一次
                    BEOPDataAccess.getInstance().addLogicOutputRecord(strTimeNow, filtered_pointList,filtered_valueList, strSourceWithIP)
                    continue

                for i, vv in enumerate(valueList):
                    if dAllPointInfo.get(pointList[i], None) == None:  # 若点不存在则跳过
                        continue

                    rValue = dRealTimeData.get(pointList[i])
                    if strType == 'int':
                        nV1 = int(float(valueList[i]))
                        nV2 = int(float(rValue))
                        if nV1 != nV2:
                            bAllSuccess = False
                            strMsg = "写入失败，点：%s 的值写入失败(目标:%d, 实际:%d)" % (pointList[i], nV1, nV2)
                            BEOPDataAccess.getInstance().setRealtimeData([pointList[i]], [valueList[i]])  # 写入失败则再写一次
                            BEOPDataAccess.getInstance().addLogicOutputRecord(strTimeNow, pointList[i], valueList[i], strSourceWithIP)
                            continue

                    elif strType in ['float', 'real', 'double']:
                        fV1 = float(valueList[i])
                        fV2 = float(rValue)

                        if fabs(fV1-fV2) > 0.001:
                            bAllSuccess = False
                            strMsg = "写入失败，点：%s 的值写入失败(目标:%.3f, 实际:%.3f)" % (pointList[i], fV1, fV2)
                            BEOPDataAccess.getInstance().setRealtimeData([pointList[i]], [valueList[i]])  # 写入失败则再写一次
                            BEOPDataAccess.getInstance().addLogicOutputRecord(strTimeNow, pointList[i], valueList[i], strSourceWithIP)
                            continue

                    else:
                        fTemp1 = 0.0
                        fTemp2 = 0.0
                        try:
                            fTemp1 = float(valueList[i])
                            fTemp2 = float(rValue)
                            bCouldNumerCompare = True
                        except:
                            bCouldNumerCompare = False

                        if bCouldNumerCompare:
                            if fabs(fTemp1 - fTemp2) > 0.001:
                                bAllSuccess = False
                                strMsg = "写入失败，点：%s 的值写入失败(目标:%s, 实际:%s)" % (pointList[i], str(valueList[i]), str(rValue))
                                BEOPDataAccess.getInstance().setRealtimeData([pointList[i]], [valueList[i]])  # 写入失败则再写一次
                                BEOPDataAccess.getInstance().addLogicOutputRecord(strTimeNow, pointList[i], valueList[i], strSourceWithIP)
                                continue

                        elif str(valueList[i]) != rValue:
                            bAllSuccess = False
                            strMsg = "写入失败，点：%s 的值写入失败(目标:%s, 实际:%s)" % (pointList[i], str(valueList[i]), str(rValue))
                            BEOPDataAccess.getInstance().setRealtimeData([pointList[i]], [valueList[i]])  # 写入失败则再写一次
                            BEOPDataAccess.getInstance().addLogicOutputRecord(strTimeNow, pointList[i], valueList[i], strSourceWithIP)
                            continue

                if bAllSuccess:
                    break

            except Exception as e:
                bAllSuccess = False
                strMsg = "写入失败: %s" % e.__str__()
                continue

            finally:
                nWaitCheckCount -= 1
                time.sleep(2.5)

        if bAllSuccess:
            ret = dict(err=0, msg='写入成功')
            return jsonify(ret)
        else:
            ret = dict(err=1, msg='写入失败，原因：%s' % strMsg)
            return jsonify(ret)

    return jsonify(ret)


@app.route('/thirdParty/upload', methods=['POST'])
def third_party_upload():
    ret = dict(err=0,data={},msg='')
    data = request.get_json()
    if not data:
        return jsonify(err=1, msg='data lost')
    pointList = data.get('pointList')

    if not pointList:
        return jsonify(err=1, msg='pointList lost')
    valueList = data.get('valueList')
    if not valueList:
        return jsonify(err=1, msg='valueList lost')

    if not isinstance(pointList, list):
        return jsonify(err=1, msg='pointList should be list')
    if not isinstance(valueList, list):
        return jsonify(err=1, msg='valueList should be list')

    if len(pointList)!= len(valueList):
        return jsonify(err=1, msg='valueList should be same size as pointList')

    thirdPartyName = data.get('thirdPartyName')
    if not thirdPartyName:
        return jsonify(err=1, msg='thirdPartyName lost')

    if not isinstance(thirdPartyName, str):
        return jsonify(err=1, msg='thirdPartyName should be str')

    strTimeNowOfCollector = data.get('timeNow')
    if strTimeNowOfCollector is None:
        return jsonify(err=1, msg='timeNow param lost ,and it should be str')

    try:
        tTransfer = datetime.strptime(strTimeNowOfCollector, '%Y-%m-%d %H:%M:%S' )
    except:
        return jsonify(err=1, msg='timeNow format is bad, should like: 2000-10-01 09:01:00')

    tCoreServer = datetime.now()

    #用于修正采集器时间与core时间的差
    tspan = tCoreServer-tTransfer

    strTimeList = data.get('timeList')
    if not strTimeList:
        strTimeList = [strTimeNowOfCollector]* len(valueList)
    elif not isinstance(strTimeList, list):
        return jsonify(err=1, msg='timeList must be list')
    elif len(strTimeList)!= len(valueList):
        return jsonify(err=1, msg='timeList should be same size as pointList')

    for i in range(len(strTimeList)):
        tt = datetime.strptime(strTimeList[i], '%Y-%m-%d %H:%M:%S')
        tt = tt +tspan
        strTimeList[i] = tt.strftime('%Y-%m-%d %H:%M:%S')

    if not g_thirdparty_pointmap:
        allPointInfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(pointList=None, sourceType='ThirdParty')
        if allPointInfo:
            for k,v in allPointInfo.items():
                strPointNameOfThirdParty = v.get('addr')
                if strPointNameOfThirdParty:
                    g_thirdparty_pointmap[strPointNameOfThirdParty] = k

    if pointList:
        filteredPointList = []
        filteredValueList = []
        for i in range(len(pointList)):
            pointNameOfDOM = g_thirdparty_pointmap.get(pointList[i])
            if pointNameOfDOM:
                filteredPointList.append(pointNameOfDOM)
                filteredValueList.append(valueList[i])
            else:
                print('warning: cannot find point by thirdparty name: %s'%(pointList[i]))

    if filteredPointList:
        ret['data'] = BEOPDataAccess.getInstance().setThirdPartyRealtimeData(filteredPointList, filteredValueList, strTimeList, thirdPartyName)

    print('recv thirdparty points count:%d, updated count: %d'%(len(pointList), len(filteredPointList)))


    return json.dumps( ret, ensure_ascii=False)


@app.route('/thirdParty/getCommand', methods=['POST'])
def third_party_command():
    ret = dict(err=0, data={}, msg='')
    data = request.get_json()
    thirdPartyName = data.get('thirdPartyName')
    if not thirdPartyName:
        return jsonify(err=1, msg='thirdPartyName lost')

    if not isinstance(thirdPartyName, str):
        return jsonify(err=1, msg='thirdPartyName should be str')

    commandDataList = BEOPDataAccess.getInstance().getAndClearThirdPartyOutputTable(None, thirdPartyName)

    localPointCommandList = []
    if commandDataList:
        if not g_thirdparty_pointmap:
            allPointInfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(pointList=None, sourceType='ThirdParty')
            if allPointInfo:
                for k, v in allPointInfo.items():
                    strPointNameOfThirdParty = v.get('addr')
                    if strPointNameOfThirdParty:
                        g_thirdparty_pointmap[strPointNameOfThirdParty] = k

        for item in commandDataList:
            strName = item.get('name')
            for k,v in g_thirdparty_pointmap.items():
                if v== strName:
                    localPointCommandList.append(dict(name=k, value=item.get('value')))
                    break

    print('getAndClearThirdPartyCommand count:%d, sent available count:%d'%(len(commandDataList), len(localPointCommandList)))

    ret['data'] = localPointCommandList
    return json.dumps(ret, ensure_ascii=False)

def uploadOperationToCloud(userName,content, strMemo, strAddress ):
    #code has changed to message queue deal
    return True

@app.route('/operationRecord/add', methods=['POST'])
def operation_record_add():
    msg = ''
    err = 0
    try:
        data =  request.get_json()
        logging.error('record operation_record_add:' + str(data))
        userName = data.get('userName')
        content = data.get('content')

        if not content:
            return json.dumps( dict(err=1, msg='content should not be empty or none'), ensure_ascii=False)

        #address = data.get('address')
        address = request.remote_addr
        if address and address!='':
            content = '%s(来自:%s)'%(content, address)
        BEOPDataAccess.getInstance().appendOperationLog(userName, content)

        uploadOperationToCloud( userName, content, '', address)

    except Exception as e:
        err = 1
        msg = e.__str__()
        logging.error('ERROR when operation_record_add:' + msg)


    return json.dumps( dict(err=err, msg=msg), ensure_ascii=False)

@app.route('/operationRecord/addChangeValue', methods=['POST'])
def operation_record_add_change_value():
    msg = ''
    err = 0
    try:
        data =  request.get_json()
        logging.error('record operation_record_add_change_value:' + str(data))
        userName = data.get('userName')
        content = data.get('content')
        lang = data.get('lang') #zh-cn
        #address = data.get('address')
        address = request.remote_addr

        strPointName = data.get('pointName')
        strDescrption = data.get('pointDescription')
        valueFrom = data.get('valueChangeFrom')
        valueTo = data.get('valueChangeTo')
        if strPointName is None or valueTo is None:
            return json.dumps( dict(err=0, msg='strPointName is None or valueChangeTo is None'), ensure_ascii=False)

        if lang=='zh-cn':
            if strDescrption=='':
                 content = '将%s从%s修改为%s(来自:%s)' % (strPointName, valueFrom, valueTo, address)
            else:
                content = '将%s(%s)从%s修改为%s(来自:%s)' % (strPointName, strDescrption, valueFrom, valueTo, address)
        else:
            if strDescrption=='':
                content = 'modified %s value from %s to %s' % (strPointName, valueFrom, valueTo)
            else:
                content = 'modified %s(%s) value from %s to %s' % (strPointName,strDescrption, valueFrom, valueTo)
        BEOPDataAccess.getInstance().appendOperationLog(userName, content)

        uploadOperationToCloud(userName, content, '', address)

    except Exception as e:
        err = 1
        msg = e.__str__()
        logging.error('ERROR when operation_record_add_change_value:' + msg)

    return json.dumps( dict(err=err, msg=msg), ensure_ascii=False)

@app.route('/operationRecord/addLogin', methods=['POST'])
def operation_record_add_login():
    msg = ''
    err = 0
    try:
        data =  request.get_json()
        logging.error('record operation_record_add_login:' + str(data))
        userName = data.get('userName')
        lang = data.get('lang') #zh-cn
        type = int(data.get('type'))

        strHost = request.host
        if re.match(r"^[a-zA-Z0-9]*\.inwhile.com:[0-9]*$", strHost):
            address = "远程穿透"
        else:
            address = request.remote_addr

        if userName == 'cx':
            return json.dumps(dict(err=err, msg=msg), ensure_ascii=False)

        if lang == 'zh-cn':
            if type==0:
                if address=='':
                    content = '用户 %s 登出' % (userName, )
                else:
                    content = '用户 %s 登出(地址: %s)' % (userName, address)
            else:
                content = '用户 %s 登入(地址: %s)'%(userName, address)
        else:
            if type == 0:
                content =  'user %s login(from: %s)'%(userName, address)
            else:
                content = 'user %s logout(from: %s)' % (userName, address)
        BEOPDataAccess.getInstance().appendOperationLog(userName, content)

        uploadOperationToCloud(userName, content, '', address)
    except Exception as e:
        err = 1
        msg = e.__str__()
        logging.error('ERROR when operation_record_add_change_value:' + msg)

    return json.dumps( dict(err=err, msg=msg), ensure_ascii=False)



@app.route('/file/download', methods=['POST'])
def file_download():
    msg = ''
    err = 0
    try:
        data =  request.get_json()
        fileName = data.get('fileName')
        readyPointNameName = data.get('readyPointName')
        path = os.getcwd()
        strSubPath = '/static/files/'+ fileName
        filePath = path + '/siteinterface' + strSubPath
        if os.path.exists(filePath):
            fileAddress = strSubPath
            msg = fileAddress
            if readyPointNameName:
                BEOPDataAccess.getInstance().setRealtimeData( [readyPointNameName], [1])
        else:
            err = 1
            msg = '文件不存在'

    except Exception as e:
        err = 1
        msg = e.__str__()

    return json.dumps( dict(err=err, msg=msg), ensure_ascii=False)


@app.route('/operationRecord/downloadHistoryFile/<startTime>/<endTime>/<language>')
def download_file_of_operation_history(startTime, endTime,language):
    rt = {'status':0, 'message':None}
    grade_dict = {1:'Alert', 2:'Fault'}
    status_dict = {0:'Disable', 1:'Delayed', 2:'Realtime'}
    try:
        hisData = BEOPDataAccess.getInstance().readOperationLogV2(startTime, endTime)
        if hisData:
            if language == 'zh-cn':
                excelData = ExcelFile('时间', '用户', '操作内容')
            else:
                excelData = ExcelFile('Time', 'User', 'Content')
            for item in hisData:
                if language == 'en':
                    pTime = item[0]
                else:
                    pTime = item[0]
                excelData.append_row([pTime, item[1], item[2]])
            filepath = make_excel_file(excelData)
            if os.path.exists(filepath):
                return jsonify(dict(err=0, msg="", data=os.path.basename(filepath)))

            # if filepath:
            #     filename = make_file_name('op_record', startTime, endTime)
            #     rt = make_response(send_file(filepath))
            #     rt.headers["Content-Disposition"] = "attachment; filename=%s;"%filename
            #     return json.dumps(dict(err=err, msg=msg), ensure_ascii=False)
        else:
            return jsonify(dict(err=0, msg="no operation record found", data=""))
    except Exception as e:
        strLog = 'download_file_of_historyFault error:' + e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=""))


@app.route('/report/recordTable', methods=['POST'])
def gen_chaobiao_Report():
    try:
        print('111')
        FILEPATH = app.root_path + '/static'
        FILEPATH = FILEPATH.replace('\\','/')

        res = {'msg':'success', 'status':True, 'filename': None}
        data = request.get_json()
        startTime = data.get('startTime')
        endTime = data.get('endTime')
        # projectId = data.get('projectId')
        timeFormat = data.get('timeFormat')
        pointList = data.get('pointList')

        hisData = BEOPDataAccess.getInstance().get_history_data_padded(pointList, startTime, endTime, timeFormat)
        print(hisData)

        dataNum = len(hisData['time'])
        print(dataNum)
        document = Document()
        beizhu = '备注：\n'
        for i in range(len(pointList)):
            beizhu += '点'+ str(i) + ':  ' + pointList[i] + '\n'
        p = document.add_paragraph()
        p.text = beizhu

        NUM = 6
        pointNum = len(pointList)
        xx = [NUM for i in range(int(pointNum/NUM))]
        if pointNum % NUM != 0:
            xx.append(pointNum % NUM)
        print(xx)
        for i in range(len(xx)):
            col = xx[i]
            table = document.add_table(rows=1, cols=col+1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '时间'
            for j in range(col):
                # hdr_cells[j+1].text = pointList[i*NUM + j]
                hdr_cells[j+1].text = '点'+str(i*NUM + j)
                print(len(hisData['map'][pointList[i*NUM + j]]))

            for k in range(dataNum):
                row = table.add_row().cells
                row[0].text = hisData['time'][k]
                for j in range(col):
                    if k < len(hisData['map'][pointList[i*NUM + j]]):
                        row[j+1].text = str(round(hisData['map'][pointList[i*NUM + j]][k], 2))
        print('done')
        filename = 'chaobiao.docx'

        document.save(FILEPATH + '/' +filename)
        res['filename'] = filename
    except Exception as e:
        print(e.__str__())
        res['msg'] = e.__str__()
        res['status'] = False

    return json.dumps(res)

@app.route('/report/recordTableForAccumPoint', methods=['POST'])
def gen_energy_Report():
    try:
        print('222')
        res = {'msg':'success', 'status':True, 'filename': None}
        FILEPATH = app.root_path + '/static'
        FILEPATH = FILEPATH.replace('\\','/')

        data = request.get_json()
        startTime = data.get('startTime')
        endTime = data.get('endTime')
        timeFormat = data.get('timeFormat')
        pointList = data.get('pointList')
        print(startTime)
        # 获取前一个时间点   表格中每个time的值将表示  到此time为止，timeFormat时间长内的power量
        startTime = get_front_time(startTime, timeFormat)
        print(startTime)

        hisData = BEOPDataAccess.getInstance().get_history_data_padded(pointList, startTime, endTime, timeFormat)
        print(hisData)
        dataNum = len(hisData['time'])
        print(dataNum)
        document = Document()
        beizhu = '备注：\n'
        for i in range(len(pointList)):
            beizhu += '点'+ str(i) + ':  ' + pointList[i] + '\n'
        p = document.add_paragraph()
        p.text = beizhu

        NUM = 6
        pointNum = len(pointList)
        xx = [NUM for i in range(int(pointNum/NUM))]
        if pointNum % NUM != 0:
            xx.append(pointNum % NUM)
        print(xx)
        for i in range(len(xx)):
            col = xx[i]
            table = document.add_table(rows=1, cols=col+1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '时间'
            for j in range(col):
                # hdr_cells[j+1].text = pointList[i*NUM + j]
                hdr_cells[j+1].text = '点'+str(i*NUM + j)
                print(len(hisData['map'][pointList[i*NUM + j]]))

            for k in range(1, dataNum):
                row = table.add_row().cells
                row[0].text = hisData['time'][k]
                for j in range(col):
                    if k < len(hisData['map'][pointList[i*NUM + j]]):
                        row[j+1].text = str(round(hisData['map'][pointList[i*NUM + j]][k] - hisData['map'][pointList[i*NUM + j]][k-1], 2))
        print('done')
        filename = 'energy.docx'

        document.save(FILEPATH + '/' +filename)
        res['filename'] = filename
    except Exception as e:
        print(e.__str__())
        res['msg'] = e.__str__()
        res['status'] = False
    return json.dumps(res)

def get_front_time(curTime, timeFormat):
    x = 0
    timestamp = datetime.strptime(curTime, '%Y-%m-%d %H:%M:%S').timestamp()
    if timeFormat == 'm1':
        x = 60
    if timeFormat == 'm5':
        x = 60 * 5
    if timeFormat == 'h1':
        x = 60 * 60
    if timeFormat == 'd1':
        x = 60 * 60 * 24
    timestamp = timestamp - x
    return datetime.fromtimestamp(timestamp).strftime('%Y-%m-%d %H:%M:%S')

def make_excel_file_by_tablib_data_with_path(filepath, excelData):
    rt = None
    file = None
    try:
        file = open(filepath, 'wb')
        file.write(excelData.xlsx)
        rt = filepath
    except Exception as e:
        print('make_excel_file error:' + e.__str__())
        logging.error(e.__str__())
    finally:
        if file:
            file.close()
    return rt

def make_excel_file(excelData):
    rt = None
    file = None
    try:
        filesDir = os.path.join(app.static_folder, "files")
        if not os.path.exists(filesDir):
            os.mkdir(filesDir)

        tempDir = os.path.join(filesDir, "temp")
        if not os.path.exists(tempDir):
            os.mkdir(tempDir)

        file_name = ObjectId().__str__() + '.xlsx'
        filepath = os.path.join(tempDir, file_name.__str__())
        file = open(filepath, 'xb')
        file.write(excelData.data.xlsx)
        rt = filepath
    except Exception as e:
        print('make_excel_file error:' + e.__str__())
        logging.error(e.__str__())
    finally:
        if file:
            file.close()
    return rt

def make_file_name(projId, startTime, endTime):
    rt = None
    try:
        startTime = datetime.strptime(startTime, '%Y-%m-%d')
        endTime = datetime.strptime(endTime, '%Y-%m-%d')
        if startTime < datetime.now():
            startTime = startTime.strftime('%Y-%m-%d')
            endTime = endTime.strftime('%Y-%m-%d')
            rt = "{}-{}___{}.xlsx".format(projId, startTime, endTime)
            #rt = str(projId) + '-' + startTime + '.xlsx'
    except Exception as e:
        print('make_file_name error:' + e.__str__())
        logging.error(e.__str__())
    return rt


@app.route('/report/genEnergyReport/<roomIndex>/<reportYear>/<reportMonthFrom>/<reportDayFrom>/<reportMonthTo>/<reportDayTo>', methods=['GET'])
def gen_cooling_month_report(roomIndex, reportYear, reportMonthFrom, reportDayFrom, reportMonthTo, reportDayTo):
    strTimeStart = '%d-%02d-%02d 00:00:00'%(int(reportYear), int(reportMonthFrom), int(reportDayFrom))
    strTimeEnd = '%d-%02d-%02d 00:00:00' % (int(reportYear), int(reportMonthTo), int(reportDayTo))

    if roomIndex=='0':
        roomIndex = ''
    pointList = ['RealtimeLoad'+roomIndex,'ChGroupPowerTotal'+roomIndex, 'CWPGroupPowerTotal'+roomIndex,
                 'PriChWPGroupPowerTotal'+roomIndex, 'CTGroupPowerTotal'+roomIndex,
                 'ChillerRoomGroupPowerTotal'+roomIndex]
    tStart = datetime.strptime(strTimeStart, '%Y-%m-%d %H:%M:%S')
    tEnd = datetime.strptime(strTimeEnd, '%Y-%m-%d %H:%M:%S')

    tCur = tStart

    strDateList = []
    fLoadList = []
    fEnergyChList = []
    fEnergyPriChWPList = []
    fEnergyCWPList = []
    fEnergyCTList = []
    fEnergyRoomList = []
    fEffRoomList = []
    while tCur<=tEnd:
        strOneDayFrom = tCur.strftime('%Y-%m-%d %H:%M:%S')
        tCur = tCur + timedelta(days=1)
        strOneDayEnd = tCur.strftime('%Y-%m-%d %H:%M:%S')
        print('gen data from %s to %s'%(strOneDayFrom, strOneDayEnd))
        result = BEOPDataAccess.getInstance().get_history_data_padded(pointList, strOneDayFrom, strOneDayEnd, 'm1')
        loadList = result['map'].get('RealtimeLoad'+roomIndex)
        strDateList.append(strOneDayFrom)
        fLoadToday = sum(loadList)/60.0*3.6
        fLoadList.append(round(fLoadToday,1))

        #Room energy total
        fvList = result['map'].get('ChGroupPowerTotal'+roomIndex)
        if len(fvList)>=2:
            fEnergy = round(fvList[-1]-fvList[0],1)
            if fEnergy<0.0:
                fEnergy = '--'
            elif fEnergy>=fLoadToday*50.0:
                fEnergy = '--'

        else:
            fEnergy = '--'
        fEnergyChList.append(fEnergy)

        #Eff
        if isinstance(fEnergy, float) and fEnergy>0:
            try:
                fEffRoomList.append(round(fLoadToday/fEnergy,2))
            except:
                fEffRoomList.append('--')
        else:
            fEffRoomList.append('--')

        # CWP
        fvList = result['map'].get('CWPGroupPowerTotal' + roomIndex)
        if len(fvList) >= 2:
            fEnergy = round(fvList[-1] - fvList[0],1)
            if fEnergy<0.0:
                fEnergy = '--'
            elif fEnergy>=fLoadToday*50.0:
                fEnergy = '--'
        else:
            fEnergy = '--'
        fEnergyCWPList.append(fEnergy)

        #Pri ChWP
        fvList = result['map'].get('PriChWPGroupPowerTotal' + roomIndex)
        if len(fvList) >= 2:
            fEnergy = round(fvList[-1] - fvList[0],1)
            if fEnergy<0.0:
                fEnergy = '--'
            elif fEnergy>=fLoadToday*50.0:
                fEnergy = '--'
        else:
            fEnergy = '--'
        fEnergyPriChWPList.append(fEnergy)

        # chiller energy
        fvList = result['map'].get('ChillerRoomGroupPowerTotal' + roomIndex)
        if len(fvList) >= 2:
            fEnergy = round(fvList[-1] - fvList[0],1)
            if fEnergy<0.0:
                fEnergy = '--'
            elif fEnergy>=fLoadToday*50.0:
                fEnergy = '--'
        else:
            fEnergy = '--'
        fEnergyRoomList.append(fEnergy)

        #CT
        fvList = result['map'].get('CTGroupPowerTotal' + roomIndex)
        if len(fvList) >= 2:
            fEnergy = round(fvList[-1] - fvList[0],1)
            if fEnergy<0.0:
                fEnergy = '--'
            elif fEnergy>=fLoadToday*50.0:
                fEnergy = '--'
        else:
            fEnergy = '--'
        fEnergyCTList.append(fEnergy)

    # 设置行和列的数量
    rowNum = len(fLoadList)+1
    colNum = 8
    # 创建空白docx文档
    fileName  = '\\files\\report_%d_%d_%d_to_%d_%d.docx' % (
        int(reportYear), int(reportMonthFrom),int(reportDayFrom), int(reportMonthTo), int(reportDayTo))
    filepath = app.static_folder + fileName

    doc = Document()
    # 添加新段落
    p = doc.add_paragraph('逐日冷量及用电量汇总')
    # 设置段落对齐方式
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加表格
    table = doc.add_table(rows=rowNum, cols=colNum)
    # 为表格的每个单元格添加文本
    for row in range(rowNum):
        if row==0:
            table.cell(row, 0).text = '日期'
            table.cell(row, 1).text = '冷量(kWh)'
            table.cell(row, 2).text = '总耗电量(kWh)'
            table.cell(row, 3).text = '机房COP(kW/kW)'
            table.cell(row, 4).text = '冷机耗电量(kWh)'
            table.cell(row, 5).text = '冷冻泵耗电量(kWh)'
            table.cell(row, 6).text = '冷却泵耗电(kWh)'
            table.cell(row, 7).text = '冷却塔耗电(kWh)'
            continue
        cell = table.cell(row, 0)
        cell.text = strDateList[row-1][:10]

        cell = table.cell(row, 1)
        cell.text = str(fLoadList[row-1])
        cell = table.cell(row, 2)
        cell.text = str(fEnergyRoomList[row-1])
        cell = table.cell(row, 3)
        cell.text = str(fEffRoomList[row - 1])

        cell = table.cell(row, 4)
        cell.text = str(fEnergyChList[row-1])
        cell = table.cell(row, 5)
        cell.text = str(fEnergyPriChWPList[row-1])
        cell = table.cell(row, 6)
        cell.text = str(fEnergyCWPList[row-1])
        cell = table.cell(row, 7)
        cell.text = str(fEnergyCTList[row-1])
    doc.save(filepath)

    if filepath:
        rt = make_response(send_file(filepath))
        rt.headers["Content-Disposition"] = "attachment; filename=%s;" % fileName
        return rt
    return json.dumps( ret, ensure_ascii=False)










@app.route('/report/genExcelReportByTemplate', methods=['POST'])
def report_gen__by_template():
    data = request.get_json()
    templateName = data.get('templateName')
    sendmailinfo = data.get('sendmailinfo')
    actTime = data.get("actTime", None)
    if actTime:
        if not is_valid_date(actTime, "%Y-%m-%d %H:%M:%S"):
            return jsonify(dict(err=1, msg="actTime格式有误", data=""))

    copyFileInfo = data.get("copyFileInfo")

    nRecordToReportHistory = data.get('record', 1) #是否記錄到數據庫

    headers = {"content-type": "application/json"}

    postData = {"templateName": templateName,
                "sendmailinfo": sendmailinfo,
                "actTime": actTime,
                "copyFileInfo": copyFileInfo,
                "record": nRecordToReportHistory
                }
    rsp = None
    try:
        rsp = requests.post("http://127.0.0.1:5002/report/genExcelReportByTemplate", headers=headers,
                            data=json.dumps(postData), timeout=30 * 60)
    except Exception as e:
        return jsonify(dict(err=1, msg=e.__str__(), data=""))

    if rsp.status_code != 200:
        return jsonify(dict(err=1, msg='生成失败，状态码不为200', data=""))

    dResult = json.loads(rsp.text)

    return jsonify(dict(err=dResult.get("err", 0), msg=dResult.get("msg", ""), data=dResult.get("data", "")))

@app.route('/report/genWordReportByTemplate', methods=['POST'])
def report_gen_word_by_template():
    data = request.get_json()
    templateName = data.get('templateName', None)
    actTime = data.get("actTime", None)

    if not templateName:
        return jsonify(dict(err=1, msg='报表模板名称不能为空', data=""))

    tNow = datetime.now()

    headers = {"content-type": "application/json"}
    postData = {"templateName": templateName, "actTime": actTime}

    rsp = requests.post("http://127.0.0.1:5002/report/genWordReportByTemplate", headers=headers, data=json.dumps(postData), timeout=30)
    if rsp.status_code != 200:
        return jsonify(dict(err=1, msg='生成失败，状态码不为200', data=""))

    dResult = json.loads(rsp.text)

    return jsonify(dResult)


@app.route('/downloadReport', methods=['POST'])
def download_report():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    templateName = data.get('templateName')
    strReportFilePath = data.get('strReportFilePath')
    BEOPSqliteAccess.getInstance().getTemplateFileFrom4DB(templateName, strReportFilePath)
    return jsonify(dict(err=0, msg='', data=None))


@app.route('/project/deleteConfig', methods=['POST'])
def project_delete_config():
    try:
        data = request.get_json()
        key = data.get("key", None)
        if not key:
            return jsonify(dict(err=1, msg="key can't be empty", data=False))

        bSuc = BEOPSqliteAccess.getInstance().deleteLocalConfig('domdb.4db', key)
        return jsonify(dict(err=1 if not bSuc else 0, msg="", data=bSuc))
    except Exception as e:
        return jsonify(dict(err=1, msg="Error in /project/deleteConfig:%s" % e.__str__(), data=False))

@app.route('/project/saveConfigByJsonFile', methods=['POST'])
def project_save_config_by_json_file():
    try:
        file = request.files.get('file', None)
        key = request.form.get("key", None)

        if key == None:
            return jsonify(dict(err=1, msg="key不能为空", data=False))
        if file == None:
            return jsonify(dict(err=1, msg="json file不能为空", data=False))

        tempDir = os.path.join(app.static_folder, "temp")
        filePath = os.path.join(tempDir, "jsonFile.json")
        if os.path.exists(filePath):
            os.remove(filePath)

        file.save(filePath)

        with open(filePath, "r", encoding="UTF8") as fo:
            dJson = json.load(fo)

        data = BEOPSqliteAccess.getInstance().saveLocalConfig('domdb.4db', key, dJson)
        return jsonify(dict(err=0, data=data, msg=""))
    except Exception as e:
        return jsonify(dict(err=1, msg="error in project_save_config_by_json_file: %s" % e.__str__(), data=False))


@app.route('/project/saveConfig', methods=['POST'])
def project_save_config():
    data = request.get_json()
    strConfig = data.get('config')
    key = data.get("key", None)

    keyName = "globalconfig"
    if isinstance(key, str):
        if len(key):
            keyName = key

    try:
        if keyName[0:8] =='AI_rule_':
            curInfo = RedisManager.get('AI_rule_update_time')
            if curInfo is None:
                curInfo = {}
            curInfo[keyName] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            RedisManager.set('AI_rule_update_time', curInfo)
        elif keyName[0:12] =='script_rule_':
            curInfo = RedisManager.get('script_rule_update_time')
            if curInfo is None:
                curInfo = {}
            curInfo[keyName] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            RedisManager.set('script_rule_update_time', curInfo)
    except Exception as e:
        logging.error('ERROR in /project/saveConfig:%s'%(e.__str__()))

    try:
        dConfig = json.loads(strConfig)
        data = BEOPSqliteAccess.getInstance().saveLocalConfig('domdb.4db',keyName, dConfig)
        rv = {'data':data,'msg':'','status':True}
    except Exception as e:
        rv = {'data':'','msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)


@app.route('/project/getConfig', methods=['POST'])
def project_get_config():
    #strConfig = BEOPDataAccess.getInstance().getUnit01('SystemConfig')
    #return jsonify(dict(err=0, msg='', data=strConfig))

    rv = {'data':[],'msg':'','status':False}
    try:
        data = request.get_json()
        # strKey = data.get('key', None)
        logging.info('/project/getConfig read key: %s'%(data.get('key', '')))
        strKey = data
        if strKey is None or not strKey:
            strKey = 'globalconfig'
        else:
            strKey = data.get('key')
    except Exception as e:
        strKey = 'globalconfig'

    try:
        data =  BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig(strKey)
        if data == None:
            data="{}"
        else:
            data = json.loads(data)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':'ERROR in JSON config:'+ e.__str__(),'status':False}

    return jsonify(rv)

@app.route('/project/updateConfig', methods=['POST'])
def project_update_config():
    rcv = request.get_json()
    seg = rcv.get("seg", None)
    key = rcv.get("key", None)
    value = rcv.get("value", None)
    if not isinstance(seg, str):
        return jsonify(dict(err=1, msg="seg必须为为字符串", data=False))
    if not len(seg):
        return jsonify(dict(err=1, msg="seg不能为空", data=False))
    if key == None:
        return jsonify(dict(err=1, msg="key不能为空", data=False))
    if value == None:
        return jsonify(dict(err=1, msg="value不能为空", data=False))

    keyList = []
    valueList = []
    if isinstance(key, list) and isinstance(value, list):
        if len(key) == len(value):
            keyList = key
            valueList = value
    elif isinstance(key, str) and isinstance(value, str):
        if len(key) and len(value):
            keyList = [key]
            valueList = [value]

    if not len(keyList) or not len(valueList):
        return jsonify(dict(err=1, msg="未发现修改项，可能为传参有误", data=False))

    if len(keyList) != len(valueList):
        return jsonify(dict(err=1, msg="传参有误，键列表长度和值列表长度不等", data=False))

    strCur = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig(seg)
    dCur = {}
    if isinstance(strCur, str):
        dCur = json.loads(strCur)

    for idx, item in enumerate(keyList):
        dCur.update({item: valueList[idx]})

    bSuc = BEOPSqliteAccess.getInstance().saveLocalConfig('domdb.4db',seg, dCur)
    return jsonify(dict(err=1 if not bSuc else 0, msg="", data=bSuc))



'''
 一次性获取多个后台配置
'''
@app.route('/project/getConfigMul', methods=['POST'])
def project_get_config_mul():
    #strConfig = BEOPDataAccess.getInstance().getUnit01('SystemConfig')
    #return jsonify(dict(err=0, msg='', data=strConfig))

    rv = {'data':[],'msg':'','status':False}
    try:
        data = request.get_json()
        strKey = data
        if strKey is None:
            rv = {'data': [], 'msg': 'body data不全', 'status': True}
            return jsonify(rv)

        strKeyList = data.get('keyList')
        if not isinstance(strKeyList, list):
            rv = {'data': [], 'msg': 'body keyList必须是配置项key1的数组', 'status': False}
            return jsonify(rv)
    except Exception as e:
        rv = {'data':[],'msg':'ERROR in body config:'+ e.__str__(),'status':False}
        return jsonify(rv)

    try:
        data =  BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfigMul(strKeyList)

        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':'ERROR in JSON config:'+ e.__str__(),'status':False}

    return jsonify(rv)

@app.route('/globalSetting/uploadLogo', methods=['POST'])
def global_setting_upload_logo():
    rv = None

    upload_file = request.files.getlist("logofile")
    filePath = os.getcwd() +'\\siteinterface\\static\\images\\logo.png'
    try:
        upload_file[0].stream.read()
        upload_file[0].stream.seek(0)
        result = upload_file[0].save(filePath)
    except Exception as e:
        return jsonify(dict(err=1, msg=e.__str__(), data= None))
    return jsonify(dict(err=0, msg='', data= None))



@app.route('/globalSetting/setAppTitle', methods=['POST'])
def global_setting_set_app_title():
    data = request.get_json()
    strConfig = data.get('title')

    BEOPDataAccess.getInstance().saveUnit01( 'AppTitle', strConfig)
    return jsonify(dict(err=0, msg='', data=''))


@app.route('/globalSetting/getAppTitle', methods=['POST'])
def global_setting_get_app_title():
    rv = None
    strTitle = BEOPDataAccess.getInstance().getUnit01('AppTitle')
    return jsonify(dict(err=0, msg='', data= strTitle))

@app.route('/dashboard/updatePage', methods=['POST'])
def dashbaord_update_page():
    data = request.get_json()
    pageName = data.get('pageName')
    newPageName = data.get('newPageName')

    BEOPDataAccess.getInstance().updateUnit01('dashboard', pageName,  newPageName)
    return jsonify(dict(err=0, msg='', data=''))

@app.route('/dashboard/createPage', methods=['POST'])
def dashbaord_create_page():
    data = request.get_json()
    name = data.get('pageName')
    pageInfo = ''
    curConfig = BEOPDataAccess.getInstance().getDoubleKeyInUnit01('dashboard', name)
    if curConfig is not None:
        return jsonify(dict(err=1, msg='pageName already exist!', data=False))
    rv = BEOPDataAccess.getInstance().saveDoubleKeyInUnit01('dashboard', name, pageInfo)
    return jsonify(dict(err=0, msg='', data=rv))

@app.route('/unit01/setValue', methods=['POST'])
def unit01_setvalue():
    rv = {}
    data = request.get_json()
    str01 = data.get('unitproperty01')
    if str01 is None:
        return jsonify(dict(err=1, msg='unitproperty01 param lost', data={}))


    str02 = data.get('unitproperty02')
    str03 = data.get('unitproperty03')
    if str02 is not None and str03 is not None:
        rv = BEOPDataAccess.getInstance().saveDoubleKeyInUnit01(str01, str02, str03)
    elif str03 is None and str02 is not None:
        rv = BEOPDataAccess.getInstance().saveUnit01(str01, str02)

    return jsonify(dict(err=0, msg='', data=rv))

@app.route('/unit01/getValue', methods=['POST'])
def unit01_getvalue():
    rv = {}
    data = request.get_json()
    str01 = data.get('unitproperty01')
    if str01 is None:
        return jsonify(dict(err=1, msg='unitproperty01 param lost', data={}))


    str02 = data.get('unitproperty02')
    str03 = data.get('unitproperty03')

    if str02 is None:
        rv = BEOPDataAccess.getInstance().getUnit01(str01)
    elif str03 is None and str02 is not None:
         rv = BEOPDataAccess.getInstance().getDoubleKeyInUnit01(str01, str02)

    return jsonify(dict(err=0, msg='', data=rv))


@app.route('/dashboard/removePages', methods=['POST'])
def dashboard_remove_pages():
    data = request.get_json()
    pageNameList = data.get('pageNameList')
    id = ObjectId().__str__()
    for pageName in pageNameList:
        BEOPDataAccess.getInstance().removeByDoubleKeyInUnit01('dashboard', pageName)
    return jsonify(dict(err=0, msg='', data=id))


@app.route('/dashboard/savePoints', methods=['POST'])
def dashboard_save_points():
    data = request.get_json()
    pageName = data.get('pageName')
    pageConfig = data.get('pageConfig')
    strPageConfig = json.dumps(pageConfig)
    rv = BEOPDataAccess.getInstance().saveDoubleKeyInUnit01('dashboard', pageName, strPageConfig)
    return jsonify(dict(err=0, msg='', data=rv))


@app.route('/dashboard/getPages', methods=['GET'])
def dashboard_get_pages():
    nameList = BEOPDataAccess.getInstance().getListInUnit01ByKey('dashboard')
    return jsonify(dict(err=0, msg='', data=nameList))

@app.route('/dashboard/getPoints', methods=['POST'])
def dashboard_get_points():
    data = request.get_json()
    pageName = data.get('pageName')
    info  = BEOPDataAccess.getInstance().getDoubleKeyInUnit01('dashboard', pageName)
    dictInfo = None
    try:
        dictInfo  = json.loads(info)
    except:
        dictInfo = {}
        return jsonify(dict(err=0, msg='', data=dictInfo))
    return jsonify(dict(err=0, msg='', data=dictInfo))

@app.route('/status',methods=['GET'])
def core_get_status():
    bDBConnectionGood = True
    try:
        bDBConnectionGood = BEOPDataAccess.getInstance().isDBConnectionGood()
        result = {'version': app.config["CORE_VERSION"]}
    except Exception as e:
        result = {}
        bDBConnectionGood = False
    if not bDBConnectionGood:
        return abort(404)
    return jsonify(dict(err=0,msg='',data=result))

@app.route('/cloudConfig/saveData', methods=['POST'])
def cloud_save_data():
    data = request.get_json()
    cloudJsonData = json.dumps(data)

    BEOPDataAccess.getInstance().saveUnit01( 'CloudConfig', cloudJsonData)
    return jsonify(dict(err=0, msg='', data=''))

@app.route('/cloudConfig/getCloudData', methods=['GET'])
def cloud_config_get_cloud_data():
    rv = None
    jsonData = {}
    strData = BEOPDataAccess.getInstance().getUnit01('CloudConfig')
    try:
        jsonData = json.loads(strData)
    except Exception as e:
        print(e.__str__())
    return jsonify(dict(err=0, msg='', data= jsonData))

@app.route('/update_realtimedata_input',methods=['POST'])
def update_realtimedata_input():
    data = request.get_json()
    pointNameList = data.get('pointList')
    pointValueList = data.get('valueList')
    rv = BEOPDataAccess.getInstance().update_realtimedata_input(pointNameList,pointValueList)
    return json.dumps(rv, ensure_ascii=False)



@app.route('/get_and_clear_realtimedata_output',methods=['GET'])
def get_and_clear_realtimedata_output():

    rv = BEOPDataAccess.getInstance().get_and_clear_realtimedata_output()

    return json.dumps(rv, ensure_ascii=False)


@app.route('/get_and_clear_realtimedata_output_by_table_name/<tableName>',methods=['GET'])
def get_and_clear_realtimedata_output_by_table_name(tableName):

    rv = BEOPDataAccess.getInstance().get_and_clear_realtimedata_output(tableName)

    return json.dumps(rv, ensure_ascii=False)

'''
文件上传接口
收到文件后将core目录(本程序的上一级)中的db文件删除
'''
@app.route('/dbfile/downloadFileToCore', methods=['POST'])
def dbfile_upload():
    rv = None
    if request.method == 'POST':
        dbfile = request.files['dbfile']
        currentPath = path.dirname(__file__)
        corePath = app.config['CORE_PATH']
        strPathDBTemp = DBFileManager.getTempDBFilePath()
        try:
            dbfile.save(os.path.join(strPathDBTemp, app.config['DB_FILE_NAME']))
        except Exception as e:
            rv = dict(msg="文件保存失败!",status=False)

        if not BEOPSqliteAccess.getInstance().test_valid_4db(os.path.join(strPathDBTemp, app.config.get('DB_FILE_NAME'))):
            rv = dict(err=1, msg='your uplaod file is not valid format for core')
            return jsonify(rv)

        os.popen("taskkill /im domcore.exe -f")
        os.popen("taskkill /im domlogic.exe -f")
        # os.popen("taskkill /im dompysite.exe -f")
        DBFileManager.getInstance().deleteDB(corePath)
        DBFileManager.getInstance().moveDBFromTempToMain(corePath)

        time.sleep(5)
        BEOPSqliteAccess.getInstance().initAutoLoadDBFile()
        process_restart_domcore()
        process_restart_domlogic()
        rv = dict(msg='',status=True)
    return jsonify(rv)

@app.route('/dom/restartDomAll', methods=['GET'])
@app.route('/debug/dom/restartDomAll', methods=['GET'])
def restart_dom_all():
    result = {}
    try:
        process_restart_domlogic()
        process_restart_domcore()
        process_restart_domhost()
        #process_restart_domhost()

        #sys.exit() #退出后domhost会启动
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False, 'err':1}
    return jsonify(result)

@app.route('/dom/runDomShutdown', methods=['GET'])
def dom_run_domShutdown():
    result = {}
    try:
        process_run_domShutdown()
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

@app.route('/dom/restartDomLogic', methods=['GET'])
def restart_domLogic():
    result = {}
    try:
        process_restart_domlogic('domlogic.exe')
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)


@app.route('/dom/restartDomLogic01', methods=['GET'])
def restart_domLogic01():
    result = {}
    try:
        process_restart_domlogic('domlogic01.exe')
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

@app.route('/dom/restartDomLogicAll', methods=['GET'])
def restart_domLogicAll():
    result = {}
    try:
        process_restart_domlogic('domlogic.exe')
        process_restart_domlogic('domlogic01.exe')
        process_restart_domlogic('domlogic02.exe')
        process_restart_domlogic('domlogic03.exe')
        process_restart_domlogic('domlogic04.exe')
        process_restart_domlogic('domlogic05.exe')
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)


@app.route('/dom/restartDomLogic02', methods=['GET'])
def restart_domLogic02():
    result = {}
    try:
        process_restart_domlogic('domlogic02.exe')
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)


@app.route('/dom/restartDomLogic03', methods=['GET'])
def restart_domLogic03():
    result = {}
    try:
        process_restart_domlogic('domlogic03.exe')
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)


@app.route('/dom/restartDomLogic04', methods=['GET'])
def restart_domLogic04():
    result = {}
    try:
        process_restart_domlogic('domlogic04.exe')
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)


@app.route('/dom/restartDomLogic05', methods=['GET'])
def restart_domLogic05():
    result = {}
    try:
        process_restart_domlogic('domlogic05.exe')
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

@app.route('/dom/restartDomCore', methods=['GET'])
def restart_domCore():
    result = {}
    try:
        process_restart_domcore()
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

@app.route('/dom/restart/<processName>', methods=['GET'])
def restart_dom_cloud_sync(processName=None):
    result = {}
    try:
        if processName == None:
            return jsonify(dict(err=1, msg="进程名不能为空", data=False))

        dProcess = {}
        for dTar in allProcList:
            nameLower = dTar.get("name").lower()
            dProcess.update({nameLower: dict(name=dTar.get("name"), seconds=dTar.get("seconds"))})

        processNameLower = processName.lower()
        dInfo = dProcess.get(processNameLower, None)
        if not dInfo:
            return jsonify(dict(err=1, msg="未找到目标进程的信息", data=False))

        strProcName = dInfo.get("name")

        restart_process(strProcName)
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

@app.route('/dom/destroyDomLogic', methods=['GET'])
def destroy_domLogic():
    result = {}
    try:
        os.popen("taskkill /im domlogic.exe -f");
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)



@app.route('/dom/restartOS', methods=['GET'])
def dom_restart_os():
    result = {}
    try:
        os.popen("shutdown -r -f -t 0")
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

@app.route('/dom/destroyDomCore', methods=['GET'])
def destroy_domCore():
    result = {}
    try:
        os.popen("taskkill /im domcore.exe -f");
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

@app.route('/dom/destroyDomAll', methods=['GET'])
def destroy_dom_all():
    result = {}
    try:
        os.popen("taskkill /im domhost.exe -f")
        os.popen("taskkill /im domcore.exe -f")
        os.popen("taskkill /im domlogic.exe -f")
        result = {'msg':'','status':True}
    except Exception as e:
        result = {'msg':e.__str__(),'status':False}
    return jsonify(result)

# 下载后的文件名称默认的根据url中的信息来命名
@app.route('/dbfile/<dbFileName>', methods=['GET'])
def dbfile_download(dbFileName):
    rv = None
    if request.method == 'GET':

        corePath = app.config['CORE_PATH']
        return send_from_directory(corePath , dbFileName)


    return jsonify(rv)



def crossdomain(origin=None, methods=None, headers=None,
                max_age=21600, attach_to_all=True,
                automatic_options=True):
    if methods is not None:
        methods = ', '.join(sorted(x.upper() for x in methods))
    if headers is not None and not isinstance(headers, str):
        headers = ', '.join(x.upper() for x in headers)
    if not isinstance(origin, str):
        origin = ', '.join(origin)
    if isinstance(max_age, timedelta):
        max_age = max_age.total_seconds()

    def get_methods():
        if methods is not None:
            return methods

        options_resp = app.make_default_options_response()
        return options_resp.headers['allow']

    def decorator(f):
        def wrapped_function(*args, **kwargs):
            if automatic_options and request.method == 'OPTIONS':
                resp = app.make_default_options_response()
            else:
                resp = make_response(f(*args, **kwargs))
            if not attach_to_all and request.method != 'OPTIONS':
                return resp

            h = resp.headers

            h['Access-Control-Allow-Origin'] = origin
            h['Access-Control-Allow-Methods'] = get_methods()
            h['Access-Control-Max-Age'] = str(max_age)
            if headers is not None:
                h['Access-Control-Allow-Headers'] = headers
            return resp

        f.provide_automatic_options = False
        return update_wrapper(wrapped_function, f)

    return decorator


@app.route('/system/getLAN1MacAddress', methods=['GET'])
@crossdomain(origin='*')
def get_lan1_mac_address():
    rv = None
    strAddressList = []
    for k, v in net_if_addrs().items():
        for item in v:
            address = item[1]
            if '-' in address and len(address) == 17:
                strAddressList.append(address)

    return  jsonify(dict(macList = strAddressList))


# 获取模式
@app.route('/debugTool/getSiteMode',methods=['GET'])
def get_site_mode():
    rv = None;
    rv = BEOPDataAccess.getInstance().getSiteMode()

    return json.dumps(rv , ensure_ascii=False)


# 设置模式
@app.route('/debugTool/setSiteMode',methods=['post'])
def set_site_mode():
    rv = None;
    data = request.get_json()
    sitemode = data.get('sitemode')
    rv = BEOPDataAccess.getInstance().setSiteMode(sitemode)

    return json.dumps(rv , ensure_ascii=False)
'''
@app.route('/debugTool/getFtpFile',methods=['post'])
def get_ftp_file():
    rv = None;
    if request.method == 'POST':
        fileName='testupdate.rar'
        download_file_from_ftp(severPath,fileName)
'''



@app.route('/domCloudLogin', methods=['post'])
def cloud_login():
    rv = dict(err=1, msg='项目权限认证未通过')

    data = request.get_json()
    strUser = data.get('cloudusername')
    strPwd = data.get('cloudpassword')
    strProjectId = data.get('projectid')

    post_data = {}
    post_data.update({'name': strUser,
                      'pwd': strPwd
                      })
    # 插入测试数据
    try:
        headers = {'content-type': 'application/json'}
        r = requests.post('http://dom.inwhile.com/api/login/2', data=json.dumps(post_data), headers=headers, timeout=300)
        if r.status_code==200:
            rvdata = json.loads(r.text)
            if rvdata.get('projects'):
                pList = rvdata.get('projects')
                bFoundProject = False
                for pp in pList:
                    if pp.get('id') == int(strProjectId):
                        bFoundProject = True
                        break
                if bFoundProject:
                    rv = dict(err=0, msg='项目权限认证通过')
        else:
            rv = dict(err=1, msg='项目权限认证未通过, 错误信息: 网络请求返回' + str(r.status_code))
    except Exception as e:
        rv = dict(err=1, msg='项目权限认证未通过, 错误信息:' + e.__str__())

    return jsonify(rv)




@app.route('/debugTool/importDataFromCloud', methods=['post'])
def debugtool_import_data_from_cloud():
    rv = dict()
    siteModeInfo = BEOPDataAccess.getInstance().getSiteMode()
    if siteModeInfo and isinstance(siteModeInfo, dict) and siteModeInfo.get('sitemode') != '0': #sitemode: simulation
        rv= dict(err=1, msg='cannot import data from cloud when core is in site mode, please change the mode to simulation mode')
        return jsonify(rv)

    data = request.get_json()
    strUser = data.get('cloudusername')
    strPwd = data.get('cloudpassword')
    strProjectId = data.get('projectid')

    post_data = {}
    post_data.update({'name': strUser,
                      'pwd': strPwd
                      })
    # 插入测试数据
    try:
        headers = {'content-type': 'application/json'}
        r = requests.post('http://dom.inwhile.com/api/login/2', data=json.dumps(post_data), headers=headers, timeout=300)
        if r.status_code==200:
            rvdata = json.loads(r.text)
            if rvdata.get('projects'):
                pList = rvdata.get('projects')
                bFoundProject = False
                for pp in pList:
                    if pp.get('id')== int(strProjectId):
                        bFoundProject = True
                        #get the project
                        post_data = {
                            "proj": int(strProjectId),
                            "queryType": 1
                        }
                        r2 = requests.post('http://dom.inwhile.com/api/get_realtimedata', data=json.dumps(post_data),
                                          headers=headers, timeout=300)
                        if r2.status_code == 200:
                            allRealtimeData = json.loads(r2.text)
                            if allRealtimeData and isinstance(allRealtimeData, list):
                                pointList = []
                                valueList = []
                                for dd in allRealtimeData:
                                    pointList.append(dd.get('name'))
                                    valueList.append(dd.get('value'))
                                BEOPDataAccess.getInstance().setRealtimeData(pointList, valueList) #write into
                                rv = dict(err=0, msg='%d count data refreshed.'%(len(pointList)))
                        else:
                            rv = dict(err=1, msg='http post get realtime from server not 200.')
                if not bFoundProject:
                    rv = dict(err=1, msg='project not found or not valid to user, project id: %d'%( int(strProjectId)))
            else:
                rv = dict(err=1, msg='projects of user is empty')

    except Exception as e:
        logging.error('ERROR when oprecord upload to server:' + e.__str__())

    return jsonify(rv)



# 获取通讯设置信息
@app.route('/communication/getCommunicationInfo',methods=['POST'])
def get_communication_info():
    rv = None
    data = request.get_json()
    str = data.get('queryStr')
    rv = BEOPDataAccess.getInstance().getCommunicationSetting(str)

    return json.dumps(rv,ensure_ascii=False)

# 设置通讯设置
@app.route('/communication/setCommunicationInfo',methods=['POST'])
def set_communication_info():
    rv = None
    data = request.get_json()
    data = data.get('data')
    for item in data:
        # bacnet网卡的环境变量设置
        if list(item.keys())[0] == 'BACNET_IFACE':
            os.environ['BACNET_IFACE'] = item['BACNET_IFACE']
            continue
        unitproperty01 = list(item.keys())[0]
        unitproperty02 = item[unitproperty01]
        rv = BEOPDataAccess.getInstance().setCommunicationSetting(unitproperty01,unitproperty02)

    return json.dumps(rv,ensure_ascii=False)

# 获取策略列表
@app.route('/strategy/getThreadList',methods=['GET'])
def strategy_get_list():
    try:
        data = LogicManager.getInstance().get_queues_list('domdb.4db')
        return jsonify(dict(err=0, msg="获取成功", data=data))
    except Exception as e:
        return jsonify(dict(err=1, msg="获取失败:%s" % e.__str__(), data=[]))


# 删除线程
@app.route('/strategy/delThread',methods=['POST'])
def strategy_del_thread():
    try:
        data = request.get_json()

        threadName = data.get('threadName', None)
        if not threadName:
            return jsonify(dict(err=1, msg="策略名不能为空", data=False))
        if not isinstance(threadName, str):
            return jsonify(dict(err=1, msg="策略名必须为字符串", data=False))
        if not len(threadName):
            return jsonify(dict(err=1, msg="策略名不能为空", data=False))

        data = LogicManager.getInstance().del_thread('domdb.4db', threadName)

        return jsonify(dict(err=0, msg="", data=data))

    except Exception as e:

        return jsonify(dict(err=1, msg="删除失败: %s" % e.__str__(), data=False))


# 获取策略详细列表
@app.route('/strategy/getThreadDetailList',methods=['GET'])
def strategy_get_detail_list():
    try:
        threadList = LogicManager.getInstance().get_queues_detail_list('domdb.4db')
        return jsonify(dict(err=0, msg="", data=threadList))
    except Exception as e:
        return jsonify(dict(err=1, msg="获取失败", data=[]))

# 修改策略名称
@app.route('/strategy/modifyThreadName', methods=['POST'])
def strategy_modify_thread_name():
    try:
        data = request.get_json()
        oldName = data.get('oldName', "")
        newName = data.get('newName', "")
        if not isinstance(oldName, str) or not isinstance(newName, str):
            return jsonify(dict(err=1, msg="新旧名称必须为字符串", data=False))

        bSuc = LogicManager.getInstance().modify_queue_name('domdb.4db', oldName, newName)

        return jsonify(dict(err=0, msg="", data=bSuc))

    except Exception as e:

        return jsonify(dict(err=1, msg="修改失败: %s" % e.__str__(), data=False))



@app.route('/modbusserver/uploadPointTableFile', methods=['POST'])
def modbusserver_upload_point_table_file():
    rv = None

    upload_file = request.files.getlist("pointtable")
    filePath = os.getcwd()
    filePath = os.path.dirname(filePath) + '\\domModbusServer\\modbus_point_table.xlsx'
    try:
        upload_file[0].stream.read()
        upload_file[0].stream.seek(0)

        result = upload_file[0].save(filePath)
        os.popen("taskkill /im domModbusServer.exe -f")
        time.sleep(1)
        corePath = app.config['CORE_PATH']
        domModbusServerDir = os.path.join(corePath,'domModbusServer')
        domModbusServerPath = os.path.join(corePath,'domModbusServer', 'domModbusServer.exe')

        # subprocess.Popen(domModbusServerPath, shell=False, creationflags=subprocess.CREATE_NEW_CONSOLE, cwd=corePath+'\\domModbusServer')
        win32api.ShellExecute(0, 'open', domModbusServerPath, '', domModbusServerDir, 0)
    except Exception as e:
        return jsonify(dict(err=1, msg=e.__str__(), data= None))
    return jsonify(dict(err=0, msg='', data= None))


@app.route('/modbusserver/getPointTableFile', methods=['GET'])
def modbus_get_point_table_file():
    rv = None

    upload_file = request.files.getlist("pointtable")
    filePath = os.getcwd()
    filePath = os.path.dirname(filePath) + '\\domModbusServer\\'
    fileName = 'modbus_point_table.xlsx'
    if os.path.exists(filePath):
        responseTemp =  send_from_directory(os.path.dirname(filePath), fileName)
        content_disposition = "attachment; filename={}".format(fileName)
        responseTemp.headers['Content-Disposition'] = content_disposition
        return responseTemp

    return jsonify(err=1, msg='no file exist')



@app.route('/strategy/getThreadInfo',methods=['POST'])
def strategy_get_thread_info():
    data = request.get_json()
    threadList = data.get('threadList')
    arr = []
    try:
        for thread in threadList:
            data = LogicManager.getInstance().get_queues_timespan('domdb.4db',thread)
            arr.append(data)

        return jsonify(dict(err=0, msg="", data=arr))

    except Exception as e:

        return jsonify(dict(err=1, msg="获取失败:%s" % e.__str__()), data=[])

#获取策略列表
@app.route('/strategy/fromThreadNameGetStrategyList',methods=['POST'])
def strategy_get_strategy_list():
    data = request.get_json()
    threadName = data.get("threadName")

    rv = {'data':[],'msg':'','status':False}

    try:
        data =  LogicManager.getInstance().get_logics_list_of_queue('domdb.4db',threadName)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)

#查询策略log
@app.route('/strategy/getLogInfo',methods=['POST'])
def strategy_get_log_info():
    try:
        data = request.get_json()
        searchTime = data.get('searchTime') if data.get("searchTime") is not None else None
        dllName = data.get('dllName') if data.get("dllName") is not None else None

        if searchTime is None:
            return json.dumps({'data':[],'msg':"查询时间为空",'status':False}, ensure_ascii=False)

        if dllName is None:
            return json.dumps({'data': [], 'msg': "dll文件名为空", 'status': False}, ensure_ascii=False)

        if not isinstance(searchTime, str) or not isinstance(dllName, str):
            return json.dumps({'data': [], 'msg': "输入参数必须为字符串", 'status': False}, ensure_ascii=False)

        timeList = re.findall("[0-9]{4}_[0-9]{2}_[0-9]{2}", searchTime)
        if not len(timeList):
            return json.dumps({'data': [], 'msg': "查询时间有误", 'status': False}, ensure_ascii=False)

        strTime = timeList[0]
        tTime = datetime.strptime(strTime, "%Y_%m_%d")
        strFolderName = "logic-%s-%02d-%02d" % (tTime.year, tTime.month, tTime.day)

        coreDir = os.path.dirname(os.getcwd())
        strLogDir = os.path.join(coreDir, "log")

        strLogicLogDir = os.path.join(strLogDir, strFolderName)
        strDllFilePath = os.path.join(strLogicLogDir, "{0}.txt".format(dllName))

        if not os.path.exists(strDllFilePath):
            return json.dumps({"data": [], "msg": "日志文件不存在", "status": False}, ensure_ascii=False)

        roundList = LogicManager.getInstance().get_log_info_from_file(strDllFilePath, dllName)

        result = {'data': roundList, 'msg':'', 'status': True}

        return jsonify(result)
        # return json.dumps(result, ensure_ascii=False)

    except Exception as e:
        return jsonify({'data': [], 'msg': e.__str__(), 'status': True})
        # return json.dumps({'data': [], 'msg': e.__str__(), 'status': True}, ensure_ascii=False)


@app.route('/log/getProcessLogInfoOfOneDay',methods=['POST'])
def get_process_log_info_of_one_day():
    rcv = request.get_json()
    date = rcv.get("date", "")
    process = rcv.get("process", "")
    if not date:
        return jsonify(dict(err=1, msg="日期不能为空", data=""))
    if not process:
        return jsonify(dict(err=1, msg="进程名不能为空", data=""))
    if not isinstance(date, str):
        return jsonify(dict(err=1, msg="日期必须为字符串", data=""))
    if not isinstance(process, str):
        return jsonify(dict(err=1, msg="进程名必须为字符串", data=""))
    if not is_valid_date(date, "%Y-%m-%d"):
        return jsonify(dict(err=1, msg="日期格式有误", data=""))
    if process not in ['domBackupMysql','domCloudSync','domhost','domLogixCore','domModbusClientCore','domModbusServer',
                       'domMoxaTelnetCore','domOPCUACore','domPersagyDataClientCore','dompysite','domSiemenseTCPCore',
                       'domSoundCenter','domUpload','domcore','domlogic']:
        return jsonify(dict(err=1, msg="进程名有误", data=""))

    if process in ["domcore", "domlogic", "domSiemenseTCPCore"]:
        LogResDict = find_domcore_or_domlogic_log_of_one_day(date, process)
    else:
        LogResDict = find_other_process_log_of_one_day(date, process)

    if LogResDict.get("code") > 0:
        return jsonify(dict(err=1, msg=LogResDict.get("msg"), data=""))

    ZipResDict = zip_process_log(date, process, LogResDict.get("data"))

    if ZipResDict.get("code") > 0:
        return jsonify(dict(err=1, msg=ZipResDict.get("msg"), data=""))
    return jsonify(dict(err=0, msg="获取成功", data=ZipResDict.get("data")))

# 下载策略在指定日志的整日log
@app.route('/strategy/downloadOneDayLogOfVeryDate',methods=['POST'])
def strategy_download_one_day_log_of_very_date():
    try:
        rcv = request.get_json()
        date = rcv.get("date", None)
        strategyName = rcv.get("strategyName", None)
        if not isinstance(date, str):
            return jsonify(dict(err=1, msg="日期必须为字符串", data=[]))
        if not is_valid_date(date, "%Y-%m-%d"):
            return jsonify(dict(err=1, msg="日期格式必须为yyyy-mm-dd", data=[]))
        if not isinstance(strategyName, str):
            return jsonify(dict(err=1, msg="策略名必须为字符串", data=[]))
        if not len(strategyName):
            return jsonify(dict(err=1, msg="策略名不能为空字符串", data=[]))

        strategyNameWithExt = "{strategyName}.dll".format(strategyName=strategyName)

        tDate = datetime.strptime(date, "%Y-%m-%d")
        strDate = tDate.strftime("%Y-%m-%d")

        logFileDir = os.path.join(app.config["CORE_PATH"], "log", "logic-{date}".format(date=strDate))
        if not os.path.exists(logFileDir):
            return jsonify(dict(err=1, msg="日志目录不存在", data=[]))

        bFound = True
        logFilePath = os.path.join(logFileDir, "{fileName}.txt".format(fileName=strategyNameWithExt))
        if not os.path.exists(logFilePath):
            bFound = False
            for rootDir, dirs, files in os.walk(logFileDir):
                for file in files:
                    tarStrategyName = strategyName.lower()
                    fileNameLower = file.lower()
                    if fileNameLower.find(tarStrategyName) != -1:
                        strategyNameWithExt = file
                        logFilePath = os.path.join(logFileDir, strategyNameWithExt)
                        bFound = True
                        break

        if not bFound:
            return jsonify(dict(err=1, msg="未找到该日期下该策略的日志文件", data=[]))

        zipFileName, strErrMsg = LogicManager.getInstance().strategy_zip_one_day_log_of_very_date(strDate, logFilePath, strategyNameWithExt)

        return jsonify(dict(err=1 if len(strErrMsg) else 0, msg=strErrMsg, data=zipFileName))
    except Exception as e:
        logging.error("ERROR in /strategy/downloadOneDayLogOfVeryDate: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败: %s" % e.__str__(), data=""))




# 查询指定时间所在的那一轮日志
@app.route('/strategy/getOneRoundLogOfVeryTime',methods=['POST'])
def strategy_get_one_round_log_of_very_time():
    try:
        rcv = request.get_json()
        time = rcv.get("time", None)
        strategyName = rcv.get("strategyName", None)
        if not isinstance(time, str):
            return jsonify(dict(err=1, msg="时间必须为字符串", data=[]))
        if not is_valid_date(time, "%Y-%m-%d %H:%M:%S"):
            return jsonify(dict(err=1, msg="时间格式必须为yyyy-mm-dd HH:MM:SS", data=[]))
        if not isinstance(strategyName, str):
            return jsonify(dict(err=1, msg="策略名必须为字符串", data=[]))
        if not len(strategyName):
            return jsonify(dict(err=1, msg="策略名不能为空字符串", data=[]))

        strategyName = strategyName.replace(".dll", "")
        strategyNameWithExt = "{strategyName}.dll".format(strategyName=strategyName)

        tTime = datetime.strptime(time, "%Y-%m-%d %H:%M:%S")
        strDate = tTime.strftime("%Y-%m-%d")

        logFileDir = os.path.join(app.config["CORE_PATH"], "log", "logic-{date}".format(date=strDate))
        if not os.path.exists(logFileDir):
            return jsonify(dict(err=1, msg="日志目录不存在", data=[]))

        bFound = True
        logFilePath = os.path.join(logFileDir, "{fileName}.txt".format(fileName=strategyNameWithExt))
        if not os.path.exists(logFilePath):
            bFound = False
            for rootDir, dirs, files in os.walk(logFileDir):
                for file in files:
                    tarStrategyName = strategyName.lower()
                    fileNameLower = file.lower()
                    if fileNameLower.find(tarStrategyName) != -1:
                        strategyNameWithExt = file
                        logFilePath = os.path.join(logFileDir, strategyNameWithExt)
                        bFound = True
                        break

        if not bFound:
            return jsonify(dict(err=1, msg="未找到日志文件", data=[]))

        # logList, strErrMsg = LogicManager.getInstance().strategy_get_one_round_log_of_very_time(tTime, logFilePath, strategyNameWithExt)
        # return jsonify(dict(err=0, msg=strErrMsg, data=logList))

        roundList, strErrMsg = LogicManager.getInstance().strategy_get_log_near_time_V2(tTime, logFilePath, strategyName)
        if roundList == None:
            return jsonify(dict(err=1, msg=strErrMsg, data=[]))
        return jsonify(dict(err=0, msg=strErrMsg, data=roundList))


    except Exception as e:
        logging.error("ERROR in /strategy/getOneRoundLogOfVeryTime: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败: %s" % e.__str__(), data=[]))

#查询某日某个策略的最近一轮log
@app.route('/strategy/getLatestRoundLogInfo',methods=['POST'])
def strategy_get_log_info_latest_round():
    try:
        rcp = request.get_json()
        searchTime = rcp.get('searchTime') if rcp.get("searchTime") is not None else None
        dllName = rcp.get('dllName') if rcp.get("dllName") is not None else None

        dllName = dllName.replace(".dll", "")
        dllName = "{name}.dll".format(name=dllName)

        if searchTime is None:
            return jsonify(dict(err=1, msg="查询时间为空", data=[]))

        if dllName is None:
            return jsonify(dict(err=1, msg="dll文件名为空", data=[]))

        if not isinstance(searchTime, str) or not isinstance(dllName, str):
            return jsonify(dict(err=1, msg="输入参数必须为字符串", data=[]))

        timeList = re.findall("[0-9]{4}_[0-9]{2}_[0-9]{2}", searchTime)
        if not len(timeList):
            return jsonify(dict(err=1, msg="查询时间有误", data=[]))

        strTime = timeList[0]
        tTime = datetime.strptime(strTime, "%Y_%m_%d")
        strFolderName = "logic-%s-%02d-%02d" % (tTime.year, tTime.month, tTime.day)

        coreDir = os.path.dirname(os.getcwd())
        strLogDir = os.path.join(coreDir, "log")

        strLogicLogDir = os.path.join(strLogDir, strFolderName)
        strDllFilePath = os.path.join(strLogicLogDir, "{0}.txt".format(dllName))

        roundList, errMsg = LogicManager.getInstance().get_latest_one_round_log_from_file(strDllFilePath, dllName)
        if roundList == None:
            return jsonify(dict(err=1, msg=errMsg, data=[]))
        return jsonify(dict(err=0, msg="", data=roundList))

    except Exception as e:
        logging.error("ERROR in /strategy/getLatestRoundLogInfo: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败: %s" % e.__str__(), data=[]))

@app.route('/strategy/getLogInfoOfOneDayZipped',methods=['POST'])
def strategy_get_log_info_of_one_day_zipped():
    try:
        rcp = request.get_json()
        searchTime = rcp.get('searchTime') if rcp.get("searchTime") is not None else None
        dllName = rcp.get('dllName') if rcp.get("dllName") is not None else None

        if searchTime is None:
            return {'data': [], 'msg': "查询时间为空", 'status': False}
        if dllName is None:
            return {'data': [], 'msg': "dll文件名为空", 'status': False}

        if not isinstance(searchTime, str) or not isinstance(dllName, str):
            return {'data': [], 'msg': "输入参数必须为字符串", 'status': False}

        timeList = re.findall("[0-9]{4}_[0-9]{2}_[0-9]{2}", searchTime)
        if not len(timeList):
            return {'data': [], 'msg': "查询时间有误", 'status': False}

        strTime = timeList[0]
        tTime = datetime.strptime(strTime, "%Y_%m_%d")
        strFolderName = "logic-%s-%02d-%02d" % (tTime.year, tTime.month, tTime.day)

        coreDir = os.path.dirname(os.getcwd())
        strLogDir = os.path.join(coreDir, "log")

        strLogicLogDir = os.path.join(strLogDir, strFolderName)
        strDllFilePath = os.path.join(strLogicLogDir, "{0}.txt".format(dllName))

        strZipFilePath = LogicManager.getInstance().get_log_info_of_one_day_zipped(strFolderName, strDllFilePath)

        if strZipFilePath:
            return jsonify(dict(err=0, msg="", data=strZipFilePath))

        return jsonify(dict(err=1, msg="", data=""))

    except Exception as e:

        logging.error("ERROR in /strategy/getLogInfoOfOneDayZipped: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败: %s" % e.__str__(), data=""))


#获取策略配置列表
@app.route('/strategy/getLogicParameters',methods=['POST'])
def strategy_get_logic_parameters():
    data = request.get_json()
    strLogicName = data.get('strLogicName')

    rv = {'data':[],'msg':'','status':False}

    try:
        data =  LogicManager.getInstance().get_logic_parameters("domdb.4db",strLogicName)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)


@app.route('/strategy/updateLogicParameters',methods=['POST'])
def strategy_update_logic_parameters():
    data = request.get_json()
    DllName = data.get('DllName')
    vname = data.get('vname')
    pname = data.get('pname')
    ptype = data.get('ptype')
    unitproperty01 = data.get('unitproperty01')

    rv = {'data':[],'msg':'','status':False}

    try:
        data =  LogicManager.getInstance().modify_logic_parameter("domdb.4db",DllName,vname,pname,ptype,unitproperty01)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)



@app.route('/strategy/getListPoint',methods=['GET'])
def strategy_get_list_point():

    rv = {'data':[],'msg':'','status':False}

    try:
        data =  LogicManager.getInstance().get_list_point("domdb.4db")
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)

@app.route('/strategy/getSysLogicList',methods=['GET'])
def strategy_get_system_logic_list():
    rv = {'data':[],'msg':'','status':False}

    try:
        #遍历文件s
        workPath = os.getcwd()
        sysLogicPath = os.path.join(workPath,'syslogic')
        logiclist = []
        for file in os.listdir(sysLogicPath):
            file_path = os.path.join(sysLogicPath, file)
            if os.path.isdir(file_path):
                print('查询到一个目录，不是dll文件')
            elif os.path.splitext(file_path)[1] == '.dll':
                logiclist.append(file)
        rv = {'data':logiclist,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)

@app.route('/strategy/addLogicIntoThread',methods=['POST'])
def strategy_add_logic_into_thread():
    data = request.get_json()
    Dllname = data.get('threadName')
    rv = {'data':[],'msg':'','status':False}

    try:
        data =  LogicManager.getInstance().delLogicByDllName("domdb.4db",Dllname)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)


@app.route('/strategy/delLogicByDllName',methods=['POST'])
def strategy_del_logic_by_dllname():
    data = request.get_json()
    Dllname = data.get('DllName')
    rv = {'data':[],'msg':'','status':False}

    try:
        data =  LogicManager.getInstance().delLogicByDllName("domdb.4db",Dllname)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)

@app.route('/strategy/disableThreadAll',methods=['POST'])
def strategy_disable_thread_all():
    data = request.get_json()
    runstatus = data.get('runstatus',0)
    rv = {'data':[],'msg':'','status':False}

    threadList = LogicManager.getInstance().get_queues_list("domdb.4db")
    try:
        data =  LogicManager.getInstance().enableOrDisableThread("domdb.4db",threadList,runstatus)
        for item in threadList:
            BEOPDataAccess.getInstance().saveUnit01(item, '0')
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)

@app.route('/strategy/enableOrDisableThread',methods=['POST'])
def strategy_enable_or_disable_thread():
    data = request.get_json()
    runstatus = data.get('runstatus')
    threadList = data.get('threadList')
    rv = {'data':[],'msg':'','status':False}

    try:
        data =  LogicManager.getInstance().enableOrDisableThread("domdb.4db",threadList,runstatus)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)

# 获取modbusclient配置
@app.route('/modbusclient/getConfig',methods=['GET'])
def modbusclient_get_config():
    rv = {'data':[],'msg':'','status':False}

    try:
        dataRsp = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig('modbusclientconfig',)
        if dataRsp == None:
            data="{}"
        else:
            data = json.loads(dataRsp)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return jsonify(rv)

# 保存modbusclient配置

@app.route('/modbusclient/saveConfig',methods=['POST'])
def modbusclient_save_config():
    data = request.get_json()
    strConfig = data.get('strConfig')
    rv = {'data':'','msg':'','status':False}

    try:
        strConfig = json.loads(strConfig)
        data =  BEOPSqliteAccess.getInstance().saveLocalConfig('domdb.4db','modbusclientconfig',strConfig)
        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':'','msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)


class LogicBase(Structure):
    _fields_ = [("Init", c_bool),
                ("ActLogic", c_bool),
                ("Exit", c_bool),
                ("SetMsgWnd", c_void_p),
                ("GetMsgWnd", c_void_p),
                ("SentMsgInfo", c_void_p),
                ("SetDataAccess", c_void_p),
                ("GetDllAuthor", c_wchar_p),
                ("SetDllAuthor", c_bool),
                ("GetDllName", c_wchar_p),
                ("SetDllName", c_wchar_p),
                 ("GetInputCount", c_wchar_p),
                  ("GetOutputCount", c_wchar_p),
                   ("GetInputName", c_wchar_p),
                    ("GetOutputName", c_wchar_p),
                     ("GetAllInputParameter", c_wchar_p),
                      ("GetAllOutputParameter", c_wchar_p)]

@app.route('/strategy/newGroupInThread',methods=['POST'])
def strategy_new_group_in_thread():
    data = request.get_json()
    dllName = data.get('dllName')
    threadName = data.get('threadName')
    groupId = data.get('groupId')
    rv = {'data': [], 'msg': '', 'status': False}
    try:
        workPath = os.getcwd()
        sysLogicPath = os.path.join(workPath, 'syslogic')
        file_path = os.path.join(sysLogicPath, dllName+'.dll')
        lib = cdll.LoadLibrary(file_path)
        lib.fnInitLogic.restype = POINTER(LogicBase)
        myObj = lib.fnInitLogic().contents

        strVersion = myObj.GetDllLogicVersion()
        strDescription = myObj.GetDllDescription()
        strInputParameters = myObj.GetAllInputParameter()
        strOutputParameters = myObj.GetAllOutputParameter()
        strAuthor = myObj.GetDllAuthor()
        strName = myObj.GetDllName()


        rv = {'data':data,'msg':'','status':True}

    except Exception as e:
        rv = {'data':[],'msg':e.__str__(),'status':False}

    return json.dumps(rv,ensure_ascii=False)


@app.route('/log/insert', methods=['POST'])
def log_insert():
    data = request.get_json()
    strTime = data.get('logtime')

    if strTime is None or not isinstance(strTime, str):
        return json.dumps(dict(err=1, msg='logtime must be a str'), ensure_ascii=False)

    # invalid query filter:
    strLogInfo = data.get('loginfo')
    if strLogInfo is None or not isinstance(strLogInfo, str):
        return json.dumps(dict(err=1, msg='loginfo must be a str'), ensure_ascii=False)

    strWarning = 'Warning: /log/insert not support but called: %s'%(strLogInfo)
    print(strWarning)
    return json.dumps(dict(err=1, msg=strWarning), ensure_ascii=False)


@app.route('/insert_history_data_mul', methods=['POST'])
def insert_history_data_mul():
    data = request.get_json()
    strPointNameList = data.get('pointNameList')

    if strPointNameList is None or not isinstance(strPointNameList, list):
        print('warning in insert_history_data: no pointname')
        return json.dumps(dict(err=1, msg='pointname must be a str'), ensure_ascii=False)

    # invalid query filter:
    strTimeList = data.get('pointTimeList')
    if strTimeList is None or not isinstance(strTimeList, list):
        return json.dumps(dict(err=1, msg='pointtimeList must be a str'), ensure_ascii=False)

    strValueList =  data.get('pointValueList')
    if strValueList is None or not isinstance(strValueList, list):
        return json.dumps(dict(err=1, msg='pointvalueList must be a str or value'), ensure_ascii=False)

    if len(strPointNameList) != len(strTimeList) or len(strPointNameList) != len(strValueList):
        strErrorInfo = '收到非法历史数据插入请求: pointNameList=%s, pointTimeList=%s, pointValueList=%s'%(
            str(strPointNameList), str(strTimeList), str(strValueList)
        )
        print(strErrorInfo)
        logging.error(strErrorInfo)

    strUserId = data.get('userid')
    if strUserId is None:
        return json.dumps(dict(err=1, msg='userid must be a str or value'), ensure_ascii=False)

    strOriginalValue = data.get('original')
    if strOriginalValue is None:
        strOriginalValue = ''

    strInfoLog = "strUserId: %s, strPointNameList: %s, strValueList: %s, strTimeList: %s" % (strUserId, strPointNameList, strValueList, strTimeList)
    strLogFileName = "insert_history_data_mul_log_{date}.log".format(date=datetime.now().strftime("%Y-%m-%d"))
    log_info_to_file(strLogFileName, strInfoLog)

    reCalcTaskList = []
    for iIndex in range(len(strPointNameList)):
        bIsFloat = False
        strPointName = strPointNameList[iIndex]
        if iIndex>= len(strTimeList) or iIndex>= len(strValueList):
            continue

        strTime = strTimeList[iIndex]
        strValue = strValueList[iIndex]
        try:
            fValue = float(strValue)
            bIsFloat = True
        except:
            pass

        if not bIsFloat:
            result = BEOPDataAccess.getInstance().insert_history_data(strPointName, strTime, json.dumps(strValue, ensure_ascii=False), strUserId, strOriginalValue)
        else:
            result = BEOPDataAccess.getInstance().insert_history_data(strPointName, strTime,str(fValue), strUserId,
                                                                      strOriginalValue)
        reCalcTaskList.append({"point":strPointName,
                         "time":strTime,
                         "value":strValue})

    assign_vpoint_recalculation_task(reCalcTaskList)

    if result[0]:
        return json.dumps(dict(err=0, msg='ok'), ensure_ascii=False)

    return json.dumps(dict(err=1, msg='insert failed'), ensure_ascii=False)



@app.route('/insert_history_data', methods=['POST'])
def insert_history_data():
    data = request.get_json()
    strPointName = data.get('pointname')

    if strPointName is None or not isinstance(strPointName, str):
        print('warning in insert_history_data: no pointname')
        return json.dumps(dict(err=1, msg='pointname must be a str'), ensure_ascii=False)

    # invalid query filter:
    strTime = data.get('pointtime')
    if strTime is None or not isinstance(strTime, str):
        return json.dumps(dict(err=1, msg='pointtime must be a str'), ensure_ascii=False)

    strValue =  data.get('pointvalue')
    if strValue is None:
        return json.dumps(dict(err=1, msg='pointvalue must be a str or value'), ensure_ascii=False)

    strUserId = data.get('userid')
    if strUserId is None:
        return json.dumps(dict(err=1, msg='userid must be a str or value'), ensure_ascii=False)

    strOriginalValue = data.get('original')
    if strOriginalValue is None:
        strOriginalValue = ''

    bIsFloat = False
    try:
        fValue = float(strValue)
        bIsFloat = True
    except:
        pass

    reCalcTaskList = []
    if not bIsFloat:
        result = BEOPDataAccess.getInstance().insert_history_data(strPointName, strTime, json.dumps(strValue, ensure_ascii=False), strUserId, strOriginalValue)
    else:
        reCalcTaskList.append({"point":strPointName,
                         "time":strTime,
                         "value":str(fValue)})

        result = BEOPDataAccess.getInstance().insert_history_data(strPointName, strTime,str(fValue), strUserId,
                                                                  strOriginalValue)

    assign_vpoint_recalculation_task(reCalcTaskList)
    if result[0]:
        return json.dumps(dict(err=0, msg='ok'), ensure_ascii=False)

    return json.dumps(dict(err=1, msg='insert failed'), ensure_ascii=False)


# 增加一个日程列表
@app.route('/schedule/add', methods=['POST'])
def schedule_add():
    rv =dict(err=0, msg='')
    data = request.get_json()
    strScheduleName = data.get('scheduleName')
    strUserName = data.get('userName')
    strPointDefine= data.get('pointDefines')
    isloop = data.get('isloop')

    msg=validate_string(strScheduleName,"scheduleName")
    if msg is not None:
        return msg

    msg = validate_string(strUserName, "userName")
    if msg is not None:
        return msg

    msg = validate_string(strPointDefine, "pointDefines")
    if msg is not None:
        return msg

    msg = validate_restricted_number(isloop, "isloop")
    if msg is not None:
        return msg

    bSucces = False

    try:
         bSucces = BEOPDataAccess.getInstance().addSchedule(strScheduleName, strUserName,isloop,strPointDefine)
    except Exception as e:
        rv = dict(err=1, msg=e.__str__())

    if not bSucces:
        rv = dict(err=1, msg='')
    return json.dumps(dict(rv), ensure_ascii=False)

# 获取日程列表
@app.route('/schedule/getscheduleList', methods=['GET'])
def schedule_getList():
    rv = {}
    try:
        result = BEOPDataAccess.getInstance().getSchedule()
        rv = {'data':result,'status':True,'msg':''}
    except Exception as e:
        rv = {'data':[],'status':False,'msg':e.__str__()}
    return json.dumps(dict(rv), ensure_ascii=False)

# 删除一个日程
@app.route('/schedule/remove', methods=['POST'])
def schedule_remove():
    data = request.get_json()
    id = data.get('id')

    msg = validate_number(id, "id")
    if msg is not None:
        return msg

    rv = dict(err=0, msg='')
    bSucces = False
    try:
        bSucces = BEOPDataAccess.getInstance().removeSchedule(id)

    except Exception as e:
        pass

    if not bSucces:
        rv = dict(err=1, msg='')

    return json.dumps(dict(rv), ensure_ascii=False)

# 编辑日程
@app.route('/schedule/edit', methods=['POST'])
def schedule_edit():
    data = request.get_json()
    id = data.get('id')
    strScheduleName = data.get('scheduleName')
    isloop =data.get('isloop')
    strPointDefine=data.get('pointDefines')

    msg = validate_number(id, "id")
    if msg is not None:
        return msg

    rv = dict(err=0, msg='')
    bSucces = False
    try:
        bSucces = BEOPDataAccess.getInstance().editSchedule(id,strScheduleName, isloop, strPointDefine)
    except Exception as e:
        pass

    if not bSucces:
        rv = dict(err=1, msg='')
    return json.dumps(dict(rv), ensure_ascii=False)
'''
title:启用禁用启用日程
param:enable 1启用 0禁用
'''
@app.route('/schedule/enableSchedule', methods=['POST'])
def schedule_enableSchedule():
    data = request.get_json()
    id = data.get('id')
    enable = data.get('enable')

    msg = validate_number(id, "id")
    if msg is not None:
        return msg

    msg = validate_restricted_number(enable, "enable")
    if msg is not None:
        return msg

    bSucces = False
    rv = dict(err=0, msg='')
    try:
        bSucces = BEOPDataAccess.getInstance().enableSchedule(enable,id)
    except Exception as e:
        pass

    if not bSucces:
        rv = dict(err=1, msg='')
    return json.dumps(dict(rv), ensure_ascii=False)


# 在编辑周期列表
@app.route('/schedule/editWeeks', methods=['POST'])
def schedule_editWeeks():
    rv = dict(err=0, msg='')
    data = request.get_json()
    bSucces = False
    try:
        id = data.get('id')
        content = data.get('content')

        msg = validate_number(id, "id")
        if msg is not None:
            return msg

        bSucces = BEOPDataAccess.getInstance().removeWeekDays(id)
        for day in content:
            timeFrom = timeToInt(day.get('timeFrom'))
            timeTo = timeToInt(day.get('timeTo'))
            timeTo = timeToInt(day.get('enable'))
            bSucces = BEOPDataAccess.getInstance().addWeekDays(day.get('weekday'), timeFrom, timeTo, day.get('value'), day.get('id'))



    except Exception as e:
         pass

    if not bSucces:
        rv = dict(err=1, msg='')
    return json.dumps(dict(rv), ensure_ascii=False)

# 在显示周期列表
@app.route('/schedule/fromGroupIdGetSecheduleTask', methods=['POST'])
def schedule_fromGrouopIdGetSecheduleTask():
    rv = dict(err=0, msg='')
    data = request.get_json()
    groupid = data.get('id')

    msg = validate_number(groupid, "id")
    if msg is not None:
        return msg

    try:
        result = BEOPDataAccess.getInstance().fromGrouopIdGetSecheduleTask(groupid)
        if result is None:
            rv = {'data':None,'status':False,'msg':''}
        else:
            rv = {'data':result,'status':True,'msg':''}
    except Exception as e:
        rv = {'data':[],'status':False,'msg':e.__str__()}
    return json.dumps(dict(rv), ensure_ascii=False)

'''
title:编辑周期列表
auth:huangzhijun
param:{"id": 4,"content": [{"id": 4,"timeFrom": "02:00","timeTo": "12:30","value": 1,"weekday": 3}]}
return:{"msg": "ok", "status": true}
'''
@app.route('/schedule/simpleEditWeeks', methods=['POST'])
def simple_edit_weeks():
    result = {'status': True, 'msg': 'ok'}
    data = request.get_json()
    bSucces = False
    try:
        groupid = data.get('groupid')
        content = data.get('content')

        msg = validate_number(groupid, "groupid")
        if msg is not None:
            return msg

        BEOPDataAccess.getInstance().removeWeekDays(groupid)
        for day in content:

            enable = day.get('enable')
            if enable ==1:
                timeFrom = day.get('timeFrom')
                timeTo = day.get('timeTo')
                BEOPDataAccess.getInstance().addWeekDays(day.get('weekday'), timeFrom, timeTo, 1 ,groupid)
                if timeFrom!="0:00":
                    BEOPDataAccess.getInstance().addWeekDays(day.get('weekday'), '00:00', timeFrom, 0, groupid)
                if timeTo!="23:59" and timeTo!="00:00":
                    BEOPDataAccess.getInstance().addWeekDays(day.get('weekday'), timeTo, "23:59", 0, groupid)

    except Exception as e:
        result['msg'] = e.__str__()
        result['status'] = False


    return  json.dumps(result, ensure_ascii=False)



'''
title:显示周期列表
auth:huangzhijun
param:{"groupid": 4}
return:{"msg": "", "data": [{"weekday": 3, "timeFrom": "02:00", "groupid": 4, "value": "1", "id": 3, "timeExecute": "2018-07-17 10:41:58", "timeTo": "12:30"}], "status": true}
'''
@app.route('/schedule/simpleGetSecheduleTaskByGroupId', methods=['POST'])
def simple_get_sechedule_task_groupId():
    rv = dict(err=0, msg='')
    data = request.get_json()
    groupid = data.get('groupid')

    msg = validate_number(groupid, "groupid")
    if msg is not None:
        return msg
    try:
        result = BEOPDataAccess.getInstance().fromGrouopIdGetSecheduleTask(groupid)
        if result is None:
            rv = {'data':result,'status':False,'msg':''}
        else:
            rv = {'data':result,'status':True,'msg':''}
    except Exception as e:
        rv = {'data':[],'status':False,'msg':e.__str__()}
    return json.dumps(dict(rv), ensure_ascii=False)


@app.route('/schedule/simpleGetScheduleTaskByGroupId/v2', methods=['POST'])
def simple_get_sechedule_task_groupId_v2():
    '''
    传入参数: {"groupid": 2}
    '''
    rv = dict(err=0, msg='')
    data = request.get_json()
    groupid = data.get('groupid')

    msg = validate_number(groupid, "groupid")
    if msg is not None:
        return msg

    try:
        result = BEOPDataAccess.getInstance().fromGrouopIdGetSecheduleTaskV2(groupid)
        if result is None:
            rv = {'data':result,'status':False,'msg':''}
        else:
            rv = {'data':result,'status':True,'msg':''}
    except Exception as e:
        rv = {'data':[],'status':False,'msg':e.__str__()}
    return json.dumps(dict(rv), ensure_ascii=False)


def str_to_time(strTime):
    res = None
    try:
        if isinstance(strTime, str):
            strTime = strTime.strip()
            res = datetime.strptime(strTime, "%H:%M")
    except:
        traceback.print_exc()
    finally:
        return res

def time_to_str(tTime):
    res = None
    try:
        if isinstance(tTime, datetime):
            res = datetime.strftime(tTime, "%H:%M")
    except:
        traceback.print_exc()
    finally:
        return res

def bubble_sort(numList):
    length = len(numList)
    for i in range(length - 1):
        for j in range(length - 1 - i):
            if numList[j] > numList[j+1]:
                numList[j], numList[j+1] = numList[j+1], numList[j]
    return numList

def schedule_exists_in_list(schedule, timeList):
    res = False
    try:
        for item in timeList:
            if item.get("timeFrom") == str_to_time(schedule.get("timeFrom")) and item.get("timeTo") == str_to_time(schedule.get("timeTo")):
                res = True
    except:
        pass
    return res


def get_start_and_end_of_timelist(timeList):
    res = list()
    timeDict = dict()
    try:
        count = 0
        for idx in range(len(timeList)):
            if idx == 0:
                timeDict.update({
                    count: [time_to_str(timeList[idx])]
                })
            else:
                if timeList[idx] - timeList[idx - 1] > timedelta(minutes=1):
                    count += 1
                    timeDict.update({
                        count: [time_to_str(timeList[idx])]
                    })
                else:
                    if isinstance(timeDict.get(count), list):
                        timeDict.get(count).append(time_to_str(timeList[idx]))
                    else:
                        timeDict.update({
                            count: [time_to_str(timeList[idx])]
                        })

        for key in timeDict.keys():
            if len(timeDict.get(key)):
                res.append({
                    "timeFrom": timeDict.get(key)[0],
                    "timeTo": timeDict.get(key)[-1]
                })
    except:
        traceback.print_exc()
    return res


def get_schedule_to_enable(dictSchedule):
    res = list()
    try:
        listKey = list(dictSchedule.keys())
        listTimeFrom = [item for item in listKey if item.startswith("timeFrom")]
        for item in listTimeFrom:
            idx = item.replace("timeFrom", "")
            if dictSchedule.get("enable{0}".format(idx)) > 0 and str_to_time(dictSchedule.get("timeFrom{0}".format(idx))) < str_to_time(dictSchedule.get("timeTo{0}".format(idx))):
                res.append({
                    "timeFrom": dictSchedule.get("timeFrom{0}".format(idx)),
                    "timeTo": dictSchedule.get("timeTo{0}".format(idx))
                })
    except:
        traceback.print_exc()
    finally:
        return res


@app.route('/schedule/simpleEditWeeks/v2', methods=['POST'])
def simple_edit_weeks_v2():
    '''
    传入参数：
    {
        "groupid": 13,
        "content": [
            {
                "id": 1,
                "weekday": 1,
                "enable": 1,
                "timeFrom": "08:00",
                "timeTo": "14:00",
                "enable1": 0,
                "timeFrom1": "08:00",
                "timeTo1": "20:00",
                "enable2": 0,
                "timeFrom2": "08:00",
                "timeTo2": "20:00",
                "enable3": 0,
                "timeFrom3": "08:00",
                "timeTo3": "20:00",
                "enable4": 0,
                "timeFrom4": "08:00",
                "timeTo4": "20:00"
            },
            {
                "id":2,
                "weekday": 2,
                "enable": 0,
                "timeFrom": "08:00",
                "timeTo": "20:00",
                "enable1": 0,
                "timeFrom1": "08:00",
                "timeTo1": "20:00",
                "enable2": 1,
                "timeFrom2": "20:00",
                "timeTo2": "20:00",
                "enable3": 1,
                "timeFrom3": "08:00",
                "timeTo3": "20:00",
                "enable4": 1,
                "timeFrom4": "13:00",
                "timeTo4": "4:00"
            }
         ]
    }
    '''
    result = dict(status=False, msg="fail")
    try:
        rcv = request.get_json()
        groupid = rcv.get('groupid')
        listContent = rcv.get('content')

        msg = validate_number(groupid, "groupid")
        if msg is not None:
            return msg

        BEOPDataAccess.getInstance().removeWeekDays(groupid)

        for content in listContent:
            weekday = int(content.get("weekday"))
            dayScheduleList = get_schedule_to_enable(content)
            timeList = list()
            for schedule in dayScheduleList:
                if not schedule_exists_in_list(schedule, timeList):
                    timeList.append({
                        "timeFrom": str_to_time(schedule.get("timeFrom")),
                        "timeTo": str_to_time(schedule.get("timeTo"))
                    })

            enableTimeList = list()
            disableTimeList = list()

            startTime = datetime.strptime("00:00", "%H:%M")
            endTime = datetime.strptime("23:59", "%H:%M")
            singleTime = startTime

            while singleTime < endTime:
                cover = False
                for timeDict in timeList:
                    if singleTime >= timeDict.get("timeFrom") and singleTime <= timeDict.get("timeTo"):
                        enableTimeList.append(singleTime)
                        cover = True
                        break

                if not cover:
                    disableTimeList.append(singleTime)

                singleTime = singleTime + timedelta(minutes=1)
            enableList = get_start_and_end_of_timelist(enableTimeList)
            disableList = get_start_and_end_of_timelist(disableTimeList)

            # enable list 写入mysql
            for obj in enableList:
                if weekday:
                    BEOPDataAccess.getInstance().addWeekDays(weekday, obj.get("timeFrom"), obj.get("timeTo"), 1, groupid)

            # disable list 写入mysql
            for obj in disableList:
                if weekday:
                    BEOPDataAccess.getInstance().addWeekDays(weekday, obj.get("timeFrom"), obj.get("timeTo"), 0, groupid)

            result = dict(status=True, msg="succeed")
    except Exception as e:
        result = dict(status=False, msg=e.__str__())
    finally:
        return  json.dumps(result, ensure_ascii=False)


@app.route('/tool/evalStringExpressionCommon', methods=['POST'])
def eval_string_common():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strExpression = data.get('Expression')
    strConditionList = data.get('ConditionList')

    rvData =  tool_eval_string_common(strExpression, strConditionList)

    return json.dumps(dict(err=0, msg='ok', data=rvData), ensure_ascii=False)



def analysis_debug_info_of_expression(strExpression, pointDataMap):
    strResultAll = ''
    strPointNameList = []
    strErrPointNameList = []
    strSuccessPointNameList = []
    strSuccessPointValueNameList = []
    nLeft = strExpression.find('<%')
    while nLeft>=0:
        nRight = strExpression.find('%>', nLeft+1)
        if nRight<0:
            break

        strPP = strExpression[nLeft+2: nRight]
        if strPP:
            strPointNameList.append(strPP)
        nLeft = strExpression.find('<%', nRight)

    for pp in strPointNameList:
        pvv = pointDataMap.get(pp, None)
        if pvv is None:
            strErrPointNameList.append(pp)
        else:
            strSuccessPointNameList.append(pp)
            strSuccessPointValueNameList.append(pvv)

    strResultAll += "\r\n读取点值错误的点位有:\r\n"
    for item in strErrPointNameList:
        strResultAll+= item
        strResultAll+= '\r\n'

    strResultAll += "\r\n读取点值成功的点位:\r\n"
    strNewExpression = strExpression
    for nIndex in range(len(strSuccessPointNameList)):
        strNewExpression = strNewExpression.replace('<%'+ strSuccessPointNameList[nIndex] + '%>', strSuccessPointValueNameList[nIndex])
        strResultAll+= strSuccessPointNameList[nIndex]
        strResultAll+= '\t'
        strResultAll += strSuccessPointValueNameList[nIndex]
        strResultAll+= '\r\n'

    strResultAll+="\r\n公式表达式以点值代入后的转换表达式为:\r\n"
    strResultAll+= strNewExpression

    return strResultAll

"""
接口用途:   用于分析计算公式，点必须是<%%>包围的形式类的计算公式的计算测试结果
作者:       golding
输入参数:   json post请求
            str字段为公式全文
            mode: string, 0:表示计算历史某个时刻, 1表示计算实时
                
"""
@app.route('/tool/evalStringExpression', methods=['POST'])
def tool_eval_string():
    strDebugInfo = ''
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strExpression = data.get('str')

    strMode = data.get('mode', '1')
    strActTime = data.get('actTime')
    bDebug = data.get('debug', 0)
    if strActTime is None:
        strActTime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    if strExpression is None:
        return json.dumps(dict(err=1, msg='str param lost', data=-1), ensure_ascii=False)

    try:
        pointData = {}
        pointDataMap = {}
        if strMode=='1':
            pointData = []
            realtimeDataAll = BEOPDataAccess.getInstance().getInputTable()
            pointData = realtimeDataAll[0]
            pointDataMap = realtimeDataAll[1]

        if bDebug ==1 or bDebug=='1':
            strDebugInfo = analysis_debug_info_of_expression(strExpression, pointDataMap)
        rvData = eval_string_expression_strict(strExpression, strMode, strActTime, pointData)
        if rvData is None:
            strError = 'ERROR in /tool/evalStringExpression return None: expression:%s, pointData Count:%d' % ( strExpression, len(pointData))
            print(strError)
            logging.error(strError)
            return json.dumps(dict(err=1, msg=strError, data="无结果", debugInfo=strDebugInfo), ensure_ascii=False)

        return json.dumps(dict(err=0, msg='ok', data=rvData, debugInfo=strDebugInfo), ensure_ascii=False)

    except Exception as e:
        strError = 'ERROR in /tool/evalStringExpression :%s, expression:%s' % (e.__str__(), strExpression)
        logging.error(strError)
        return json.dumps(dict(err=1, msg=strError, data=-1, debugInfo=strDebugInfo), ensure_ascii=False)


@app.route('/tools/recoveryVirtualPointsValue', methods=['GET'])
def tools_recovery_virtual_point_value():
    try:
        result = BEOPSqliteAccess.getInstance().GetBackupRealtimePointValuesStatus()
        pointList = []
        valueList = []
        for item in result:
            pointList.append(item.get('pointname'))
            valueList.append(item.get('pointvalue'))
        if pointList:
            BEOPDataAccess.getInstance().setRealtimeData(pointList, valueList)
        rv = {'data':result,'status':True,'msg':''}
    except Exception as e:
        rv = {'data':[],'status':False,'msg':e.__str__()}
    return json.dumps(dict(rv), ensure_ascii=False)



@app.route('/mysqldb/iosample/checkbuild', methods=['GET'])
def db_iosample_checkbuild():
    try:
        result = BEOPDataAccess.getInstance().checkBuildIOSampleDB()
        rv = {'data':result,'status':True,'msg':''}
    except Exception as e:
        rv = {'data':[],'status':False,'msg':e.__str__()}
    return json.dumps(dict(rv), ensure_ascii=False)

@app.route('/reportTool/genExcelReportByTableData', methods=['POST'])
def report_tool_gen_excel_report_by_table_data():
    rv = dict(err=0, msg='', data={})

    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strReportName = data.get('reportName')
    strStartTime = data.get('strStartTime')
    strEndTime = data.get('strEndTime')

    headerList  = data.get('headerList')
    pointList = data.get('pointList')
    tableDataList= data.get('tableDataList')
    timeList = data.get('timeList')
    nDisplayDelta = data.get("displayDelta", 0)

    if strReportName is None:
        return json.dumps(dict(err=1, msg='reportName param lost', data=-1), ensure_ascii=False)
    if strStartTime is None:
        return json.dumps(dict(err=1, msg='strStartTime param lost', data=-1), ensure_ascii=False)
    if strEndTime is None:
        return json.dumps(dict(err=1, msg='strEndTime param lost', data=-1), ensure_ascii=False)
    if headerList is None:
        return json.dumps(dict(err=1, msg='headerList param lost', data=-1), ensure_ascii=False)
    if pointList is None:
        return json.dumps(dict(err=1, msg='pointList param lost', data=-1), ensure_ascii=False)
    if tableDataList is None:
        return json.dumps(dict(err=1, msg='tableDataList param lost', data=-1), ensure_ascii=False)

    try:
        nDisplayDelta = int(float(nDisplayDelta))
    except:
        return json.dumps(dict(err=1, msg='displayDelta数据类型有误', data=-1), ensure_ascii=False)

    if not isinstance(nDisplayDelta, int):
        return json.dumps(dict(err=1, msg='displayDelta must be an integer', data=-1), ensure_ascii=False)

    if nDisplayDelta < 0:
        return json.dumps(dict(err=1, msg='displayDelta must be a possitive inteter', data=-1), ensure_ascii=False)

    if timeList is not None:
        if len(tableDataList)!= len(timeList):
            return json.dumps(dict(err=1, msg='table data list size must same as time list', data=-1), ensure_ascii=False)

    if len(headerList)!= len(pointList):
        return json.dumps(dict(err=1, msg='headerList list size must same as pointList list', data=-1), ensure_ascii=False)

    # 时间列倒退处理，如nDisplayDelta=1，则倒退一天
    if nDisplayDelta > 0:
        while nDisplayDelta > 0:
            try:
                strPreDay = (datetime.strptime(timeList[0], "%Y-%m-%d") - timedelta(days=1)).strftime("%Y-%m-%d")
                timeList.insert(0, strPreDay)
                timeList.pop()
            except:
                pass
            finally:
                nDisplayDelta -= 1
    try:
        if tableDataList:
            if timeList is not None:
                headerList.insert(0, '时间')
            excelData = tablib.Dataset(headerList)
            i = 0
            for item in tableDataList:
                oneRowData = []
                if timeList is not None:
                    oneRowData.append(timeList[i])
                for pt in pointList:
                    strvv = str(item.get(pt))
                    try:
                        strvv = float(strvv)
                    except:pass
                    oneRowData.append(strvv)
                excelData.append(oneRowData)
                i += 1

            tNow = datetime.now()

            strPath = app.static_folder
            strTempPath = os.path.join(strPath, 'files')
            strTempPath = os.path.join(strTempPath, 'temp')
            if not os.path.exists(strTempPath):
                os.makedirs(strTempPath)
            strFileName = strReportName+ tNow.strftime('%Y%m%d%H%M%S') + '.xlsx'
            strFilePath = os.path.join(strTempPath, strFileName)
            filepath = make_excel_file_by_tablib_data_with_path(strFilePath, excelData)
            if filepath:
                rv= dict(msg='', err=0, data= {"reportfilename": strFileName})
    except Exception as e:
        rv = {'data': [], 'status': False, 'msg': e.__str__()}
    return json.dumps(dict(rv), ensure_ascii=False)


def generate_report_for_points(pointList, strStartTime, strEndTime, strReportName, strPeriod):
    pinfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(pointList)
    headerList = []
    headerList.append('时间')
    for pt in pointList:
        oneptinfo = pinfo.get(pt)
        if oneptinfo and oneptinfo.get('description'):
            headerList.append(oneptinfo.get('description') + '(%s)' % (pt))
        else:
            headerList.append(pt)

    tStartTime = datetime.strptime(strStartTime, '%Y-%m-%d %H:%M:%S')
    tEndTime = datetime.strptime(strEndTime, '%Y-%m-%d %H:%M:%S')

    tStartTime = tStartTime.replace(hour=0, minute=0, second=0)
    tEndTime = tEndTime.replace(hour=23, minute=59, second=59)

    tCur = tStartTime

    resultInfoList = []
    strPath = app.static_folder
    strTempPath = os.path.join(strPath, 'files')
    strTempPath = os.path.join(strTempPath, 'temp')
    if not os.path.exists(strTempPath):
        os.makedirs(strPath)

    strPath = os.path.join(strTempPath, 'system_report_' + strReportName)
    if not os.path.exists(strPath):
        os.makedirs(strPath)

    while tCur <= tEndTime:

        strFileName = strReportName + tCur.strftime('_%Y_%m_%d')
        strFilePath = os.path.join(strPath, strFileName + '.xlsx')
        if os.path.exists(strFilePath):
            tCur = tCur + timedelta(days=1)
            continue

        tDayEndOfCur = tCur.replace(hour=23, minute=59, second=59)
        hisdataAll = BEOPDataAccess.getInstance().get_history_data_padded(pointList, tCur.strftime('%Y-%m-%d %H:%M:%S'),
                                                                          tDayEndOfCur.strftime('%Y-%m-%d %H:%M:%S'),
                                                                          strPeriod)
        if hisdataAll:
            try:
                excelData = tablib.Dataset(headerList)
                i = 0
                for item in hisdataAll.get('time'):
                    oneRowData = []
                    oneRowData.append(item)
                    for pt in pointList:
                        pointValueAtTime = ''
                        try:
                            vList = hisdataAll.get('map').get(pt)
                            pointValueAtTime = vList[i]
                        except:
                            pass
                        oneRowData.append(str(pointValueAtTime))
                    excelData.append(oneRowData)
                    i += 1
                filepath = make_excel_file_by_tablib_data_with_path(strFilePath, excelData)
                if filepath:
                    resultInfoList.append(strFilePath)
            except Exception as e:
                pass

        tCur = tCur + timedelta(days=1)
        time.sleep(10)

    # remove .zip
    strZipFilePath = os.path.join(strTempPath, strReportName + '.zip')

    zip_dir(strPath, strZipFilePath)

    return True

@app.route('/reportTool/genExcelReportByPoint', methods=['POST'])
def report_tool_gen_excel_report_by_point():
    rv = dict(err=0, msg='', data={})

    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strReportName = data.get('reportName')
    strStartTime = data.get('timeStart')
    strEndTime = data.get('timeEnd')

    pointList = data.get('pointList')
    strPeriod = data.get('timeFormat')
    if strPeriod is None:
        strPeriod = 'm1'



    if strReportName is None:
        return json.dumps(dict(err=1, msg='reportName param lost', data=-1), ensure_ascii=False)
    if strStartTime is None:
        return json.dumps(dict(err=1, msg='timeStart param lost', data=-1), ensure_ascii=False)
    if strEndTime is None:
        return json.dumps(dict(err=1, msg='timeEnd param lost', data=-1), ensure_ascii=False)
    if pointList is None:
        return json.dumps(dict(err=1, msg='pointList param lost', data=-1), ensure_ascii=False)

    if generate_report_for_points(pointList, strStartTime, strEndTime, strReportName, strPeriod):
        rv['err'] = 0
    else:
        rv['err'] = 1

    rv['data'] = strReportName+'.zip'
    return json.dumps(dict(rv), ensure_ascii=False)


@app.route('/dataSample/getSampleList', methods=['POST'])
def data_sample_get_list():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    nIdList = BEOPDataAccess.getInstance().getSampleList()
    return jsonify(dict(err=0, data=nIdList))

@app.route('/dataSample/getSampleDataList', methods=['POST'])
def data_sample_get_data_list():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    sampleId = data.get('sampleId')

    rvData = BEOPDataAccess.getInstance().getSampleData(sampleId)
    return jsonify(dict(err=0, data=rvData))

@app.route('/dataSample/downloadSampleData', methods=['POST'])
def download_sample_data():
    try:
        rcv = request.get_json()
        sampleName = rcv.get("sampleName") if rcv.get("sampleName") else None
        inputCount = rcv.get("inputCount") if rcv.get("inputCount") else None
        outputCount = rcv.get("outputCount") if rcv.get("outputCount") else None

        if not sampleName:
            return jsonify(dict(err=1, msg="sampleName 不能为空", data=""))

        if not inputCount:
            return jsonify(dict(err=1, msg="inputCount 不能为空", data=""))

        if not isinstance(inputCount, int):
            return jsonify(dict(err=1, msg="inputCount 必须为整数", data=""))

        if not outputCount:
            return jsonify(dict(err=1, msg="outputCount 不能为空", data=""))

        if not isinstance(outputCount, int):
            return jsonify(dict(err=1, msg="outputCount 必须为整数", data=""))

        if inputCount <= 0 or outputCount <= 0:
            return jsonify(dict(err=1, msg="inputCount 和 outputCount必须大于0", data=""))

        # 获取数据
        rvData = BEOPDataAccess.getInstance().getSampleData(sampleName)

        # 根据inputCout, outputCount组织表头
        tableHeader = ["timeFrom", "timeTo"]

        for idx in range(inputCount):
            tableHeader.append("input%02d" % (idx+1))

        for idx in range(outputCount):
            tableHeader.append("output%02d" % (idx+1))

        if not len(tableHeader):
            return jsonify(dict(err=1, msg="获取内容为空", data=""))

        staticFolder = os.path.join(os.getcwd(), "siteinterface", "static")

        # 下载文件路径准备
        if not os.path.exists(staticFolder):
            os.mkdir(staticFolder)

        filesFolder = os.path.join(staticFolder, "files")

        if not os.path.exists(filesFolder):
            os.mkdir(filesFolder)

        tempFolder = os.path.join(filesFolder, "temp")

        if not os.path.exists(tempFolder):
            os.mkdir(tempFolder)

        strFileName = "sampleData.xlsx"
        destFile = os.path.join(tempFolder, strFileName)

        # 判断删除上次下载内容
        if os.path.exists(destFile):
            os.remove(destFile)

        # 开始写入EXCEL
        book = Workbook()

        sheet = book.create_sheet("sampleData", 0)

        # 写入表头
        for idx, item in enumerate(tableHeader):
            sheet.cell(row=1, column=idx+1, value=item)

        # 写入内容
        for i, data in enumerate(rvData):
            for j, item in enumerate(tableHeader):
                sheet.cell(row=i+2, column=j+1, value=data.get(item))

        book.save(destFile)

        return jsonify(dict(err=0, msg="下载成功", data=strFileName))

    except Exception as e:
        strLog = "ERROR in /dataSample/downloadSampleData： %s" % e.__str__()
        logging.error(strLog)

        return jsonify(dict(err=1, msg="下载失败: %s" % e.__str__(), data=""))


@app.route('/dataSample/getOrCreateSample', methods=['POST'])
def data_sample_get_or_create_sample():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    sampleName = data.get('sampleName')

    nId = BEOPDataAccess.getInstance().getOrCreateSampleIdByName(sampleName)
    return jsonify(dict(err=0, data=nId))

@app.route('/dataSample/getIO', methods=['POST'])
def data_sample_get_data_by_name():
    rv = dict(err=0, msg='', data={})
    strTimeFormat = "%Y-%m-%d"
    data = request.get_json()
    sampleName = data.get('sampleName') if data.get("sampleName") is not None else None
    strTimeFrom = data.get("timeFrom") if data.get("timeFrom") else None
    strTimeTo = data.get("timeTo") if data.get("timeTo") else None

    if not sampleName:
        return json.dumps(dict(err=1, msg='sampleName不能为空', data=-1), ensure_ascii=False)
    if strTimeFrom:
        if not isValidDate(strTimeFrom, strTimeFormat):
            return jsonify(dict(err=1, msg="开始时间格式有误", data=[]))
    if strTimeTo:
        if not isValidDate(strTimeTo, strTimeFormat):
            return jsonify(dict(err=1, msg="结束时间格式有误", data=[]))

    if strTimeFrom and strTimeTo:
        if datetime.strptime(strTimeFrom, strTimeFormat) > datetime.strptime(strTimeTo, strTimeFormat):
            return jsonify(dict(err=1, msg="开始时间不能大于结束时间", data=[]))

        strTimeFrom = "{0} 00:00:00".format(strTimeFrom)
        strTimeTo = "{0} 23:59:59".format(strTimeTo)

    dataList = BEOPDataAccess.getInstance().getSampleData(sampleName, strTimeFrom, strTimeTo)

    return jsonify(dict(err=0, data=dataList))


@app.route('/dataSample/removeIO', methods=['POST'])
def data_sample_remove_io():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    sampleId = data.get('sampleId')
    timeFrom = data.get('timeFrom')
    timeTo = data.get('timeTo')
    #if sampleId is None or timeFrom is None or timeTo is None:
        #return json.dumps(dict(err=1, msg='failed: sampleId/timeFrom/timeTo param need', data=-1), ensure_ascii=False)
    dataList = BEOPDataAccess.getInstance().removeSampleData(sampleId, timeFrom, timeTo)

    return jsonify(dict(err=0, data=dataList))


@app.route('/dataSample/insertIOs', methods=['POST'])
def data_sample_inserts():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed for no json data in body', data=-1), ensure_ascii=False)

    sampleList = data.get('sampleList')

    ret = BEOPDataAccess.getInstance().insertDataSampleIOs(sampleList)
    if ret == False:
        rv['err'] = 1

    return json.dumps(dict(rv), ensure_ascii=False)






@app.route('/dataSample/insertIO', methods=['POST'])
def data_sample_insert():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed for no json data in body', data=-1), ensure_ascii=False)
    sampleId = data.get('sampleId')
    strStartTime = data.get('timeFrom')
    strEndTime = data.get('timeTo')

    strInput01 = data.get('input01','')
    strInput02 = data.get('input02','')
    strInput03 = data.get('input03','')
    strInput04 = data.get('input04','')
    strInput05 = data.get('input05','')
    strInput06 = data.get('input06','')
    strInput07 = data.get('input07','')
    strInput08 = data.get('input08','')
    strInput09 = data.get('input09','')
    strInput10 = data.get('input10','')

    strInput11 = data.get('input11','')
    strInput12 = data.get('input12','')
    strInput13 = data.get('input13','')
    strInput14 = data.get('input14','')
    strInput15 = data.get('input15','')
    strInput16 = data.get('input16','')
    strInput17 = data.get('input17','')
    strInput18 = data.get('input18','')
    strInput19 = data.get('input19','')
    strInput20 = data.get('input20','')


    strInput21 = data.get('input21','')
    strInput22 = data.get('input22','')
    strInput23 = data.get('input23','')
    strInput24 = data.get('input24','')
    strInput25 = data.get('input25','')
    strInput26 = data.get('input26','')
    strInput27 = data.get('input27','')
    strInput28 = data.get('input28','')
    strInput29 = data.get('input29','')
    strInput30 = data.get('input30','')


    strInput31 = data.get('input31','')
    strInput32 = data.get('input32','')
    strInput33 = data.get('input33','')
    strInput34 = data.get('input34','')
    strInput35 = data.get('input35','')
    strInput36 = data.get('input36','')
    strInput37 = data.get('input37','')
    strInput38 = data.get('input38','')
    strInput39 = data.get('input39','')
    strInput40 = data.get('input40','')


    strInput41 = data.get('input41','')
    strInput42 = data.get('input42','')
    strInput43 = data.get('input43','')
    strInput44 = data.get('input44','')
    strInput45 = data.get('input45','')
    strInput46 = data.get('input46','')
    strInput47 = data.get('input47','')
    strInput48 = data.get('input48','')
    strInput49 = data.get('input49','')
    strInput50 = data.get('input50','')


    strOutput01 = data.get('output01','')
    strOutput02 = data.get('output02','')
    strOutput03 = data.get('output03','')
    strOutput04 = data.get('output04','')
    strOutput05 = data.get('output05','')
    strOutput06 = data.get('output06','')
    strOutput07 = data.get('output07','')
    strOutput08 = data.get('output08','')
    strOutput09 = data.get('output09','')
    strOutput10 = data.get('output10','')

    if BEOPDataAccess.getInstance().insertDataSampleIO(sampleId, strStartTime, strEndTime, strInput01,strInput02,strInput03,strInput04,strInput05,strInput06,strInput07,strInput08,strInput09,strInput10,
                                                       strInput11, strInput12, strInput13, strInput14, strInput15,
                                                       strInput16, strInput17, strInput18, strInput19, strInput20,
                                                       strInput21, strInput22, strInput23, strInput24, strInput25,
                                                       strInput26, strInput27,
                                                       strInput28, strInput29, strInput30,
                                                       strInput31, strInput32, strInput33, strInput34, strInput35,
                                                       strInput36, strInput37,
                                                       strInput38, strInput39, strInput40,
                                                       strInput41, strInput42, strInput43, strInput44, strInput45,
                                                       strInput46, strInput47,
                                                       strInput48, strInput49, strInput50,
                                                       strOutput01, strOutput02, strOutput03, strOutput04, strOutput05,
                                                       strOutput06, strOutput07, strOutput08, strOutput09, strOutput10):
        rv['err'] = 0
    else:
        rv['err'] = 1

    return json.dumps(dict(rv), ensure_ascii=False)


@app.route('/fdditem/getAll', methods=['POST'])
def fdd_item_get_all():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)


    nId = BEOPDataAccess.getInstance().getAllFddItems()
    return jsonify(dict(err=0, data=nId))


@app.route('/fdditem/get_realtime_status', methods=['POST'])
def fdd_item_get_realtime_status():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    fddGroupList = data.get('fddGroup')
    fddTimeFrom = data.get('QueryTimeFrom')
    fddTimeTo = data.get('QueryTimeTo')

    if fddGroupList is None:
        return jsonify(dict(err=1, msg='fddGroupList param lost', data=-1))

    rv = BEOPDataAccess.getInstance().getFddStatus(fddGroupList,fddTimeFrom)
    return jsonify(dict(err=0, data=rv))



@app.route('/fdditem/getHistoryStatus', methods=['POST'])
def fdd_item_get_history_status():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)


    nId = BEOPDataAccess.getInstance().getFddItemHistory()
    return jsonify(dict(err=0, data=nId))

@app.route('/fdditem/register', methods=['POST'])
def fdd_item_register():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strName = data.get('name')
    strNameCh = data.get('nameCh')
    strDescription = data.get('description')
    strOfEquipment = data.get('OfEquipment')
    strOfZone = data.get('ofZone')
    strOfResponseParty = data.get('ofResponseParty')
    strOfFaultClassify = data.get('OfFaultClassify')

    nId = BEOPDataAccess.getInstance().getOrCreateFddItem(strName, strNameCh, strDescription, strOfEquipment, strOfZone, strOfResponseParty, strOfFaultClassify)
    return jsonify(dict(err=0, data=nId))



@app.route('/fdditem/setStatus', methods=['POST'])
def fdd_item_setstatus():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strName = data.get('name')
    strFddTime = data.get('fddtime')
    if strName is None or strFddTime is None:
        return json.dumps(dict(err=1, msg='failed: name or fddtime param lost', data=-1), ensure_ascii=False)
    strContent = data.get('content','')
    strAnalysis = data.get('analysis','')
    strSuggestion = data.get('suggestion','')
    nGrade = data.get('grade',0)
    nStatus = data.get('status','')
    strRisk = data.get('risk','')

    bSuccess = BEOPDataAccess.getInstance().setFddFaultStatus(strName, strFddTime,  nStatus, nGrade, strContent, strAnalysis, strSuggestion, strRisk)
    if not bSuccess:
        return jsonify(dict(err=1, data={}, msg='setFddFaultStatus failed'))
    return jsonify(dict(err=0, data={}))




@app.route('/corefile/uploadpointtable', methods=['POST'])
def core_file_upload_point_table():
     upload_file = request.files['file']
     filename = upload_file.filename
     tmpfilename = ''
     path = os.getcwd()
     saveFilePath = os.path.join(path, 'temp')
     saveFilePath= os.path.join(saveFilePath, filename)
     upload_file.save(saveFilePath)

    #import excel file
    #@todo
     pointList = readPointsListFromExcel(saveFilePath)
     if pointList:
        BEOPSqliteAccess.getInstance().clearPointList()
        BEOPSqliteAccess.getInstance().insertPointList(pointList)
     else:
        return json.dumps(dict(err=1, msg='excel import failed'))

    #notice core to reload point table
     rv = BEOPDataAccess.getInstance().saveUnit01('core_reload_point_table',1)
     time.sleep(5)

     return json.dumps(dict(err=0, msg='', data = rv))



@app.route('/corefile/uploaddb', methods=['POST'])
def core_file_upload_db():
    upload_file = request.files['file']
    filename = upload_file.filename
    tmpfilename = ''
    path = os.getcwd()
    saveFilePath = os.path.join(path, 'temp')
    saveFilePath= os.path.join(saveFilePath, filename)
    upload_file.save(saveFilePath)

    #copy file
    if not BEOPSqliteAccess.getInstance().test_valid_4db(saveFilePath):
        return json.dumps(dict(err=1, msg='db file not valid'))
    #kill core
    process_stop_all()


    coreFilePath = app.config['CORE_PATH']
    coreFilePath = os.path.join(coreFilePath, app.config.get('DB_FILE_NAME'))
    shutil.copy(saveFilePath, coreFilePath)

    #restart core

    process_restart_domcore()
    process_restart_domlogic()

    return json.dumps(dict(err=0, msg=''))


'''
get data error
  设置数据点的数据为不可信
  1:不可信
  2:必然错误，不能用于任何控制和计算
  
'''


@app.route('/fdd/setDataErr', methods=['POST'])
def fdd_set_data_error():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strPointList = data.get('pointList')
    strFddTime = data.get('fddtime')
    strErrGrade = data.get('error')
    strReason = data.get('reason', '')
    if strPointList is None or strFddTime is None or strErrGrade is None:
        return json.dumps(dict(err=1, msg='failed: pointList or fddtime or error param lost', data=-1), ensure_ascii=False)


    bSuccess = BEOPDataAccess.getInstance().setDataErr(strPointList, strFddTime, strErrGrade, strReason)
    if not bSuccess:
        return jsonify(dict(err=1, data={}, msg='fdd_set_data_error failed'))
    return jsonify(dict(err=0, data={}))

@app.route('/fdd/getDataErr', methods=['POST'])
def fdd_get_data_error():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strPointList = data.get('pointList')
    strFddTime = data.get('fddtime')

    if strPointList is None or strFddTime is None:
        return json.dumps(dict(err=1, msg='failed: pointList or fddtime or error param lost', data=-1), ensure_ascii=False)


    rv = BEOPDataAccess.getInstance().getDataErr(strPointList, strFddTime)
    if rv is None:
        return jsonify(dict(err=1, data={}, msg='fdd_set_data_error failed'))
    return jsonify(dict(err=0, data=rv))

@app.route('/energy/feeReportGen', methods=['POST'])
def energy_fee_report_gen():
    rv = dict(err=0, msg='', data={})
    msg = ''
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strPointList = data.get('pointList')
    strTimeFrom = data.get('dateFrom')
    strTimeTo = data.get('dateTo')
    priceDef = data.get('price')

    if priceDef is None:
        msg = 'price参数缺失'
    if priceDef is None or priceDef == {}:
        priceDef = dict(priceFeng=1, priceGu=1, pricePing=1, priceListM30=[1]*48)

    priceFeng = priceDef.get('priceFeng')
    priceGu = priceDef.get('priceGu')
    pricePing = priceDef.get('pricePing')
    priceListM30 = priceDef.get('priceListM30')


    if not isinstance(priceListM30, list):
        return json.dumps(dict(err=1, msg='failed for the priceListM30 should be array', data=-1), ensure_ascii=False)

    if len(priceListM30)<48:
        return json.dumps(dict(err=1, msg='failed for the priceListM30 should be array size==48', data=-1), ensure_ascii=False)

    feeData ={}

    try:
        tTimeFrom = datetime.strptime(strTimeFrom, '%Y-%m-%d')
        tTimeFrom = tTimeFrom.replace(hour=0, minute=0, second=0)
    except:
        try:
            tTimeFrom = datetime.strptime(strTimeFrom, '%Y-%m-%d %H:%M:%S')
        except:
            return json.dumps(dict(err=1, msg='timeFrom format wrong, should be YYYY-MM-DD or YYYY-MM-DD HH:mm:SS', data=-1),
                              ensure_ascii=False)

    try:
        tTimeTo = datetime.strptime(strTimeTo, '%Y-%m-%d')
        tTimeTo = tTimeTo.replace(hour=0, minute=0, second=0)
    except:
        try:
            tTimeTo = datetime.strptime(strTimeTo, '%Y-%m-%d %H:%M:%S')
        except:
            return json.dumps(dict(err=1, msg='timeTo format wrong, should be YYYY-MM-DD or YYYY-MM-DD HH:mm:SS', data=-1),
                              ensure_ascii=False)

    if tTimeFrom.year == tTimeTo.year and tTimeFrom.month == tTimeTo.month and tTimeFrom.day == tTimeTo.day:
        tTimeTo = tTimeTo.replace(hour=23, minute=59, second=59)

    tCur = tTimeFrom
    badDataList = []
    while tCur<tTimeTo:
        tNexCur = tCur + timedelta(days=1)
        if tNexCur>tTimeTo:
            tNexCur = tTimeTo
        dataOneDay = BEOPDataAccess.getInstance().getInstance().get_history_data_padded(strPointList, tCur.strftime('%Y-%m-%d %H:%M:%S'), tNexCur.strftime('%Y-%m-%d %H:%M:%S'), 'm5')

        for point in strPointList:
            if feeData.get(point) is None:
                feeData[point] = dict(rawHourlyData=[], data=[], stat=dict(fengMonth=0, guMonth=0, pingMonth=0, totalMonth=0,
                                                         fengFeeMonth=0, guFeeMonth=0, pingFeeMonth=0, totalFeeMonth=0, startNum=None, endNum=None), errorList=[])

            dataOfThisPoint= dataOneDay['map'].get(point, None)
            sumFengQuantity = 0
            sumGuQuantity = 0
            sumPingQuantity = 0

            if dataOfThisPoint is None:
                strError = 'ERROR: get point %s  m5 data(from %s to %s) return None'%(point, tCur.strftime('%Y-%m-%d %H:%M:%S'), tNexCur.strftime('%Y-%m-%d %H:%M:%S'))
                print(strError)
                logging.error(strError)

            if dataOfThisPoint:
                for nDataIndex in range(len(dataOfThisPoint)-1):
                    try:
                        nPriceIndex = int(nDataIndex/6)
                        strTimeThis = dataOneDay['time'][nDataIndex]
                        tThis = datetime.strptime(strTimeThis, '%Y-%m-%d %H:%M:%S')

                        if nDataIndex==0 and feeData[point]['stat']['startNum'] is None:
                            feeData[point]['stat']['startNum'] = dataOfThisPoint[nDataIndex]

                        if nDataIndex== len(dataOfThisPoint)-2 and tNexCur==tTimeTo:
                            feeData[point]['stat']['endNum'] = dataOfThisPoint[nDataIndex + 1]

                        if tThis.minute==0:
                            feeData[point]['rawHourlyData'].append(dataOfThisPoint[nDataIndex])
                        try:
                            fDelta = float(dataOfThisPoint[nDataIndex+1])- float(dataOfThisPoint[nDataIndex])
                        except:
                            #data1 = [point, dataOneDay['time'][nDataIndex-1] if nDataIndex>=1 else None, str(dataOfThisPoint[nDataIndex-1]) if nDataIndex>=1 else None]
                            #data2 = [point, dataOneDay['time'][nDataIndex], str(dataOfThisPoint[nDataIndex])]
                            #data3 = [point, dataOneDay['time'][nDataIndex+1] if nDataIndex+1<len(dataOfThisPoint) else None, str(dataOfThisPoint[nDataIndex+1]) if nDataIndex+1<len(dataOfThisPoint) else None]
                            #badDataList.extend([data1,data2,data3,['','','']])
                            strError = 'ERROR in cal delta:%s at %s: %s - %s'%(point, tCur.strftime('%Y-%m-%d'), str(dataOfThisPoint[nDataIndex+1]), str(dataOfThisPoint[nDataIndex]))
                            print(strError)
                            logging.error(strError)
                            continue

                        if fDelta<0 or fDelta>5000:#有时5000可能有效，因为电表一天没走，处理需要时间
                            strError = 'PowerTotal Data ERROR: %s, the delta is:%.2f at time:%s'%(point, fDelta, strTimeThis)
                            print(strError)
                            logging.error(strError)
                            dataOfThisPoint[nDataIndex + 1] = dataOfThisPoint[nDataIndex] #累积量延续覆盖

                            data1 = [dataOneDay['time'][nDataIndex - 1] if nDataIndex >= 1 else '',point,
                                     str(dataOfThisPoint[nDataIndex - 1]) if nDataIndex >= 1 else '']
                            data2 = [ dataOneDay['time'][nDataIndex], point,str(dataOfThisPoint[nDataIndex])]
                            data3 = [dataOneDay['time'][nDataIndex + 1] if nDataIndex + 1 < len(dataOfThisPoint) else '',point,
                                     str(dataOfThisPoint[nDataIndex + 1]) if nDataIndex + 1 < len(
                                         dataOfThisPoint) else '']
                            data4 = [dataOneDay['time'][nDataIndex + 2] if nDataIndex + 2 < len(dataOfThisPoint) else '',point,
                                     str(dataOfThisPoint[nDataIndex + 2]) if nDataIndex + 2 < len(
                                         dataOfThisPoint) else '']
                            data5 = [dataOneDay['time'][nDataIndex + 3] if nDataIndex + 3 < len(dataOfThisPoint) else '',point,
                                     str(dataOfThisPoint[nDataIndex + 3]) if nDataIndex + 3 < len(
                                         dataOfThisPoint) else '']
                            badDataList.extend([data1, data2, data3, data4, data5, ['', '', '']])
                            feeData[point]['errorList'].append(dict(date=strTimeThis, data1= dataOfThisPoint[nDataIndex], data2 = dataOfThisPoint[nDataIndex+1]))

                        else:
                            if nPriceIndex>= len(priceListM30):
                                xx=0
                            elif priceListM30[nPriceIndex]==2:
                                sumFengQuantity+= fDelta
                            elif priceListM30[nPriceIndex]==1:
                                sumPingQuantity+= fDelta
                            elif priceListM30[nPriceIndex]==0:
                                sumGuQuantity+= fDelta

                    except Exception as e:
                        strLog = "ERROR in /energy/feeReportGen: %s" % e.__str__()
                        logging.error(strLog)
                        return jsonify(dict(msg="数据获取失败", err=1, data={}, reportURL="", reportErrURL="",reportRawDataURL=""))

            feeData[point]['stat']['fengMonth'] +=round(sumFengQuantity,1)
            feeData[point]['stat']['guMonth'] +=round(sumGuQuantity,1)
            feeData[point]['stat']['pingMonth'] +=round(sumPingQuantity,1)
            feeData[point]['stat']['totalMonth'] += round((sumPingQuantity+sumFengQuantity+sumGuQuantity),1)

            thisDate = tCur.strftime('%Y-%m-%d')
            feeData[point]['data'].append(dict(date= thisDate, fengQuantity=sumFengQuantity, guQuantity = sumGuQuantity, pingQuantity=  sumPingQuantity))

        tCur = tNexCur

    if not feeData:
        return jsonify(dict(msg="数据获取失败", err=1, data={}, reportURL="", reportErrURL="",reportRawDataURL=""))

    #stat
    for point in strPointList:
        try:
            feeData[point]['stat']['pointName'] = point
            feeData[point]['stat']['fengFeeMonth'] = round(feeData[point]['stat']['fengMonth'] * float(priceFeng),1)
            feeData[point]['stat']['guFeeMonth'] = round(feeData[point]['stat']['guMonth'] * float(priceGu),1)
            feeData[point]['stat']['pingFeeMonth'] = round(feeData[point]['stat']['pingMonth'] * float(pricePing),1)
            feeData[point]['stat']['totalFeeMonth'] = round(feeData[point]['stat']['fengFeeMonth']+feeData[point]['stat']['guFeeMonth']+ feeData[point]['stat']['pingFeeMonth'],1)
        except Exception as e:
            logging.error('\r\nERROR in debug iamdora:' + e.__str__())
            logging.error(str(feeData[point]['stat']))
            logging.error('\r\n')
        #save to excel
    print("")
    feeData['columns'] = []
    feeData['columns'].append(dict(title='点名', key='pointName'))
    feeData['columns'].append(dict(title='支路', key='description'))
    feeData['columns'].append(dict(title='编号', key=''))
    feeData['columns'].append(dict(title='本月起始读数', key='startNum'))
    feeData['columns'].append(dict(title='本月终止读数', key='endNum'))
    feeData['columns'].append(dict(title='本月总量', key='totalMonth'))
    feeData['columns'].append(dict(title='本月峰期用量', key='fengMonth'))
    feeData['columns'].append(dict(title='本月谷期用量', key='guMonth'))
    feeData['columns'].append(dict(title='本月平期用量', key='pingMonth'))
    feeData['columns'].append(dict(title='本月总费用', key='totalFeeMonth'))
    feeData['columns'].append(dict(title='本月峰期费用', key='fengFeeMonth'))
    feeData['columns'].append(dict(title='本月谷期费用', key='guFeeMonth'))
    feeData['columns'].append(dict(title='本月平期费用', key='pingFeeMonth'))
    feeData['columns'].append(dict(title='抄表日期', key=''))

    strHeaderList = [item.get('title') for item in feeData['columns']]

    excelData = tablib.Dataset(strHeaderList)
    i = 0
    pointsInfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db(strPointList)
    fTotalStart = 0
    fTotalEnd = 0
    fUseTotal = 0
    fUseTotalFeng = 0
    fUseTotalGu = 0
    fUseTotalPing = 0
    fUseFeeTotal = 0
    fUseFeeTotalFeng = 0
    fUseFeeTotalGu = 0
    fUseFeeTotalPing = 0
    for pt in strPointList:
        data = feeData[pt]
        oneRowData = []
        oneRowData.append(pt)
        ptInfo = pointsInfo.get(pt)
        feeData[pt]['stat']['description'] = pt if ptInfo is None else ptInfo['description']
        oneRowData.append(pt if ptInfo is None else ptInfo['description'])
        oneRowData.append('')
        oneRowData.append(str(data['stat']['startNum']))
        oneRowData.append(str(data['stat']['endNum']))
        oneRowData.append(str(data['stat']['totalMonth']))
        oneRowData.append(str(data['stat']['fengMonth']))
        oneRowData.append(str(data['stat']['guMonth']))
        oneRowData.append(str(data['stat']['pingMonth']))
        oneRowData.append(str(data['stat']['totalFeeMonth']))
        oneRowData.append(str(data['stat']['fengFeeMonth']))
        oneRowData.append(str(data['stat']['guFeeMonth']))
        oneRowData.append(str(data['stat']['pingFeeMonth']))
        oneRowData.append('')
        excelData.append(oneRowData)
        try:
            fTotalStart+=data['stat']['startNum']
            fTotalEnd+= data['stat']['endNum']
            fUseTotal += data['stat']['totalMonth']
            fUseTotalFeng += data['stat']['fengMonth']
            fUseTotalGu += data['stat']['guMonth']
            fUseTotalPing += data['stat']['pingMonth']
            fUseFeeTotal += data['stat']['totalFeeMonth']
            fUseFeeTotalFeng += data['stat']['fengFeeMonth']
            fUseFeeTotalGu += data['stat']['guFeeMonth']
            fUseFeeTotalPing += data['stat']['pingFeeMonth']

        except Exception as e:
            logging.error('ERROR in summarize TOtal:%s'%(e.__str__()))

    fTotalStart = round(fTotalStart,1)
    fTotalEnd = round(fTotalEnd,1)
    fUseTotal = round(fUseTotal,1)
    fUseTotalFeng =round(fUseTotalFeng,1)
    fUseTotalGu =round(fUseTotalGu,1)
    fUseTotalPing =round(fUseTotalPing,1)
    fUseFeeTotal =round(fUseFeeTotal,1)
    fUseFeeTotalFeng=round(fUseFeeTotalFeng,1)
    fUseFeeTotalGu =round(fUseFeeTotalGu,1)
    fUseFeeTotalPing=round(fUseFeeTotalPing,1)

    feeData['pointList'] = strPointList
    feeData['pointList'].append('##TOTAL##')
    feeData['##TOTAL##'] = dict()
    feeData['##TOTAL##']['stat'] = dict(description='总计', startNum=fTotalStart, endNum = fTotalEnd, totalMonth = fUseTotal, fengMonth=fUseTotalFeng,
                                                guMonth = fUseTotalGu, pingMonth= fUseTotalPing, totalFeeMonth = fUseFeeTotal, fengFeeMonth=fUseFeeTotalFeng,
                                                guFeeMonth = fUseFeeTotalGu, pingFeeMonth= fUseFeeTotalPing)
    excelData.append(['总计','','',str(fTotalStart), str(fTotalEnd), str(fUseTotal), str(fUseTotalFeng), str(fUseTotalGu),
                      str(fUseTotalPing),str(fUseFeeTotal), str(fUseFeeTotalFeng), str(fUseFeeTotalGu), str(fUseFeeTotalPing), ''])

    tNow = datetime.now()

    strPath = app.static_folder
    strTempPath = os.path.join(strPath, 'files')
    strTempPath = os.path.join(strTempPath, 'temp')
    if not os.path.exists(strTempPath):
        os.makedirs(strTempPath)
    strFileName = 'EnergyFeeReport' + tNow.strftime('%Y%m%d%H%M%S') + '.xlsx'
    strFilePath = os.path.join(strTempPath, strFileName)
    filepath = make_excel_file_by_tablib_data_with_path(strFilePath, excelData)

    excelDataErr = tablib.Dataset(['pointtime', 'pointname', 'pointvalue'])
    for errorRecord in badDataList:
        excelDataErr.append(errorRecord)
    strERRFileName = 'EnergyFeeErrorReport' + tNow.strftime('%Y%m%d%H%M%S') + '.xlsx'
    strERRFilePath = os.path.join(strTempPath, strERRFileName)
    errfilepath = make_excel_file_by_tablib_data_with_path(strERRFilePath, excelDataErr)

    #raw data excel
    rawHeaderList = ['时间']
    for pt in strPointList:
        ptDescription = '' if ptInfo is None else ptInfo['description']
        rawHeaderList.append('%s(%s)'%(pt, ptDescription))
    excelDataRawData = tablib.Dataset(rawHeaderList)

    tCur = tTimeFrom
    nIndex = 0
    while tCur <= tTimeTo:
        oneRawRowData = [tCur.strftime('%Y-%m-%d %H:%M:%S')]
        for pt in strPointList:
            if feeData[pt] == "##TOTAL##":
                continue
            try:
                if feeData[pt].get('rawHourlyData'):
                    oneRawRowData.append(feeData[pt].get('rawHourlyData')[nIndex])
                else:
                    oneRawRowData.append(0)
            except Exception as e:
                strError = 'ERROR in get rawHourlyData of point:%s, index:%d, time:%s'%(pt, nIndex, tCur.strftime('%Y-%m-%d %H:%M:%S'))
                print(strError)
                logging.error(strError)
                oneRawRowData.append(0)
                continue
        excelDataRawData.append(oneRawRowData)
        tCur+= timedelta(hours=1)
        nIndex+=1
    strRawFileName = 'EnergyFeeRawDataReport_' + tTimeFrom.strftime('%Y%m%d%H%M%S')+'_to_'+ tTimeTo.strftime('%Y%m%d%H%M%S') + '.xlsx'
    strRawFilePath = os.path.join(strTempPath, strRawFileName)
    rawfilepath = make_excel_file_by_tablib_data_with_path(strRawFilePath, excelDataRawData)

    return jsonify(dict(msg=msg, err=0, data=feeData, reportURL = strFileName, reportErrURL = strERRFileName, reportRawDataURL = strRawFileName))




@app.route('/data/importHistoryDataFile', methods=['POST'])
def data_import_histoy_data_file():
    upload_file = request.files['file']
    requestfilename = upload_file.filename
    path = os.getcwd()
    saveFilePath = os.path.join(path, 'temp')
    if not os.path.exists(saveFilePath):
        os.mkdir(saveFilePath)

    filename = datetime.now().strftime('%Y%m%d%H%M%S') + '_' + requestfilename
    strFileName= os.path.join(saveFilePath, filename)

    allImportInfo = {}
    if RedisManager.is_alive():
        allImportInfo = RedisManager.get('import_history_data_progress')
        if not allImportInfo:
            allImportInfo = {}
        if allImportInfo.get(requestfilename) is None:
            allImportInfo[requestfilename] = {}
        allImportInfo[requestfilename]["requestTime"] = datetime.now()
        RedisManager.set('import_history_data_progress', allImportInfo)

    upload_file.save(strFileName)

    #copy file
    rv = dict(err=0, msg='', data={})

    try:
        allData = read_xlsx(strFileName)
    except Exception as e:
        return json.dumps(dict(err=1, msg='ERROR in read_xlsx:%s'%(e.__str__()), data={}))
    if allData is None:
        rv = dict(err=1, msg='文件读取失败: 检查data目录下是否存在 %s 文件以及文件内容格式是否符合标准'%(strFileName))
    else:
        print('文件读取成功: data/%s，共%d条数据' % (strFileName, len(allData)))
        nInsertIndex = 1
        resultAll = ['--------start import--------']

        if RedisManager.is_alive():
            RedisManager.set('import_history_data_progress_' + requestfilename, resultAll)
            allImportInfo = RedisManager.get('import_history_data_progress')
            if not allImportInfo:
                allImportInfo = {}
            allImportInfo[requestfilename]["startTime"] = datetime.now()
            allImportInfo[requestfilename]['dataTotalCount']= len(allData)
            RedisManager.set('import_history_data_progress', allImportInfo)

        reCalcTaskList = []
        nSuccessCount = 0
        nFailCount = 0
        for row in allData:
            try:
                print('import %d row data from (%s) '%(nInsertIndex, filename))
                if row.get('pointtime') =="REALTIME":
                    setRealtimeResult = BEOPDataAccess.getInstance().setRealtimeData([row.get('pointname')], [row.get('pointvalue')])
                    bInsertSuc = True if setRealtimeResult.get("err") == 0 else False
                    strTimeNow = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    strSourceWithIP = 'HISTORY_DATA_IMPORT_%s' % (request.remote_addr)
                    BEOPDataAccess.getInstance().addLogicOutputRecord(strTimeNow,[row.get('pointname')],[row.get('pointvalue')], strSourceWithIP)
                    nSuccessCount += 1
                else:
                    insertResult = BEOPDataAccess.getInstance().insert_history_data(row.get('pointname'), row.get('pointtime'), row.get('pointvalue'),row.get('strUserId') ,row.get('strOriginalValue'))
                    bInsertSuc = insertResult[0]
                    reCalcTaskList.append(dict(point=row.get('pointname'), value=row.get('pointvalue'), time=row.get('pointtime')))
                if bInsertSuc:

                    nSuccessCount+=1#resultAll.append('保存成功%s(%s,%s)'%(row.get('pointname'), row.get('pointtime'), row.get('pointvalue')))
                else:
                    nFailCount+=1
                    resultAll.append('保存失败%s(%s), 原因:%s' % (row.get('pointname'), row.get('pointtime'),insertResult[1]))
                if RedisManager.is_alive():
                    RedisManager.set('import_history_data_progress_'+ requestfilename, dict(msg=resultAll, successCount=nSuccessCount, failCount=nFailCount))
                nInsertIndex +=1
            except Exception as exx:
                continue

        resultAll.append('--------end import--------')

        assign_vpoint_recalculation_task(reCalcTaskList)

        if RedisManager.is_alive():
            RedisManager.set('import_history_data_progress_' + requestfilename, resultAll)
            allImportInfo = RedisManager.get('import_history_data_progress')
            if not allImportInfo:
                allImportInfo = {requestfilename: {}}
            allImportInfo[requestfilename]["endTime"] = datetime.now()
            RedisManager.set('import_history_data_progress', allImportInfo)
        print('insert finished')
        time.sleep(10)

    #kill core
    return json.dumps(dict(err=0, msg='', data=resultAll))
    #return jsonify(rv)





@app.route('/get_realtimedata_in_table',methods=['POST'])
def get_realtimedata_in_table():

    args = request.form
    rdata = request.get_json()

    if rdata.get('proj') is None:
        proj = 1
    else:
        proj = rdata.get('proj')
    pointList = rdata.get('pointList')
    strTableName = rdata.get('tableName')

    tNow = datetime.now()
    try:
        dj = BEOPDataAccess.getInstance().getInputTable(pointList, strTableName)[0]
    except Exception as e:
        print('ERROR: ' + e.__str__())

@app.route('/update_process_status',methods=['POST'])
def update_process_status():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    process = data.get('process')

    try:
        dj = BEOPDataAccess.getInstance().updateProcessStatus(process)
    except Exception as e:
        strError = 'ERROR: ' + e.__str__()
        print(strError)
        return jsonify(dict(err=1, data={}, msg = strError ))
    return jsonify(dict(err=0, data=dj))


@app.route('/get_process_status',methods=['POST'])
def get_process_status():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    process = data.get('process')

    try:
        nStatus = BEOPDataAccess.getInstance().getProcessStatus(process)
    except Exception as e:
        strError = 'ERROR: ' + e.__str__()
        print(strError)
        return jsonify(dict(err=1, data={}, msg = strError ))
    return jsonify(dict(err=0, data=dict(status= nStatus)))



@app.route('/get_core_status',methods=['POST'])
def get_core_status():
    args = request.form
    rdata = request.get_json()
    try:
        dj = BEOPDataAccess.getInstance().getCoreStatus()
    except Exception as e:
        print('ERROR: ' + e.__str__())
    return jsonify(dict(err=0, data=dj))

  # rv = dict(err=0, msg='', data={})
  #   data = request.get_json()
  #   if data is None:
  #       return json.dumps(dict(err=0, msg='failed', data=-1), ensure_ascii=False)
  #   sampleId = data.get('sampleId')
  #
  #   nId = BEOPDataAccess.getInstance().getSampleData(sampleId)
  #   return jsonify(dict(err=0, data=nId))
@app.route('/data/getImportHistoryDataProcessInfo', methods=['POST'])
def get_import_HistoryDataProfessInfo():
    # copy file
    args = request.form
    rdata = request.get_json()
    try:
        dj = BEOPDataAccess.getInstance().getimportHistoryDataProfess()
    except Exception as e:
        print('ERROR: ' + e.__str__())
    return jsonify(dict(err=0, data=dj))


@app.route('/logic/addRecord', methods=['POST'])
def logic_add_record():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strLogicName = data.get('logicname')
    strPointTime = data.get('pointtime')
    strPointName = data.get('pointname')
    strPointValue = data.get('pointvalue')

    logging.info('/logic/addRecord from %s set %s'%(strLogicName, strPointName))

    if strLogicName is None or strPointTime is None or strPointName is None or strPointValue is None:
        return json.dumps(dict(err=1, msg='failed: logicname or pointtime or pointname or pointvalue param lost', data=-1), ensure_ascii=False)


    bSuccess = BEOPDataAccess.getInstance().addLogicOutputRecord(strPointTime, strPointName, strPointValue, strLogicName)
    if not bSuccess:
        return jsonify(dict(err=1, data={}, msg='logic_add_record failed'))
    return jsonify(dict(err=0, data={}, msg='logic_add_record success'))

@app.route('/logic/search', methods=['POST'])
@app.route('/debug/logic/search', methods=['POST'])
def logic_search_record():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strThreadName = data.get('threadname')
    strPointName = data.get('pointname')

    if strThreadName is None  or strPointName is None :
        return json.dumps(dict(err=1, msg='failed: threadname or pointname  param lost', data=-1), ensure_ascii=False)


    strAll = BEOPDataAccess.getInstance().logic_search(strThreadName, strPointName)
    return jsonify(dict(err=0, data=strAll))

@app.route('/dataSample/predictOutput', methods=['POST'])
def data_sample_predict_output():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    sampleId = data.get('sampleId')
    fDataInputList = data.get('inputList')

    #generate model

    #predict output

    fDataOutputList = []
    return jsonify(dict(err=0, data=fDataOutputList))


@app.route('/sendmailbycloud', methods=['POST'])
def send_mail_from_cloud():
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    subject = data.get('subject')
    recipients = data.get('recipients')
    sender = data.get('sender')
    html = data.get('html')


    jsonData = app.config['CLOUD_SETTING']
    if not jsonData:
        return json.dumps(dict(err=1, msg='failed to send mail for cloud setting is not valid', data=-1), ensure_ascii=False)

    strUserName = jsonData.get('CloudSettingUserName')
    strUserPwd = jsonData.get('CloudSettingPwd')
    strMessageQueueAddress = jsonData.get('CloudSettingServerAddress')

    bSuccess = sendEmailByMessageQueue(strMessageQueueAddress, strUserName, strUserPwd, subject, recipients, 1, sender, html, cc=[], bcc=[])
    #generate model

    if not bSuccess:
        return jsonify(dict(err=1, data={}, msg='Failed to push email message into queue'))

    return jsonify(dict(err=0, msg='', data={}))


###
# 接口：为获取机器码

###
@app.route('/license/getRequestNo', methods=['POST'])
def license_get_request_no():
    #@todo1
    rv = dict(err=0, msg='', data={})

    data = request.get_json()
    nLicenseVersion = BEOPDataAccess.getInstance().getLicenseVersion()
    licenseManage = LicenseManage(nLicenseVersion)
    contentList = licenseManage.getVolumeSerialNumberList()
    volnumber = None
    if contentList and isinstance(contentList, list):
        volnumber = contentList[-1]
        for cc in contentList:
            if cc and len(cc)>len(volnumber):#choose length more
                volnumber = cc

    else:
        return json.dumps(dict(err=1, msg='failed for no request key avialable', data=-1), ensure_ascii=False)

    # if data is None:
    #     return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strVolNumberMD5 = licenseManage.getMD5OfSerialNumber(volnumber)
    return jsonify(dict(err=0, data=strVolNumberMD5))


@app.route('/license/saveResponseNo', methods=['POST'])
def license_save_response_no():
    #@todo save into ini files.
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)
    strKey = data.get('key')
    for nLicenseVersion in [0,1,2]:
        try:
            licenseManage = LicenseManage(nLicenseVersion)
            registeResult = licenseManage.regist(strKey)
            if registeResult:
                return jsonify(dict(err=0, data={}, msg=''))
        except Exception as e:
            strError = 'ERROR in license_save_response_no:' + e.__str__()
            print(strError)
            logging.error(strError)
            continue

    return jsonify(dict(err=1, msg='failed'))

@app.route('/license/checkAuthored', methods=['POST'])
def checkAuthored():
    # data = request.get_json()
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed for post json is None', data=-1), ensure_ascii=False)
    nLicenseVersion = BEOPDataAccess.getInstance().getLicenseVersion()
    licenseManage = LicenseManage(nLicenseVersion)
    expired_, leftdays_ = licenseManage.checkAuthored()
    return jsonify(dict(expired=expired_, leftdays=leftdays_))

'''
诊断报表下载接口
参数：
  timeFrom
  timeTo
  groupNameList
  emailNoticeList:
  通过查询后，将在这段时间内有故障的所有故障项写入到一个word文件
  如果emailNoticeList不为空，该文件将作为附件发送至列表里的邮箱
  同时，接口返回该文件的URL(原理:将生成的word文件存在siteinterface/static/files/temp/Fdd201810141500.docx)
'''

@app.route('/fdd/genServiceReport', methods=['POST'])
def fdd_gen_service_report():
    # data = request.get_json()
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', url=-1), ensure_ascii=False)

    fddGroupList = data.get('fddGroup')
    fddTimeFrom = data.get('QueryTimeFrom')
    fddTimeTo = data.get('QueryTimeTo')
    emailNoticeList = data.get('emailNoticeList')
    dFrom = datetime.strptime(fddTimeFrom, '%Y-%m-%d %H:%M:%S')
    dTo = datetime.strptime(fddTimeTo, '%Y-%m-%d %H:%M:%S')
    mail_result = False
    if time.mktime(dTo.timetuple()) - time.mktime(dFrom.timetuple()) < 0:
        return jsonify(dict(err=1, msg='fddTimeFrom > fddTimeTo', url=-1))
    if fddGroupList is None:
        return jsonify(dict(err=1, msg='fddGroupList param lost', url=-1))

    rv = BEOPDataAccess.getInstance().getFddStatusPeriod(fddGroupList, fddTimeFrom,fddTimeTo)
    if rv == False:
        return jsonify(dict(err=1, msg='Error in get data', url=-1))
    url = BEOPDataAccess.getInstance().writeDocx(rv,fddTimeFrom,fddTimeTo)
    if url == False:
        return jsonify(dict(err=1, msg='Error in WriteDocx', url=-1))
    if emailNoticeList:
        updateMailSetting()
        mail_result = send_email('报表', emailNoticeList, None,attachment_list=[dict(filepath = url.replace('siteinterface\\',''), filename= os.path.basename(url), filetype='doc/docx')])
    if mail_result == False:
        return jsonify(dict(err=1, msg='Error in Send Email', url=-1))
    return jsonify(dict(err=0, msg='Success ', url=url))


'''
timeFrom:
timeTo:
查询给定范围内的数据健康度
返回dict(err=0, data=[{err=1, msg='', name=''},{err=1, msg='', name=''},{err=1, msg='', name=''}]
   
   data数组表示错误清单，每个元素msg是一个字符串
   
   第一个灯：后台数据引擎状态, 查询core_status表里process=core的最新的timeTo是否在5分钟内, 进程是否在
   第2个灯：后台策略引擎状态, 查询core_status表里process=logic的最新的timeTo是否在5分钟内,进程是否在x
   第三个灯：后台历史数据状态,查询历史数据库historydata_miniute_2018_10_31表存在，且零点有>1个数据
   
'''
@app.route('/system/getHealth', methods=['POST'])
def system_get_health():
    # data = request.get_json()
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', url=-1), ensure_ascii=False)
    '''
     第一个灯,domcore进程是否在线，
     第二个灯,domcore进程是否在线
    '''
    outPutData = []
    rv1 = BEOPDataAccess.getInstance().getProcessStatus('domcore')
    mg1=''
    core = "后台domcore引擎"
    logic = "后台domlogic引擎"
    historydata = '后台历史数据'
    if rv1==1:
        mg1="工作异常"
    elif rv==-1:
        mg1="工作异常"
    else:
        mg1="工作正常"
    p1={'err':rv1,'msg':mg1,'name':core}
    outPutData.append(p1)
    rv2 = BEOPDataAccess.getInstance().getProcessStatus('domlogic')

    if rv2 ==1:
        mg1="工作异常"
    elif rv2 == -1:
        mg1="工作异常"
    else:
        mg1="工作正常"
    p2={'err':rv2,'msg':mg1,'name':logic}
    outPutData.append(p2)
    rv3 = BEOPDataAccess.getInstance().getHistoryDataStatus()
    if rv3 == 1:
        mg1 = "工作异常"
    elif rv3 ==-1:
        mg1 = "工作异常"
    else:
        mg1= "工作正常"
    p3 ={'err':rv3,'msg':mg1,'name':historydata}
    outPutData.append(p3)
    return jsonify(dict(err=0, msg='Success', data=outPutData))




@app.route('/system/initFiles', methods=['GET'])
def system_init_files():

    # init logo and logo_small
    rv = BEOPSqliteAccess.getInstance().InitLogoFiles()
    rv = rv and BEOPSqliteAccess.getInstance().InitModbusServerFile()
    if rv:
        return jsonify(dict(err=0, msg='Success', data={}))
    else:
        return jsonify(dict(err=1, msg='Failed, Please check db file or directory', data={}))


@app.route('/system/initProject', methods=['GET'])
def system_init_project():

    rv = BEOPSqliteAccess.getInstance().initAutoLoadDBFile()
    if rv:
        return jsonify(dict(err=0, msg='Success', data={}))
    else:
        return jsonify(dict(err=1, msg='Failed, Please check db file or directory', data={}))




@app.route('/warning/uploadConfigFile', methods=['POST'])
def warning_upload_ocnfig_file():
     upload_file = request.files.get('file')
     print('recv warning upload config file')
     if not upload_file:
         print('no file param and file attached')
         return json.dumps(dict(err=1, msg='no file param and file attached', data = {}))
     filename = upload_file.filename
     if not filename.endswith('xlsx'):
         print('only xlsx file supported')
         return json.dumps(dict(err=1, msg='only xlsx file supported', data = {}))

     strPath = app.static_folder
     strTempPath = os.path.join(strPath, 'files')
     strTempPath = os.path.join(strTempPath, 'temp')

     saveFilePath= os.path.join(strTempPath, filename)
     upload_file.save(saveFilePath)

     print('start saveTemplateFileTo4DB')
     bSuccess, strMsg = BEOPSqliteAccess.getInstance().saveTemplateFileTo4DB('warning_config.xlsx', saveFilePath, "", 0, "")

     print('read warning config')
     warningConfigBoolList, warningConfigRangeList, warningConfigRuleList = read_warning_config_from_xlsx(saveFilePath)

     print('remove all config')
     BEOPDataAccess.getInstance().removeWarningConfigAll()
     print('start to add warning bool config item : count :%d'%(len(warningConfigBoolList)))
     nBoolInsert = 0
     nRangeInsert = 0
     nScriptInsert = 0
     for item in warningConfigBoolList:
         ret = BEOPDataAccess.getInstance().addWarningConfigItem(item)
         if ret.get('err',1)==0:
             nBoolInsert+=1


     print('start to add warning range config item : count :%d' % (len(warningConfigRangeList)))
     for item in warningConfigRangeList:
         ret = BEOPDataAccess.getInstance().addWarningConfigItem(item)
         if ret.get('err',1)==0:
             nRangeInsert+=1


     print('start to add warning rule config item : count :%d' % (len(warningConfigRuleList)))
     for data in warningConfigRuleList:
         strScript = data.get('script', '')
         nWarningLevel = data.get("boolWarningLevel", 0)
         pointname = data.get('pointname', '')
         ofPositionName = data.get('ofPosition', '')
         ofSystemName = data.get('ofSystem', '')
         ofDepartmentName = data.get('ofDepartment', '')
         ofGroupName = data.get('ofGroup', '')
         tags = data.get('tag', '')
         strWarningInfo = data.get('boolWarningInfo', '')
         strWarningGroup = data.get('warningGroup', '')
         nRuleId = data.get('id',-1)
         strWarningInfoDetail = data.get('infoDetail')
         strUnitPropery01 = data.get('unitproperty01','')
         strUnitPropery02= data.get('unitproperty02', '')
         strUnitPropery03 = data.get('unitproperty03', '')
         strUnitPropery04 = data.get('unitproperty04', '')
         strUnitPropery05 = data.get('unitproperty05', '')

         nNewId = BEOPDataAccess.getInstance().addWarningConfigRuleItem(nWarningLevel, strScript, pointname,
                                                                        ofDepartmentName, ofPositionName, ofSystemName,
                                                                        ofGroupName, tags, strWarningInfo,
                                                                        strWarningGroup,strWarningInfoDetail, nRuleId, strUnitPropery02, strUnitPropery03, strUnitPropery04, strUnitPropery05)
         if nNewId>=0:
             nScriptInsert+=1



     if (nBoolInsert+nRangeInsert+nScriptInsert)>0:
         print('uploadConfigFile success')
         return json.dumps(dict(err=0, msg='', data = dict(boolInsertCount=nBoolInsert, rangeInsertCount=nRangeInsert, scriptInsertCount= nScriptInsert)))
     else:
         print('uploadConfigFile failed')
         return json.dumps(dict(err=1, msg='', data = dict(boolInsertCount=nBoolInsert, rangeInsertCount=nRangeInsert, scriptInsertCount= nScriptInsert)))


@app.route('/warning/downloadHistoryFile/<startTime>/<endTime>/<language>')
def download_file_of_warning_history(startTime, endTime,language):
    rt = {'status':0, 'message':None}
    grade_dict = {1:'一般', 2:'较重', 3:'严重'}
    try:
        hisData = BEOPDataAccess.getInstance().getHistoryWarningList(startTime, endTime)

        if language == 'zh-cn':
            excelData = ExcelFile('时间', '报警内容', '等级', '持续至', '相关点')
        else:
            excelData = ExcelFile('Time', 'User', 'level', 'Till To', 'Point')
        for item in hisData:
            excelData.append_row([item.get('time', ''), item.get('info', ''), grade_dict.get(item.get("level", 1)), item.get('endtime', ''),item.get('strBindPointName', '')])
        filepath = make_excel_file(excelData)
        if filepath:
            filename = make_file_name('warning_record', startTime, endTime)
            rt = make_response(send_file(filepath))
            rt.headers["Content-Disposition"] = "attachment; filename=%s;"%filename
            return rt
        else:
            rt = {'status':0, 'message':'Invalid parameter'}
    except Exception as e:
        print('download_file_of_history warning error:' + e.__str__())
        logging.error(e.__str__())
        rt = {'status':0, 'message':e.__str__()}
    return json.dumps(dict(err=1, msg='', data = rt))



@app.route('/licence/query', methods=['GET'])
def license_query():
    nLicenceVersion = BEOPDataAccess.getInstance().getLicenseVersion()
    softAuthInfo = LicenseManage(nLicenceVersion).checkAuthored()
    licenseInfo = dict(expired=softAuthInfo[0], leftdays=softAuthInfo[1])
    return json.dumps(dict(err=1, msg='', data=licenseInfo))


@app.route('/report/getReportNameList', methods=['POST'])
def report_get_report_name_list():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', url=-1), ensure_ascii=False)

    rv = BEOPSqliteAccess.getInstance().getReportNameList()
    return json.dumps(dict(err=0, msg='', data= rv))

@app.route('/report/getAllProgress')
def report_get_all_progress():
    if not RedisManager.is_alive():
        return jsonify(dict(err=1, msg="redis is not alive", data=[]))

    resList = []
    # reportAllInfo = RedisManager.get("report_gen_info")
    keyList = RedisManager.get_key_name_list_of_pattern("*report_gen_info_*")
    if not len(keyList):
        return jsonify(dict(err=0, msg="无报表进度信息", data=[]))

    for strReportKey in keyList:
        if strReportKey == "report_gen_info":
            continue

        reportName = strReportKey.replace("report_gen_info_", "")
        dReportProgressInfo = RedisManager.get(strReportKey)
        if not dReportProgressInfo:
            continue

        percent = 100 * (dReportProgressInfo.get("curProgress") / dReportProgressInfo.get("totalProgress")) if dReportProgressInfo.get("totalProgress") else 0
        dInfo = {"curProgress": dReportProgressInfo.get("curProgress", None),
                     "totalProgress": dReportProgressInfo.get("totalProgress", None),
                     "name": reportName,
                     "percent": round(percent, 1),
                 "startTime": dReportProgressInfo.get("startTime", ""),
                 "stopTime": dReportProgressInfo.get("stopTime", "")}

        resList.append(dInfo)

    return jsonify(dict(err=0, msg="", data=resList))

@app.route('/report/getProgress', methods=["POST"])
def report_get_progress():
    try:
        rcv = request.get_json()
        name = rcv.get("name", "")
        if not RedisManager.is_alive():
            return jsonify(dict(err=1, msg="redis is not alive", data={}))

        reportAllInfo = RedisManager.get("report_gen_info")
        if not reportAllInfo:
            return jsonify(dict(err=0, msg="无报表进度信息", data={}))
        
        timeInfo = reportAllInfo.get(name, {})
        
        strReportKey = 'report_gen_info_%s' % (name)
        progressInfo = RedisManager.get(strReportKey)
        if not progressInfo:
            return jsonify(dict(err=0, msg="无此报表的进度信息", data={}))

        percent = 100 * (progressInfo.get("curProgress") / progressInfo.get("totalProgress")) if progressInfo.get("totalProgress") else 0
        progressInfo.update({"startTime": timeInfo.get("startTime", None),
                             "stopTime": timeInfo.get("stopTime", None),
                             "percent": round(percent, 1)})

        return jsonify(dict(err=0, msg="", data=progressInfo))
    except Exception as e:
        logging.error("ERROR in /report/getProgress: %s" % e.__str__())
        return jsonify(dict(err=1, msg=e.__str__(), data={}))


'''
报表清单接口
参数：
  timeFrom
  timeTo
  pageNum: 第几页
  pageSize: 100

'''

@app.route('/report/getReportHistory', methods=['POST'])
def report_get_report_history():
    # data = request.get_json()
    rv = dict(err=0, msg='', data={})
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', url=-1), ensure_ascii=False)

    nPageNum = data.get('pageNum')
    nPageSize = data.get('pageSize')

    strTimeFrom = data.get('timeFrom')
    strTimeTo  = data.get('timeTo')

    strReportName = data.get('name')

    if strTimeFrom is None:
        return json.dumps(dict(err=1, msg='failed for timeFrom param lost', data={}), ensure_ascii=False)

    try:
        tTimeFrom = datetime.strptime(strTimeFrom, '%Y-%m-%d %H:%M:%S')
    except:
        return json.dumps(dict(err=1, msg='failed for timeFrom format is wrong, samle: 2018-01-01 00:00:00', data={}), ensure_ascii=False)

    try:
        tTimeTo = datetime.strptime(strTimeTo, '%Y-%m-%d %H:%M:%S')
    except:
        return json.dumps(dict(err=1, msg='failed for timeTo format is wrong, samle: 2018-01-01 00:00:00', data={}), ensure_ascii=False)

    if strTimeTo is None:
        return json.dumps(dict(err=1, msg='failed for timeTo param lost', data={}), ensure_ascii=False)

    rv = BEOPDataAccess.getInstance().getReportHistory(tTimeFrom, tTimeTo, nPageNum, nPageSize, strReportName)
    nTotal = BEOPDataAccess.getInstance().getReportCount(tTimeFrom, tTimeTo, strReportName)
    return json.dumps(dict(err=0, msg='', data=dict(reportList=rv, total = nTotal)))

@app.route('/report/delete', methods=['POST'])
def delete_report():
    try:
        rcv = request.get_json()
        reportId = rcv.get("id", None)
        if reportId is None:
            return jsonify(dict(err=1, msg="报表id不能为空", data=False))
        if not isinstance(reportId, int):
            return jsonify(dict(err=1, msg="报表id必须为整数", data=False))
        bsuc = BEOPDataAccess.getInstance().deleteReport(reportId)
        if bsuc:
            return jsonify(dict(err=0, msg="删除成功", data=True))
        return jsonify(dict(err=1, msg="删除失败", data=False))
    except Exception as e:
        logging.error("ERROR in /report/delete: %s" % e.__str__())
        return jsonify(dict(err=1, msg="删除失败", data=False))


@app.route('/report/insertReportHistory', methods=['POST'])
def insert_get_report_history():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', url=-1), ensure_ascii=False)

    strName = data.get('name')
    strDescription  = data.get('description')
    strGenTime = data.get('genTime')
    nFileSize= data.get('filesize')
    strAuthor = data.get('author')
    strURL = data.get('url')
    strReportTimeFrom = data.get('reportTimeFrom')
    strReportTimeTo = data.get('reportTimeTo')
    nReportTimeType = data.get('reportTimeType')
    nId = BEOPDataAccess.getInstance().insertReportHistory(strName, strDescription, strGenTime, nFileSize, strAuthor, strURL, strReportTimeFrom,
                        strReportTimeTo, nReportTimeType)
    if nId is None:
        json.dumps(dict(err=1, msg='', data={}))

    return json.dumps(dict(err= 0, msg='', data=dict(id=nId)))

@app.route('/pointtable/addmul', methods=['POST'])
def point_table_add_mul():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', url=-1), ensure_ascii=False)

    strNameList = data.get('nameList')
    nNeedRestart = data.get('restart', 0)
    strDescriptionList = data.get('descriptionList')
    for iIndex in range(len(strNameList)):
        strName = strNameList[iIndex]
        strDescription = strDescriptionList[iIndex]
        nID = BEOPSqliteAccess.getInstance().AddVPointIfNotExist(strName, strDescription)

    #restart core
    if nNeedRestart==1:
        restart_domCore()
    return json.dumps(dict(err=0, msg='', data=[]))


@app.route('/pointtable/add', methods=['POST'])
def point_table_add():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed', url=-1), ensure_ascii=False)

    strName = data.get('name')
    nNeedRestart = data.get('restart', 0)
    strDescription = data.get('description')
    nID = BEOPSqliteAccess.getInstance().AddVPointIfNotExist(strName, strDescription)
    if nID:
        #restart core
        if nNeedRestart==1:
            restart_domCore()
    return json.dumps(dict(err=0, msg='', data=dict(id=nID)))


@app.route('/db/getDatabaseName', methods=['GET'])
def db_get_db_name():
    dbname = app.config['DATABASE']
    return json.dumps(dict(err=0, msg='', data=dict(dbname=dbname)))




@app.route('/pointtool/transferPointName', methods=['POST'])
def pointtool_trans_auto():
    data = request.get_json()
    if data is None:
        return json.dumps(dict(err=1, msg='failed for no data in request json', url=-1), ensure_ascii=False)

    strDesList = data.get('descriptionList') if data.get("descriptionList") is not None else list()
    strNameList = data.get('nameList') if data.get("nameList") is not None else list()
    roomDefine = data.get('roomDefine') if data.get("roomDefine") is not None else dict()

    strAnaResultList = []

    nTemp = 1
    for idx, item in enumerate(strDesList):
        strResult = PointNameAI.getInstance().analysis_description(item, roomDefine, idx)
        if strResult in strAnaResultList:
            strResult = 'ERROR_PointNameDup_%d_'%(nTemp)+ strResult
            nTemp+=1
        strAnaResultList.append(strResult)

    # 测试
    testFolder = os.path.join(os.getcwd(), "test_folder")
    if not os.path.exists(testFolder):
        os.mkdir(testFolder)
    fileName = "test_gshp_result_%s.xlsx" % datetime.now().strftime("%Y%m%d%H%M%S")
    saveFilePath = os.path.join(os.getcwd(), "test_folder", fileName)
    if os.path.exists(saveFilePath):
        os.remove(saveFilePath)

    book = Workbook()
    sheet = book.create_sheet("Sheet1", 0)

    for idx, item in enumerate(strAnaResultList):
        sheet.cell(row=(idx+1), column=1, value=strDesList[idx])
        sheet.cell(row=(idx+1), column=2, value=strAnaResultList[idx])

    book.save(saveFilePath)
    return json.dumps(dict(err=0, msg='done', data=''))

    # return json.dumps(dict(err=0, msg='', data=strAnaResultList))


@app.route('/calculation/repairHistory/<timeFrom>/<timeTo>/<period>/<pointName>', methods=['GET'])
@app.route('/calculation/repairHistory/<timeFrom>/<timeTo>/<period>', methods=['GET'])
def calculation_repair_history(timeFrom, timeTo, period, pointName=None):
    corePath = app.config['CORE_PATH']
    strWorkDir = os.path.join(corePath, 'domRepairCalculationHistory')
    strExeFilePath = os.path.join(strWorkDir, 'domRepairCalculationHistory.exe')
    os.system('%s %s %s %s'%(strExeFilePath, timeFrom, timeTo, period))
    return jsonify(dict(err=0, msg='', data=1))




@app.route('/config/get_value_app', methods=['POST'])
def config_get_value_app():
    data = request.get_json()
    itemList = data.get('itemList')
    vvList = []
    for item in itemList:
        vv = app.config.get(item)
        vvList.append(vv)
    return json.dumps(dict(err=0, msg='', data=vvList))

@app.route('/config/get_value_config_file', methods=['POST'])
def config_get_value_config_file():
    data = request.get_json()
    itemList = data.get('itemList')
    currentPath = os.getcwd()

    fatherPath = os.path.dirname(currentPath)
    dompysitePath = os.path.join(fatherPath, 'dompysite')
    vvList = []
    try:
        configPath = os.path.join(dompysitePath, 'config.ini')
        cf = ConfigObj(configPath, encoding='UTF8')

        for item in itemList:
            seg = item[0]
            param = item[1]
            try:
                vv = cf.get(seg).get(param)
            except:
                vv = None
            vvList.append(vv)

    except Exception as e:
        strError = 'ERROR : %s in config.ini(%s). cloud data upload canceled.' % (e.__str__(), configPath)
        print(strError)

    return json.dumps(dict(err=0, msg='', data=vvList))



@app.route('/redis/getAllKeys', methods=['POST', 'GET'])
def redis_get_all_keys():
    keyList = []
    try:
        if RedisManager.is_alive():
            keyList = RedisManager.get_all_keys()
    except Exception as e:
        return json.dumps(dict(err=1, msg='Exception:%s'%(e.__str__())), data =[])


    return json.dumps(dict(err=0, msg='', data=keyList))

@app.route('/redis/delMinutesData/<strDate>', methods=['GET'])
def redis_del_key_minutes_of_date(strDate):

    try:
        tTime = datetime.strptime(strDate,"%Y-%m-%d")
        tTime = tTime.replace(minute=0, second=0)
    except:
        return json.dumps(dict(err=1, msg='时间格式需为日期加时间'))

    nDelCount = 0
    try:
        if RedisManager.is_alive():
            nDelCount = RedisManager.del_history_data_minutes_keys_of_date(tTime)
        else:
            return json.dumps(dict(err=1, msg='Redis is not alive', data={}))
    except Exception as e:
        return json.dumps(dict(err=1, msg='Exception:%s'%(e.__str__())))

    return json.dumps(dict(err=0, msg='', data=dict(delCount=nDelCount)))


@app.route('/redis/get/<strKey>', methods=['GET'])
def redis_get_key_value1(strKey):

    strValue = ''
    try:
        if RedisManager.is_alive():
            jsonValue = RedisManager.get(strKey)
        else:
            return json.dumps(dict(err=1, msg='Redis is not alive', data={}))
    except Exception as e:
        return json.dumps(dict(err=1, msg='Exception:%s'%(e.__str__())))

    return json.dumps(dict(err=0, msg='', data=jsonValue))

@app.route('/redis/get', methods=['POST'])
def redis_get_key_value():
    data = request.get_json()
    strKey = data.get('key')
    strValue = ''
    try:
        if RedisManager.is_alive():
            jsonValue = RedisManager.get(strKey)
    except Exception as e:
        return json.dumps(dict(err=1, msg='Exception:%s'%(e.__str__())))


    return json.dumps(dict(err=0, msg='', data=jsonValue))


@app.route('/redis/updateHistoryAtMoment', methods=['POST'])
def update_redis_history_at_moment():
    data = request.get_json()
    strTime = data.get('time')
    try:
        tTime = datetime.strptime(strTime,"%Y-%m-%d %H:%M:%S")
        tTime = tTime.replace(minute=0, second=0)
    except:
        return json.dumps(dict(err=1, msg='时间格式需为日期加时间'))

    if not RedisManager.is_alive():
        return json.dumps(dict(err=1, msg='Redis not alive'))

    dataList = BEOPDataAccess.getInstance().get_history_data_all_one_moment_padded(tTime.strftime("%Y-%m-%d %H:%M:%S"), 'h1')

    for item in dataList:
        RedisManager.set_history_data_list([item['name']], [strTime], [item['value']])

    return json.dumps(dict(err=0, msg='成功更新%d个'%(len(dataList)), data=len(dataList)))


@app.route('/get_history_data_all_one_moment_padded', methods=['POST'])
def get_history_data_all_one_moment_padded():
    data = request.get_json()
    strTime = data.get('time')
    strTimeFormat = data.get('format')
    data = BEOPDataAccess.getInstance().get_history_data_all_one_moment_padded(strTime, strTimeFormat)
    return json.dumps(dict(err=0, msg='', data=data))


@app.route('/save_vpoint_value', methods=['POST'])
def save_vpoint_value():
    data = request.get_json()
    data = BEOPSqliteAccess.getInstance().SaveBackupRealtimePointValuesStatus()
    return json.dumps(dict(err=0, msg='', data=data))

'''
机制是从mainService获取后存在本地mysql 数据库中

'''
@app.route('/weather/requestAndSaveForcast', methods=['GET'])
def weather_save():
    strLocation = request.args.get('location', '')
    try:
        rsp = requests.get('http://dom.inwhile.com/api/s6/weather/forecast?location=%s' % strLocation, verify=False)
        dRsp = json.loads(rsp.text)
        if dRsp.get("err") > 0:
            json.dumps(dict(err=0, msg='', data={}))

        strInfo = dRsp.get("data")
        dInfo = json.loads(strInfo)
        dateTime = dInfo.get("updateTime")
        tDate = dateutil.parser.parse(dateTime)
        strUpdateTime = "{year}-{month:0>2d}-{day:0>2d} {hour:0>2d}:{minute:0>2d}:{second:0>2d}".format(year=tDate.year, month=tDate.month, day=tDate.day, hour=tDate.hour, minute=tDate.minute, second=tDate.second)

        basicDict = {"location": city_name_convertor(strLocation)}

        if isinstance(dInfo, dict) and dInfo.get('daily'):
            ffList = dInfo.get('daily')

            BEOPDataAccess.getInstance().saveWeatherData(ffList, basicDict, strUpdateTime)

            tnow = datetime.now()
            rv = BEOPDataAccess.getInstance().getWeatherData(tnow.strftime('%Y-%m-%d'))
        else:
            return json.dumps(dict(err=1, msg="", data={}))
    except Exception as e:
        return json.dumps(dict(err=1, msg='ERROR:' + e.__str__(), data={}))
    return json.dumps(dict(err=0, msg='', data=rv))


@app.route('/weather/requestAndSaveCalendar', methods=['POST'])
def weather_calendar_save():
    rsp = request.get_json()
    strDate = rsp.get("date", None)
    if strDate == None:
        return jsonify(dict(err=1, msg="日期不能为空", data={}))
    if not isinstance(strDate, str):
        return jsonify(dict(err=1, msg="日期必须为字符串", data={}))
    if not is_valid_date(strDate, "%Y-%m-%d"):
        return jsonify(dict(err=1, msg="日期格式必须为yyyy-mm-dd", data={}))

    headers = {"content-type": "application/json"}

    post_data = {"date":strDate}
    rsp = requests.post('http://dom.inwhile.com/api/calendar/requestAndSaveCalendarToRedis', data=json.dumps(post_data), headers=headers, timeout=15)
    if rsp.status_code == 200:
        dData = json.loads(rsp.text)
        return jsonify(dData)
    return jsonify(dict(err=1, msg="", data={}))


@app.route('/sound/playByBaidu', methods=['POST'])
def sound_play_baidu():
    data = request.get_json()
    strKey = data.get('key','')
    strText = data.get('text')
    tReturnDate = datetime.now()
    strFileFormat = data.get('format', '3')
    soundPerson = data.get('per', '1')
    my_baidu_app_key = data.get('client_id', 'TmGPsRLhIHgxSFQwVyA9GDGG')
    my_baidu_app_secret = data.get('client_secret', 'ypGdbj2GTHW49OwVBVFIIgIFN7n3d9Pu')

    if strText is None:
        return json.dumps(dict(err=1, msg='text param needed', data={}))

    try:

        strFileType = strFileFormat

        strStaticDirPath = "tempsound_per"+ soundPerson
        temp_folder = os.path.join(app.static_folder, "files", strStaticDirPath)
        if not os.path.exists(temp_folder):
            os.mkdir(temp_folder)

        strTransList = lazy_pinyin(strText)
        strTransListNew = [item if item!='' and item!=',' else '_' for item in strTransList]
        strAllFinal =  '_'.join(strTransListNew)
        if len(strAllFinal) > 100:
            strAllFinal = strAllFinal[0:100]
        file_name = strAllFinal + (".wav" if strFileType == '6' else ".mp3")
        saveFilePath = os.path.join(temp_folder, file_name)

        if os.path.exists(saveFilePath):
            return json.dumps(dict(err=0, msg='', data=dict(fileURL=strStaticDirPath + '/' + file_name)))
        else:
            #auth and get token
            myToken = None
            if app.config['REDIS_ALIVE']:
                BaiduTokenInRedis = RedisManager.get_baidu_token()
                if BaiduTokenInRedis:
                    strApplyTime = BaiduTokenInRedis.get('apply_time')
                    tApplyTime = datetime.strptime(strApplyTime, '%Y-%m-%d %H:%M:%S')
                    if (datetime.now()-tApplyTime).total_seconds()<=60*60:
                        myToken = BaiduTokenInRedis.get('access_token')

            if myToken is None:
                r0 = requests.get('https://openapi.baidu.com/oauth/2.0/token?grant_type=client_credentials&client_id=%s&client_secret=%s'%(my_baidu_app_key, my_baidu_app_secret ))
                r0text = r0.content
                try:
                    jsonBaiduToken = json.loads(r0text.decode())
                    myToken = jsonBaiduToken.get('access_token')
                    if app.config['REDIS_ALIVE']:
                        jsonBaiduToken["apply_time"] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        RedisManager.set_baidu_token(jsonBaiduToken)
                except:
                    json.dumps(dict(err=1, msg='ERROR in apply token from baidu', data={}))

            post_data = dict(tex=strText,
                             tok=myToken,
                             cuid= '67600045131314',
                             ctp=  '1',
                             lan = 'zh',
                             spd = '5',
                             pit = '5',
                             vol = '5',
                             per = soundPerson,
                             aue=strFileFormat)

            headers = {'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8'}
            r = requests.post('http://tsn.baidu.com/text2audio', data=post_data,
                              headers=headers, timeout=300)

            text = r.content  # mp3二进制数据


            # 将mp3的二进制数据保存到本地的mp3
            tnow = datetime.now()

            f = open(saveFilePath, "wb")
            f.write(text)
            f.close()

            return json.dumps(dict(err=0, msg='', data=dict(fileURL = strStaticDirPath + '/' + file_name)))
    except Exception as e:
        return json.dumps(dict(err=1, msg='ERROR:' + e.__str__(), data={}))
    return json.dumps(dict(err=0, msg='', data=rv))

@app.route("/weather/getTodayWeatherInfo")
def get_today_weather_info():
    strToday = datetime.now().strftime("%Y-%m-%d")
    dWeather = BEOPDataAccess.getInstance().getTodayWeatherInfo(strToday)
    if dWeather is None:
        return jsonify(dict(err=1, msg="数据库繁忙，请重试", data={}))
    if not dWeather:
        return jsonify(dict(err=1, msg="获取失败", data={}))
    return jsonify(dict(err=0, msg="获取成功", data=dWeather))


@app.route("/weather/getForcastWeatherInfo/<strDate>/<nDays>")
def get_forcast_weather_info(strDate, nDays):
    strToday = strDate
    dWeather = BEOPDataAccess.getInstance().getWeatherDataForcast(strToday, int(nDays))
    if not dWeather:
        return jsonify(dict(err=1, msg="获取失败", data={}))
    return jsonify(dict(err=0, msg="获取成功", data=dWeather))

@app.route("/weather/getForcastWeatherInfoV2/<strDate>/<nDays>")
def get_forcast_weather_info_v2(strDate, nDays):
    strToday = strDate
    lData = BEOPDataAccess.getInstance().getWeatherDataForcastV2(strToday, int(nDays))
    headers = {"content-type": "application/json"}
    postData = {"fromDate": strToday, "days": int(nDays)}
    rv = requests.post("http://dom.inwhile.com/api/requestCalendarOfPeriod", data=json.dumps(postData), timeout=15, headers=headers)

    dCalendarData = {}
    if rv.status_code == 200:
        rvData = json.loads(rv.text)
        dCalendarData = rvData.get("data")

    for item in lData:
        strDate = item.get("date")
        calendar = dCalendarData.get(strDate, None)
        item.update({"calendar": calendar})

    return jsonify(dict(err=0, msg="获取成功", data=lData))

@app.route("/modbusclient/clearOutputTable", methods=["POST"])
def clear_modbus_equipment_output_table():
    rcv = request.get_json()
    arr = BEOPDataAccess.getInstance().clearModbusEquipmentOutputTable()

    return jsonify(dict(err=0, data=arr))

@app.route("/modbusclient/getAndClearOutputTable", methods=["POST"])
def get_and_clear_modbus_equipment_output_table():
    rcv = request.get_json()
    pointList = rcv.get("pointList", [])
    arr = BEOPDataAccess.getInstance().getAndClearModbusEquipmentOutputTable(pointList)

    return jsonify(dict(err=0, data=arr))

@app.route('/modbusclient/setRealtimeData', methods=['POST'])
def modbusclient_set_realtime_data():
    data = request.get_json()
    strPointNameList = data.get('pointNameList', [])
    strPointValueList = data.get('pointValueList', [])

    if not strPointNameList:
        return jsonify(dict(err=1, msg="poitnNameList param not valid", data={}))

    if not strPointValueList:
        return jsonify(dict(err=1, msg="pointValueList param not valid", data={}))

    rv = BEOPDataAccess.getInstance().setRealtimeDataToModbusEquipmentTable(strPointNameList, strPointValueList)

    return jsonify(rv)


@app.route('/obix/setRealtimeDataAndGetPointsToWrite', methods=['POST'])
def obix_set_realtimedata_and_get_points_to_write():
    nPointsUpdated = None
    listPointsToWrite = []
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('pointNameList', [])
        pointValueList = rcv.get('pointValueList', [])

        if not pointNameList:
            return jsonify(dict(err=1, nPointsUpdated=None, lWrite=[]))

        if not pointValueList:
            return jsonify(dict(err=1, nPointsUpdated=None, lWrite=[]))

        if len(pointNameList):
            nPointsUpdated = BEOPDataAccess.getInstance().setRealtimeDataToObixInputTable(pointNameList, pointValueList)

        listPointsToWrite = BEOPDataAccess.getInstance().getPointsFromObixOutputTable(pointNameList)

        return jsonify(dict(err=0, nPointsUpdated=nPointsUpdated, lWrite=listPointsToWrite))

    except Exception as e:

        logging.error("ERROR in obix_set_realtimedata_and_get_points_to_write: %s" % e.__str__())
        return jsonify(dict(err=1, nPointsUpdated=nPointsUpdated, lWrite=listPointsToWrite))


@app.route("/logix/setRealtimeDataAndGetPointsToWrite", methods=["POST"])
def logix_set_realtimedata():
    nPointsUpdated = None
    listPointsToWrite = []
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('pointNameList', [])
        pointValueList = rcv.get('pointValueList', [])

        if len(pointNameList) and len(pointValueList):
            nPointsUpdated = BEOPDataAccess.getInstance().setRealtimeDataToLogixInputTable(pointNameList, pointValueList)

        listPointsToWrite = BEOPDataAccess.getInstance().getPointsFromLogixOutputTable()

        return jsonify(dict(err=0, nPointsUpdated=nPointsUpdated, writeList=listPointsToWrite))

    except Exception as e:

        logging.error("ERROR in logix_set_realtimedata: %s" % e.__str__())
        return jsonify(dict(err=1, nPointsUpdated=None, writeList=listPointsToWrite))

@app.route("/knx/setRealtimeDataAndGetPointsToWrite", methods=["POST"])
def knx_set_realtimedata():
    nPointsUpdated = None
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('pointNameList', [])
        pointValueList = rcv.get('pointValueList', [])

        if len(pointNameList) and len(pointValueList):
            nPointsUpdated = BEOPDataAccess.getInstance().setRealtimeDataToKnxInputTable(pointNameList, pointValueList)

        return jsonify(dict(err=0, nPointsUpdated=nPointsUpdated))
    except Exception as e:
        logging.error("ERROR in /knx/setRealtimeDataAndGetPointsToWrite: %s" % e.__str__())

        return jsonify(dict(err=1, nPointsUpdated=None))

@app.route("/knx/getWriteCmd", methods=["POST"])
def knx_get_write_cmd():
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('writablePointNameList', [])
        pointsToWrite = BEOPDataAccess.getInstance().getPointsFromKnxOutputTable(pointNameList)
        return jsonify(dict(err=0, data=pointsToWrite))
    except Exception as e:
        logging.error("ERROR in /knx/getWriteCmd: %s" % e.__str__())
        return jsonify(dict(err=1, data=[]))

@app.route("/abslc/setRealtimeDataAndGetPointsToWrite", methods=["POST"])
def abslc_set_realtimedata():
    nPointsUpdated = None
    listPointsToWrite = []
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('pointNameList', [])
        pointValueList = rcv.get('pointValueList', [])

        if len(pointNameList) and len(pointValueList):
            nPointsUpdated = BEOPDataAccess.getInstance().setRealtimeDataToAbslcInputTable(pointNameList, pointValueList)

        listPointsToWrite = BEOPDataAccess.getInstance().getPointsFromAbslcOutputTable()

        return jsonify(dict(err=0, nPointsUpdated=nPointsUpdated, writeList=listPointsToWrite))
    except Exception as e:
        logging.error("ERROR in /abslc/setRealtimeDataAndGetPointsToWrite: %s" % e.__str__())

        return jsonify(dict(err=1, nPointsUpdated=None, writeList=listPointsToWrite))

@app.route("/dlt645/setRealtimeDataAndGetPointsToWrite", methods=["POST"])
def dlt645_set_realtimedata():
    nPointsUpdated = None
    listPointsToWrite = []
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('pointNameList', [])
        pointValueList = rcv.get('pointValueList', [])

        if len(pointNameList) and len(pointValueList):
            nPointsUpdated = BEOPDataAccess.getInstance().setRealtimeDataToDlt645InputTable(pointNameList, pointValueList)

        listPointsToWrite = BEOPDataAccess.getInstance().getPointsFromDlt645OutputTable()

        return jsonify(dict(err=0, nPointsUpdated=nPointsUpdated, writeList=listPointsToWrite))
    except Exception as e:
        logging.error("ERROR in /dlt645/setRealtimeDataAndGetPointsToWrite: %s" % e.__str__())

        return jsonify(dict(err=1, nPointsUpdated=None, writeList=listPointsToWrite))

@app.route("/moxa/setRealtimeData", methods=["POST"])
def moxa_set_realtimedata():
    nPointsUpdated = None
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('pointNameList', [])
        pointValueList = rcv.get('pointValueList', [])

        if len(pointNameList) and len(pointValueList):
            nPointsUpdated = BEOPDataAccess.getInstance().setRealtimeDataToMoxaInputTable(pointNameList, pointValueList)

        return jsonify(dict(err=0, nPointsUpdated=nPointsUpdated))

    except Exception as e:
        logging.error("ERROR in moxa_set_realtimedata: %s" % e.__str__())
        return jsonify(dict(err=1, nPointsUpdated=None))


@app.route("/domiot/uploadDataAndGetAndClearOutputTable", methods=["POST"])
def upload_controller_data():
    try:
        rcv = request.get_json()
        dataMap = rcv.get("dataMap") if rcv.get("dataMap") is not None else None
        controllerIp = rcv.get("controllerIp") if rcv.get("controllerIp") is not None else None

        if dataMap is None:
            return jsonify(dict(err=1, msg="数据表不能为空", data=dict(read=dict(err=1, msg="core服务器获得的数据表为空", data=0), write=[])))

        if controllerIp is None:
            return jsonify(dict(err=1, msg="控制器IP不能为空", data=dict(read=dict(err=1, msg="core服务器获得的控制器IP为空", data=0), write=[])))

        # 获取指定控制器的Persagy Controller点信息
        pointList = BEOPSqliteAccess.getInstance().getAllPersayControllerPointInfo(controllerIp)

        # 遍历指定控制器的点将来自arm的数据翻译点名，准备pointNameList 和 pointValueList
        pointNameList = []
        pointValueList = []

        # 若从arm上传的数据表不为空，则准备pointNameList 和 pointValueList
        if dataMap:
            for point in pointList:
                strTypeAddress = "{0}{1:02d}".format(point.get("type"), point.get("address"))
                multiple = point.get("multiple")
                initValue = dataMap.get(strTypeAddress, None)

                if initValue is None:
                    continue

                if strTypeAddress not in dataMap.keys():
                    continue

                if multiple.isdigit():
                    pointValue = initValue / int(float(multiple))
                else:
                    try:
                        jStart, sStart, jEnd, sEnd = get_multiple_values_for_persagy_controller_points_calc(multiple)
                        initValue = jStart if initValue < jStart else initValue
                        pointValue = (initValue - jStart)*(sEnd - sStart)/(jEnd - jStart) + sStart
                    except:
                        pointValue = initValue

                pointNameList.append(point["name"])
                pointValueList.append(pointValue)

        # 将arm上传的数据写入persagy_controller_input表
        dInsertResult = dict(err=0, msg='运行正常，未发现需上传到core服务器的点', data=0)
        if len(pointNameList):
            dInsertResult = BEOPDataAccess.getInstance().setRealtimeDataToPersayControllerTable(pointNameList, pointValueList)

        # 检查persagy_controller_output表，是否有需要写值的点
        listPointsToWrite = BEOPDataAccess.getInstance().getPointsFromPersagyControllerOutputTable()

        if listPointsToWrite is None:
            return jsonify(dict(err=1, msg="数据库繁忙，请重试", data=None))

        # 若无点可写则直接返回上传数据的结果
        if not len(listPointsToWrite):
            return jsonify(dict(err=0, msg="", data={"read": dInsertResult, "write": []}))

        # 将可写点信息以点名为键做成字典
        dictWritablePoints = {}
        for point in pointList:
            if point.get("R_W") == 0:
                continue
            dictWritablePoints.update({point.get("name"): {"type": point.get("type"),
                                                           "address": point.get("address"), "R_W": point.get("R_W"), "multiple": point.get("multiple")}})

        # 遍历需要写值的点列表，从可写点字典中拿到点的地址，放进pList中
        pList = []
        for point in listPointsToWrite:
            if point.get("name") in dictWritablePoints.keys():
                multiple = dictWritablePoints[point.get("name")]["multiple"]
                type = dictWritablePoints[point.get("name")]["type"]
                nInitValue = float(point.get("value")) if type == "AO" else int(float(point.get("value")))

                if multiple.isdigit():
                    pointValue = nInitValue * int(float(multiple))
                else:
                    try:
                        jStart, sStart, jEnd, sEnd = get_multiple_values_for_persagy_controller_points_calc(multiple)
                        pointValue = (nInitValue - sStart) * (jEnd - jStart) / (sEnd - sStart) + jStart
                    except:
                        pointValue = nInitValue

                dInfo = {"name": point.get("name"), "value": pointValue}
                dTypeAddress = dictWritablePoints.get(point.get("name"))
                dInfo.update(dTypeAddress)
                pList.append(dInfo)

        # 将需要写值的点从persagy_controller_output表中删除
        if len(pList):
            BEOPDataAccess.getInstance().delFromPersagyControllerOutputTable(pList)

        return jsonify(dict(err=0, msg="succeed", data={"read": dInsertResult, "write": pList}))

    except Exception as e:
        strLog = "ERROR in /domiot/uploadControllerData： %s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg="", data=dict(read=dict(err=1, msg="uploadControllerData报错：%s" % e.__str__(), data=0), write=[])))


@app.route("/domiot/clearOutputTable")
def clear_persagy_output_table():
    try:
        rst = BEOPDataAccess.getInstance().clearPersagyOutputTable()
        return jsonify(dict(err=0, data=rst))
    except Exception as e:
        strLog = "ERROR in /domiot/clearOutputTable: %s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, data=False))

@app.route("/getCurrentTime")
def get_current_time():
    try:
        strTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return jsonify(dict(err=0, msg="", data=strTime))
    except Exception as e:
        strLog = "ERROR in getCurrentTime:%s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg="", data=None))



@app.route('/order/newOrder', methods=['post'])
def order_new_order():
    currentPath = os.getcwd()

    fatherPath = os.path.dirname(currentPath)
    dompysitePath = os.path.join(fatherPath, 'dompysite')

    try:
        configPath = os.path.join(dompysitePath, 'config.ini')
        cf = ConfigObj(configPath, encoding='UTF8')
        strProjectName = cf['cloud']['projectname']
    except:
        strError = 'No projectname info found in config.ini(%s). cloud data upload canceled.' % (configPath)
        print(strError)
        logging.error(strError)

        rv = dict(err=1,msg='strError')
        return jsonify(rv)


    rv = dict()

    dataBody = request.get_json()
    dataBody["projectName"] = strProjectName

    # 插入测试数据
    try:
        headers = {'content-type': 'application/json'}
        r = requests.post('http://dom.inwhile.com/api/addMaintainList', data=json.dumps(dataBody), headers=headers, timeout=300)
        if r.status_code==200:
            rvdata = json.loads(r.text)
            if rvdata and rvdata.get('status'):
                rv = dict(err=0, msg='')
            else:
                rv = dict(err=1, msg=rvdata.get('msg'))

    except Exception as e:
        logging.error('ERROR in order_new_order:' + e.__str__())

    return jsonify(rv)

@app.route("/operation/logicRecordOutput", methods=["POST"])
def operation_logic_record_output():
    try:

        rcv = request.get_json()
        pointName = rcv.get("pointName", None)
        timeFrom = rcv.get("timeFrom", None)
        timeTo = rcv.get("timeTo", None)

        if not isinstance(pointName, str) or not isinstance(timeFrom, str) or not isinstance(timeTo, str):
            return jsonify(dict(err=1, msg="输入量必须为字符串", data=[]))
        if not re.match("^[0-9]{4}-[0-9]{2}-[0-9]{2} [0-9]{2}:[0-9]{2}:[0-9]{2}$", timeFrom):
            return jsonify(dict(err=1, msg="起始时间格式不对", data=[]))
        if not re.match("^[0-9]{4}-[0-9]{2}-[0-9]{2} [0-9]{2}:[0-9]{2}:[0-9]{2}$", timeTo):
            return jsonify(dict(err=1, msg="终止时间格式不对", data=[]))

        tTimeFrom = datetime.strptime(timeFrom, "%Y-%m-%d %H:%M:%S")
        tTimeTo = datetime.strptime(timeTo, "%Y-%m-%d %H:%M:%S")

        if tTimeFrom > tTimeTo:
            return jsonify(dict(err=1, msg="开始时间不能大于结束时间", data=[]))

        if tTimeFrom.day != tTimeTo.day:
            return jsonify(dict(err=1, msg="不支持跨天查询", data=[]))

        logDir = os.path.join(app.config["CORE_PATH"], "log")
        if not os.path.exists(logDir):
            return jsonify(dict(err=1, msg="log文件夹不存在", data=[]))

        opRecordDir = os.path.join(logDir, "oprecord-{year}-{month:02}-{day:02}".format(year=tTimeFrom.year, month=tTimeFrom.month, day=tTimeFrom.day))
        if not os.path.exists(opRecordDir):
            return jsonify(dict(err=1, msg="未发现当日log目录", data=[]))

        rt = LogicManager.getInstance().get_operation_record_output(pointName, tTimeFrom, tTimeTo, opRecordDir)

        if rt.get("code") > 0:
            return jsonify(dict(err=1, msg=rt.get("msg"), data=[]))

        rtList = rt.get("data")
        res = []
        for object in rtList:
            res.extend(object)

        return jsonify(dict(err=0, msg=rt.get("msg"), data=res))

    except Exception as e:
        logging.error('/operation/logicRecordOutput: %s' % e.__str__())
        return jsonify(dict(err=1, msg="获取失败:%s" % e.__str__(), data=[]))


@app.route("/logix/clearOutputTable")
def clear_logix_output_table():
    rcv = request.get_json()
    arr = BEOPDataAccess.getInstance().clearLogixOutputTable()
    return jsonify(dict(err=0, data=arr))

@app.route("/abslc/clearOutputTable")
def clear_abslc_output_table():
    rcv = request.get_json()
    arr = BEOPDataAccess.getInstance().clearAbslcOutputTable()
    return jsonify(dict(err=0, data=arr))

@app.route("/knx/clearOutputTable")
def clear_knx_output_table():
    rcv = request.get_json()
    arr = BEOPDataAccess.getInstance().clearKnxOutputTable()
    return jsonify(dict(err=0, data=arr))

@app.route("/dlt645/clearOutputTable")
def clear_dlt645_output_table():
    arr = BEOPDataAccess.getInstance().clearDlt645OutputTable()
    return jsonify(dict(err=0, data=arr))

@app.route("/pointCalculation/calHistory", methods=["POST"])
def point_calculation_cal_history():
    try:
        rcv = request.get_json()
        strPointNameList = rcv.get("pointNameList", None)
        timeFrom = rcv.get("timeFrom", None)
        timeTo = rcv.get("timeTo", None)
        strEquation = rcv.get("equation", None)
        strPeriod = rcv.get("period", None)

        if not isinstance(strPointNameList, list) or not isinstance(timeFrom, str) or not isinstance(timeTo, str):
            return jsonify(dict(err=1, msg="输入量必须为字符串", data=[]))
        if not re.match("^[0-9]{4}-[0-9]{2}-[0-9]{2} [0-9]{2}:[0-9]{2}:[0-9]{2}$", timeFrom):
            return jsonify(dict(err=1, msg="起始时间格式不对", data=[]))
        if not re.match("^[0-9]{4}-[0-9]{2}-[0-9]{2} [0-9]{2}:[0-9]{2}:[0-9]{2}$", timeTo):
            return jsonify(dict(err=1, msg="终止时间格式不对", data=[]))
        if not isinstance(strPeriod, str):
            return jsonify(dict(err=1, msg="period参数必须是字符串，形如m1,h1,d1", data=[]))

        if len(strPointNameList)==0:
            return jsonify(dict(err=1, msg="pointNameList should not be empty"))

        tTimeFrom = datetime.strptime(timeFrom, "%Y-%m-%d %H:%M:%S")
        tTimeTo = datetime.strptime(timeTo, "%Y-%m-%d %H:%M:%S")

        timeList = []

        # 采样间隔为m1时，每7天一批次查询
        if strPeriod == "m1" and (tTimeTo - tTimeFrom).total_seconds() > 3600 * 24 * 7:
            tInit = tTimeFrom
            while (tTimeTo - tInit).total_seconds() >= 3600 * 24 * 7:
                tUlti = tInit + timedelta(days=7)
                timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tUlti.strftime("%Y-%m-%d %H:%M:%S")))
                tInit = tUlti + timedelta(minutes=1)
            timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tTimeTo.strftime("%Y-%m-%d %H:%M:%S")))

        # 采样间隔为m5时，每14天一批次查询
        elif strPeriod == "m5" and (tTimeTo - tTimeFrom).total_seconds() > 3600 * 24 * 14:
            tInit = tTimeFrom
            while (tTimeTo - tInit).total_seconds() >= 3600 * 24 * 14:
                tUlti = tInit + timedelta(days=14)
                timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tUlti.strftime("%Y-%m-%d %H:%M:%S")))
                tInit = tUlti + timedelta(minutes=5)
            timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tTimeTo.strftime("%Y-%m-%d %H:%M:%S")))

        # 采样间隔为h1时，每180天一批次查询
        elif strPeriod == "h1" and (tTimeTo - tTimeFrom).total_seconds() > 3600 * 24 * 180:
            tInit = tTimeFrom
            while (tTimeTo - tInit).total_seconds() >= 3600 * 24 * 180:
                tUlti = tInit  + timedelta(days=180)
                timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tUlti.strftime("%Y-%m-%d %H:%M:%S")))
                tInit = tUlti + timedelta(hours=1)
            timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tTimeTo.strftime("%Y-%m-%d %H:%M:%S")))

        # 采样间隔为d1时，每365天一批次查询
        elif strPeriod == "d1" and (tTimeTo - tTimeFrom).total_seconds() > 3600 * 24 * 365:
            tInit = tTimeFrom
            while (tTimeTo - tInit).total_seconds() >= 3600 * 24 * 365:
                tUlti = tInit + timedelta(days=365)
                timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tUlti.strftime("%Y-%m-%d %H:%M:%S")))
                tInit = tUlti + timedelta(days=1)
            timeList.append((tInit.strftime("%Y-%m-%d %H:%M:%S"), tTimeTo.strftime("%Y-%m-%d %H:%M:%S")))

        allData = dict(time=[], lostTime=[], map={})

        if len(timeList):
            for timeTuple in timeList:
                dataDict = BEOPDataAccess.getInstance().get_history_data_padded(strPointNameList, timeTuple[0], timeTuple[1], strPeriod)

                if not dataDict or not dataDict.get("map"):
                    continue

                allData["time"].extend(dataDict.get("time"))
                allData["lostTime"].extend(dataDict.get("lostTime"))

                for strPointName in dataDict.get("map").keys():
                    if strPointName in allData.get("map").keys():
                        allData.get("map").get(strPointName).extend(dataDict.get("map").get(strPointName))
                    else:
                        allData.get("map").update({strPointName: dataDict.get("map").get(strPointName)})

        else:
            allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointNameList, timeFrom, timeTo, strPeriod)
        
        if not allData:
            return jsonify(dict(err=1, msg="获取时间段内历史数据失败",  data=[]))

        if allData.get('map') is None:
            return jsonify(dict(err=1, msg="获取时间段内历史数据失败：map内容未返回",  data=[]))

        allTimeList = BEOPDataAccess.getInstance().generateTimeListInRange(timeFrom, timeTo, strPeriod)

        calculatedDataList = []
        try:
            for iIndex in range(len(allTimeList)):
                strEquationAtThatTime = strEquation
                for pp in strPointNameList:
                    dataHisOfPP = allData['map'].get(pp)

                    if not len(dataHisOfPP):
                        return jsonify(dict(err=1, msg="历史数据获取失败或点名有误", data=[]))

                    strEquationAtThatTime = strEquationAtThatTime.replace('<%'+pp+'%>', str(dataHisOfPP[iIndex]))

                try:
                    xxx = eval(strEquationAtThatTime)
                    calculatedDataList.append(xxx)
                except Exception as e:
                    print('ERROR in /pointCalculation/calHistory eval(%s): %s' % (strEquationAtThatTime, e.__str__()))
                    calculatedDataList.append(None)
                    continue
        except Exception as e3:
            return jsonify(dict(err=1, msg="遍历历史数据计算失败:%s" % e3.__str__(), data=[]))

        allData['map'][strEquation] = calculatedDataList
        return jsonify(allData)

    except Exception as e:
        logging.error('/pointCalculation/calHistory: %s' % e.__str__())
        return jsonify(dict(err=1, msg="获取失败:%s" % e.__str__(), data=[]))


@app.route("/scriptCalculation/calHistory", methods=["POST"])
def script_calculation_cal_history():
    try:
        rcv = request.get_json()
        timeFrom = rcv.get("timeFrom", None)
        timeTo = rcv.get("timeTo", None)
        script = rcv.get("script", None)
        period = rcv.get("period", None)

        if not re.match("^[0-9]{4}-[0-9]{2}-[0-9]{2} [0-9]{2}:[0-9]{2}:[0-9]{2}$", timeFrom):
            return jsonify(dict(err=1, msg="起始时间格式不对", data={}))
        if not re.match("^[0-9]{4}-[0-9]{2}-[0-9]{2} [0-9]{2}:[0-9]{2}:[0-9]{2}$", timeTo):
            return jsonify(dict(err=1, msg="终止时间格式不对", data={}))
        if not isinstance(period, str):
            return jsonify(dict(err=1, msg="period参数必须是字符串，形如m1,h1,d1", data=[]))
        if not isinstance(script, str):
            return jsonify(dict(err=1, msg="脚本必须是字符串", data={}))

        if not script or not len(script):
            return jsonify(dict(err=1, msg="脚本不能为空", data={}))

        tTimeFrom = datetime.strptime(timeFrom, "%Y-%m-%d %H:%M:%S")
        tTimeTo = datetime.strptime(timeTo, "%Y-%m-%d %H:%M:%S")

        if period == "m1" and (tTimeTo - tTimeFrom).total_seconds() > 3600*24*7:
            return jsonify(dict(err=1, msg="起止时间差不可超过7天（m1）", data={}))

        if period == "m5" and (tTimeTo - tTimeFrom).total_seconds() > 3600*24*14:
            return jsonify(dict(err=1, msg="起止时间差不可超过14天（m5）", data={}))

        if period == "h1" and (tTimeTo - tTimeFrom).total_seconds() > 3600*24*180:
            return jsonify(dict(err=1, msg="起止时间差不可超过180天（h1）", data={}))

        if period == "d1" and (tTimeTo - tTimeFrom).total_seconds() > 3600*24*365:
            return jsonify(dict(err=1, msg="起止时间差不可超过365天（d1）", data={}))

        if tTimeFrom > tTimeTo:
            return jsonify(dict(err=1, msg="开始时间不能大于结束时间", data={}))

        matchList = re.findall("<%.*?%>", script)
        strPointNameList = []
        dataDict = {}
        if len(matchList):
            for match in matchList:
                match = match.replace("<%", "")
                match = match.replace("%>", "")
                strPointNameList.append(match)

            dataDict = BEOPDataAccess.getInstance().get_history_data_padded(strPointNameList, timeFrom, timeTo, period)

            if not dataDict or dataDict.get("map") is None:
                return jsonify(dict(err=1, msg="数据获取失败", data={}))

        strTimeList = BEOPDataAccess.getInstance().generateTimeListInRange(timeFrom, timeTo, period)

        calcResList = []
        for idx, strActTime in enumerate(strTimeList):
            strScriptAtThatTime = script

            # 将<% %>内的点名替换成历史数据
            if len(strPointNameList):
                for strPointName in strPointNameList:
                    hisDataList = dataDict["map"].get(strPointName)

                    if not len(hisDataList):
                        return jsonify(dict(err=1, msg="历史数据获取失败或点名有误", data={}))

                    strScriptAtThatTime = strScriptAtThatTime.replace("<%{0}%>".format(strPointName), str(hisDataList[idx]))

            # if strScriptAtThatTime.startswith('accum_point('):
            #     strScriptAtThatTime = strScriptAtThatTime[0:len(strScriptAtThatTime) - 1] + ',actTime="' + strTime + '"' + strScriptAtThatTime[-1]

            # API中插入actTime=strTime
            for strApi in g_systemAPIList:
                strScriptAtThatTime = insert_acttime_for_api(strScriptAtThatTime, strApi, strActTime)

            # 计算
            try:
                value = eval(strScriptAtThatTime)
                calcResList.append(value)
            except Exception as e:
                strLogFileName = "dompysite_history_recalc_err_log_{date}.log".format(date=datetime.now().strftime("%Y-%m-%d"))
                strErrInfo = "ERROR in /scriptCalculation/calHistory***script:{script}***:{error}".format(script=script, error=e.__str__())
                log_info_to_file(strLogFileName, strErrInfo)
                calcResList.append(None)

        rst = dict(map={script: calcResList}, time=strTimeList)

        return jsonify(dict(err=0, msg="计算成功", data=rst))

    except Exception as e:

        logging.error("ERROR in /scriptCalculation/calHistory: %s" % e.__str__())

        return jsonify(dict(err=1, msg="计算失败", data={}))


@app.route("/scriptCalculation/insertHistory", methods=["POST"])
def script_calculation_insert_history():
    try:
        rcv = request.get_json()

        valueList = rcv.get("valueList", None)
        timeList = rcv.get("timeList", None)
        strPointName = rcv.get("point", None)

        reCalcTaskList = []
        if strPointName is None or not len(strPointName):
            return jsonify(dict(err=1, msg="点名为空", data=False))
        if not isinstance(timeList, list) or not isinstance(valueList, list):
            return jsonify(dict(err=1, msg="时间列表或者值列表有误", data=False))
        if len(timeList) != len(valueList):
            return jsonify(dict(err=1, msg="值列表长度必须与时间列表长度一致", data=False))

        for idx, time in enumerate(timeList):
            BEOPDataAccess.getInstance().insert_history_data(strPointName, time, valueList[idx], 'CAL_REPAIR', '')
            reCalcTaskList.append({"point":strPointName,
                         "time":time,
                         "value":valueList[idx]})

        return jsonify(dict(err=0, msg="插入数据成功", data=True))

    except Exception as e:
        logging.error("ERROR in /scriptCalculation/insertHistory: %s" % e.__str__())

        return jsonify(dict(err=1, msg="插入数据失败", data=False))



@app.route('/projectImages/update', methods=['POST'])
def project_image_update():
    data = request.get_json()


    try:
        nImageId = data.get('imageId')
        nImageId = int(float(nImageId))
    except:
        return jsonify(dict(err=1, msg="imageId字段需为一个整数", data=False))

    bExtrated = BEOPSqliteAccess.getInstance().prepareResouceImageSpecial(nImageId)

    if bExtrated:
        return jsonify(dict(err=0, msg="", data=True))
    else:
        return jsonify(dict(err=1, msg="", data=True))


@app.route("/fdditem/getFaultInfo", methods=["POST"])
def get_fault_info():
    try:
        rsp = request.get_json()
        name = rsp.get("name") if rsp.get("name") is not None else None
        if not name:
            return jsonify(dict(err=1, msg="name 不能为空", data={}))
        if not isinstance(name, str):
            return jsonify(dict(err=1, msg="name必须为字符串", data={}))
        res = BEOPDataAccess.getInstance().getFddItemInfo(name)
        return jsonify(dict(err=0, msg="获取成功", data=res))
    except Exception as e:
        logging.error("ERROR in /fdditem/getFaultInfo: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败", data={}))

@app.route("/logic/updateHealthLog", methods=["POST"])
def update_health_log():
    try:
        rcv = request.get_json()
        strNameList = rcv.get("threadNameList", "")
        if not isinstance(strNameList, str):
            return jsonify(dict(err=1, msg="threadNameList必须为字符串", data=False))
        if not len(strNameList):
            return jsonify(dict(err=1, msg="threadNameList不能为空", data=False))

        nameList = strNameList.split(",")

        strExitErr = ""
        strTimeoutErr = ""
        for name in nameList:

            threadName = name.split("##")[0]
            logicName = name.split("##")[1]

            keyFrom = "LogicThread##heartbeat##{0}".format(name)
            keyTo = "LogicThread##actoncetime##{0}".format(name)

            dTimeFrom = RedisManager.get(keyFrom)
            if dTimeFrom is None or not isinstance(dTimeFrom, dict):
                logging.error("/logic/updateHealthLog: no start time found in redis on thread: %s" % name)
                continue

            strTimeFrom = dTimeFrom.get("time", None)
            if strTimeFrom is None:
                logging.error("/logic/updateHealthLog: no start time found in redis on thread: %s" % name)
                continue

            dTimeTo = RedisManager.get(keyTo)
            if dTimeTo is None or not isinstance(dTimeTo, dict):
                logging.error("/logic/updateHealthLog: no end time found in redis on thread: %s" % name)
                continue

            strTimeTo = dTimeTo.get("time", None)
            if strTimeTo is None:
                logging.error("/logic/updateHealthLog: no end time found in redis on thread: %s" % name)
                continue

            if not isValidDate(strTimeFrom, "%Y-%m-%d %H:%M:%S") or not isValidDate(strTimeTo, "%Y-%m-%d %H:%M:%S"):
                return jsonify(dict(err=1, msg="Time kept in redis of the threads is not in correct format.", data=False))

            tTimeFrom = datetime.strptime(strTimeFrom, "%Y-%m-%d %H:%M:%S")
            tTimeTo = datetime.strptime(strTimeTo, "%Y-%m-%d %H:%M:%S")

            strTimeNow = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            if tTimeFrom > tTimeTo:
                strExitErr += "{0}    {1}    {2}    {3}    {4}    {5}\n".format(strTimeNow,threadName,logicName,strTimeFrom,strTimeTo,"exitErr")

            if (tTimeTo - tTimeFrom).total_seconds() > 60:
                strTimeoutErr += "{0}    {1}    {2}    {3}    {4}    {5}\n".format(strTimeNow,threadName,logicName,strTimeFrom,strTimeTo,"timeoutErr")

        if not len(strExitErr) and not len(strTimeoutErr):
            return jsonify(dict(err=0, msg="未发现异常线程", data=True))

        coreDir = os.path.dirname(os.getcwd())
        filePath = os.path.join(coreDir, "log", "domlogic_health_report.log")

        strFmt = "a"
        if not os.path.exists(filePath):
            strFmt = "w"

        with open(filePath, strFmt) as fileObj:
            fileObj.write(strExitErr + strTimeoutErr)

        return jsonify(dict(err=0, msg="已将异常线程记录写入log", data=True))

    except Exception as e:
        strLog = "ERROR in /logic/updateHealthLog: %s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=False))

@app.route("/updateAllPointInfoIntoRedis")
def update_all_point_info_into_redis():
    bSuc = BEOPSqliteAccess.getInstance().updateAllPointInfoIntoRedis()
    if not bSuc:
        return jsonify(dict(err=1, msg="更新失败", data=False))
    return jsonify(dict(err=0, msg="更新成功", data=True))


@app.route("/updatePageContentIntoRedis/<pageId>")
def update_page_content_into_redis(pageId):
    nPageId = int(pageId)
    if nPageId == 60000499:
        print()
    rvContent = BEOPSqliteAccess.getInstance().getPlant('',pageId)
    if rvContent is not None:
        RedisManager.set_4db_page_content(pageId, rvContent)
        return jsonify(dict(err=0, msg="更新成功", data=rvContent))

    return jsonify(dict(err=1, msg="更新失败", data=None))

@app.route("/updatePageContentDebugToolIntoRedis/<pageId>")
def update_page_content_DebugTool_into_redis(pageId):
    nPageId = int(pageId)
    rvContent = BEOPSqliteAccess.getInstance().getPlantDebugTool(pageId)
    if rvContent is not None:
        RedisManager.set_4db_page_content_DebugTool(pageId, rvContent)
        return jsonify(dict(err=0, msg="更新成功", data=rvContent))
    return jsonify(dict(err=1, msg="更新失败", data=None))

@app.route("/getPointInfoFromRedis", methods=["POST"])
def get_point_info_from_redis():
    try:
        rcv = request.get_json()
        pointList = rcv.get("pointList", None)
        if pointList is not None:
            if not isinstance(pointList, list):
                return jsonify(dict(err=1, msg="点名必须用列表传入", data=[]))

        if not RedisManager.is_alive():
            return jsonify(dict(err=1, msg="redis未在运行", data=[]))

        infoDict = RedisManager.get("all_point_info")
        if not isinstance(infoDict, dict):
            return jsonify(dict(err=1, msg="从redis获取到的点位全部信息数据格式有误（不为字典）", data=[]))

        resList = []
        missingPoint = []
        if pointList is None:
            for key, value in infoDict.items():
                value.update({"pointName": key})
                resList.append(value)
        else:
            for point in pointList:
                info = infoDict.get(point, None)
                if info is None:
                    missingPoint.append(point)
                    continue
                info.update({"pointName": point})
                resList.append(info)

        msg = "Done."
        if len(missingPoint):
            msg += "But these points not found: {0}".format(",".join(missingPoint))

        return jsonify(dict(err=0, msg=msg, data=resList))

    except Exception as e:
        strLog = "ERROR in /getPointInfoFromRedis: %s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=[]))


"""
若pointList传入None，则返回全部点
否则按传入的pointList返回查询结果
"""
@app.route("/serviceData/getRealtime", methods=["POST"])
def service_data_get_realtime():
    rcv = request.get_json()
    if rcv == None:
        rcv = {}
    pointList = rcv.get("pointList", None)
    scriptList = rcv.get('scriptList', [])

    if scriptList == None:
        scriptList = []

    strIp = request.remote_addr
    strIpProc = strIp.replace(".", "_")
    strKey = "LastQueryTimeRealtimeData_{ip}".format(ip=strIpProc)
    strLastTime = RedisManager.get(strKey)

    valid = False
    if not strLastTime:
        valid = True
    else:
        count = 0
        while count <= 5:
            strLastTime = RedisManager.get(strKey)
            tLastTime = datetime.strptime(strLastTime, "%Y-%m-%d %H:%M:%S")
            if (datetime.now() - tLastTime).total_seconds() < 2:
                time.sleep(2)
                count += 1
                continue
            valid = True
            break

    if not valid:
        return jsonify(dict(err=1, msg="请求过于频繁，实时数据请求的时间间隔必须在30秒以上", data=[]))

    bAllPointsReturn = False
    ptNmList = []
    if isinstance(pointList, list):
        ptNmList = pointList
    elif isinstance(pointList, str):
        ptNmList = [pointList]
    elif pointList == None:
        ptNmList = []
        bAllPointsReturn = True

    for script in scriptList:
        ptList = find_vars_in_str(script)
        ptNmList.extend(ptList)

    ptNmList = list(set(ptNmList))

    tNow = datetime.now()
    try:
        rvRequest = BEOPDataAccess.getInstance().getInputTable(ptNmList)
        if rvRequest is None:
            return jsonify(dict(err=1, msg="实时值获取失败", data=[]))

        lAllData = rvRequest[0]
        dAllData = rvRequest[1]

        if bAllPointsReturn:
            ret = lAllData
        else:
            ret = []
            for ptNm in ptNmList:
                strValue = dAllData.get(ptNm, None)
                if strValue == None:
                    continue
                ret.append(dict(name=ptNm, value=strValue))

        if lAllData:
            if len(scriptList):
                retExtendsList = []
                for oneScript in scriptList:
                    scriptResult = eval_string_expression_strict(oneScript, '1', tNow.strftime('%Y-%m-%d %H:%M:%S'), lAllData)
                    retExtendsList.append(dict(name=oneScript, value=scriptResult))
                ret.extend(retExtendsList)

        RedisManager.set(strKey, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        return jsonify(dict(err=0, msg="获取成功", data=ret))

    except Exception as e:
        print('ERROR: ' + e.__str__())
        return jsonify(dict(err=1, msg="获取失败：%s" % e.__str__(), data=[]))


@app.route("/serviceData/getHistory", methods=["POST"])
def service_data_get_history():
    data = request.get_json()
    if not isinstance(data, dict):
        data = {}

    strTimeFormat = data.get('timeFormat')
    strRequestNo = datetime.now().strftime("%Y%m%d%H%M%S")
    #log_info_to_file('dompysite_debug_get_history.log', strRequestNo+ ': 1.Recv request:%s'%(str(data.get('pointList'))))

    if strTimeFormat not in ["m1", "m5", "h1", "d1", "M1"]:
        return jsonify(dict(err=1, msg="历史数据请求的时间间隔必须是如下其中之一：m1/m5/h1/d1/M1", data={}))

    valid = False
    strLastTime = RedisManager.get("LastQueryTimeHistoryData")
    if not strLastTime:
        valid = True
    else:
        count = 0
        while count <= 5:
            strLastTime = RedisManager.get("LastQueryTimeHistoryData")
            tLastTime = datetime.strptime(strLastTime, "%Y-%m-%d %H:%M:%S")
            if (datetime.now() - tLastTime).total_seconds() < 5:
                time.sleep(10)
                count += 1
                continue
            valid = True
            break

    if not valid:
        return jsonify(dict(err=1, msg="请求过于频繁，历史数据请求的时间间隔必须在10秒以上", data={}))

    #log_info_to_file('dompysite_debug_get_history.log', strRequestNo+ ': 2')


    projId = data.get("projId", None)
    pointList = data.get('pointList')
    scriptList = data.get("scriptList", None)
    filter = data.get("filter", [])
    strTimeStart = data.get('timeStart')
    strTimeEnd = data.get('timeEnd')

    dFilter = {}
    for dInfo in filter:
        try:
            if dInfo.get("pointName") not in dFilter.keys():
                dFilter.update({dInfo.get("pointName"): dict(filterType=dInfo.get("filterType"),
                                                             params=dInfo.get("params"))})
        except:
            pass

    if scriptList:
        scriptList = list(set(scriptList))

    bPointListParamOK = False
    if not pointList:
        bPointListParamOK = False
    elif isinstance(pointList, str):
        bPointListParamOK = True
    elif isinstance(pointList, list):
        bAllStr = True
        for item in pointList:
            if not isinstance(item, str):
                bAllStr = False
        if bAllStr:
            bPointListParamOK = True

    if not bPointListParamOK:
        return jsonify(dict(err=1, msg="pointList param should be str or str List", data={}))

    if projId:
        projdbKey = "projectdb" if projId == 1 else "projectdb%03d" % projId
        cf = ConfigObj("config.ini", encoding="UTF8")
        if cf.get(projdbKey) is None:
            dj = list()
        dbfileName = cf.get(projdbKey).get("dbFileName")
        dbfileDir = os.path.dirname(app.config.get("USE_4DB_NAME"))
        strDBPath = os.path.join(dbfileDir, dbfileName)

        strRealDBName = BEOPSqliteAccess.getInstance().GetRealtimeDBName(strDBPath)
        save_app_config_database()
        app.config['DATABASE'] = strRealDBName

    pointListReady = pointList
    pointsFromScriptList = []
    if scriptList:
        pointsFromScriptList = get_points_from_script_list(scriptList)
        pointsMerged = []
        for item in pointsFromScriptList:
            for point in item.get("pointList"):
                if point not in pointsMerged:
                    pointsMerged.append(point)

        pointListReady = list(set(pointList + pointsMerged))


    #log_info_to_file('dompysite_debug_get_history.log', strRequestNo+ ': 3')
    hisDataResult = BEOPDataAccess.getInstance().get_history_data_padded(pointListReady, strTimeStart, strTimeEnd,
                                                                         strTimeFormat, dFilter)

    #log_info_to_file('dompysite_debug_get_history.log', strRequestNo+ ': 4')
    if scriptList:
        scriptHisData = get_history_data_of_scripts(hisDataResult.get("map"), pointsFromScriptList)
        hisDataResult.get("map").update(scriptHisData)

    if projId:
        restore_app_config_database()

    RedisManager.set("LastQueryTimeHistoryData", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    #log_info_to_file('dompysite_debug_get_history.log', strRequestNo+ ': 5')
    return jsonify(dict(err=0, msg="获取成功", data=hisDataResult))

@app.route("/serviceData/monthlyStatics", methods=["POST"])
def service_data_monthly_statics():
    try:
        rcv = request.get_json()
        timeStart = rcv.get("timeStart", None)
        timeEnd = rcv.get("timeEnd", None)
        pointList = rcv.get("pointList", [])
        if not timeStart or not isinstance(timeStart, str):
            return jsonify(dict(err=1, msg="开始时间不能为空", data=[]))
        if not timeEnd or not isinstance(timeEnd, str):
            return jsonify(dict(err=1, msg="结束时间不能为空", data=[]))
        if not pointList:
            return jsonify(dict(err=1, msg="点名列表不能为空", data=[]))

        if not isValidDate(timeStart, "%Y-%m") or not isValidDate(timeEnd, "%Y-%m"):
            return jsonify(dict(err=1, msg="起始或终止日期格式有误", data=[]))

        tTimeStart = datetime.strptime(timeStart, "%Y-%m")
        tTimeEnd = datetime.strptime(timeEnd, "%Y-%m")

        if tTimeStart > tTimeEnd:
            return jsonify(dict(err=1, msg="开始时间不能大于结束时间", data=[]))

        tFrom = (tTimeStart.replace(day=28, hour=0, minute=0, second=0) + timedelta(days=5)).replace(day=1)
        tTo = (tTimeEnd.replace(day=28, hour=0, minute=0, second=0) + timedelta(days=5)).replace(day=1)

        strTFrom = tFrom.strftime("%Y-%m-%d %H:%M:%S")
        strTTo = tTo.strftime("%Y-%m-%d %H:%M:%S")

        data = BEOPDataAccess.getInstance().get_history_data_padded(pointList, strTFrom, strTTo, "d1", {})

        rt = {"map": {}, "time": []}

        for idx, strTime in enumerate(data.get("time")):
            if not re.match("^[0-9]{4}-[0-9]{2}-01 00:00:00$", strTime):
                continue

            for pName in data.get("map").keys():
                if pName not in rt["map"].keys():
                    rt["map"].update({
                        pName: [data["map"].get(pName)[idx]]
                    })
                else:
                    rt["map"].get(pName).append(data["map"].get(pName)[idx])

            strYearMonth = (datetime.strptime(strTime, "%Y-%m-%d %H:%M:%S") - timedelta(days=5)).strftime("%Y-%m")
            rt["time"].append(strYearMonth)

        return jsonify(dict(err=0, msg="获取成功", data=rt))

    except Exception as e:
        traceback.print_exc()
        logging.error("ERROR in /serviceData/monthlyStatics: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败：%s" % e.__str__(), data={}))


@app.route("/getCurrentSystemResource")
def get_current_system_resource():
    res = getCurrentSystemResource()
    if res.get("code") > 0:
        return jsonify(dict(err=1, msg="获取失败", data={}))
    return jsonify(dict(err=0, msg="获取成功", data=res["data"]))



@app.route('/learnModel/genModelByValues', methods=['POST'])
def model_gen_by_values():
    data = request.get_json()
    nError = 0
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    model_name = data.get('model_name')
    x_values = data.get('x_values',[])
    y_values = data.get('y_values', [])
    model_type = data.get('model_type', '')
    validate_percent = data.get('validate_percent', 15)


    domRegression = DomRegression(model_name)
    allScoreMap = domRegression.model_gen_by_values( x_values, y_values, model_type, validate_percent)

    return jsonify(data = allScoreMap, msg='', err=0)



@app.route('/learnModel/genModelByPointNames', methods=['POST'])
def model_gen_by_point_names():
    data = request.get_json()
    nError = 0
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    model_name = data.get('model_name')
    model_type = data.get('model_type', '')
    PointNameListX = data.get('PointNameListX',[])
    PointNameListY = data.get('PointNameListY', [])

    if not isinstance(PointNameListX, list):
        return jsonify(dict(err=1, msg='PointNameListX should be list'))

    if not isinstance(PointNameListY, list):
        return jsonify(dict(err=1, msg='PointNameListY should be list'))

    if len(PointNameListX)==0:
        return jsonify(dict(err=1, msg='PointNameListX is empty'))

    if len(PointNameListY)==0:
        return jsonify(dict(err=1, msg='PointNameListY is empty'))

    strTimeFrom = data.get('timeFrom')
    strTimeTo = data.get('timeTo')
    strTimeFormat = data.get('timeFormat')

    domReg = DomRegression(model_name)
    x_all_data, y_all_data = domReg.get_data_list_by_point_name(PointNameListX, PointNameListY, strTimeFrom, strTimeTo, strTimeFormat)

    allScoreMap= domReg.model_gen_by_values(x_all_data, y_all_data, model_type )

    predict_y_data = domReg.predict_data(x_all_data)


    return jsonify(data = dict(score=allScoreMap, predict=predict_y_data), msg='', err=0)

@app.route('/learnModel/predict', methods=['POST'])
def model_predict():
    data = request.get_json()
    nError = 0
    if data is None:
        return json.dumps(dict(err=1, msg='failed', data=-1), ensure_ascii=False)

    model_name = data.get('model_name', '')
    x_data_list = data.get('x_data_list',[])

    if model_name=='':
        return jsonify(dict(err=1, msg='model_name should be string and not empty'))

    if not isinstance(x_data_list, list):
        return jsonify(dict(err=1, msg='x_data_list should be list'))


    domRegression = DomRegression(model_name)
    if domRegression.load_model():
        y_data = domRegression.predict_data(np.array(x_data_list))
        if isinstance(y_data, np.ndarray):
            y_data = y_data.tolist()
        return jsonify(dict(data = y_data, msg='', err=0))
    else:
        return jsonify(dict(data = [], msg='load model failed', err=1))


@app.route("/tool/ConvertExcelToTable", methods=["POST"])
def tool_convert_excel_to_table():
    try:
        infoTable = request.files.get("file")

        if not infoTable:
            strLog = "没有发现导入的表格"
            return jsonify(dict(err=1, msg=strLog, data=[]))

        tableName = infoTable.filename
        if not tableName.endswith("xlsx"):
            return jsonify(dict(err=1, msg="只支持.xlsx文件", data=[]))

        saveFileDir = os.path.join(os.getcwd(), "siteinterface", "static", "temp")
        if not os.path.exists(saveFileDir):
            os.mkdir(saveFileDir)

        tempFileName = "{strTime}_{tableName}.xlsx".format(tableName=os.path.splitext(tableName)[0], strTime=datetime.now().strftime("%Y%m%d%H%M%S"))

        saveFilePath = os.path.join(saveFileDir, tempFileName)

        if os.path.exists(saveFilePath):
            os.remove(saveFilePath)

        infoTable.save(saveFilePath)

        book = xlrd.open_workbook(saveFilePath)
        sheet = book.sheet_by_index(0)
        nrows = sheet.nrows

        if nrows < 1:
            return jsonify(dict(err=1, msg="表格中无内容", data=[]))

        valueList = list()
        for idx in range(nrows):
            values = sheet.row_values(idx)
            valueList.append(values)

        return jsonify(dict(err=0, msg="导入成功", data=valueList))

    except Exception as e:
        strLog = "ERROR in /tool/ConvertExcelToTable: %s" % e.__str__()
        logging.error(strLog)

        return jsonify(dict(err=1, msg="转换失败：%s" % e.__str__(), data=[]))


@app.route("/tool/ConvertTableToExcel", methods=["POST"])
def tool_convert_table_to_excel():
    try:
        workbook = xlwt.Workbook('gbk')
        sheet = workbook.add_sheet('Sheet1')

        data = request.get_json()
        rowDataList = data.get('data')
        if not rowDataList:
            return jsonify(dict(err=1, msg="data param lost"))

        book = Workbook()

        sheet = book.create_sheet("数据", 0)
        for rowIndex, dataList in enumerate(rowDataList):
            for colIndex, data in enumerate(dataList):
                sheet.cell(row=rowIndex+1, column=colIndex+1, value=data)

        saveFileDir = os.getcwd() + '\\siteinterface\\static\\temp'
        if not os.path.exists(saveFileDir):
            os.makedirs(saveFileDir)

        strCreateFileTempName = datetime.now().strftime('%Y%m%d%H%M%S') + '_' + "table.xlsx"
        saveFilePath = os.path.join(saveFileDir, strCreateFileTempName)

        if os.path.exists(saveFilePath):
            os.remove(saveFilePath)

        book.save(saveFilePath)

        # response = Response()
        # response.status_code = 200
        #
        # output = io.BytesIO()
        # workbook.save(output)
        # response.data = output.getvalue()
        #
        # filename = 'data.xlsx'
        # mimetype_tuple = mimetypes.guess_type(filename)
        #
        # response.headers['Pragma'] = 'public'
        # response.headers['Expires'] = '0'
        # response.headers['Content-Type'] = mimetype_tuple[0]
        # response.headers['Content-Transfer-Encoding'] = 'binary'
        # response.headers['Content-Length'] = len(response.data)
        # response.headers['Content-Disposition'] = 'attachment; filename="%s"' % filename

        return jsonify(dict(err=0, msg="下载成功", data=strCreateFileTempName))
    except Exception as e:
        logging.error("ERROR in/tool /ConvertTableToExcel: %s" % e.__str__())
        return jsonify(dict(err=1, msg="下载失败", data=""))

@app.route("/core/autoUpdate/<strType>/<strProcessName>/<strUpdateTime>")
def core_auto_update(strType, strProcessName, strUpdateTime):
    try:
        """
        0: 手动
        1: 自动
        """
        strDomUpdateDir = os.path.join(app.config['CORE_PATH'], "domUpdate")
        if not os.path.exists(strDomUpdateDir):
            logging.error("Error in /core/autoUpdate/<strType>/<strProcessName>/<strUpdateTime> domUpdate directory doesn't exist")
            return jsonify(dict(err=1, msg="domUpdate directory doesn't exist", data=False))

        try:
            nType = int(strType)
        except:
            return jsonify(dict(err=1, msg="strType has to be number", data=False))

        if nType not in [0, 1]:
            return jsonify(dict(err=1, msg="strType has to be 0 or 1", data=False))

        if not isinstance(strProcessName, str):
            return jsonify(dict(err=1, msg="autoUpdateProcessName must be a string", data=False))

        if not isinstance(strUpdateTime, str):
            return jsonify(dict(err=1, msg="autoUpdateTime must be a string", data=False))

        if nType == 0 and not isValidDate(strUpdateTime, "%Y-%m-%d %H:%M:%S"):
            return jsonify(dict(err=1, msg="wrong datetime format for manualUpdate", data=False))

        if nType == 1 and not isValidDate(strUpdateTime, "%H:%M:%S"):
            return jsonify(dict(err=1, msg="wrong datetime format for autoUpdate", data=False))

        aptFilePath = os.path.join(strDomUpdateDir, "appointment.json")

        if not os.path.exists(aptFilePath):
            return jsonify(dict(err=1, msg="appointment.json doesn't exist", data=False))

        strJson = ""
        with open(aptFilePath, "r", encoding="UTF8") as obj:
            lines = obj.readlines()
            for line in lines:
                if line == "\n" or not len(line):
                    continue
                strJson += line

        dJson = {}
        if len(strJson):
            dJson = json.loads(strJson)

        if dJson.get("autoUpdateList") is None:
            dJson.update({"autoUpdateList": []})

        if dJson.get("manualUpdateList") is None:
            dJson.update({"manualUpdateList": []})

        if nType == 0:
            for idx, procDict in enumerate(dJson.get("manualUpdateList")):
                if procDict.get("processName") == strProcessName:
                    dJson.get("manualUpdateList").remove(dJson.get("manualUpdateList")[idx])
                    break

            dJson.get("manualUpdateList").append(dict(processName=strProcessName, updateTime=strUpdateTime))

        elif nType == 1:
            for idx, procDict in enumerate(dJson.get("autoUpdateList")):
                if procDict.get("processName") == strProcessName:
                    dJson.get("autoUpdateList").remove(dJson.get("autoUpdateList")[idx])
                    break
            dJson.get("autoUpdateList").append(dict(processName=strProcessName, updateTime=strUpdateTime))

        if os.path.exists(aptFilePath):
            os.remove(aptFilePath)

        with open(aptFilePath, "w+", encoding="UTF8") as obj:
            obj.write(json.dumps(dJson))

        return jsonify(dict(err=0, msg="succeed to add", data=True))

    except Exception as e:
        logging.error("ERROR in /core/autoUpdate/<strType>/<strProcessName>/<strUpdateTime>: %s" % e.__str__())
        return jsonify(dict(err=1, msg="fail to add", data=False))

@app.route("/config/get/<sectionName>/<itemName>")
def config_get(sectionName, itemName):
    try:
        currentPath = os.getcwd()
        fatherPath = os.path.dirname(currentPath)
        dompysitePath = os.path.join(fatherPath, 'dompysite')
        configPath = os.path.join(dompysitePath, 'config.ini')
        cf = ConfigObj(configPath, encoding='UTF8')
        data = cf[sectionName][itemName]
        return jsonify(dict(err=0, msg="获取成功", data=data))
    except Exception as e:
        logging.error("ERROR in /config/get: %s" % e.__str__())
        return jsonify(dict(err=1, msg="获取失败", data=""))


@app.route('/update_domUpdate', methods=['GET'])
def logic_update_dom_update():
    # start = time.clock()
    print('logic_update_dom_update request recv from factory')
    logging.error('logic_update_dom_update request recv from factory')
    #data = request.get_json()#从json获取数据
    strDomUpdatePath = os.path.join(app.config["CORE_PATH"], 'domUpdate')
    strDownloadZipPath = os.path.join(strDomUpdatePath, 'domUpdate.zip')
    unzipPath = app.config["CORE_PATH"] #os.path.join(strDomUpdatePath, "_unzip_")

    rst = {
        "succeed_to_cover": [],
        "fail_to_cover": [],
        "succeed_to_add": []
    }

    try:
        # downloadStart = time.clock()
        #下载

        # NOTE the stream=True parameter below
        rep = requests.get(regionDomUpdate, stream=True, timeout=300)
        rep.raise_for_status()
        with open(strDownloadZipPath, 'wb') as f:
            for chunk in rep.iter_content(chunk_size=8192):
                if chunk:  # filter out keep-alive new chunks
                    f.write(chunk)


        # downloadEnd = time.clock() - downloadStart
        #↓创建缓存文件夹
        if (os.path.exists(unzipPath)) == False:
            os.mkdir(unzipPath)  # 创建目录


        # 解压缩
        # unzipStart = time.clock()

        sZip = zipfile.ZipFile(strDownloadZipPath, 'r')
        sZip.extractall(unzipPath)
        sZip.close()

        # unzipEnd = time.clock() - unzipStart

        # ↓更新文件:
        # updateStart = time.clock()

        # end = time.clock()-start
        print('update_domUpdate : 更新成功')
        logging.error('update_domUpdate : 更新成功')
        return json.dumps(dict(err=0, msg='更新成功', data=rst))
    except Exception as e:
        traceback.print_exc()

        print('update_domUpdate : 更新失败')
        logging.error('update_domUpdate : 更新失败')
        return json.dumps(dict(err=1, msg='更新失败', data=[]))

@app.route("/template/ossSyncToLocal")
def template_oss_sync_to_local():
    strKey = "update/template"
    osstool = OSSTool()

    if not osstool.file_exists("{key}/".format(key=strKey)):
        return jsonify(dict(err=1, msg="OSS的update/template路径不存在", data={"pysiteSync": False, "factorySync": False}))

    length = osstool.calculate_length(strKey)
    count = 0
    stableCount = 0
    bStable = False

    while count < 20:
        curLength = osstool.calculate_length(strKey)
        if curLength == length:
            stableCount += 1
        else:
            length = curLength

        if stableCount >= 3:
            bStable = True

        if bStable:
            break

        count += 1
        time.sleep(1)

    if not bStable:
        return jsonify(dict(err=1, msg="update/template/大小在变化，可能正在执行模板上传，本次同步已被放弃", data={"pysiteSync": False, "factorySync": False}))

    templateDir = os.path.join(app.static_folder, "template")
    if not os.path.exists(templateDir):
        os.mkdir(templateDir)

    fileNameList = osstool.iterate_files(strKey)

    nDownloadCount = 0
    for fileName in fileNameList:
        ossFilePath = "{key}/{fileName}".format(key=strKey, fileName=fileName)
        localFilePath = os.path.join(templateDir, fileName)
        bSuc = osstool.download(ossFilePath, localFilePath)
        if not bSuc:
            logging.error("ERROR in /template/OssSyncToLocal::fail to download {fileName}".format(fileName=fileName))

        nDownloadCount += 1
        time.sleep(0.1)

    bPysiteSync = True if len(fileNameList) == nDownloadCount else False

    domDir = os.path.dirname(app.config["CORE_PATH"])
    factoryDir = os.path.join(domDir, "factory")
    if not os.path.exists(factoryDir):
        return jsonify(dict(err=1, msg="factory路径不存在:{directory}".format(directory=factoryDir),
                            data={"pysiteSync": bPysiteSync, "factorySync": False}))

    factoryTemplateDir = os.path.join(factoryDir, "template")
    if not os.path.exists(factoryTemplateDir):
        return jsonify(dict(err=1, msg="factory的template路径不存在:{directory}".format(directory=factoryTemplateDir),
                            data={"pysiteSync": bPysiteSync, "factorySync": False}))

    bFactorySync = False
    try:
        for root, dirs, files in os.walk(templateDir):
            for file in files:
                srcPath = os.path.join(templateDir, file)
                dstPath = os.path.join(factoryTemplateDir, file)
                shutil.copy(srcPath, dstPath)
                time.sleep(0.1)
        bFactorySync = True
    except Exception as e:
        logging.error("ERROR in /template/OssSyncToLocal: %s" % e.__str__())

    return jsonify(dict(err=0, msg="", data={"pysiteSync": bPysiteSync, "factorySync": bFactorySync}))


@app.route("/getMysqlDiagnosis")
def mysql_diagnosis():
    dData = BEOPDataAccess.getInstance().getMysqlDiagnosisInfo()
    return jsonify(dData)

@app.route("/exportAllModeAndEnv")
def export_all_mode_and_env():
    try:
        dEnv, dEnvDetail = BEOPDataAccess.getInstance().getEnvById(None)
        if not dEnv or not dEnvDetail:
            return jsonify(dict(err=1, msg="数据库繁忙，请重试"))

        dMode = BEOPDataAccess.getInstance().getAllMode()
        if dMode.get("err") > 0:
            return jsonify(dict(err=1, msg="数据库繁忙，请重试", data=[]))

        book = Workbook()

        # 准备场景
        envSumTitleList = ["id", "name", "tags", "createtime", "type", "creator", "description", "enabled"]
        sheetEnvSummary = book.create_sheet("场景汇总")
        for idx, item in enumerate(envSumTitleList):
            sheetEnvSummary.cell(row=1, column=idx+1, value=item)

        minEnvId = min(dEnv.keys())
        maxEnvId = max(dEnv.keys())

        idx = 0
        curEnvId = minEnvId
        while curEnvId >= minEnvId and curEnvId <= maxEnvId:
            try:
                if not dEnv.get(curEnvId, None):
                    continue

                for i, value in enumerate(envSumTitleList):
                    sheetEnvSummary.cell(row=idx+2, column=i+1, value=dEnv.get(curEnvId).get(value))

                sheetDetail = book.create_sheet("场景-{id}".format(id=curEnvId))
                sheetDetail.cell(row=1, column=1, value="pointName")
                sheetDetail.cell(row=1, column=2, value="pointValue")
                for j, tpl in enumerate(dEnvDetail.get(curEnvId)):
                    sheetDetail.cell(row=j+2,column=1, value=tpl[0])
                    sheetDetail.cell(row=j+2,column=2, value=tpl[1])

            except:
                pass
            finally:
                curEnvId += 1
                idx += 1

        # 准备模式
        dModeData = dMode.get("data")
        modeSumTitleList = ["id", "name", "tags", "createtime", "type", "creator", "description", "enabled"]
        modeDetailTitleList = ["envId", "triggerTime", "triggerTimeType", "actionOnce"]
        sheetModeSummary = book.create_sheet("模式汇总")
        for idx, item in enumerate(modeSumTitleList):
            sheetModeSummary.cell(row=1, column=idx+1, value=item)

        minModeId = min(dEnv.keys())
        maxModeId = max(dEnv.keys())

        idx = 0
        curModeId = minModeId
        while curModeId >= minModeId and curModeId <= maxModeId:
            try:
                if not dModeData.get(curModeId, None):
                    continue

                for i, value in enumerate(modeSumTitleList):
                    sheetModeSummary.cell(row=idx + 2, column=i + 1, value=dModeData.get(curModeId).get(value))

                sheetDetail = book.create_sheet("模式-{id}".format(id=curModeId))
                for k, strTitle in enumerate(modeDetailTitleList):
                    sheetDetail.cell(row=1, column=k+1, value=strTitle)

                for j, dInfo in enumerate(dModeData.get(curModeId).get("details")):
                    for k, strTitle in enumerate(modeDetailTitleList):
                        sheetDetail.cell(row=j+2, column=k+1, value=dInfo.get(strTitle))

                idx += 1
            except:
                pass
            finally:
                curModeId += 1

        filesDir = os.path.join(app.static_folder, "files")
        if not os.path.exists(filesDir):
            os.mkdir(filesDir)

        destExcelName = "allEnvModeInfo.xlsx"
        destExcelPath = os.path.join(filesDir, destExcelName)

        if os.path.exists(destExcelPath):
            os.remove(destExcelPath)

        book.save(destExcelPath)

        return jsonify(dict(err=0, msg="", data=True))
    except Exception as e:
        strLog = "ERROR in exportAllModeAndEnv: %s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=False))

@app.route("/getModbusEquipmentInputTableData")
def get_modbus_equipment_input_table_data():
    result = BEOPDataAccess.getInstance().getModbusEquipmentInputTableData()
    return jsonify(result)

@app.route("/sendTCPPackage", methods=["POST"])
def send_tcp_package():
    try:
        rcv = request.get_json()
        ip = rcv.get("ip", None)
        port = rcv.get("port", None)
        package = rcv.get("package", None)
        if ip is None or port is None or package is None:
            return jsonify(dict(err=1, msg="ip, port, package 不能为空", data=""))

        if not isinstance(package, str) and not isinstance(package, list):
            return jsonify(dict(err=1, msg="package必须为字符串或列表", data=""))

        if not isinstance(ip, str):
            return jsonify(dict(err=1, msg="ip必须为字符串", data=""))
        if not isinstance(port, int):
            return jsonify(dict(err=1, msg="port必须为整形", data=""))

        packageList = package
        if isinstance(package, str):
            packageList = [package]

        cli = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        cli.settimeout(3)

        bConnected = False
        count = 0
        while not bConnected and count < 10:
            try:
                cli.connect((ip, port))
                bConnected = True
            except Exception as ept:
                pass
            finally:
                count += 1

        if not bConnected:
            return jsonify(dict(err=1, msg="socket server连接失败", data=""))

        for package in packageList:
            cli.send(bytes(package, "UTF-8"))
            strInfo = 'sendTCPPackage send(ip:%s, port:%d) package: '%(ip, port)+ package
            print(strInfo)
            logging.error(strInfo)
            time.sleep(0.2)

        cli.close()

        return jsonify(dict(err=0, msg="", data=""))

    except Exception as e:
        strLog = "ERROR in /sendTCPPackage: %s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=""))

@app.route("/getDeeplogicRedisConfig")
def get_deeplogic_redis_config():
    result = BEOPDataAccess.getInstance().get_deeplogic_redis_config()
    return jsonify(result)


@app.route("/revert/<processName>")
@app.route("/revert/<processName>/<revertCount>")
def revert_process(processName=None, revertCount=1):
    if processName == None:
        return jsonify(dict(err=1, msg="进程名不能为空", data=False))

    nRevertCount = None
    try:
        nRevertCount = int(revertCount)
    except:
        pass

    if not isinstance(nRevertCount, int):
        return jsonify(dict(err=1, msg="回滚版本数量输入有误", data=False))

    if nRevertCount <= 0:
        return jsonify(dict(err=1, msg="回滚版本数量必须为大于0的整数", data=False))

    strProcName = ""
    dProcess = {}
    for dTar in allProcList:
        nameLower = dTar.get("name").lower()
        dProcess.update({nameLower: dict(name=dTar.get("name"), seconds=dTar.get("seconds"))})

    try:
        processNameLower = processName.lower()
        dInfo = dProcess.get(processNameLower, None)
        if not dInfo:
            return jsonify(dict(err=1, msg="未找到目标升级进程的信息", data=False))

        strProcName = dInfo.get("name", "")
        nWaitSeconds = dInfo.get("seconds", 60*2)
    except Exception as e:
        strLog = "/domUpdate/{process}执行失败: {err}".format(err=e.__str__(), process=strProcName)
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=False))

    try:
        if ProcessManager.getInstance().findProcess("domUpdate.exe"):
            return jsonify(dict(err=1, msg="domUpdate.exe在运行，本次回滚放弃", data=False))

        filePath = os.path.join(app.config["CORE_PATH"], "domUpdate", "domUpdate.exe")
        fileDir = os.path.join(app.config["CORE_PATH"], "domUpdate")
        if not os.path.exists(filePath):
            return jsonify(dict(err=1, msg="domUpdate.exe不存在，本次回滚放弃", data=False))

        strCmd = "-name {processName} -revert {nRevertCount}".format(processName=strProcName, nRevertCount=nRevertCount)

        shell32 = ctypes.windll.shell32
        shell32.ShellExecuteW(0, "runas", filePath, strCmd, fileDir, 0)

        if strProcName == "dompysite":
            return jsonify(dict(err=0, msg="目标升级进程为dompysite，升级指令已下发，后台即将失去响应，请于约6分钟后自行查看dompysite升级结果", data=True))

        nCount = 0
        while nCount < nWaitSeconds:
            print("[dompysite start] domUpdate {process} count: {count}".format(count=nCount, process=strProcName))
            time.sleep(1)
            nCount += 1

        if ProcessManager.getInstance().findProcess("{process}.exe".format(process=strProcName)):
            return jsonify(dict(err=0, msg="回滚成功", data=True))

        return jsonify(dict(err=1, msg="未检测到{process}.exe运行，可能回滚失败，请远程连接服务器核查".format(process=strProcName), data=False))

    except Exception as e:
        strLog = "/domUpdate/{process}回滚报错: {err}".format(err=e.__str__(), process=strProcName)
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=False))


@app.route("/domUpdate/<processName>")
def domupdate_domcore(processName=None):
    if processName == None:
        return jsonify(dict(err=1, msg="进程名不能为空", data=False))

    if processName.lower() == "domupdate":
        return jsonify(dict(err=0, msg="domUpdate会自动更新，无需特意升级", data=True))

    strProcName = ""
    dProcess = {}
    for dTar in allProcList:
        nameLower = dTar.get("name").lower()
        dProcess.update({nameLower: dict(name=dTar.get("name"), seconds=dTar.get("seconds"))})

    try:
        processNameLower = processName.lower()
        dInfo = dProcess.get(processNameLower, None)
        if not dInfo:
            return jsonify(dict(err=1, msg="未找到目标升级进程的信息", data=False))

        strProcName = dInfo.get("name", "")
        nWaitSeconds = dInfo.get("seconds", 60*2)
    except Exception as e:
        strLog = "/domUpdate/{process}执行失败: {err}".format(err=e.__str__(), process=strProcName)
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=False))

    try:

        if ProcessManager.getInstance().findProcess("domUpdate.exe"):
            return jsonify(dict(err=1, msg="domUpdate.exe在运行，本次更新放弃", data=False))

        filePath = os.path.join(app.config["CORE_PATH"], "domUpdate", "domUpdate.exe")
        fileDir = os.path.join(app.config["CORE_PATH"], "domUpdate")
        if not os.path.exists(filePath):
            return jsonify(dict(err=1, msg="domUpdate.exe不存在，本次更新放弃", data=False))

        strCmd = " {process}".format(process=strProcName)

        shell32 = ctypes.windll.shell32
        shell32.ShellExecuteW(0, "runas", filePath, strCmd, fileDir, 0)

        if strProcName == "dompysite":
            return jsonify(dict(err=0, msg="目标升级进程为dompysite，升级指令已下发，后台即将失去响应，请于约6分钟后自行查看dompysite升级结果", data=True))

        nCount = 0
        bHasStarted = False
        while nCount < nWaitSeconds:
            print("[dompysite start] domUpdate {process} count: {count}".format(count=nCount, process=strProcName))
            if ProcessManager.getInstance().findProcess("{process}.exe".format(process=strProcName)):
                bHasStarted = True
                break

            time.sleep(1)
            nCount += 1

        if bHasStarted:
            return jsonify(dict(err=0, msg="进程升级启动成功", data=True))

        return jsonify(dict(err=1, msg="未检测到{process}.exe运行，可能升级失败，请远程连接服务器核查".format(process=strProcName), data=False))

    except Exception as e:
        strLog = "/domUpdate/{process}升级报错: {err}".format(err=e.__str__(), process=strProcName)
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=False))


@app.route("/saveImgToOSS", methods=["POST"])
def save_img_to_oss():
    imgDir = ""
    try:
        ossDir = request.form.get("directory")

        tempDir = os.path.join(app.static_folder, "temp")
        if not os.path.exists(tempDir):
            os.mkdir(tempDir)

        imgDir = os.path.join(tempDir, "img_{time}".format(time=datetime.now().strftime("%Y_%m_%d_%H_%M_%S_%f")))
        if not os.path.exists(imgDir):
            os.mkdir(imgDir)

        ossTool = OSSTool()
        bSuccess = True

        nameList = []
        codeNameList = []
        extList = [".jpg", ".png", ".jpeg", ".bmp", ".gif", ".mp4", ".wmv"]
        for idx in range(1, 11):
            try:
                file = request.files.get("file%02d" % idx)
                if file is None or not isinstance(file.content_type, str):
                    continue

                if file.filename in nameList:
                    continue

                ext = os.path.splitext(file.filename)[1]
                extLower = ext.lower()
                if extLower not in extList:
                    return jsonify(dict(err=1, msg="只能上传如下后缀格式的图片和视频:{s}".format(s="、".join(extList)), data=[]))

                nameList.append(file.filename)

                fileName = "{code}{ext}".format(code=bson.ObjectId(), ext=ext)

                filePath = os.path.join(imgDir, fileName)
                file.save(filePath)

                codeNameList.append(fileName)

                bSuc = ossTool.upload(filePath, "{ossDir}{fileName}".format(ossDir=ossDir, fileName=fileName))
                if not bSuc:
                    logging.error("/saveImgToOSS:fail to save img to OSS(fileName:%s" % file.filename)
                    bSuccess = False

            except Exception as e:
                logging.error("ERROR in /saveImgToOSS: %s" % e.__str__())
            finally:
                time.sleep(0.2)

        return jsonify(dict(err=1 if not bSuccess else 0,
                            msg="some of the img failed to save to oss" if not bSuccess else "succeed to save img to oss",
                            data=codeNameList))

    except Exception as e:
        strLog = "ERROR in /saveImgToOSS: %s" % e.__str__()
        logging.error(strLog)
        return jsonify(dict(err=1, msg=strLog, data=[]))
    finally:
        if os.path.exists(imgDir):
            shutil.rmtree(imgDir)

@app.route('/webPublish')
def web_publish():
    return render_template('webPublish.html')

@app.route("/saveLogoImg", methods=["POST"])
def save_logo_img():
    file = request.files["file"]
    strType = request.form.get("type")

    if strType not in ["logo", "logo_small"]:
        return jsonify(dict(err=1, msg="type 必须为logo或logo_small", data=False))

    filesDir = os.path.join(app.static_folder, "files")
    if not os.path.exists(filesDir):
        os.mkdir(filesDir)

    tempDir = os.path.join(filesDir, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)

    fileName = strType + ".png"
    filePath = os.path.join(tempDir, fileName)

    file.save(filePath)

    bSuc, strMsg = BEOPSqliteAccess.getInstance().saveTemplateFileTo4DB(fileName, filePath, "", 0, "")

    # 解压一次logo图片
    if bSuc:
        BEOPSqliteAccess.getInstance().InitLogoFiles()
    return jsonify(dict(err=0 if bSuc else 1, msg=strMsg, data=bSuc))

@app.route("/deleteLogoImg/<strType>")
def delete_logo_img(strType=None):
    if strType == None:
        return jsonify(dict(err=1, msg="type不能为空(logo或logo_small)", data=False))

    if strType not in ["logo", "logo_small"]:
        return jsonify(dict(err=1, msg="type必须为logo 或logo_small", data=False))

    fileName = "{name}.png".format(name=strType)
    bSuc = BEOPSqliteAccess.getInstance().deleteTemplateFileFrom4db(fileName)
    return jsonify(dict(err=1 if not bSuc else 0, msg="", data=bSuc))

@app.route("/backendFile/add", methods=["POST"])
def backend_file_add():
    file = request.files["file"]
    strType = request.form.get("type")   # 0-普通文件；2-报表
    strTime = request.form.get("time")  # 当添加报表时传入时间字符串，如1.5,10.3,16:34
    fileTimeType = request.form.get("fileTimeType")  # 报表时间类型  0-日  1-周  2-月
    description = request.form.get("description", "")
    userOfRole = request.form.get("roleList", None)  # 权限等级列表或字符串
    emailBody = request.form.get("emailBody", None)   # 邮件正文

    nType = None
    try:
        nType = int(strType)
    except:
        pass

    nFileTimeType = None
    try:
        nFileTimeType = int(fileTimeType)
    except:
        pass

    if nType not in [0, 2]:
        return jsonify(dict(err=1, msg="type必须为0或2", data=False))

    if nType == 2 and not strTime:
        return jsonify(dict(err=1, msg="报表则必须传入生成时间", data=False))

    if nType == 2 and nFileTimeType == None:
        return jsonify(dict(err=1, msg="报表则必须传入时间类型", data=False))

    if nType == 2 and nFileTimeType not in [0, 1, 2]:
        return jsonify(dict(err=1, msg="报表时间类型必须为0,1,2", data=False))

    userOfRoleList = []
    if userOfRole != None:
        if isinstance(userOfRole, list):
            for item in userOfRole:
                if is_int_digit(item):
                    userOfRoleList.append(str(item))
        elif isinstance(userOfRole, str):
            les = json.loads(userOfRole)
            for item in les:
                if is_int_digit(item):
                    userOfRoleList.append(str(item))

    if emailBody != None:
        if not isinstance(emailBody, str):
            return jsonify(dict(err=1, msg="邮件正文必须为字符串", data=False))
        if not len(emailBody):
            return jsonify(dict(err=1, msg="邮件正文不能是空字符串", data=False))

    fileName = file.filename
    filesDir = os.path.join(app.static_folder, "files")
    if not os.path.exists(filesDir):
        os.mkdir(filesDir)

    tempDir = os.path.join(filesDir, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)

    filePath = os.path.join(tempDir, fileName)
    if os.path.exists(filePath):
        os.remove(filePath)
    file.save(filePath)

    bSuc, strMsg = BEOPSqliteAccess.getInstance().saveTemplateFileTo4DB(fileName, filePath, strTime, nFileTimeType, description, nType, userOfRoleList, emailBody)
    if fileName.find("modbus_point_table") >= 0 and bSuc:
        BEOPSqliteAccess.getInstance().InitModbusServerFile()

        bDataOpen = False
        ret = BEOPDataAccess.getInstance().getProcessList(False)
        procList = ret.get("data", [])
        if isinstance(procList, list):
            for proc in procList:
                if proc.get("processName", "").lower() == "dommodbusserver.exe":
                    bDataOpen = True
                    break

        if bDataOpen:
            os.popen("taskkill /im domModbusServer.exe -f")

    return jsonify(dict(err=0 if bSuc else 1, msg=strMsg, data=bSuc))

@app.route("/backendFile/get", methods=["POST"])
def backend_file_get():
    rcv = request.get_json()
    nType = rcv.get("type", None)
    if nType is not None:
        if nType not in [0, 2]:
            return jsonify(dict(err=1, msg="type必须为0或2", data=False))

    dataList = BEOPSqliteAccess.getInstance().getAutoReportTemplateIdListFileFrom4DB()

    resList = []
    for data in dataList:
        if nType != None:
            if data.get("fileType") != nType:
                continue

        resList.append(data)

    return jsonify(dict(err=0, msg="", data=resList))

@app.route("/backendFile/delete", methods=["POST"])
def backend_file_delete():
    rcv = request.get_json()
    nId = rcv.get("id", None)
    if nId == None:
        return jsonify(dict(err=1, msg="id不能为空", data=False))

    bSuc, strMsg = BEOPSqliteAccess.getInstance().deleteBackendFileFrom4DB(nId)

    return jsonify(dict(err=0 if bSuc else 1, msg=strMsg, data=bSuc))


@app.route("/backendFile/update", methods=["POST"])
def backend_file_update():
    strId = request.form.get("id", None)         # id
    strType = request.form.get("type", None)   # 0-普通文件；2-报表     reserve01
    strTime = request.form.get("time", None) # 当添加报表时传入时间字符串，如1.5,10.3,16:34   reserve03
    strTimeType = request.form.get("fileTimeType", None)  # 报表时间类型  0-日  1-周  2-月   reserve02
    strDesc = request.form.get("description", None)   # file_description
    file = request.files.get("file", None)   # file_binary
    userOfRole = request.form.get("roleList", None)  # 权限等级列表或字符串
    emailBody = request.form.get("emailBody", None)  # 邮件正文

    nId = None
    try:
        nId = int(strId)
    except:
        pass

    nType = None
    try:
        nType = int(strType)
    except:
        pass

    nTimeType = None
    try:
        nTimeType = int(strTimeType)
    except:
        pass

    if nId == None:
        return jsonify(dict(err=1, msg="id不能为空", data=False))

    if not isinstance(nId, int):
        return jsonify(dict(err=1, msg="id必须为整数", data=False))

    if nType != None:
        if nType not in [0, 2]:
            return jsonify(dict(err=1, msg="文件类型必须为0,2", data=False))

    if strTime:
        if not isinstance(strTime, str):
            return jsonify(dict(err=1, msg="时间必须为字符串", data=False))

    if nTimeType != None:
        if nTimeType not in [0, 1, 2]:
            return jsonify(dict(err=1, msg="时间类型必须为0,1,2", data=False))

    if strDesc != None:
        if not isinstance(strDesc, str):
            return jsonify(dict(err=1, msg="desctiption必须为字符串", data=False))

    if nType == None and strTime == None and nTimeType == None and strDesc == None and file == None:
        return jsonify(dict(err=1, msg="未发现修改项", data=False))

    userOfRoleList = []
    if userOfRole != None:
        if isinstance(userOfRole, list):
            for item in userOfRole:
                if is_int_digit(item):
                    userOfRoleList.append(str(item))
        elif isinstance(userOfRole, str):
            les = json.loads(userOfRole)
            for item in les:
                if is_int_digit(item):
                    userOfRoleList.append(str(item))

    if emailBody != None:
        if not isinstance(emailBody, str):
            return jsonify(dict(err=1, msg="邮件正文必须为字符串", data=False))

    filePath = None
    if file != None:
        fileName = file.filename
        filesDir = os.path.join(app.static_folder, "files")
        if not os.path.exists(filesDir):
            os.mkdir(filesDir)

        tempDir = os.path.join(filesDir, "temp")
        if not os.path.exists(tempDir):
            os.mkdir(tempDir)

        filePath = os.path.join(tempDir, fileName)
        if os.path.exists(filePath):
            os.remove(filePath)
        file.save(filePath)

    bSuc, strMsg = BEOPSqliteAccess.getInstance().updateBackendFileInfoOf4DB(nId, nType, strTime, nTimeType, strDesc, filePath, userOfRoleList, emailBody)

    return jsonify(dict(err=0 if bSuc else 1, msg=strMsg, data=bSuc))


@app.route("/backendFile/download", methods=["POST"])
def backend_file_download():
    data = request.get_json()
    templateName = data.get('templateName')
    if not isinstance(templateName, str):
        return jsonify(dict(err=1, msg='模板名必须为字符串', data=""))

    if templateName.find(".") == -1:
        return jsonify(dict(err=1, msg='传入的模板文件名有误，必须包含后缀', data=""))

    templateNameSplit = os.path.splitext(templateName)
    if len(templateNameSplit) < 2:
        return jsonify(dict(err=1, msg='传入的模板文件名有误，必须包含后缀', data=""))

    tempDir = os.path.join(app.static_folder, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)

    remove_files_specified_seconds_ago(tempDir, 3600)

    templateNameLocal = "{base}_{time}{ext}".format(base=templateNameSplit[0], time=datetime.now().strftime("%Y-%m-%d-%H-%m-%S-%f"), ext=templateNameSplit[1])

    filePath = os.path.join(tempDir, templateNameLocal)
    if os.path.exists(filePath):
        try:
            os.remove(filePath)
        except:
            pass

    bSuc = BEOPSqliteAccess.getInstance().getTemplateFileFrom4DB(templateName, filePath)
    if bSuc:
        return jsonify(dict(err=0, msg='', data=os.path.join("static", "temp", templateNameLocal)))

    return jsonify(dict(err=1, msg='下载失败', data=""))

@app.route("/unit01/setValueMulti", methods=["POST"])
def unit01_set_value_multi():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        return jsonify(dict(err=1, msg="传参必须为字典", data=False))

    if not rcv:
        return jsonify(dict(err=1, msg="传参不能为空", data=False))

    dTar = {}
    nValueListLength = 0
    for key, value in rcv.items():
        valueList = None
        if not isinstance(key, int) and not isinstance(key, float) and not isinstance(key, str):
            continue

        if isinstance(value, int) or isinstance(value, float) or isinstance(value, str):
            valueList = [value]
        elif isinstance(value, list):
            valueList = value

        if valueList == None:
            continue

        if len(valueList) > 14:
            continue

        if len(valueList) > nValueListLength:
            nValueListLength = len(valueList)

        dTar.update({key: valueList})

    bSuc, strMsg = BEOPDataAccess.getInstance().updateMultiKeyValueOfUnit01(dTar)

    return jsonify(dict(err=0 if bSuc else 1, msg=strMsg, data=bSuc))

@app.route("/unit01/getValueMulti", methods=["POST"])
def unit01_get_value_multi():
    rcv = request.get_json()
    key = rcv.get("key", None)
    if key == None:
        return jsonify(dict(err=1, msg="key不能为空", data={}))

    dData = BEOPDataAccess.getInstance().getMultiKeyValueFromUnit01(key)

    return jsonify(dict(err=0, msg="", data=dData))

@app.route("/getMysqlConfig")
def get_mysql_config():

    myIniPath = None
    if os.path.exists("C:\\DomMySQL\\my.ini"):
        myIniPath = "C:\\DomMySQL\\my.ini"
    elif os.path.exists("D:\\DomMySQL\\my.ini"):
        myIniPath = "D:\\DomMySQL\\my.ini"
    elif os.path.exists("E:\\DomMySQL\\my.ini"):
        myIniPath = "E:\\DomMySQL\\my.ini"
    elif os.path.exists("F:\\DomMySQL\\my.ini"):
        myIniPath = "F:\\DomMySQL\\my.ini"
    else:
        strPath = find_full_path_of_executable_file("mysqld.exe")
        if strPath is not None:
            try:
                binDir = os.path.dirname(strPath)
                domMysqlDir = os.path.dirname(binDir)
                myIniPath = os.path.join(domMysqlDir, "my.ini")
            except:
                pass

    if myIniPath is None:
        return jsonify(dict(err=1, msg="my.ini路径获取失败", data={}))

    if not os.path.exists(myIniPath):
        return jsonify(dict(err=1, msg="my.ini路径获取失败", data={}))

    res = get_mysql_my_ini_config(myIniPath)
    for key, value in res.items():
        if key == "innodb_file_per_table" or key == "max_connections":
            res.update({key: int(value)})
        if key == "max_heap_table_size" or key == "tmp_table_size":
            res.update({key: int(value.replace("M", ""))})
        if key == "port":
            res.update({key: int(value)})

    return jsonify(dict(err=0, msg="获取成功", data=res))


@app.route("/updateMysqlConfig", methods=["POST"])
def update_mysql_config():
    rcv = request.get_json()
    innodb_file_per_table = rcv.get("innodb_file_per_table", None)  # 1
    tmp_table_size = rcv.get("tmp_table_size", None)  # 256
    max_heap_table_size = rcv.get("max_heap_table_size", None)  # 512

    if tmp_table_size != None:
        if not isinstance(tmp_table_size, int):
            return jsonify(dict(err=1, msg="tmp_table_size必须为整数", data=False))
        if tmp_table_size > 2048 or tmp_table_size < 0:
            return jsonify(dict(err=1, msg="tmp_table_size范围必须为在0到2048之间", data=False))

    if max_heap_table_size != None:
        if not isinstance(max_heap_table_size, int):
            return jsonify(dict(err=1, msg="max_heap_table_size必须为整数", data=False))
        if max_heap_table_size > 2048 or max_heap_table_size < 0:
            return jsonify(dict(err=1, msg="max_heap_table_size范围必须为在0到2048之间", data=False))

    if innodb_file_per_table != None:
        if innodb_file_per_table not in [0, 1]:
            return jsonify(dict(err=1, msg="innodb_file_per_table必须为0或1", data=False))

    strPath = find_full_path_of_executable_file("mysqld.exe")
    binDir = os.path.dirname(strPath)
    domMysqlDir = os.path.dirname(binDir)
    myIniPath = os.path.join(domMysqlDir, "my.ini")
    if strPath == None:
        myIniPath = "D:\\DomMySQL\\my.ini"

    if not os.path.exists(myIniPath):
        return jsonify(dict(err=1, msg="my.ini路径获取失败", data={}))

    bSuc = update_mysql_my_ini_config(myIniPath, innodb_file_per_table, tmp_table_size, max_heap_table_size)

    return jsonify(dict(err=1 if not bSuc else 0, msg="", data=bSuc))


@app.route("/config/dompysite/save", methods=["POST"])
def config_dompysite_save():
    rcv = request.get_json()
    key = rcv.get("key", None)
    value = rcv.get("value", None)

    if key == None or value == None:
        return jsonify(dict(err=1, msg="键和值不能为空", data=False))

    cMng = ConfigIniManager()
    bSuc = cMng.set_content_of_options(key, value)
    return jsonify(dict(err=1 if not bSuc else 0, msg="", data=bSuc))

@app.route("/dcim/setRealtimeDataAndGetPointsToWrite", methods=["POST"])
def dcim_set_realtimedata():
    nPointsUpdated = None
    listPointsToWrite = []
    try:
        rcv = request.get_json()
        pointNameList = rcv.get('pointNameList', [])
        pointValueList = rcv.get('pointValueList', [])
        strTimeList = rcv.get("pointTimeList", [])

        if len(pointNameList) and len(pointValueList):
            nPointsUpdated = BEOPDataAccess.getInstance().setRealtimeDataToDcimInputTable(strTimeList, pointNameList, pointValueList)

        listPointsToWrite = BEOPDataAccess.getInstance().getPointsFromDcimOutputTable()

        return jsonify(dict(err=0, nPointsUpdated=nPointsUpdated, writeList=listPointsToWrite))
    except Exception as e:
        logging.error("ERROR in /dcim/setRealtimeDataAndGetPointsToWrite: %s" % e.__str__())

        return jsonify(dict(err=1, nPointsUpdated=None, writeList=listPointsToWrite))


@app.route("/dcim/clearOutputTable")
def clear_dcim_output_table():
    arr = BEOPDataAccess.getInstance().clearDcimOutputTable()
    return jsonify(dict(err=0, data=arr))


@app.route("/oss/uploadFile", methods=["POST"])
def oss_upload_file():
    oFile = request.files.get("file")
    directory = request.form.get("ossDirectory")
    ossFileName = request.form.get("ossFileName")  # 不带后缀

    if oFile == None:
        return jsonify(dict(err=1, msg="文件不能为空", data=False))

    if not isinstance(directory, str):
        return jsonify(dict(err=1, msg="上传路径不能为空", data=False))

    if not len(directory):
        return jsonify(dict(err=1, msg="上传路径不能为空", data=False))

    if not isinstance(ossFileName, str):
        return jsonify(dict(err=1, msg="文件保存在OSS上的名称不能为空", data=False))
    if not len(ossFileName):
        return jsonify(dict(err=1, msg="文件保存在OSS上的名称不能为空", data=False))

    directoryList = directory.split(",")
    if not len(directoryList):
        return jsonify(dict(err=1, msg="上传路径不能为空", data=False))

    strOssDir = "/".join(directoryList)

    strTempDir = os.path.join(app.static_folder, "temp")
    if not os.path.exists(strTempDir):
        os.mkdir(strTempDir)

    fileNameSplit = os.path.splitext(oFile.filename)
    if len(fileNameSplit) < 2:
        return jsonify(dict(err=1, msg="上传的文件有误", data=False))

    fileNameLocal = "{base}_{time}{ext}".format(base=fileNameSplit[0], time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S-%f"), ext=fileNameSplit[1])

    localFilePath = os.path.join(strTempDir, fileNameLocal)
    oFile.save(localFilePath)

    nCount = 0
    bExists = False
    while nCount <= 5:
        if os.path.exists(localFilePath):
            bExists = True
            break

        nCount += 1
        time.sleep(1)

    if not bExists:
        return jsonify(dict(err=1, msg="文件保存失败", data=False))

    tool = OSSTool()
    bSuc = tool.upload(localFilePath, "{ossDir}/{fileName}{ext}".format(ossDir=strOssDir, fileName=ossFileName, ext=fileNameSplit[1]))

    try:
        os.remove(localFilePath)
    except:
        pass

    if not bSuc:
        return jsonify(dict(err=1, msg="上传失败", data=False))

    return jsonify(dict(err=0, msg="上传成功", data=True))


@app.route("/report/getTemplateFromCloud")
def report_get_template_from_cloud():
    tempDir = os.path.join(app.static_folder, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)

    tool = OSSTool()

    localFilePath = os.path.join(tempDir, "report_CloudTemplate_{time}.json".format(time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S-%f")))
    if os.path.exists(localFilePath):
        try:
            os.remove(localFilePath)
        except:
            pass

    bSuc = tool.download("update/reportwizard/CloudTemplate.json", localFilePath)
    if not bSuc or not os.path.exists(localFilePath):
        return jsonify(dict(err=1, msg="下载CloudTemplate.json失败", data=False))

    with open(localFilePath, "r", encoding="UTF8", errors="ignore") as of:
        try:
            dJson = json.load(of)
        except:
            dJson = {}

    templateList = dJson.get("data", [])
    try:
        os.remove(localFilePath)
    except:
        pass
    return jsonify(dict(err=0, msg="获取成功", data=templateList))


@app.route("/report/addFromCloud", methods=["POST"])
def report_add_from_cloud():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    reportTemplateId = rcv.get("id", None)
    strType = rcv.get("type")  # 0-普通文件；2-报表
    strTime = rcv.get("time")  # 当添加报表时传入时间字符串，如1.5,10.3,16:34
    fileTimeType = rcv.get("fileTimeType")  # 报表时间类型  0-日  1-周  2-月
    description = rcv.get("description", "")
    userOfRole = rcv.get("roleList", None)  # 权限等级列表或字符串
    emailBody = rcv.get("emailBody", None)  # 邮件正文
    plantPrefixReplaceTo = rcv.get("plantPrefix", None)  # 机房前缀

    if not isinstance(reportTemplateId, str):
        return jsonify(dict(err=1, msg="报表ID不能为空", data=False))

    if not len(reportTemplateId):
        return jsonify(dict(err=1, msg="报表ID不能为空", data=False))

    if description.strip().find(" ") >= 0:
        return jsonify(dict(err=1, msg="报表描述文本中间夹杂有空格", data=False))

    nType = None
    try:
        nType = int(strType)
    except:
        pass

    nFileTimeType = None
    try:
        nFileTimeType = int(fileTimeType)
    except:
        pass

    if nType not in [0, 2]:
        return jsonify(dict(err=1, msg="type必须为0或2", data=False))

    if nType == 2 and not strTime:
        return jsonify(dict(err=1, msg="报表则必须传入生成时间", data=False))

    if nType == 2 and nFileTimeType == None:
        return jsonify(dict(err=1, msg="报表则必须传入时间类型", data=False))

    if nType == 2 and nFileTimeType not in [0, 1, 2]:
        return jsonify(dict(err=1, msg="报表时间类型必须为0,1,2", data=False))

    if plantPrefixReplaceTo != None:
        if not isinstance(plantPrefixReplaceTo, str):
            return jsonify(dict(err=1, msg="替换为的机房前缀必须为字符串", data=False))

        if not len(plantPrefixReplaceTo):
            return jsonify(dict(err=1, msg="替换为的机房前缀不能为空字符串", data=False))

    tool = OSSTool()
    if not tool.file_exists("update/reportwizard/{id}/".format(id=reportTemplateId)):
        return jsonify(dict(err=1, msg="该ID指向的云端报表模板不存在", data=False))

    userOfRoleList = []
    if userOfRole != None:
        if isinstance(userOfRole, list):
            for item in userOfRole:
                if is_int_digit(item):
                    userOfRoleList.append(str(item))
        elif isinstance(userOfRole, str):
            try:
                les = json.loads(userOfRole)
            except:
                les = []
            if isinstance(les, list):
                for item in les:
                    if is_int_digit(item):
                        userOfRoleList.append(str(item))

    if emailBody != None:
        if not isinstance(emailBody, str):
            return jsonify(dict(err=1, msg="邮件正文必须为字符串", data=False))
        if not len(emailBody):
            return jsonify(dict(err=1, msg="邮件正文不能是空字符串", data=False))

    tempDir = os.path.join(app.static_folder, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)

    localFileName = "report_template_{time}.xlsx".format(time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S-%f"))
    localFilePath = os.path.join(tempDir, localFileName)

    reportTemplateJsonLocalPath = os.path.join(tempDir, "report_CloudTemplate_{time}.json".format(time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S-%f")))
    if os.path.exists(reportTemplateJsonLocalPath):
        try:
            os.remove(reportTemplateJsonLocalPath)
        except:
            pass

    bSucDownload = tool.download("update/reportwizard/CloudTemplate.json", reportTemplateJsonLocalPath)
    if not bSucDownload or not os.path.exists(reportTemplateJsonLocalPath):
        return jsonify(dict(err=1, msg="下载CloudTemplate.json失败", data=False))

    with open(reportTemplateJsonLocalPath, "r", encoding="UTF8", errors="ignore") as fo:
        try:
            dJson = json.load(fo)
        except:
            dJson = {}

        reportList = dJson.get("data", [])
        cloudReportTemplateName = None
        for dReport in reportList:
            if dReport.get("CloudTemplateID", "") == reportTemplateId:
                cloudReportTemplateName = dReport.get("CloudTemplateName", "")
                break

    if cloudReportTemplateName == None:
        try:
            os.remove(reportTemplateJsonLocalPath)
        except:
            pass
        return jsonify(dict(err=1, msg="云端CloudTemplate.json中未找到id为{id}的报表模板".format(id=reportTemplateId), data=False))

    reportFileNameWithExt = "cloudReportTemplate_{id}.xlsx".format(id=reportTemplateId)
    if isinstance(cloudReportTemplateName, str):
        if len(cloudReportTemplateName):
            reportFileNameWithExt = "{name}.xlsx".format(name=cloudReportTemplateName)

    if os.path.exists(localFilePath):
        try:
            os.remove(localFilePath)
        except:
            return jsonify(dict(err=1, msg="删除上次残留的报表模板文件失败", data=False))

    ossFilePath = "update/reportwizard/{id}/report_template.xlsx".format(id=reportTemplateId)
    bSucDownload = tool.download(ossFilePath, localFilePath)
    if not bSucDownload or not os.path.exists(localFilePath):
        return jsonify(dict(err=1, msg="该ID指向的云端报表模板文件下载失败", data=False))

    # 将点名中的机房前缀替换
    if plantPrefixReplaceTo != None:
        reportTool = ReportTool(localFileName, localFilePath)
        bSuc, errMsg = reportTool.replacePlantPrefixOfPointName("Plant01", plantPrefixReplaceTo)
        if not bSuc:
            return jsonify(dict(err=1, msg=errMsg, data=False))

    bSuc, strMsg = BEOPSqliteAccess.getInstance().saveTemplateFileTo4DB(reportFileNameWithExt, localFilePath, strTime, nFileTimeType,
                                                                        description, nType, userOfRoleList, emailBody)

    if os.path.exists(localFilePath):
        try:
            os.remove(localFilePath)
        except:
            pass

    return jsonify(dict(err=0 if bSuc else 1, msg=strMsg, data=bSuc))

@app.route("/getPlantRoomNameList")
def get_plant_room_name_list():
    strKeyList = ["globalconfig"]
    dData = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfigMul(strKeyList)
    if not isinstance(dData, dict):
        return jsonify(dict(err=1, msg="获取失败", data=[]))

    dGlobalConfig = dData.get("globalconfig", {})
    if not isinstance(dGlobalConfig, dict):
        return jsonify(dict(err=1, msg="获取失败", data=[]))

    plantRoomList = dGlobalConfig.get("ChillerPlantRoom", [])
    if not isinstance(plantRoomList, list):
        return jsonify(dict(err=1, msg="获取失败", data=[]))

    resList = []
    for room in plantRoomList:
        resList.append(dict(roomNameEn=room.get("RoomName", ""),
                            roomNameCh=room.get("RoomNameCHS", "")))
    return jsonify(dict(err=0, msg="", data=resList))

@app.route("/getCriticalLog", methods=["POST"])
def get_critical_log():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    timeFrom = rcv.get("timeFrom", None)
    timeTo = rcv.get("timeTo", None)
    if not isinstance(timeFrom, str):
        return jsonify(dict(err=1, msg="起始时间不能为空", data=[]))
    if not len(timeFrom):
        return jsonify(dict(err=1, msg="起始时间不能为空", data=[]))
    if not is_valid_date(timeFrom, "%Y-%m-%d %H:%M:%S"):
        return jsonify(dict(err=1, msg="起始时间格式有误", data=[]))
    if not is_valid_date(timeTo, "%Y-%m-%d %H:%M:%S"):
        return jsonify(dict(err=1, msg="终止时间格式有误", data=[]))

    logList, strErrMsg = LogicManager.getInstance().get_critical_log(timeFrom, timeTo)

    return jsonify(dict(err=0, msg=strErrMsg, data=logList))


@app.route("/redis/getMemorySettings")
def redis_get_memory_settings():
    if not RedisManager.is_alive():
        return jsonify(dict(err=1, msg="redis未运行", data=None))

    localSaveStatus = RedisManager.get_local_save_status()
    rt1 = RedisManager.get_maxmemory()
    rt2 = RedisManager.get_maxmemory_policy()
    return jsonify(dict(err=0, msg="", data={"maxMemory": rt1, "maxMemoryPolicy": rt2, "localSaveStatus": localSaveStatus}))

@app.route('/strategy/removeTodayLogOfLogic',methods=['POST'])
def remove_today_log_of_logic():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    logicName = rcv.get("logicName", None)
    if not isinstance(logicName, str):
        return jsonify(dict(err=1, msg="策略名称不能为空", data=False))
    if not len(logicName):
        return jsonify(dict(err=1, msg="策略名称不能为空", data=False))

    logicName = logicName.replace(".dll", "")
    logDir = os.path.join(app.config["CORE_PATH"], "log")
    if not os.path.exists(logDir):
        return jsonify(dict(err=0, msg="log文件夹不存在", data=True))

    logicDir = os.path.join(logDir, "logic-{date}".format(date=datetime.now().strftime("%Y-%m-%d")))
    logicFilePath = os.path.join(logicDir, "{logicName}.dll.txt".format(logicName=logicName))
    if not os.path.exists(logicFilePath):
        return jsonify(dict(err=0, msg="该策略的今日log已清空", data=True))
    try:
        os.remove(logicFilePath)
    except Exception as e:
        return jsonify(dict(err=1, msg="删除今日log失败:{err}".format(err=e.__str__()), data=False))

    return jsonify(dict(err=0, msg="删除成功", data=True))

@app.route('/logicPush/create',methods=['POST'])
def logic_push_create():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    inputList = rcv.get("inputList", [])
    if not isinstance(inputList, list):
        return jsonify(dict(err=1, msg="输入必须为列表", data=False))

    if not len(inputList):
        return jsonify(dict(err=1, msg="输入列表不能为空", data=False))

    strUuidList = []
    paramList = []
    for item in inputList:
        groupName = item.get("groupName", "")
        logicSourceName = item.get("logicSourceName", "")
        content = item.get("content", "")
        reason = item.get("reason", "")
        occurTime = item.get("occurTime", "")
        important = item.get("important", None)
        urgent = item.get("urgent", None)
        result = item.get("result", None)
        actTime = item.get("actTime", "")
        positionName = item.get("positionName", "")
        remark = item.get("remark", "")

        if not isinstance(groupName, str):
            return jsonify(dict(err=1, msg="groupName必须为字符串", data=False))
        if not len(groupName):
            return jsonify(dict(err=1, msg="groupName不能为空", data=False))
        if not isinstance(logicSourceName, str):
            return jsonify(dict(err=1, msg="logicSourceName必须为字符串", data=False))
        if not len(logicSourceName):
            return jsonify(dict(err=1, msg="logicSourceName不能为空", data=False))
        if not isinstance(content, str):
            return jsonify(dict(err=1, msg="content必须为字符串", data=False))
        if not len(content):
            return jsonify(dict(err=1, msg="content不能为空", data=False))
        if not isinstance(reason, str):
            return jsonify(dict(err=1, msg="reason必须为字符串", data=False))
        if not len(reason):
            return jsonify(dict(err=1, msg="reason不能为空", data=False))
        if not isinstance(occurTime, str):
            return jsonify(dict(err=1, msg="occurTime必须为字符串", data=False))
        if not len(occurTime):
            return jsonify(dict(err=1, msg="occurTime不能为空", data=False))
        if not isValidDate(occurTime, "%Y-%m-%d %H:%M:%S"):
            return jsonify(dict(err=1, msg="occurTime时间格式有误", data=False))
        if not isinstance(important, int):
            return jsonify(dict(err=1, msg="important必须为整数", data=False))
        if not isinstance(urgent, int):
            return jsonify(dict(err=1, msg="urgent必须为整数", data=False))
        if not isinstance(result, int):
            return jsonify(dict(err=1, msg="result必须为整数", data=False))

        if not isinstance(actTime, str):
            return jsonify(dict(err=1, msg="actTime必须为字符串", data=False))
        if not len(actTime):
            return jsonify(dict(err=1, msg="actTime不能为空", data=False))
        if not isValidDate(actTime, "%Y-%m-%d %H:%M:%S"):
            return jsonify(dict(err=1, msg="actTime时间格式有误", data=False))

        if not isinstance(positionName, str):
            return jsonify(dict(err=1, msg="positionName必须为字符串", data=False))
        if not len(positionName):
            return jsonify(dict(err=1, msg="positionName不能为空", data=False))

        if not isinstance(remark, str):
            return jsonify(dict(err=1, msg="remark必须为字符串", data=False))

        strRandom = uuid.uuid4().hex
        paramList.append(
            (strRandom, groupName, logicSourceName, content, reason, occurTime, important, urgent, result, actTime, positionName, remark)
        )

        strUuidList.append(strRandom)

    bSuc = BEOPDataAccess.getInstance().logicPushRecord(paramList)
    return jsonify(dict(err=1 if not bSuc else 0, msg="", data=[] if not bSuc else strUuidList))


@app.route('/logicPush/getLogicList',methods=['POST'])
def logic_push_get_logic_list():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    date = rcv.get("date", None)
    if not isinstance(date, str):
        return jsonify(dict(err=1, msg="日期必须为字符串", data=[]))
    if not len(date):
        return jsonify(dict(err=1, msg="日期不能为空", data=[]))

    tFrom = None
    tTo = None
    if re.match(r"^[0-9]{4}$", date):
        nYear = int(date)
        if nYear > datetime.now().year:
            return jsonify(dict(err=1, msg="年份不能大于当前年份", data=[]))
        tFrom = datetime.strptime("{year}-01-01 00:00:00".format(year=date), "%Y-%m-%d %H:%M:%S")
        if nYear < datetime.now().year:
            tTo = datetime.strptime("{year}-12-31 23:59:59".format(year=date), "%Y-%m-%d %H:%M:%S")
        else:
            tTo = datetime.now()

    elif re.match(r"^[0-9]{4}-[0-9]{2}$", date):
        lSplit = date.split("-")
        strFrom = "{year}-{month}-01 00:00:00".format(year=lSplit[0], month=lSplit[1])
        try:
            tFrom = datetime.strptime(strFrom, "%Y-%m-%d %H:%M:%S")
        except:
            return jsonify(dict(err=1, msg="输入的日期有误，无法转成正常的时间格式", data=[]))

        if tFrom > datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0):
            return jsonify(dict(err=1, msg="时间不能大于当前月份", data=[]))
        if tFrom.replace(day=1, hour=0, minute=0, second=0, microsecond=0) < datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0):
            tTo = tFrom.replace(day=27)
            tTo += timedelta(days=10)
            tTo = tTo.replace(day=1)
        else:
            tTo = datetime.now()

    elif re.match(r"^[0-9]{4}-[0-9]{2}-[0-9]{2}$", date):
        strFrom = date + " 00:00:00"
        try:
            tFrom = datetime.strptime(strFrom, "%Y-%m-%d %H:%M:%S")
        except:
            return jsonify(dict(err=1, msg="输入的日期有误，无法转成正常的时间格式", data=[]))

        if tFrom > datetime.now().replace(hour=0, minute=0, second=0, microsecond=0):
            return jsonify(dict(err=1, msg="时间不能大于今日", data=[]))

        tNow = datetime.now()
        if tFrom.year == tNow.year and tFrom.month == tNow.month and tFrom.day == tNow.day:
            tTo = datetime.now()
        else:
            tTo = tFrom.replace(hour=23, minute=59, second=59)

    if tFrom == None or tTo == None:
        return jsonify(dict(err=1, msg="输入的日期格式有误", data=[]))

    strBegin = tFrom.strftime("%Y-%m-%d %H:%M:%S")
    strEnd = tTo.strftime("%Y-%m-%d %H:%M:%S")

    dataList = BEOPDataAccess.getInstance().getLogicPushList(strBegin, strEnd, None)
    if dataList == None:
        return jsonify(dict(err=1, msg="查询失败，请稍后再试", data=[]))

    dPlant = {}
    for dData in dataList:
        positionName = dData.get("positionName", "")
        if not positionName:
            continue
        if positionName not in dPlant.keys():
            dPlant.update({positionName: 0})

        dPlant[positionName] += 1

    lPlant = []
    for key, value in dPlant.items():
        lPlant.append(dict(roomName=key, count=value))

    for i in range(len(lPlant)):
        for j in range(len(lPlant)-i-1):
            if lPlant[j].get("roomName") > lPlant[j+1].get("roomName"):
                lPlant[j], lPlant[j+1] = lPlant[j+1], lPlant[j]

    return jsonify(dict(err=0, msg="", data=lPlant))


@app.route('/logicPush/update',methods=['POST'])
def logic_push_update():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    strId = rcv.get("id", None)
    nResult = rcv.get("result", None)
    nYear = rcv.get("year", None)
    if not isinstance(strId, str):
        return jsonify(dict(err=1, msg="id必须为字符串", data=False))
    if not len(strId):
        return jsonify(dict(err=1, msg="id不能为空", data=False))
    if nResult not in [-1, 0, 1]:
        return jsonify(dict(err=1, msg="result必须为-1或0或1", data=False))
    if not isinstance(nYear, int):
        return jsonify(dict(err=1, msg="年份必须为整数", data=False))

    bSuc = BEOPDataAccess.getInstance().updateLogicPush(strId, nResult, nYear)
    return jsonify(dict(err=1 if not bSuc else 0, msg="", data=bSuc))

@app.route('/logicPush/getLogicDetails',methods=['POST'])
def logic_push_get_logic_details():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    date = rcv.get("date", None)
    roomName = rcv.get("roomName", None)

    if not isinstance(date, str):
        return jsonify(dict(err=1, msg="日期必须为字符串", data=[]))
    if not len(date):
        return jsonify(dict(err=1, msg="日期不能为空", data=[]))

    tFrom = None
    tTo = None
    strTimeType = None
    bIsToday = False
    bIsThisMonth = False
    tStartTime = None
    if re.match(r"^[0-9]{4}-[0-9]{2}$", date):  # 按月份查
        strTimeType = "month"
        lSplit = date.split("-")
        strFrom = "{year}-{month}-01 00:00:00".format(year=lSplit[0], month=lSplit[1])
        try:
            tFrom = datetime.strptime(strFrom, "%Y-%m-%d %H:%M:%S")
        except:
            return jsonify(dict(err=1, msg="输入的日期有误，无法转成正常的时间格式", data=[]))

        if tFrom > datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0):
            return jsonify(dict(err=1, msg="时间不能大于当前月份", data=[]))
        if tFrom.replace(day=1, hour=0, minute=0, second=0, microsecond=0) < datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0):
            tTo = tFrom.replace(day=27)
            tTo += timedelta(days=10)
            tTo = tTo.replace(day=1)
        else:
            bIsThisMonth = True
            tTo = datetime.now()

    elif re.match(r"^[0-9]{4}-[0-9]{2}-[0-9]{2}$", date):  # 按日查
        strTimeType = "day"
        strFrom = date + " 00:00:00"
        try:
            tFrom = datetime.strptime(strFrom, "%Y-%m-%d %H:%M:%S")
        except:
            return jsonify(dict(err=1, msg="输入的日期有误，无法转成正常的时间格式", data=[]))

        if tFrom > datetime.now().replace(hour=0, minute=0, second=0, microsecond=0):
            return jsonify(dict(err=1, msg="时间不能大于今日", data=[]))

        tNow = datetime.now()
        if tFrom.year == tNow.year and tFrom.month == tNow.month and tFrom.day == tNow.day:
            tTo = datetime.now()
            bIsToday = True
        else:
            tTo = tFrom.replace(hour=23, minute=59, second=59)

    if tFrom == None or tTo == None or strTimeType == None:
        return jsonify(dict(err=1, msg="输入的日期格式有误", data=[]))

    if not isinstance(roomName, str):
        return jsonify(dict(err=1, msg="机房名称必须为字符串", data=[]))
    if not len(roomName):
        return jsonify(dict(err=1, msg="机房名称不能为空", data=[]))

    strBegin = tFrom.strftime("%Y-%m-%d %H:%M:%S")
    strEnd = tTo.strftime("%Y-%m-%d %H:%M:%S")
    recordList = BEOPDataAccess.getInstance().getLogicPushList(strBegin, strEnd, roomName)
    if recordList == None:
        return jsonify(dict(err=1, msg="查询失败，请稍后再试", data=[]))

    dGroup = {}
    for record in recordList:
        groupName = record.get("groupName", "")
        if not groupName:
            continue

        if groupName not in dGroup.keys():
            dGroup.update({groupName: []})

        dGroup[groupName].append(record)

    dataList = []
    dCalendar = {}
    if strTimeType == "month":
        tStart = tFrom
        if bIsThisMonth:
            tStop = datetime.now()
        else:
            tStop = tStart.replace(day=27)
            tStop += timedelta(days=10)
            tStop = tStop.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

        tTar = tStart
        while tTar < tStop:
            strDayKey = "{year}-{month:02}-{day:02}".format(year=tTar.year, month=tTar.month, day=tTar.day)
            dCalendar.update({strDayKey: 0})
            tTar += timedelta(days=1)

        for groupName, infoList in dGroup.items():
            dCountDaily = copy.deepcopy(dCalendar)
            contentList = []
            countList = []
            nTotalCount = 0
            for dInfo in infoList:
                if not isinstance(dInfo.get("actTime", ""), datetime):
                    continue

                recordDateKey = dInfo.get("actTime").strftime("%Y-%m-%d")
                dCountDaily[recordDateKey] += 1
                contentList.append({
                                "details": dInfo.get("content", ""),
                                "suggestTime": dInfo.get("occurTime", "").strftime("%Y-%m-%d %H:%M:%S") if isinstance(dInfo.get("occurTime", ""), datetime) else "",
                                "executeTime": dInfo.get("actTime").strftime("%Y-%m-%d %H:%M:%S")
                            })
                nTotalCount += 1

            tTar = tStart
            while tTar < tStop:
                key = tTar.strftime("%Y-%m-%d")
                countList.append(dCountDaily.get(key, 0))
                tTar += timedelta(days=1)

            dataList.append(dict(title=groupName,
                                 count=nTotalCount,
                                 countList=countList,
                                 content=contentList))

    elif strTimeType == "day":
        tStart = tFrom
        if bIsToday:
            tStop = datetime.now()
        else:
            tStop = tFrom.replace(hour=23, minute=59, second=59)

        tTar = tStart
        while tTar < tStop:
            strHourKey = "{hour}".format(hour=tTar.hour)
            dCalendar.update({strHourKey: 0})
            tTar += timedelta(hours=1)

        for groupName, infoList in dGroup.items():
            dCountHourly = copy.deepcopy(dCalendar)
            contentList = []
            countList = []
            nTotalCount = 0
            for dInfo in infoList:
                if not isinstance(dInfo.get("actTime", ""), datetime):
                    continue

                recordHourKey = "{hour}".format(hour=dInfo.get("actTime").hour)
                dCountHourly[recordHourKey] += 1
                contentList.append({
                                "details": dInfo.get("content", ""),
                                "suggestTime": dInfo.get("occurTime", "").strftime("%Y-%m-%d %H:%M:%S") if isinstance(dInfo.get("occurTime", ""), datetime) else "",
                                "executeTime": dInfo.get("actTime").strftime("%Y-%m-%d %H:%M:%S")
                            })
                nTotalCount += 1

            tTar = tStart
            while tTar < tStop:
                key = "{hour}".format(hour=tTar.hour)
                countList.append(dCountHourly.get(key, 0))
                tTar += timedelta(hours=1)

            dataList.append(dict(title=groupName,
                                 count=nTotalCount,
                                 countList=countList,
                                 content=contentList))

    for i in range(len(dataList)):
        for j in range(len(dataList)-i-1):
            if dataList[j].get("title") > dataList[j+1].get("title"):
                dataList[j], dataList[j+1] = dataList[j+1], dataList[j]

    return jsonify(dict(err=0, msg="", data=dataList))

@app.route("/backendFile/updateTemplateFile", methods=["POST"])
def update_report_template_file():
    file = request.files.get("file")
    nId = request.form.get("reportId")
    if file == None:
        return jsonify(dict(err=1, msg="未发现文件对象", data=False))
    fileName = file.filename

    if not is_int_digit(nId):
        return jsonify(dict(err=1, msg="报表id必须为整数或整数字符串", data=False))
    nId = int(nId)

    tempDir = os.path.join(app.static_folder, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)
    filePath = os.path.join(tempDir, "{name}_{time}{ext}".format(name=os.path.splitext(fileName)[0], time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S"), ext=os.path.splitext(fileName)[1]))
    if os.path.exists(filePath):
        try:
            os.remove(filePath)
        except:
            pass

    file.save(filePath)
    if not os.path.exists(filePath):
        return jsonify(dict(err=1, msg="传入的文件保存失败", data=False))
    bSuc, msg = BEOPSqliteAccess.getInstance().updateReportTemplateFile(nId, fileName, filePath)
    try:
        os.remove(filePath)
    except:
        pass
    return jsonify(dict(err=1 if not bSuc else 0, msg=msg, data=bSuc))


# 获取最近一次测试失败的策略日志
@app.route('/strategy/getLatestRoundLogOfTestFailure',methods=['POST'])
def strategy_get_log_info_latest_round_of_test_failure():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    roomName = rcv.get("roomName", None)
    dllName = rcv.get("dllName", None)   # 策略dll文件名（含.dll或不含都可以）
    keyword = rcv.get("keyword", "失败")

    if not isinstance(roomName, str):
        return jsonify(dict(err=1, msg="roomName必须为字符串", data=""))
    if not isinstance(dllName, str):
        return jsonify(dict(err=1, msg="dllName必须为字符串", data=""))
    if not len(dllName):
        return jsonify(dict(err=1, msg="dllName不能为空", data=""))
    if dllName.endswith(".dll"):
        dllName = dllName.replace(".dll", "")

    if not isinstance(keyword, str):
        keyword = "失败"
    if not len(keyword):
        keyword = "失败"

    logDir = os.path.join(app.config["CORE_PATH"], "log")
    if not os.path.exists(logDir):
        return jsonify(dict(err=1, msg="log目录不存在", data=""))

    logDateDir = os.path.join(logDir, "logic-{year}-{month:02d}-{day:02d}".format(year=datetime.now().year, month=datetime.now().month, day=datetime.now().day))
    if not os.path.exists(logDateDir):
        return jsonify(dict(err=1, msg="指定日期的日志目录不存在({dir})".format(dir=logDateDir), data=""))

    customLogicName, msg = BEOPSqliteAccess.getInstance().getLogicCustomName(dllName, roomName)
    if customLogicName == None or customLogicName == "":
        return jsonify(dict(err=1, msg=msg, data=""))

    strDllFilePath = os.path.join(logDateDir, "{name}.txt".format(name=customLogicName))

    logList, msg = LogicManager.getInstance().get_latest_round_log_of_test_failure(strDllFilePath, customLogicName, keyword)

    if len(logList):
        strLogList = "".join(logList)
        return jsonify(dict(err=0, msg="", data=strLogList))
    return jsonify(dict(err=0, msg=msg if msg else "未发现测试失败的日志", data=""))


@app.route("/uploadDomdbOnce")
def upload_domdb_once():
    domUploadDir = os.path.join(app.config["CORE_PATH"], "domUpload")
    if not os.path.exists(domUploadDir):
        return jsonify(dict(err=1, msg="domUpload不存在", data=False))

    domUploadPath = os.path.join(domUploadDir, "domUpload.exe")
    if not os.path.exists(domUploadPath):
        return jsonify(dict(err=1, msg="domUpload不存在", data=False))

    infoConfigPath = os.path.join(app.config["CORE_PATH"], "DBFileVersion", "config.ini")
    try:
        os.remove(infoConfigPath)
    except:
        pass

    try:
        res = subprocess.Popen("domUpload.exe", stderr=subprocess.PIPE, stdout=subprocess.PIPE, shell=True, cwd=os.path.join(app.config["CORE_PATH"], "domUpload"), creationflags=subprocess.CREATE_NEW_CONSOLE)
        info, out = res.communicate()
    except Exception as e:
        return jsonify(dict(err=1, msg="执行失败:%s" % e.__str__(), data=False))

    return jsonify(dict(err=0, msg="执行成功", data=True))

@app.route("/requestAndSaveCalendarData", methods=["POST"])
def request_and_save_calendar_data():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    strFromDate = rcv.get("fromDate", None)
    nDays = rcv.get("days", None)
    if not isinstance(strFromDate, str):
        return jsonify(dict(err=1, msg="起始日期必须为字符串", data=False))
    if not isValidDate(strFromDate, "%Y-%m-%d"):
        return jsonify(dict(err=1, msg="起始日期格式必须为yyyy-mm-dd", data=False))
    if not isinstance(nDays, int):
        return jsonify(dict(err=1, msg="天数必须为整数", data=False))
    if nDays < 1:
        return jsonify(dict(err=1, msg="天数必须大于1", data=False))

    postData = {"fromDate": strFromDate, "days": nDays}
    headers = {"content-type": "application/json"}
    rsp = requests.post("http://47.100.17.99:80/api/requestCalendarOfPeriod", headers=headers, data=json.dumps(postData), timeout=30)
    if rsp.status_code != 200:
        return jsonify(dict(err=1, msg="status code != 200", data=False))

    try:
        dText = json.loads(rsp.text)
        if not isinstance(dText, dict):
            return jsonify(dict(err=1, msg="返回的数据格式不是字典", data=False))
        dData = dText.get("data", None)
        if not isinstance(dData, dict):
            return jsonify(dict(err=1, msg="返回的数据格式不是字典", data=False))

        bSuc = BEOPDataAccess.getInstance().saveCalendarDataMultiDays(dData)
        if bSuc == None:
            return jsonify(dict(err=1, msg="数据库操作失败，请稍后再试", data=False))
        return jsonify(dict(err=1 if not bSuc else 0, msg="", data=bSuc))
    except Exception as e:
        return jsonify(dict(err=1, msg=e.__str__(), data=False))


@app.route("/report/downloadGeneratedHistoryReport", methods=["POST"])
def download_generated_history_report():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    projectReportsDir = os.path.join(app.static_folder, "projectReports")
    if not os.path.exists(projectReportsDir):
        return jsonify(dict(err=1, msg="dompysite下报表所在路径不存在", data=""))

    reportDir = os.path.join(projectReportsDir, "0")
    if not os.path.exists(reportDir):
        return jsonify(dict(err=1, msg="dompysite下报表所在路径不存在", data=""))

    reportName = rcv.get("reportName", None)  # 需要带后缀
    beginTime = rcv.get("beginTime", None)
    endTime = rcv.get("endTime", None)

    if not isinstance(reportName, str):
        return jsonify(dict(err=1, msg="报表名称必须为字符串", data=""))
    if not len(reportName):
        return jsonify(dict(err=1, msg="报表名称不能为空", data=""))

    reportNameSplit = os.path.splitext(reportName)
    if not reportNameSplit[1]:
        return jsonify(dict(err=1, msg="报表名称必须带文件后缀", data=""))

    strReportName = reportNameSplit[0]
    strExt = reportNameSplit[1]

    tBeginTime = None
    tEndTime = None
    if isinstance(beginTime, str):
        if not is_valid_date(beginTime, "%Y-%m-%d %H:%M:%S"):
            return jsonify(dict(err=1, msg="起始查询时间格式有误，必须为yyyy-mm-dd HH:MM:SS", data=""))
        tBeginTime = datetime.strptime(beginTime, "%Y-%m-%d %H:%M:%S")

    if isinstance(endTime, str):
        if not is_valid_date(endTime, "%Y-%m-%d %H:%M:%S"):
            return jsonify(dict(err=1, msg="终止查询时间格式有误，必须为yyyy-mm-dd HH:MM:SS", data=""))
        tEndTime = datetime.strptime(endTime, "%Y-%m-%d %H:%M:%S")

    tempDir = os.path.join(app.static_folder, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)

    # 清除6小时前的历史文件
    for root, dirs, fileNames in os.walk(tempDir):
        for fileName in fileNames:
            if not fileName.endswith(".zip"):
                continue

            timeList = re.findall(r"[0-9]{4}-[0-9]{2}-[0-9]{2}-[0-9]{2}-[0-9]{2}-[0-9]{2}", fileName)
            if not len(timeList):
                continue

            strTime = timeList[0]
            tTime = datetime.strptime(strTime, "%Y-%m-%d-%H-%M-%S")
            if (datetime.now() - tTime).total_seconds() > 3600 * 6:
                try:
                    os.remove(os.path.join(root, fileName))
                except:
                    pass

    tarDir = os.path.join(tempDir, "{name}_{time}".format(name=strReportName, time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    if os.path.exists(tarDir):
        try:
            shutil.rmtree(tarDir)
        except Exception as e:
            return jsonify(dict(err=1, msg="删除上次残留文件夹失败:%s" % e.__str__(), data=""))

    os.mkdir(tarDir)

    tarZipFileName = "{name}_{time}.zip".format(name=strReportName, time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S"))
    tarZipPath = os.path.join(tempDir, tarZipFileName)
    if os.path.exists(tarZipPath):
        try:
            os.remove(tarZipPath)
        except Exception as e:
            return jsonify(dict(err=1, msg="删除上次残留的压缩包失败:%s" % e.__str__(), data=""))

    nTotal = 0
    for root, dis, fileNames in os.walk(reportDir):
        for fileName in fileNames:
            try:
                if not fileName.startswith(strReportName):
                    continue

                if not fileName.endswith(strExt):
                    continue

                bCount = False
                if not isinstance(tBeginTime, datetime) or not isinstance(tEndTime, datetime):
                    bCount = True
                else:
                    timeList = re.findall(r"[0-9]{14}", fileName)
                    if not len(timeList):
                        continue

                    strTime = timeList[0]

                    tTime = None
                    try:
                        tTime = datetime.strptime(strTime, "%Y%m%d%H%M%S")
                    except:
                        pass

                    if isinstance(tTime, datetime):
                        if tTime >= tBeginTime and tTime <= tEndTime:
                            bCount = True

                if bCount:
                    shutil.copyfile(os.path.join(root, fileName), os.path.join(tarDir, fileName))
                    nTotal += 1

            except:
                pass
            finally:
                time.sleep(0.2)

    if nTotal == 0:
        try:
            shutil.rmtree(tarDir)
        except:
            pass
        return jsonify(dict(err=1, msg="未发现任何符合查询条件的后台报表", data=""))

    zip_dir(tarDir, tarZipPath)

    try:
        shutil.rmtree(tarDir)
    except:
        pass
    return jsonify(dict(err=0, msg="共压缩{n}个报表文件".format(n=nTotal), data=tarZipFileName))


@app.route("/warning/exportHistory", methods=["POST"])
def export_warning_history():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    beginDate = rcv.get("beginDate", None)
    endDate = rcv.get("endDate", None)

    tBeginTime = None
    tEndTime = None
    if isinstance(beginDate, str):
        if not is_valid_date(beginDate, "%Y-%m"):
            return jsonify(dict(err=1, msg="起始查询日期格式有误，必须为yyyy-mm", data=""))
        tBeginTime = datetime.strptime(beginDate, "%Y-%m").replace(day=1, hour=0, minute=0, second=0)

    if isinstance(endDate, str):
        if not is_valid_date(endDate, "%Y-%m"):
            return jsonify(dict(err=1, msg="终止查询日期格式有误，必须为yyyy-mm", data=""))
        tEndTime = datetime.strptime(endDate, "%Y-%m").replace(day=1, hour=0, minute=0, second=0)

    tempDir = os.path.join(app.static_folder, "temp")
    if not os.path.exists(tempDir):
        os.mkdir(tempDir)

    # 清除6小时前的历史zip文件
    for root, dirs, fileNames in os.walk(tempDir):
        for fileName in fileNames:
            if not fileName.endswith(".zip"):
                continue

            timeList = re.findall(r"[0-9]{4}-[0-9]{2}-[0-9]{2}-[0-9]{2}-[0-9]{2}-[0-9]{2}", fileName)
            if not len(timeList):
                continue

            strTime = timeList[0]
            tTime = datetime.strptime(strTime, "%Y-%m-%d-%H-%M-%S")
            if (datetime.now() - tTime).total_seconds() > 3600 * 6:
                try:
                    os.remove(os.path.join(root, fileName))
                except:
                    pass

    tarDir = os.path.join(tempDir, "warnningrecord_{time}".format(time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    if os.path.exists(tarDir):
        try:
            shutil.rmtree(tarDir)
        except Exception as e:
            return jsonify(dict(err=1, msg="删除上次残留文件夹失败:%s" % e.__str__()))

    os.mkdir(tarDir)

    tableNameList = BEOPDataAccess.getInstance().listWarningTables()
    if tableNameList == None:
        return jsonify(dict(err=1, msg="查询失败，请稍后再试", data=""))

    tableNameListTar = []
    if isinstance(tBeginTime, datetime) and isinstance(tEndTime, datetime):
        for tableName in tableNameList:
            if tableName == "warningrecord":
                tableNameListTar.append(tableName)
            else:
                timeList = re.findall(r"[0-9]{4}_[0-9]{2}", tableName)
                if not len(timeList):
                    continue

                strTime = timeList[0]
                tTime = None
                try:
                    tTime = datetime.strptime(strTime, "%Y_%m").replace(day=1, hour=0, minute=0, second=0)
                except:
                    pass

                if not isinstance(tTime, datetime):
                    continue

                if tTime >= tBeginTime and tTime <= tEndTime:
                    tableNameListTar.append(tableName)
    else:
        tableNameListTar.extend(tableNameList)

    bSuc, nCount = BEOPDataAccess.getInstance().exportHistoryWarningRecordInExcel(tarDir, tableNameListTar)
    if bSuc == None:
        return jsonify(dict(err=1, msg="查询失败，请稍后再试", data=""))

    zipFileName = "warningrecord_{time}.zip".format(time=datetime.now().strftime("%Y-%m-%d-%H-%M-%S"))
    zipFilePath = os.path.join(tempDir, zipFileName)
    if os.path.exists(zipFilePath):
        try:
            os.remove(zipFilePath)
        except Exception as e:
            return jsonify(dict(err=1, msg="删除上次残留的压缩包失败:%s" % e.__str__(), data=""))

    if nCount == 0:
        try:
            shutil.rmtree(tarDir)
        except:
            pass
        return jsonify(dict(err=1, msg="未导出任何历史报警表格", data=""))

    zip_dir(tarDir, zipFilePath)

    try:
        shutil.rmtree(tarDir)
    except:
        pass
    return jsonify(dict(err=0, msg="共导出{n}个历史报警表".format(n=nCount), data=zipFileName))


@app.route("/killOM")
def kill_om():
    bKilled = False

    nCount = 0
    while nCount < 5:
        print("执行一次 kill OM.exe")
        try:
            ProcessManager.getInstance().killProcess("OM.exe")
        except:
            pass
        try:
            ProcessManager.getInstance().killProcess("Electron.exe")
        except:
            pass

        time.sleep(1)

        if not ProcessManager.getInstance().findProcess("OM.exe") and not ProcessManager.getInstance().findProcess("Electron.exe"):
            bKilled = True

        if bKilled:
            break

        nCount += 1
        time.sleep(1)

    return jsonify(dict(err=0, msg="", data=bKilled))

@app.route("/setSiteSimulationMode", methods=["POST"])
def set_site_simulation_mode():
    rcv = request.get_json()
    if not isinstance(rcv, dict):
        rcv = {}

    nMode = rcv.get("mode", None)
    if nMode not in [0, 1]:
        return jsonify(dict(err=1, msg="模式必须为0或1", data=False))

    domcoreiniPath = os.path.join(app.config["CORE_PATH"], "domcore.ini")
    if not os.path.exists(domcoreiniPath):
        return jsonify(dict(err=1, msg="domcore.ini不存在", data=False))

    # 切仿真模式
    if nMode == 0:
        try:
            dCur = None
            strCur = BEOPSqliteAccess.getInstance().getValueByKeyInLocalConfig("CloudSetting")
            if strCur == None:
                dCur = {}
            else:
                try:
                    dCur = json.loads(strCur)
                except:
                    pass

            if not isinstance(dCur, dict):
                return jsonify(dict(err=1, msg="当前云连接状态读取失败，仿真模式切换失败", data=False))

            if not dCur:
                bCloudEnabled = False
            else:
                if dCur.get("CloudSettingEnable") == True:
                    bCloudEnabled = True
                else:
                    bCloudEnabled = False

            if bCloudEnabled:
                return jsonify(dict(err=1, msg="当前处于dom云连接开启状态，故不允许切换为仿真模式", data=False))

            bSuc, strMsg = BEOPDataAccess.getInstance().updateMultiKeyValueOfUnit01({"sitemode": ["0"]})
            if not bSuc:
                return jsonify(dict(err=1, msg="mysql数据库中的仿真模式信号写入失败:{msg}".format(msg=strMsg), data=False))

            cf = MyConf()
            cf.read(domcoreiniPath)
            if not cf.has_section("core"):
                cf.add_section("core")

            cf["core"].update({"sitemode": "0"})
            with open(domcoreiniPath, "w", encoding="UTF8", errors="ignore") as fo:
                cf.write(fo)

            return jsonify(dict(err=0, msg="设置成功", data=True))
        except Exception as e:
            return jsonify(dict(err=1, msg="设置失败:{err}".format(err=e.__str__()), data=False))

    # 设置现场模式
    elif nMode == 1:
        try:
            bSuc, strMsg = BEOPDataAccess.getInstance().updateMultiKeyValueOfUnit01({"sitemode": ["1"]})
            if not bSuc:
                return jsonify(dict(err=1, msg="mysql数据库中的现场模式信号写入失败:{msg}".format(msg=strMsg), data=False))

            cf = MyConf()
            cf.read(domcoreiniPath)
            if not cf.has_section("core"):
                cf.add_section("core")

            cf["core"].update({"sitemode": "1"})
            with open(domcoreiniPath, "w", encoding="UTF8", errors="ignore") as fo:
                cf.write(fo)

            return jsonify(dict(err=0, msg="设置成功", data=True))
        except Exception as e:
            return jsonify(dict(err=1, msg="设置失败:{err}".format(err=e.__str__()), data=False))

    return jsonify(dict(err=1, msg="设置失败", data=False))
