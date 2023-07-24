# -*- coding: utf-8 -*-

"""
Routes and views for the flask application.
"""
import copy

from siteinterface import app
from flask import g, json
from math import floor, ceil
import os, sys
import mysql.connector
from math import floor, ceil
from datetime import datetime,timedelta
import time
import logging
import zipfile
from docx.shared import Inches
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from collections import OrderedDict
from siteinterface.BEOPMySqlDBContainer import *
import mysql.connector.locales.eng.client_error
from siteinterface.commonUtils import *
import datetime as datetime_
import traceback
from siteinterface.sqlite_manager import SqliteManager
from siteinterface.RedisManager import RedisManager
import xlrd
from siteinterface.commonUtils import get_number_format, standardAssetTmplDef, isValidDate
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.styles.borders import Border, Side


def log_info_to_file(strLogFileName, strLogInfo):
    strCorePath = os.path.dirname(os.getcwd())
    strLogPath = os.path.join(strCorePath, 'dompysite')
    strLogPath = os.path.join(strLogPath, 'log')
    strLogFilePath = os.path.join(strLogPath, strLogFileName)

    tNow = datetime.now()
    with open(strLogFilePath, 'a+') as f:
        f.write(tNow.strftime('%Y-%m-%d %H:%M:%S')+ '    ' + strLogInfo + '\n')

def get_history_data_padded_log_info(strFileName, strLog):
    strLogDir = os.path.join(os.getcwd(), "log")
    if not os.path.exists(strLogDir):
        os.mkdir(strLogDir)

    hisDataLogDir = os.path.join(strLogDir, "get_history_data_padded")
    if not os.path.exists(hisDataLogDir):
        os.mkdir(hisDataLogDir)

    logFilePath = os.path.join(hisDataLogDir, strFileName)

    tNow = datetime.now()
    with open(logFilePath, 'a+') as f:
        f.write(tNow.strftime('%Y-%m-%d %H:%M:%S') + '    ' + strLog + '\n')

class BEOPDataAccess:
    __instance = None


    def __init__(self):
        self._data = dict()
        self._pointList = dict()
        self._db = None
        self._cur = None
        self._mysqlDbContainer = BEOPMySqlDBContainer('POOL_MAIN', 'dompysite')
        self._mysqlDbContainerRealTimePool = BEOPMySqlDBContainer('POOL_REALTIME_TABLE', 'dompysite01')
        self._mysqlDbContainerModbusClientPool = BEOPMySqlDBContainer('POOL_MODBUS_CLIENT', 'dompysite02')
        self._lastUpdatedTime = None

    @classmethod
    def getInstance(self):
        if(self.__instance == None):
            print('Init database access connections.')
            self.__instance = BEOPDataAccess()
        return self.__instance

    def isNeedUpdate(self):
        if self._lastUpdatedTime is None:
            return True
        tNow = datetime.now()
        tdelta = tNow - self._lastUpdatedTime
        if tdelta.total_seconds()>=10:
            return True

        return False

    def updateData(self):
        if not self.isNeedUpdate():
            return True
        self.getPointList()
        dataCached=self.getCachedData()
        dataCurrent=self.getInputTable()[0]
        self._data=self.mergeRealtimeData(dataCurrent,dataCached)
        return True


    def getInputTableFromMySQL(self, pointList=None, strTableName= None):
        if strTableName is None:
            strTableName = app.config['TABLE_INPUT']
        strSQL ='select pointname, pointvalue from %s order by pointname'%(strTableName)
        #print(strSQL)
        rv = self._mysqlDbContainerRealTimePool.op_db_query(app.config['DATABASE'], strSQL)
        #print(rv)
        retReal = []
        retAll = {}

        nRetryCount = 5
        while rv is None and nRetryCount>0:
            time.sleep(1)
            rv = self._mysqlDbContainerRealTimePool.op_db_query(app.config['DATABASE'], strSQL)
            nRetryCount-=1

        if rv is None:
            return None

        strCurTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        bGetAll = False
        if pointList is None:
            bGetAll = True
        if isinstance(pointList, list) and len(pointList)==0:
            bGetAll = True
        if not bGetAll:
            for item in rv:
                if isinstance(item[0],bytes) and isinstance(item[1], bytes):
                    strItem = item[0].decode()
                    strValue  = item[1].decode()
                    if strItem in pointList:
                        retReal.append(dict(name=strItem, value=strValue, time=strCurTime))
                    retAll[strItem] =  strValue
                elif isinstance(item[0],str):
                    if item[0] in pointList:
                        retReal.append(dict(name=item[0], value=item[1], time=strCurTime))
                    retAll[item[0]] = item[1]
        else:
            for item in rv:
                retReal.append(dict(name=item[0], value=item[1], time=strCurTime))
                retAll[item[0]] = item[1]

        return retReal, retAll

    def getInputTable(self, pointList=None, strTableName= None):
        if app.config['REDIS_ALIVE'] and strTableName is None:
            strUpdateTime = RedisManager.get_realtime_data_update_time_string()
            if strUpdateTime is not None:
                try:
                    tUpdateTime = datetime.strptime(strUpdateTime, '%Y-%m-%d %H:%M:%S')
                    tDelta = datetime.now()-tUpdateTime
                    if tDelta.total_seconds()<=60:
                        pointDataMap = RedisManager.get_realtime_data()
                        if pointDataMap is not None and isinstance(pointDataMap, dict):
                            pointDataListReturn = []
                            pointDataDictReturn = {}
                            strCurTime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            for k, v in pointDataMap.items():
                                if pointList and isinstance(pointList, list) and len(pointList) > 0:
                                    if k in pointList:
                                        pointDataListReturn.append(dict(name=k, value=v, time=strCurTime))
                                        pointDataDictReturn.update({k: v})
                                else:
                                    pointDataListReturn.append(dict(name=k, value=v, time=strCurTime))
                                    pointDataDictReturn.update({k: v})

                            if pointDataMap:
                                return pointDataListReturn, pointDataDictReturn
                    else:
                        logging.error('ERROR in REDIS getInput: time delayed too long, realtime_data_update_time:%s, delayed seconds:%d'%(strUpdateTime, int(tDelta.total_seconds())))
                except Exception as ee:
                    logging.error('ERROR in REDIS getInput: Exception:%s' % (ee.__str__()))

        return self.getInputTableFromMySQL(pointList, strTableName)

    def getInputTableAsListWithTime(self, pointList=None):
        retReal = []
        strSQL ='select time,pointname, pointvalue from %s order by pointname'%(app.config['TABLE_INPUT'])
        rv = self._mysqlDbContainerRealTimePool.op_db_query(app.config['DATABASE'], strSQL)
        if rv is None:
            return retReal

        #print(str(pointList))
        if pointList is not None:
            for item in rv:
                if item[0]=='ChOnOff01':
                    a=1
                if (item[0] in pointList):
                    #print('%s : %s'%(str(item[0]), str(item[1])))
                    retReal.append([item[0].strftime('%Y-%m-%d %H:%M:%S'),item[1], item[2]])
        else:
            for item in rv:
                retReal.append([item[0].strftime('%Y-%m-%d %H:%M:%S'),item[1], item[2]])
        return retReal

    def getInputTableAsDictListWithTimeV1(self, pointList=None):
        dataList, dataDict = self.getInputTable(pointList)
        if not dataDict or not isinstance(dataDict, dict):
            return {}
        return dataDict

    def getInputTableAsDictListWithTime(self, pointList=None):
        strSQL ='select time,pointname, pointvalue from %s order by pointname'%(app.config['TABLE_INPUT'])
        rv = self._mysqlDbContainerRealTimePool.op_db_query(app.config['DATABASE'], strSQL)
        retReal = []
        if rv is None:
            return retReal
        #print(str(pointList))
        if pointList is not None:
            for item in rv:
                if item[1]=='ChOnOff01':
                    a=1
                if (item[1] in pointList):
                    #print('%s : %s'%(str(item[0]), str(item[1])))
                    retReal.append(dict(t=item[0].strftime('%Y-%m-%d %H:%M:%S'),pn=item[1], pv=item[2]))
        else:
            for item in rv:
                retReal.append(dict(t=item[0].strftime('%Y-%m-%d %H:%M:%S'),pn=item[1], pv=item[2]))
        return retReal

    # 根据点名关键字获取实时数据
    def getInputTableByKeyWord(self, keyWordList, targetPage, pageNum):
        try:
            result = {'data':None, 'totalNum':-1, 'msg': 'success'}

            offset = (targetPage-1) * pageNum
            strSQL = 'select time, pointname, pointvalue from ' +app.config['TABLE_INPUT'];
            if len(keyWordList) != 0:
                # strSQL += ' where '
                first = True
                for i in range(len(keyWordList)):
                    if keyWordList[i].strip() != '':
                        if first == True:
                            strSQL += ' where pointname like "%'+keyWordList[i].strip()+'%" '
                            first = False
                        else:
                            strSQL += ' and pointname like "%'+keyWordList[i].strip()+'%" '
            strSQL += ' order by pointname'
            rv = self._mysqlDbContainerRealTimePool.op_db_query(app.config['DATABASE'], strSQL)
            if rv is None:
                result['msg'] = 'db query error'
                return result

            total = len(rv)
            # strSQL += ' limit '+str(offset)+','+str(pageNum)

            # rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)
            retReal = []
            for i in range(pageNum):
                if i+offset < total:
                    item = rv[i+offset]
                    retReal.append(dict(time=item[0].strftime('%Y-%m-%d %H:%M:%S'),name=item[1], value=item[2]))
                else:
                    break
            result['data'] = retReal
            result['totalNum'] = total
        except Exception as e:
            print('ERROR: ' + e.__str__())
            result['msg'] = e.__str__()

        return result

    def getSiteMode(self):
        result = {'sitemode':None,'msg':'','status':False}

        try:
            strSQL = 'select unitproperty02 from `%s` where unitproperty01="sitemode" ' %(app.config['TABLE_UNIT01'])
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'],strSQL)
            if rv is None:
                result['msg'] = 'db query return None in getSiteMode'
                return result
            result['status'] = True
            result['sitemode'] = rv[0][0]
        except Exception as e:
            result['msg'] = e.__str__()
        return result

    def setSiteMode(self,sitemode):
        result = { 'data':[],'msg':'','status':False }

        try:
            strSQL = 'update `%s` set unitproperty02="%s" where  unitproperty01="sitemode" ' %(app.config['TABLE_UNIT01'],sitemode)
            rv = self._mysqlDbContainer.op_db_update(app.config['DATABASE'],strSQL)
            result['status'] = rv
        except Exception as e:
            result['msg'] = e.__str__()
        return result

    def getLicenseVersion(self,):
        strSQL = 'select unitproperty02 from `%s` where unitproperty01="licVer" ' % (app.config['TABLE_UNIT01'])
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)
        if rv is None:
            return 0
        try:
            return int(float(rv[0][0]))
        except:
            pass
        return 0

    def saveLicenseVersion(self, nVersion):
        self.saveUnit01('licVer', str(nVersion))

    def getProcessList(self, bGetAll):
        result = { 'data':[],'msg':'','status':False }
        try:
            strSQL = ""
            if not bGetAll:
                strSQL = 'select unitproperty02,unitproperty03,unitproperty04,unitproperty01 from `%s` where unitproperty01="hostprocess" ' %(app.config['TABLE_UNIT01'])
            else:
                strSQL = 'select unitproperty02,unitproperty03,unitproperty04,unitproperty01 from `%s` where unitproperty01="hostprocess" or unitproperty01="taskprocess"' % (
                app.config['TABLE_UNIT01'])
            if not strSQL:
                return { 'data':[],'msg':'','status':False }

            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'],strSQL)
            if rv is None:
                result['msg'] = 'db query error in getProcessList'
                result['status'] = False
                return result
            data=[]
            for item in rv:
                data.append(dict(directoryName=item[0],processName=item[1],processStatus=item[2],processType=item[3]))
            result = { 'data':data,'msg':'','status':True }
        except Exception as e:
            print('ERROR in getProcessList:'+ e.__str__())
            result['msg'] = e.__str__()
            result['status'] = False
        return result

    def commonGetFields(self,unitproperty01):
        strSQL = 'select unitproperty02 from `%s` where unitproperty01="%s" ' %(app.config['TABLE_UNIT01'],unitproperty01)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)
        if rv is None:
            return None
        return rv[0][0]

    def getModbusSetting(self):
        fieldList = [
            "modbusreadonebyone",
            "modbusidinterval",
            "modbusmutilcount" ,
            "modbustimeout",
            "modbuspollinterval",
            "co3pcmdinterval",
            "co3ptimeoutl",
            "co3prollinterval"]

        result = { 'data':{},'msg':'','status':False }
        try:
            data = {}
            for item in fieldList:
                data[item] = self.commonGetFields(item)

            result = { 'data':data,'msg':'','status':True }
        except Exception as e:
            print('ERROR:' + e.__str__())
            result['msg'] = e.__str__()
            result['status'] = False
        return result

    def getOPCSetting(self):
        fieldList=["opcclientmaxpoint",
                    "opclanguageid",
                    "opcmutilcount",
                    "opcupdaterate",
                    "opccmdsleep",
                    "opcmultiadd",
                    "opcpollsleep",
                    "opcreconnect",
                    "opccmdsleepfromdevice",
                    "opccheckreconnect",
                    "opcpollsleepfromdevice",
                    "OPCServerIP",
                    "opcasync20interval",
                    "opcreconnectignore",
                    "opcmainpollsleep",
                    "opcasync20mode",
                   "enableSecurity",
                   "disableCheckQuality",
                   "opcclientthread"
                   ]


        result = { 'data':{},'msg':'','status':False }
        try:
            data = {}
            for item in fieldList:
                data[item] = self.commonGetFields(item)

            result = { 'data':data,'msg':'','status':True }
        except Exception as e:
            print('ERROR:' + e.__str__())
            result['msg'] = e.__str__()
            result['status'] = False
        return result

    def getDatabaseSetting(self):
        fieldList = [
            "modbusreadonebyone",
            "modbusidinterval",
            "modbusmutilcount"
        ]


        result = { 'data':{},'msg':'','status':False }
        try:
            data = {}
            for item in fieldList:
                data[item] = self.commonGetFields(item)

            result = { 'data':data,'msg':'','status':True }
        except Exception as e:
            print('ERROR:' + e.__str__())
            result['msg'] = e.__str__()
            result['status'] = False
        return result

    def getBacnet(self):
        BACNET_IFACE = 'BACNET_IFACE'

        fieldList = [
            "readinterval",
            "readtypeinterval",
            "readcmdinterval",
            "readbacnetmode",
            "readtimeout"
        ]
        result = { 'data':{},'msg':'','status':False }
        try:
            data = {}
            value = os.environ.get(BACNET_IFACE)
            if value == None:
                data[BACNET_IFACE] = ''
            else:
                data[BACNET_IFACE] = value

            for item in fieldList:
                data[item] = self.commonGetFields(item)

            result = { 'data':data,'msg':'','status':True }

        except Exception as e:
            print('ERROR:' + e.__str__())
            result['msg'] = e.__str__()
            result['status'] = False
        return result

    # 根据一个查询字段获取不同分类的通讯设置
    def getCommunicationSetting(self,str):
        if str == 'modbus':
            return self.getModbusSetting()
        elif str == "OPC":
            return self.getOPCSetting()
        elif str == 'database':
            return self.getDatabaseSetting()
        elif str == 'bacnet':
            return self.getBacnet()

    # 设置通讯设置
    def setCommunicationSetting(self,unitproperty01,unitproperty02):

        result = { 'data':{},'msg':'','status':False }
        try:
            strSQL='UPDATE `%s` SET unitproperty02=%s  WHERE unitproperty01="%s" ' %(app.config['TABLE_UNIT01'],unitproperty02,unitproperty01)
            rv = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strSQL)
            result = { 'data':rv,'msg':'','status':True }
        except Exception as e:
            print('ERROR:'+ e.__str__())
            result['msg'] = e.__str__()
            result['status'] = False
        return result

    #增加守护进程
    def addProcessByName(self,directory,processName,strType, exeTime):
        result = { 'data':None,'msg':'','status':None }
        status = None
        queryList = []
        paramList = []
        try:
            if strType == "hostprocess":
                strDel = "DELETE FROM unit01 WHERE unitproperty02 = %s AND unitproperty03 = %s"
                queryList.append(strDel)
                paramList.append((directory, processName))

                strInsert = 'insert into unit01 (unitproperty01,unitproperty02,unitproperty03,unitproperty04) values ("hostprocess",%s,%s,1)'
                queryList.append(strInsert)
                paramList.append((directory, processName))

            elif strType == "taskprocess":
                strDel = "DELETE FROM unit01 WHERE unitproperty02 = %s AND unitproperty03 = %s AND unitproperty04=%s"
                queryList.append(strDel)
                paramList.append((directory, processName, exeTime))

                strInsert = 'insert into unit01 (unitproperty01,unitproperty02,unitproperty03,unitproperty04) values ("taskprocess",%s,%s,%s)'
                queryList.append(strInsert)
                paramList.append((directory, processName, exeTime))

            if len(queryList) and len(paramList):
                status = self._mysqlDbContainer.op_db_transaction_update_many(app.config['DATABASE'], queryList, paramList)

            result = {'data': None,'msg':'','status': status}

        except Exception as e:
            print('ERROR:'+ e.__str__())
            result['msg'] = e.__str__()

            result['status'] = False
        return result

    #删除守护进程
    def delProcessList(self,processList, strType):
        result = {'data':None,'msg':'','status':None}
        if processList is None or not isinstance(processList, list):
            return result
        try:

            point_join_str = '(\'%s\')' % '\',\''.join(processList)
            strSQL = 'delete from `%s` where unitproperty01="%s" and unitproperty03 in %s ' %(app.config['TABLE_UNIT01'], strType, point_join_str)
            status = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strSQL)
            result= { 'data':None,'msg':'','status':status }
        except Exception as e:
            print('ERROR: ' + e.__str__())
            result['msg'] = e.__str__()

        return result

    #根据点名关键字获取日志信息
    def searchLogByKeyWord(self, keyWordList, timeFrom, timeTo, targetPage, pageSize):
        try:
            result = {'data':None, 'totalNum':-1, 'msg': 'success'}
            delta = timedelta(days=1)
            dtimeFrom = datetime.strptime(timeFrom, '%Y-%m-%d %H:%M:%S') #datetime 开始时间
            ymdFrom = dtimeFrom.strftime('%Y-%m-%d') #字符串表示天 开始

            dtimeTo = datetime.strptime(timeTo, '%Y-%m-%d %H:%M:%S') #datetime 结束时间
            ymdTo = dtimeTo.strftime('%Y-%m-%d') #字符串表示天  结束

            last = ' '
            if len(keyWordList) != 0:
                # strSQL += ' where '
                first = True
                for i in range(len(keyWordList)):
                    if keyWordList[i].strip() != '':
                        if first == True:
                            last += ' where loginfo like "%'+keyWordList[i].strip()+'%" '
                            first = False
                        else:
                            last += ' and loginfo like "%'+keyWordList[i].strip()+'%" '
            last += ' order by time desc'

            to = dtimeTo
            data = []

            while (ymdTo != ymdFrom):
                #更新开始时间

                #sql  查 ymdTo天的数据

                query = 'select * from log_'+ymdTo.replace('-','_')+last
                rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], query)
                if rv is not None:
                    for item in rv:
                        data.append(dict(time=item[0], loginfo=item[1]))
                #更新结束时间
                to -= delta
                ymdTo = to.strftime('%Y-%m-%d')
                # to = datetime.strptime(ymdTo + ' 23:59:59', '%Y-%m-%d %H:%M:%S') #datetime 结束时间
            query = 'select * from log_'+ymdFrom.replace('-','_')+last
            rv = self._mysqlDbContainer.op_db_query('domlog', query)
            if rv:
                for item in rv:
                    data.append(dict(time=item[0], loginfo=item[1]))

            retReal = []
            st = 0
            ed = len(data)-1
            while(st < len(data) and data[st]['time'] > dtimeTo):
                st += 1
            while(ed >= 0 and data[ed]['time'] < dtimeFrom):
                ed -= 1
            total = ed-st

            if total < 0:
                total = 0
            a = st
            while(a<=ed):
                data[a]['time'] = data[a]['time'].strftime('%Y-%m-%d %H:%M:%S')
                a += 1

            if targetPage==-1 and pageSize==-1:
                while(st<=ed):
                    retReal.append(data[st])
                    st += 1
            else:
                offset = (targetPage-1) * pageSize
                st += offset
                i = 0
                while(st <= ed and i<pageSize):
                    retReal.append(data[st])
                    st += 1
                    i += 1

            result['data'] = retReal
            result['totalNum'] = total
        except Exception as e:
            print('ERROR:' + e.__str__())
            result['msg'] = e.__str__()

        return result
    # 没必要。。。
    # def checkTableExists(self, tablename):
    #     query = 'select count(*) from information_schema.tables where table_name="'+tablename+'"'
    #     rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], query)
    #     if rv[0] == 1:
    #         return True
    #     return False

    def updateInputTable(self, pointName , pointValue):
        strInputTableName = app.config['TABLE_INPUT']
        q = 'Update '+ strInputTableName +'set time= now(), pointvalue = "%s" where pointname = "%s"' % (pointValue, pointName)
        return self.op_db_update(q)

    def appendOutputTable(self, pair):
        q='delete from `%s` where pointname = "%s"' %(app.config['TABLE_OUTPUT'],pair[0])
        bSuccess1 = self._mysqlDbContainer.op_db_update(q)
        q='insert into `%s` (pointname,pointvalue) values ("%s","%s")' %(app.config['TABLE_OUTPUT'],pair[0],pair[1])
        bSuccess2 = self._mysqlDbContainer.op_db_update(q)
        return bSuccess1 and bSuccess2

    def isWriteable(self, point):
        q='select * from `%s` where unitproperty02="%s" and unitproperty06="%s"'%(app.config['TABLE_POINT'],point,'W')
        rv = self._mysqlDbContainer.op_db_query(q)
        if rv is None:
            return False
        return len(rv) > 0


    def appendOperationLog(self, userName, strOperation):
        if isinstance(strOperation, str):
            q = 'insert into `%s` (user,OptRemark) values ("%s", "%s")'%(app.config['TABLE_OP'], userName, strOperation)
            return self._mysqlDbContainer.op_db_update(app.config['DATABASE'], q)

        elif isinstance(strOperation, list):
            strQuery = "INSERT INTO " + app.config['TABLE_OP'] + "(user, OptRemark) VALUES (%s, %s)"
            paramList = []
            for item in strOperation:
                paramList.append(
                    (userName, item)
                )
            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(paramList))
            return bSuc


    def readOperationLogInTimeRange(self, strTimeFrom, strTimeTo):

        q = 'select RecordTime,user,OptRemark from operation_record where RecordTime >= "%s" and RecordTime <= "%s" and OptRemark not like "程序%%" and OptRemark not like "%%s3db%%" order by RecordTime' %(strTimeFrom, strTimeTo)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'],q)

        ret = []
        '''
        ('time')
        strUser = op.get('user')
        strContent = op.get('content')
        strReason = op.get('reason')
        '''
        if rv:
            for item in rv:
                ret.append(dict(time=item[0].strftime('%Y-%m-%d %H:%M:%S'),user=item[1], content=item[2], reason=''))

        return ret

    def getUserOfRoleByNameList(self, userNameList, userIdList=[]):
        dNameRole = {"cx": 10, "admin": 3, "guest": 1}
        dIdRole = {-1: 10, 0: 3, 9999: 1}

        strQuery = "SELECT userid, username, userofrole FROM beopuser WHERE username in (" + ",".join(["'{0}'".format(name) for name in userNameList]) + ")"
        if len(userIdList):
            strQuery += " OR userid in (" + ",".join([str(nId) for nId in userIdList]) + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)

        if items == None:
            return dNameRole, dIdRole

        for item in items:
            try:
                userName = item[1].lower()
                dNameRole.update({userName: int(float(item[2]))})
                dIdRole.update({int(float(item[0])): int(float(item[2]))})
            except:
                pass

        return dNameRole, dIdRole

    def readOperationLogV2(self, strStart, strEnd, userId=None, nType=None):
        recordList = []

        userNameList = ["admin", "cx"]
        if nType != None:
            userList = BEOPDataAccess.getInstance().get_all_users()
            for dUser in userList:
                if dUser.get("username", None):
                    if dUser.get("username") not in userNameList:
                        userNameList.append(dUser.get("username"))

        try:
            tFrom = datetime.strptime(strStart, "%Y-%m-%d").replace(hour=0, minute=0, second=0, microsecond=0)
            tTo = datetime.strptime(strEnd, "%Y-%m-%d").replace(hour=23, minute=59, second=59, microsecond=0)

            strFrom = tFrom.strftime("%Y-%m-%d %H:%M:%S")
            strTo = tTo.strftime("%Y-%m-%d %H:%M:%S")

            strQuery = 'select RecordTime, user, OptRemark from operation_record where RecordTime >= %s and RecordTime <= %s order by RecordTime'

            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (strFrom, strTo))
            if items == None:
                return []

            allRecords = []
            enUserNameList = []
            userIdList = []
            for item in items:
                allRecords.append(
                    [item[0].strftime('%Y-%m-%d %H:%M:%S'), item[1], item[2]]
                )

                if item[1] not in enUserNameList:
                    enUserNameList.append(item[1])

            if isinstance(userId, str):  # userId是om本地用户名
                if userId not in enUserNameList:
                    enUserNameList.append(userId)
            elif isinstance(userId, int):  # userId是本地用户ID
                userIdList.append(userId)

            dNameRole, dIdRole = self.getUserOfRoleByNameList(enUserNameList, userIdList)

            nCurUserRole = 1
            if userId in ["cx"]:
                nCurUserRole = 10
            elif isinstance(userId, str):
                nCurUserRole = dNameRole.get(userId, 1)
            elif isinstance(userId, int):
                nCurUserRole = dIdRole.get(userId, 1)
            elif userId == None:
                nCurUserRole = 1

            for record in allRecords:
                userName = record[1].lower()
                if record[1].find("云端-") != -1:
                    userName = record[1].replace("云端-", "")

                nRole = dNameRole.get(userName, None)

                bValidRole = False
                if userId in [-1, "cx"]:   # 超级管理员
                    bValidRole = True
                else:
                    # 2022-07-27 如果是策略的日志则认为是访客日志，可被权限等级>=3的用户访问
                    if nRole == None:
                        if nCurUserRole >= 3:
                            bValidRole = True
                    else:
                        if nRole <= nCurUserRole:
                            bValidRole = True

                bValidClassify = False
                if nType == None:
                    bValidClassify = True
                else:
                    if nType == 0:
                        if (record[2].find("登入") >= 0 or record[2].find("登出") >= 0) and record[1] in userNameList:
                            bValidClassify = True
                    elif nType == 1:
                        if record[1] in userNameList and record[2].find("登入") == -1 and record[2].find("登出") == -1:
                            bValidClassify = True
                    elif nType == 2:
                        if record[1] not in userNameList:
                            bValidClassify = True

                if bValidRole and bValidClassify:
                    recordList.append(record)

            return recordList
        except Exception as e:
            logging.error("ERROR in readOperationLogV2: %s" % e.__str__())
            return []

    def readOperationLog(self, strStart, strEnd, userId=None):
        try:
            ret = []
            tFrom = datetime.strptime(strStart, "%Y-%m-%d").replace(hour=0, minute=0, second=0, microsecond=0)
            tTo = datetime.strptime(strEnd, "%Y-%m-%d").replace(hour=23, minute=59, second=59, microsecond=0)

            strFrom = tFrom.strftime("%Y-%m-%d %H:%M:%S")
            strTo = tTo.strftime("%Y-%m-%d %H:%M:%S")

            strQuery = None
            param = None

            # 0是超级用户cx
            if userId in [0]:
                strQuery = 'select RecordTime,user,OptRemark from operation_record where RecordTime >= %s and ' \
                           'RecordTime <= %s and OptRemark not like "%程序%" and OptRemark not like "%3db%" order by RecordTime'
                param = (strFrom, strTo)
            else:
                # userId是om本地用户ID
                if isinstance(userId, int):
                    strQuery = 'SELECT op.RecordTime, op.user, op.OptRemark FROM operation_record op LEFT JOIN beopuser be ON' \
                               ' op.user = be.username WHERE be.userofrole <= (SELECT userofrole FROM beopuser WHERE userid=%s) ' \
                               'AND op.RecordTime >= %s AND op.RecordTime <= %s AND op.OptRemark NOT LIKE "%程序%" AND ' \
                               'op.OptRemark NOT LIKE "%3db%" ORDER BY op.RecordTime'
                    param = (userId, strFrom, strTo)

                # userId是om本地用户名
                elif isinstance(userId, str):
                    # userId 不存在，按访客处理
                    if self.get_user_id_by_name(userId) == None:
                        strQuery = 'SELECT op.RecordTime, op.user, op.OptRemark FROM operation_record op LEFT JOIN beopuser be ON' \
                                   ' op.user = be.username WHERE be.userofrole <= 1 ' \
                                   'AND op.RecordTime >= %s AND op.RecordTime <= %s AND op.OptRemark NOT LIKE "%程序%" AND ' \
                                   'op.OptRemark NOT LIKE "%3db%" ORDER BY op.RecordTime'
                        param = (strFrom, strTo)

                    # 若userId存在则按userid的权限处理
                    else:
                        strQuery = 'SELECT op.RecordTime, op.user, op.OptRemark FROM operation_record op LEFT JOIN beopuser be ON' \
                                   ' op.user = be.username WHERE be.userofrole <= (SELECT userofrole FROM beopuser WHERE username=%s) ' \
                                   'AND op.RecordTime >= %s AND op.RecordTime <= %s AND op.OptRemark NOT LIKE "%程序%" AND ' \
                                   'op.OptRemark NOT LIKE "%3db%" ORDER BY op.RecordTime'
                        param = (userId, strFrom, strTo)

                # 若userId为None则按访客权限处理
                elif userId == None:
                    strQuery = 'SELECT op.RecordTime, op.user, op.OptRemark FROM operation_record op LEFT JOIN beopuser be ON' \
                               ' op.user = be.username WHERE be.userofrole <= 1 ' \
                               'AND op.RecordTime >= %s AND op.RecordTime <= %s AND op.OptRemark NOT LIKE "%程序%" AND ' \
                               'op.OptRemark NOT LIKE "%3db%" ORDER BY op.RecordTime'
                    param = (strFrom, strTo)

            rv = []
            if strQuery and param:
                rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=param)

            if rv:
                for item in rv:
                    ret.append([item[0].strftime('%Y-%m-%d %H:%M:%S'),item[1], item[2]])

            return ret
        except Exception as e:
            logging.error("ERROR in readOperationLog: %s" % e.__str__())
            return []

    def getHistoryDataFromh1(self, dbname, id, timeStart, timeEnd):
        data = []
        try:
            startTime = datetime.strptime(timeStart,'%Y-%m-%d %H:%M:%S')
            endTime = datetime.strptime(timeEnd,'%Y-%m-%d %H:%M:%S')
        except ValueError as e:
            strErrorInfo = 'Time format is Wrong: %s, %s' % (timeStart, timeEnd)
            logging.error(strErrorInfo)
            logging.error(e)
            return data

        flagTime = startTime

        while True:
            if flagTime.year > endTime.year: break
            if flagTime.year == endTime.year and flagTime.month > endTime.month: break
            yDay = flagTime.strftime("%Y_%m")
            try:
                r = None
                if startTime.year == endTime.year and startTime.month == endTime.month:
                    strQ = 'SELECT time, value FROM historydata_hour_' + yDay + ' where time >= \'%s\' and time <= \'%s\' and pointname = \'%s\''%(timeStart, timeEnd, id,)
                    r = self._mysqlDbContainer.op_db_query(strQ)
                elif flagTime.year == startTime.year and flagTime.month == startTime.month:
                    strQ = 'SELECT time, value FROM historydata_hour_' + yDay + ' where time >= \'%s\' and pointname = \'%s\''%(timeStart, id,)
                    r = self._mysqlDbContainer.op_db_query(strQ)
                elif flagTime.year == endTime.year and flagTime.month == endTime.month:
                    strQ = 'SELECT time, value FROM historydata_hour_' + yDay + ' where time <= \'%s\' and pointname = \'%s\''%(timeEnd, id,)
                    r = self._mysqlDbContainer.op_db_query(strQ)
                else:
                    strQ = 'SELECT time, value FROM historydata_hour_' + yDay + ' where pointname = \'%s\'' % (id,)
                    r = self._mysqlDbContainer.op_db_query(strQ)

                if r:
                    for x in r:
                        data.append([x[0].strftime('%Y-%m-%d %H:%M:%S'), x[1]])
            except Exception as e:
                logging.error(e)
            finally:
                flagTime += timedelta(31)
        return data

    def get_history_data_all_one_moment(self, strTime, strTimeFormat):
        rv = []
        try:
            start_time = datetime.strptime(strTime, '%Y-%m-%d %H:%M:%S')
        except ValueError as e:
            logging.error('Time format is Wrong: %s' % (strTime))
            logging.error(e)

        if strTimeFormat == 'm1':
            tablename = 'historydata_minute_' + start_time.strftime('%Y_%m_%d')
        elif strTimeFormat == 'm5':
            tablename = 'historydata_5minute_' + start_time.strftime('%Y_%m_%d')
        elif strTimeFormat == 'h1':
            tablename = 'historydata_hour_' + start_time.strftime('%Y_%m')
        elif strTimeFormat == 'd1':
            tablename = 'historydata_day_' + start_time.strftime('%Y')
        try:
            strQ = 'SELECT time, pointname, value FROM ' + tablename + ' where time = \'%s\'' % (start_time,)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)
            if rv is None:
                return []
            return rv
        except Exception as e:
            logging.error(e)

        return rv

    def get_history_data_all_one_moment_padded(self, strTime, strTimeFormat):
        rvFilteredMap = {}
        rvAll=  []
        try:
            start_time = datetime.strptime(strTime, '%Y-%m-%d %H:%M:%S')
            start_time_look_before = start_time - timedelta(days=3)
        except ValueError as e:
            logging.error('Time format is Wrong: %s' % (strTime))
            logging.error(e)

        if strTimeFormat == 'm1':
            tablename = 'historydata_minute_' + start_time.strftime('%Y_%m_%d')
        elif strTimeFormat == 'm5':
            tablename = 'historydata_5minute_' + start_time.strftime('%Y_%m_%d')
        elif strTimeFormat == 'h1':
            tablename = 'historydata_hour_' + start_time.strftime('%Y_%m')
        elif strTimeFormat == 'd1':
            tablename = 'historydata_day_' + start_time.strftime('%Y')
        try:
            strQ = 'SELECT time, pointname, value FROM ' + tablename + ' where time <= \'%s\' and time>=\'%s\' order by time' % (start_time, start_time_look_before)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)
            if rv:
                for item in rv:
                    rvFilteredMap[item[1]] = item[2]

            for k,v in rvFilteredMap.items():
                rvAll.append(dict(name=k, value=v))

            return rvAll
        except Exception as e:
            logging.error(e)

        return rvAll

    def get_history_data_from_per_period(self, point_list, time_start, time_end, strTimeFormat='m1'):
        data = {}
        try:
            start_time = datetime.strptime(time_start,'%Y-%m-%d %H:%M:%S')
            end_time = datetime.strptime(time_end,'%Y-%m-%d %H:%M:%S')
        except ValueError as e:
            logging.error('Time format is Wrong: %s, %s' % (time_start, time_end))
            logging.error(e)
            return data
        if isinstance(point_list, list) is False:
            point_list = [point_list]
        flag_time = start_time
        start_time_date = start_time.strftime('%Y_%m_%d')
        end_time_date = end_time.strftime('%Y_%m_%d')
        point_join_str = '(\'%s\')' % '\',\''.join(point_list)

        for point in point_list:
            data[point] = []

        while True:
            tablename = None
            flag_time_date = flag_time.strftime('%Y_%m_%d')
            flag_array = flag_time_date.split('_')
            if flag_time_date > end_time_date: break;
            if strTimeFormat=='m1':
                tablename = 'historydata_minute_' + flag_time_date
            elif strTimeFormat=='m5':
                tablename = 'historydata_5minute_' + flag_time_date
            elif strTimeFormat == 'h1':
                tablename = 'historydata_hour_' + flag_array[0] + '_' + flag_array[1]
            elif strTimeFormat == 'd1' or strTimeFormat=='M1':
                tablename = 'historydata_day_' + flag_array[0]
            elif strTimeFormat == "s5":
                tablename = 'historydata_5second_' + flag_time_date

            if not tablename:
                continue

            try:
                r = None
                if start_time_date == end_time_date:
                    if len(point_list)==1:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where time >= \'%s\' and time <= \'%s\' and pointname = \'%s\' order by time' % (
                        time_start, time_end, point_list[0])
                    else:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where time >= \'%s\' and time <= \'%s\' and pointname in %s order by time'%(time_start, time_end, point_join_str)

                    r = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)

                elif flag_time_date == start_time_date:
                    if len(point_list)==1:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where time >= \'%s\' and pointname = \'%s\' order by time' % (time_start, point_list[0])
                    else:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where time >= \'%s\' and pointname in %s order by time' % (time_start, point_join_str)
                    r = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)

                elif flag_time_date == end_time_date:
                    if len(point_list)==1:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where time <= \'%s\' and pointname = \'%s\' order by time' % (time_end, point_list[0])
                    else:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where time <= \'%s\' and pointname in %s order by time' % (time_end, point_join_str)
                    r = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)

                else:
                    if len(point_list) == 1:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where pointname = \'%s\' order by time' % (point_list[0])
                    else:
                        strQ = 'SELECT pointname, time, value FROM ' + tablename + ' where pointname in %s order by time' % (point_join_str,)
                    r = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)

                if r is not None:
                    for x in r:
                        num = x[2]
                        try:
                            num = round(float(num) * 1000)/1000
                        except Exception as err:
                            pass

                        try:
                            strTimeRow = x[1].strftime('%Y-%m-%d %H:%M:%S')
                        except:
                            continue

                        if strTimeFormat=='M1':
                            if x[1].day!=1:
                                continue

                        strKey = x[0]
                        if not strKey in data.keys():
                            for kk in data.keys():
                                if kk.lower()== x[0].lower():
                                    strKey = kk
                        data[strKey].append([strTimeRow, num])
            except Exception as e:
                strError = 'ERROR in get_history_data_from_per_period: '+ e.__str__()
                print(strError)
                logging.error(strError)
            finally:
                if strTimeFormat == 'm1':
                    flag_time += timedelta(1)
                elif strTimeFormat == 'm5':
                    flag_time += timedelta(1)
                elif strTimeFormat == 'h1':
                    flag_time = time_get_first_day_of_next_month(flag_time)
                elif strTimeFormat == 'd1' or strTimeFormat=='M1':
                    flag_time = time_get_first_day_of_next_year(flag_time)
                elif strTimeFormat == "s5":
                    flag_time += timedelta(1)

        return data


    def checkTableExist(self, tableName):
        q='show tables like "%s"'%tableName
        rv = self._mysqlDbContainer.op_db_query(q)
        if rv is None:
            return False
        return len(rv) > 0

    def time_get_months_between(self, ts, te):
        nMonths = 0
        tCur = ts
        while tCur<te:
            tCur+=timedelta(32)
            tCur = tCur.replace(day=1)
            nMonths+=1

        return nMonths

    def time_get_next_month_begin(self, tt):
        tnext = tt.replace(day=1)
        tnext = tt+timedelta(days=32)
        tnext = tnext.replace(day=1, hour=0, minute=0, second=0)
        return tnext

    def time_get_pre_month_begin(self, tt):
        tprev = tt.replace(day=1)
        tprev = tt-timedelta(days=2)
        tprev = tprev.replace(day=1, hour=0, minute=0, second=0)
        return tprev

    def padData(self, data, timeStart, timeEnd, timeFormat):
        if len(data) <=0:#这里原先是<3,但问题是只查到一个数据点时没有1补齐，显示很难看
            return data

        tS = datetime.strptime(timeStart,'%Y-%m-%d %H:%M:%S')
        tE = datetime.strptime(timeEnd,'%Y-%m-%d %H:%M:%S')
        if timeFormat == 's5':
            ts = tS.replace(second=0) + timedelta(seconds=ceil(tS.second / 5) * 5)
            te = tE.replace(second=0) + timedelta(seconds=floor(tE.second / 5) * 5)
            l = floor((te - ts).total_seconds() / 5) + 1
            if l < 2:
                return data
            if len(data) < l:
                e = True
            else:
                e = False
            tv = [None] * l
            vv = [None] * l
            ev = [False] * l
            for x in data:
                t = datetime.strptime(x[0],'%Y-%m-%d %H:%M:%S')
                s = floor((t - tS).total_seconds() / 5)
                s = max(s, 0)
                if s < len(tv):
                    tv[s] = t
                    vv[s] = x[1]
            for i in range(1,l):
                if tv[i] is None:
                    if tv[i - 1] is not None:
                        tv[i] = tv[i - 1] + timedelta(seconds=5)
                        vv[i] = vv[i - 1]
                        ev[i] = True
            for i in range(l - 1,0,-1):
                if tv[i - 1] is None:
                    tv[i - 1] = tv[i] - timedelta(seconds=5)
                    vv[i - 1] = vv[i]
                    ev[i - 1] = True
            rv = []
            for x in range(l):
                rv.append([tv[x],vv[x]])
        if timeFormat == 'm1':
            ts = tS.replace(minute=0, second=0) + timedelta(seconds=ceil(tS.minute / 1) * 1 * 60)
            te = tE.replace(minute=0, second=0) + timedelta(seconds=floor(tE.minute / 1) * 1 * 60)
            l = floor((te - ts).total_seconds() / 1 / 60) + 1
            if l < 2:
                return data
            if len(data) < l:
                e = True
            else:
                e = False
            tv = [None] * l
            vv = [None] * l
            ev = [False] * l
            for x in data:
                t = datetime.strptime(x[0],'%Y-%m-%d %H:%M:%S')
                s = floor((t - tS).total_seconds() / 1 / 60)
                s = max(s, 0)
                if s < len(tv):
                    tv[s] = t
                    vv[s] = x[1]
            for i in range(1,l):
                if tv[i] is None:
                    if tv[i - 1] is not None:
                        tv[i] = tv[i - 1] + timedelta(seconds=60 * 1)
                        vv[i] = vv[i - 1]
                        ev[i] = True
            for i in range(l - 1,0,-1):
                if tv[i - 1] is None:
                    tv[i - 1] = tv[i] - timedelta(seconds=60 * 1)
                    vv[i - 1] = vv[i]
                    ev[i - 1] = True
            rv = []
            for x in range(l):
                rv.append([tv[x], vv[x]])
        if timeFormat == 'm5':
            ts = tS.replace(minute=0, second=0) + timedelta(seconds=ceil(tS.minute / 5) * 5 * 60)
            te = tE.replace(minute=0, second=0) + timedelta(seconds=floor(tE.minute / 5) * 5 * 60)
            l = floor((te - ts).total_seconds() / 5 / 60) + 1
            if l < 2:
                return data
            if len(data) < l:
                e = True
            else:
                e = False
            tv = [None] * l
            vv = [None] * l
            ev = [False] * l
            for x in data:
                t = datetime.strptime(x[0],'%Y-%m-%d %H:%M:%S')
                s = floor((t - tS).total_seconds() / 5 / 60)
                s = max(s, 0)
                if s < len(tv):
                    tv[s] = t
                    vv[s] = x[1]
            for i in range(1,l):
                if tv[i] is None:
                    if tv[i - 1] is not None:
                        tv[i] = tv[i - 1] + timedelta(seconds=60 * 5)
                        vv[i] = vv[i - 1]
                        ev[i] = True
            for i in range(l - 1,0,-1):
                if tv[i - 1] is None:
                    tv[i - 1] = tv[i] - timedelta(seconds=60 * 5)
                    vv[i - 1] = vv[i]
                    ev[i - 1] = True
            rv = []
            for x in range(l):
                rv.append([tv[x],vv[x]])
        if timeFormat == 'm30':
            ts = tS.replace(minute=0, second=0) + timedelta(seconds=ceil(tS.minute / 30) * 5 * 60)
            te = tE.replace(minute=0, second=0) + timedelta(seconds=floor(tE.minute / 30) * 5 * 60)
            l = floor((te - ts).total_seconds() / 30 / 60) + 1
            if l < 2:
                return data
            if len(data) < l:
                e = True
            else:
                e = False
            tv = [None] * l
            vv = [None] * l
            ev = [False] * l
            for x in data:
                t = datetime.strptime(x[0],'%Y-%m-%d %H:%M:%S')
                s = floor((t - tS).total_seconds() / 30 / 60)
                s = max(s, 0)
                if s < len(tv):
                    tv[s] = t
                    vv[s] = x[1]
            for i in range(1,l):
                if tv[i] is None:
                    if tv[i - 1] is not None:
                        tv[i] = tv[i - 1] + timedelta(seconds=60 * 30)
                        vv[i] = vv[i - 1]
                        ev[i] = True
            for i in range(l - 1,0,-1):
                if tv[i - 1] is None:
                    tv[i - 1] = tv[i] - timedelta(seconds=60 * 30)
                    vv[i - 1] = vv[i]
                    ev[i - 1] = True
            rv = []
            for x in range(l):
                rv.append([tv[x], vv[x]])
        if timeFormat == 'h1':
            ts = tS.replace(minute=0, second=0) + timedelta(seconds=ceil(tS.minute / 60) * 60 * 60)
            te = tE.replace(minute=0, second=0) + timedelta(seconds=floor(tE.minute / 60) * 60 * 60)
            l = floor((te - ts).total_seconds() / 60 / 60) + 1
            if l < 2:
                return data
            if len(data) < l:
                e = True
            else:
                e = False
            tv = [None] * l
            vv = [None] * l
            ev = [False] * l
            for x in data:
                t = datetime.strptime(x[0],'%Y-%m-%d %H:%M:%S')
                s = floor((t - tS).total_seconds() / 60 / 60)
                s= max(s, 0)
                if s<len(tv):
                    try:
                        tv[s] = t
                        vv[s] = x[1]
                    except Exception as e:
                        print(e.__str__())
            for i in range(1,l):
                if tv[i] is None:
                    if tv[i - 1] is not None:
                        tv[i] = tv[i - 1] + timedelta(seconds=60 * 60)
                        vv[i] = vv[i - 1]
                        ev[i] = True
            for i in range(l - 1,0,-1):
                if tv[i - 1] is None:
                    tv[i - 1] = tv[i] - timedelta(seconds=60 * 60)
                    vv[i - 1] = vv[i]
                    ev[i - 1] = True
            rv = []
            for x in range(l):
                rv.append([tv[x], vv[x]])
        if timeFormat == 'd1':
            ts = tS.replace(hour=0, minute=0, second=0) + timedelta(days=ceil(tS.hour / 24))
            te = tE.replace(hour=0, minute=0, second=0) + timedelta(days=floor(tE.hour / 24))
            l = (te - ts).days + 1
            if l < 2:
                return data
            if len(data) < l:
                e = True
            else:
                e = False
            tv = [None] * l
            vv = [None] * l
            ev = [False] * l
            for x in data:
                t = datetime.strptime(x[0],'%Y-%m-%d %H:%M:%S')
                s = (t - tS).days
                s = max(s, 0)
                if s < len(tv):
                    tv[s] = t
                    vv[s] = x[1]
            for i in range(1,l):
                if tv[i] is None:
                    if tv[i - 1] is not None:
                        tv[i] = tv[i - 1] + timedelta(1)
                        vv[i] = vv[i - 1]
                        ev[i] = True
            for i in range(l - 1,0,-1):
                if tv[i - 1] is None:
                    tv[i - 1] = tv[i] - timedelta(1)
                    vv[i - 1] = vv[i]
                    ev[i - 1] = True
            rv = []
            for x in range(l):
                rv.append([tv[x], vv[x]])
        if timeFormat == 'M1':
            ts = tS.replace(day=1, hour=0, minute=0, second=0)
            te = tE.replace(day=1, hour=0, minute=0, second=0)
            l = self.time_get_months_between(ts, te) + 1
            if l < 2:
                return data
            if len(data) < l:
                e = True
            else:
                e = False
            tv = [None] * l
            vv = [None] * l
            ev = [False] * l
            for x in data:
                t = datetime.strptime(x[0],'%Y-%m-%d %H:%M:%S')
                s = self.time_get_months_between(tS, t)
                s = max(s, 0)
                if s < len(tv):
                    tv[s] = t
                    vv[s] = x[1]
            for i in range(1,l):
                if tv[i] is None:
                    if tv[i - 1] is not None:
                        tv[i] = self.time_get_next_month_begin(tv[i - 1])
                        vv[i] = vv[i - 1]
                        ev[i] = True
            for i in range(l - 1,0,-1):
                if tv[i - 1] is None:
                    tv[i - 1] = self.time_get_pre_month_begin(tv[i])
                    vv[i - 1] = vv[i]
                    ev[i - 1] = True
            rv = []
            for x in range(l):
                rv.append([tv[x], vv[x]])
        return rv


    def get_history_data_from_redis(self, pointList, startTime, endTime, strTimeFormat):

        startTime = startTime.replace(second=0, microsecond=0)

        endTime = endTime.replace(second=0, microsecond=0)

        tspanStartToNow = datetime.now()-startTime
        if tspanStartToNow.total_seconds()>60*60*24*3: #3天前的数据不从redis里查
            if strTimeFormat=='m1' or strTimeFormat=='m5':
                if startTime!=endTime:
                    return None
                elif startTime.minute!=0:
                    return None
        if startTime==endTime and startTime.minute==0 and strTimeFormat in ["m1", "m5"]:
            strTimeFormat = 'h1'

        if strTimeFormat=='m1' or strTimeFormat=='m5':
            if (endTime-startTime).total_seconds()>=60*30: #不让去redis取超过30次，反而而不如去数据库里拿了
                return None

            rvStart = RedisManager.get_history_data_minutes(startTime, pointList)

            tSpan = timedelta(minutes=1)
            if strTimeFormat == 'm5':
                tSpan = timedelta(minutes=5)
            if rvStart:
                rvAll = {}
                bCheckStartData = True
                for iIndex in range(len(pointList)):
                    pt  = pointList[iIndex]
                    if rvAll.get(pt) is None:
                        rvAll[pt] = []
                    if rvStart[iIndex] is None:
                        bCheckStartData = False
                        break

                    #try to transfer to float
                    floatValue = get_number_format(rvStart[iIndex])
                    rvAll[pt].append([startTime.strftime('%Y-%m-%d %H:%M:%S'), floatValue])

                if bCheckStartData:
                    tCur = startTime + tSpan
                    while tCur<=endTime:
                        rvData = RedisManager.get_history_data_minutes(tCur, pointList)

                        for iIndex in range(len(pointList)):
                            pt = pointList[iIndex]
                            if rvData is None or rvData[iIndex] is None:
                                lastRefValue = None
                                lastRefValue = get_number_format(rvAll[pt][-1][1])

                                rvAll[pt].append([tCur.strftime('%Y-%m-%d %H:%M:%S'), lastRefValue])
                            else:
                                floatValue =  get_number_format(rvData[iIndex])
                                rvAll[pt].append([tCur.strftime('%Y-%m-%d %H:%M:%S'), floatValue])

                        tCur = tCur + tSpan
                    print('get_history_data_from_redis minutes success:%s, %s, %s to %s'%(strTimeFormat, str(pointList), startTime.strftime('%Y-%m-%d %H:%M:%S'), endTime.strftime('%Y-%m-%d %H:%M:%S')))
                    return rvAll

            return None



        rvAll = {}
        if isinstance(pointList, list) is False:
            pointList = [pointList]


        for pt in pointList:
            rvAll[pt] = []


        if startTime.minute!=0:
            return None

        tSpan = timedelta(hours=1)
        if strTimeFormat=='d1':
            tSpan = timedelta(days=1)

        for pt in pointList:
            tCur = startTime
            while tCur<=endTime:
                rv =  RedisManager.get_history_data(pt, tCur)
                if rv is None:
                    return None
                if isinstance(rv, str) and rv=='':
                    return None
                rvAll[pt].append([tCur.strftime('%Y-%m-%d %H:%M:%S'), rv])
                tCur+=tSpan

        return rvAll


    def get_history_data_padded(self, pointList, strTimeStart, strTimeEnd, strTimeFormat, dFilter={}, debugLogFile=None):
        result = []
        # invalid query filter:
        now_time = datetime.now()

        if strTimeStart is None or strTimeEnd is None or strTimeFormat is None:
            return {'error': 1, 'msg': '参数填写不全！'}

        if isinstance(pointList, str):
            pointList = [pointList]

        if not pointList:
            return {'error': 1, 'msg': 'pointList is empty！'}

        startTime = None
        endTime = None
        formatPattern = None

        try:

            startTime = datetime.strptime(strTimeStart, '%Y-%m-%d %H:%M:%S')
            if strTimeFormat=="m1":
                startTime = startTime.replace(second = 0)
            elif strTimeFormat=="m5":
                startTime = startTime.replace(minute = ceil(startTime.minute/5)*5)
            elif strTimeFormat=="h1":
                startTime = startTime.replace(minute = 0, second=0)
            elif strTimeFormat=="d1":
                startTime = startTime.replace(hour = 0,minute = 0, second = 0)
            elif strTimeFormat=="M1":
                startTime = startTime.replace(day=1, hour=0, minute=0, second=0)
            elif strTimeFormat == "s5":
                startTime = find_next_nearest_time_in_5_seconds(startTime)

            endTime = datetime.strptime(strTimeEnd, '%Y-%m-%d %H:%M:%S')
            if startTime>now_time:
                return {'error': 1, 'msg': '查询起始时间({startTime})大于当前时间({currentTime})，若当前时间有误，请检查服务器当前时间设置'.format(startTime=startTime.strftime("%Y-%m-%d %H:%M:%S"),
                                                                                                                                           currentTime=now_time.strftime("%Y-%m-%d %H:%M:%S"))}
            if endTime>now_time:
                endTime = now_time
            if startTime>endTime:
                return {'error': 1, 'msg': '查询起始时间({startTime})大于查询终止时间({endTime})'.format(startTime=startTime.strftime("%Y-%m-%d %H:%M:%S"),
                                                                                            endTime=endTime.strftime("%Y-%m-%d %H:%M:%S"))}

            if strTimeFormat in ["m1", "m5", "h1", "d1", "M1"]:
                endTime = endTime.replace(second=0)
                strTimeStart = startTime.strftime('%Y-%m-%d %H:%M:00')
                strTimeEnd = endTime.strftime('%Y-%m-%d %H:%M:00')
            elif strTimeFormat in ["s5"]:
                endTime = find_last_nearest_time_in_5_seconds(endTime)
                strTimeStart = startTime.strftime('%Y-%m-%d %H:%M:%S')
                strTimeEnd = endTime.strftime('%Y-%m-%d %H:%M:%S')

        except:
            return {'error': 1, 'msg': '时间格式有误！'}

        if startTime > endTime:
            print('get_history_data_padded:error params get_history_data_padded:[startTime:%s],[endTime:%s]'%(strTimeStart,strTimeEnd))
            logging.error('get_history_data_padded:error params get_history_data_padded:[startTime:%s],[endTime:%s]'%(strTimeStart,strTimeEnd))
            return {'error': 1, 'msg': '开始时间不能大于结束时间！'}
        if endTime > now_time:
            endTime = now_time
        if strTimeFormat == 'm1':
            formatPattern = '%Y-%m-%d %H:%M:00'
            if (endTime - startTime).days > 7:
                print('error: time range too long for m1 period data query')
                return {'error': 1, 'msg': '取样间隔为"m1"时，最多支持查询7天的数据！'}
        elif strTimeFormat == 'm5':
            formatPattern = '%Y-%m-%d %H:%M:00'
            if (endTime - startTime).days > 14:
                print('error: time range too long for m5 period data query :' + strTimeStart + strTimeEnd)
                return {'error': 1, 'msg': '取样间隔为"m5"时，最多支持查询14天的数据！'}
        elif strTimeFormat == 'h1':
            formatPattern = '%Y-%m-%d %H:00:00'
            if (endTime - startTime).days > 180:
                print('error: time range too long for h1 period data query ')
                return {'error': 1, 'msg': '取样间隔为"h1"时，最多支持查询180天的数据！'}
        elif strTimeFormat == 'd1':
            formatPattern = '%Y-%m-%d 00:00:00'
            if (endTime - startTime).days > 1000:
                print('error: time range too long for d1 period data query ')
                return {'error': 1, 'msg': '取样间隔为"d1"时，最多支持查询1000天的数据！'}
        elif strTimeFormat == 'M1':
            formatPattern = '%Y-%m-01 00:00:00'
            if (endTime - startTime).days > 2000:
                print('error: time range too long for M1 period data query ')
                return {'error': 1, 'msg': '取样间隔为"M1"时，最多支持查询1000天的数据！'}
        elif strTimeFormat == "s5":
            formatPattern = '%Y-%m-%d %H:%M:%S'
            if (endTime - startTime).total_seconds() > 3600 * 6:
                print('error: time range too long for s5 period data query')
                return {'error': 1, 'msg': '取样间隔为"s5"时，最多支持查询6小时的数据！'}
        else:
            print('error: time period format not supported')
            return {'error': 1, 'msg': '取样间隔配"%s"置不支持！'%(strTimeFormat,)}


        #如果是起点终点时刻相同，且为整小时，那么先尝试从redis取，全部成功则跳过从历史数据库查询
        bAllGotFromRedis = False
        bRedisAlive = RedisManager.is_alive()

        rv = None
        try:
            if bRedisAlive and strTimeFormat not in ["s5"]:  # 5秒钟数据不进redis查
                if debugLogFile:
                    log_info_to_file(debugLogFile, '3.1 start to read get_history_data_from_redis')
                rv = self.get_history_data_from_redis(pointList, startTime, endTime, strTimeFormat)
                if debugLogFile:
                    log_info_to_file(debugLogFile, '3.2 finish read get_history_data_from_redis')

                    get_history_data_padded_log_info(debugLogFile, "get_history_data_padded::data from redis: {data}".format(data=rv))

            if rv is not None:
                bAllGotFromRedis = True
        except:
            bAllGotFromRedis = False
            pass

        if debugLogFile:
            get_history_data_padded_log_info(debugLogFile, "get_history_data_padded::if succeed to get from redis: {data}".format(data=bAllGotFromRedis))

        if not bAllGotFromRedis:
            tBeginOfPeriod = startTime.replace(hour=0, minute=0, second=0)

            if strTimeFormat =='h1' or strTimeFormat=='d1' or strTimeFormat=='M1':
                tBeginOfPeriod = time_get_first_day_of_this_month(startTime)
            if debugLogFile:
                log_info_to_file(debugLogFile, '3.3 start read get_history_data_from_per_period')

            rv = self.get_history_data_from_per_period(pointList, tBeginOfPeriod.strftime('%Y-%m-%d %H:%M:%S'), strTimeEnd, strTimeFormat)

            if debugLogFile:
                get_history_data_padded_log_info(debugLogFile, "get_history_data_padded::data from mysql: {data}".format(data=rv))

            if debugLogFile:
                log_info_to_file(debugLogFile, '3.4 finish read get_history_data_from_per_period')

        if len(rv.keys()) == 0:
            return {'error': 1, 'msg': '未查询到历史数据'}

        generatedTimeList = self.generateTimeListInRange(startTime, endTime, strTimeFormat)
        hisDataDic = {'map': {}}
        maxTimeList = []
        for key, value in rv.items():
            trimedValue = self.prePadData(value, strTimeStart, strTimeEnd)
            paddedValue = self.padData(trimedValue, strTimeStart, strTimeEnd, strTimeFormat)

            filteredValue = paddedValue
            if dFilter.get(key):
                filteredValue = self.filter_data(dFilter.get(key).get("filterType"), dFilter.get(key).get("params"), paddedValue)

            hisDataDic['map'][key] = [x[1] for x in filteredValue]

            timeList = [item[0] if isinstance(item[0], str) else item[0].strftime(formatPattern) for item in filteredValue]

            if len(timeList) > len(maxTimeList):
                maxTimeList = timeList

            if bRedisAlive and not bAllGotFromRedis:
                resultValueList = hisDataDic['map'][key]
                if isinstance(resultValueList, list) and len(resultValueList)==0:
                    RedisManager.set_history_hour_data_list_for_one_point(key, generatedTimeList, ['']*len(generatedTimeList))
                else:
                    RedisManager.set_history_hour_data_list_for_one_point(key, timeList, resultValueList)

        if "time" not in hisDataDic.keys() and len(maxTimeList):
            hisDataDic['time'] = maxTimeList
        else:
            hisDataDic['time'] = generatedTimeList

        hisDataDic['lostTime'] = self.queryLostTimeInRange(startTime, endTime, strTimeFormat)
        return hisDataDic

    def queryFirstRecordTimeInCoreStatus(self,):
        lostIndexList = []
        ret = []

        strSQL = 'select min(timeFrom) from core_status where process="core"'
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)

        if result is None:
            return None

        if isinstance(result, list) and len(result)==0:
            return None

        return result[0][0]

    def queryLostTimeInRange(self, startTime, endTime, strTimeFormat):
        lostIndexList = []
        ret = []
        return ret  #20210926 太多这个mysql查询了，每次历史数据查询都会触发这个

        strSQL = 'select timeFrom, timeTo from core_status where process="core" and status01=1 and timeTo >= "%s" and timeFrom <= "%s" order by timeFrom' % (startTime.strftime('%Y-%m-%d %H:%M:%S'), endTime.strftime('%Y-%m-%d %H:%M:%S'))
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)



        goodDataTimeList= []
        if result is None:
            return []

        if isinstance(result, list) and len(result)==0:
            return []

        for item in result:
            goodDataTimeList.append(dict(timeFrom=item[0], timeTo=item[1]))

        tRecordGoodTimeStart = self.queryFirstRecordTimeInCoreStatus()

        tCur = startTime
        nTimeIndex = 0
        while tCur <= endTime:
            if strTimeFormat == 'm1':
                tCur += timedelta(seconds=60)
            elif strTimeFormat == 'm5':
                tCur += timedelta(seconds=60 * 5)
            elif strTimeFormat == 'h1':
                tCur += timedelta(hours=1)
            elif strTimeFormat == 'd1':
                tCur += timedelta(days=1)
            elif strTimeFormat == "s5":
                tCur += timedelta(seconds=5)

            bGood = False
            for tt in goodDataTimeList:
                if tRecordGoodTimeStart is None:
                    bGood = True
                elif tCur<= tt['timeTo'] and tCur>= tt['timeFrom']:
                    bGood = True
                elif tCur<= tRecordGoodTimeStart:
                    bGood = True

            if not bGood:
                lostIndexList.append(nTimeIndex)
            nTimeIndex+=1

        return lostIndexList


        return []


    def generateTimeListInRange(self, startTime, endTime, strTimeFormat):
        strTimeList = []
        try:
            if isinstance(startTime, str):
                startTime = datetime.strptime(startTime,'%Y-%m-%d %H:%M:%S')
            if isinstance(endTime, str):
                endTime = datetime.strptime(endTime, '%Y-%m-%d %H:%M:%S')
        except:
            return strTimeList

        tCur = startTime
        while tCur<=endTime:
            strTimeList.append(tCur.strftime('%Y-%m-%d %H:%M:%S'))
            if strTimeFormat=='m1':
                tCur+=timedelta(seconds=60)
            elif strTimeFormat=='m5':
                tCur+=timedelta(seconds=60*5)
            elif strTimeFormat=='h1':
                tCur+=timedelta(hours=1)
            elif strTimeFormat=='d1':
                tCur+=timedelta(days=1)
            elif strTimeFormat=='M1':
                tCur+=timedelta(days=32)
                tCur = tCur.replace(day=1)
            elif strTimeFormat == "s5":
                tCur += timedelta(seconds=5)

        return strTimeList

    def prePadData(self, valueList,strTimeStart, strTimeEnd,):
        if len(valueList)==0:
            return valueList

        strt0 = valueList[0][0]
        if strt0== strTimeStart:
            return valueList

        tStart = datetime.strptime(strTimeStart,'%Y-%m-%d %H:%M:%S')
        tEnd = datetime.strptime(strTimeEnd,'%Y-%m-%d %H:%M:%S')

        strte = valueList[-1][0]
        t0 = datetime.strptime(strt0,'%Y-%m-%d %H:%M:%S')
        te = datetime.strptime(strte,'%Y-%m-%d %H:%M:%S')

        newValueList = []
        itemPre = None
        for item in valueList:

            item_t = datetime.strptime(item[0], '%Y-%m-%d %H:%M:%S')
            if item_t< tStart:
                itemPre = item
                continue
            elif item_t>tStart and itemPre is not None and len(newValueList)==0:  #注意不是item_t>=tStart
                newValueList.append([strTimeStart, itemPre[1]])
                if item_t<=tEnd:#add this ,否则小时上查数据会多出一条来,20181205
                    newValueList.append([item[0], item[1]])#fixed bug, please add this
            elif item_t>tEnd:
                continue
            else:
                newValueList.append([item[0], item[1]])

        if len(newValueList)==0 and te<tStart:
            newValueList.append([strTimeStart, valueList[-1][1]])
        return newValueList

    def getThirdPartyInputTable(self, strThirdPartyName):
        strThirdPartyName = 'thirdparty'

        strSQL ='select pointname, pointvalue from realtimedata_input_%s'%(strThirdPartyName)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)
        retReal = []
        #print(str(pointList))
        if rv:
            for item in rv:
                retReal.append(dict(name=item[0], value=item[1]))

        return retReal

    def TestGetThirdPartyOutRealtimeData(self,strThirdPartyName):
        dbname = app.config['DATABASE']
        tablename = 'realtimedata_output_%s' % (strThirdPartyName)
        q = 'select count(1) from '+ tablename
        rv = self._mysqlDbContainer.op_db_query(dbname, q)
        if rv:
            return rv[0][0]
        else:
            return ''

    def TestSetThirdPartyOutRealtimeData(self, strPointNameList, strPointValueList, strThirdPartyName):
        dbname = app.config['DATABASE']
        tablename =  'realtimedata_output_%s'%(strThirdPartyName)
        if len(strPointNameList) <= 0:
            return dict(err=1, msg='no points')
        dbname = app.config['DATABASE']
        nDecimal = app.config.get('DECIMAL',None)
        if nDecimal is None:
            nDecimal = 2

        strNewPointValueList = []
        for pv in strPointValueList:
            try:
                pv = float(pv)
            except:pass
            if isinstance(pv, float):
                pv = round(pv, nDecimal)
            strNewPointValueList.append(pv)
        q = 'replace into ' + tablename + ' values(%s, %s)'
        alldata = []
        nIndex = 0
        tNow = datetime.now()
        for nIndex in range(len(strPointNameList)):
            alldata.append(tuple([strPointNameList[nIndex], str(strPointValueList[nIndex])]))
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainer.op_db_update_many(dbname, q, params)
        if bSuccess:
            return dict(err=0, msg=strNewPointValueList)
        else:
            return dict(err=1, msg='insert failed')



    def TestGetProcessTime(self,ProcessName):
        dbname = app.config['DATABASE']
        tablename = 'core_status'
        q = 'select timeTo from '+ tablename + ' where process = %s'
        rv = self._mysqlDbContainer.op_db_query(dbname, q, (ProcessName,))
        if rv:
            return rv[0][0]
        else:
            return ''


    def TestGetThirdPartyRealtimeData(self, strPointName, strThirdPartyName):
        dbname = app.config['DATABASE']
        tablename = 'realtimedata_input_%s' % (strThirdPartyName)
        q = 'select pointvalue from '+ tablename + ' where pointname = %s'
        rv = self._mysqlDbContainerRealTimePool.op_db_query(dbname,q,(strPointName,))
        if rv:
            return rv[0][0]
        else:
            return ''

    def setThirdPartyRealtimeData(self, strPointNameList, strPointValueList, strTimeList,  strThirdPartyName):
        dbname = app.config['DATABASE']
        strThirdPartyName = 'thirdparty'
        tablename =  'realtimedata_input_%s'%(strThirdPartyName)

        if len(strPointNameList) <= 0:
            return dict(err=1, msg='no points')
        dbname = app.config['DATABASE']
        nDecimal = app.config.get('DECIMAL', None)
        if nDecimal is None:
            nDecimal = 2

        strNewPointValueList = []
        for pv in strPointValueList:
            try:
                pv = float(pv)
            except:pass
            if isinstance(pv, float):
                pv = round(pv, nDecimal)
            strNewPointValueList.append(pv)
        q = 'replace into ' + tablename + ' values(%s,%s, %s)'
        alldata = []
        nIndex = 0
        tNow = datetime.now()
        strTimeNow = tNow.strftime('%Y-%m-%d %H:%M:%S')
        for nIndex in range(len(strPointNameList)):
            alldata.append(tuple([strTimeNow, strPointNameList[nIndex], str(strNewPointValueList[nIndex])]))
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainer.op_db_update_many(dbname, q, params)
        if bSuccess:
            return dict(err=0, msg=strNewPointValueList)
        else:
            return dict(err=1, msg='insert failed')

    def getAndClearThirdPartyOutputTable(self, pointList, strThirdPartyName):
        strThirdPartyName = 'thirdparty'

        strSQL ='select pointname, pointvalue from realtimedata_output_%s'%(strThirdPartyName)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)
        retReal = []
        #print(str(pointList))
        if rv:
            for item in rv:
                if not pointList:
                    retReal.append(dict(name=item[0], value=item[1]))
                elif pointList and item[0] in pointList:
                    retReal.append(dict(name=item[0], value=item[1]))

        if len(retReal)>0:
            self.delFromThirdPartyOutputTable(retReal,strThirdPartyName)

        logging.error('[ThirdParty] Commands Output: %s'%(str(retReal)))
        return retReal

    def delFromThirdPartyOutputTable(self, retReal,  strThirdPartyName):
        dbname = app.config['DATABASE']
        strTableName = 'realtimedata_output_%s' % (strThirdPartyName)
        q = 'delete from ' + strTableName + ' where pointname = %s and pointvalue =  %s'
        alldata = []
        nIndex = 0
        for item in retReal:
            alldata.append(tuple([item.get('name'), item.get('value')]))
            nIndex += 1
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainer.op_db_update_many(dbname, q, params)
        return bSuccess

    def setRealtimeData(self, strPointNameList, strPointValueList):

        #logging.error("[IMPORTANT]setRealtimeData: %s-%s" % (strPointNameList, strPointValueList))

        if isinstance(strPointNameList,list) and isinstance(strPointValueList, list):
            if len(strPointNameList)!= len(strPointValueList):
                return dict(err=1, msg='name size != value size')
            elif len(strPointValueList)==0:
                return dict(err=1, msg='value size = 0')
        dbname = app.config['DATABASE']
        nBatchSize = 300
        nIndexStart = 0
        nIndexEnd = min(nBatchSize, len(strPointNameList))
        while nIndexEnd <= len(strPointNameList) and nIndexStart<=nIndexEnd:
            q = 'insert into `realtimedata_output` values(%s, %s)'
            alldata = []
            nIndex = 0
            strPointNameListSub = strPointNameList[nIndexStart: nIndexEnd]
            strPointValueListSub = strPointValueList[nIndexStart: nIndexEnd]
            for strPointName in strPointNameListSub:
                if strPointValueListSub[nIndex] is None:
                    logging.error('ERROR in setRealtimeData request by omsite: pointname:%s, pointvalue is None'%(strPointName))

                strVV = str(strPointValueListSub[nIndex])
                alldata.append(tuple([strPointNameListSub[nIndex], strVV.encode('utf-8')]))
                nIndex +=1
            params = tuple(alldata)

            ##logging.error("[IMPORTANT]setRealtimeData params: %s" % (params,))

            bSuccess =self._mysqlDbContainer.op_db_update_many(dbname, q, params)

            #logging.error("[IMPORTANT]setRealtimeData:success: %s" % bSuccess)

            if bSuccess:
                nIndexStart = nIndexEnd+1
                nIndexEnd = min(nIndexStart+nBatchSize, len(strPointNameList))
                continue
            else:
                return dict(err=1, msg='命令已经发送处理，请耐心等待')

        return dict(err=0, msg='succeed')

    def setRealtimeDataOutputWireless(self, strPointNameList, strPointValueList):
        dbname = app.config['DATABASE']
        q = 'insert into `realtimedata_output_wireless` values(%s, %s)'
        alldata = []
        nIndex = 0
        for strPointName in strPointNameList:
            alldata.append(tuple([strPointNameList[nIndex], strPointValueList[nIndex]]))
            nIndex +=1
        params = tuple(alldata)
        bSuccess =self._mysqlDbContainer.op_db_update_many(dbname, q, params )
        if bSuccess:
            return dict(err=0, msg='')
        else:
            return dict(err=1, msg='insert failed')


    def saveRealtimeDatas(self, rtDatas):
        dbname = app.config['DATABASE']
        q = 'delete * from realtimedata_input'
        self._mysqlDbContainerRealTimePool.op_db_update(dbname, q, ())
        q = 'insert into `realtimedata_input` values("%s", "%s" , "%s")'
        alldata = []
        for key,data in rtDatas:
            alldata.append(tuple([key, data[0], data[1]]))
        params = tuple(alldata)
        bSuccess =self._mysqlDbContainerRealTimePool.op_db_update_many(dbname, q, params )

    def getServerOption(self, strOption):
        return '0'

    def validate_user(self, user_name, user_pwd):
        ''' 用户验证 '''
        sql = 'SELECT userid, username, userpwd, userofrole FROM beopuser WHERE BINARY `username`="%s"' % (user_name)
        result = []
        try:
            result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
            if result is None:
                return dict(err=1, msg='查询失败，请稍后再试')
        except Exception as e:
            logging.error('validate_user error:' + e.__str__())
            return dict(err=1, msg="查询失败，请稍后再试")

        if len(result)==0:
            return dict(err=1, msg='用户名不存在，请注意大小写严格区分')
        else:
            result = result[0]
            if result[2] == user_pwd:
                return dict(err=0, msg='', data=dict(id=result[0], name=result[1], role=int(result[3])))
            else:
                return dict(err=2, msg='密码错误')
    
    def add_user(self, user_info):
        if not user_info.get("username", None):
            return dict(success=False, msg="用户名不能为空", userid=None)

        sql = 'select count(userid) from beopuser where username="%s"' % user_info['username']
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
        if result is None:
            return -1
        if result[0][0] > 0:
            # 用户名已存在
            return -1

        sql = 'select userid from beopuser'
        infoList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
        nMaxId = 0
        if infoList:
            userIdList = []
            for info in infoList:
                try:
                    userIdList.append(int(info[0]))
                except:
                    pass

            nMaxId = max(userIdList)

        # sql = 'select max(userid) from beopuser'
        # result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
        # try:
        #     if result and len(result)>0:
        #         if len(result[0])>0:
        #             nMaxId = int(result[0][0])
        # except:
        #     nMaxId = 0

        nNewUserId = nMaxId + 1
        userName = user_info.get("username", "")
        userFullName = user_info.get("userfullname", "")
        userSex = user_info.get("usersex", "")
        userMobile = user_info.get("usermobile", "")
        userPassword = user_info.get("userpwd", "")
        userOfRole = user_info.get("userofrole", None)

        if not userName:
            return dict(success=False, msg="用户名不能为空", userid=None)
        if not userPassword:
            return dict(success=False, msg="密码不能为空", userid=None)
        if not userOfRole:
            return dict(success=False, msg="权限等级不能为空", userid=None)
        
        sql = 'insert into beopuser (userid, username,userfullname,usersex,usermobile,userpwd, userofrole) values(%d, "%s","%s","%s","%s", "%s", "%s")' % \
             (nNewUserId, userName, userFullName, userSex, userMobile, userPassword, userOfRole)
        try:
            result = self._mysqlDbContainer.op_db_insert(app.config['DATABASE'], sql)
        except Exception as e:
            print('add_user error:' + e.__str__())
            logging.error('add_user error:' + e.__str__())
        return {
            'success': result['success'],
            'userid': nNewUserId
        }
    
    def delete_users(self, user_ids):
        ''' 删除指定的用户记录 '''
        if isinstance(user_ids, list) is False:
            user_ids = [user_ids]

        sql = 'delete from beopuser where userid in (%s)' % ','.join([str(id) for id in user_ids])
        result = False
        try:
            result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        except Exception as e:
            print('delete_users error:' + e.__str__())
            logging.error('delete_users error:' + e.__str__())
        return result

    def modify_user(self, userid, user_info):
        ''' 修改用户记录 '''
        sql_builder = []
        for k, v in user_info.items():
            if isinstance(v, int):
                sql_builder.append('%s=%d' % (k, v))
            else:
                sql_builder.append('%s="%s"' % (k, v))
        sql = 'update beopuser set ' + ','.join(sql_builder) + ' where userid=%s' % (userid, )
        try:
            result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        except Exception as e:
            print('modify_user error:' + e.__str__())
            logging.error('modify_user error:' + e.__str__())
        return result
    
    def get_all_users(self):
        ''' 获取所有用户记录 '''
        sql = 'select userid, username, userofrole, userfullname, usersex, usermobile, useremail from beopuser'

        arr = [{"userid": -1,
                "username": "cx",
                "userofrole": "10",
                "username_zh": "系统内置根用户",
                "usersex": "",
                "usermobile": "",
                "useremail": ""}]
        try:
            result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
            if result:
                for item in result:
                    sex = "" if item[4] == None else item[4]
                    mobile = "" if item[5] == None else item[5]
                    strEmail = "" if item[6] == None else item[6]

                    arr.append({
                        "userid": item[0],
                        "username": item[1],
                        "userofrole": item[2],
                        "username_zh": item[3],
                        "usersex": sex,
                        "usermobile": mobile,
                        "useremail": strEmail
                    })
        except Exception as e:
            print('get_all_users error:' + e.__str__())
            logging.error('get_all_users error:' + e.__str__())
        return arr

    def is_user_read_only(self, strSource):
        nUserId = self.get_user_id_by_name(strSource)
        if nUserId is not None:
            userInfo = self.get_user_info(nUserId)
            if userInfo:
                nRole = userInfo.get('userofrole')
                if nRole == 0:
                    return True

        return False

    def get_user_id_by_name(self, strName):
        if strName == "cx":
            return -1

        sql = 'select userid  from beopuser where username = "%s"' % (strName)
        try:
            result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
            if result is None:
                return None
            if not len(result):
                return None
            if not len(result[0]):
                return None
            userId = int(result[0][0])
            return userId

        except Exception as e:
            print('get_user_info error:' + e.__str__())
            logging.error('get_user_info error:' + e.__str__())
        return None

    def get_user_info(self, userId):
        if userId == -1:
            return {
                    "username": "cx",
                    "userofrole": 10,
                    "userfullname": "内置根用户",
                    "usersex": "",
                    "usermobile": "",
                    "useremail": ""
                }

        sql = 'select userid, username, userofrole, userfullname, usersex, usermobile, useremail from beopuser where userid = %d'%(userId)
        info = {}
        try:
            result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
            arr = []
            if result and len(result)>0:
                item = result[0]

                sex = "" if item[4] == None else item[4]
                mobile = "" if item[5] == None else item[5]
                strEmail = "" if item[6] == None else item[6]

                info = {
                    "username": item[1],
                    "userofrole": item[2],
                    "userfullname":item[3],
                    "usersex": sex,
                    "usermobile": mobile,
                    "useremail": strEmail
                }
        except Exception as e:
            print('get_user_info error:' + e.__str__())
            logging.error('get_user_info error:' + e.__str__())
        return info

    def getProjS3db(self, proj_id):
        return app.config['S3DB_NAME']


    def updateWarningConfigId(self, pointname, type, script, newId):
        if script is None:
            sql = 'update warning_config set id = %d where pointname = "%s" and boolwarning = %d and script is NULL'  % (
            newId, pointname, type)
        else:
            sql = 'update warning_config set id = %d where pointname = "%s" and boolwarning = %d and script = "%s"' % (newId, pointname, type, script)
        bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)

        if not bSuccess:
            return False

        if RedisManager.is_alive():
            allWarningConfigList = self.getAllWarningConfig()
            RedisManager.set_warning_config(allWarningConfigList)
        return True

    def addWarningConfigItem(self, data):
        ret = dict(err=0, msg='')

        strOfDepartment = data.get('ofDepartment', '')
        strOfGroup = data.get('ofGroup', '')
        strOfPosition = data.get('ofPosition', '')
        strOfSystem = data.get('ofSystem', '')
        strTag = data.get('tag', '')

        hhlimit = None if data.get("hhlimit") == "" else data.get("hhlimit")
        hlimit = None if data.get("hlimit") == "" else data.get("hlimit")
        llimit = None if data.get("llimit") == "" else data.get("llimit")
        lllimit = None if data.get("lllimit") == "" else data.get("lllimit")

        nNewId = data.get('id', -1)
        if nNewId<=0:
            nNewId = self.getMaxIdInTable('warning_config')+1
        strQuery = 'insert into warning_config(HHEnable,HEnable, LENable,LLenable,HHLimit,HLimit,LLimit,LLLimit,pointname,' \
          'HHwarninginfo,Hwarninginfo,Lwarninginfo,LLwarninginfo,boolwarning,boolwarninginfo,boolwarninglevel,unitproperty01,id, ofDepartment, ofGroup, ofPosition, ofSystem, tag,unitproperty02,unitproperty03,unitproperty04,unitproperty05) ' \
          'values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)'

        param = (data.get("hhenable"), data.get("henable"), data.get("lenable"), data.get("llenable"), hhlimit,
                 hlimit, llimit, lllimit, data.get("pointname"), data.get("hhinfo"),data.get("hinfo"),data.get("linfo"),
                 data.get("llinfo"), data.get("type"), data.get("boolWarningInfo"), data.get("boolWarningLevel"),
                 data.get("warningGroup"), nNewId, strOfDepartment, strOfGroup, strOfPosition, strOfSystem,strTag,
        data.get('unitproperty02',''),data.get('unitproperty03',''),data.get('unitproperty04',''),data.get('unitproperty05',''))

        bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, parameter=param)

        if not bSuccess:
            ret['err'] =1

        if RedisManager.is_alive():
            allWarningConfigList = self.getAllWarningConfig()
            RedisManager.set_warning_config(allWarningConfigList)
        return ret

    def removeWarningConfigAll(self):
        sql = 'delete from warning_config'
        bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        if RedisManager.is_alive():
            allWarningConfigList = self.getAllWarningConfig()
            RedisManager.set_warning_config(allWarningConfigList)
        return bSuccess

    def removeWarningConfigItemById(self, nID):
        sql = 'delete from warning_config where id = %d' % (nID)
        bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        if RedisManager.is_alive():
            allWarningConfigList = self.getAllWarningConfig()
            RedisManager.set_warning_config(allWarningConfigList)
        return bSuccess

    def removeWarningConfigItem(self, strPointName, nWarningType):
        ret = dict(err=0, msg='')
        sql = 'delete from warning_config where pointname = "%s" and boolwarning = %d'%(strPointName, nWarningType)
        bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        if not bSuccess:
            ret['err'] = 1

        if RedisManager.is_alive():
            allWarningConfigList = self.getAllWarningConfig()
            RedisManager.set_warning_config(allWarningConfigList)
        return ret

    def addWarningConfigRuleItem(self,nWarningLevel,  strScript, pointname, ofDepartmentName, ofPositionName, ofSystemName, ofGroupName, tags, strWarningInfo, strWarningGroupCustom,strInfoDetail, ruleId,
                                 strUnitProperty02, strUnitProperty03,  strUnitProperty04,  strUnitProperty05):
        ret = dict(err=0, msg='')
        if ruleId<=0:
            ruleId  = self.getMaxIdInTable('warning_config')+1

        q = 'insert into  warning_config(id,boolwarning,boolWarningLevel, boolwarninginfo, script, ofDepartment, ofPosition, ofSystem, ofGroup,pointname, unitproperty01, tag, infodetail, unitproperty02, unitproperty03, unitproperty04, unitproperty05) values (%s,%s, %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)'


        params = (ruleId,3, nWarningLevel, strWarningInfo,  strScript, ofDepartmentName, ofPositionName, ofSystemName, ofGroupName, pointname, strWarningGroupCustom, tags, strInfoDetail, strUnitProperty02, strUnitProperty03,  strUnitProperty04,  strUnitProperty05)

        bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], q, params)
        if not bSuccess:
            return -1

        if RedisManager.is_alive():
            allWarningConfigList = self.getAllWarningConfig()
            RedisManager.set_warning_config(allWarningConfigList)
        return ruleId

    def getAllWarningConfig(self):
        ret = []

        strSQL = 'select HHEnable, HEnable, LEnable, LLEnable, HHLimit, HLimit, LLimit, LLLimit, pointname, HHwarninginfo,Hwarninginfo, Lwarninginfo, LLwarninginfo, boolwarning, boolwarninginfo,' \
                 'boolwarninglevel, unitproperty01, unitproperty02,unitproperty03,unitproperty04,unitproperty05,' \
                 'script,ofPosition,ofSystem,ofDepartment,ofGroup,tag,id, infodetail from warning_config'
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)
        if result is None:
            return ret

        for item in result:
            nType = 1
            if item[13] is not None:
                nType = int(item[13])

            nID = -1
            if item[27] is not None:
                nID = int(item[27])

            if nType<3:
                ret.append(dict(hhenable=item[0], henable = item[1], lenable = item[2], llenable = item[3], hhlimit = item[4], hlimit = item[5],
                                llimit = item[6], lllimit= item[7], pointname= item[8], hhinfo = item[9], hinfo = item[10], linfo = item[11], llinfo= item[12], type= nType, boolWarningInfo = item[14], boolWarningLevel = item[15],
                                warningGroup = item[16], id=nID,ofPosition=item[22], ofSystem=item[23], ofDepartment=item[24], ofGroup=item[25],
                                tag=item[26], unitproperty01= item[16], unitproperty02= item[17],unitproperty03= item[18],unitproperty04= item[19],unitproperty05= item[20]))
            elif nType==3: #rule
                ret.append(dict(type=nType, boolWarningInfo=item[14], script=item[21], ofPosition=item[22], ofSystem=item[23], ofDepartment=item[24], ofGroup=item[25],
                                tag=item[26], infoDetail = item[28],  id= nID, pointname=item[8], boolWarningLevel= item[15],warningGroup = item[16],
                                unitproperty01=item[16], unitproperty02=item[17], unitproperty03=item[18],unitproperty04=item[19], unitproperty05=item[20]))



        return ret

    def editWarningConfigItem(self, data):
        ret = dict(err=0, msg='')

        self.removeWarningConfigItem(data.get('pointname'), data.get('type'))
        self.addWarningConfigItem(data)

        return ret

    def editWarningConfigRuleItemById(self, nID, nWarningLevel,  strScript, pointname, ofDepartmentName, ofPositionName, ofSystemName, ofGroupName, tags, strWarningInfo, strWarningGroupCustom):
        q = 'update  warning_config set boolWarningLevel = %s, script = %s, ofDepartment=%s, ofPosition=%s, ofSystem=%s, ofGroup=%s,pointname=%s ,tag=%s, boolwarninginfo = %s' \
            ', unitproperty01= %s where id = %s'

        params = ( nWarningLevel, strScript, ofDepartmentName, ofPositionName, ofSystemName, ofGroupName, pointname, tags, strWarningInfo, strWarningGroupCustom,  nID)

        bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], q, params)
        if RedisManager.is_alive():
            allWarningConfigList = self.getAllWarningConfig()
            RedisManager.set_warning_config(allWarningConfigList)
        return bSuccess

    '''
    ['pointname', 'warningInfo', 'startTime', 'endTime']
    '''
    def getWarningRecordTableName(self, ttime=None):
        return 'warningrecord'
        #if ttime is None:
        #    ttime = datetime.now()
        #return 'warningrecord_%s'%(ttime.strftime('%Y'))


    def getRealtimeWarningList(self, nSeconds, bUnClosed = False):
        ret = []

        tnow = datetime.now()
        tEnd = tnow - timedelta(seconds= nSeconds)

        strEndTime = tEnd.strftime('%Y-%m-%d %H:%M:%S')
        strSQL = 'select * from %s where endtime > "%s"' %(self.getWarningRecordTableName(tnow), strEndTime)
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)

        if result is None:
            return ret

        for item in result:
            nID = -1
            nRuleId = -1
            try:
                nID = item[8]
                nRuleId = item[9]
            except:
                pass

            strOfPosition = ''
            strOfSystem=''
            strOfDepartment=''
            strOfGroup=''
            strTag=''
            strInfoDetail = ''
            strUnitProperty03 = ''
            strUnitProperty04 = ''
            strUnitProperty05 = ''
            strGoodRange = ""
            strPointValue = ""

            nConfirmStatus = item[5]
            nConfirmUserId = item[6]

            strConfirmRemark = ""
            tConfirmOpTime = None

            nIgnoreStatus = 0
            nIgnoreUserId = None

            strIgnoreRemark = ""
            tIgnoreOpTime = None
            tIgnoreToTime = None

            nCloseStatus = 0
            nCloseUserId = None
            strCloseRemark = ""
            tCloseOpTime = None

            if len(item)>=15:
                strOfPosition = item[10]
                strOfSystem = item[11]
                strOfDepartment = item[12]
                strOfGroup = item[13]
                strTag = item[14]
            if len(item)>=16:
                strInfoDetail = item[15]
            if len(item)>=18:
                strGoodRange = item[16]
                strPointValue = item[17]
            if len(item)>=21:
                strUnitProperty03 = item[18]
                strUnitProperty04 = item[19]
                strUnitProperty05 = item[20]

            # 21：confirmRemark
            if len(item) >= 22:
                if isinstance(item[21], str):
                    strConfirmRemark = item[21]
                else:
                    strConfirmRemark = ""

            if len(item) >= 23:
                if isinstance(item[22], datetime):
                    tConfirmOpTime = item[22]
                elif isinstance(item[22], str):
                    try:
                        tConfirmOpTime = datetime.strptime(item[22], "%Y-%m-%d %H:%M:%S")
                    except:
                        pass

            if len(item) >= 24:
                nIgnoreStatus = item[23]

            if len(item) >= 25:
                nIgnoreUserId = item[24]

            if len(item) >= 26:
                if isinstance(item[25], str):
                    strIgnoreRemark = item[25]
                else:
                    strIgnoreRemark = ""

            if len(item) >= 27:
                if isinstance(item[26], datetime):
                    tIgnoreOpTime = item[26]
                elif isinstance(item[26], str):
                    try:
                        tIgnoreOpTime = datetime.strptime(item[26], "%Y-%m-%d %H:%M:%S")
                    except:
                        pass

            if len(item) >= 28:
                if isinstance(item[27], datetime):
                    tIgnoreToTime = item[27]
                elif isinstance(item[27], str):
                    try:
                        tIgnoreToTime = datetime.strptime(item[27], "%Y-%m-%d %H:%M:%S")
                    except:
                        pass

            #if len(item) >= 29:
            #    nCloseStatus = item[28]

            #if len(item) >= 30:
            #    nCloseUserId = item[29]

            if len(item) >= 31:
                if isinstance(item[30], str):
                    strCloseRemark = item[30]
                else:
                    strCloseRemark = ""

            nCloseStatus = 0
            if len(item) >= 32:
                if isinstance(item[31], datetime):
                    tCloseOpTime = item[31]
                    try:
                        if tCloseOpTime > item[0]:
                            nCloseStatus = 1
                    except:
                        nCloseStatus = 0
                elif isinstance(item[31], str):
                    try:
                        tCloseOpTime = datetime.strptime(item[31], "%Y-%m-%d %H:%M:%S")
                        if tCloseOpTime > item[0].strftime('%Y-%m-%d %H:%M:%S'):
                            nCloseStatus = 1
                    except:
                        nCloseStatus = 0
            if bUnClosed==True and nCloseStatus==1:
                continue
            ret.append(dict(time=item[0].strftime('%Y-%m-%d %H:%M:%S'), info = item[2], level = item[3], endtime = item[4].strftime('%Y-%m-%d %H:%M:%S'),
                            strBindPointName = item[7], id=nID, ruleId=nRuleId, ofPosition=strOfPosition, ofSystem= strOfSystem, ofDepartment=strOfDepartment,
                            ofGroup=strOfGroup, tag= strTag, infoDetail=strInfoDetail, goodRange=strGoodRange, strBindPointValue = strPointValue,
                            strUnitProperty01 = strGoodRange, strUnitProperty02 = strPointValue, strUnitProperty03 = strUnitProperty03, strUnitProperty04= strUnitProperty04,strUnitProperty05 = strUnitProperty05,
                            nConfirmStatus=nConfirmStatus, nConfirmUserId=nConfirmUserId, strConfirmRemark=strConfirmRemark,tConfirmOpTime=tConfirmOpTime,
                            nIgnoreStatus=nIgnoreStatus,nIgnoreUserId=nIgnoreUserId,strIgnoreRemark=strIgnoreRemark,tIgnoreOpTime=tIgnoreOpTime, tIgnoreToTime=tIgnoreToTime,
                            nCloseStatus=nCloseStatus, nCloseUserId=nCloseUserId, strCloseRemark=strCloseRemark,tCloseOpTime=tCloseOpTime))

        return ret

    def updateWarningRecordEndTimeToNow(self, strInfo, strHappenTime, strInfoDetail):
        strTimeNow = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        strSQL = 'update %s set endtime = "%s", infodetail = "%s" where time =  "%s" and info =  "%s"' % ( self.getWarningRecordTableName(),
            strTimeNow,strInfoDetail,  strHappenTime, strInfo
        )
        result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strSQL)
        return result

    def insertWarningRecord(self, strInfo, nLevel, strBindPointName, nRuleId, strOfPosition,
                            strOfSystem, strOfDepartment, strOfGroup, strTag, strInfoDetail,
                            strGoodRange = "", strPointValue = "", strUnitProperty03 = "",strUnitProperty04 = "",strUnitProperty05 = ""):
        strWarningRecordTableName = self.getWarningRecordTableName()
        curActiveWarningList = self.getRealtimeWarningList(60*60*24, True)
        nRetryCount = 5
        while curActiveWarningList is None and nRetryCount>0:
            curActiveWarningList = self.getRealtimeWarningList(60 * 60*24, True)
            nRetryCount-=1
            time.sleep(1)

        if nRetryCount==0:
            return False

        bFound = False
        strHappenTime = None
        for item in curActiveWarningList:
            if item.get('info')== strInfo and item.get('nCloseStatus')!=1:
                strHappenTime = item.get('time')
                bFound = True

        if bFound:
            return self.updateWarningRecordEndTimeToNow(strInfo, strHappenTime, strInfoDetail)
        else:
            strTimeNow = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            newId = self.getMaxIdInTable(strWarningRecordTableName)+1
            strSQL = 'insert into %s(time,code,info,level,endtime,confirmed,confirmeduser,bindpointname, id, ruleId, ' \
                     'ofPosition, ofSystem, ofDepartment, ofGroup, tag, infodetail, unitproperty01, unitproperty02, unitproperty03, unitproperty04, unitproperty05) values("%s","%s","%s","%s","%s","%s","%s","%s", %d, %d,"%s","%s","%s","%s","%s","%s","%s","%s","%s","%s","%s")' % ( strWarningRecordTableName,
                strTimeNow, '0', strInfo, str(nLevel), strTimeNow, '0','',strBindPointName, newId, nRuleId, strOfPosition, strOfSystem, strOfDepartment, strOfGroup, strTag, strInfoDetail,strGoodRange, strPointValue, strUnitProperty03, strUnitProperty04, strUnitProperty05
            )
            result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strSQL)
            return result


    def getWarningGroupList(self,):
        ret = []
        sql = 'select distinct unitproperty01 from warning_config'
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
        if result is None:
            return ret
        for item in result:
            ret.append(item[0])
        return ret

    def getHistoryWarningList(self, strTimeStart, strTimeEnd):
        ret = []

        strTimeFrom = "{0} 00:00:00".format(strTimeStart)
        strTimeTo = "{0} 23:59:59".format(strTimeEnd)

        strWarningRecordTableName = self.getWarningRecordTableName()
        strSQL = 'select * from %s where time >= "%s" and time <= "%s"' % (strWarningRecordTableName,  strTimeFrom, strTimeTo)
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)

        if result is None:
            return ret

        for item in result:
            ret.append(dict(time=item[0].strftime('%Y-%m-%d %H:%M:%S'), info = item[2], level = item[3], endtime = item[4].strftime('%Y-%m-%d %H:%M:%S'), strBindPointName = item[7]))
        return ret

    def getUnit01(self, strKey):
        ret = None
        strSQL = 'select unitproperty02 from unit01 where unitproperty01 = "%s"' % (strKey,)
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)

        if result is None:
            return None

        for item in result:
            ret = item[0]
            break
        return ret

    def getListInUnit01ByKey(self, strKey):
        ret = []
        strSQL = 'select unitproperty02 from unit01 where unitproperty01 = "%s"' % (strKey,)
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)

        if result is None:
            return ret

        for item in result:
            ret.append(item[0])

        return ret

    def saveUnit01(self, strKey, strConfig):
        curConfig = self.getUnit01(strKey)
        sql = ''
        result = None
        try:
            if curConfig is None:
                sql = 'insert into unit01(unitproperty01,unitproperty02) values(%s,%s)'
                result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql, (strKey, strConfig))
            # unit01 表里插入
            else:
                sql = 'update unit01 set unitproperty02 = %s  where unitproperty01 = %s'
                result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'],sql, (strConfig, strKey))

        except Exception as e:
            print('saveUnit01 error:' + e.__str__())
            logging.error('saveUnit01 error:' + e.__str__())
        return result

    def removeByDoubleKeyInUnit01(self, strKey, strKey2):
        curConfig = self.getDoubleKeyInUnit01(strKey, strKey2)
        sql = ''
        result = None
        if curConfig is not None:
            sql = 'delete from unit01 where unitproperty01 ="%s" and unitproperty02 = "%s"' % (
            strKey, strKey2)
        try:
            result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        except Exception as e:
            print('removeByDoubleKeyInUnit01 error:' + e.__str__())
            logging.error('removeByDoubleKeyInUnit01 error:' + e.__str__())
        return result

    def removeByKeyInUnit01(self, strKey):
        curConfig = self.getUnit01(strKey)
        sql = ''
        result = None
        if curConfig is not None:
            sql = 'delete from unit01 where unitproperty01 ="%s"' % (strKey)
        try:
            result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        except Exception as e:
            print('removeByDoubleKeyInUnit01 error:' + e.__str__())
            logging.error('removeByDoubleKeyInUnit01 error:' + e.__str__())
        return result

    def saveDoubleKeyInUnit01(self, strKey, strKey2, strConfig):
        curConfig = self.getDoubleKeyInUnit01(strKey, strKey2)
        sql = ''
        result = None
        try:
            if curConfig is None:
                sql = 'insert into unit01(unitproperty01,unitproperty02,unitproperty03) values(%s,%s,%s)'
                result = self._mysqlDbContainer.op_db_insert(app.config['DATABASE'], sql, (strKey, strKey2, strConfig))
            # unit01 表里插入
            else:
                sql = 'update unit01 set unitproperty03 = %s  where unitproperty01 = %s and unitproperty02 = %s'
                result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql, (strConfig, strKey, strKey2))

        except Exception as e:
            print('saveUnit01 error:' + e.__str__())
            logging.error('saveUnit01 error:' + e.__str__())
        return result

    def updateUnit01(self, strKey, strKey2, strNewKye2):
        sql = ''
        try:
            sql = 'update unit01 set unitproperty02 = %s where unitproperty01 = %s and unitproperty02 = %s'
            result = self._mysqlDbContainer.op_db_update(app.config['DATABASE'],sql, (strNewKye2,  strKey, strKey2))

        except Exception as e:
            print('updateUnit01 error:' + e.__str__())
            logging.error('updateUnit01 error:' + e.__str__())
        return result

    def getDoubleKeyInUnit01(self, strKey, strKey2):
        ret = None
        strSQL = 'select unitproperty03 from unit01 where unitproperty01 = "%s" and unitproperty02= "%s"'%(strKey, strKey2)
        # 差距呢beop数据库中的数据，查询语句strSQL
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)

        if result is None:
            return ret

        for item in result:
            ret = item[0]
            break
        return ret

    def pointExistInRealtimeInput(self, pointName):
        dbname = app.config['DATABASE']
        q = 'select pointname from realtimedata_input where pointname = %s'
        params = tuple([pointName])
        rv = self._mysqlDbContainerRealTimePool.op_db_query(dbname, q, params)
        if rv and len(rv)>0:
            return True
        return False


    def update_realtimedata_input(self, pointList, valueList):
        iIndex = 0
        dbname = app.config['DATABASE']
        for point in pointList:
            if self.pointExistInRealtimeInput(point):
                q = 'update `realtimedata_input` set pointvalue = %s where pointname = %s'
                params = tuple([valueList[iIndex], point])
                bSuccess = self._mysqlDbContainerRealTimePool.op_db_update(dbname, q, params)
            else:
                q = 'insert into `realtimedata_input`(time,pointname,pointvalue) values(%s, %s, %s)'
                params = tuple([datetime.now().strftime('%Y-%m-%d %H:%M:%S'), point, valueList[iIndex]])
                bSuccess = self._mysqlDbContainerRealTimePool.op_db_insert(dbname, q, params)

            iIndex +=1
        return iIndex

    def get_and_clear_realtimedata_output(self, tableName = None):

        dbname = app.config['DATABASE']
        if tableName is None:
            tableName = 'realtimedata_output'
        q = 'select pointname, pointvalue from `%s`'%(tableName)

        rv = self._mysqlDbContainer.op_db_query(dbname, q,)
        if rv is None:
            return []
        dataAll = []
        allDeleteSQLData = []
        sqlDel = 'delete from ' + tableName + ' where pointname = "%s" and pointvalue= "%s"'
        for item in rv:
            dataAll.append(dict(pointname=item[0], pointvalue=item[1]))


            allDeleteSQLData.append(tuple([item[0], item[1]]))
        params = tuple(allDeleteSQLData)
        bSuccess = self._mysqlDbContainer.op_db_update_many(dbname, sqlDel, params)
        return dataAll

    def getDBVersion(self):
        dbname = app.config['DATABASE']
        q = 'select incocontent from beopinfo where infoname = "%s"'%('version')

        rv = self._mysqlDbContainer.op_db_query(dbname, q, )
        if rv and len(rv)>0:
            return rv[0][0]

        return '0'

    def updateDBVersion(self, strVersion):
        dbname = app.config['DATABASE']
        q = 'select incocontent from beopinfo where infoname = "%s"' % ('version')

        rv = self._mysqlDbContainer.op_db_query(dbname, q, )
        if rv and len(rv) > 0:
            q = 'update beopinfo set incocontent = "%s" where infoname = "%s"'%(strVersion, 'version')
        else:
            q = 'insert into beopinfo(infoname,incocontent) values("%s", "%s")'%('version', strVersion)

        return self._mysqlDbContainer.op_db_update(dbname, q, )

    def createRealtimeinputTablesIfNotExist(self):
        dbname = app.config['DATABASE']
        bSuccess = True

        q = 'CREATE TABLE IF NOT EXISTS  `realtimedata_input_%s` (\
                          `time` timestamp NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\
                          `pointname` varchar(64) NOT NULL,\
                          `pointvalue` varchar(2560) NOT NULL\
                        ) ENGINE=InnoDB DEFAULT CHARSET=utf8; ' % ('vpoint')
        bSuccess = bSuccess and self._mysqlDbContainerRealTimePool.op_db_update(dbname, q, )

        for strItem in ['modbus', 'opc', 'bacnet', 'siemens_direct', 'wireless']:
            q = 'CREATE TABLE IF NOT EXISTS  `realtimedata_input_%s` (\
                  `time` timestamp NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\
                  `pointname` varchar(64) NOT NULL,\
                  `pointvalue` varchar(2560) NOT NULL\
                ) ENGINE=MEMORY DEFAULT CHARSET=utf8; '%(strItem)

            bSuccess = bSuccess and self._mysqlDbContainerRealTimePool.op_db_update(dbname, q, )
        return bSuccess


    def createRealtimeOutputTablesIfNotExist(self):
        dbname = app.config['DATABASE']
        bSuccess = True
        for strItem in ['modbus', 'opc', 'bacnet', 'siemens_direct', 'wireless', 'vpoint']:
            q = '\
                CREATE TABLE IF NOT EXISTS  `realtimedata_output_%s` (\
                  `pointname` varchar(64) NOT NULL,\
                  `pointvalue` varchar(2560) NOT NULL\
                ) ENGINE=MEMORY DEFAULT CHARSET=utf8; '%(strItem)
            bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )
        return bSuccess

    def createRealtimeOutputTablesFromCloudIfNotExist(self):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = '\
            CREATE TABLE IF NOT EXISTS  `realtimedata_output_%s` (\
              `pointname` varchar(64) NOT NULL,\
              `pointvalue` varchar(2560) NOT NULL\
            ) ENGINE=MEMORY DEFAULT CHARSET=utf8; '%('cloud')
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )
        return bSuccess

    def getMySQLDataDir(self):
        rv = None
        dbname = app.config['DATABASE']
        q = 'show variables like "datadir"'

        rv = self._mysqlDbContainer.op_db_query(dbname, q, )
        try:
            if rv and len(rv) > 0:
                rv =  rv[0][1]
        except:
            rv = None

        return rv

    # update logic parameter to mysql realtimely.
    def  InsertLogicParameters(self, strThreadName, strLogicName, strSetType, strVariableName, strLinkName, strLinkType, strCondition):
        dbname = app.config['DATABASE']
        bSuccess = True
        sql = 'insert into \
               unit02(unitproperty01, unitproperty02, unitproperty03, unitproperty04, unitproperty05, unitproperty06, unitproperty07) \
               values("%s", "%s", "%s", "%s", "%s", "%s", "%s")'%(strThreadName, strLogicName, strSetType, strVariableName, strLinkName, strLinkType, strCondition)

        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, sql, )

        return bSuccess


    def remove_history_data_in_time_range(self, strTimeFrom, strTimeTo, strPointNameList, strPeriod, strWhere=None):
         try:
            t_time = datetime.strptime(strTimeFrom,'%Y-%m-%d %H:%M:00')

            flag_time_date = t_time.strftime('%Y_%m_%d')

            strTime_m1 = t_time.strftime('%Y-%m-%d %H:%M:00')

            tablename = 'historydata_minute_' + flag_time_date

            if strPeriod=="m5":
                tablename = 'historydata_5minute_' + flag_time_date
            elif strPeriod == "m1":
                tablename = 'historydata_minute_' + flag_time_date
            elif strPeriod == "h1":
                flag_time_date = t_time.strftime('%Y_%m')
                tablename = 'historydata_hour_' + flag_time_date
            elif strPeriod == "d1":
                flag_time_date = t_time.strftime('%Y')
                tablename = 'historydata_day_' + flag_time_date

            if strWhere is None or len(strWhere)<=0 or strWhere== " ":
                strQ = 'delete from %s where pointname in (%s) and time>= \'%s\' and time<=\'%s\'' % (tablename, str(strPointNameList).replace('[', '').replace(']', ''), strTimeFrom, strTimeTo)
            else:
                strQ = 'delete from %s where pointname in (%s) and time>= \'%s\' and time<=\'%s\' and %s' % (
                tablename, str(strPointNameList).replace('[', '').replace(']', ''), strTimeFrom, strTimeTo, strWhere)

            logging.error('remove_history_data_in_time_range: where:%s, sql command:%s'%( strWhere, strQ))
            r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
            if not r:
                return False



         except Exception as e:
             logging.error(e)
             print('ERROR in remove_history_data:' + e.__str__())
             return False

         return True

    def remove_history_data(self, strTimeAt, strPointName):
         try:
            t_time = datetime.strptime(strTimeAt,'%Y-%m-%d %H:%M:00')

            flag_time_date = t_time.strftime('%Y_%m_%d')
            flag_array = flag_time_date.split('_')

            strTime_m1 = t_time.strftime('%Y-%m-%d %H:%M:00')

            tablename = 'historydata_minute_' + flag_time_date
            self.createHistoryTablesIfNotExist(tablename)
            strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_m1 )
            r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
            if not r:
                return False

            if t_time.minute%5==0:
                tablename = 'historydata_5minute_' + flag_time_date
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_m1)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False

            if t_time.minute==0:
                flag_time_date = t_time.strftime('%Y_%m')
                tablename = 'historydata_hour_' + flag_time_date
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_m1)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False

            if t_time.hour==0 and t_time.minute==0:
                flag_time_date = t_time.strftime('%Y')
                tablename = 'historydata_day_' + flag_time_date
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_m1)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False

         except Exception as e:
             logging.error(e)
             print('ERROR in remove_history_data:' + e.__str__())

         return True


    def modify_history_data(self, strTimeAt, strPointName, strPointValue):
        self.remove_history_data(strTimeAt, strPointName)
        try:
            t_time = datetime.strptime(strTimeAt,'%Y-%m-%d %H:%M:00')

            flag_time_date = t_time.strftime('%Y_%m_%d')
            flag_array = flag_time_date.split('_')

            strTime_m1 = t_time.strftime('%Y-%m-%d %H:%M:00')

            tablename = 'historydata_minute_' + flag_time_date
            self.createHistoryTablesIfNotExist(tablename)
            strQ = 'insert into %s (time, pointname, value) values(\'%s\',\'%s\',\'%s\')' % (tablename, strTimeAt, strPointName, strPointValue )
            r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
            if not r:
                return False

            if t_time.minute%5==0:
                tablename = 'historydata_5minute_' + flag_time_date
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'insert into %s (time, pointname, value) values(\'%s\',\'%s\',\'%s\')' % (
                tablename, strTimeAt, strPointName, strPointValue)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False

            if t_time.minute==0:
                flag_time_date = t_time.strftime('%Y_%m')
                tablename = 'historydata_hour_' + flag_time_date
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'insert into %s (time, pointname, value) values(\'%s\',\'%s\',\'%s\')' % (
                tablename, strTimeAt, strPointName, strPointValue)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False

            if t_time.hour==0 and t_time.minute==0:
                flag_time_date = t_time.strftime('%Y')
                self.createHistoryTablesIfNotExist(tablename)
                tablename = 'historydata_day_' + flag_time_date
                strQ = 'insert into %s (time, pointname, value) values(\'%s\',\'%s\',\'%s\')' % (
                tablename, strTimeAt, strPointName, strPointValue)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False

            if RedisManager.is_alive():
                RedisManager.set_history_data_list([strPointName], [strTimeAt], [strPointValue])
        except Exception as e:
            logging.error(e)

        return True

    def insert_history_data(self, strPointName, strTime, strValue,strUserId, strOriginalValue):
        data = {}
        try:
            t_time = datetime.strptime(strTime,'%Y-%m-%d %H:%M:%S')
            # t_time = t_time.replace(second=0)  需要支持s5
        except ValueError as e:
            logging.error('Time format is Wrong: %s' % (strTime))
            logging.error(e)
            return False, 'time format wrong'

        try:

            flag_time_date = t_time.strftime('%Y_%m_%d')
            flag_array = flag_time_date.split('_')

            # 1分钟表
            strTime_m1 = t_time.strftime('%Y-%m-%d %H:%M:00')
            tablename = 'historydata_minute_' + flag_time_date
            self.createHistoryTablesIfNotExist(tablename)
            strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_m1 )
            r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
            if not r:
                return False, 'Delete historydata_minute_ Failed'

            strQ = 'insert into %s(pointname, time, value) values( \'%s\',  \'%s\' ,\'%s\')' % (tablename, strPointName, strTime_m1, strValue)
            r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
            if not r:
                return False
            elif RedisManager.is_alive():
                RedisManager.set_history_data_list([strPointName], [strTime], [strValue])

            # 5秒钟表
            if t_time.second % 5 == 0:
                strTime_s5 = t_time.strftime('%Y-%m-%d %H:%M:%S')
                tablename = 'historydata_5second_' + flag_time_date
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_s5)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'delete historydata_5second_ failed'
                strQ = 'insert into %s (pointname, time, value) values( \'%s\',  \'%s\', \'%s\')' % (tablename, strPointName, strTime_s5, strValue)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'insert historydata_5second_ failed'

            # 5分钟表
            if t_time.minute%5==0:
                t_time_m5 = t_time.replace(minute= 5*int(t_time.minute/5))
                strTime_m5 = t_time_m5.strftime('%Y-%m-%d %H:%M:00')
                tablename = 'historydata_5minute_' + flag_time_date
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_m5)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'delete historydata_5minute_ failed'
                strQ = 'insert into %s(pointname, time, value) values( \'%s\',  \'%s\', \'%s\')' % (
                tablename, strPointName, strTime_m5, strValue)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'insert historydata_5minute_ failed'

            # 小时表
            if t_time.minute==0:
                strTime_h1 = t_time.strftime('%Y-%m-%d %H:00:00')
                tablename = 'historydata_hour_' + flag_array[0] + '_' + flag_array[1]
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_h1)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'delete hour table failed'
                strQ = 'insert into %s(pointname, time, value) values( \'%s\',  \'%s\', \'%s\')' % (
                tablename, strPointName, strTime_h1, strValue)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'insert hour table failed'

            # 天表
            if t_time.hour==0 and t_time.minute==0:
                strTime_d1 = t_time.strftime('%Y-%m-%d 00:00:00')
                tablename = 'historydata_day_' + flag_array[0]
                self.createHistoryTablesIfNotExist(tablename)
                strQ = 'delete from %s where pointname=\'%s\' and time= \'%s\'' % (tablename, strPointName, strTime_d1)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'delete day table failed'
                strQ = 'insert into %s(pointname, time, value) values( \'%s\',  \'%s\', \'%s\')' % (
                    tablename, strPointName, strTime_d1, strValue)
                r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
                if not r:
                    return False, 'insert day table failed'

            if RedisManager.is_alive():
                RedisManager.set_history_data_list([strPointName], [strTime], [strValue])
            #strQ = 'insert into %s(pointname, time, pointvalue, userid, originalvalue) values( \'%s\',  \'%s\', \'%s\', \'%s\', \'%s\')' % (
            #    'manual_historydata_modification_record', strPointName, strTime, strValue, strUserId, strOriginalValue)
            #r = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQ)
            #if not r:
            #    return False

        except Exception as e:
            logging.error(e)
            return False, 'Exception: %s'%(e.__str__())

        return True, ''


    def createHistoryTablesIfNotExist(self, strHistoryDataTableName):
        dbname = app.config['DATABASE']
        bSuccess = True

        q = 'CREATE TABLE IF NOT EXISTS  `%s`.`%s` (\
                          `time` timestamp NOT NULL DEFAULT CURRENT_TIMESTAMP,\
                          `pointname` varchar(255) NOT NULL,\
                          `value` varchar(20000) NOT NULL,\
                         INDEX `Index_t_n` (`time`,`pointname`) \
                        ) ENGINE=InnoDB DEFAULT CHARSET=utf8; ' % (app.config['DATABASE'], strHistoryDataTableName)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )
        return bSuccess

    def addSchedule(self, strScheduleName, strUserName,isloop= 0, strPointDefine= ''):
        dbname = app.config['DATABASE']
        bSuccess = True
        # loop是mysql的关键保留字，这里替换为isloop
        q = 'INSERT INTO `schedule_list` (`type`, `name`, `point`,`loop`, `enable`, `author`) VALUES(0,%s,%s,%s,0,%s)'
        params = (strScheduleName, str(strPointDefine) ,str(isloop), strUserName)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, params)
        return bSuccess

    def getSchedule(self):
        dbname = app.config['DATABASE']
        result = None
        q = 'select id,`type`,`name`,point,`loop`,`enable`,`author` from  schedule_list'
        result = self._mysqlDbContainer.op_db_query(dbname,q)

        if result is None:
            return None

        data = []
        for item in result:
            data.append({
                "id" : item[0],
                'type': item[1],
                'name': item[2],
                'point': item[3],
                'isloop': item[4],
                'enable' : item[5],
                'author' : item[6]
            })


        return data


    def removeSchedule(self, id):
        dbname = app.config['DATABASE']
        bSuccess = True
        cSuccess = True
        q='Delete from schedule_list where id=%d' %(id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q)

        cSuccess= self.removeWeekDays(id)

        return bSuccess and cSuccess

    def editSchedule(self, id,strScheduleName, isloop, strPointDefine= ''):
        dbname = app.config['DATABASE']
        bSuccess = True
        q='Update schedule_list set name="%s",point="%s",`loop`=%d where id="%d"' %(strScheduleName, str(strPointDefine) ,isloop,id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q)
        return bSuccess

    def enableSchedule(self,enable, id):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'update schedule_list set enable=%d where id=%d'% (enable,id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q)
        return bSuccess

    def addWeekDays(self, weekday, timeFrom, timeTo, value, id):
        strHM = timeFrom.split(':')
        if len(strHM)<2:
            return
        timeFrom =  strHM[0]+':'+ strHM[1]

        strHM = timeTo.split(':')
        if len(strHM) < 2:
            return

        timeTo = strHM[0] + ':' + strHM[1]

        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'Insert into schedule_info_weeky(weekday,timeFrom,timeTo,value,groupid,timeExecute)values(%d,"%s","%s","%s",%d,DATE_FORMAT(now(),\"%%Y-%%m-%%d %%H:%%i:%%s\")) ' % (weekday, timeFrom, timeTo, value, id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q)
        return bSuccess

    def removeWeekDays(self, id):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'Delete from schedule_info_weeky where groupid=%d' % (id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q)
        return bSuccess

    def fromGrouopIdGetSecheduleTask(self,groupid):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'select id,weekday,timeFrom,timeTo,value,groupid,timeExecute from schedule_info_weeky where groupid=%d' % (groupid)
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result is None:
            return None
        data = []
        days = [1, 2, 3, 4, 5, 6, 7]
        for day in days:
            rv = self.existsDays( day, result)
            if rv['status'] == True :
                data.append(rv['data'])
            else:data.append({
                "id": day,
                'weekday': day,
                'timeFrom': '08:00',
                'timeTo': '20:00',
                'value': '',
                'groupid': groupid,
                'timeExecute': '',
                'enable': 0
            })
        return data

    def fromGrouopIdGetSecheduleTaskV2(self,groupid):
        res = list()
        try:
            dbname = app.config['DATABASE']
            # q = 'select id,weekday,timeFrom,timeTo,value,groupid,timeExecute from schedule_info_weeky where groupid=%d and value=1 order by timeFrom asc' % (groupid)
            q = """SELECT
                        id,
                        weekday,
                        timeFrom,
                        timeTo,
                        value,
                        groupid,
                        timeExecute
                    FROM
                        schedule_info_weeky
                    WHERE
                        groupid =% d
                    AND
                    VALUE
                        = 1
                    ORDER BY
                        timeFrom ASC""" % (groupid)
            schedulePool = self._mysqlDbContainer.op_db_query(dbname, q)

            if schedulePool is None:
                return None

            days = [1, 2, 3, 4, 5, 6, 7]

            for day in days:
                dictScheduleDay = dict(id=day, weekday=day)
                count = 0
                for schedule in schedulePool:
                    if schedule[1] == day:
                        enableKey = "enable{0}".format(count) if count > 0 else "enable"
                        timeFromKey = "timeFrom{0}".format(count) if count > 0 else "timeFrom"
                        timeToKey = "timeTo{0}".format(count) if count > 0 else "timeTo"

                        dictScheduleDay.update({
                            enableKey: 1,
                            timeFromKey: intToTime(schedule[2]),
                            timeToKey: intToTime(schedule[3])
                        })

                        count += 1

                while count < 5:
                    defaultTime = "08:00"
                    listKeys = list(dictScheduleDay.keys())

                    listKeysProcessed = list()
                    if len(listKeys):
                        listKeysProcessed = [item.replace("timeTo", "") for item in listKeys if item.startswith("timeTo")]

                    if len(listKeysProcessed):
                        try:
                            maxTimeTo = max(listKeysProcessed)
                            if maxTimeTo.isdigit() or maxTimeTo == "":
                                defaultTime = dictScheduleDay.get("timeTo{0}".format(maxTimeTo))
                        except:
                            traceback.print_exc()

                    enableKey = "enable{0}".format(count) if count > 0 else "enable"
                    timeFromKey = "timeFrom{0}".format(count) if count > 0 else "timeFrom"
                    timeToKey = "timeTo{0}".format(count) if count > 0 else "timeTo"

                    dictScheduleDay.update({
                        enableKey: 0,
                        timeFromKey: defaultTime,
                        timeToKey: defaultTime
                    })

                    count += 1

                res.append(dictScheduleDay)

        except:
            traceback.print_exc()
        return res

    def existsDays(self,day,result):
        rv ={"status":False,"data":None}
        for item in result:
            if day == item[1]:
                rv['status'] = True
                rv['data'] = {
                "id": item[0],
                'weekday': item[1],
                'timeFrom': intToTime(item[2]),
                'timeTo': intToTime(item[3]),
                'value': item[4],
                'groupid': item[5],
                'timeExecute': item[6],
                'enable': 1
                }
                return rv

        return rv

    def existsDaysV2(self, day, result):
        rv = {"status": False, "data": None}
        timeRangesOfOneDay = []
        for item in result:
            if day == item[1]:
                oneRecord = dict()
                oneRecord.update({"status": True})
                oneRecord.update({
                    "data": {
                    "id": item[0],
                    'weekday': item[1],
                    'timeFrom': intToTime(item[2]),
                    'timeTo': intToTime(item[3]),
                    'value': item[4],
                    'groupid': item[5],
                    'timeExecute': item[6],
                    'enable': 1}
                })
                timeRangesOfOneDay.append(oneRecord)
        return timeRangesOfOneDay

    def checkBuildFDDDB(self):
        dbname = app.config['DATABASE']
        bSuccess = True

        q = 'CREATE TABLE IF NOT EXISTS `fdd_list_item` (\
								    `name` varchar(255) NOT NULL,\
									`name_ch` varchar(255) NOT NULL,\
									`enabled` int(10) unsigned NOT NULL DEFAULT 0,\
									`createtime` timestamp NOT NULL DEFAULT CURRENT_TIMESTAMP,\
									`updatetime` timestamp,\
									`description` text,\
									`ofEquipment` varchar(255) DEFAULT NULL,\
									`ofZone` varchar(255) DEFAULT NULL,\
									`ofResponseParty` varchar(255) DEFAULT NULL,\
									`ofFaultClassify` varchar(255) DEFAULT NULL,\
									PRIMARY KEY (`name`)\
									) ENGINE=InnoDB DEFAULT CHARSET=utf8;'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )
        return bSuccess


    def checkCreateIfNotExistFDDTable(self, tTime):
        dbname = app.config['DATABASE']
        bSuccess = True
        q= 'CREATE TABLE IF NOT EXISTS `fdd_result_%s` (`fddname` varchar(255) NOT NULL,\
								  `fddtime` timestamp NOT NULL DEFAULT CURRENT_TIMESTAMP,\
								  `faultStatus` int(10) NOT NULL,\
								  `grade` int(10) NOT NULL,\
								   `content` varchar(2560) NOT NULL,\
								    `analysis` varchar(2560) NOT NULL,\
									 `suggestion` varchar(2560) NOT NULL,\
									  `risk` varchar(2560) NOT NULL,\
                                    `day_history` varchar(2560) NOT NULL,\
								  PRIMARY KEY (`fddname`,`fddtime`)) \
								  ENGINE=InnoDB DEFAULT CHARSET=utf8;'%(tTime.strftime('%Y_%m'))

        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )
        return bSuccess

    #暂时不分表，数据错误，记录起点和终点时间，按道理数据不会过分庞大
    def checkCreateIfNotExistFddDataErrorTable(self):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'CREATE TABLE IF NOT EXISTS `fdd_data_error` (`id` INT NOT NULL,\
                                    `pointname` varchar(255) NOT NULL,\
                                  `timeFrom` timestamp NOT NULL,\
                                  `timeTo` timestamp NOT NULL,\
                                  `errorgrade` int(10) NOT NULL,\
                                    `reason`  TEXT,\
                                 `rs1`  TEXT, `rs2`  TEXT,`rs3`  TEXT,`rs4`  TEXT,`rs5`  TEXT,\
                                  PRIMARY KEY (`id`)) \
                                  ENGINE=InnoDB DEFAULT CHARSET=utf8;'#%(strTime.strftime('%Y'))

        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )
        return bSuccess

    def checkBuildIOSampleDB(self):
        dbname = app.config['DATABASE']
        bSuccess = True

        q = 'CREATE TABLE IF NOT EXISTS  `data_sample_list` (\
                                 `id` INT NOT NULL,\
                                  `time` timestamp NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\
                                  `name` varchar(256) NOT NULL,\
                                  `author` varchar(256) NOT NULL\
                                ) ENGINE=InnoDB DEFAULT CHARSET=utf8; '
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )


        q = 'CREATE TABLE IF NOT EXISTS  `data_sample_io_define` (\
                       `sampleId` INT NOT NULL,\
                        `timeFrom` timestamp,\
                        `timeTo` timestamp,\
                      `input01` varchar(256),`input02` varchar(256),`input03` varchar(256),`input04` varchar(256),`input05` varchar(256), \
                      `input06` varchar(256),`input07` varchar(256),`input08` varchar(256),`input09` varchar(256),`input10` varchar(256), \
                        `input11` varchar(256),`input12` varchar(256),`input13` varchar(256),`input14` varchar(256),`input15` varchar(256), \
                       `input16` varchar(256),`input17` varchar(256),`input18` varchar(256),`input19` varchar(256),`input20` varchar(256), \
        `input21` varchar(256),`input22` varchar(256),`input23` varchar(256),`input24` varchar(256),`input25` varchar(256), \
                              `input26` varchar(256),`input27` varchar(256),`input28` varchar(256),`input29` varchar(256),`input30` varchar(256), \
        `input31` varchar(256),`input32` varchar(256),`input33` varchar(256),`input34` varchar(256),`input35` varchar(256), \
                              `input36` varchar(256),`input37` varchar(256),`input38` varchar(256),`input39` varchar(256),`input40` varchar(256), \
        `input41` varchar(256),`input42` varchar(256),`input43` varchar(256),`input44` varchar(256),`input45` varchar(256), \
                              `input46` varchar(256),`input47` varchar(256),`input48` varchar(256),`input49` varchar(256),`input50` varchar(256), \
                        `output01` varchar(256),`output02` varchar(256),`output03` varchar(256),`output04` varchar(256),`output05` varchar(256), \
                        `output06` varchar(256),`output07` varchar(256),`output08` varchar(256),`output09` varchar(256),`output10` varchar(256), \
                     `inserttime` timestamp, KEY `Index_1` (`sampleId`,`timeFrom`,`timeTo`)\
                    ) ENGINE=InnoDB DEFAULT CHARSET=utf8; '

        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbname, q, )
        return bSuccess

    def getSampleData(self, strSampleName, strTimeFrom, strTimeTo):
        dbname = app.config['DATABASE']
        dataList =[]
        nID = BEOPDataAccess.getInstance().getOrCreateSampleIdByName(strSampleName)
        if isinstance(nID, int) and nID < 0:
            return dataList

        strQuery = 'select input01, input02, input03, input04, input05, input06,input07,input08, input09, input10,' \
            'input11, input12, input13, input14, input15, input16,input17,input18, input19, input20,' \
            ' input21, input22, input23, input24, input25, input26,input27,input28, input29, input30,' \
            ' input31, input32, input33, input34, input35, input36,input37,input38, input39, input40,' \
            ' input41, input42, input43, input44, input45, input46,input47,input48, input49, input50,' \
            ' output01, output02,output03, output04, output05, output06, output07, output08,output09,output10,' \
            'timeFrom, timeTo from data_sample_io_define where sampleId = %s'
        params = (nID,)

        if strTimeFrom and strTimeTo:
            strQuery = 'select input01, input02, input03, input04, input05, input06,input07,input08, input09, input10,' \
                'input11, input12, input13, input14, input15, input16,input17,input18, input19, input20,' \
                ' input21, input22, input23, input24, input25, input26,input27,input28, input29, input30,' \
                ' input31, input32, input33, input34, input35, input36,input37,input38, input39, input40,' \
                ' input41, input42, input43, input44, input45, input46,input47,input48, input49, input50,' \
                ' output01, output02,output03, output04, output05, output06, output07, output08,output09,output10,' \
                'timeFrom, timeTo from data_sample_io_define where sampleId = %s and timeFrom > %s and timeTo < %s'
            params = (nID, strTimeFrom, strTimeTo)

        result = self._mysqlDbContainer.op_db_query(dbname, strQuery, params)
        if result and isinstance(result, list):
            for item in result:
                dataList.append(dict(
                    input01 = item[0],
                    input02=item[1],
                    input03=item[2],
                    input04=item[3],
                    input05=item[4],
                    input06=item[5],
                    input07=item[6],
                    input08=item[7],
                    input09=item[8],
                    input10=item[9],
                    input11=item[10],
                    input12=item[11],
                    input13=item[12],
                    input14=item[13],
                    input15=item[14],
                    input16=item[15],
                    input17=item[16],
                    input18=item[17],
                    input19=item[18],
                    input20=item[19],
                    input21=item[20],
                    input22=item[21],
                    input23=item[22],
                    input24=item[23],
                    input25=item[24],
                    input26=item[25],
                    input27=item[26],
                    input28=item[27],
                    input29=item[28],
                    input30=item[29],
                    input31=item[30],
                    input32=item[31],
                    input33=item[32],
                    input34=item[33],
                    input35=item[34],
                    input36=item[35],
                    input37=item[36],
                    input38=item[37],
                    input39=item[38],
                    input40=item[39],
                    input41=item[40],
                    input42=item[41],
                    input43=item[42],
                    input44=item[43],
                    input45=item[44],
                    input46=item[45],
                    input47=item[46],
                    input48=item[47],
                    input49=item[48],
                    input50=item[49],

                    output01=item[50],
                    output02=item[51],
                    output03=item[52],
                    output04=item[53],
                    output05=item[54],
                    output06=item[55],
                    output07=item[56],
                    output08=item[57],
                    output09=item[58],
                    output10=item[59],
                    timeFrom=str(item[60]),
                    timeTo=str(item[61]),
                ) )

        return dataList

    def insertDataSampleIOs(self, sampleList):
        dbname = app.config['DATABASE']
        bSuccess = True

        alldata = []
        for item in sampleList:
            alldata.append(tuple(item))
        params = str(alldata)[1:-1]

        q = 'insert into data_sample_io_define(sampleId, timeFrom, timeTo, input01, input02, input03, input04, input05, input06,input07,input08, input09, input10,' \
            'input11, input12, input13, input14, input15, input16,input17,input18, input19, input20,' \
            ' input21, input22, input23, input24, input25, input26,input27,input28, input29, input30,' \
            ' input31, input32, input33, input34, input35, input36,input37,input38, input39, input40,' \
            ' input41, input42, input43, input44, input45, input46,input47,input48, input49, input50,' \
            ' output01, output02,output03, output04, output05, output06, output07, output08,output09,output10) values %s' % params

        bSuccess = self._mysqlDbContainer.op_db_update_many(dbname, q, ())
        return bSuccess


    def insertDataSampleIO(self, sampleId, strStartTime, strEndTime, strInput01,strInput02,strInput03,strInput04,strInput05,strInput06,strInput07,strInput08,strInput09,strInput10,
                           strInput11, strInput12, strInput13, strInput14, strInput15, strInput16, strInput17,strInput18, strInput19, strInput20,
                           strInput21, strInput22, strInput23, strInput24, strInput25, strInput26, strInput27,
                           strInput28, strInput29, strInput30,
                           strInput31, strInput32, strInput33, strInput34, strInput35, strInput36, strInput37,
                           strInput38, strInput39, strInput40,
                           strInput41, strInput42, strInput43, strInput44, strInput45, strInput46, strInput47,
                           strInput48, strInput49, strInput50,
                           strOutput01, strOutput02, strOutput03, strOutput04, strOutput05, strOutput06, strOutput07, strOutput08, strOutput09, strOutput10):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'insert into data_sample_io_define(sampleId, timeFrom, timeTo, input01, input02, input03, input04, input05, input06,input07,input08, input09, input10,' \
            'input11, input12, input13, input14, input15, input16,input17,input18, input19, input20,' \
            ' input21, input22, input23, input24, input25, input26,input27,input28, input29, input30,' \
            ' input31, input32, input33, input34, input35, input36,input37,input38, input39, input40,' \
            ' input41, input42, input43, input44, input45, input46,input47,input48, input49, input50,' \
            ' output01, output02,output03, output04, output05, output06, output07, output08,output09,output10) values' \
            '(%s, %s, %s,    %s, %s, %s, %s, %s,%s, %s, %s, %s, %s,   %s, %s, %s, %s, %s,%s, %s, %s, %s, %s,  %s, %s, %s, %s, %s,%s, %s, %s, %s, %s,  %s, %s, %s, %s, %s,%s, %s, %s, %s, %s,  %s, %s, %s, %s, %s,%s, %s, %s, %s, %s,  %s, %s, %s, %s, %s,%s, %s, %s, %s, %s) '

        params = (str(sampleId), strStartTime,strEndTime, strInput01,strInput02,strInput03,strInput04,strInput05,strInput06,strInput07,strInput08,strInput09,strInput10,strInput11, strInput12, strInput13, strInput14, strInput15, strInput16, strInput17,strInput18, strInput19, strInput20,
                           strInput21, strInput22, strInput23, strInput24, strInput25, strInput26, strInput27, strInput28, strInput29, strInput30,
                           strInput31, strInput32, strInput33, strInput34, strInput35, strInput36, strInput37,
                           strInput38, strInput39, strInput40,  strInput41, strInput42, strInput43, strInput44, strInput45, strInput46, strInput47,
                           strInput48, strInput49, strInput50,strOutput01, strOutput02, strOutput03, strOutput04, strOutput05,  strOutput06, strOutput07, strOutput08, strOutput09, strOutput10)
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, q, params)
        return bSuccess

    def getMaxIdInTable(self, tableName):
        sql = 'select max(id) from %s'%(tableName)
        result = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
        nMaxId = 0
        try:
            if result and len(result) > 0:
                if len(result[0]) > 0:
                    nMaxId = int(result[0][0])
        except:
            nMaxId = 0
        return  nMaxId

    def getSampleList(self, ):
        dbname = app.config['DATABASE']
        bSuccess = True
        nId = -1
        allSampleList = []
        q = 'select id,time,author,name from data_sample_list'
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result and isinstance(result, list) and len(result) > 0:
            for item in result:
                allSampleList.append(dict(id=item[0], time=item[1], author=item[2], name=item[3]))

        return allSampleList

    def getAllDataOfSample(self, nSamleId):
        dbname = app.config['DATABASE']

        allSampleDataList = []
        q = 'select * from data_sample_io_define where sampleId = %d'%(int(nSamleId))
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result and isinstance(result, list) and len(result) > 0:
            for item in result:
                allSampleDataList.append(item)

        return allSampleDataList

    def getOrCreateSampleIdByName(self, strSampleName):
        dbname = app.config['DATABASE']
        bSuccess = True
        nId = -1

        q = 'select id,time,author from data_sample_list where name= "%s"' % (strSampleName)
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result and isinstance(result, list) and len(result)>0:
            nId = result[0][0]
        else:
            nId = self.getMaxIdInTable('data_sample_list')+1
            q = 'insert into data_sample_list(id, name, time, author) values(%s, %s, %s, %s)'
            params = ( str(nId), strSampleName, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), '')

            bSuccess = self._mysqlDbContainer.op_db_update(dbname, q, params)
        return nId

    def removeSampleData(self, nSampleId, strTimeFrom, strTimeTo):
        dbname = app.config['DATABASE']

        allSampleDataList = []
        if strTimeFrom is None and strTimeTo is None:
            q = 'delete from data_sample_io_define where sampleId = %d' % (int(nSampleId))
        else:
            q = 'delete from data_sample_io_define where sampleId = %d and timeFrom>= \'%s\' and timeTo<=\'%s\'' % (int(nSampleId), strTimeFrom, strTimeTo)
        return self._mysqlDbContainer.op_db_update(dbname, q)

    '''
    根据FddNameList数组，返回同样大小的数组，每个元素是每个FddName对应的诊断内容结果状态
    '''
    def getFddRealtimeStatusMul(self, fddNameList,fddTimeFrom):


        tTime = datetime_.date(int(fddTimeFrom[0:4]), int(fddTimeFrom[5:7]), int(fddTimeFrom[8:10]))

        self.checkCreateIfNotExistFDDTable(tTime)

        dbname = app.config['DATABASE']
        bSuccess = True
        nId = -1
        strTableName = 'fdd_result_' + time.strftime('%Y_%m',time.strptime(fddTimeFrom, '%Y-%m-%d %H:%M:%S'))

        q = 'select faultStatus, grade, content, analysis, suggestion,risk, day_history, fddname from ' + strTableName + ' where fddname in (%s) and fddtime= DATE_FORMAT("%s", "%%Y-%%m-%%d %%H:%%i:%%S")' % \
            (str(fddNameList).replace('[', '').replace(']', ''), fddTimeFrom)
        params = ()
        result = self._mysqlDbContainer.op_db_query(dbname, q, params)
        if result is None:
            return None

        allInfo = {}
        for oneRecord in result:
            try:
                strHis = oneRecord[6]
                nFaultHistory = json.loads(strHis)
                statusInfo = dict(fddtime= fddTimeFrom, faultStatus= oneRecord[0],
                                  grade = oneRecord[1], content = oneRecord[2], analysis = oneRecord[3], suggestion = oneRecord[4], risk= oneRecord[5],
                                  history = nFaultHistory)

                strOneFddName = oneRecord[7]
                allInfo[strOneFddName] = statusInfo
            except:
                print('ERROR in getFddItemHistory')

        return allInfo

    def getFddRealtimeStatus(self, fddName):
        statusInfo = {}
        tTime = datetime.now()
        self.checkCreateIfNotExistFDDTable(tTime)

        dbname = app.config['DATABASE']
        bSuccess = True
        nId = -1
        strTableName = 'fdd_result_' + tTime.strftime('%Y_%m')
        q = 'select faultStatus, grade, content, analysis, suggestion,risk, day_history from ' + strTableName + ' where fddname= "%s" and fddtime= DATE_FORMAT("%s","%%Y-%%m-%%d %%H:%%i:%%S")'%( fddName, tTime.strftime('%Y-%m-%d 00:00:00').replace('11','09'))
        params = ()
        result = self._mysqlDbContainer.op_db_query(dbname, q, params)

        if result:
            try:
                strHis = result[0][6]
                nFaultHistory = json.loads(strHis)
                statusInfo = dict(fddtime= tTime.strftime('%Y-%m-%d 00:00:00'), faultStatus= result[0][0], grade = result[0][1], content = result[0][2], analysis = result[0][3], suggestion = result[0][4], risk=result[0][5],
                                  history = nFaultHistory)
            except:
                print('ERROR in getFddItemHistory')

        return statusInfo


    def getOneFddStatus(self, strFddName):
        fddDefine = self.getFddItemInfo(strFddName)
        fddDefine["fddInfo"] = self.getFddRealtimeStatus(strFddName)
        return fddDefine
        return dict(name='FDDTempSupplyPointWrong',name_ch='皮带松动故障诊断',description='故障描述', ofEquipment='AHU-1-2', ofZone='老车身车间', ofResponseParty='维修班组', ofFaultClassify='设备老化导致的故障类',
               enabled=1, updatetime='2018-09-12 23:00:00', fddInfo = dict(fddtime='2018-09-12 23:00:00', faultStatus=1, grade=3, content='推测皮带松动',
                                                                           analysis='送风温度为25度，风压偏低，水阀全开，仍然供冷不足', suggestion='建议检查皮带', risk='2500', history=
                                                                           [0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0]))
    # def getMulFddRealtimeStatus(self, fddNameList,fddTimeFrom):
    #
    #
    #     tTime = datetime_.date(int(fddTimeFrom[0:4]), int(fddTimeFrom[5:7]), int(fddTimeFrom[8:10]))
    #     print(tTime.strftime('%Y-%m'))
    #     self.checkCreateIfNotExistFDDTable(tTime)
    #
    #     dbname = app.config['DATABASE']
    #     bSuccess = True
    #     nId = -1
    #     rv = []
    #     strTableName = 'fdd_result_' + tTime.strftime('%Y_%m')
    #     q = 'select faultStatus, grade, content, analysis, suggestion,risk, day_history from ' + strTableName + ' where fddname in %s and fddtime= %s'
    #     params = (fddNameList, tTime.strftime('%Y-%m-%d 00:00:00'))
    #     resultlist = self._mysqlDbContainer.op_db_query(dbname, q, params)
    #     if resultlist:
    #         try:
    #             for i in range(len(resultlist)):
    #                 result = resultlist[i]
    #                 strHis = result[6]
    #                 nFaultHistory = json.loads(strHis)
    #                 statusInfo = dict(fddtime= tTime.strftime('%Y-%m-%d 00:00:00'), faultStatus= result[0], grade = result[1], content = result[2], analysis = result[3], suggestion = result[4], risk=result[5],
    #                               history = nFaultHistory)
    #                 rv.append(statusInfo)
    #         except:
    #             print('ERROR in getFddItemHistory')
    #
    #     return rv

    def plot(self,history_data):
        file_name = ('siteinterface\\static\\files\\temp\\history%s.png' % datetime.now()).replace(':','_').replace('-','_').replace(' ','_')
        xlist = [i for i in range(len(history_data))]
        try:
            x = np.array(xlist)
            y = np.array(history_data)
            plt.xlim((0, len(history_data)))
            plt.ylim((0, 1.1))
            plt.bar(x, y, color='r', width=0.2, edgecolor='r')
            plt.yticks([])
            plt.xticks([])
            plt.subplots_adjust(top=1, bottom=0, right=1, left=0, hspace=0, wspace=0)
            plt.margins(0, 0)
            plt.rcParams['savefig.dpi'] = 100

            plt.savefig(file_name)
            plt.clf()
        except Exception as e :
            print(e.__str__())
            print('Error in generate plot')
            return False
        return file_name

    def writeDocx(self,data,fddTimeFrom,fddTimeTo):

        file_name = 'siteinterface\\static\\files\\temp\\Fdd'+time.strftime('%Y%m%d%H%M%S')+'.docx'
        # file_name_zip = 'siteinterface\\static\\files\\temp\\报表.zip'
        if os.path.exists(file_name):
            os.remove(file_name)

        doc = docx.Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal'].font.size = 15
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(u'故障诊断报表')
        run.font.size = Pt(20)

        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(u'(从%s至%s)' % (fddTimeFrom, fddTimeTo))
        run.font.size = Pt(10)

        table_menu = ['故障名称', '最近状态', '故障持续时间', '故障趋势']
        table = doc.add_table(rows=len(data)+1, cols=4)
        table.columns[2].width = Inches(1.2)
        table.columns[3].width = Inches(2.5)
        for i in range(len(table_menu)):
            run = table.cell(0, i).paragraphs[0].add_run(table_menu[i])
            run.font.name = u'宋体'
            run.font.size = 140000
            table.cell(0, i).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        row = 1
        for one_day in data:
            infolist=[]
            infolist.append(one_day['fddname'])
            infolist.append(one_day['content'])
            infolist.append('%.2f h' % (sum(one_day['history']) / 4.0))
            # errorInfo = one_day['content']
            # errorTime = '%.2f h'%(sum(one_day['history'])/4.0)
            # errorhistory = one_day['sidtory']
            for col in range(len(infolist)):
                run = table.cell(row, col).paragraphs[0].add_run(infolist[col])
                run.font.name = u'宋体'
                run.font.size = 140000
                table.cell(row, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            history_plot_path = self.plot(one_day['history'])
            run = table.cell(row, 3).paragraphs[0].add_run()
            try:
                if history_plot_path != False:
                    run.add_picture(history_plot_path, width=2200000, height=800000)
                    table.cell(row, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    os.remove(history_plot_path)
            except Exception as e:
                print(str(e))
                print('Error in get hitory plot')
                return False
            row += 1
        doc.save(file_name)
        # azip = zipfile.ZipFile(file_name_zip, 'w')
        # azip.write(file_name,arcname=os.path.basename(file_name))
        # azip.close()
        return file_name

    def getMulFddItemInfo(self, fddNameGroupMap,fddNameList):
        dbname = app.config['DATABASE']
        bSuccess = True
        fddResult = []
        data = []
        q = 'select name, name_ch, description, ofEquipment, ofZone, ofResponseParty, ofFaultClassify, enabled, updatetime from fdd_list_item where name in (%s)'%(str(fddNameList).replace('[', '').replace(']', ''))
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result and isinstance(result, list) and len(result) > 0:
            groupname_old = fddNameGroupMap.get(result[0][0],None)
        if result and isinstance(result, list) and len(result) > 0:
            for item in result:
                groupname = fddNameGroupMap.get(item[0],None)
                if groupname != groupname_old:
                    data.append(dict(fddResult=fddResult))
                    groupname_old = groupname
                    fddResult = []
                fddResult.append(
                    dict(name=item[0], name_ch=item[1], description=item[2], ofEquipment=item[3], ofZone=item[4],
                         ofResponseParty=item[5], ofFaultClassify=item[6],
                         enabled=item[7], updatetime=item[8]))
            if fddResult:
                data.append(dict(fddResult=fddResult))
        return data

    def getMulFddStatusPeriod(self,fddNameList, fddTimeFrom,fddTimeTo):
        tTimeFrom = datetime_.date(int(fddTimeFrom[0:4]), int(fddTimeFrom[5:7]), int(fddTimeFrom[8:10]))
        self.checkCreateIfNotExistFDDTable(tTimeFrom)

        dbname = app.config['DATABASE']
        bSuccess = True
        nId = -1
        strTableName = 'fdd_result_' + time.strftime('%Y_%m', time.strptime(fddTimeFrom, '%Y-%m-%d %H:%M:%S'))

        q = 'select faultStatus, grade, content, analysis, suggestion,risk, day_history, fddname, DATE_FORMAT(fddtime, "%Y-%m-%d %H:%i:%S") as fddtime from ' + strTableName + ' where fddname in (%s) and DATE_FORMAT(fddtime, "%%Y-%%m-%%d %%H:%%i:%%S") BETWEEN "%s" and "%s" ' \
                                                                                                                         'and day_history != "[0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]"' % \
            (str(fddNameList).replace('[', '').replace(']', ''), fddTimeFrom, fddTimeTo)
        params = ()
        try:
            result = self._mysqlDbContainer.op_db_query(dbname, q, params)
            if result is None:
                return None
        except:
            print('Input Time Error')
            return False
        allInfo = []
        for oneRecord in result:
            try:
                strHis = oneRecord[6]
                nFaultHistory = json.loads(strHis)
                statusInfo = dict(fddtime=oneRecord[8],faultStatus=oneRecord[0],
                                  grade=oneRecord[1], content=oneRecord[2], analysis=oneRecord[3],
                                  suggestion=oneRecord[4], risk=oneRecord[5],
                                  history=nFaultHistory,fddname=oneRecord[7])

                allInfo.append(statusInfo)
            except:
                print('ERROR in getFddItemHistory')
                return False
        return allInfo

    def getMulFddStatus(self,fddNameGroupMap,strFddNameList,fddTimeFrom):
        data = self.getMulFddItemInfo(fddNameGroupMap,strFddNameList)
        FddInfoList = self.getFddRealtimeStatusMul(strFddNameList,fddTimeFrom)
        if FddInfoList is None:
            return None
        fddIndoDefault = dict(fddtime= fddTimeFrom, faultStatus= 0,
                                  grade = 5, content = '无', analysis = '无', suggestion = '无', risk= '无',
                                  history = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0])
        for group in data:
            for item in group['fddResult']:
                item['fddInfo'] = FddInfoList.get(item['name'], fddIndoDefault)
            # if i>= len(FddRealtimeStatusList):
            #     print('ERROR in getMulFddStatus: FddRealtimeStatusList size< fddDefineList size')
            #     break


        return data
    # 获取一段时间的数据
    def getFddStatusPeriod(self,groupInfoList,fddTimeFrom,FddTimeTo):
        fddNameList = []

        for group in groupInfoList:
            groupName = group.get('groupName')
            for itemName in group.get('fddList'):
                fddNameList.append(groupName + itemName)

        data = self.getMulFddStatusPeriod(fddNameList, fddTimeFrom,FddTimeTo)
        return data


    def getFddStatus(self, groupInfoList,fddTimeFrom):
        testDataList = []
        fddNameList = []

        i = 0
        fddNameGroupMap = {}
        for group in groupInfoList:

            groupName = group.get('groupName')
            for itemName in group.get('fddList'):
                fddNameList.append(groupName+itemName)
                fddNameGroupMap[groupName+itemName] = groupName

        data = self.getMulFddStatus(fddNameGroupMap, fddNameList,fddTimeFrom)

        if data is None:
            return None

        for group in data:
            strFDDName = group['fddResult'][0]['name']
            group['groupName'] = fddNameGroupMap.get(strFDDName, None)
            if group['groupName'] is None:
                continue

            nTotalFault = 0
            fddResult = group['fddResult']
            for i in fddResult:
                if i and i.get('fddInfo') and i.get('fddInfo').get('faultStatus') == 1:
                     nTotalFault+=1
            group['groupFddFaultCount'] = nTotalFault

            # 一次获取一个数据以下代码
            # nTotalFault = 0
            #
            # for fdd in fddNameList:
            #     oneResult = self.getOneFddStatus(fdd)
            #     groupResult['fddResult'].append(oneResult)
            #     if oneResult and oneResult.get('fddInfo') and oneResult.get('fddInfo').get('faultStatus')==1:
            #         nTotalFault+=1
            #
            # groupResult['groupFddFaultCount'] = nTotalFault
            # testDataList.append(groupResult)
        return data

        testDataList = [dict(groupName='AHU-1', groupFddFaultCount=3, fddResult= [
            dict(name='FDDTempSupplyPointWrong',name_ch='皮带松动故障诊断',description='故障描述', ofEquipment='AHU-1-2', ofZone='老车身车间', ofResponseParty='维修班组', ofFaultClassify='设备老化导致的故障类',
               enabled=1, updatetime='2018-09-12 23:00:00', fddInfo = dict(fddtime='2018-09-12 23:00:00', faultStatus=1, grade=3, content='推测皮带松动',
                                                                           analysis='送风温度为25度，风压偏低，水阀全开，仍然供冷不足', suggestion='建议检查皮带', risk='2500', history=
                                                                           [0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,1,0,0,0,0]))
        ]),
                        dict(groupName='AHU-3', groupFddFaultCount=2   , fddResult=[
                            dict(name='FDDTempSupplyPointWrong', name_ch='皮带松动故障诊断', description='故障描述',
                                 ofEquipment='AHU-1-2', ofZone='老车身车间', ofResponseParty='维修班组',
                                 ofFaultClassify='设备老化导致的故障类',
                                 enabled=1, updatetime='2018-09-12 23:00:00',
                                 fddInfo=dict(fddtime='2018-09-12 23:00:00', faultStatus=1, grade=3, content='推测皮带松动',
                                              analysis='送风温度为25度，风压偏低，水阀全开，仍然供冷不足', suggestion='建议检查皮带', risk='2500',
                                              history=
                                              [0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1,
                                               1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0,
                                               0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1,
                                               1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0]))
                        ]),
                        dict(groupName='AHU-2', groupFddFaultCount=4, fddResult=[
                            dict(name='FDDTempSupplyPointWrong', name_ch='皮带松动故障诊断', description='故障描述',
                                 ofEquipment='AHU-2-2', ofZone='老车身车间', ofResponseParty='维修班组',
                                 ofFaultClassify='设备老化导致的故障类',
                                 enabled=1, updatetime='2018-09-12 23:00:00',
                                 fddInfo=dict(fddtime='2018-09-12 23:00:00', faultStatus=1, grade=3, content='推测皮带松动',
                                              analysis='送风温度为25度，风压偏低，水阀全开，仍然供冷不足', suggestion='建议检查皮带', risk='2500',
                                              history=
                                              [0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1,
                                               1, 1, 1, 1, 0, 0, 0, 0, 1, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0,
                                               0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1,
                                               1, 1, 1, 1, 0, 1, 0, 0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0])),
                            dict(name='FDDTempSupplyPointWrong', name_ch='皮带松动故障诊断', description='故障描述',
                                 ofEquipment='AHU-2-3', ofZone='老车身车间', ofResponseParty='维修班组',
                                 ofFaultClassify='设备老化导致的故障类',
                                 enabled=1, updatetime='2018-09-12 23:00:00',
                                 fddInfo=dict(fddtime='2018-09-12 23:00:00', faultStatus=0, grade=3, content='推测皮带松动',
                                              analysis='送风温度为25度，风压偏低，水阀全开，仍然供冷不足', suggestion='建议检查皮带', risk='2500',
                                              history=
                                              [0, 0, 0, 1, 1, 0, 0, 1, 1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1,
                                               1, 1, 1, 1, 1, 1, 1, 0, 1, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0,
                                               0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1,
                                               1, 1, 1, 1, 0, 1, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 0, 0, 0, 0]))
                        ]),
                        ]
        return testDataList

    def getAllFddItems(self):
        dbname = app.config['DATABASE']
        bSuccess = True
        rv = []

        q = 'select name, name_ch, description, ofEquipment, ofZone, ofResponseParty, ofFaultClassify, enabled, updatetime from fdd_list_item'
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result and isinstance(result, list) and len(result) > 0:
            for item in result:
                rv.append(dict(name=item[0], name_ch = item[1], description=item[2], ofEquipment=item[3], ofZone=item[4], ofResponseParty=item[5], ofFaultClassify=item[6],
                               enabled=item[7], updatetime=item[8]))
        return rv

    def getFddItemInfo(self, fddName):
        dbname = app.config['DATABASE']
        bSuccess = True
        rv = {}

        q = 'select name, name_ch, description, ofEquipment, ofZone, ofResponseParty, ofFaultClassify, enabled, updatetime, createtime from fdd_list_item where name= "%s"'%(fddName)
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result and isinstance(result, list) and len(result) > 0:
            item =  result[0]
            return dict(name=item[0], name_ch=item[1], description=item[2], ofEquipment=item[3], ofZone=item[4],
                               ofResponseParty=item[5], ofFaultClassify=item[6],
                               enabled=item[7], updatetime=item[8].strftime("%Y-%m-%d %H:%M:%S"), createtime=item[9].strftime("%Y-%m-%d %H:%M:%S"))
        return {}


    def getOrCreateFddItem(self, strName, strNameCh, strDescription, strOfEquipment, strOfZone, strOfResponseParty, strOfFaultClassify):
        dbname = app.config['DATABASE']
        bSuccess = True

        q = 'select name from fdd_list_item where name= "%s"' % (strName)
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result and isinstance(result, list) and len(result)>0:
            return bSuccess
        else:
            q = 'insert into fdd_list_item(name, name_ch, description, ofEquipment, ofZone, ofResponseParty, ofFaultClassify) values(%s, %s, %s, %s, %s, %s, %s)'
            params = ( strName, strNameCh, strDescription, strOfEquipment, strOfZone, strOfResponseParty, strOfFaultClassify)

            bSuccess = self._mysqlDbContainer.op_db_update(dbname, q, params)
        return bSuccess

    def getFddItemHistory(self, strFddName, strTime):
        statusList = []
        tTime = datetime.strptime(strTime, '%Y-%m-%d %H:%M:%S')
        self.checkCreateIfNotExistFDDTable(tTime)

        strTime = tTime.strftime('%Y-%m-%d 00:00:00')

        dbname = app.config['DATABASE']
        bSuccess = True
        nId = -1
        strTableName = 'fdd_result_' + tTime.strftime('%Y_%m')
        q = 'select day_history from ' + strTableName + ' where fddname= %s and fddtime= %s'
        params = (strFddName, strTime)
        result = self._mysqlDbContainer.op_db_query(dbname, q, params)

        if result:
            try:
                strHis = result[0][0]
                statusList = json.loads(strHis)
            except:
                print('ERROR in getFddItemHistory')

        if len(statusList)<24*4:
            statusList = [0]*24*4
        return statusList




    def setFddFaultStatus(self,strName, strTime,  nFaultStatus, nGrade, strContent, strAnalysis, strSuggestion, strRisk):
        tTime= datetime.strptime(strTime,  '%Y-%m-%d %H:%M:%S')
        self.checkCreateIfNotExistFDDTable(tTime)

        dayHistoryList = self.getFddItemHistory(strName, strTime)

        nIndex = tTime.hour*4+ int(tTime.minute/15)
        dayHistoryList[nIndex] = nFaultStatus
        strDayHistory = json.dumps(dayHistoryList)

        dbname = app.config['DATABASE']
        bSuccess = True
        nId = -1
        strTableName = 'fdd_result_'+ tTime.strftime('%Y_%m')
        strTimeOnlyDate = tTime.strftime('%Y-%m-%d 00:00:00')
        q = 'delete from '+strTableName +' where fddname= %s and fddtime= %s'
        params=(strName, strTimeOnlyDate)
        result = self._mysqlDbContainer.op_db_update(dbname, q, params)

        q = 'insert into '+ strTableName +' (fddname, fddtime, faultStatus, grade, content, analysis, suggestion, risk, day_history) values(%s, %s, %s, %s, %s, %s, %s, %s, %s)'
        params = (strName, strTimeOnlyDate,  nFaultStatus, nGrade, strContent, strAnalysis, strSuggestion, strRisk, strDayHistory)

        bSuccess = self._mysqlDbContainer.op_db_update(dbname, q, params)
        return bSuccess

    def setDataErr(self, pointList, strTime, strErrorGrade, strReason = ''):

        bSuccess = True
        for point in pointList:
            bSuccess = bSuccess and self.setDataErrOnePoint(point, strTime, strErrorGrade, strReason)

        return bSuccess


    def setDataErrOnePoint(self, pointName, strTime, strErrorGrade, strReason):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'select id, pointname, timeFrom, timeTo from fdd_data_error where pointname= "%s" and timeFrom<= "%s" and timeTo>= "%s"' % ( pointName, strTime, strTime)
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result is None:
            return False

        tTime = datetime.strptime(strTime, '%Y-%m-%d %H:%M:%S')
        tUpdateTime = tTime+ timedelta(minutes = 5) #如果数据有问题，那么五分钟内都会作废
        strUpdateTime = tUpdateTime.strftime( '%Y-%m-%d %H:%M:%S')

        if result and isinstance(result, list) and len(result) > 0:
            idCur = int(result[0][0])
            q = 'update fdd_data_error set timeTo = "%s" where id = %d' %(strUpdateTime, idCur)
            bSuccess = self._mysqlDbContainer.op_db_update(dbname, q, ())
        else:
            newId = self.getMaxIdInTable('fdd_data_error')+1
            tStartTime = tTime  - timedelta(minutes = 5) #如果数据有问题，那么五分钟前都会作废
            strNewStartTime = tStartTime.strftime( '%Y-%m-%d %H:%M:%S')
            q = 'insert into fdd_data_error(id, pointname, timeFrom, timeTo, errorgrade, reason) values(%s, %s, %s, %s, %s, %s)'
            params = (  newId, pointName, strNewStartTime, strUpdateTime, strErrorGrade, strReason)

            bSuccess = self._mysqlDbContainer.op_db_update(dbname, q, params)

        return bSuccess

    def getDataErr(self, pointNameList, strFddTime):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'select id, pointname, timeFrom, timeTo, errorgrade, reason from fdd_data_error where pointname in (%s) and timeFrom<= "%s" and timeTo>= "%s"' % (str(pointNameList).replace('[', '').replace(']', ''), strFddTime, strFddTime)

        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result is None :
            return None

        rv = {}
        if isinstance(result, list) and len(result) > 0:
            for item in result:
                ptName = item[1]
                rv[ptName] = dict(error=1, grade = item[4], reason = item[5], timefrom = item[2], timeTo = item[3])

        return rv

    def getAndClearModbusEquipmentOutputTable(self, pointList):

        strSQL ='select pointname, pointvalue from realtimedata_output_modbus_equipment'
        rv = self._mysqlDbContainerModbusClientPool.op_db_query(app.config['DATABASE'], strSQL)
        if rv is None:
            return []
        retReal = []
        #print(str(pointList))
        for item in rv:
            if pointList:
                if item[0] in pointList:
                    retReal.append(dict(name=item[0], value=item[1]))
            else:
                retReal.append(dict(name=item[0], value=item[1]))
        if len(retReal)>0:
            self.delFromModbusEquipmentOutputTable(retReal)
        return retReal

    def delFromModbusEquipmentOutputTable(self, retReal):
        dbname = app.config['DATABASE']
        q = 'delete from realtimedata_output_modbus_equipment where pointname = %s and pointvalue =  %s'
        alldata = []
        nIndex = 0
        for item in retReal:
            alldata.append(tuple([item.get('name'), item.get('value')]))
            nIndex += 1
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainerModbusClientPool.op_db_update_many(dbname, q, params)
        return bSuccess

    def clearModbusEquipmentOutputTable(self):
        dbname = app.config['DATABASE']
        q = 'delete from realtimedata_output_modbus_equipment'
        alldata = []
        nIndex = 0

        bSuccess = self._mysqlDbContainerModbusClientPool.op_db_update(dbname, q, )
        return bSuccess

    def getModbusEquipmentInputTable(self, pointList):
        strSQL ='select pointname, pointvalue from realtimedata_input_modbus_equipment order by pointname'
        rv = self._mysqlDbContainerModbusClientPool.op_db_query(app.config['DATABASE'], strSQL)
        if rv is None:
            return None
        retReal = []
        #print(str(pointList))
        if pointList is not None:
            for item in rv:
                if item[0]=='ChOnOff01':
                    a=1
                if (item[0] in pointList):
                    #print('%s : %s'%(str(item[0]), str(item[1])))
                    retReal.append(dict(name=item[0], value=item[1]))
        else:
            for item in rv:
                retReal.append(dict(name=item[0], value=item[1]))
        return retReal

    def isDBConnectionGood(self):
        dbname = app.config['DATABASE']
        nRetryCount = 0
        bTestResult = self._mysqlDbContainer.test_db_connection_is_living(dbname)
        if bTestResult is None:
            bTestResult = False
        while not bTestResult and nRetryCount<=5:
            time.sleep(10)
            bTestResult = self._mysqlDbContainer.test_db_connection_is_living(dbname)
            if bTestResult is None:
                bTestResult = False
            nRetryCount+=1

        if not bTestResult:
            return False

        return True

    def updateProcessStatus(self, processName):
        dbname = app.config['DATABASE']

        q = 'select id,timeFrom, timeTo,process from core_status where process = %s order by timeFrom DESC'
        param = (processName, )

        result = self._mysqlDbContainer.op_db_query(dbname, q, param)
        if result is None:
            return False

        tNow = datetime.now()
        if not result:
            nNewID = self.getMaxIdInTable('core_status') +1
            q = 'insert into core_status(id, timeFrom, timeTo, process, status) values(%s, %s, %s, %s, 1)'
            param = (nNewID, tNow.strftime('%Y-%m-%d %H:%M:%S'), tNow.strftime('%Y-%m-%d %H:%M:%S'), processName)
            return self._mysqlDbContainer.op_db_update(dbname, q, param)
        else:
            try:
                if isinstance(result, list) and len(result) > 1:
                    sqlDel ='delete from core_status where process = %s and timeFrom< %s'
                    sqlDelParams = (processName, result[0][1].strftime('%Y-%m-%d %H:%M:%S'))
                    self._mysqlDbContainer.op_db_update(dbname, sqlDel, sqlDelParams)
            except:
                pass
            q = 'update core_status set timeTo = %s where process = %s'
            param = (tNow.strftime('%Y-%m-%d %H:%M:%S'), processName)
            return self._mysqlDbContainer.op_db_update(dbname, q, param)

    def getProcessStatus(self, processName, nOnlineTimeoutMinutes=5):
        nStatus = 0
        dbname = app.config['DATABASE']

        q = 'select id,timeFrom, timeTo,process from core_status where process = %s and status =1 order by timeTo desc '
        param = (processName,)

        result = self._mysqlDbContainer.op_db_query(dbname, q, param)

        if result is None:
            return -1

        tNow = datetime.now()
        if not result:
            nStatus =  -1
        else :
            tTimeTo = result[0][2]
            tspan =  datetime.now()-tTimeTo
            if tspan.total_seconds()>nOnlineTimeoutMinutes*60:
                nStatus =  1
            else:
                nStatus =  0
        return nStatus

    def getHistoryDataStatus(self):
        nStatus = 0
        dbname = app.config['DATABASE']
        tNow = datetime.now()
        table = 'historydata_minute_' + tNow.strftime('%Y_%m_%d')
        q = 'select count(1) from information_schema.tables where table_name = %s';
        param = (table,)
        result = self._mysqlDbContainer.op_db_query(dbname,q,param)
        if result is None:
            nStatus = -1
        q = 'select count(1) from '+ table + ' where time = %s'
        param= (tNow.strftime('%Y-%m-%d 00:00:00'),)
        result =  self._mysqlDbContainer.op_db_query(dbname,q,param)
        if result is None:
            return -1

        if len(result)==0 :
            nStatus = 1
        else:
            nStatus = 0
        return nStatus

    def getCoreStatus(self):
        dbname = app.config['DATABASE']
        bSuccess = True
        q = 'select * from unit01 where unitproperty01 = "LogicStatusBegin" or unitproperty01 = "LogicStatusEnd" '

        result = self._mysqlDbContainer.op_db_query(dbname, q)
        if result is None:
            return None

        rv = {}
        if isinstance(result, list) and len(result) > 0:
            for item in result:
                strDllName =  item[1]

                if not rv.get(strDllName):
                    rv[strDllName] = dict(timeBegin="", timeEnd="", timeCostSeconds=-1, actBeginId=-1, actEndId=-1)
                iter = item[0]
                if iter=='LogicStatusBegin':
                    #rv[strDllName]['timeBegin'] = item[2].strftime('%Y-%m-%d %H:%M:%S')
                    rv[strDllName]['timeBegin'] = item[2]
                    rv[strDllName]['actBeginId'] = item[3]

                elif iter == 'LogicStatusEnd':
                    rv[strDllName]['timeEnd'] = item[2]
                    rv[strDllName]['actEndId'] = item[3]
                    rv[strDllName]['timeCostSeconds'] = item[4]

        return rv
    def getimportHistoryDataProfess(self):
        dbname = app.config['DATABASE']
        bSuccess = True
        #q = 'select unitproperty03 from unit01  where unitproperty01 = "import_history_data_progress" and unitproperty02= "%s"' % (filename)
        q= 'select unitproperty02,unitproperty03 from unit01  where unitproperty01 = "import_history_data_progress"'
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        rv={}
        if result is None:
            return None

        if isinstance(result, list) and len(result) > 0:
            for item in result:
                filename= item[0]
                rv[filename] = dict(ProcessID="")
                if(item[1] is None):
                   item[1] = '0'
                rv[filename]['ProcessID'] = item[1]

        return rv

    def addLogicOutputRecord(self, strPointTime, strPointName, strPointValue, strLogicName):
        if isinstance(strPointName, list) and isinstance(strPointValue, list):
            strPointNameList = strPointName
            strPointValueList = strPointValue
        else:
            strPointNameList = [strPointName]
            strPointValueList = [strPointValue]
        return self.addLogicOutputRecordMul(strPointTime, strPointNameList, strPointValueList, strLogicName)


    def addLogicOutputRecordMul(self, strPointTime, strPointNameList, strPointValueList, source="Unknown"):
        if not isinstance(source, str):
            source = "Unknown"

        if not len(source):
            source = "Unknown"

        result = {'data': None, 'msg': '', 'status': None}
        coreDir = app.config["CORE_PATH"]
        tnow = datetime.now()
        logDir = os.path.join(coreDir, 'log')
        if not os.path.exists(logDir):
            os.mkdir(logDir)

        opRecordDir = os.path.join(logDir, "oprecord-"+ tnow.strftime("%Y-%m-%d"))
        if not os.path.exists(opRecordDir):
            os.mkdir(opRecordDir)

        opRecordFilePath = os.path.join(opRecordDir, "Source_{source}.txt".format(source=source))
        with open(opRecordFilePath, "a", encoding="UTF8", errors="ignore") as fo:
            for nIdx, pointName in enumerate(strPointNameList):
                try:
                    fo.write("\n%s\t%s\t%s\n" % (strPointTime, pointName, strPointValueList[nIdx]))
                except:
                    pass

        result['status'] = True
        return result



    def logic_search(self, strThreadName, strPointName):
        dbname = app.config['DATABASE']
        bSuccess = True
        # q = 'select unitproperty03 from unit01  where unitproperty01 = "import_history_data_progress" and unitproperty02= "%s"' % (filename)
        if strThreadName =='' and strPointName=='':
            q = 'select logicname,pointtime, pointname, pointvalue from logic_output_point_record'
        elif strPointName !='':
            q = 'select logicname,pointtime, pointname, pointvalue from logic_output_point_record  where pointname = "%s"'%(strPointName)
        elif strThreadName!='':
           q = 'select logicname,pointtime, pointname, pointvalue from logic_output_point_record  where logicname like "[Thread:%s]%%"'%(strThreadName)
        result = self._mysqlDbContainer.op_db_query(dbname, q)
        rv = {}
        if result is None:
            return None

        strAll = ''
        if isinstance(result, list) and len(result) > 0:
            for item in result:
                oneStr = 'LogicName: %s'%(item[0])
                oneStr += ' '
                oneStr+= 'PointName: %s'%(item[2])
                oneStr += ' '
                oneStr += 'Pointtime: %s'%(item[1])
                oneStr += ' '
                oneStr += 'Pointvalue: %s'%(item[3])
                oneStr += '\r\n'
                strAll +=oneStr

        return strAll

    def getReportCount(self, tFrom, tTo, strReportName):
        dbname = app.config['DATABASE']
        bSuccess = True
        # q = 'select unitproperty03 from unit01  where unitproperty01 = "import_history_data_progress" and unitproperty02= "%s"' % (filename)
        if strReportName is None or strReportName =="":
            q = 'select count(*) from report_history where reportTimeFrom>= %s and reportTimeTo<=%s'
            param = (tFrom.strftime('%Y-%m-%d %H:%M:%S'), tTo.strftime('%Y-%m-%d %H:%M:%S'))
        else:
            q = 'select count(*) from report_history where reportTimeFrom>= %s and reportTimeTo<=%s and description = %s'
            param = (tFrom.strftime('%Y-%m-%d %H:%M:%S'), tTo.strftime('%Y-%m-%d %H:%M:%S'), strReportName)
        result = self._mysqlDbContainer.op_db_query(dbname, q, param)
        rv = {}
        if not result:
            return 0

        return result[0][0]

    def getReportNameList(self,):
        dbname = app.config['DATABASE']
        bSuccess = True
        # q = 'select unitproperty03 from unit01  where unitproperty01 = "import_history_data_progress" and unitproperty02= "%s"' % (filename)
        q = 'select DISTINCT description from report_history'
        param = ()
        result = self._mysqlDbContainer.op_db_query(dbname, q, param)

        if result is None:
            return []
        rv = []
        if isinstance(result, list) and len(result) > 0:
            for item in result:
                strNN = item[0]
                strNN = strNN.strip()
                if strNN:
                    rv.append(strNN)

        return rv

    def deleteReport(self, reportId):
        try:
            strQuery = "DELETE FROM report_history WHERE id = %s"
            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, parameter=(reportId,))
            return bSuc
        except Exception as e:
            logging.error("ERROR in deleteReport: %s" % e.__str__())
            return False

    def getReportHistory(self, tFrom, tTo, nPageNum, nPageSize, strReportName = None):
        dbname = app.config['DATABASE']
        bSuccess = True
        if nPageNum<=0:
            nPageNum = 1
        # q = 'select unitproperty03 from unit01  where unitproperty01 = "import_history_data_progress" and unitproperty02= "%s"' % (filename)
        if strReportName is None or strReportName =="":
            # q = 'select id,name, description, gentime, filesize,author,url, reportTimeFrom, reportTimeTo, reportTimeType from report_history where reportTimeFrom>= %s and reportTimeTo<=%s' \
                # 'limit '+ str((nPageNum-1)*nPageSize)+ ',' + str(nPageSize)
            q = 'select id,name, description, gentime, filesize,author,url, reportTimeFrom, reportTimeTo, reportTimeType from report_history where gentime >= %s and gentime <= %s' \
                'limit ' + str((nPageNum - 1) * nPageSize) + ',' + str(nPageSize)
            param = (tFrom.strftime('%Y-%m-%d %H:%M:%S'), tTo.strftime('%Y-%m-%d %H:%M:%S'))
        else:
            # q = 'select id,name, description, gentime, filesize,author,url, reportTimeFrom, reportTimeTo, reportTimeType from report_history where reportTimeFrom>= %s and reportTimeTo<=%s and description = %s'\
                # 'limit ' + str((nPageNum - 1) * nPageSize) + ',' + str(nPageSize)
            q = 'select id,name, description, gentime, filesize,author,url, reportTimeFrom, reportTimeTo, reportTimeType from report_history where gentime >= %s and gentime <=%s and description = %s' \
                'limit ' + str((nPageNum - 1) * nPageSize) + ',' + str(nPageSize)
            param = (tFrom.strftime('%Y-%m-%d %H:%M:%S'), tTo.strftime('%Y-%m-%d %H:%M:%S'), strReportName)
        result = self._mysqlDbContainer.op_db_query(dbname, q,param)
        rv = {}
        if result is None:
            return None

        allReportList = []
        if isinstance(result, list) and len(result) > 0:
            for item in result:
                try:
                    try:
                        strGenTime = item[3].strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        strGenTime = ''

                    try:
                        strTimeFrom = item[7].strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        strTimeFrom = ''

                    try:
                        strTimeTo = item[8].strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        strTimeTo = ''

                    allReportList.append(dict(id= item[0],
                                              name=item[1],
                                              description = item[2],
                                              gentime = strGenTime,
                                              filesize = item[4],
                                              author= item[5],
                                              url=  item[6],
                                              reportTimeFrom = strTimeFrom,
                                              reportTimeTo = strTimeTo,
                                              reportTimeType = item[9]
                                              ))
                except:
                    pass

        return allReportList



    def insertReportHistory(self, strName, strDescription, strGenTime, nFileSize, strAuthor, strURL, strReportTimeFrom, strReportTimeTo, nReportTimeType):
        dbname = app.config['DATABASE']
        bSuccess = True

        nId = self.getMaxIdInTable('report_history')+1
        # q = 'select unitproperty03 from unit01  where unitproperty01 = "import_history_data_progress" and unitproperty02= "%s"' % (filename)
        q = 'insert into report_history(id,name, description, gentime, filesize,author,url, reportTimeFrom, reportTimeTo, reportTimeType) values(%s, %s,%s,%s,%s,%s,%s,%s,%s,%s)'
        param = (nId, strName, strDescription, strGenTime, nFileSize, strAuthor, strURL, strReportTimeFrom, strReportTimeTo, nReportTimeType)
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, q,param)

        if bSuccess:
            return nId
        else:
            return None

    def delete_log_table(self, strTableName):
        dbname = 'domlog'
        q = 'drop table %s'%(strTableName)
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, q,())
        return bSuccess

    def createEnv(self, envid, ena, name, description, tags, creator, createtime):
        strQuery = '''INSERT INTO env (
                            id,
                            enabled,
                            name,
                            description,
                            tags,
                            creator,
                            createtime
                        )
                        VALUES
                            (%s, %s ,%s ,%s ,%s ,%s ,%s)'''

        bSuc = self._mysqlDbContainer.op_db_insert("beopdata", strQuery, parameter=(
        envid, ena, name, description, tags, creator, createtime))
        return bSuc

    def checkIfEnvNameExists(self, envName):
        try:
            strQuery = """SELECT
                            count(name) AS num
                        FROM
                            env
                        WHERE
                            name = %s"""
            arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(envName,))
            if arr is None:
                return dict(code=1, data=None)


            if not arr:
                return dict(code=0, data=False)
            if not len(arr):
                return dict(code=0, data=False)
            if not len(arr[0]):
                return dict(code=0, data=False)
            if arr[0][0] > 0:
                return dict(code=0, data=True)
            return dict(code=0, data=False)
        except:
            return dict(code=1, data=None)

    def createMode(self, envid, type, ena, name, description, tags, creator, createtime):
        strQuery = '''INSERT INTO mode (
                            id,
                            type,
                            enabled,
                            name,
                            description,
                            tags,
                            creator,
                            createtime
                        )
                        VALUES
                            (%s, %s, %s ,%s ,%s ,%s ,%s ,%s)'''

        bSuc = self._mysqlDbContainer.op_db_insert("beopdata", strQuery, parameter=(
            envid, type, ena, name, description, tags, creator, createtime))
        return bSuc

    def getModeInfoForCopy(self, modeId):
        try:
            strQueryMode = "SELECT type, enabled, description, tags FROM mode WHERE id = %s"
            strQueryModeDetail = "SELECT modeid, triggerTime, triggerTimeType, triggerTurn, triggerEnvId, triggerPointNameList, triggerValue, actionOnce FROM mode_detail WHERE modeid = %s"
            modeInfo = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQueryMode, parameter=(modeId,))
            if modeInfo is None:
                return None
            modeDetailInfo = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQueryModeDetail, parameter=(modeId,))
            if modeDetailInfo is None:
                return None
            return modeInfo[0], modeDetailInfo
        except Exception as e:
            strLog = "ERROR in getModeInfoForCopy: %s" % e.__str__()
            return ()

    def insertMultiModeDetail(self, modeid, modeDetailList):
        try:
            paramList =[]
            for modeDetail in modeDetailList:
                modeDetailList = list(modeDetail)
                modeDetailList[0] = modeid
                paramList.append(tuple(modeDetailList))

            strQuery = "INSERT INTO mode_detail (modeid, triggerTime, triggerTimeType, triggerTurn, triggerEnvId, triggerPointNameList, triggerValue, actionOnce) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)"
            bSuc = self._mysqlDbContainer.op_db_update_many(app.config["DATABASE"], strQuery, parameter=(tuple(paramList)))
            return bSuc
        except Exception as e:
            strLog = "ERROR in insertMultiModeDetail: %s" % e.__str__()
            logging.error(strLog)
            return False

    def nameUsedByOtherMode(self, strModeName, modeId):
        try:
            strQuery = """SELECT count(*) FROM mode WHERE name = %s and id <> %s"""
            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(strModeName, modeId))
            if arr is None:
                return dict(code=1, data=None)

            if not arr:
                return dict(code=1, data=None)
            if not len(arr):
                return dict(code=1, data=None)
            if not len(arr[0]):
                return dict(code=1, data=None)
            if arr[0][0] == 0:
                return dict(code=0, data=False)
            else:
                return dict(code=0, data=True)
        except:
            return dict(code=1, data=None)

    def nameUsedByOtherEnv(self, strEnvName, envId):
        try:
            strQuery = """SELECT count(*) FROM env WHERE name = %s and id <> %s"""
            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(strEnvName, envId))
            if arr is None:
                return dict(code=1, data=None)
            if not arr:
                return dict(code=1, data=None)
            if not len(arr):
                return dict(code=1, data=None)
            if not len(arr[0]):
                return dict(code=1, data=None)
            if arr[0][0] == 0:
                return dict(code=0, data=False)
            else:
                return dict(code=0, data=True)
        except:
            return dict(code=1, data=None)


    def removeEnv(self, envid):
        strDeleteEnv = """DELETE
                    FROM
                        env
                    WHERE
                        id = %s"""

        strDeleteEnvDetail = """DELETE
                            FROM
                                env_detail
                            WHERE
                                envid = %s"""

        bSucEnv = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strDeleteEnv, parameter=(envid,))
        bSucEnvDetail = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strDeleteEnvDetail, parameter=(envid,))

        return bSucEnv and bSucEnvDetail

    def removeMode(self, envid):
        strQuery = """UPDATE mode
                    SET enabled = 0
                    WHERE
                        id = %s"""

        bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(envid,))
        return bSuc

    def editEnv(self, paramList, strSetList, envid):
        try:
            strSet = ",".join(strSetList)
            strQuery = "UPDATE env SET {0} WHERE id=%s".format(strSet)
            paramList.append(envid)
            bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(tuple(paramList)))
            return bSuc
        except Exception as e:
            strLog = "ERROR in editEnv: %s" % e.__str__()
            logging.error(strLog)
            return False

    def editMode(self, paramList, strSetList, modeId):
        try:
            strSet = ",".join(strSetList)
            strQuery = """UPDATE mode SET {0} WHERE id=%s""".format(strSet)
            paramList.append(modeId)
            bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(tuple(paramList)))
            return bSuc
        except Exception as e:
            strLog = "ERROR in editMode: %s" % e.__str__()
            logging.error(strLog)
            return False

    def getAllEnv(self):
        strQuery = """SELECT
                        *
                    FROM
                        env
                    WHERE
                        enabled > 0
                    ORDER BY
                        id ASC"""

        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)
        return arr

    def updateModeTypeInModeCalendar(self, modeId, type):
        try:
            strUpdate = "UPDATE mode_calendar SET ofSystem=%s WHERE modeid=%s AND ofDate >= %s"
            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strUpdate, parameter=(type, modeId, datetime.now().strftime("%Y-%m-%d")))
            return bSuc
        except Exception as e:
            logging.error("ERROR in updateModeTypeInModeCalendar: %s" % e.__str__())
            return False
        
    def getAllMode(self):

        strQuery = "SELECT id, type, enabled, name, description, tags, creator, createtime FROM mode"

        items = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)
        if items is None:
            return dict(err=1, msg="数据库繁忙", data={})

        if isinstance(items, list):
            if not len(items):
                return dict(err=0, msg="", data={})

        dMode = {}
        modeIdList = []
        for item in items:
            modeId = item[0]
            modeIdList.append(modeId)
            dMode.update({modeId: dict(
                    id=modeId,
                    type=item[1],
                    enabled=item[2],
                    name=item[3],
                    description=item[4],
                    tags=item[5],
                    creator=item[6],
                    createtime=item[7],
                    details=[]
                )})

        clauseList = ["modeid={id}".format(id=modeId) for modeId in modeIdList]

        if len(modeIdList):
            strQueryDetail = "SELECT triggerTime, triggerTimeType, triggerEnvId, actionOnce, modeid FROM mode_detail WHERE "
            strQueryDetail += " or ".join(clauseList)

            items = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQueryDetail)
            if items is None:
                return dict(err=1, msg="数据库繁忙", data={})

            for item in items:
                modeId = item[4]
                if dMode.get(modeId, None) == None:
                    continue

                dMode.get(modeId).get("details").append(dict(triggerTime=item[0], triggerTimeType=item[1], envId=item[2], actionOnce=item[3]))

            for modeId, dInfo in dMode.items():
                detailList = dInfo.get("details")
                for i in range(len(detailList)):
                    for j in range(len(detailList)-i-1):
                        try:
                            tTimePrev = datetime.strptime(detailList[j].get("triggerTime"), "%H:%M")
                            tTimeAfter = datetime.strptime(detailList[j+1].get("triggerTime"), "%H:%M")
                            if tTimePrev > tTimeAfter:
                                detailList[j], detailList[j+1] = detailList[j+1], detailList[j]
                        except:
                            pass
            return dict(err=0, msg="", data=dMode)
        return dict(err=0, msg="", data={})

    def getEnvById(self, envid):
        if envid is None:
            strQuery = "SELECT id, type, enabled, name, description, tags, creator, createtime FROM env"
            strQueryDetail = "SELECT envid, pointname, pointvalue FROM env_detail"
            arrEnv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            arrDetail = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQueryDetail)
        else:
            strQuery = """SELECT id, type, enabled, name, description, tags, creator, createtime FROM env WHERE id = %s"""
            strQueryDetail = """SELECT envid, pointname, pointvalue FROM env_detail WHERE envid = %s"""
            arrEnv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(envid,))
            arrDetail = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQueryDetail, parameter=(envid,))

        if arrEnv is None:
            return None, None

        if arrDetail is None:
            return None, None

        dEnv = {}
        dEnvDetail = {}
        for item in arrEnv:
            dEnv.update({item[0]: dict(id=item[0], type=item[1], enabled=item[2], name=item[3], description=item[4], tags=item[5], creator=item[6], createtime=item[7])})

        for item in arrDetail:
            if item[0] not in dEnvDetail.keys():
                dEnvDetail.update({item[0]: []})

            dEnvDetail.get(item[0]).append(
                (item[1], item[2])
            )

        return dEnv, dEnvDetail

    def getAllTagsFromEnv(self):
        try:
            strQuery = "SELECT tags FROM env"
            itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            proList = []
            for item in itemList:
                if item and len(item):
                    if not len(item[0]):
                        continue
                    proList = proList + item[0].split(";")

            # 按tag的首字母ascii码升序排列
            tagList = list(set(proList))
            for i in range(len(tagList)):
                for j in range(len(tagList)-i-1):
                    try:
                        if ord(tagList[j][0]) > ord(tagList[j+1][0]):
                            tagList[j], tagList[j+1] = tagList[j+1], tagList[j]
                    except:
                        pass

            return tagList
        except Exception as e:
            logging.error("ERORR in getAllTagsFromEnv: %s" % e.__str__())
            return None

    def getEnvByTags(self, tagList):
        try:
            whereClauseList = []
            for tag in tagList:
                whereClauseList.append("tags like '%{0}%'".format(tag))

            strWhereClause = " or ".join(whereClauseList)

            strQuery = "select * from env where {0} order by id asc".format(strWhereClause)

            envList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)

            return envList

        except Exception as e:
            logging.error("ERROR in getEnvByTags: %s" % e.__str__())
            return None

    def bubbleSortModeDetail(self, objectList):
        try:
            if not len(objectList):
                return []

            n = len(objectList)
            for i in range(n):
                for j in range(0, n-i-1):
                    try:
                        if datetime.strptime(objectList[j][3], "%H:%S") > datetime.strptime(objectList[j+1][3], "%H:%S"):
                            objectList[j], objectList[j+1] = objectList[j+1], objectList[j]
                    except:
                        continue

            return objectList
        except:
            return []

    def getModeById(self, modeId):
        strQuery = """SELECT * FROM mode WHERE id = %s"""

        dbname = app.config['DATABASE']
        nDbVersion = 3
        try:
            nDbVersion = app.config["DATABASE_VERSION"][dbname]
        except:
            pass
        strQueryDetail = """SELECT md.modeid, md.triggerTime, md.triggerEnvId, md.triggerTimeType, e.name 
        FROM mode_detail md left join env e on md.triggerEnvId = e.id WHERE md.modeid = %s"""
        if nDbVersion >= 4:
            strQueryDetail = """SELECT md.modeid, md.triggerTime, md.triggerEnvId, md.triggerTimeType, e.name, md.actionOnce FROM mode_detail md left join env e on md.triggerEnvId = e.id WHERE md.modeid = %s"""

        arrMode = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(modeId,))
        if arrMode is None:
            return None

        arrDetail = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQueryDetail, parameter=(modeId,))

        if arrDetail is None:
            return [arrMode, list()]

        if not len(arrDetail):
            return [arrMode, list()]

        triggerPointList = [list(item)[1] for item in arrDetail if list(item)[3] == 0]
        triggerPointValueList = BEOPDataAccess.getInstance().getInputTable(triggerPointList)

        if triggerPointValueList is None:
            return [arrMode, list()]

        triggerPointDict = {}
        if len(triggerPointList):
            for dData in triggerPointValueList[0]:
                triggerPointDict.update({dData.get("name"): dData.get("value")})

        detailList = []
        for item in arrDetail:
            try:
                itemList = list(item)

                modeId = itemList[0]
                triggerTimeOrPointName = itemList[1]
                envId = itemList[2]
                triggerTimeType = itemList[3]
                envName = itemList[4]
                systemTimePointName = triggerTimeOrPointName if triggerTimeType == 0 else ""

                triggerTime = triggerPointDict.get(triggerTimeOrPointName) if triggerTimeType == 0 else triggerTimeOrPointName

                actionOnce = None
                try:
                    actionOnce = itemList[5]
                except:
                    pass

                # 若时间类型为2（隔日时间）则减去24小时
                if triggerTimeType == 2:
                    hour = int(triggerTime.split(":")[0])
                    minute = int(triggerTime.split(":")[1])
                    newHour = hour - 24 if hour - 24 >= 0 else 0
                    triggerTime = "%02d:%02d" % (newHour, minute)

                detailList.append([modeId, envId, triggerTimeType, triggerTime, envName, systemTimePointName, actionOnce])

            except:
                pass

        # 按场景的触发时间冒泡排序
        detailListSorted = self.bubbleSortModeDetail(detailList)

        return [arrMode, detailListSorted]

    def countPointNameInEnvDetail(self, envid, pointName):
        strQuery = """SELECT
                        count(pointname) AS count
                    FROM
                        env_detail
                    WHERE
                        envid = %s
                    AND pointname = %s"""

        arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(envid, pointName))
        return arr

    def envIdExistsInEnv(self, envid):
        strQuery = """SELECT
                        count(id) AS count
                    FROM
                        env
                    WHERE
                        id = %s"""
        arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(envid,))
        return arr


    def modeIdExistsInMode(self, modeId):
        strQuery = """SELECT
                        count(id) AS count
                    FROM
                        mode
                    WHERE
                        id = %s"""
        arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(modeId,))
        return arr

    def updateEnvDetail(self, envId, envList):
        strQueryList = []
        paramList = []
        strQuery = "DELETE FROM env_detail WHERE envid = %s"

        strQueryList.append(strQuery)
        paramList.append(tuple((envId,),))


        strQuery = """INSERT INTO env_detail (envid, pointname, pointvalue)
                                VALUES
                                    (%s, %s, %s)"""

        strQueryList.append(strQuery)
        paramList.append(tuple(envList))

        if len(envList):
            arr = self._mysqlDbContainer.op_db_transaction_update_many(app.config['DATABASE'], strQueryList, paramList)
        else:
            arr = True
        return arr


    def insertIntoModeDetail(self, modeId, triggerTime, timeType, envId, actionOnce):
        try:
            dbname = app.config['DATABASE']
            nDbVersion = 3
            try:
                nDbVersion = app.config["DATABASE_VERSION"][dbname]
            except:
                pass

            if nDbVersion >= 4 and actionOnce == None:
                actionOnce = 0

            strQuery = """SELECT pointname, pointvalue FROM env_detail WHERE envid = %s"""

            strInsert = """INSERT INTO mode_detail (modeid, triggerTime, triggerTimeType, triggerEnvId, triggerPointNameList) VALUES (%s, %s, %s, %s, %s)"""
            if nDbVersion >= 4:
                strInsert = """INSERT INTO mode_detail (modeid, triggerTime, triggerTimeType, triggerEnvId, triggerPointNameList, actionOnce) VALUES (%s, %s, %s, %s, %s, %s)"""

            arrQuery = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(envId,))

            if arrQuery is None:
                return dict(code=1, msg="保存失败", data=False)

            # 场景中没有配置点也可以允许加入mode_detail表里
            # if not arrQuery:
            #     return dict(code=1, msg="该场景无详情内容", data=False)
            # if not len(arrQuery):
            #     return dict(code=1, msg="该场景无详情内容", data=False)

            pointNameList = []
            valueList = []
            for item in arrQuery:
                pointNameList.append(str(item[0]))
                # valueList.append(str(item[1]))

            strPointNameList = ",".join(pointNameList)
            # strValueList = ",".join(valueList)

            insertParam = (modeId, triggerTime, timeType, envId, strPointNameList)
            if nDbVersion >= 4:
                insertParam = (modeId, triggerTime, timeType, envId, strPointNameList, actionOnce)

            arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strInsert, parameter=insertParam)

            if arr:
                return dict(code=0, msg="保存成功", data=True)
            return dict(code=1, msg="保存失败", data=False)

        except Exception as e:

            return dict(code=1, msg="保存失败：%s" % e.__str__(), data=False)

    def updateModeDetail(self, modeId, oldTime, oldTimeType, oldEnvId, newTime, newTimeType, newEnvId, actionOnce):
        try:
            dbname = app.config['DATABASE']
            nDbVersion = 3
            try:
                nDbVersion = app.config["DATABASE_VERSION"][dbname]
            except:
                pass

            if nDbVersion >= 4 and actionOnce != None:
                if actionOnce not in [0, 1]:
                    return dict(code=1, msg="actionOnce必须为0或者1", data=False)

            strQueryEnvDetail = """SELECT pointname, pointvalue FROM env_detail WHERE envid = %s"""

            strUpdate = """UPDATE mode_detail SET triggerTime = %s, triggerTimeType = %s, triggerEnvId = %s, triggerPointNameList = %s WHERE modeid = %s AND triggerTime = %s AND triggerEnvId = %s"""
            if nDbVersion >= 4 and actionOnce in [0, 1]:
                strUpdate = """UPDATE mode_detail SET triggerTime = %s, triggerTimeType = %s, triggerEnvId = %s, triggerPointNameList = %s, actionOnce=%s WHERE modeid = %s AND triggerTime = %s AND triggerEnvId = %s"""

            strCount = """SELECT count(*) FROM mode_detail WHERE modeid = %s AND triggerTime = %s AND triggerEnvId = %s"""
            param = (modeId, newTime, newEnvId)
            if nDbVersion >= 4 and actionOnce in [0, 1]:
                strCount = """SELECT count(*) FROM mode_detail WHERE modeid = %s AND triggerTime = %s AND triggerEnvId = %s AND actionOnce=%s"""
                param = (modeId, newTime, newEnvId, actionOnce)

            arrCount = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strCount, parameter=param)

            if arrCount is None:
                return dict(code=1, msg="数据库繁忙，请重试", data=False)

            if not arrCount:
                return dict(code=1, msg="查询当前模式内容的记录条数失败", data=False)
            if not len(arrCount):
                return dict(code=1, msg="查询当前模式内容的记录条数失败", data=False)
            if not len(arrCount[0]):
                return dict(code=1, msg="查询当前模式内容的记录条数失败", data=False)

            if arrCount[0][0] > 0:
                return dict(code=1, msg="修改失败，因为期望的场景已存在", data=False)

            # 取新envId的detail
            arrEnvDetail = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQueryEnvDetail,
                                                              parameter=(newEnvId,))

            if arrEnvDetail is None:
                return dict(code=1, msg="数据库繁忙，请重试", data=False)

            envDetailList = arrEnvDetail
            if not arrEnvDetail or not len(arrEnvDetail):
                envDetailList = []

            # 将点名列表和值列表拼接成字符串
            pointNameList = []
            valueList = []
            for item in envDetailList:
                pointNameList.append(str(item[0]))
                # valueList.append(str(item[1]))

            strPointNameList = ",".join(pointNameList)
            # strValueList = ",".join(valueList)

            paramList = (newTime, newTimeType, newEnvId, strPointNameList, modeId, oldTime, oldEnvId)
            if nDbVersion >= 4 and actionOnce in [0, 1]:
                paramList = (newTime, newTimeType, newEnvId, strPointNameList, actionOnce, modeId, oldTime, oldEnvId)

            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strUpdate, parameter=paramList)

            if not bSuc:
                return dict(code=1, msg="修改失败: 修改原场景时失败", data=False)
            return dict(code=0, msg="修改成功", data=True)

        except Exception as e:
            return dict(code=1, msg="修改失败: %s" % e.__str__(), data=False)

    def deleteFromModeDetail(self, modeid):
        strQuery = """DELETE
                        FROM
                            mode_detail
                        WHERE
                            modeid = %s"""
        arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(modeid,))
        return arr

    def addEnvSchedule(self, name, envid, author):
        strQuery = 'INSERT INTO `schedule_list` (`type`, `name`, `point`,`loop`, `enable`, `author`) VALUES(1,%s,%s,0,0,%s)'
        arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(name, str(envid), author))
        return arr


    def bindModeToCalendar(self, strDate, modeId, author, nOfSystem):
        strQuery = 'INSERT INTO `mode_calendar` (`ofDate`, `modeid`, `creator`, `ofSystem`) VALUES(%s,%s,%s,%s)'
        arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery,
                                                  parameter=(strDate, str(modeId), author, str(nOfSystem)))
        return arr

    def editEnvSchedule(self, name, envid, author, id):
        strQuery = 'Update schedule_list set name="%s",point="%s", author="%s" where id="%d"' % (name, str(envid), author, id)
        arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery)
        return arr

    def envScheduleExists(self, id):
        res = False
        try:
            strQuery = """SELECT
                                count(*) AS count
                            FROM
                                schedule_list
                            WHERE
                                id = %s
                            AND type = 1"""
            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(id,))

            if arr and len(arr):
                if len(arr[0]):
                    if arr[0][0] > 0:
                        res = True

        except:
            pass
        finally:
            return res

    def getAllSchedule(self):
        strQuery = """SELECT
                            *
                        FROM
                            schedule_list
                        WHERE
                            type = 1"""
        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)

        return arr

    def getScheduleById(self, id):
        strQuery = """SELECT
                        *
                    FROM
                        schedule_list
                    WHERE
                        type = 1
                    AND
                        id = %s"""

        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(id,))
        return arr

    def getAllEnvSchedulePlan(self):
        strQuery = """SELECT
                            siw.id AS planId,
                            siw.weekday AS weekday,
                            siw.timeFrom AS time,
                            siw.groupid AS scheduleId,
                            sl.name AS scheduleName
                        FROM
                            schedule_list sl
                        LEFT JOIN schedule_info_weeky siw ON sl.id = siw.groupid
                        WHERE
                            type = 1"""
        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)

        if arr is None:
            return None

        rtList = list()

        if arr and len(arr):
            for item in arr:
                try:
                    rtList.append({
                        "planId": item[0],
                        "weekday": item[1],
                        "time": item[2],
                        "scheduleId": item[3],
                        "scheduleName": item[4]
                    })
                except:
                    pass

        return rtList

    def searchEnvByName(self, keyword):
        strQuery = """SELECT
                            *
                        FROM
                            env
                        WHERE
                            name LIKE '%{0}%'
                        AND enabled > 0""".format(keyword)
        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)
        return arr

    def getCurrentMode(self):
        res = None
        strQuery = """SELECT
                        unitproperty02 AS info
                    FROM
                        unit01"""
        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)

        if arr is None:
            return None

        if arr and len(arr):
            if len(arr[14]):
                try:
                    res = int(arr[14][0])
                except:
                    pass

        return res

    def insertEnvSchedulePlan(self, planList):
        res = False
        try:
            strQuery = """INSERT INTO schedule_info_weeky (
                            weekday,
                            timeFrom,
                            timeTo,
                            value,
                            groupid,
                            timeExecute
                        )
                        VALUES
                            (
                                %s,
                                %s,
                                %s,
                                %s,
                                %s,
                                %s
                            )"""
            arr = self._mysqlDbContainer.op_db_update_many(app.config["DATABASE"], strQuery, tuple(planList))
            if arr:
                res = True
        except:
            pass
        finally:
            return res

    def insertEnvSchedule(self, scheduleList):
        res = False
        try:
            strQuery = """INSERT INTO schedule_list (
                            type,
                            name,
                            point,
                            `loop`,
                            enable,
                            author
                        )
                        VALUES
                            (%s, %s, %s, %s, %s, %s)"""

            arr = self._mysqlDbContainer.op_db_update_many(app.config["DATABASE"], strQuery, tuple(scheduleList))
            # 这里需返回所有插入记录时新生成的自增id (日程ID)，回传给下面插入计划的
            if arr:
                res = True
        except:
            pass
        finally:
            return res

    def removeEnvSchedulePlan(self, planId):
        res = False
        try:
            strQuery = """DELETE
                        FROM
                            schedule_info_weeky
                        WHERE
                            id = %s"""

            arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(planId,))

            if arr:
                res = True
        except:
            pass

        finally:
            return res

    def removeMode(self, modeId):
        res = False
        try:
            strQuery1 = """DELETE
                            FROM
                                mode
                            WHERE
                                id = %s"""

            strQuery2 = """DELETE
                            FROM
                                mode_detail
                            WHERE
                                modeid = %s"""

            arr1 = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery1, parameter=(modeId,))

            arr2 = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery2, parameter=(modeId,))

            res = arr1 and arr2

        except:
            pass
        finally:
            return res

    def getRealtimeDataByPointName(self, strPointName):
        res = None
        try:
            strQuery = """SELECT
                                pointvalue
                            FROM
                                realtimedata_input
                            WHERE
                                pointname = %s"""

            arr = self._mysqlDbContainerRealTimePool.op_db_query(app.config["DATABASE"], strQuery, parameter=(strPointName,))

            if arr is None:
                res = None
                return res

            if len(arr):
                if len(arr[0]):
                    res = arr[0][0]

        except:
            pass
        finally:
            return res

    def insertIntoModeDetailMulti(self, dataList):
        res = False
        try:
            strQuery = """INSERT INTO mode_detail (
                                modeid,
                                triggerTime,
                                triggerEnvId
                            )
                            VALUES
                                (%s, %s, %s)"""

            arr = self._mysqlDbContainer.op_db_update_many(app.config["DATABASE"], strQuery, tuple(dataList))

            res = arr
        except:
            pass
        finally:
            return res

    def removeModeDetailContent(self, modeId, triggerTimeOrPointName, triggerEnvId):
        res = False
        try:
            strQuery = """DELETE
                            FROM
                                mode_detail
                            WHERE
                                modeid = %s
                            AND triggerTime = %s
                            AND triggerEnvId = %s"""

            bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(modeId, triggerTimeOrPointName, triggerEnvId))

            res = bSuc

        except:
            pass
        finally:
            return res


    def saveWeatherData(self, forecastList, basicDict, strUpdateTime):
        dbname = app.config['DATABASE']
        for item in forecastList:
            strDate = item.get("fxDate")
            dInfo = {
                "date": strDate,
                "basic": basicDict,
                "tmp_max": item.get("tempMax"),
                "tmp_min": item.get("tempMin"),
                "cond_code_d": item.get("iconDay"),
                "cond_code_n": item.get("iconNight"),
                "hum": item.get("humidity"),
                "update": {
                    "loc": strUpdateTime
                },
                "cond_txt_d": item.get("textDay"),
                "cond_txt_n": item.get("textNight"),
                "pcpn": item.get("precip"),
                "pres": item.get("pressure"),
                "vis": item.get("vis"),
                "sr": item.get("sunrise"),
                "ss": item.get("sunset"),
                "uv_index": item.get("uvIndex"),
                "wind_spd_d": item.get("windSpeedDay"),
                "wind_spd_n": item.get("windSpeedNight"),
                "wind_dir_d": item.get("windDirDay"),
                "wind_dir_n": item.get("windDirNight"),
                "wind_sc_d": item.get("windScaleDay"),
                "wind_sc_n": item.get("windScaleNight")
            }

            q = 'select ofDate, forcast, realtime  from weather_calendar where ofDate = %s '

            strForcast = str(dInfo)

            param = (strDate,)

            result = self._mysqlDbContainer.op_db_query(dbname, q, param)
            if result is None:
                return False
            tNow = datetime.now()
            if not result:
                q = 'insert into weather_calendar(ofDate, forcast, raw_update_time) values(%s, %s, %s)'
                param = (strDate, strForcast, tNow.strftime('%Y-%m-%d %H:%M:%S'))
                self._mysqlDbContainer.op_db_update(dbname, q, param)
            else:
                q = 'update weather_calendar set forcast = %s, raw_update_time=%s where ofDate = %s'
                param = (strForcast,  tNow.strftime('%Y-%m-%d %H:%M:%S'), strDate)
                self._mysqlDbContainer.op_db_update(dbname, q, param)
        return True

    def getWeatherData(self, strDate):
        dbname = app.config['DATABASE']

        q = 'select forcast  from weather_calendar where ofDate = %s '

        param = (strDate,)

        result = self._mysqlDbContainer.op_db_query(dbname, q, param)
        if result is None:
            return None

        tNow = datetime.now()
        if not result:
            return {}
        else:
            try:
                strForest = result[0][0]
                strForest = strForest.replace('\'', '"')
                jsonForest = json.loads(strForest)
            except:
                jsonForest = None

        return jsonForest

    def getWeatherDataForcastV2(self, strFromDate, nDays=30):
        rv = []
        dbname = app.config['DATABASE']

        tFromDate = datetime.strptime(strFromDate, '%Y-%m-%d')
        tToDate = tFromDate + timedelta(days=nDays)

        strToDate = tToDate.strftime('%Y-%m-%d')

        q = 'select ofDate, forcast from weather_calendar where ofDate >= %s and ofDate<=%s'

        param = (strFromDate, strToDate)

        result = self._mysqlDbContainer.op_db_query(dbname, q, param)

        if result is None:
            return None

        if not result:
            return rv
        else:
            try:
                for item in result:
                    strForest = item[1]
                    if strForest:
                        strForest = strForest.replace('\'', '"')
                        jsonForest = json.loads(strForest)
                        try:
                            strDate = jsonForest.get("date")
                            tDate = datetime.strptime(strDate, "%Y-%m-%d")
                            strDateNew = "{0}月{1}日".format(tDate.month, tDate.day)
                            jsonForest["date"] = strDateNew
                        except:
                            pass
                    else:
                        jsonForest = None

                    rv.append(dict(forcast=jsonForest, date=item[0].strftime('%Y-%m-%d')))
            except Exception as e:
                logging.error('ERROR in getWeatherDataForcastV2:' + e.__str__())
                jsonForest = None

        return rv

    def getWeatherDataForcast(self, strFromDate, nDays=31):
        rv = []
        dbname = app.config['DATABASE']

        tFromDate = datetime.strptime(strFromDate, '%Y-%m-%d')
        tToDate = tFromDate + timedelta(days=nDays)

        strToDate = tToDate.strftime('%Y-%m-%d')

        q = 'select ofDate, forcast,calendar  from weather_calendar where ofDate >= %s and ofDate<=%s'

        param = (strFromDate, strToDate)

        result = self._mysqlDbContainer.op_db_query(dbname, q, param)

        if result is None:
            return None

        if not result:
            return rv
        else:
            try:
                for item in result:
                    strForest = item[1]
                    if strForest:
                        strForest = strForest.replace('\'', '"')
                        jsonForest = json.loads(strForest)
                        try:
                            strDate = jsonForest.get("date")
                            tDate = datetime.strptime(strDate, "%Y-%m-%d")
                            strDateNew = "{0}月{1}日".format(tDate.month, tDate.day)
                            jsonForest["date"] = strDateNew
                        except:
                            pass
                    else:
                        jsonForest = None

                    strCalendar = item[2]
                    if strCalendar:
                        strCalendar = strCalendar.replace('\'', '"')
                        jsonCalendar = json.loads(strCalendar)
                    else:
                        jsonCalendar = None
                    rv.append(dict(forcast = jsonForest, date= item[0].strftime('%Y-%m-%d'), calendar=jsonCalendar))
            except Exception as e:
                logging.error('ERROR in getWeatherDataForcast:' + e.__str__())
                jsonForest = None

        return rv


    def getCalendarUpdatedDate(self):
        try:
            strQuery = 'SELECT ofDate, calendar FROM weather_calendar WHERE ofDate >= curdate() order by ofDate'
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)

            if items is None:
                return None
            if not len(items):
                return None

            tLatest = None
            for item in items:
                tTime = item[0]
                if item[1] is None or item[1] == "":
                    break
                if tLatest:
                    if isinstance(tTime, str):
                        tTime = datetime.strptime(item[0], "%Y-%m-%d %H:%M:%S")
                    if tTime != tLatest + timedelta(days=1):
                        break
                tLatest = tTime
            return tLatest
        except Exception as e:
            logging.error("ERROR in getCalendarUpdatedDate: %s" % e.__str__())
            return None


    def getCalendarOfDay(self, tDate):
        dbname = app.config['DATABASE']

        strOfDate = tDate.strftime('%Y-%m-%d')
        q = 'select calendar from weather_calendar where ofDate = %s '

        param = (strOfDate,)

        result = self._mysqlDbContainer.op_db_query(dbname, q, param)

        if result:
            strRV = result[0][0]
            if strRV:
                strRV = strRV.replace('\'', '"')
                return strRV

        return None

    def saveCalendarData(self, tDate, infoCalendar):
        dbname = app.config['DATABASE']

        strOfDate = tDate.strftime('%Y-%m-%d')
        q = 'select ofDate,calendar  from weather_calendar where ofDate = %s '
        strCalendarInfo = str(infoCalendar)

        param = (strOfDate,)

        result = self._mysqlDbContainer.op_db_query(dbname, q, param)
        if result is None:
            return None

        tNow = datetime.now()
        if not result:
            q = 'insert into weather_calendar(ofDate, calendar, raw_update_time) values(%s, %s, %s)'
            param = (strOfDate, strCalendarInfo, tNow.strftime('%Y-%m-%d %H:%M:%S'))
            self._mysqlDbContainer.op_db_update(dbname, q, param)
        else:
            q = 'update weather_calendar set calendar = %s, realtime=%s where ofDate = %s'
            param = (strCalendarInfo, tNow.strftime('%Y-%m-%d %H:%M:%S'), strOfDate)
            self._mysqlDbContainer.op_db_update(dbname, q, param)

        return True
    # {"2023-07-03": {}, "2023-07-04": {}}
    def saveCalendarDataMultiDays(self, dCalendarInfo):
        tNow = datetime.now()
        sqlList = []
        for strDate in dCalendarInfo.keys():
            sqlList.append("ofDate='{0}'".format(strDate))

        strQuery = "SELECT ofDate, calendar FROM weather_calendar WHERE " + " OR ".join(sqlList)
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        dValid = {}    # 日历数据正常的日期
        dInvalid = {}  # 日历数据不正常的日期
        for item in items:
            bValid = True
            if item[1] == None or item[1] == "":
                bValid = False

            dateKey = None
            if isinstance(item[0], datetime):
                dateKey = item[0].strftime("%Y-%m-%d")
            elif isinstance(item[0], str):
                if isValidDate(item[0], "%Y-%m-%d %H:%M:%S"):
                    dateKey = datetime.strptime(item[0], "%Y-%m-%d %H:%M:%S").strftime("%Y-%m-%d")
                elif isValidDate(item[0], "%Y-%m-%d"):
                    dateKey = item[0]

            if dateKey == None:
                continue

            if bValid:
                dValid.update({dateKey: item[1]})
            else:
                dInvalid.update({dateKey: None})

        updateList = []
        insertList = []
        for strDate, dCalendar in dCalendarInfo.items():
            if not isinstance(dCalendar, dict):
                continue

            if not dCalendar:
                continue

            if strDate in dValid.keys():
                continue

            if strDate in dInvalid.keys():
                updateList.append(
                    (str(dCalendar), tNow.strftime('%Y-%m-%d %H:%M:%S'), strDate)
                )
                continue

            insertList.append(
                (strDate, str(dCalendar), tNow.strftime('%Y-%m-%d %H:%M:%S'))
            )

        strUpdate = "UPDATE weather_calendar SET calendar = %s, realtime=%s WHERE ofDate = %s"
        strInsert = "INSERT INTO weather_calendar (ofDate, calendar, realtime) VALUES (%s, %s, %s)"

        bsuc1 = bsuc2 = True
        if len(updateList):
            bsuc1 = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strUpdate, tuple(updateList))

        if len(insertList):
            bsuc2 = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strInsert, tuple(insertList))

        return bsuc1 and bsuc2


    def PointExists(self, strPointName):
        res = False
        try:
            strQuery = """SELECT
                                count(*) AS count
                            FROM
                                realtimedata_input
                            WHERE
                                pointname = %s"""

            arr =  self._mysqlDbContainerRealTimePool.op_db_query(app.config["DATABASE"], strQuery, parameter=(strPointName,))
            if arr and len(arr):
                if len(arr[0]):
                    if arr[0][0] > 0:
                        res = True
        except:
            pass
        finally:
            return res


    def getSysLogicParameters(self, logicName):
        res = None
        try:
            strQuery = """SELECT
                            param_input
                        FROM
                            sys_logic
                        WHERE
                            name = %s"""

            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, (logicName,))
            if arr:
                if len(arr):
                    if len(arr[0]):
                        res = arr[0][0]

        except:
            pass
        finally:
            return res

    def getSysLogicList(self):
        logicList = list()
        try:
            strQuery = '''SELECT
                            *
                        FROM
                            sys_logic'''

            sqlData = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)
            if sqlData is None:
                return None

            for data in sqlData:
                res_dict = dict()
                logic_name = data[0] if data[0] else ""
                author = data[1] if data[1] else ""
                version = data[2] if data[2] else ""
                group_name = data[3] if data[3] else ""
                order_index = data[4] if data[4] else 0
                description = data[5] if data[5] else ""

                param_input = ""
                try:
                    param_input = data[6] if data[6] else ""
                except:
                    pass

                param_output = ""
                try:
                    param_output = data[7] if data[7] else ""
                except:
                    pass

                res_dict.update({
                    "logic_name": logic_name,
                    "author": author,
                    "version": version,
                    "group_name": group_name,
                    "order_index": order_index,
                    "description": description,
                    "param_input": param_input,
                    "param_output": param_output
                })

                logicList.append(res_dict)

        except:
            pass
        finally:
            return logicList

    def getCalendarWithWeather(self, strTimeFrom, strTimeTo):
        try:
            strQuery = """SELECT
                            ofDate,
                            forcast
                        FROM
                            weather_calendar
                        WHERE
                            ofDate >= %s
                        AND ofDate <= %s
                        ORDER BY
                            ofDate"""

            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(strTimeFrom, strTimeTo))
            if arr is None:
                return None

            if not len(arr):
                return dict()

            dataDict = dict()
            for item in arr:

                if not len(item):
                    continue

                tDate = item[0]

                infoDict = eval(item[1]) if item[1] is not None else {}

                if not infoDict:
                    continue

                weatherDict = {
                    "date": infoDict.get("date") if infoDict.get("date") else None,
                    "weather": infoDict.get("cond_txt_n") if infoDict.get("cond_txt_n") else None,
                    "temp_max": float(infoDict.get("tmp_max")) if infoDict.get("tmp_max") else None,
                    "temp_min": float(infoDict.get("tmp_min")) if infoDict.get("tmp_min") else None,
                    "hum": infoDict.get("hum") if infoDict.get("hum") else None,
                    "sunrise": infoDict.get("sr") if infoDict.get("sr") else None,
                    "sunset": infoDict.get("ss") if infoDict.get("ss") else None
                }

                dataDict[tDate.strftime("%Y-%m-%d")] = weatherDict

            return dataDict

        except:
            return dict()


    def getSystemTimeDefine(self):
        try:
            if app.config.get("USE_4DB_FILE_FORMAT"):
                domdb_4db = app.config.get("USE_4DB_NAME")
                if not os.path.exists(domdb_4db):
                    return dict(err=1)

            with SqliteManager(domdb_4db) as dbm:
                strQuery = """SELECT
                                *
                            FROM
                                local_config
                            WHERE
                                name = 'system_time_define'"""

                arr = dbm.exec_query(strQuery)

            if len(arr):
                content = arr[0].get("content").decode("gbk")
                if content:
                    resList = list()
                    infoDict = eval(content)
                    timeTypeList = list(infoDict.keys())
                    for type in timeTypeList:
                        resList.append({type: infoDict.get(type).get("point")})

                    return dict(err=0, data=resList)

            return dict(err=2)
        except Exception as e:

            return dict(err=3, msg=e.__str__())

    def getModeByType(self, nType=None):
        res = []
        try:
            strQuery = """SELECT
                            id,
                            type,
                            enabled,
                            name,
                            description,
                            tags,
                            creator,
                            createtime
                        FROM
                            mode order by name"""
            param = ()

            if nType is not None:
                strQuery = """SELECT
                                id,
                                type,
                                enabled,
                                name,
                                description,
                                tags,
                                creator,
                                createtime
                            FROM
                                mode
                            WHERE
                                type = %s  order by name"""
                param = (nType,)

            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=param)
            if arr is None:
                dict(code=1, msg="数据库繁忙，请重试", data=[])

            if not arr:
                return dict(code=0, msg="该查询条件下无数据", data=[])

            if not len(arr):
                return dict(code=0, msg="该查询条件下无数据", data=[])

            for item in arr:
                res.append({
                    "modeId": item[0],
                    "type": item[1],
                    "enabled": item[2],
                    "name": item[3],
                    "description": item[4],
                    "tags": item[5],
                    "creator": item[6],
                    "createtime": datetime.strftime(item[7], "%Y-%m-%d %H:%M:%S")
                })

            return dict(code=0, msg="获取成功", data=res)

        except Exception as e:

            return dict(code=2, msg=e.__str__(), data=[])

    def getModeByTypeList(self, typeList):
        modeDict = {}
        try:
            strTypeList = ",".join([str(nType) for nType in typeList])

            strQuery = """SELECT id,type,enabled,name,description,tags,creator,createtime FROM mode WHERE type in (%s)""" % strTypeList

            itemList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)
            if itemList is None:
                return dict(code=1, msg="数据库繁忙，请重试", data={})

            for item in itemList:
                dInfo = {
                    "modeId": item[0],
                    "type": item[1],
                    "enabled": item[2],
                    "name": item[3],
                    "description": item[4],
                    "tags": item[5],
                    "creator": item[6],
                    "createtime": datetime.strftime(item[7], "%Y-%m-%d %H:%M:%S")
                }
                if item[1] in modeDict.keys():
                    modeDict.get(item[1]).append(dInfo)
                else:
                    modeDict.update({
                        item[1]: [dInfo]
                    })

            return dict(code=0, msg="获取成功", data=modeDict)

        except Exception as e:
            return dict(code=1, msg="获取失败: %s" % e.__str__(), data={})


    def getModeByTime(self, strTime):
        try:
            strQuery = """SELECT
                                modeid
                            FROM
                                mode_calendar
                            WHERE
                                ofDate = %s"""


            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(strTime, ))

            if arr is None:
                return dict(code=2, msg='数据库繁忙，请重试', data=[])

            if not arr:
                return dict(code=0, msg="该查询条件下无数据", data=[])

            if not len(arr):
                return dict(code=0, msg="该查询条件下无数据", data=[])

            modeIdList = []
            for item in arr:
                modeIdList.append(item[0])

            return dict(code=0, msg="获取成功", data=modeIdList)

        except Exception as e:

            return dict(code=2, msg=e.__str__(), data=[])

    def getModeByTypeAndTime(self, strTime, nType):
        try:
            strQuery = """SELECT
                                modeid
                            FROM
                                mode_calendar
                            WHERE
                                ofDate = %s
                            AND ofSystem = %s"""

            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(strTime, nType))
            if arr is None:
                return dict(code=2, msg='数据库繁忙，请重试', data=None)
            if not arr:
                return dict(code=0, msg="该查询条件下无数据", data=None)

            if not len(arr):
                return dict(code=0, msg="该查询条件下无数据", data=None)

            modeId = arr[0][0]

            return dict(code=0, msg="获取成功", data=modeId)

        except Exception as e:

            return dict(code=2, msg=e.__str__(), data=None)

    def getModeByTypeListAndTime(self, strTime, typeList):
        modeDict = {}
        try:
            strTypeList = ",".join([str(nType) for nType in typeList])
            strQuery = """SELECT ofSystem, modeid FROM mode_calendar WHERE ofDate = '%s' and ofSystem in (%s)""" % (strTime, strTypeList)

            itemList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)

            if itemList is None:
                return dict(code=2, msg='数据库繁忙，请重试', data=[])

            for item in itemList:
                modeDict.update({item[0]: item[1]})

            return dict(code=0, msg="获取成功", data=modeDict)

        except Exception as e:

            return dict(code=2, msg=e.__str__(), data={})

    def getTypeOfMode(self, modeId):
        try:
            strQuery = "SELECT type FROM mode WHERE id = %s"
            rt = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(modeId,))
            ofSystem = rt[0][0]
            return int(ofSystem)
        except Exception as e:
            logging.error("ERROR in getTypeOfMode: %s" % e.__str__())
            return None


    def insertIntoModeCalendar(self, date, modeId, creator, type):
        try:
            if self.getTypeOfMode(modeId) != type:
                return dict(code=1, msg="传入的模式type有误", data=False)

            strQuery = """INSERT INTO mode_calendar (ofDate, modeid, creator, ofSystem) VALUES
                                (
                                    %s,
                                    %s,
                                    %s,
                                    %s
                                )"""

            strQueryCheck = """SELECT ofSystem FROM mode_calendar WHERE ofDate = %s AND modeid = %s"""

            strQueryDelete = """DELETE FROM mode_calendar WHERE ofDate = %s AND ofSystem = %s"""
            strQueryDelete2 = "DELETE FROM mode_calendar WHERE ofDate=%s AND modeid=%s AND ofSystem=%s"

            arrCheck = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQueryCheck, parameter=(date, modeId))
            if arrCheck is None:
                return dict(code=2, msg='数据库繁忙，请重试', data=False)

            if len(arrCheck):
                if len(arrCheck[0]):
                    if arrCheck[0][0] == type:
                        return dict(code=0, msg="该日期下该模式已存在", data=True)
                    else:
                        bDelete = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQueryDelete2,
                                                                        parameter=(date, modeId, arrCheck[0][0]))
                        if not bDelete:
                            return dict(code=1, msg='保存失败', data=False)

            # if not arrCheck:
            #     return dict(code=1, msg="保存失败", data=False)
            # if not len(arrCheck):
            #     return dict(code=1, msg="保存失败", data=False)
            # if not len(arrCheck[0]):
            #     return dict(code=1, msg="保存失败", data=False)

            # # 判断该记录已经存在于表中则直接返回“保存成功”
            # if arrCheck[0][0] > 0:
            #     return dict(code=0, msg="保存成功", data=True)

            # 将该系统该日期下的模式删除
            arrDelete = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQueryDelete, parameter=(date, type))
            if not arrDelete:
                return dict(code=1, msg="保存失败", data=False)

            # 插入表
            arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(date, modeId, creator, type))
            if not arr:
                return dict(code=1, msg="保存失败", data=False)

            return dict(code=0, msg="保存成功", data=True)

        except Exception as e:

            return dict(code=1, msg="保存失败：%s" % e.__str__(), data=False)


    def batchInsertIntoModeCalendar(self, year, month, fromDay, toDay, modeId, creator, type, wdList):
        try:
            # 生成待写入表中的数据列表、生成日期列表
            params = []
            dateList = []
            idx = fromDay
            while idx <= toDay:
                strTemp = "%s-%02d-%02d" % (year, month, idx)
                try:
                    tTemp = datetime.strptime(strTemp,  "%Y-%m-%d")
                    if tTemp.weekday() in wdList:
                        params.append((strTemp, modeId, creator, type))
                        dateList.append("'{0}'".format(tTemp.strftime("%Y-%m-%d")))
                except:
                    params.append((strTemp, modeId, creator, type))
                idx += 1

            # 删除本月所有（或工作日或周末）该type下的模式
            strDeleteClause = "({0})".format(",".join(dateList))
            strDelete = """DELETE FROM mode_calendar WHERE ofDate in %s and ofSystem = %s""" % (strDeleteClause, type)
            arrDelete = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strDelete)
            if not arrDelete:
                return dict(code=1, msg="保存失败", data=False)

            # 写入表中
            strInsert = """INSERT INTO mode_calendar (ofDate, modeid, creator, ofSystem) VALUES (%s, %s, %s, %s)"""
            arrInsert = self._mysqlDbContainer.op_db_update_many(app.config["DATABASE"], strInsert,
                                                              parameter=tuple(params))

            if not arrInsert:
                return dict(code=1, msg="保存失败", data=False)

            return dict(code=0, msg="保存成功", data=True)

        except Exception as e:

            return dict(code=1, msg="获取失败：%s" % e.__str__(), data=[])

    def getModeOfTypeOnCertainDate(self, nType, date):
        try:
            strQuery = """select mc.ofDate, mc.modeid, mc.ofSystem, m.name from mode m left join mode_calendar mc 
                            on m.id = mc.modeid where mc.ofDate = %s and mc.ofSystem = %s"""
            params = (date, nType)
            itemList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=params)
            if itemList is None:
                return dict(code=1, msg='数据库繁忙，请重试', data={})
            if not len(itemList):
                return dict(code=1, msg='该类型该日期下无模式', data={})
            if not len(itemList[0]):
                return dict(code=1, msg='该类型该日期下无模式', data={})

            return dict(code=0, msg="", data={"modeId": itemList[0][1], "modeName": itemList[0][3], "type": itemList[0][2]})
        except Exception as e:
            logging.error("ERROR in getModeOfTypeOnCertainDate: %s" % e.__str__())
            return dict(code=1, data={})

    def getCalendarWithMode(self, strTimeFrom, strTimeTo, nType):
        try:
            strQuery = """select mc.ofDate, mc.modeid, mc.ofSystem, m.name from mode m left join mode_calendar mc 
                            on m.id = mc.modeid where mc.ofDate >= %s and mc.ofDate <= %s order by mc.ofDate"""
            params = (strTimeFrom, strTimeTo)

            strWeatherQuery = """select ofDate, forcast from weather_calendar where ofDate >=%s and ofDate <=%s"""
            weatherParams = (strTimeFrom, strTimeTo)

            if nType is not None:
                strQuery = """select mc.ofDate, mc.modeid, mc.ofSystem, m.name from mode m left join mode_calendar mc 
                on m.id = mc.modeid where mc.ofDate >= %s and mc.ofDate <= %s and mc.ofSystem = %s order by mc.ofDate"""

                params = (strTimeFrom, strTimeTo, nType)

            itemList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=params)
            weatherList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strWeatherQuery, parameter=weatherParams)

            if itemList is None:
                return dict(code=2, msg='数据库繁忙，请重试', data=None)

            if weatherList is None:
                return dict(code=2, msg='数据库繁忙，请重试', data=None)

            weatherDict = {}
            for weather in weatherList:
                try:
                    if len(weather) < 2:
                        continue

                    strDate = weather[0].strftime("%Y-%m-%d")

                    weatherInfoDict = {}
                    if weather[1]:
                        weatherInfoDict = eval(weather[1])

                    if not weatherInfoDict:
                        continue

                    basic = weatherInfoDict.get("basic", "")
                    location = basic.get("location", "") if basic else ""
                    weatherDict.update({
                        strDate: {
                            "code": weatherInfoDict.get("cond_code_d"),
                            "tempMax": weatherInfoDict.get("tmp_max"),
                            "tempMin": weatherInfoDict.get("tmp_min"),
                            "hum": weatherInfoDict.get("hum"),
                            "desc": weatherInfoDict.get("cond_txt_d"),
                            "windSc": weatherInfoDict.get("wind_sc"),
                            "windDir": weatherInfoDict.get("wind_dir"),
                            "location": location
                        }
                    })
                except Exception as e:
                    logging.error("ERROR in get weather data of /calendar/getCalendarWithMode: %s" % e.__str__())

            dInfo = {}
            for item in itemList:
                date = item[0].strftime("%Y-%m-%d")
                dModeInfo = {
                    "modeId": item[1],
                    "modeName": item[3],
                    "type": item[2]
                }

                if date not in dInfo.keys():
                    dInfo.update({date: [dModeInfo]})
                else:
                    dInfo.get(date).append(dModeInfo)

            return dict(code=0, data=(dInfo, weatherDict))

        except Exception as e:

            return dict(code=1, msg="获取失败：%s" % e.__str__(), data=[])

    def removeModeFromCalendar(self, modeId, date):
        try:
            strQuery = """DELETE
                            FROM
                                mode_calendar
                            WHERE
                                ofDate = %s
                            AND modeid = %s"""

            arr = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, parameter=(date, modeId))

            if arr:
                return dict(code=0, msg="删除成功")
            else:
                return dict(code=1, msg="删除失败")

        except Exception as e:
            return dict(code=1, msg="删除失败:%s" % e.__str__())


    def batchRemoveModeFromCalendar(self, modeId, tFrom, tTo):
        try:
            strDelete = "DELETE FROM mode_calendar WHERE modeid=%s AND ofDate >= %s AND ofDate <= %s"
            bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strDelete, parameter=(modeId,
                                                                                                    tFrom.strftime("%Y-%m-%d"),
                                                                                                    tTo.strftime("%Y-%m-%d")))
            if bSuc is None:
                return dict(code=1, msg="数据库繁忙，请稍后重试")
            if bSuc:
                return dict(code=0, msg="解绑成功")
            return dict(code=1, msg="解绑失败")
        except Exception as e:
            return dict(code=1, msg="解绑失败:%s" % e.__str__())

    def batchRemoveModeFromCalendarThisMonth(self, nType, tFrom, tTo):
        try:

            strDelete = "DELETE FROM mode_calendar WHERE ofSystem=%s AND ofDate >= %s AND ofDate <= %s"
            param = (nType, tFrom.strftime("%Y-%m-%d"), tTo.strftime("%Y-%m-%d"))
            if nType is None:
                strDelete = "DELETE FROM mode_calendar WHERE ofDate >= %s AND ofDate <= %s"
                param = (tFrom.strftime("%Y-%m-%d"),tTo.strftime("%Y-%m-%d"))

            bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strDelete, parameter=param)
            if bSuc is None:
                return dict(code=1, msg="数据库繁忙，请稍后重试")
            if bSuc:
                return dict(code=0, msg="解绑成功")
            return dict(code=1, msg="解绑失败")
        except Exception as e:
            return dict(code=1, msg="解绑失败:%s" % e.__str__())


    def getCurrentActiveModeName(self, nType):
        try:
            strQuery = """SELECT
                            mode.name
                        FROM
                            mode_calendar
                        LEFT JOIN mode ON mode_calendar.modeid = mode.id
                        WHERE
                            mode_calendar.ofSystem = %s
                        AND mode_calendar.ofDate = %s"""
            strDate = datetime.now().strftime("%Y-%m-%d")
            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(nType, strDate))
            if arr is None:
                return None

            if not arr:
                return None
            if not len(arr):
                return None
            if not len(arr[0]):
                return None
            return arr[0][0]
        except Exception as e:
            strLog = "ERROR in getCurrentActiveModeName: %s" % e.__str__()
            logging.error(strLog)
            return None

    def getAllFixPosInfo(self, pageId):
        fixList = []
        strQuery = "SELECT pcf.fixId, pcf.posX, pcf.posY, pcf.visible, f.urgent FROM page_contain_fix pcf LEFT JOIN fix f ON pcf.fixId = f.id WHERE pageId = %s"
        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(pageId,))
        if arr is None:
            return None

        for item in arr:
            fixList.append({
                "fixId": item[0],
                "posX": item[1],
                "posY": item[2],
                "visible": item[3],
                "urgent": item[4]
            })
        return fixList

    def createFix(self, fixId, reportTime, importance, urgent, content, reportUser, solveUser, energyEffects, x, y, visible, pageId, reportTitle):
        nResult = 0  # 创建时默认置为未解决（0-未解决；1-已解决；2-已关闭）
        strInsertFix = "INSERT INTO fix (id, reportTime, importance, urgent, content, reportUser, solveUser, energyEffects, title, result) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        strInsertPageContainFix = "INSERT INTO page_contain_fix (pageId, fixId, posX, posY, visible) VALUE (%s, %s, %s, %s, %s)"

        dFix = self._mysqlDbContainer.op_db_insert(app.config["DATABASE"], strInsertFix, parameter=(fixId, reportTime, importance, urgent, content, reportUser, solveUser, energyEffects, reportTitle, nResult))
        dPageContainFix = self._mysqlDbContainer.op_db_insert(app.config["DATABASE"], strInsertPageContainFix, parameter=(pageId, fixId, x, y, visible))

        return dFix.get("success") and dPageContainFix.get("success")

    def removeFix(self, fixId):
        strDeleteFix = "DELETE FROM fix WHERE id = %s"
        strDeletePageContainFix = "DELETE FROM page_contain_fix WHERE fixId = %s"
        bFix = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strDeleteFix, parameter=(fixId,))
        bPageContainFix = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strDeletePageContainFix, parameter=(fixId,))
        return bFix and bPageContainFix

    def fixIdExistsInFix(self, fixId):
        strQuery = """SELECT
                        count(id) AS count
                    FROM
                        fix
                    WHERE
                        id = %s"""
        arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(fixId,))
        return arr

    def modifyFix(self, fixId, paramDict):
        try:
            strQuery = "SELECT f.reportTime, f.importance, f.urgent, f.content, f.result, f.closeTime, f.reportUser, " \
                       "f.solveUser, f.energyEffects, pcf.posX, pcf.posY, pcf.visible, f.title FROM fix f " \
                       "LEFT JOIN page_contain_fix pcf ON f.id = pcf.fixId WHERE f.id = %s"

            strUpdate = "UPDATE fix f LEFT JOIN page_contain_fix pcf ON f.id = pcf.fixId SET f.reportTime=%s, " \
                        "f.importance=%s, f.urgent=%s, f.content=%s, f.result=%s, f.closeTime=%s, f.reportUser=%s, " \
                        "f.solveUser=%s, f.energyEffects=%s, pcf.posX=%s, pcf.posY=%s, pcf.visible=%s, f.title=%s WHERE pcf.fixId = %s"

            arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(fixId,))

            if arr is None:
                return False

            if not arr:
                return False
            if not len(arr):
                return False
            if not len(arr[0]):
                return False

            paramDictExists = {}
            fieldList = ["reportTime", "importance", "urgent", "content", "result", "closeTime",
                         "reportUser", "solveUser", "energyEffects", "x", "y", "visible", "title"]
            for idx, item in enumerate(fieldList):
                paramDictExists[item] = arr[0][idx]

            # 替换需要update的字段内容
            for key in paramDict.keys():
                paramDictExists[key] = paramDict[key]

            readyList = []
            for field in fieldList:
                readyList.append(paramDictExists.get(field))

            readyList.append(fixId)

            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strUpdate, parameter=tuple(readyList))

            return bSuc
        except Exception as e:
            strLog = "ERROR in modifyFix: %s" % e.__str__()
            logging.error(strLog)
            return False

    def getFixById(self, fixId):
        try:
            strQuery = "SELECT f.reportTime, f.importance, f.urgent, f.content, f.result, f.closeTime, f.reportUser, " \
                       "f.solveUser, f.energyEffects, pcf.posX, pcf.posY, pcf.visible, f.title FROM fix f " \
                       "LEFT JOIN page_contain_fix pcf ON f.id = pcf.fixId WHERE f.id = %s"

            arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(fixId,))

            if arr is None:
                return None

            if not arr:
                return None
            if not len(arr):
                return None
            if not len(arr[0]):
                return None

            levelDict = {0: "低", 1: "中", 2: "高"}
            resultDict = {0: "未解决", 1: "已解决"}
            data = {
                "reportTime": arr[0][0].strftime("%Y-%m-%d %H:%M:%S"),
                "importance": levelDict.get(arr[0][1], "低"),
                "urgent": levelDict.get(arr[0][2], "低"),
                "content": arr[0][3],
                "result": resultDict.get(arr[0][4], None),
                "closeTime": None if arr[0][5] < arr[0][0] else arr[0][5].strftime("%Y-%m-%d %H:%M:%S"),
                "reportUser": arr[0][6],
                "solveUser": arr[0][7],
                "energyEffects": levelDict.get(arr[0][8], "低"),
                "x": arr[0][9],
                "y": arr[0][10],
                "visible": arr[0][11],
                "title": arr[0][12]
            }

            return data
        except Exception as e:
            strLog = "ERROR in getFixById: %s" % e.__str__()
            logging.error(strLog)
            return None

    def getTodayWeatherInfo(self, strToday):
        try:
            strQuery = "SELECT forcast FROM weather_calendar WHERE ofDate = %s"
            arr = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(strToday,))
            if arr is None:
                return None

            if not len(arr):
                return {}
            if not len(arr[0]):
                return {}
            if not arr[0][0]:
                return {}

            temp = None
            res = BEOPDataAccess.getInstance().getInputTable(["WeatherTdryNow"])[0]
            if len(res):
                temp = res[0].get("value")

            dInfo = eval(arr[0][0])
            basic = dInfo.get("basic", "")
            location = basic.get("location", "") if basic else ""

            curHour = datetime.now().hour

            windSc = dInfo.get("wind_sc_n") if (curHour >= 18 and curHour <= 23) or (curHour >= 0 and curHour <= 6) else dInfo.get("wind_sc_d")

            weather = {
                "code": dInfo.get("cond_code_d"),
                "tempMax": dInfo.get("tmp_max"),
                "tempMin": dInfo.get("tmp_min"),
                "hum": dInfo.get("hum"),
                "temp": temp,
                "desc": dInfo.get("cond_txt_d"),
                "windSc": windSc,
                "windDir": dInfo.get("wind_dir"),
                "location": location
            }
            return weather
        except Exception as e:
            strLog = "ERROR in getTodayWeatherInfo: %s" % e.__str__()
            logging.error(strLog)
            return {}

    def pointExistsInSchedule(self, strPointName):
        try:
            strToday = datetime.now().strftime("%Y-%m-%d")
            strQuery = "select triggerpointNameList from mode_calendar mc left join mode_detail md on mc.modeid = md.modeid where mc.ofDate = %s"
            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(strToday,))
            if arr is None:
                return None

            if not arr:
                return None
            if not len(arr):
                return None

            strPoints = ""
            for tpl in arr:
                strPoints += tpl[0]
            pointNameList = strPoints.split(",")
            if not strPointName in pointNameList:
                return None

            currentValue = BEOPDataAccess.getInstance().getRealtimeDataByPointName(strPointName)

            return currentValue
        except Exception as e:
            strLog = "ERROR in pointExistsInSchedule: %s" % e.__str__()
            logging.error(strLog)
            return ""

    def getFixContentById(self, fixId):
        try:
            strQuery = "SELECT reportTime, importance, urgent, content, result, closeTime, reportUser, solveUser, energyEffects, title FROM fix WHERE  id = %s"
            arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(fixId,))

            if arr is None:
                return None

            if not arr:
                return None
            if not len(arr):
                return None
            return arr[0]
        except Exception as e:
            strLog = "ERROR in getFixContentById: %s" % e.__str__()
            logging.error(strLog)
            return None

    def getFixContentByResult(self, nResult=None):
        strQuery = "SELECT reportTime, content, urgent, importance, energyEffects, reportUser, title, result, id, closeTime FROM fix"
        if nResult != None:
            strQuery = "SELECT reportTime, content, urgent, importance, energyEffects, reportUser, title, result, id, closeTime FROM fix" \
                       " where result = %s" % nResult
        arr = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)
        if arr is None:
            return None

        if not len(arr):
            return None
        return arr

    def getFixByPeriod(self, strTimeFrom, strTimeTo):
        strQuery = "SELECT reportTime, importance, urgent, content, result, closeTime, reportUser, solveUser, " \
                   "energyEffects, id, title from fix where reportTime > %s and reportTime < %s"
        itemList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery, parameter=(strTimeFrom, strTimeTo))

        if itemList is None:
            return None
        if len(itemList) == 0:
            return []

        dataList = []
        levelDict = {0: "低", 1: "中", 2: "高"}
        resultDict = {0: "未解决", 1: "已解决", 2: "已关闭"}
        for item in itemList:
            data = {
                "reportTime": item[0].strftime("%Y-%m-%d %H:%M:%S"),
                "importance": levelDict.get(item[1], None),
                "urgent": levelDict.get(item[2], None),
                "content": item[3],
                "result": resultDict.get(item[4], "未解决"),
                "closeTime": None if item[5] < item[0] else item[5].strftime("%Y-%m-%d %H:%M:%S"),
                "reportUser": item[6],
                "solveUser": item[7],
                "energyEffects": levelDict.get(item[8], None),
                "fixId": item[9],
                "title": item[10]
            }
            dataList.append(data)
        return dataList

    def keywordSearchFix(self, keyword):
        try:
            strQuery = 'SELECT reportTime, importance, urgent, content, result, closeTime, reportUser, solveUser, ' \
                       'energyEffects, id, title FROM fix WHERE title LIKE "%' + keyword + '%" OR content LIKE "%' + keyword + '%"'

            itemList = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)

            if itemList is None:
                return None
            if len(itemList) == 0:
                return []

            dataList = []
            levelDict = {0: "低", 1: "中", 2: "高"}
            resultDict = {0: "未解决", 1: "已解决"}
            for item in itemList:
                data = {
                    "reportTime": item[0].strftime("%Y-%m-%d %H:%M:%S"),
                    "importance": levelDict.get(item[1], None),
                    "urgent": levelDict.get(item[2], None),
                    "content": item[3],
                    "result": resultDict.get(item[4], None),
                    "closeTime": None if item[5] < item[0] else item[5].strftime("%Y-%m-%d %H:%M:%S"),
                    "reportUser": item[6],
                    "solveUser": item[7],
                    "energyEffects": levelDict.get(item[8], None),
                    "fixId": item[9],
                    "title": item[10]
                }
                dataList.append(data)
            return dataList
        except Exception as e:
            strLog = "ERROR in keywordSearchFix: %s" % e.__str__()
            logging.error(strLog)
            return None

    def setRealtimeDataToModbusEquipmentTable(self, strPointNameList, strPointValueList):
        if len(strPointNameList) <= 0:
            return dict(err=1, msg='no points')
        dbname = app.config['DATABASE']
        nDecimal = app.config['DECIMAL']
        if nDecimal is None:
            nDecimal = 2

        strNewPointValueList = []
        for pv in strPointValueList:
            if isinstance(pv, float):
                pv = round(pv, nDecimal)
            strNewPointValueList.append(pv)
        q = 'replace into realtimedata_input_modbus_equipment values(%s,%s, %s)'
        alldata = []
        nIndex = 0
        tNow = datetime.now()
        strTimeNow = tNow.strftime('%Y-%m-%d %H:%M:%S')
        for nIndex in range(len(strPointNameList)):
            alldata.append(tuple([strTimeNow, strPointNameList[nIndex], str(strPointValueList[nIndex])]))
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainerModbusClientPool.op_db_update_many(dbname, q, params)
        if bSuccess:
            return dict(err=0, msg='')
        else:
            return dict(err=1, msg='insert failed')

    def setRealtimeDataToObixInputTable(self, strPointNameList, strPointValueList):
        if len(strPointNameList) <= 0:
            return None
        dbname = app.config['DATABASE']
        nDecimal = app.config['DECIMAL']
        if nDecimal is None:
            nDecimal = 2

        strNewPointValueList = []
        for pv in strPointValueList:
            if isinstance(pv, float):
                pv = round(pv, nDecimal)
            strNewPointValueList.append(pv)
        q = 'replace into realtimedata_input_obix values(%s,%s, %s)'
        alldata = []

        tNow = datetime.now()
        strTimeNow = tNow.strftime('%Y-%m-%d %H:%M:%S')
        for nIndex in range(len(strPointNameList)):
            alldata.append(tuple([strTimeNow, strPointNameList[nIndex], str(strPointValueList[nIndex])]))
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainerRealTimePool.op_db_update_many(dbname, q, params)
        if bSuccess:
            return len(strPointNameList)
        return None

    def setRealtimeDataToLogixInputTable(self, strPointNameList, strPointValueList):
        if not len(strPointNameList):
            return False
        try:
            strDel = 'delete from realtimedata_input_logix where pointname in (%s)'

            strIns = "INSERT INTO realtimedata_input_logix VALUES (%s, %s, %s)"

            paramDel = [(pointName,) for pointName in strPointNameList]

            paramIns = [(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pointName, strPointValueList[idx])
                        for idx, pointName in enumerate(strPointNameList)]

            bSuc = self._mysqlDbContainerRealTimePool.op_db_transaction_update_many(app.config['DATABASE'], [strDel, strIns],
                                                                        [tuple(paramDel), tuple(paramIns)])
            if bSuc:
                return len(strPointNameList)
            return None

        except Exception as e:
            logging.error("ERROR in setRealtimeDataToLogixInputTable: %s" % e.__str__())

            return None

    def setRealtimeDataToAbslcInputTable(self, strPointNameList, strPointValueList):
        if not len(strPointNameList):
            return False
        try:
            strDel = 'delete from realtimedata_input_abslc where pointname in (%s)'

            strIns = "INSERT INTO realtimedata_input_abslc VALUES (%s, %s, %s)"

            paramDel = [(pointName,) for pointName in strPointNameList]

            paramIns = [(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pointName, strPointValueList[idx])
                        for idx, pointName in enumerate(strPointNameList)]

            bSuc = self._mysqlDbContainerRealTimePool.op_db_transaction_update_many(app.config['DATABASE'], [strDel, strIns],
                                                                        [tuple(paramDel), tuple(paramIns)])
            if bSuc:
                return len(strPointNameList)
            return None

        except Exception as e:
            logging.error("ERROR in setRealtimeDataToAbslcInputTable: %s" % e.__str__())

            return None

    def setRealtimeDataToKnxInputTable(self, strPointNameList, strPointValueList):
        if not len(strPointNameList):
            return False
        try:
            strDel = 'delete from realtimedata_input_knx where pointname in (%s)'

            strIns = "INSERT INTO realtimedata_input_knx VALUES (%s, %s, %s)"

            paramDel = [(pointName,) for pointName in strPointNameList]

            paramIns = [(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pointName, strPointValueList[idx])
                        for idx, pointName in enumerate(strPointNameList)]

            bSuc = self._mysqlDbContainerRealTimePool.op_db_transaction_update_many(app.config['DATABASE'], [strDel, strIns],
                                                                        [tuple(paramDel), tuple(paramIns)])
            if bSuc:
                return len(strPointNameList)
            return None

        except Exception as e:
            logging.error("ERROR in setRealtimeDataToKnxInputTable: %s" % e.__str__())

            return None

    def setRealtimeDataToDlt645InputTable(self, strPointNameList, strPointValueList):
        if not len(strPointNameList):
            return False
        try:
            strDel = 'delete from realtimedata_input_dlt645 where pointname in (%s)'

            strIns = "INSERT INTO realtimedata_input_dlt645 VALUES (%s, %s, %s)"

            paramDel = [(pointName,) for pointName in strPointNameList]

            paramIns = [(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pointName, strPointValueList[idx])
                        for idx, pointName in enumerate(strPointNameList)]

            bSuc = self._mysqlDbContainerRealTimePool.op_db_transaction_update_many(app.config['DATABASE'], [strDel, strIns],
                                                                        [tuple(paramDel), tuple(paramIns)])
            if bSuc:
                return len(strPointNameList)
            return None
        except Exception as e:
            logging.error("ERROR in setRealtimeDataToDlt645InputTable: %s" % e.__str__())
            return None

    def setRealtimeDataToDcimInputTable(self, strTimeList, strPointNameList, strPointValueList):
        if not len(strPointNameList):
            return False

        try:
            strQuery = "REPLACE INTO realtimedata_input_dcim VALUES (%s, %s, %s)"

            dataList = []
            for idx, pointName in enumerate(strPointNameList):
                dataList.append(
                    (strTimeList[idx], pointName, strPointValueList[idx])
                )

            bSuc = self._mysqlDbContainerModbusClientPool.op_db_update_many(app.config['DATABASE'], strQuery, tuple(dataList))
            if bSuc:
                return len(strPointNameList)
            return None
        except Exception as e:
            logging.error("ERROR in setRealtimeDataToDcimInputTable: %s" % e.__str__())
            return None

    def setRealtimeDataToMoxaInputTable(self, strPointNameList, strPointValueList):
        if not len(strPointNameList):
            return False

        try:
            strDel = 'delete from realtimedata_input_moxa_tcp_server where pointname in (%s)'
            strIns = "INSERT INTO realtimedata_input_moxa_tcp_server VALUES (%s, %s, %s)"

            paramDel = [(pointName,) for pointName in strPointNameList]

            paramIns = [(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pointName, strPointValueList[idx])
                        for idx, pointName in enumerate(strPointNameList)]

            bSuc = self._mysqlDbContainerRealTimePool.op_db_transaction_update_many(app.config['DATABASE'], [strDel, strIns],
                                                                        [tuple(paramDel), tuple(paramIns)])
            if bSuc:
                return len(strPointNameList)
            return None

        except Exception as e:
            logging.error("ERROR in setRealtimeDataToMoxaInputTable: %s" % e.__str__())

            return None


    def setRealtimeDataToPersayControllerTable(self, strPointNameList, strPointValueList):
        if len(strPointNameList) <= 0:
            return dict(err=1, msg='no points')
        dbname = app.config['DATABASE']
        nDecimal = app.config['DECIMAL']
        if nDecimal is None:
            nDecimal = 2

        strNewPointValueList = []
        for pv in strPointValueList:
            if isinstance(pv, float):
                pv = round(pv, nDecimal)
            strNewPointValueList.append(pv)
        q = 'replace into realtimedata_input_persagy_controller values(%s, %s, %s)'
        alldata = []
        nIndex = 0
        tNow = datetime.now()
        strTimeNow = tNow.strftime('%Y-%m-%d %H:%M:%S')
        for nIndex in range(len(strPointNameList)):
            alldata.append(tuple([strTimeNow, strPointNameList[nIndex], str(strPointValueList[nIndex])]))
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainerRealTimePool.op_db_update_many(dbname, q, params)

        if bSuccess:
            return dict(err=0, msg='成功写入persagy_input表', data=len(strPointNameList))
        else:
            return dict(err=1, msg='core服务器写入persagy_input表报错', data=0)


    def delFromPersagyControllerOutputTable(self, pList):
        dbname = app.config['DATABASE']
        strQuery = 'delete from realtimedata_output_persagy_controller where pointname = %s'
        alldata = []
        nIndex = 0
        for item in pList:
            alldata.append(tuple([item.get('name')]))
            nIndex += 1
        params = tuple(alldata)
        bSuccess = self._mysqlDbContainer.op_db_update_many(dbname, strQuery, params)
        return bSuccess

    def clearPersagyOutputTable(self):
        dbname = app.config['DATABASE']
        strQuery = 'delete from realtimedata_output_persagy_controller'
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, strQuery)
        return bSuccess

    def getPointsFromPersagyControllerOutputTable(self):
        pointList = []
        strSQL = 'select pointname, pointvalue from realtimedata_output_persagy_controller'
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strSQL)

        if items is None:
            return None

        for item in items:
            pointList.append(dict(name=item[0], value=item[1]))
        return pointList

    def getPointsFromObixOutputTable(self, pointList):
        strQuery = "SELECT pointname, pointvalue FROM realtimedata_output_obix"
        pointsToWrite = []
        try:
            itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if not itemList:
                return []
            if not isinstance(itemList, list):
                return []
            if not len(itemList):
                return []
            for item in itemList:
                if item[0] in pointList:
                    pointsToWrite.append(dict(name=item[0], value=item[1]))

            self.deleteFromObixOutputTable(pointsToWrite)

            return pointsToWrite
        except Exception as e:
            logging.error("ERROR in getPointsFromObixOutputTable: %s" % e.__str__())
            return []

    def getPointsFromLogixOutputTable(self):
        strQuery = "SELECT pointname, pointvalue FROM realtimedata_output_logix"
        pointsToWrite = []
        try:
            itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if not itemList:
                return []
            if not isinstance(itemList, list):
                return []
            if not len(itemList):
                return []
            for item in itemList:
                pointsToWrite.append([item[0], item[1]])

            self.deleteFromLogixOutputTable(pointsToWrite)

            return pointsToWrite
        except Exception as e:
            logging.error("ERROR in getPointsFromLogixOutputTable: %s" % e.__str__())
            return []

    def getPointsFromDlt645OutputTable(self):
        strQuery = "SELECT pointname, pointvalue FROM realtimedata_output_dlt645"
        pointsToWrite = []
        try:
            itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if not itemList:
                return []
            if not isinstance(itemList, list):
                return []
            if not len(itemList):
                return []
            for item in itemList:
                pointsToWrite.append([item[0], item[1]])

            self.deleteFromDlt645OutputTable(pointsToWrite)

            return pointsToWrite
        except Exception as e:
            logging.error("ERROR in getPointsFromDlt645OutputTable: %s" % e.__str__())
            return []

    def getPointsFromDcimOutputTable(self):
        strQuery = "SELECT pointname, pointvalue FROM realtimedata_output_dcim"
        pointsToWrite = []
        try:
            itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if not itemList:
                return []
            if not isinstance(itemList, list):
                return []
            if not len(itemList):
                return []
            for item in itemList:
                pointsToWrite.append([item[0], item[1]])

            self.deleteFromDcimOutputTable(pointsToWrite)

            return pointsToWrite
        except Exception as e:
            logging.error("ERROR in getPointsFromDcimOutputTable: %s" % e.__str__())
            return []

    def getPointsFromAbslcOutputTable(self):
        strQuery = "SELECT pointname, pointvalue FROM realtimedata_output_abslc"
        pointsToWrite = []
        try:
            itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if not itemList:
                return []
            if not isinstance(itemList, list):
                return []
            if not len(itemList):
                return []
            for item in itemList:
                pointsToWrite.append([item[0], item[1]])

            self.deleteFromAbslcOutputTable(pointsToWrite)

            return pointsToWrite
        except Exception as e:
            logging.error("ERROR in getPointsFromAbslcOutputTable: %s" % e.__str__())
            return []

    def getPointsFromKnxOutputTable(self, writablePointList):
        strQuery = "SELECT pointname, pointvalue FROM realtimedata_output_knx"
        pointsToWrite = []
        try:
            itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if not itemList:
                return []
            if not isinstance(itemList, list):
                return []
            if not len(itemList):
                return []
            for item in itemList:
                if not len(writablePointList):
                    pointsToWrite.append([item[0], item[1]])
                    continue

                if item[0] in writablePointList:
                    pointsToWrite.append([item[0], item[1]])

            self.deleteFromKnxOutputTable(pointsToWrite)

            return pointsToWrite
        except Exception as e:
            logging.error("ERROR in getPointsFromKnxOutputTable: %s" % e.__str__())
            return []

    def deleteFromObixOutputTable(self, points):
        try:
            params = []
            for point in points:
                params.append((point.get("name"), point.get("value")))

            strQuery = "DELETE FROM realtimedata_output_obix WHERE pointname = %s AND pointvalue = %s"

            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, parameter=tuple(params))

            return bSuc

        except Exception as e:
            logging.error("ERROR in deleteFromObixOutputTable: %s" % e.__str__())
            return False

    def deleteFromLogixOutputTable(self, points):
        try:

            strQuery = "DELETE FROM realtimedata_output_logix WHERE pointname = %s AND pointvalue = %s"

            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, parameter=tuple(points))

            return bSuc

        except Exception as e:
            logging.error("ERROR in deleteFromLogixOutputTable: %s" % e.__str__())
            return False

    def deleteFromAbslcOutputTable(self, points):
        try:
            strQuery = "DELETE FROM realtimedata_output_abslc WHERE pointname = %s AND pointvalue = %s"
            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, parameter=tuple(points))
            return bSuc
        except Exception as e:
            logging.error("ERROR in deleteFromAbslcOutputTable: %s" % e.__str__())
            return False

    def deleteFromKnxOutputTable(self, points):
        try:
            strQuery = "DELETE FROM realtimedata_output_knx WHERE pointname = %s AND pointvalue = %s"

            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, parameter=tuple(points))

            return bSuc

        except Exception as e:
            logging.error("ERROR in deleteFromKnxOutputTable: %s" % e.__str__())
            return False

    def deleteFromDcimOutputTable(self, points):
        try:
            strQuery = "DELETE FROM realtimedata_output_dcim WHERE pointname = %s AND pointvalue = %s"
            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, parameter=tuple(points))
            return bSuc
        except Exception as e:
            logging.error("ERROR in deleteFromDcimOutputTable: %s" % e.__str__())
            return False

    def deleteFromDlt645OutputTable(self, points):
        try:
            strQuery = "DELETE FROM realtimedata_output_dlt645 WHERE pointname = %s AND pointvalue = %s"
            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, parameter=tuple(points))
            return bSuc
        except Exception as e:
            logging.error("ERROR in deleteFromDlt645OutputTable: %s" % e.__str__())
            return False

    def getOperationLogicRecordOutput(self, pointName, timeFrom, timeTo):
        strQuery = "SELECT pointtime, logicname, pointname, pointvalue FROM logic_output_point_record " \
                   "WHERE pointname = %s AND pointtime >= %s AND pointtime <= %s ORDER BY pointtime DESC"
        itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(pointName, timeFrom, timeTo))
        if not itemList:
            return dict(code=1, data=[], msg="数据获取失败")
        if not len(itemList):
            return dict(code=1, data=[], msg="获取到的数据为空")
        dataList = [dict(time=item[0], logicName=item[1], pointName=item[2], value=item[3]) for item in itemList]
        return dict(code=0, data=dataList, msg="获取成功")

    def getOperationLogicRecordOutputV2(self, pointName, tTimeFrom, tTimeTo, opRecordDir):
        strMsg = ""
        rtList = []

        for root, dirs, strFileNames in os.walk(opRecordDir):
            for strFileName in strFileNames:

                if not strFileName.endswith(".txt"):
                    continue

                patPoint = r"^" + pointName + "\.txt$"
                patSource = r"^Source_.*\.txt$"
                if re.match(patPoint, strFileName):
                    tFrom = datetime.now()
                    with open(os.path.join(opRecordDir, strFileName), "r", encoding="UTF8", errors="ignore") as fo:
                        lines = fo.readlines()
                    for line in lines:
                        if line == "\n":
                            continue

                        if not line:
                            continue

                        try:
                            items = line.replace("\n", "").split("\t")
                            validObjectList = [obj for obj in items if len(obj)]
                            if len(validObjectList) < 3:
                                continue

                            strTime = validObjectList[0].strip()
                            operator = validObjectList[1].strip()
                            value = validObjectList[2].strip()

                            tTime = datetime.strptime(strTime, "%Y-%m-%d %H:%M:%S")
                            if tTime >= tTimeFrom and tTime <= tTimeTo:
                                rtList.append(dict(time=strTime, logicName=operator, pointName=pointName, value=value))
                        except:
                            pass

                    logging.error("getOperationLogicRecordOutputV2::pointName::%s::tDelta:%s" % (strFileName, (datetime.now() - tFrom).total_seconds()))

                elif re.match(patSource, strFileName):
                    tFrom = datetime.now()
                    source = strFileName.replace("Source_", "").replace(".txt", "")
                    with open(os.path.join(opRecordDir, strFileName), "r", encoding="UTF8", errors="ignore") as fo:
                        lines = fo.readlines()
                    for line in lines:
                        if line == "\n":
                            continue

                        if not line:
                            continue

                        try:
                            items = line.replace("\n", "").split("\t")
                            validObjectList = [obj for obj in items if len(obj)]
                            if len(validObjectList) < 3:
                                continue

                            strTime = validObjectList[0].strip()
                            linePointName = validObjectList[1].strip()
                            value = validObjectList[2].strip()

                            tTime = datetime.strptime(strTime, "%Y-%m-%d %H:%M:%S")

                            if linePointName != pointName:
                                continue

                            if tTime >= tTimeFrom and tTime <= tTimeTo:
                                rtList.append(dict(time=strTime, logicName=source, pointName=linePointName, value=value))
                        except:
                            pass

                    logging.error("getOperationLogicRecordOutputV2::source_::%s::tDelta:%s" % (strFileName, (datetime.now() - tFrom).total_seconds()))

        return dict(code=0, msg="获取成功" if not len(strMsg) else strMsg, data=rtList)

    def clearLogixOutputTable(self):
        dbname = app.config['DATABASE']
        strQuery = 'delete from realtimedata_output_logix'
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, strQuery)
        return bSuccess

    def clearAbslcOutputTable(self):
        dbname = app.config['DATABASE']
        strQuery = 'delete from realtimedata_output_abslc'
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, strQuery)
        return bSuccess

    def clearKnxOutputTable(self):
        dbname = app.config['DATABASE']
        strQuery = 'delete from realtimedata_output_knx'
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, strQuery)
        return bSuccess

    def clearDlt645OutputTable(self):
        dbname = app.config['DATABASE']
        strQuery = 'delete from realtimedata_output_dlt645'
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, strQuery)
        return bSuccess

    def clearDcimOutputTable(self):
        dbname = app.config['DATABASE']
        strQuery = 'delete from realtimedata_output_dcim'
        bSuccess = self._mysqlDbContainer.op_db_update(dbname, strQuery)
        return bSuccess

    def getFaultStatusInfo(self, strTimeFrom, strTimeTo):
        try:

            strQuery = "SELECT fddName, status, modifyTime, orderId, opContentData FROM fdd_work_order WHERE createTime>=%s AND createTime<=%s"

            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(strTimeFrom, strTimeTo))

            rt = {}
            for item in items:
                if item[3] not in rt.keys():
                    rt.update({item[3]: [dict(fddName=item[0], status=item[1], modifyTime=item[2], opContent=item[4])]})
                else:
                    rt.get(item[3]).append(dict(fddName=item[0], status=item[1], modifyTime=item[2], opContent=item[4]))

            return dict(code=0, data=rt)

        except Exception as e:

            logging.error("ERROR in getFddStatusDict: %s" % e.__str__())
            return dict(code=1, data={})

    # 查询在时间范围内发生过更改的工单
    def getFaultModifiedInTimeRange(self, strBeginTime, strEndTime):
        strQuery = "SELECT orderId FROM fdd_work_order WHERE modifyTime >= %s AND modifyTime <= %s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (strBeginTime, strEndTime))
        if items == None:
            return None

        nIdList = []
        for item in items:
            if len(item):
                if item[0] not in nIdList:
                    nIdList.append(item[0])
        return nIdList


    def getFaultInfoOfCertainOrderId(self, nOrderIdList):
        try:
            if not len(nOrderIdList):
                return {}

            strQuery = "SELECT fddName, opUserName, opType, opContentData, orderId, modifyTime, detail, status, title, ownUser, img FROM fdd_work_order WHERE orderId in (" + ",".join([str(nId) for nId in nOrderIdList]) + ")"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)

            if items is None:
                return None
            elif not len(items):
                return {}

            rt = {}
            for item in items:
                nOrderId = item[4]
                if nOrderId not in rt.keys():
                    rt.update({nOrderId: []})
                rt[nOrderId].append(dict(fddName=item[0], opUserName=item[1], opType=item[2], opContentData=item[3],
                                                modifyTime=item[5], detail=item[6], status=item[7], name=item[8], owner=item[9], img=item[10]))

            return rt
        except Exception as e:
            logging.error("ERROR in getFaultInfo: %s" % e.__str__())
            return {}


    def getFaultInfo(self, startTime=None, endTime=None, orderId=None):
        try:
            if startTime == None and endTime == None and orderId == None:
                return dict(code=1, msg="startTime, endTime, orderId都为None", data={})

            strQuery = "SELECT fddName, opUserName, opType, opContentData, orderId, modifyTime, detail, status, title, ownUser, img FROM fdd_work_order WHERE modifyTime >= %s AND modifyTime <= %s"
            param = (startTime, endTime)
            if orderId != None:
                if isinstance(orderId, int):
                    strQuery = "SELECT fddName, opUserName, opType, opContentData, orderId, modifyTime, detail, status, title, ownUser, img FROM fdd_work_order WHERE orderId=%s"
                    param = (orderId,)
                elif isinstance(orderId, list) and len(orderId):
                    strQuery += " OR orderId in (" + ",".join([str(nId) for nId in orderId]) + ")"

            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=param)

            if items is None:
                return dict(code=1, msg="获取失败", data={})
            elif not len(items):
                return dict(code=0, msg="无数据", data={})

            rt = {}
            for item in items:
                if item[4] not in rt.keys():
                    rt.update({item[4]: [dict(fddName=item[0], opUserName=item[1], opType=item[2], opContentData=item[3],
                                              modifyTime=item[5], detail=item[6], status=item[7], name=item[8], owner=item[9], img=item[10])]})
                else:
                    rt.get(item[4]).append(dict(fddName=item[0], opUserName=item[1], opType=item[2], opContentData=item[3],
                                                modifyTime=item[5], detail=item[6], status=item[7], name=item[8], owner=item[9], img=item[10]))
            return dict(code=0, msg="", data=rt)

        except Exception as e:
            logging.error("ERROR in getFaultInfo: %s" % e.__str__())
            return dict(code=1, msg="", data={})

    def getAllFaultOrderIds(self):
        try:
            strQuery = "SELECT DISTINCT orderId FROM fdd_work_order"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            rt = []
            if items is None:
                return dict(code=1, msg="获取失败", data=[])
            elif not len(items):
                return dict(code=0, msg="无数据", data=[])
            else:
                for item in items:
                    rt.append(item[0])
            return dict(code=0, msg="", data=rt)
        except Exception as e:
            logging.error("ERROR in getAllFaultOrderIds: %s" % e.__str__())
            return dict(code=1, msg="", data=[])

    def getFddNameAndCurrentStatus(self, orderId):
        try:
            # strQuery = "SELECT fwo.fddName,fwo.`status` FROM (SELECT fddName, max(modifyTime) AS mmt FROM fdd_work_order GROUP BY fddName) AS B LEFT JOIN fdd_work_order fwo on fwo.modifyTime=B.mmt AND B.fddName=fwo.fddName where fwo.orderId=%s"
            strQuery = "SELECT fddName, modifyTime, status FROM fdd_work_order WHERE orderId=%s"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
            if items is None:
                return dict(code=1, msg="获取失败", data=(None, None))
            if not len(items):
                return dict(code=1, msg="无数据", data=(None, None))
            lTar = items[0]
            for item in items:
                if item[1] > lTar[1]:
                    lTar = item
            return dict(code=0, msg="", data=(lTar[0], lTar[2]))

        except Exception as e:
            logging.error('ERROR in getFddNameAndCurrentStatus: %s' % e.__str__())
            return dict(code=1, data=None)

    def getWorkOrderCurrentEstimatedTime(self, orderId):
        strQuery = "SELECT opContentData FROM fdd_work_order WHERE orderId=%s ORDER BY modifyTime ASC"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
        if items is None:
            return dict(code=1, msg="获取失败", data=None)
        if not len(items):
            return dict(code=1, msg="无数据", data=None)

        estimatedTime = None
        for item in items:
            if not item[0]:
                continue
            if not isinstance(item[0], str):
                continue

            strItem = eval(item[0])
            if strItem.get("estimatedTime", None):
                estimatedTime = datetime.strptime(strItem.get("estimatedTime"), "%Y-%m-%d %H:%M:%S")
        return dict(code=0, msg="", data=estimatedTime)

    def getFddNameAndCurrentStatusByFddName(self, fddName):
        try:
            strQuery = "SELECT fwo.fddName,fwo.`status` FROM (SELECT fddName, max(modifyTime) AS mmt FROM fdd_work_order GROUP BY fddName) AS B LEFT JOIN fdd_work_order fwo on fwo.modifyTime=B.mmt AND B.fddName=fwo.fddName where fwo.fddName=%s"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(fddName,))
            if items is None:
                return dict(code=1, msg="获取失败", data=None)
            elif not len(items):
                return dict(code=0, msg="无数据", data=(fddName, -2))   # 表示无此故障
            else:
                return dict(code=0, data=items[0])
        except Exception as e:
            logging.error('ERROR in getFddNameAndCurrentStatus: %s' % e.__str__())
            return dict(code=1, data=None)

    def getStartTimeOfFault(self, orderId):
        try:
            strQuery = "select opContentData, modifyTime from fdd_work_order where opContentData like '%changeTo%' and orderId = %s"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
            if items is None:
                return None
            elif not len(items):
                return None
            else:
                tTime = items[0][1]
                for item in items:
                    if eval(item[0]).get("changeTo") == 1 and item[1] > tTime:
                        tTime = item[1]
                return tTime

        except Exception as e:
            logging.error("ERROR in getStartTimeOfFault: %s" % e.__str__())
            return None

    def getCreateTimeOfOrder(self, orderId):
        try:
            strQuery = "SELECT modifyTime FROM fdd_work_order WHERE orderId=%s AND opType=-1"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
            if items is None:
                return None
            elif not len(items):
                return None
            elif not len(items[0]):
                return None
            return items[0][0]
        except Exception as e:
            logging.error("ERROR in getCreateTimeOfOrder: %s" % e.__str__())
            return None


    def calcFaultDuration(self, orderId):
        try:
            # strQuery = "SELECT opContentData, modifyTime FROM fdd_work_order WHERE opContentData like '%changeTo%' and orderId =%s"
            strQuery = "SELECT opContentData, modifyTime, opType FROM fdd_work_order WHERE orderId =%s ORDER BY modifyTime ASC"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
            if items is None:
                return None
            elif not len(items):
                return None

            eventList = []
            tInitial = items[0][1]
            for item in items:
                if not len(item):
                    continue
                if item[0] == None:
                    continue

                dOpContent = eval(item[0])
                if not isinstance(dOpContent, dict):
                    continue
                if dOpContent.get("changeTo", None) is not None:
                    eventList.append(
                        (dOpContent.get("changeTo"), item[1])
                    )

            eventList.sort(key=lambda x: x[1], reverse=False)

            tStart = None
            tEnd = None
            durationSeconds = 0
            for event in eventList:
                if event[0] == 1:
                    tStart = event[1]
                elif event[0] == 4:
                    tEnd = event[1]
                    if tStart and tEnd:
                        durationSeconds += (tEnd - tStart).total_seconds()
                        tStart = None
                        tEnd = None

            if tEnd == None:
                tEnd = datetime.now()
            if tStart == None:
                tStart = tInitial

            durationSeconds += (tEnd - tStart).total_seconds()

            return durationSeconds
        except Exception as e:
            strLog = "ERROR in calcFaultDuration: %s" % e.__str__()
            logging.error(strLog)
            return None

    def enableFault(self, faultLabelPoint, nEnable):
        try:

            pDataList, pDataDict = self.getInputTable([faultLabelPoint])

            if faultLabelPoint not in pDataDict.keys():
                return dict(code=1, msg="未找到FaultLabel点", data=False)

            strInfo = pDataDict.get(faultLabelPoint, "")

            jsonInfo = {}
            try:
                jsonInfo = json.loads(strInfo)
            except:
                pass

            dInfo = {} if not isinstance(jsonInfo, dict) else jsonInfo

            dInfo.update({"Enabled": nEnable})

            nErr = self.setRealtimeData([faultLabelPoint], [json.dumps(dInfo, ensure_ascii=False)]).get("err", 1)
            if nErr == 1:
                return dict(code=1, msg="故障状态更新失败", data=False)

            count = 0
            bSuc = False
            while count < 3:
                pDataList, pDataDict = self.getInputTable([faultLabelPoint])
                strInfo = pDataDict.get(faultLabelPoint, "")
                jsonInfo = json.loads(strInfo)
                nResEnabled = jsonInfo.get("Enabled", None)
                if nResEnabled == nEnable:
                    bSuc = True
                    break

                count += 1
                time.sleep(2)

            if not bSuc:
                return dict(code=1, msg="更新失败，未检测到修改后的故障状态", data=False)

            return dict(code=0, msg="更新成功", data=True)

        except Exception as e:
            strLog = 'ERROR in enableFault: %s' % e.__str__()
            logging.error(strLog)
            return dict(code=1, msg="更新失败：%s" % strLog, data=False)

    def getFaultLatestProcessor(self, orderId):
        strQuery = "SELECT opContentData FROM fdd_work_order WHERE opContentData like '%toUserName%' AND orderId=%s ORDER BY modifyTime ASC"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
        if items is None:
            return ""
        elif not len(items):
            return ""

        processor = ""
        for item in items:
            if not len(item):
                continue
            if item[0] is None:
                continue
            strItem = eval(item[0])
            if strItem.get("toUserName", None):
                processor = strItem.get("toUserName", "")
        return processor

    def processFault(self, fddName, orderId, userName, curStatus, targetStatus, processor, estimatedTime, reason,
                     opType, faultLabelPoint, adminUserName, curFaultOwner, strMaintainanceFaultPointName, faultCreator, role):
        try:
            modifyTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            createTime = BEOPDataAccess.getInstance().getCreateTimeOfOrder(orderId)
            if not createTime:
                return dict(code=1, data=False, msg="查询工单的创建时间失败")

            if curStatus == 0 and targetStatus == 1 and opType == 0 and role == "admin":
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, 0, %s, %s, %s, %s, %s, %s)"
                param = (fddName, userName, str({"changeTo": 1, "toUserName": processor, "estimatedTime": estimatedTime}), orderId, modifyTime, targetStatus, createTime, processor)
            elif curStatus == 1 and targetStatus == 1 and opType == 0 and role == "admin":
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, 0, %s, %s, %s, %s, %s, %s)"
                param = (fddName, userName, str({"changeTo": 1, "toUserName": processor, "estimatedTime": estimatedTime}),orderId, modifyTime, targetStatus, createTime, processor)
            elif curStatus == 4 and targetStatus == 1 and opType == 0 and role == "admin":
                strProcessor = BEOPDataAccess.getInstance().getFaultLatestProcessor(orderId)
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, 0, %s, %s, %s, %s, %s, %s)"
                param = (fddName, userName, str({"changeTo": 1}), orderId, modifyTime, targetStatus, createTime, strProcessor)
            elif curStatus in [0, 1] and targetStatus == 4 and opType == 0 and role == "admin":
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) values (%s, %s, 0, %s, %s, %s, %s, %s, %s)"
                param = (fddName, userName, str({"changeTo": 4, "reason": reason}), orderId, modifyTime, targetStatus, createTime, adminUserName)
            elif curStatus == 2 and targetStatus == 3 and opType == 0 and role == "admin":
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime) VALUES (%s, %s, 0, %s, %s, %s, %s, %s)"
                param = (fddName, userName, str({"changeTo": 3}), orderId, modifyTime, targetStatus, createTime)
            elif curStatus == 1 and targetStatus == 1 and opType == 1 and role == "admin":
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, 1, %s, %s, %s, %s, %s, %s)"
                param = (fddName, userName, str({"estimatedTime": estimatedTime}), orderId, modifyTime, targetStatus, createTime, curFaultOwner)
            elif curStatus == 1 and targetStatus == 2 and opType == 0 and role == "executor":
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, 0, %s, %s, %s, %s, %s, %s)"
                tStartTime = BEOPDataAccess.getInstance().getStartTimeOfFault(orderId)
                if not tStartTime:
                    logging.error("工单(id:%s)无开始时间" % orderId)
                    tStartTime = None
                    # return dict(code=1, msg="该工单无开始时间", data=False)

                durationSeconds = BEOPDataAccess.getInstance().calcFaultDuration(orderId)
                if durationSeconds is None:
                    return dict(code=1, msg="耗时计算失败", data=False)

                param = (fddName, userName, str({"changeTo": 2, "duration": durationSeconds, "conclusion": reason}), orderId, modifyTime, targetStatus, createTime, adminUserName)
            elif targetStatus == 5 and opType == 0 and role == "admin":
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime) VALUES (%s, %s, %s,%s,%s,%s,%s,%s)"
                param = (fddName, userName, opType, str({"changeTo": 5}), orderId, modifyTime, targetStatus, createTime)

            elif curStatus == 0 and targetStatus == 2 and opType == 0 and faultCreator == userName:
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, 0, %s, %s, %s, %s, %s, %s)"
                tStartTime = BEOPDataAccess.getInstance().getStartTimeOfFault(orderId)
                if not tStartTime:
                    logging.error("工单(id:%s)无开始时间" % orderId)
                    tStartTime = None

                durationSeconds = BEOPDataAccess.getInstance().calcFaultDuration(orderId)
                if durationSeconds is None:
                    logging.error("工单(id:%s)计算耗时失败" % orderId)
                    durationSeconds = None

                param = (fddName, userName, str({"changeTo": 2, "duration": durationSeconds, "conclusion": reason}), orderId, modifyTime, targetStatus,createTime, adminUserName)

            elif curStatus in [0, 1, 2, 4] and targetStatus == 3 and opType == 0 and role == "admin":
                strOwner = ""
                strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"

                dOpContent = {"changeTo": 3, "conclusion": reason}
                durationSeconds = BEOPDataAccess.getInstance().calcFaultDuration(orderId)
                if durationSeconds is None:
                    logging.error("工单(id:%s)计算耗时失败" % orderId)

                dOpContent.update({"duration": durationSeconds})

                param = (fddName, userName, opType, str(dOpContent), orderId, modifyTime, targetStatus, createTime, strOwner)

            else:
                return dict(code=1, msg="根据工单的当前状态，该操作无效", data=False)

            bSucOpWorkOrder = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, parameter=param)
            if not bSucOpWorkOrder:
                return dict(code=1, msg="更新工单状态失败", data=False)

            if targetStatus == 2 and fddName != strMaintainanceFaultPointName:
                pDataList, pDataDict = self.getInputTable([faultLabelPoint])
                if faultLabelPoint not in pDataDict.keys():
                    return dict(code=1, msg="未找到FaultLabel点，所以故障标记内容未更新，但工单状态已更新", data=False)
                strInfo = pDataDict.get(faultLabelPoint, "")
                jsonInfo = {}
                try:
                    jsonInfo = json.loads(strInfo)
                except:
                    pass

                dInfo = {} if not isinstance(jsonInfo, dict) else jsonInfo

                if "Reason" not in dInfo.keys():
                    dInfo.update({"Reason": {}})

                dInfo.get("Reason").update({datetime.now().strftime("%Y-%m-%d"): reason})

                nErr = self.setRealtimeData([faultLabelPoint], [json.dumps(dInfo, ensure_ascii=False)]).get("err", 1)
                if nErr == 1:
                    return dict(code=1, msg="更新故障标记内容失败，但工单状态已更新", data=False)

                count = 0
                bSuc = False
                while count < 3:
                    pDataList, pDataDict = self.getInputTable([faultLabelPoint])
                    strInfo = pDataDict.get(faultLabelPoint, "")
                    dInfoGet = json.loads(strInfo)

                    if dInfoGet == dInfo:
                        bSuc = True
                        break

                    count += 1
                    time.sleep(2)

                if not bSuc:
                    return dict(code=1, msg="更新失败，未检测到故障总结内容的更新", data=False)

            return dict(code=0, msg="更新成功", data=True)

        except Exception as e:
            strLog = 'ERROR in processFault: %s' % e.__str__()
            logging.error(strLog)
            return dict(code=1, msg="更新失败：%s" % strLog, data=False)

    def addFault(self, fddName, newOrderId, strOrderName, detail):
        try:
            fddNameAndCurrStatus = BEOPDataAccess.getInstance().getFddNameAndCurrentStatusByFddName(fddName)
            if fddNameAndCurrStatus.get("code") > 0:
                return dict(code=1, msg="数据查询失败", data=False)

            if fddNameAndCurrStatus.get("data")[1] in [0, 1, 2, 4]:
                return dict(code=1, msg="工单已存在", data=False)

            modifyTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            strInsert = "INSERT INTO fdd_work_order (fddName, opUserName, opType, orderId, modifyTime, title, detail, status, createTime) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
            bSuccess = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strInsert, parameter=(fddName, 'system', -1, newOrderId, modifyTime, strOrderName, detail, 0, modifyTime))

            return dict(code=0 if bSuccess else 1, msg="", data=bSuccess)
        except Exception as e:
            logging.error('ERROR in addFault: %s' % e.__str__())
            return dict(code=1, msg="添加失败", data=False)

    def addMaintainanceFault(self, creator, processor, name, detail, estimatedTime, orderId, MaintainanceFaultPointName, adminUserName, strImg):
        try:
            modifyTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            strCreate = "INSERT INTO fdd_work_order (fddName, opUserName, opType, orderId, modifyTime, title, detail, status, createTime, ownUser, img) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"
            paramCreate = (MaintainanceFaultPointName, creator, -1, orderId, modifyTime, name, detail, 0, modifyTime, adminUserName, strImg)

            bCreate = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strCreate, parameter=paramCreate)
            if not bCreate:
                return dict(code=1, msg="创建失败", data=False)

            strAssign = None
            paramAssign = None
            if processor and not estimatedTime:
                tAssignTime = datetime.strptime(modifyTime, "%Y-%m-%d %H:%M:%S") + timedelta(seconds=1)
                strAssignTime = tAssignTime.strftime("%Y-%m-%d %H:%M:%S")
                strAssign = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
                paramAssign = (MaintainanceFaultPointName, creator, 0, str({"changeTo": 1, "toUserName": processor}), orderId, strAssignTime, 1, modifyTime, processor)

            elif processor and estimatedTime:
                tAssignTime = datetime.strptime(modifyTime, "%Y-%m-%d %H:%M:%S") + timedelta(seconds=1)
                strAssignTime = tAssignTime.strftime("%Y-%m-%d %H:%M:%S")
                strAssign = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
                paramAssign = (MaintainanceFaultPointName, creator, 0, str({"changeTo": 1, "toUserName": processor, "estimatedTime": estimatedTime}), orderId, strAssignTime, 1, modifyTime, processor)

            # 只有预计完成时间，没有指派人
            elif estimatedTime and not processor:
                tAssignTime = datetime.strptime(modifyTime, "%Y-%m-%d %H:%M:%S") + timedelta(seconds=1)
                strAssignTime = tAssignTime.strftime("%Y-%m-%d %H:%M:%S")
                strAssign = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
                paramAssign = (MaintainanceFaultPointName, creator, 0, str({"estimatedTime": estimatedTime}), orderId, strAssignTime, 0, modifyTime, processor)

            bAssign = None
            if strAssign and paramAssign:
                bAssign = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strAssign, parameter=paramAssign)

            if bAssign == False:
                return dict(code=0, msg="创建成功但分派失败，请手动分派", data=True)
            if bAssign == True:
                return dict(code=0, msg="创建且分派成功", data=True)

            return dict(code=0, msg="创建成功", data=True)

        except Exception as e:
            logging.error('ERROR in addMaintainanceFault: %s' % e.__str__())
            return dict(code=1, msg="创建失败", data=False)

    def editWorkOrder(self, orderId, orderName, detail, processor, estimatedTime, userName, fddName, curStatus):

        nameDetailList = []
        if len(detail):
            nameDetailList.append("detail='{0}'".format(detail))
        if len(orderName):
            nameDetailList.append("title='{0}'".format(orderName))

        strUpdate = ""
        if len(nameDetailList):
            strUpdate = "UPDATE fdd_work_order SET " + ",".join(nameDetailList) + " WHERE orderId=" + str(orderId) + " AND opType=-1"

        errList = []
        if len(strUpdate):
            bUpdate = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strUpdate)
            if not bUpdate:
                errList.append("故障名称/详情更新失败")

        if curStatus != 1 and not len(processor) and not len(estimatedTime):
            msg = "修改成功"
            if len(errList):
                msg = ";".join(errList)
            return dict(err=0 if not len(errList) else 1, msg=msg, data=False if len(errList) else True)

        if curStatus != 1 and (len(processor) or len(estimatedTime)):
            errList.append("当前工单状态不是进行中，故无法修改指派人和预计完成时间（但名称和详情可以修改）")
            return dict(err=1 if len(errList) else 0, msg=errList, data=False if len(errList) else True)

        if not len(processor) and not len(estimatedTime):
            return dict(err=1 if len(errList) else 0, msg=errList, data=False if len(errList) else True)

        dUpdateInfo = {"changeTo": 1}
        if len(processor):
            dUpdateInfo.update({"toUserName": processor})
        if len(estimatedTime):
            dUpdateInfo.update({"estimatedTime": estimatedTime})

        modifyTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        createTime = self.getCreateTimeOfOrder(orderId)
        if not createTime:
            errList.append("获取工单创建时间失败")
            return dict(err=1 if len(errList) else 0, msg=errList, data=False if len(errList) else True)

        nOpType = 1 if len(estimatedTime) else 0
        strInsert = "INSERT INTO fdd_work_order (fddName, opUserName, opType, opContentData, orderId, modifyTime, status, createTime, ownUser) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
        param = (fddName, userName, nOpType, str(dUpdateInfo), orderId, modifyTime, curStatus, createTime, processor)

        bSucOpWorkOrder = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strInsert, parameter=param)
        if not bSucOpWorkOrder:
            errList.append("修改处理人（或预计完成时间）失败")

        return dict(err=1 if len(errList) else 0, msg=errList, data=False if len(errList) else True)



    def filter_data(self, filterType, params, timeValueList):
        try:
            rtList = []
            logicalMap = {"(": ">", "[": ">=", ")": "<", "]": "<="}

            if filterType == "==":
                for timeValue in timeValueList:
                    if timeValue[1] == params:
                        rtList.append(timeValue)

            elif filterType == "InRange" or filterType == "NotInRange":
                strStart = params.split(",")[0].strip()
                strStop = params.split(",")[1].strip()

                nStart = float(re.findall("[\d.]+", strStart.strip())[0])
                nStop = float(re.findall("[\d.]+", strStop.strip())[0])

                lLogical = logicalMap.get(strStart[0], None)
                rLogical = logicalMap.get(strStop[-1], None)

                if lLogical and rLogical:
                    for timeValue in timeValueList:

                        expression = "{value} {lLogical} {nStart} and {value} {rLogical} {nStop}".format(value=timeValue[1],
                                                                                                  lLogical=lLogical,
                                                                                                  nStart=nStart,
                                                                                                  rLogical=rLogical,
                                                                                                  nStop=nStop)
                        eExpress = eval(expression)

                        if filterType == "InRange" and eval(expression) is True:
                            rtList.append(timeValue)

                        elif filterType == "NotInRange" and eval(expression) is False:
                            rtList.append(timeValue)

            elif filterType == ">":
                for timeValue in timeValueList:
                    if timeValue[1] > float(params):
                        rtList.append(timeValue)

            elif filterType == ">=":
                for timeValue in timeValueList:
                    if timeValue[1] >= float(params):
                        rtList.append(timeValue)

            elif filterType == "<":
                for timeValue in timeValueList:
                    if timeValue[1] < float(params):
                        rtList.append(timeValue)

            elif filterType == "<=":
                for timeValue in timeValueList:
                    if timeValue[1] <= float(params):
                        rtList.append(timeValue)

            return rtList

        except Exception as exp:
            logging.error("ERROR in filter_data: %s" % exp.__str__())
            return []

    def getModeTypeDict(self):
        try:
            result = {}
            strQuery = "SELECT id, type FROM mode"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if items is None:
                return dict(code=1, msg="数据库繁忙，请重试", data={})
            for item in items:
                if item[1] not in result.keys():
                    result.update({item[1]: [item[0]]})
                else:
                    result.get(item[1]).append(item[0])
            return dict(code=0, msg="获取成功", data=result)
        except Exception as e:
            logging.error("ERROR in getModeTypeDict: %s" % e.__str__())
            return dict(code=1, msg="获取失败", data={})

    def batchDeleteMode(self, nType, modeIdList):
        try:
            strDelete1 = "DELETE FROM mode WHERE type=%s"
            strDelete2 = "DELETE FROM mode_detail WHERE modeid in (%s)" % ",".join([str(num) for num in modeIdList])
            strDelete3 = "DELETE FROM mode_calendar WHERE modeid in (%s) AND ofDate >= '%s'" % (",".join([str(num) for num in modeIdList]),
                                                                                              datetime.now().strftime("%Y-%m-%d"))

            bSuc1 = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strDelete1, parameter=(nType,))
            if not bSuc1:
                return False

            bSuc2 = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strDelete2)
            if not bSuc2:
                return False

            bSuc3 = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strDelete3)

            if not bSuc3:
                return False
            return True
        except Exception as e:
            logging.error("ERROR in batchDeleteMode: %s" % e.__str__())
            return False

    def getPointLatestTriggerTime(self, strPointName):
        strQuery = "select mc.ofDate, md.modeid, ed.envid, ed.pointname, md.triggerTime, ed.pointvalue from " \
                   "env_detail ed left join mode_detail md on ed.envid = md.triggerEnvId left join mode_calendar mc on " \
                   "md.modeid = mc.modeid where ed.pointname = %s and mc.ofDate <= %s and mc.ofDate >= curDate() order by mc.ofDate asc"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(strPointName, (datetime.now() + timedelta(days=2)).strftime("%Y-%m-%d")))

        if items is None:
            return dict(code=1, msg="数据库繁忙，请重试", data=())

        timeValue = ()
        for item in items:
            try:
                if item[0] is None or item[1] is None or item[2] is None or item[3] is None or item[4] is None or item[5] is None:
                    continue

                tDate = item[0]
                strTime = item[4]
                if not re.match("^[0-9]{2}:[0-9]{2}$", item[4]):
                    pList, pMap = self.getInputTable([item[4]])
                    strTime = pMap.get(item[4])

                value = item[5]
                tTime = datetime.strptime(strTime, "%H:%M")
                tDateTime = tDate.replace(hour=tTime.hour, minute=tTime.minute, second=0)
                timeValue = (tDateTime, value)
                break
            except Exception as e:
                logging.error("ERROR in getPointTriggerTime: %s" % e.__str__())

        return dict(code=0, msg="", data=timeValue)

    def getCoreLogicVersion(self):
        strQuery = "SELECT unitproperty01, unitproperty02, unitproperty03 FROM unit01 WHERE unitproperty01 like '%version%'"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if not items:
            return dict(code=1, msg="数据库繁忙，稍后再试", version={}, startupTime={})
        if not len(items):
            return dict(code=1, msg="查询失败", version={}, startupTime={})

        dVersion = {}
        dStartupTime = {}

        for item in items:
            strStTime = ""
            if isinstance(item[2], str):
                strStTime = item[2]
            elif isinstance(item[2], datetime):
                try:
                    strStTime = item[2].strftime("%Y-%m-%d %H:%M:%S")
                except:
                    pass
            elif item[2] == None:
                strStTime = ""

            if item[0] == "version_domcore":
                dVersion.update({"domcore": item[1].replace("V", "")})
                dStartupTime.update({"domcore": strStTime})
            elif item[0] == "version_domlogic":
                dVersion.update({"domlogic": item[1].replace("V", "")})
                dStartupTime.update({"domlogic": strStTime})
            elif item[0] == "version_domSiemenseTCPCore":
                dVersion.update({"domSiemenseTCPCore": item[1].replace("V", "")})
                dStartupTime.update({"domSiemenseTCPCore": strStTime})

            elif item[0] == "version_domlogic01":
                dStartupTime.update({"domlogic01": strStTime})
            elif item[0] == "version_domlogic02":
                dStartupTime.update({"domlogic02": strStTime})
            elif item[0] == "version_domlogic03":
                dStartupTime.update({"domlogic03": strStTime})
            elif item[0] == "version_domlogic04":
                dStartupTime.update({"domlogic04": strStTime})
            elif item[0] == "version_domlogic05":
                dStartupTime.update({"domlogic05": strStTime})

        return dict(code=0, msg="", version=dVersion, startupTime=dStartupTime)

    def expandPadTimeValueHourlyList(self, strCurPointTimeList, strCurPointValueList):
        if len(strCurPointTimeList)<=1:
            return

        for iIndex in range(len(strCurPointTimeList)-1):
            strT1 = strCurPointTimeList[iIndex]
            strT2 = strCurPointTimeList[iIndex+1]
            tT1 = datetime.strptime(strT1, '%Y-%m-%d %H:%M:%S')
            tT2 = datetime.strptime(strT2, '%Y-%m-%d %H:%M:%S')
            if (tT2-tT1).total_seconds()>60*60*1.5:
                strCurPointTimeList.insert(iIndex+1, (tT1+timedelta(hours=1)).strftime('%Y-%m-%d %H:%M:%S'))
                strCurPointValueList.insert(iIndex+1, strCurPointValueList[iIndex])

    def updateRedisDataHoursOfMonth(self, tDate):
        data = {}
        tablename = 'historydata_hour_' + tDate.strftime('%Y_%m')

        try:
            r = None
            tLast3DaysBefore = tDate - timedelta(days=7)
            strQ = 'SELECT pointname, time, value FROM ' + tablename + " where time>='"+ tLast3DaysBefore.strftime('%Y-%m-%d %H:%M:%S') +"' order by pointname"
            r = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)

            strCurPointTimeList = []
            strCurPointValueList = []
            strCurPointName = None
            if r is not None:
                for x in r:
                    strPointName = x[0]
                    if strPointName is None or strPointName=='':
                        continue

                    strValue = x[2]
                    try:
                        tTime = x[1]
                        tTime = tTime.replace(second=0, microsecond=0)
                    except:
                        continue

                    if strPointName!=strCurPointName and strCurPointName is not None:
                        self.expandPadTimeValueHourlyList(strCurPointTimeList, strCurPointValueList)
                        RedisManager.set_history_hour_data_list_for_one_point(strCurPointName, strCurPointTimeList, strCurPointValueList)

                        strCurPointTimeList = [tTime.strftime('%Y-%m-%d %H:%M:%S')]
                        strCurPointValueList = [strValue]
                        strCurPointName = strPointName
                    else:
                        if strCurPointName is None:
                            strCurPointName = strPointName
                        strCurPointTimeList.append(tTime.strftime('%Y-%m-%d %H:%M:%S'))
                        strCurPointValueList.append(strValue)

                if strCurPointName and strCurPointTimeList and strCurPointValueList:
                        RedisManager.set_history_hour_data_list_for_one_point(strCurPointName, strCurPointTimeList, strCurPointValueList)



        except Exception as e:
            strError = 'ERROR in updateRedisDataMinutesOfDay: ' + e.__str__()
            print(strError)
            logging.error(strError)
            return False

        return True

    def updateRedisDataMinutesOfDay(self, tDate):
        data = {}
        tablename = 'historydata_minute_' + tDate.strftime('%Y_%m_%d')

        try:
            r = None
            strQ = 'SELECT pointname, time, value FROM ' + tablename + " order by time"
            r = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)

            strCurPointNameList = []
            strCurPointValueList = []
            tCurTime = None
            if r is not None:
                for x in r:
                    strPointName = x[0]
                    strValue = x[2]
                    try:
                        tTime = x[1]
                        tTime = tTime.replace(second=0, microsecond=0)
                    except:
                        continue

                    if tTime!=tCurTime and tCurTime is not None:
                        RedisManager.set_history_data_list_minutes(tCurTime, strCurPointNameList, strCurPointValueList)

                        strCurPointNameList = [strPointName]
                        strCurPointValueList = [strValue]
                        tCurTime = tTime
                    else:
                        if tCurTime is None:
                            tCurTime = tTime
                        strCurPointNameList.append(strPointName)
                        strCurPointValueList.append(strValue)

                if tCurTime and strCurPointValueList and strCurPointValueList:
                        RedisManager.set_history_data_list_minutes(tCurTime, strCurPointNameList, strCurPointValueList)



        except Exception as e:
            strError = 'ERROR in updateRedisDataMinutesOfDay: ' + e.__str__()
            print(strError)
            logging.error(strError)
            return False

        return True

    def updateRedisData(self, strTime, strTablePeriod='h1'):
        rv = self.get_history_data_all_one_moment_padded(strTime, strTablePeriod)
        bRedisAlive = RedisManager.is_alive()
        if bRedisAlive:
            for item in rv:
                try:
                    strPointName = item['name']
                    strPointValue = item['value']
                    RedisManager.set_history_data_list([strPointName], [strTime], [strPointValue])
                except:
                    continue
        return True

    def countUserPendingFault(self, userName, role):
        strQuery = None
        param = ()
        if role == "admin":
            strQuery = "select A.orderId, fwo.status, A.maxModTime, fwo.ownUser from (select max(modifyTime) as maxModTime, " \
                       "orderId from fdd_work_order group by orderId) as A left join fdd_work_order fwo " \
                       "on A.maxModTime = fwo.modifyTime AND A.orderId = fwo.orderId where fwo.ownUser = %s or status = 0"
            param = ("%{name}%".format(name=userName),)
        elif role == "executor":
            strQuery = "select A.orderId, fwo.status, A.maxModTime, fwo.ownUser from (select max(modifyTime) as maxModTime, " \
                       "orderId from fdd_work_order group by orderId) as A left join fdd_work_order fwo " \
                       "on A.maxModTime = fwo.modifyTime AND A.orderId = fwo.orderId where fwo.ownUser = %s"
            param = ("%{name}%".format(name=userName),)

        if not strQuery:
            return dict(code=1, msg="当前用户的权限在后台配置中有误", data=None)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=param)

        if not items:
            return dict(code=1, msg="数据库繁忙，稍后再试", data=None)
        if not len(items):
            return dict(code=1, msg="查询失败", data=None)

        count = 0
        for item in items:
            if not len(item):
                continue
            if role == "admin" and item[1] in [0, 2, 4]:
                count += 1
            elif role == "executor" and item[1] in [1]:
                count += 1

        return dict(code=0, msg="查询成功", data=count)

    def getFaultOwner(self, orderId):
        strQuery = "select A.orderId, fwo.ownUser from (select max(modifyTime) as maxModTime, orderId from fdd_work_order " \
                   "group by orderId) as A left join fdd_work_order fwo on A.maxModTime = fwo.modifyTime AND A.orderId = fwo.orderId where fwo.orderId = %s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
        if not items:
            return dict(code=1, msg="数据库繁忙，稍后再试", data=None)
        if not len(items):
            return dict(code=1, msg="查询失败", data=None)
        return dict(code=0, msg="查询成功", data=items[0][1])

    def getFaultCreator(self, orderId):
        strQuery = "SELECT opUserName FROM fdd_work_order WHERE orderId=%s AND opType=-1"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, parameter=(orderId,))
        if not items:
            return dict(code=1, msg="数据库繁忙，稍后再试", data=None)
        if not len(items):
            return dict(code=1, msg="查询失败", data=None)
        return dict(code=0, msg="查询成功", data=items[0][0])

    def getAllOngoingAndPendingFaultOrders(self):
        strQuery = "select A.orderId, A.maxModTime, fwo1.status, fwo1.fddName, fwo1.opUserName, fwo1.opType, " \
                   "fwo1.opContentData, fwo1.createTime, fwo1.ownUser, fwo2.detail, fwo2.title from " \
                   "(select max(modifyTime) as maxModTime, orderId from fdd_work_order group by orderId) as A left join" \
                   " fdd_work_order fwo1 on A.maxModTime = fwo1.modifyTime AND A.orderId = fwo1.orderId left join (select orderId, detail, title from" \
                   " fdd_work_order where opType = -1) as fwo2 on fwo1.orderId = fwo2.orderId where fwo1.status in (1, 2, 4)"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if not items:
            return dict(code=1, msg="数据库繁忙，稍后再试", data=None)
        if not len(items):
            return dict(code=1, msg="查询失败", data=None)
        return dict(code=0, msg="", data=items)

    def resetWorkOrders(self, infoList):
        strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, orderId, modifyTime, detail, status, createTime, title) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, parameter=tuple(infoList))
        return bSuc

    def getEquipmentsList(self, targetPage, pageSize, projectId):
        try:
            dbName = app.config['DATABASE']

            strQuery = 'select count(1) from equip_list where projId=(%s)'
            params = tuple([projectId])
            xx = self._mysqlDbContainer.op_db_query(dbName, strQuery, params)
            total = xx[0][0]

            offset = (targetPage - 1) * pageSize
            strQuery = 'select id,`default_equipType`, `projId`, `name`, `description`, `online_addr`, `maintenanceStatus`, `repairStatus`,' \
                       ' `repairResponsiblePerson`, `installLocation`, `communicateStatus`, `warningStatus`, `area_id`,`model_id`, `qrcode` ,`system_id` from equip_list where projId=%s limit %s,%s'
            params = (projectId, offset, pageSize)
            result = self._mysqlDbContainer.op_db_query(dbName, strQuery, params)

            records = []
            for item in result:
                tmp = {
                    'id': item[0],
                    'type': item[1],
                    'projId': item[2],
                    'name': item[3],
                    'description': item[4],
                    'online_addr': item[5],
                    'maintenanceStatus': item[6],
                    'repairStatus': item[7],
                    'repairResponsiblePerson': item[8],
                    'installLocation varchar': item[9],
                    'communicateStatus': item[10],
                    'warningStatus': item[11],
                    'area_id': item[12],
                    'model_id': item[13],
                    'qrcode': item[14],
                    'system_id': item[15]
                }
                records.append(tmp)
            return records, total

        except Exception as e:
            logging.error('getEquipmentsList error:' + e.__str__())
            return False

    def addEquipment(self, post, pageSize):
        try:
            dbName = app.config['DATABASE']
            result = {'status': True, 'id': None, 'msg': 'success'}
            strQuery = 'insert into equip_list(id,`default_equipType`, `projId`, `name`, `description`, `online_addr`, `maintenanceStatus`, `repairStatus`,' \
                       ' `repairResponsiblePerson`, `installLocation`, `communicateStatus`, `warningStatus`, `area_id`, `model_id`, `system_id`) values(null,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)'
            params = tuple(post)
            new_id = self._mysqlDbContainer.op_db_update_with_id(dbName, strQuery, params)
            print(new_id)
            print('addEquipment insert success: Equipment id: %s'%(str(new_id)))
            strQuery = 'select count(1) from equip_list where projId=(%s)'
            params = tuple([post[1]])
            total = self._mysqlDbContainer.op_db_query(dbName, strQuery, params)
            print('total=', total)
            # pageNum = math.ceil(total/pageSize)
            return new_id, total[0][0]
        except Exception as e:
            logging.error('addEquipment error:' + e.__str__())
            return False

    def delEquipments(self, projectId, delArray, pageSize, curPage):
        try:
            # 删除设备
            dbName = app.config['DATABASE']
            strQuery = 'delete from equip_list where id=-1'
            for id in delArray:
                strQuery += ' or id=%s'
            params = tuple(delArray)
            self._mysqlDbContainer.op_db_update(dbName, strQuery, params)

            data, total = self.getEquipmentsList(curPage, pageSize, projectId)
            return data, total

        except Exception as e:
            logging.error('delEquipments error:' + e.__str__())
            return False

    def delAsset(self, equip_id, project_id):
        bSuccess = True
        sql = 'delete from equip_asset where equip_id=%s and project_id=%s' % ( equip_id,project_id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)
        return bSuccess

    def modifyEquipment(self, post):
        try:
            # 修改设备
            dbName = app.config['DATABASE']
            strQuery = 'update equip_list set default_equipType=%s, projId=%s,name=%s,description=%s,online_addr=%s,' \
                       'maintenanceStatus=%s,repairStatus=%s,repairResponsiblePerson=%s,installLocation=%s,communicateStatus=%s,' \
                       'warningStatus=%s,area_id=%s,model_id=%s,system_id=%s where id=%s'
            params = tuple(post)
            self._mysqlDbContainer.op_db_update(dbName, strQuery, params)
            return True
        except Exception as e:
            logging.error('modifyEquipment error:' + e.__str__())
            return False

    def searchEquipment(self, projectId, searchKey, system_id, pageSize, targetPage):
        total = 0
        records = None
        flag = False
        try:
            # 查找关键字
            dbName = app.config['DATABASE']
            flag = True
            arr = searchKey.split(' ')
            str1 = 'select count(1) '
            strQuery = "from equip_list l"
            print(strQuery)
            paramArray = []
            strQuery += " where l.projId = %s"
            paramArray.append(projectId)

            if system_id != '':
                strQuery += " and l.system_id = %s "
                paramArray.append(system_id)

            strQuery += " and (l.name like %s "
            for i in range(1, len(arr)):
                strQuery += " or l.name like %s"
            strQuery += ')'
            arr = ['%' + x + '%' for x in arr]

            params = tuple(paramArray + arr)
            # print(strQuery % params)
            xx = self._mysqlDbContainer.op_db_query(dbName, str1 + strQuery, params)
            if len(xx) == 0:
                total = 0
                records = None
                flag = False
            else:
                total = xx[0][0]

            if total > 0:
                # flag = False
                # data = self.getEquipmentsList(1, pageSize, projectId)
                # records = data[0]
                # else:

                strQuery += 'limit %s, %s'

                strQuery = 'select l.id,l.default_equipType,l.projId,l.name,l.description,l.online_addr,l.maintenanceStatus,l.repairStatus,' \
                           'l.repairResponsiblePerson,l.installLocation,l.communicateStatus,l.warningStatus,l.area_id,l.model_id,l.qrcode,l.system_id ' + strQuery
                offset = (targetPage - 1) * pageSize
                arr.append(offset)
                arr.append(pageSize)
                params = tuple(paramArray + arr)
                data = self._mysqlDbContainer.op_db_query(dbName, strQuery, params)
                records = []

                pointListTotal = []
                projId = ''
                for index in range(len(data)):
                    item = data[index]
                    tmp = {
                        'id': item[0],
                        'type': item[1],
                        'projId': item[2],
                        'name': item[3],
                        'description': item[4],
                        'online_addr': item[5],
                        'maintenanceStatus': item[6],
                        'repairStatus': item[7],
                        'repairResponsiblePerson': item[8],
                        'installLocation varchar': item[9],
                        'communicateStatus': item[10],
                        'warningStatus': item[11],
                        'area_id': item[12],
                        'model_id': item[13],
                        'qrcode': item[14],
                        'system_id': item[15],
                        'paramCode':[]
                    }
                    records.append(tmp)

                equipIds = [str(item.get('id')) for item in records]
                equipParamsMap = self.getParamsOfEquipments(projectId, equipIds)
                for index in range(len(records)):
                    records[index]['paramCode'] = equipParamsMap.get(str(records[index].get('id')),[])
                total = len(records)

            return records, total, flag
        except Exception as e:
            logging.error('modifyEquipment error:' + e.__str__())
            return False
        finally:
            pass

    def getParamsOfEquipments(self,projectId, equipIds):
        paramMap = {}
        try:
            dbName = app.config['DATABASE']
            strQuery = "select p.equip_id, p.paramCode,p.paramName,p.minvalue,p.maxvalue,p.paramUnit,paramCommand  from equip_param p where p.equip_id IN (%s)" % (','.join(equipIds))
            data = self._mysqlDbContainer.op_db_query(dbName, strQuery)
            paramArr = []

            pointListTotal = []
            for item in data:
                equipId = str(item[0])
                paramCode = item[1]
                paramArr.append(paramCode)
                commandDict = {}
                if item[6]:
                    try:
                        commandDict = json.loads(item[6])
                    except:
                        pass
                tmp = {
                    'point': str(item[1]),
                    'name': str(item[2]),
                    'minvalue': str(item[3]),
                    'maxvalue': str(item[4]),
                    'unit': str(item[5]),
                    'paramCommand': commandDict
                }
                if paramMap.get(equipId, None) is None:
                    paramMap[equipId] = []
                paramMap[equipId].append(tmp)
                pointListTotal.append(paramCode)

            if pointListTotal:
                result = self.getInputTable(pointListTotal)
                pvs = result[0]
                pvsMap = {}
                if isinstance(pvs, list) and pvs:
                    for i in pvs:
                        pvsMap[i['name']] = i['value']

                for k,v in paramMap.items():
                    for index in range(len(v)):
                        item = v[index]
                        strPTName = item.get('point')
                        paramMap[k][index]['value'] = pvsMap.get(strPTName, '')

        except Exception as e:
            traceback.print_exc()
            strError = 'ERROR in getParamsOfEquipments'+ e.__str__()
            print(strError)
            logging.error(strError)

        return paramMap

    def getEquipmentById(self, _id):
        try:
            dbName = app.config['DATABASE']
            strQuery = 'select id,`default_equipType`, `projId`, `name`, `description`, `online_addr`, `maintenanceStatus`, `repairStatus`,' \
                       ' `repairResponsiblePerson`, `installLocation`, `communicateStatus`, `warningStatus`, `area_id`,`model_id`,  `qrcode`,`system_id` from equip_list where id=%s'
            params = (_id,)
            item = self._mysqlDbContainer.op_db_query(dbName, strQuery, params)
            if (len(item) == 0):
                return None
            item = item[0]
            data = {
                'id': item[0],
                'type': item[1],
                'projId': item[2],
                'name': item[3],
                'description': item[4],
                'online_addr': item[5],
                'maintenanceStatus': item[6],
                'repairStatus': item[7],
                'repairResponsiblePerson': item[8],
                'installLocation varchar': item[9],
                'communicateStatus': item[10],
                'warningStatus': item[11],
                'area_id': item[12],
                'model_id': item[13],
                'qrcode': item[14],
                'system_id': item[15]
            }
            return data
        except Exception as e:
            logging.error('modifyEquipment error:' + e.__str__())
            return None

    def getAssetTemplates(self,project_id):
        dbName = app.config['DATABASE']
        sql = 'select `id`,`name` from equip_asset_tmpl_def where project_id=%s' % (project_id)
        rv = self._mysqlDbContainer.op_db_query(dbName, sql)
        if rv == None:
            return {}

        data = {}
        for item in rv:
            data.update({item[0]: item[1]})
        return data

    def updateAsset(self, equip_id,project_id,list):
        bSuccess = True
        dbName = app.config['DATABASE']
        sql = 'delete from equip_asset where equip_id=%s and project_id=%s' % ( equip_id,project_id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbName, sql)

        sql = 'INSERT INTO `equip_asset`(`equip_id`,  `en_name`, `param_value`, `project_id`) VALUES (%s, %s, %s, %s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update_many(dbName, sql, list)
        return bSuccess

    def updateAssetFile(self, equip_id,list):
        bSuccess = True
        dbName = app.config['DATABASE']
        sql = 'delete from equip_asset_file where equip_id=%s ' % ( equip_id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(dbName, sql)

        sql = 'INSERT INTO `equip_asset_file`(`fileName`, `url`, `equip_id`) VALUES (%s, %s, %s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update_many(dbName, sql, list)
        return bSuccess

    def getUpdateInitAsset(self,template_id,equip_id,project_id):
        sql = 'select b.cn_name,b.en_name,a.param_value,b.ui_type,b.sort_num,b.group_num from (select * from equip_asset where equip_id=%s )a  ' \
              'right join (select * from equip_asset_tmpl where tmpl_def_id=%s) b on a.en_name=b.en_name and a.project_id=b.project_id' \
              ' and a.project_id=%s   order by sort_num asc' % (equip_id,template_id,project_id)
        dbName = app.config['DATABASE']
        rv = self._mysqlDbContainer.op_db_query(dbName, sql)
        data = []
        for item in rv:
            data.append({
                "cn_name": item[0],
                'en_name': item[1],
                'param_value': item[2],
                'ui_type': item[3],
                'sort_num': item[4],
                'group_num': item[5]
            })
        return data

    def getEquipmentTemplateId(self, strEquipType):
        strQuery = "SELECT id FROM equip_asset_tmpl_def WHERE name like '%" + strEquipType + "%'"
        # strQuery = "SELECT id FROM equip_asset_tmpl_def WEHRE name like '%{key}%'".format(key=strEquipType)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if rv is None:
            return None
        if not len(rv):
            return None
        if not len(rv[0]):
            return None
        return rv[0][0]

    def getEquipmentId(self, identity):
        strQuery = "SELECT el.id FROM equip_list el left join equip_asset ea on el.id = ea.equip_id where ea.param_value = %s"
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (identity,))
        if rv is None:
            return None
        if not len(rv):
            return None
        if not len(rv[0]):
            return None
        return int(float(rv[0][0]))

    def updateEquipQrcode(self,equip_id,equip_no):
        bSuccess = True
        sql = 'update equip_list set qrcode=%s where id=%s '
        print("3:" +str(equip_no))
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple([equip_no,equip_id]))
        return bSuccess

    def getEquipParamList(self,equip_id,pageSize, targetPage):
        sql1 = 'select count(1)  from equip_param  where 1=1 '
        sql2 = 'select `id`, `paramName`, `paramCode`, `minValue`, `maxValue`, `paramUnit` from equip_param  where 1=1 '
        result = {'status': True, 'msg': 'ok', 'data': None, 'total': -1, 'flag': True}

        dbName = app.config['DATABASE']

        paramArray = []
        strQuery = " "
        if equip_id != '':
            strQuery += " and equip_id = %s "
            paramArray.append(equip_id)

        params = tuple(paramArray)
        rv = self._mysqlDbContainer.op_db_query(dbName, sql1+strQuery, params)
        if len(rv) == 0:
            total = 0
        else:
            total = rv[0][0]

        data = []

        if(total>0):
            sql2 += strQuery + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            paramArray.append(offset)
            paramArray.append(pageSize)
            params = tuple(paramArray)
            rv = self._mysqlDbContainer.op_db_query(dbName, sql2, params)
            for item in rv:
                data.append({
                    "id": item[0],
                    'paramName': item[1],
                    'paramCode': item[2],
                    'minValue': item[3],
                    'maxValue': item[4],
                    'paramUnit': item[5]
                })

        result['data'] = data
        result['total'] = total
        return result

    def getAddInitParam(self, template_id):
        sql = 'select id, `paramName`, `paramCode`, `minValue`, `maxValue`, `paramUnit`, `sort_num` from equip_param_tmpl where tmpl_def_id=%s order by sort_num asc' % (template_id)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
        data = []
        for item in rv:
            data.append({
                "id": item[0],
                "paramName": item[1],
                'paramCode': item[2],
                'minValue': item[3],
                'maxValue': item[4],
                'paramUnit': item[5],
                'sort_num': item[6]
            })
        return data

    def getInitParam(self, equip_id):
        sql = 'select id,paramName,paramCode,minValue,`maxValue`,paramUnit from equip_param where equip_id=%s' % (equip_id)

        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql)
        data = []
        for item in rv:
            data.append({
                "id": item[0],
                "paramName": item[1],
                'paramCode': item[2],
                'minValue': item[3],
                'maxValue': item[4],
                'paramUnit': item[5]
            })
        return data

    def updateEquipParam(self, equip_id,list):
        bSuccess = True

        sql = 'delete from equip_param where equip_id=%s' % ( equip_id)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)

        sql = 'INSERT INTO `equip_param`( `equip_id`, `paramCode`,`paramName`,  `minValue`, `maxValue`, `paramUnit`, `paramCommand`) VALUES ( %s, %s, %s, %s, %s, %s, %s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], sql, list)
        return bSuccess

    def getEquipOperationList(self,searchKey,projectId,pageSize, targetPage):
        sql1 = 'select count(1)  from equip_operation  where 1=1 '
        sql2 = 'select `id`, `describe`, `operate_time`,responsible_name from equip_operation  where 1=1 '
        result = {'status': True, 'msg': 'ok', 'data': None, 'total': -1, 'flag': True}

        paramArray = []
        strQuery = " and project_id = %s"
        paramArray.append(projectId)
        keyArray = searchKey.split(' ')

        strQuery += " and (`describe` like %s "
        for i in range(1, len(keyArray)):
            strQuery += " or `describe` like %s "
        strQuery += ')'
        arr = ['%' + x + '%' for x in keyArray]

        params = tuple(paramArray + arr)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql1 + strQuery, params)
        if rv is None:
            return {'status': False, 'msg': 'ok', 'data': None, 'total': -1}
        elif not len(rv):
            return {'status': False, 'msg': 'ok', 'data': None, 'total': -1}
        if len(rv) == 0:
            total = 0
        else:
            total = rv[0][0]

        data = []

        if(total>0):
            sql2 += strQuery + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            paramArray = paramArray + arr
            paramArray.append(offset)
            paramArray.append(pageSize)
            params = tuple(paramArray)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql2, params)
            for item in rv:
                data.append({
                    "id": item[0],
                    'describe': item[1],
                    'operate_time': to_time_format_string(item[2], "%Y-%m-%d %H:%M:%S"),
                    'responsible_name': item[3]
                })

        result['data'] = data
        result['total'] = total
        return result

    def delEquipOperationList(self, delArray):
        bSuccess = True
        strQuery = 'delete from equip_operation where id=-1'
        for id in delArray:
            strQuery += ' or id=%s'
        params = tuple(delArray)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, params)
        return bSuccess

    def addEquipOperationList(self, post):
        bSuccess = True
        sql = 'INSERT INTO `equip_operation`( `describe`, `operate_time`, `responsible_name`,project_id) VALUES (%s,%s,%s,%s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(post))
        return bSuccess

    def updateEquipOperationList(self, id,describe, operate_time, responsible_name):
        bSuccess = True
        sql = 'update equip_operation set '
        flag = False
        attr = []
        fields = []

        if describe is not None and describe !="":
            attr.append(describe)
            fields.append(' `describe` = %s')
            flag = True

        if operate_time is not None and operate_time !="":
            attr.append(operate_time)
            fields.append(' operate_time = %s')
            flag = True


        if responsible_name is not None and responsible_name !="":
            attr.append(responsible_name)
            fields.append(' responsible_name = %s')
            flag = True

        if flag:
            sql +=', '.join(fields)+ ' where id= %s'
            attr.append(id)
            bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(attr))

        return bSuccess

    def addAssetTmplDef(self, post):
        bSuccess = True
        sql = 'INSERT INTO `equip_asset_tmpl_def`( `name`, `project_id`, `describe`) VALUES (%s,%s,%s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(post))
        return bSuccess

    def delAssetTmplDef(self, delArray):
        strQuery = 'delete from equip_asset_tmpl_def where id=-1'
        strQuery2 = "DELETE FROM equip_asset_tmpl WHERE tmpl_def_id in (" + ",".join([str(num) for num in delArray]) + ")"
        for id in delArray:
            strQuery += ' or id=%s'
        params = tuple(delArray)
        bSuc1 = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, params)

        bSuc2 = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery2)
        return bSuc1 and bSuc2

    def updateAssetTmplDef(self, id,name, describe):
        bSuccess = True
        sql = 'update equip_asset_tmpl_def set '
        flag = False
        attr = []
        fields = []

        if describe is not None and describe !="":
            attr.append(describe)
            fields.append(' `describe` = %s')
            flag = True

        if name is not None and name !="":
            attr.append(name)
            fields.append(' `name` = %s')
            flag = True

        if flag:
            sql +=', '.join(fields)+ ' where id= %s'
            attr.append(id)
            bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(attr))
        return bSuccess

    def checkAndAutoInitAssetTmplDef(self):
        strQuery = "SELECT id, name FROM equip_asset_tmpl_def"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items is None:
            return False

        maxId = 0
        existingTmplList = []
        existingTmplIdList = []
        for item in items:
            if item[0] > maxId:
                maxId = item[0]
            existingTmplList.append(item[1])
            existingTmplIdList.append(item[0])

        needToAddTmplNameList = []
        for tmplName in standardAssetTmplDef.keys():
            if tmplName not in existingTmplList:
                needToAddTmplNameList.append(tmplName)

        # 增加缺失的资产模板设备
        if len(needToAddTmplNameList):

            strInsert1 = "INSERT INTO `equip_asset_tmpl_def`( `name`, `project_id`, `describe`) VALUES (%s,%s,%s)"
            strInsert2 = "INSERT INTO `equip_asset_tmpl`(`tmpl_def_id`, `cn_name`, `en_name`, `ui_type`, `group_num`, `sort_num`, `project_id`) VALUES (%s,%s,%s,%s,%s,%s,%s)"
            for tmplName in needToAddTmplNameList:
                result = self._mysqlDbContainer.op_db_insert(app.config['DATABASE'], strInsert1, (tmplName, 0, tmplName))
                if not result.get("success"):
                    continue

                lastRowId = result.get("lastrowid")

                tmplId = lastRowId
                argList = standardAssetTmplDef.get(tmplName)
                paramList = []

                for arg in argList:
                    tmpList = [tmplId]
                    tmpList.extend(arg)
                    paramList.append(tuple(tmpList))

                self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strInsert2, tuple(paramList))
                time.sleep(0.5)

        # 对既有的资产模板设备进行详情字段检查，若缺失则增加
        if len(existingTmplIdList):
            strQuery = "SELECT t1.tmpl_def_id, t1.cn_name, t1.en_name, t2.name FROM equip_asset_tmpl t1 LEFT JOIN equip_asset_tmpl_def t2 ON t1.tmpl_def_id = t2.id WHERE t1.tmpl_def_id IN (" + ",".join([str(num) for num in existingTmplIdList]) + ")"
            tmplInfoList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
            if tmplInfoList == None:
                return

            dTmplChNameIdMap = {}
            dExistingTmplInfo = {}
            for tmplInfo in tmplInfoList:
                tmplChName = tmplInfo[3]
                parameter = tmplInfo[2]
                nTmplId = tmplInfo[0]
                dTmplChNameIdMap.update({tmplChName: nTmplId})

                if tmplChName not in dExistingTmplInfo.keys():
                    dExistingTmplInfo.update({tmplChName: []})

                if parameter not in dExistingTmplInfo[tmplChName]:
                    dExistingTmplInfo[tmplChName].append(parameter)

            nMaxId = self.getMaxIdInTable("equip_asset_tmpl")
            nNewId = nMaxId + 1
            paramList = []
            for tmplChName, parameterList in dExistingTmplInfo.items():
                stdList = standardAssetTmplDef.get(tmplChName, [])
                for std in stdList:
                    if std[1] not in parameterList:
                        prList = [nNewId, dTmplChNameIdMap.get(tmplChName)]
                        prList.extend(std)
                        paramList.append(prList)
                        nNewId += 1

            if len(paramList):
                strInsert = "INSERT INTO `equip_asset_tmpl`(`id`, `tmpl_def_id`, `cn_name`, `en_name`, `ui_type`, `group_num`, `sort_num`, `project_id`) VALUES (%s, %s,%s,%s,%s,%s,%s,%s)"
                self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strInsert, tuple(paramList))

        self.initAssetTemplateExcelTables()

    def getAssetTmplDefList(self,searchKey,projectId,pageSize,targetPage):
        sql1 = 'select count(1)  from equip_asset_tmpl_def  where 1=1 '
        sql2 = 'select `id`, `name`, `describe` from equip_asset_tmpl_def  where 1=1 '
        result = {'status': True, 'msg': 'ok', 'data': None, 'total': -1, 'flag': True}

        paramArray = []
        strQuery = " and project_id = %s"
        paramArray.append(projectId)
        keyArray = searchKey.split(' ')

        strQuery += " and (`name` like %s "
        for i in range(1, len(keyArray)):
            strQuery += " or `name` like %s "
        strQuery += ')'
        arr = ['%' + x + '%' for x in keyArray]

        params = tuple(paramArray + arr)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql1 + strQuery, params)
        if len(rv) == 0:
            total = 0
        else:
            total = rv[0][0]

        data = []
        if(total>0):
            sql2 += strQuery + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            paramArray = paramArray + arr
            paramArray.append(offset)
            paramArray.append(pageSize)
            params = tuple(paramArray)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql2, params)
            for item in rv:
                data.append({
                    "id": item[0],
                    'name': item[1],
                    'describe': item[2]
                })

        result['data'] = data
        result['total'] = total
        return result

    def addAssetTmpl(self, post):
        bSuccess = True
        sql = 'INSERT INTO `equip_asset_tmpl`(`tmpl_def_id`, `cn_name`, `en_name`, `ui_type`, `group_num`, `sort_num`, `project_id`) VALUES (%s,%s,%s,%s,%s,%s,%s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(post))
        return bSuccess

    def addAssetTmplMulti(self, tmplId, dataList):
        strDel = "DELETE FROM equip_asset_tmpl WHERE tmpl_def_id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strDel, (tmplId,))
        if not bSuc:
            return False

        strQuery = 'INSERT INTO `equip_asset_tmpl`(`tmpl_def_id`, `cn_name`, `en_name`, `ui_type`, `group_num`, `sort_num`, `project_id`) VALUES (%s,%s,%s,%s,%s,%s,%s)'
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(dataList))
        return bSuc

    def delAssetTmpl(self, delArray):
        bSuccess = True
        strQuery = 'delete from equip_asset_tmpl where id=-1'
        for id in delArray:
            strQuery += ' or id=%s'
        params = tuple(delArray)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, params)
        return bSuccess

    def updateAssetTmpl(self, id,cn_name, en_name,ui_type,group_num,sort_num):
        bSuccess = True
        sql = 'update equip_asset_tmpl set '
        flag = False
        attr = []
        fields = []

        if cn_name is not None and cn_name !="":
            attr.append(cn_name)
            fields.append(' `cn_name` = %s')
            flag = True

        if en_name is not None and en_name !="":
            attr.append(en_name)
            fields.append(' `en_name` = %s')
            flag = True

        if ui_type is not None and ui_type !="":
            attr.append(ui_type)
            fields.append(' `ui_type` = %s')
            flag = True

        if group_num is not None and group_num !="":
            attr.append(group_num)
            fields.append(' `group_num` = %s')
            flag = True

        if sort_num is not None and sort_num !="":
            attr.append(sort_num)
            fields.append(' `sort_num` = %s')
            flag = True

        if flag:
            sql +=', '.join(fields)+ ' where id= %s'
            attr.append(id)
            bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(attr))
        return bSuccess

    def getAssetTmplList(self,searchKey,tmpl_def_id,projectId,pageSize, targetPage):
        sql1 = 'select count(1)  from equip_asset_tmpl  where 1=1 '
        sql2 = 'select `id`, `tmpl_def_id`, `cn_name`, `en_name`, `ui_type`, `group_num`, `sort_num` from equip_asset_tmpl  where 1=1 '
        result = {'status': True, 'msg': 'ok', 'data': None, 'total': -1, 'flag': True}

        paramArray = []
        strQuery = " and project_id = %s"
        paramArray.append(projectId)
        keyArray = searchKey.split(' ')

        if tmpl_def_id != '':
            strQuery += " and tmpl_def_id = %s "
            paramArray.append(tmpl_def_id)

        strQuery += " and (`cn_name` like %s "
        for i in range(1, len(keyArray)):
            strQuery += " or `cn_name` like %s "
        strQuery += ')'
        arr = ['%' + x + '%' for x in keyArray]

        params = tuple(paramArray + arr)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql1 + strQuery, params)
        if len(rv) == 0:
            total = 0
        else:
            total = rv[0][0]

        data = []

        if(total>0):
            sql2 += strQuery + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            paramArray = paramArray + arr
            paramArray.append(offset)
            paramArray.append(pageSize)
            params = tuple(paramArray)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql2, params)
            for item in rv:
                data.append({
                    "id": item[0],
                    'tmpl_def_id': item[1],
                    'cn_name': item[2],
                    'en_name':item[3],
                    'ui_type': item[4],
                    'group_num': item[5],
                    'sort_num': item[6]
                })
        result['data'] = data
        result['total'] = total
        return result

    def addParamTmpl(self, post):
        bSuccess = True
        sql = 'INSERT INTO equip_param_tmpl(`tmpl_def_id`, `paramName`, `paramCode`, `minValue`,`maxValue`,`paramUnit`,`sort_num`)  VALUES (%s,%s,%s,%s,%s,%s,%s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(post))
        return bSuccess

    def addParamTmplMulti(self, tmplId, dataList):
        strDel = "DELETE FROM equip_param_tmpl WHERE tmpl_def_id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strDel, (tmplId,))
        if not bSuc:
            return False

        strQuery = 'INSERT INTO equip_param_tmpl(`tmpl_def_id`, `paramName`, `paramCode`, `minValue`,`maxValue`,`paramUnit`,`sort_num`)  VALUES (%s,%s,%s,%s,%s,%s,%s)'
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(dataList))
        return bSuc

    def delParamTmpl(self, delArray):
        bSuccess = True
        strQuery = 'delete from equip_param_tmpl where id=-1'
        for id in delArray:
            strQuery += ' or id=%s'
        params = tuple(delArray)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, params)
        return bSuccess

    def updateParamTmpl(self, id,paramName, paramCode ,minValue,maxValue,paramUnit,sort_num):
        bSuccess = True
        sql = 'update equip_param_tmpl set '
        flag = False
        attr = []
        fields = []

        if paramName is not None and paramName !="":
            attr.append(paramName)
            fields.append(' `paramName` = %s')
            flag = True

        if maxValue is not None and maxValue !="":
            attr.append(maxValue)
            fields.append(' `maxValue` = %s')
            flag = True

        if minValue is not None and minValue !="":
            attr.append(minValue)
            fields.append(' `minValue` = %s')
            flag = True

        if paramUnit is not None and paramUnit !="":
            attr.append(paramUnit)
            fields.append(' `paramUnit` = %s')
            flag = True

        if paramCode is not None and paramCode !="":
            attr.append(paramCode)
            fields.append(' `paramCode` = %s')
            flag = True

        if sort_num is not None and sort_num !="":
            attr.append(sort_num)
            fields.append(' `sort_num` = %s')
            flag = True

        if flag:
            sql +=', '.join(fields)+ ' where id= %s'
            attr.append(id)
            bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(attr))
        return bSuccess

    def getParamTmplList(self,searchKey,tmpl_def_id,projectId,pageSize, targetPage):
        sql1 = 'select count(1)  from equip_param_tmpl  where 1=1 '
        sql2 = 'select `id`,`tmpl_def_id`, `paramName`, `paramCode`, `minValue`,`maxValue`,`paramUnit`,`sort_num` from equip_param_tmpl  where 1=1 '
        result = {'status': True, 'msg': 'ok', 'data': None, 'total': -1, 'flag': True}

        paramArray = []
        strQuery = " "

        keyArray = searchKey.split(' ')

        if tmpl_def_id != '':
            strQuery += " and tmpl_def_id = %s "
            paramArray.append(tmpl_def_id)

        strQuery += " and (`paramName` like %s "
        for i in range(1, len(keyArray)):
            strQuery += " or `paramName` like %s "
        strQuery += ')'
        arr = ['%' + x + '%' for x in keyArray]

        params = tuple(paramArray + arr)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql1 + strQuery, params)
        if len(rv) == 0:
            total = 0
        else:
            total = rv[0][0]

        data = []

        if(total>0):
            sql2 += strQuery + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            paramArray = paramArray + arr
            paramArray.append(offset)
            paramArray.append(pageSize)
            params = tuple(paramArray)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql2, params)
            for item in rv:
                data.append({
                    "id": item[0],
                    'tmpl_def_id': item[1],
                    'paramName':item[2],
                    'paramCode': item[3],
                    'minValue': item[4],
                    'maxValue': item[5],
                    'paramUnit': item[6],
                    'sort_num': item[7]
                })
        result['data'] = data
        result['total'] = total
        return result

    def addSystem(self, post):
        bSuccess = True
        sql = 'INSERT INTO `system_list`(`system_name`, `system_desc`, `createTime`, `system_img`, `projId`)VALUES (%s, %s, now(), %s, %s)'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(post))
        return bSuccess

    def delSystem(self, delArray):
        bSuccess = True
        strQuery = 'delete from system_list where id=-1'
        for id in delArray:
            strQuery += ' or id=%s'
        params = tuple(delArray)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, params)
        return bSuccess

    def updateSystem(self, id, system_name, system_desc ,system_img):
        bSuccess = True
        sql = 'update system_list set '
        flag = False
        attr = []
        fields = []

        if system_name is not None and system_name !="":
            attr.append(system_name)
            fields.append(' `system_name` = %s')
            flag = True

        if system_desc is not None and system_desc !="":
            attr.append(system_desc)
            fields.append(' `system_desc` = %s')
            flag = True

        if system_img is not None and system_img !="":
            attr.append(system_img)
            fields.append(' `system_img` = %s')
            flag = True

        if flag:
            sql +=', '.join(fields)+ ' where id= %s'
            attr.append(id)
            bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql,tuple(attr))
        return bSuccess

    def getSystemList(self,searchKey,projectId,pageSize, targetPage):
        sql1 = 'select count(1)  from system_list  where 1=1 '
        sql2 = 'select `id`,`system_name`, `system_desc`, `createTime`, `system_img`, `projId` from system_list  where 1=1 '
        result = {'status': True, 'msg': 'ok', 'data': None, 'total': -1, 'flag': True}

        paramArray = []
        strQuery = " and projId = %s"
        paramArray.append(projectId)
        keyArray = searchKey.split(' ')

        strQuery += " and (`system_name` like %s "
        for i in range(1, len(keyArray)):
            strQuery += " or `system_name` like %s "
        strQuery += ')'
        arr = ['%' + x + '%' for x in keyArray]

        params = tuple(paramArray + arr)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql1 + strQuery, params)
        if len(rv) == 0:
            total = 0
        else:
            total = rv[0][0]

        data = []

        if(total>0):
            sql2 += strQuery + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            paramArray = paramArray + arr
            paramArray.append(offset)
            paramArray.append(pageSize)
            params = tuple(paramArray)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql2, params)
            for item in rv:
                data.append({
                    "id": item[0],
                    'system_name': item[1],
                    'system_desc':item[2],
                    'createTime': to_time_format_string(item[3],"%Y-%m-%d %H:%M:%S"),
                    'sort_num': item[4]
                })
        result['data'] = data
        result['total'] = total
        return result

    def updateEquipByExcel(self, equipArray, project_id):
        # 设备数据处理
        equipName = system_id = str_position = str_equipment_type = area_id = model_id = ''

        for item in equipArray:
            if item.get('cn_name')=='设备名称':
                equipName = item.get('value')
            elif item.get('cn_name')=='系统编码':
                system_id=item.get('value')
            elif item.get('cn_name')=='安装位置':
                str_position = str(item.get('value'))
            elif item.get('cn_name')=='设备类型':
                str_equipment_type= item.get('value')
            elif item.get('cn_name')=='模板id':
                model_id =item.get('value')
            elif item.get('cn_name')=='区域编码':
                area_id =item.get('value')

        equip_post = [str_equipment_type, project_id, equipName, "", "", 0, 0, "", str_position, 0, 0, area_id, model_id,system_id, ]
        equip_id, total = BEOPDataAccess.getInstance().addEquipment(equip_post,1)

        return equip_id,model_id

    def updateEquipAssetByExcel(self,assetArray,project_id,equip_id,model_id):
        # 资产数据处理
        if len(assetArray) > 0:
            assetTmpl = self.getUpdateInitAsset(model_id, 0, project_id)
            assetList = []
            equip_no_val = None
            for item in assetTmpl:
                cname = item.get('cn_name')
                for item2 in assetArray:
                    cn_name2 = item2.get('cn_name')
                    if cname == cn_name2:
                        en_name = item.get('en_name')
                        value = item2.get('value')
                        post = [equip_id, en_name, value, project_id]
                        assetList.append(post)
                        if en_name == 'equip_no':
                            equip_no_val = value

            self.getInstance().updateAsset(equip_id, project_id, assetList)

            if equip_no_val is not None and equip_no_val != '':
                self.updateEquipQrcode(equip_id, equip_no_val)

    def updateEquipParamByExcel(self,paramArray,project_id,equip_id,model_id):

        # 参数数据处理
        if len(paramArray) > 0:
            paramTmpl = self.getAddInitParam(model_id)
            paramList = []

            for item in paramTmpl:
                paramName = item.get('paramName')
                for item2 in paramArray:
                    cn_name2 = item2.get('cn_name')
                    if paramName == cn_name2:
                        #paramCode = item2.get('value')
                        minValue = item2.get('minValue')
                        maxValue = item2.get('maxValue')
                        unitValue = item2.get('unitValue')
                        paramCode = item2.get("paramCode")
                        paramCommand = item2.get("paramCommand")
                        post = [equip_id,paramCode,paramName, minValue, maxValue, unitValue,paramCommand]
                        paramList.append(post)

            self.updateEquipParam(equip_id, paramList)

    def importEquipmentExcel(self,filePath,project_id):
        datafile = xlrd.open_workbook(filePath)
        existingAssetNameList = self.getExistingAssetNames()
        nsheets = datafile.nsheets
        for nSheetIdx in range(nsheets):
            try:
                sheet = datafile.sheet_by_index(nSheetIdx)
                sheetName = sheet.name
                if sheetName in existingAssetNameList:
                    continue

                rows_num = sheet.nrows
                cols_num = sheet.ncols
                if rows_num > 1:
                    merged = sheet.merged_cells

                    excelArray = []

                    for (rlow, rhigh, clow, chigh) in merged:
                        if clow == 0 and chigh == 1:
                            dataArray = []
                            test = ""
                            mergedName = sheet.cell_value(rlow, 0)
                            print(test)
                            for row in range(rlow, rhigh):
                                name = sheet.cell_value(row, 1)
                                ctype2 = sheet.cell(row, 2).ctype
                                ctype3 = sheet.cell(row, 3).ctype

                                ctype4 = sheet.cell(row, 4).ctype

                                value = sheet.cell_value(row, 2)
                                minValue = sheet.cell_value(row, 3)
                                maxValue = sheet.cell_value(row, 4)
                                unitValue = sheet.cell_value(row, 5)
                                paramCode = sheet.cell_value(row,6)
                                paramCommand = sheet.cell_value(row, 7)

                                if ctype2 == 2 and value % 1 == 0:  # 如果是整形
                                    value = int(value)
                                if ctype3 == 2 and minValue % 1 == 0:  # 如果是整形
                                    minValue = int(minValue)
                                if ctype4 == 2 and maxValue % 1 == 0:  # 如果是整形
                                    maxValue = int(maxValue)

                                data = {"cn_name": name, "value": value, "minValue": minValue, "maxValue": maxValue, "unitValue": unitValue, "paramCode": paramCode, "paramCommand": paramCommand}
                                dataArray.append(data)
                            if mergedName=='设备信息':
                                excelArray.append({"equip":dataArray})
                            elif mergedName=='资产信息':
                                excelArray.append({"asset": dataArray})
                            elif mergedName=='系统参数':
                                excelArray.append({"param": dataArray})

                    if len(excelArray) > 0:
                        assetData = []
                        paramData = []
                        equip_id = None
                        model_id = None
                        for excelItem in excelArray:
                            if ('equip' in excelItem):
                                equip_id,model_id=self.updateEquipByExcel(excelItem.get('equip'),project_id)
                            if ('asset' in excelItem):
                                assetData = excelItem.get('asset')
                            if ('param' in excelItem):
                                paramData =excelItem.get('param')

                        self.updateEquipAssetByExcel(assetData, project_id,equip_id, model_id)
                        self.updateEquipParamByExcel(paramData, project_id,equip_id, model_id)
            except Exception as e:
                strError = "importEquipmentExcel ERROR: %s" % e.__str__()
                log_info_to_file("equip_asset_import_file_error_%s.log" % datetime.now().strftime("%Y-%m-%d"), strError)

    def importEquipmentExcelZip(self, filePath, project_id, fileDir):
        # zipfile解压
        z = zipfile.ZipFile(filePath, 'r')
        strNewPath = os.path.join(fileDir, datetime.now().strftime('%Y%m%d%H%M%S'))
        try:
            os.mkdir(strNewPath)
        except:
            pass
        z.extractall(path=strNewPath)
        z.close()

        for parent, dirnames, filenames in os.walk(strNewPath, followlinks=True):
            for filename in filenames:
                file_path = os.path.join(parent, filename)
                self.importEquipmentExcel(file_path, project_id)

    # 定义导入excel
    def importEquipExcel(self, filePath, project_id):
        datafile = xlrd.open_workbook(filePath, formatting_info=True)
        # 获取工作簿
        nsheets = datafile.nsheets
        sheet = datafile.sheet_by_index(0)
        rows_num = sheet.nrows
        cols_num = sheet.ncols
        if rows_num > 1:
            merged = sheet.merged_cells

            excelArray = []

            for (rlow, rhigh, clow, chigh) in merged:
                if clow == 0 and chigh == 1:
                    dataArray = []
                    test = ""
                    mergedName = sheet.cell_value(rlow, 0)
                    print(test)
                    for row in range(rlow, rhigh):
                        name = sheet.cell_value(row, 1)
                        ctype2 = sheet.cell(row, 2).ctype
                        ctype3 = sheet.cell(row, 3).ctype
                        ctype4 = sheet.cell(row, 4).ctype
                        value = sheet.cell_value(row, 2)
                        minValue = sheet.cell_value(row, 3)
                        maxValue = sheet.cell_value(row, 4)
                        unitValue = sheet.cell_value(row, 5)
                        paramCode = sheet.cell_value(row, 6)
                        paramCommand = sheet.cell_value(row, 7)

                        if ctype2 == 2 and value % 1 == 0:  # 如果是整形
                            value = int(value)
                        if ctype3 == 2 and minValue % 1 == 0:  # 如果是整形
                            minValue = int(minValue)
                        if ctype4 == 2 and maxValue % 1 == 0:  # 如果是整形
                            maxValue = int(maxValue)

                        data = {"cn_name": name, "value": value, "minValue": minValue, "maxValue": maxValue,
                                "unitValue": unitValue, "paramCode": paramCode, "paramCommand": paramCommand}
                        dataArray.append(data)
                    if mergedName == '设备信息':
                        excelArray.append({"equip": dataArray})
                    elif mergedName == '资产信息':
                        excelArray.append({"asset": dataArray})
                    elif mergedName == '系统参数':
                        excelArray.append({"param": dataArray})

            if len(excelArray) > 0:
                assetData = []
                paramData = []
                equip_id = None
                model_id = None
                for excelItem in excelArray:
                    if ('equip' in excelItem):
                        equip_id, model_id = self.updateEquipByExcel(excelItem.get('equip'), project_id)
                    if ('asset' in excelItem):
                        assetData = excelItem.get('asset')
                    if ('param' in excelItem):
                        paramData = excelItem.get('param')

                self.updateEquipAssetByExcel(assetData, project_id, equip_id, model_id)
            # self.addEquipParamByExcel(paramData, equip_id)

    def getRiskFactorList(self,searchKey,projectId,pageSize, targetPage):
        sql1 = 'select count(1)  from risk_factor  where 1=1 '
        sql2 = 'select `id`, `factorName`, `factorCode`, `type`, `num`,`deviceId`,`deviceName`,`remark`, `status` from risk_factor  where 1=1 '
        result = {'status': True, 'msg': 'ok', 'data': None, 'total': -1, 'flag': True}

        paramArray = []
        strQuery = " and project_id = %s"
        paramArray.append(projectId)
        keyArray = searchKey.split(' ')

        strQuery += " and (`factorName` like %s "
        for i in range(1, len(keyArray)):
            strQuery += " or `factorName` like %s "
        strQuery += ')'
        arr = ['%' + x + '%' for x in keyArray]

        params = tuple(paramArray + arr)
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql1 + strQuery, params)
        if len(rv) == 0:
            total = 0
        else:
            total = rv[0][0]

        data = []

        if(total>0):
            sql2 += strQuery + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            paramArray = paramArray + arr
            paramArray.append(offset)
            paramArray.append(pageSize)
            params = tuple(paramArray)
            rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sql2, params)
            for item in rv:
                data.append({
                    "id": item[0],
                    'factorName': item[1],
                    'factorCode':item[2],
                    'type': item[3],
                    'num': item[4],
                    'deviceId': item[5],
                    'deviceName': item[6],
                    'remark': item[7],
                    'status': item[8]
                })
        result['data'] = data
        result['total'] = total
        return result

    def addRiskFactorDetail(self,factor_id, riskFactorDetailData):
        bSuccess = True
        if len(riskFactorDetailData) > 0:
            list = []
            sql = 'INSERT INTO `risk_factor_detail`(`level_id`, `factor_id`, `minValue`, `maxValue`) VALUES(%s,%s,%s,%s)'
            for detail in riskFactorDetailData:
                post =[detail['level_id'],factor_id,detail['minValue'],detail['maxValue']]
                list.append(post)

            bSuccess = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], sql,list)
        return bSuccess

    def addRiskFactor(self, post, riskFactorDetailData):
        bSuccess = True
        sql = 'INSERT INTO `risk_factor`( `factorName`, `factorCode`, `type`, `num`, `remark`, `status`,`project_id`,`deviceId`,`deviceName`) ' \
              ' VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)'
        new_id = self._mysqlDbContainer.op_db_update_with_id(app.config['DATABASE'], sql, tuple(post))
        if(new_id>0):
            bSuccess = bSuccess and self.addRiskFactorDetail(new_id, riskFactorDetailData)
        return bSuccess


    def updateRiskFactor(self, id, factorName, factorCode, type, num, remark, add_detail_Data, del_detail_Data):
        bSuccess = True
        sql = 'update risk_factor set '
        flag = False
        attr = []
        fields = []

        if factorName is not None and factorName != "":
            attr.append(factorName)
            fields.append(' `factorName` = %s')
            flag = True

        if factorCode is not None and factorCode != "":
            attr.append(factorCode)
            fields.append(' `factorCode` = %s')
            flag = True

        if type is not None and type != "":
            attr.append(type)
            fields.append(' `type` = %s')
            flag = True

        if num is not None and num != "":
            attr.append(num)
            fields.append(' `num` = %s')
            flag = True

        if remark is not None and remark != "":
            attr.append(remark)
            fields.append(' `remark` = %s')
            flag = True

        if flag:
            sql += ', '.join(fields) + ' where id= %s'
            attr.append(id)
            bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql, tuple(attr))
            bSuccess = bSuccess and self.addRiskFactorDetail(id, add_detail_Data)
            bSuccess = bSuccess and self.delRiskFactorDetail(del_detail_Data)

        return bSuccess

    def delRiskFactorDetail(self,delArray):
        bSuccess = True
        sql = 'delete from risk_factor_detail where factor_id=-1'
        for id in delArray:
            sql += ' or id=%s'
        params = tuple(delArray)
        bSuccess =self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql, params)
        return bSuccess

    def getRiskFactorByFactorId(self,factor_id):
        strQuery = "select `id`, `level_id`, `factor_id`, `minValue`, `maxValue` from risk_factor_detail where factor_id=%s"
        rv = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (factor_id,))
        data = []
        for item in rv:
            data.append({
                "id": item[0],
                "level_id": item[1],
                'factor_id': item[2],
                'minValue': item[3],
                'maxValue': item[4]
            })
        return data

    def enableRiskFactor(self, status, id):
        bSuccess = True
        sql = 'update risk_factor set `status`=%s where id=%s'
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql, tuple([status, id]))
        return bSuccess

    def delRiskFactor(self, delArray):
        bSuccess = True
        sql1 = 'delete from risk_factor where id=-1'
        sql2 = 'delete from risk_factor_detail where factor_id=-1'
        for id in delArray:
            sql1 += ' or id=%s'
            sql2 += ' or factor_id=%s'
        params = tuple(delArray)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql1, params)
        bSuccess = bSuccess and self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql2, params)
        return bSuccess

    def readProjectEquipmentFromMobile(self, projectId, equipmentId):
        try:
            dbName = app.config['DATABASE']

            strQuery = "select ea.en_name, A.cn_name, ea.param_value, A.sort_num from equip_asset ea left join " \
                       "(select en_name, cn_name, sort_num from equip_asset_tmpl where project_id = %s group by en_name) as A " \
                       "on ea.en_name = A.en_name where ea.project_id = %s and ea.equip_id = %s order by A.sort_num"

            rsp = self._mysqlDbContainer.op_db_query(dbName, strQuery, (projectId, projectId, equipmentId))
            if not rsp:
                return dict(code=1, data={}, msg="获取到的数据为空")
            if not len(rsp):
                return dict(code=1, data={}, msg="获取到的数据为空")
            if not len(rsp[0]):
                return dict(code=1, data={}, msg="获取到的数据为空")

            rspData = []
            en_name_existing = []
            for item in rsp:
                if item[0] in en_name_existing:
                    continue

                rspData.append({item[1]: item[2]})
                en_name_existing.append(item[0])

            return dict(code=0, data=rspData, msg="获取成功")

        except Exception as exp:
            strLog = "ERROR in readProjectEquipmentFromMobile: %s" % exp.__str__()
            logging.error(strLog)
            return dict(code=1, data={}, msg=strLog)

    def searchArea(self, projectId, searchKey, pageSize, targetPage):
        try:
            # 查找关键字
            dbName = app.config['DATABASE']
            flag = True
            arr = searchKey.split(' ')
            str1 = 'select count(*) '
            strQuery = "from inspact_area where projId=%s and (areaName like %s"
            for i in range(1, len(arr)):
                strQuery += " or areaName like %s"
            strQuery += ')'
            arr = ['%' + x + '%' for x in arr]

            params = tuple([projectId] + arr)
            # print(strQuery % params)
            xx = self._mysqlDbContainer.op_db_query(dbName, str1 + strQuery, params)
            if len(xx) == 0:
                total = 0
            else:
                total = xx[0][0]

            strQuery += ' order by seqno limit %s, %s'
            strQuery = 'select * ' + strQuery
            offset = (targetPage - 1) * pageSize
            arr.append(offset)
            arr.append(pageSize)
            params = tuple([projectId] + arr)
            data = self._mysqlDbContainer.op_db_query(dbName, strQuery, params)
            records = []
            for item in data:
                tmp = {
                    'id': item[0],
                    'areaName': item[1],
                    'description': item[2],
                    'seqNo': item[3]
                }
                records.append(tmp)
            return records, total, flag
        except Exception as e:
            logging.error('modifyEquipment error:' + e.__str__())
            return False
        finally:
            pass

    def equipCareList(self, projectId, status,startTime,endTime, pageSize, targetPage):
        try:
            # 查找关键字
            dbName = app.config['DATABASE']
            flag = True

            totalSql = "select count(*)"
            detailSql = "select distinct a.id,`title`,`operation_instruction`,`description`,`status`,`attention`,`createTime`,(select GROUP_CONCAT(username) from care_user b where a.id=b.care_id) userNames,`deviceName`,`selTime`,`checkContent`"
            strBody = " from  care_history  a where 1=1  and projId=%s "
            params = [projectId]

            if status != '':
                strBody += " and `status` = %s "
                params.append(status)

            if startTime != '':
                strBody += " and createTime>=%s "
                params.append(startTime)

            if endTime != '':
                strBody += " and createTime<=%s "
                params.append(endTime)

            xx = self._mysqlDbContainer.op_db_query(dbName, totalSql + strBody, tuple(params))
            if len(xx) == 0:
                total = 0
            else:
                total = xx[0][0]

            listSql = detailSql + strBody + ' limit %s, %s'
            offset = (targetPage - 1) * pageSize
            params.append(offset)
            params.append(pageSize)
            data = self._mysqlDbContainer.op_db_query(dbName, listSql, tuple(params))
            records = []
            for item in data:
                records.append({
                    "id": item[0],
                    'title': item[1],
                    'operation_instruction': item[2],
                    'description': item[3],
                    'status': item[4],
                    'attention': item[5],
                    'actionTime': str(item[6]),
                    'submitTime': str(item[6]),
                    'userNames': item[7],
                    'deviceName': item[8],
                    'selTime': item[9],
                    'checkContent': item[10]
                })
            return records, total, flag
        except Exception as e:
            logging.error('modifyEquipment error:' + e.__str__())
            return False
        finally:
            pass

    # 添加巡检计划
    def addArea(self, name, description, strProjId):
        try:
            dbName = app.config['DATABASE']
            result = {'status': True, 'id': None, 'msg': 'success'}

            str = 'select max(seqno) from inspact_area where projId=%s'
            xx = self._mysqlDbContainer.op_db_query(dbName, str, (strProjId,))
            index = 0
            if len(xx) > 0 and not xx[0][0] is None:
                index = xx[0][0]
            index = index + 1

            strQuery = 'insert into inspact_area( `areaName`, `description`, `seqno`, `projId`, `status`)' \
                       ' values(%s,%s,%s,%s,%s)'
            new_id = self._mysqlDbContainer.op_db_update_with_id(dbName, strQuery,
                                                                 (name, description, index, strProjId, '1',))

            upQuery2 = 'update inspact_area set status=0 where projId=%s'
            self._mysqlDbContainer.op_db_update(dbName, upQuery2, (strProjId,))

            upQuery = 'update inspact_area set status=1 where seqNo=1 and projId=%s'
            self._mysqlDbContainer.op_db_update(dbName, upQuery, (strProjId,))

            # pageNum = math.ceil(total/pageSize)
            return new_id
        except Exception as e:
            logging.error('addInspactPlan error:' + e.__str__())
            return False

    # 删除巡检计划
    def delArea(self, delArray):
        try:
            # 删除设备

            dbName = app.config['DATABASE']
            strQuery = 'delete from inspact_area where id=-1'
            for id in delArray:
                strQuery += ' or id=%s'
            params = tuple(delArray)
            print(params)
            self._mysqlDbContainer.op_db_update(dbName, strQuery, params)
        except Exception as e:
            logging.error('delInspactPlan error:' + e.__str__())
            return False

    def getMysqlDiagnosisInfo(self):
        strQuery = "SHOW variables"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items is None:
            return dict(err=1, msg="数据库请求失败", data={})

        dData = {}
        for item in items:
            if len(item):
                dData.update({item[0]: item[1]})
        return dict(err=0, msg="", data=dData)

    def getModbusEquipmentInputTableData(self):
        strQuery = "SELECT pointname, pointvalue FROM realtimedata_input_modbus_equipment"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items is None:
            return dict(err=1, data={})

        dData = {}
        for item in items:
            dData.update({item[0]: item[1]})
        return dict(err=0, data=dData)

    def get_deeplogic_redis_config(self):
        strQuery = "SELECT unitproperty02 FROM unit01 WHERE unitproperty01='deep_logic_redis_address'"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return dict(err=1, msg="数据请求失败，请稍后重试", data={})
        if not len(items):
            return dict(err=0, msg="无redis配置信息", data={"ip": "", "port": None, "password": ""})
        strInfo = items[0][0]
        infoList = strInfo.split(",")
        if len(infoList) < 3:
            return dict(err=1, msg="配置有误", data={})
        nPort = ""
        try:
            nPort = int(float(infoList[1]))
        except:
            pass

        return dict(err=0, msg="", data={"ip": infoList[0], "port": nPort, "password": infoList[2]})

    def tmplDefIdExists(self, tmplDefId):
        strQuery = "SELECT name FROM equip_asset_tmpl_def WHERE id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (tmplDefId,))
        if items == None:
            return None
        if not len(items):
            return False
        if not len(items[0]):
            return False
        return True

    def getExistingAssetNames(self):
        strQuery = "SELECT name FROM equip_list"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return []
        if not len(items):
            return []

        nameList = []
        for item in items:
            if not len(item):
                continue

            nameList.append(item[0])
        return nameList

    def initAssetTemplateExcelTables(self):
        strErr = ""

        # 准备模板文件磁盘路径
        filesDir = os.path.join(app.static_folder, "files")
        if not os.path.exists(filesDir):
            os.mkdir(filesDir)

        assetDir = os.path.join(filesDir, "asset_template")
        if not os.path.exists(assetDir):
            os.mkdir(assetDir)

        # 单元格对齐工具
        oAlign = Alignment(horizontal="center", vertical="center")

        # 获取资产信息
        strQueryAsset = "select eat.tmpl_def_id, eat.cn_name, eat.en_name, eatd.name, eat.sort_num from equip_asset_tmpl eat left join equip_asset_tmpl_def eatd on eat.tmpl_def_id = eatd.id"
        assetInfoList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQueryAsset)
        if assetInfoList == None:
            return dict(err=1, msg="获取资产信息失败")
        if not len(assetInfoList):
            return dict(err=1, msg="资产信息为空")

        # 整理资产信息字典
        dAssetInfo = {}
        for item in assetInfoList:
            if len(item) < 5:
                continue

            tmplDefId = item[0]
            cnName = item[1]
            enName = item[2]
            tmplName = item[3]
            sortNum = item[4]

            if tmplDefId == None or cnName == None or enName == None or tmplName == None or sortNum == None:
                continue

            strKey = "%s@*@%s" % (tmplDefId, tmplName)

            if strKey not in dAssetInfo.keys():
                dAssetInfo.update({strKey: []})

            dAssetInfo.get(strKey).append([cnName, enName, sortNum])

        # 资产信息按sort_num 排序
        for tmplDefId, infoList in dAssetInfo.items():
            for i in range(len(infoList)):
                for j in range(len(infoList)-i-1):
                    if infoList[j][2] > infoList[j+1][2]:
                        infoList[j], infoList[j+1] = infoList[j+1], infoList[j]

        if not dAssetInfo:
            return dict(err=1, msg="资产信息为空")

        # 获取参数信息
        strQueryParam = "select ept.tmpl_def_id, ept.paramName, ept.paramCode, ept.minValue, ept.maxValue, ept.paramUnit, " \
                        "ept.sort_num, eatd.name  from equip_param_tmpl ept left join equip_asset_tmpl_def eatd on ept.tmpl_def_id = eatd.id"
        paramInfoList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQueryParam)

        # 整理参数信息字典
        dParamInfo = {}
        for item in paramInfoList:
            if len(item) < 8:
                continue

            tmplDefId = item[0]
            paramName = item[1]
            pointName = item[2]
            minValue = item[3]
            maxValue = item[4]
            strUnit = item[5]
            sortNum = item[6]
            tmplName = item[7]

            if tmplDefId == None or paramName == None or pointName == None or minValue == None or maxValue == None or \
                    strUnit == None or sortNum == None or tmplName == None:
                continue

            strKey = "%s@*@%s" % (tmplDefId, tmplName)

            if strKey not in dParamInfo.keys():
                dParamInfo.update({strKey: []})

            dParamInfo.get(strKey).append([paramName, pointName, minValue, maxValue, strUnit, sortNum])

        # 参数信息按sort_num排序
        for tmplDefId, infoList in dParamInfo.items():
            for i in range(len(infoList)):
                for j in range(len(infoList)-i-1):
                    if infoList[j][5] > infoList[j+1][5]:
                        infoList[j], infoList[j+1] = infoList[j+1], infoList[j]

        # 制作模板excel
        for strKey, astInfoList in dAssetInfo.items():
            try:
                if not len(astInfoList):
                    continue

                category = strKey.split("@*@")[1]
                tmplId = strKey.split("@*@")[0]
                assetName = "1#{name}".format(name=category)

                excelFileName = "模板_设备信息导入_{cate}.xlsx".format(cate=category)
                filePath = os.path.join(assetDir, excelFileName)

                book = Workbook()
                sheet = book.create_sheet(assetName, 0)

                # 填表头
                for idx, title in enumerate(["类别", "项目", "内容", "最小值", "最大值", "单位", "点名", "枚举定义"]):
                    sheet.cell(row=1, column=idx+1, value=title)

                # 合并单元格
                sheet.merge_cells(start_row=2, start_column=1, end_row=5, end_column=1)
                sheet.cell(row=2, column=1, value="设备信息")

                # 对齐
                cell = sheet.cell(2, 1)
                cell.alignment = oAlign

                for idx, content in enumerate(["设备名称", "系统编码", "区域编码", "模板id"]):
                    sheet.cell(row=idx+2, column=2, value=content)

                sheet.cell(row=2, column=3, value=assetName)
                sheet.cell(row=5, column=3, value=tmplId)

                sheet.merge_cells(start_row=7, start_column=1, end_row=6+len(astInfoList), end_column=1)
                sheet.cell(row=7, column=1, value="资产信息")

                cell = sheet.cell(7, 1)
                cell.alignment = oAlign

                for idx, content in enumerate(astInfoList):
                    chName = content[0]
                    if chName == "名称":
                        sheet.cell(row=idx+7, column=3, value=assetName)
                    sheet.cell(row=idx+7, column=2, value=content[0])

                prmInfoList = dParamInfo.get(strKey, [])
                if len(prmInfoList):
                    nParamStartRow = 8 + len(astInfoList)
                    sheet.merge_cells(start_row=nParamStartRow, start_column=1, end_row=nParamStartRow + len(prmInfoList) - 1, end_column=1)
                    sheet.cell(row=nParamStartRow, column=1, value="系统参数")

                    cell = sheet.cell(nParamStartRow, 1)
                    cell.alignment = oAlign

                    for idx, content in enumerate(prmInfoList):
                        sheet.cell(row=nParamStartRow+idx, column=2, value=content[0])
                        sheet.cell(row=nParamStartRow+idx, column=4, value=content[2])
                        sheet.cell(row=nParamStartRow + idx, column=5, value=content[3])
                        sheet.cell(row=nParamStartRow + idx, column=6, value=content[4])
                        sheet.cell(row=nParamStartRow + idx, column=7, value=content[1])

                try:
                    if os.path.exists(filePath):
                        try:
                            os.remove(filePath)
                        except:
                            pass

                    book.save(filePath)
                except Exception as err:
                    strErr += ";删除并保存既有模板文件时出错，可能这个文件未关闭（文件名:%s): err:%s" % (excelFileName, err.__str__())

            except Exception as e:
                strErr += ";生成模板文件主循环报错：%s" % e.__str__()

            finally:
                time.sleep(0.1)

        if len(strErr):
            return dict(err=1, msg=strErr)
        return dict(err=0, msg="")

    def fieldExistsInTable(self, fieldName, value, tableName):
        strQuery = "SELECT {field} FROM {table} WHERE {field}='{value}'".format(field=fieldName, table=tableName, value=value)
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return True
        if not len(items):
            return False
        if not len(items[0]):
            return False
        return True


    def addRpsClient(self, id, name, address, strCreateTime, code):
        strQuery = "INSERT INTO rps_client (id, name, address, create_time, enabled, code) VALUE (%s, %s, %s, %s, %s, %s)"
        nEnabled = 1
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id, name, address, strCreateTime, nEnabled, code))
        return bSuc

    def deleteRpsClient(self, fieldName, value):
        strQuery = "UPDATE rps_client SET enabled=0 WHERE {field}='{value}'".format(field=fieldName, value=value)
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery)
        return bSuc

    def updateRpsClient(self, id, updateList):
        strQuery = "UPDATE rps_client SET"

        paramList = []
        for item in updateList:
            paramList.append(" {field}='{value}'".format(field=item[0], value=item[1]))

        strQuery += ",".join(paramList)
        strQuery += " WHERE id=%s"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id,))
        return bSuc

    def getRpsClient(self, anchor):
        strQuery = "SELECT id, name, address, create_time, code FROM rps_client WHERE enabled=1"
        if anchor:
            strQuery += " and {field}='{value}'".format(field=anchor[0], value=anchor[1])

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None
        res = []
        for item in items:
            createTime = ""
            try:
                createTime = item[3].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            res.append({
                "id": item[0],
                "name": item[1],
                "address": item[2],
                "create_time": createTime,
                "code": item[4]
            })
        return res

    def enableRpsClient(self, id):
        strQuery = "UPDATE rps_client SET enabled=1 WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id,))
        return bSuc

    def addRpsContract(self, nId, clientId, salesId, clientContact, contractFileAddress, strCreateTime, deliveryTime, dueTime, amount, status):
        strQuery = "INSERT INTO rps_contract (id, client_id, sales_id, client_contact, attachment_id, create_time, status"
        strValues = ") VALUES (%s, %s, %s, %s, %s, %s, %s"
        params = [nId, clientId, salesId, clientContact, contractFileAddress, strCreateTime, status]
        if deliveryTime:
            strQuery += ",delivery_time"
            strValues += ",%s"
            params.append(deliveryTime)
        if dueTime:
            strQuery += ",due_time"
            strValues += ",%s"
            params.append(dueTime)
        if amount != None:
            strQuery += ",amount"
            strValues += ",%s"
            params.append(amount)
        strQuery += strValues

        strQuery += ")"

        dResult = self._mysqlDbContainer.op_db_insert(app.config['DATABASE'], strQuery, tuple(params))

        return dResult

    def deleteRpsContract(self, id):
        strQuery = "DELETE FROM rps_contract WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id,))
        return bSuc

    def getRpsAttachmentIdListOfContract(self, contractId):
        strQuery = "SELECT attachment_id FROM rps_contract WHERE id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (contractId,))
        if items == None:
            return None
        if not len(items):
            return []
        if not len(items[0]):
            return None
        attachmentIdList = items[0][0].split(",")

        res = []
        for item in attachmentIdList:
            num = None
            try:
                num = int(float(item))
            except:
                pass
            if num != None:
                res.append(num)
        return res

    def getRpsAttachmentIdListOfRetirement(self, retirementId):
        strQuery = "SELECT attachment_id FROM rps_retirement WHERE id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (retirementId,))
        if items == None:
            return None
        if not len(items):
            return []
        if not len(items[0]):
            return None
        attachmentIdList = items[0][0].split(",")

        res = []
        for item in attachmentIdList:
            num = None
            try:
                num = int(float(item))
            except:
                pass
            if num != None:
                res.append(num)
        return res

    def disableRpsAttachment(self, idList):
        if not len(idList):
            return True
        strQuery = "UPDATE rps_attachment SET enabled=0 WHERE id in (" + ",".join([str(nId) for nId in idList]) + ")"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery)
        return bSuc



    def updateRpsContract(self, id, updateList):
        strQuery = "UPDATE rps_contract SET"

        paramList = []
        for item in updateList:
            paramList.append(" {field}='{value}'".format(field=item[0], value=item[1]))

        strQuery += ",".join(paramList)
        strQuery += " WHERE id=%s"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id,))
        return bSuc

    #查询所有的合同，用checklist作为查询筛选
    def getRpsContract(self, checkList):
        strQuery = "SELECT id, client_id, sales_id, client_contact, attachment_id, create_time, delivery_time, due_time, status, creator_id, amount FROM rps_contract"

        paramList = []
        if len(checkList):
            for check in checkList:
                paramList.append("{field}='{value}'".format(field=check[0], value=check[1]))

        if len(paramList):
            strQuery += " WHERE " + " and ".join(paramList)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None
        res = []
        for item in items:
            createTime = ""
            deliveryTime = ""
            dueTime = ""
            try:
                createTime = item[5].strftime("%Y-%m-%d")
            except:
                pass
            try:
                deliveryTime = item[6].strftime("%Y-%m-%d")
            except:
                pass
            try:
                dueTime = item[7].strftime("%Y-%m-%d")
            except:
                pass

            attachment = {}
            if isinstance(item[4], str):
                fileIdList = [num for num in item[4].split(",")]
                for fileId in fileIdList:
                    try:
                        attachment.update({int(float(fileId)): {}})
                    except:
                        pass

            res.append({
                "id": item[0],
                "clientId": item[1],
                "salesId": item[2],
                "clientContact": item[3],
                "attachment": attachment,
                "createTime": createTime,
                "deliveryTime": deliveryTime,
                "dueTime": dueTime,
                "status": item[8],
                "creatorId": item[9],
                "amount": item[10]
            })
        return res

    def getRpsRetirement(self, nId):
        strQuery = "SELECT id, equip_id, status, reason, applicant_id, apply_date, warehouse_id, attachment_id, remark FROM rps_retirement WHERE enabled=1"
        paramList = []
        if nId != None:
            strQuery += " AND id=%s"
            paramList.append(nId)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(paramList))
        if items == None:
            return None

        res = []
        for item in items:
            applyTime = ""
            try:
                applyTime = item[5].strftime("%Y-%m-%d")
            except:
                pass

            attachment = {}
            if isinstance(item[7], str):
                fileIdList = [num for num in item[7].split(",")]
                for fileId in fileIdList:
                    try:
                        attachment.update({int(float(fileId)): {}})
                    except:
                        pass

            res.append({
                "id": item[0],
                "equipId": item[1],
                "status": item[2],
                "reason": item[3],
                "applicantId": item[4],
                "applyTime": applyTime,
                "warehouseId": item[6],
                "remark": item[8],
                "attachment": attachment
            })
        return res

    def addRpsWarehouse(self, insertList):
        strQuery = "INSERT INTO rps_warehouse ("

        fieldList = ["enabled"]
        paramList = [1]
        for item in insertList:
            fieldList.append(item[0])
            paramList.append(item[1])

        strQuery += ",".join(fieldList)
        strQuery += ") VALUES ("
        strQuery += ",".join(["%s"] * len(fieldList))
        strQuery += ")"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, tuple(paramList))
        return bSuc

    def deleteRpsWarehouse(self, id):
        strQuery = "UPDATE rps_warehouse SET enabled=0 WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id,))
        return bSuc

    def enableRpsWarehouse(self, id):
        strQuery = "UPDATE rps_warehouse SET enabled=1 WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id,))
        return bSuc

    def updateRpsWarehouse(self, id, updateList):
        strQuery = "UPDATE rps_warehouse SET"

        paramList = []
        for item in updateList:
            paramList.append(" {field}='{value}'".format(field=item[0], value=item[1]))

        strQuery += ",".join(paramList)
        strQuery += " WHERE id=%s"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (id,))
        return bSuc

    def getRpsWarehouse(self, checkList):
        strQuery = "SELECT id, name, address, admin_id, admin_name, type, location, code FROM rps_warehouse WHERE enabled=1"

        paramList = []
        if len(checkList):
            for item in checkList:
                paramList.append("{field}='{value}'".format(field=item[0], value=item[1]))

        if len(paramList):
            strQuery += " AND " + " AND ".join(paramList)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        res = []
        for item in items:
            lon = ""
            lat = ""
            if isinstance(item[6], str) and item[6].count(",") == 1:
                lon = item[6].split(",")[0]
                lat = item[6].split(",")[1]
            res.append({
                "id": item[0],
                "name": item[1] if isinstance(item[1], str) else "",
                "address": item[2] if isinstance(item[2], str) else "",
                "adminId": item[3] if isinstance(item[3], str) else "",
                "adminName": item[4] if isinstance(item[4], str) else "",
                "type": item[5],
                "location": dict(lon=lon, lat=lat),
                "code": item[7] if isinstance(item[7], str) else ""
            })
        return res

    def getAlreadyExistingAsset(self, assetCodeList):
        cList = []
        for assetCode in assetCodeList:
            cList.append('"{0}"'.format(assetCode))
        strQuery = 'SELECT name, online_addr FROM equip_list WHERE online_addr in (' + ','.join(cList) + ')'
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        return items

    def addRpsEquipMulti(self, postList):
        if not len(postList):
            return True

        strQuery = 'insert into equip_list(id,`default_equipType`, `projId`, `name`, `description`, `online_addr`, `maintenanceStatus`, `repairStatus`,' \
                       ' `repairResponsiblePerson`, `installLocation`, `communicateStatus`, `warningStatus`, `area_id`, `model_id`, `system_id`) values(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)'
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(postList))
        return bSuc

    def getAlreadyExistingTemplateId(self, templateIdList):
        strQuery = 'SELECT id FROM equip_asset_tmpl_def WHERE id IN (' + ','.join(templateIdList) + ')'
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        cList = []
        for item in items:
            if len(item):
                if item[0] != None:
                    cList.append(str(item[0]))
        return cList

    def keywordSearchRpsAssetCode(self, keyword):
        strQuery = "SELECT id, name, description, online_addr, maintenanceStatus, repairStatus, communicateStatus, warningStatus, model_id FROM equip_list WHERE online_addr LIKE '%" + keyword + "%'"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return items

        if not isinstance(items, list):
            return None
        res = []
        for item in items:
            res.append({
                "id": item[0],
                "name": item[1],
                "assetCode": item[3],
                "assetCodeStatus": item[2],
                "maintenanceStatus": item[4],
                "repairStatus": item[5],
                "communicateStatus": item[6],
                "warningStatus": item[7],
                "templateId": item[8]
            })
        return res

    """在表tableName中，字段fieldName的值为value的条目存在于除 字段besidesFieldName 的值为tarValue的条目之外
    """
    def fieldExistsInTableBesides(self, fieldName, value, tableName, besidesFieldName, tarValue):
        strQuery = "SELECT {field} FROM {table} WHERE {field}='{value}' AND {besidesFieldName} != '{tarValue}'".format(field=fieldName, table=tableName, value=value, besidesFieldName=besidesFieldName, tarValue=tarValue)
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return True
        if not len(items):
            return False
        if not len(items[0]):
            return False
        return True

    def getRpsRFIDByAssetCodeList(self, assetCodeList):
        if not len(assetCodeList):
            return {}

        strQuery = "SELECT asset_code, label_index, asset_template_id, rfid_code, creator_id, create_time, actuator_id FROM rps_rfid where asset_code in (" + ",".join(['"{0}"'.format(assetCode) for assetCode in assetCodeList]) + ")"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        res = {}
        for item in items:
            createTime = ""
            try:
                createTime = item[5].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            assetCode = item[0]
            if assetCode not in res.keys():
                res.update({assetCode: []})

            res[assetCode].append(
                dict(assetCode=item[0],
                    index=item[1],
                    templateId=item[2],
                    rfidCode=item[3],
                    creatorId=item[4],
                    createTime=createTime,
                    actuatorId=item[6])
            )

        return res

    def getRpsRFIDByRFIDCodeList(self, rfidCodeList):
        if not len(rfidCodeList):
            return {}

        strQuery = "SELECT asset_code, label_index, asset_template_id, rfid_code, creator_id, create_time, actuator_id FROM rps_rfid where rfid_code in (" + ",".join(['"{0}"'.format(rfidCode) for rfidCode in rfidCodeList]) + ")"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        res = {}
        for item in items:
            createTime = ""
            try:
                createTime = item[5].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            rfidCode = item[3]

            res.update({rfidCode: dict(assetCode=item[0],
                                       index=item[1],
                                       templateId=item[2],
                                       creatorId=item[4],
                                       createTime=createTime,
                                       actuatorId=item[6])})
        return res

    def getLabelIndexAndRFIDOfAsset(self, assetCode):
        strQuery = "SELECT label_index, rfid_code FROM rps_rfid WHERE asset_code=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (assetCode,))
        if items == None:
            return None, None
        indexList = []
        rfidCodeList = []
        for item in items:
            if len(item):
                indexList.append(item[0])
                rfidCodeList.append(item[1])
        return indexList, rfidCodeList

    def addRFID(self, assetCode, nLabelIndex, assetTemplateId, strFRID):
        strQuery = "INSERT INTO rps_rfid (asset_code, label_index, asset_template_id, rfid_code) VALUES (%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (assetCode, nLabelIndex, assetTemplateId, strFRID))
        return bSuc

    def searchRFID(self, rfidCode):
        strQuery = "SELECT asset_code, label_index, asset_template_id, rfid_code FROM rps_rfid WHERE rfid_code=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (rfidCode,))
        if items == None:
            return None

        res = {}
        for item in items:
            res.update({
                "assetCode": item[0],
                "labelIndex": int(item[1]),
                "assetTemplateId": int(item[2]),
                "rfidCode": item[3]
            })
        return res

    def deleteRFID(self, rfidCode):
        strQuery = "DELETE FROM rps_rfid WHERE rfid_code=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (rfidCode,))
        return bSuc

    def rfidEventSubmit(self, param):
        strQuery = "INSERT INTO rps_rfid_log (id, asset_code, rfid_code, time, receiver_id, warehouse_id, quality, quality_desc, enabled) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(param))
        return bSuc

    def addRpsAttachment(self, infoList):
        strQuery = "INSERT INTO rps_attachment (id, file_name, file_type, oss_path, create_time, enabled) VALUES (%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(infoList))
        return bSuc

    def getRpsAttatchmentInfo(self, fileIdList):
        if not len(fileIdList):
            return {}

        strQuery = "SELECT id, file_name, file_type, oss_path, create_time FROM rps_attachment"
        if len(fileIdList):
            strQuery += " WHERE id in (" + ",".join(str(fileId) for fileId in fileIdList) + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        dResult = {}
        for item in items:
            fileInfo = {
                "id": item[0],
                "fileName": item[1],
                "fileType": item[2],
                "ossPath": item[3],
                "createTime": item[4].strftime("%Y-%m-%d %H:%M:%S")
            }
            dResult.update({item[0]: fileInfo})
        return dResult

    def addAssetToContract(self, param):
        strQuery = "INSERT INTO rps_contract_asset_template (contract_id, template_id, quantity) VALUES (%s, %s, %s)"
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(param))
        return bSuc

    def deleteAssetFromContract(self, contractId, templateId):
        strQuery = "DELETE FROM rps_contract_asset_template WHERE contract_id=%s AND template_id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (contractId, templateId))
        return bSuc

    def updateAssetFromContract(self, contractId, templateId, quantity):
        # strQuery = "UPDATE rps_contract_asset_template SET quantity=%s WHERE contract_id=%s AND template_id=%s"
        strQuery = "REPLACE INTO rps_contract_asset_template (contract_id, template_id, quantity) VALUES (%s, %s, %s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (contractId, templateId, quantity))
        return bSuc

    def getAssetOfContract(self, contractId):
        strQuery = "SELECT template_id, quantity FROM rps_contract_asset_template WHERE contract_id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (contractId,))
        if items == None:
            return None

        res = {}
        for item in items:
            res.update({
                int(float(item[0])): int(float(item[1]))
            })
        return res

    #获取某个出入库单的资产详情，资产模板号:该资产模板数量
    def getAssetOfInOutWarehouse(self, nInOutId):
        strQuery = "SELECT template_id, quantity FROM rps_in_out_warehouse_detail WHERE in_out_id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (nInOutId,))
        if items == None:
            return None

        res = {}
        for item in items:
            res.update({
                int(float(item[0])): int(float(item[1]))
            })
        return res


    def getEquipAssetTmpl(self, templateIdPreCheckList):
        strQuery = "SELECT tmpl_def_id, cn_name, en_name FROM equip_asset_tmpl WHERE tmpl_def_id in (" + ",".join(templateIdPreCheckList) + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return {}
        res = {}
        for item in items:
            tmpl_def_id = int(float(item[0]))
            if tmpl_def_id not in res.keys():
                res.update({tmpl_def_id: {}})

            res.get(tmpl_def_id).update({item[1]: item[2]})

        return res

    def updateAssetMulti(self,paramList):
        if not len(paramList):
            return True

        equipIdList = []
        for param in paramList:
            equipIdList.append(param[0])

        sql = 'delete from equip_asset where equip_id in (' + ",".join([str(equipId) for equipId in equipIdList]) + ")"
        bSuc1 = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sql)

        sql = 'INSERT INTO `equip_asset`(`equip_id`,  `en_name`, `param_value`, `project_id`) VALUES (%s, %s, %s, %s)'
        bSuc2 = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], sql, tuple(paramList))
        return bSuc1 and bSuc2

    "nType: 单据类型 1-销售订单;2-入库单;3-出库单;4-物流单;5-报废单"
    def getRpsStatus(self, nType, nBillId):
        strTableName = ""
        if nType == 1:
            strTableName = "rps_contract"
        elif nType in [2, 3]:
            strTableName = "rps_in_out_warehouse"
        elif nType == 4:
            strTableName = "rps_transport"
        elif nType == 5:
            strTableName = "rps_retirement"

        if not strTableName:
            return None

        strQuery = "SELECT status FROM {table} WHERE id=%s".format(table=strTableName)
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (nBillId,))
        if items == None:
            return None
        try:
            return int(float(items[0][0]))
        except:
            return None

    def getRpsContractAssetInfo(self, contractIdList):
        strQuery = "SELECT rcat.contract_id, rcat.template_id, rcat.quantity, eatd.name, eatd.describe FROM rps_contract_asset_template rcat LEFT JOIN equip_asset_tmpl_def eatd ON rcat.template_id=eatd.id"
        if len(contractIdList):
            strQuery += " WHERE rcat.contract_id in (" + ",".join([str(nId) for nId in contractIdList]) + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        res = {}
        for item in items:
            contractId = int(float(item[0]))
            templateId = item[1]
            quantity = item[2]
            templateName = item[3]
            templateDescription = item[4]
            if contractId not in res.keys():
                res.update({contractId: []})

            res.get(contractId).append(dict(quantity=quantity,
                                            templateName=templateName,
                                            description=templateDescription,
                                            templateId=templateId))

        return res

    def addTransport(self, nId, inoutwarehouseId, departWarehouseId, destWarehouseId, estimatedArriveDate, nStatus):
        strQuery = "INSERT INTO rps_transport (id, inoutwarehouse_id, depart_warehouse_id, dest_warehouse_id, estimated_arrive_time, status) VALUES (%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (nId, inoutwarehouseId, departWarehouseId, destWarehouseId, estimatedArriveDate, nStatus))
        return bSuc
    def getRpsAttachmentIdListOfTransportReceiveImg(self, transportId):
        strQuery = "SELECT receive_delivery_img FROM rps_transport WHERE id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (transportId,))
        if items == None:
            return None
        if not isinstance(items, list) or not len(items):
            return None

        if not len(items[0]):
            return None

        if items[0][0] == None or items[0][0] == "":
            return []

        res = []
        idList = items[0][0].split(",")
        for nId in idList:
            try:
                res.append(int(float(nId)))
            except:
                pass
        return res

    def updateRpsTransportReceiveImg(self, transportId, attachmentIdList):
        strAttachment = ",".join([str(nId) for nId in attachmentIdList])
        strQuery = "UPDATE rps_transport SET receive_delivery_img = %s WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (strAttachment, transportId))
        return bSuc

    def getRpsAttachmentIdListOfTransportSignImg(self, transportId):
        strQuery = "SELECT sign_img FROM rps_transport WHERE id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (transportId,))
        if items == None:
            return None
        if not isinstance(items, list) or not len(items):
            return None

        if not len(items[0]):
            return None

        if items[0][0] == None or items[0][0] == "":
            return []

        res = []
        idList = items[0][0].split(",")
        for nId in idList:
            try:
                res.append(int(float(nId)))
            except:
                pass
        return res

    def updateRpsTransportSignImg(self, transportId, attachmentIdList):
        strAttachment = ",".join([str(nId) for nId in attachmentIdList])
        strQuery = "UPDATE rps_transport SET sign_img = %s WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (strAttachment, transportId))
        return bSuc

    def removeRpsTransport(self, transportId):
        strQuery = "DELETE FROM rps_transport WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (transportId,))
        return bSuc

    def updateRpsTransport(self, transportId, updateList):
        strQuery = "UPDATE rps_transport SET "
        setList = []
        for update in updateList:
            setList.append("{field}='{value}'".format(field=update[0], value=update[1]))
        strQuery += ",".join(setList)
        strQuery += " WHERE id=%s"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (transportId,))
        return bSuc

    def updateRpsAuditRecved(self, nid):

        strQuery = "UPDATE rps_audit SET recved=1"
        if nid != None:
            strQuery += " WHERE id=%s"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery,(nid,))
        return bSuc

    def getAllEquipAssetTmpl(self, ignoreList=[]):
        strQuery = "SELECT id, name, `describe` FROM equip_asset_tmpl_def"
        if len(ignoreList):
            strQuery = "SELECT id, name, `describe` FROM equip_asset_tmpl_def WHERE id not in (" + ",".join([str(nId) for nId in ignoreList]) + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        res = []
        for item in items:
            res.append({
                "id": item[0],
                "name": item[1],
                "description": item[2]
            })
        return res


    def getRpsTransport(self, transportId):
        strQuery = "SELECT id, inoutwarehouse_id, depart_warehouse_id, dest_warehouse_id, estimated_arrive_time, status, receive_delivery_img, sign_img FROM rps_transport"
        param = ()
        if transportId != None:
            strQuery += " WHERE id=%s"
            param = (transportId,)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, param)
        if items == None:
            return None

        res = []
        for item in items:
            nId = item[0]
            inoutwarehouseId = item[1]
            departWarehouseId = item[2]
            destWarehouseId = item[3]
            estimatedArriveTime = item[4].strftime("%Y-%m-%d") if item[4] != None and item[4] != "" else ""
            status = item[5]
            strReceive = item[6]
            strSign = item[7]

            receiveIdList = []
            if strReceive != None:
                rList = strReceive.split(",")
                for tar in rList:
                    try:
                        receiveIdList.append(int(float(tar)))
                    except:
                        pass

            signIdList = []
            if strSign != None:
                sList = strSign.split(",")
                for tar in sList:
                    try:
                        signIdList.append(int(float(tar)))
                    except:
                        pass

            res.append(dict(transportId=nId,
                            inoutwarehouseId=inoutwarehouseId,
                            departWarehouseId=departWarehouseId,
                            destWarehouseId=destWarehouseId,
                            estimatedArriveTime=estimatedArriveTime,
                            status=status,
                            receiveImgList=receiveIdList,
                            signImgList=signIdList))

        return res

    def addRpsInOutWarehouse(self, nId, warehouseId, transportId, nType, creatorId, strCreateTime,
                                                             nStatus, nEnabled, additional):
        strQuery = "INSERT INTO rps_in_out_warehouse (id, warehouse_id, transport_id, type, creator_id, create_time, status, enabled"
        strParam = "(%s,%s,%s,%s,%s,%s,%s,%s"
        paramList = [nId, warehouseId, transportId, nType, creatorId, strCreateTime, nStatus, nEnabled]
        if additional != None:
            strQuery += " ,additional"
            strParam += ",%s"
            paramList.append(additional)

        strQuery += ") VALUES "
        strParam += ")"

        strInsert = strQuery + strParam

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strInsert, tuple(paramList))
        return bSuc

    def getRpsAttachmentIdListInOutWarehouse(self, nId):
        strQuery = "SELECT attachment_id FROM rps_in_out_warehouse WHERE id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (nId,))
        if items == None:
            return None

        if not len(items):
            return None
        if not len(items[0]):
            return None
        if items[0][0] == None or items[0][0] == "":
            return []

        rawList = items[0][0].split(",")
        idList = []
        for raw in rawList:
            try:
                idList.append(int(float(raw)))
            except:
                pass
        return idList

    def updateRpsInOutWarehouse(self, nId, updateList):
        if not len(updateList):
            return True

        strQuery = "UPDATE rps_in_out_warehouse SET "
        setList = []
        for update in updateList:
            setList.append("{fieldName}='{value}'".format(fieldName=update[0], value=update[1]))

        strQuery += ",".join(setList)
        strQuery += " WHERE id=%s"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (nId,))
        return bSuc

    def addAssetToInOutWarehouse(self, paramList):
        strQuery = "INSERT INTO rps_in_out_warehouse_detail (in_out_id, template_id, quantity) VALUES (%s, %s, %s)"
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(paramList))
        return bSuc

    def deleteRpsInOutWarehouseAsset(self, inOutId, templateId=None):
        strQuery = "DELETE FROM rps_in_out_warehouse_detail WHERE in_out_id=%s"
        param = [inOutId]

        if templateId != None:
            strQuery += " AND template_id=%s"
            param.append(templateId)

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, tuple(param))
        return bSuc

    def updateRpsInOutWarehouseAssetQuantity(self, inOutId, templateId, quantity):
        # strQuery = "UPDATE rps_in_out_warehouse_detail SET quantity=%s WHERE in_out_id=%s AND template_id=%s"
        strQuery = "REPLACE INTO rps_in_out_warehouse_detail (in_out_id, template_id, quantity) VALUES (%s, %s, %s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (inOutId, templateId, quantity))
        return bSuc

    def getInOutWarehouse(self, nId):
        strQuery = "SELECT id, warehouse_id, contract_id, type, creator_id, create_time, status, attachment_id, additional FROM rps_in_out_warehouse WHERE enabled=1"
        param = []
        if nId != None:
            strQuery += " AND id=%s"
            param.append(nId)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return None

        dataList = []
        for item in items:
            attachment = {}
            if item[7] != None and item[7] != "":
                lAttach = item[7].split(",")
                for attach in lAttach:
                    try:
                        attachment.update({
                            int(float(attach)): {}
                        })
                    except:
                        pass
            dataList.append({
                "id": item[0],
                "warehouseId": item[1],
                "contractId": item[2],
                "type": item[3],
                "creatorId": item[4],
                "createTime": item[5].strftime("%Y-%m-%d %H:%M:%S") if item[5] != None else "",
                "status": item[6],
                "attachment": attachment,
                "additional": item[8],
                "asset": []
            })
        return dataList

    #批量获取指定多个出入库单的资产内容详情
    def getRpsInOutAsset(self, inOutIdList):
        if not len(inOutIdList):
            return {}

        strQuery = "SELECT riowd.in_out_id, riowd.template_id, riowd.quantity, eatd.name, eatd.describe FROM rps_in_out_warehouse_detail riowd LEFT JOIN equip_asset_tmpl_def eatd ON riowd.template_id=eatd.id WHERE riowd.in_out_id in (" + ",".join([str(nId) for nId in inOutIdList]) + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        res = {}
        for item in items:
            inOutId = item[0]
            if inOutId not in res.keys():
                res.update({inOutId: []})

            res.get(inOutId).append({
                "templateId": item[1],
                "quantity": int(float(item[2])),
                "templateName": item[3],
                "description": item[4]
            })
        return res

    def executeRpsAudit(self, billId, nType, userId, targetStatus, opinion):
        auditTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        strQuery = "UPDATE rps_audit SET status=%s, auditor_id=%s, opinion=%s, audit_time=%s WHERE bill_id=%s AND type=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (targetStatus, userId, opinion, auditTime, billId, nType))
        return bSuc

    def removeRpsAudit(self, billId, nTypeOrList):
        if not isinstance(nTypeOrList, list):
            nTypeOrList = [nTypeOrList]

        strQuery = "DELETE FROM rps_audit WHERE bill_id=%s AND type in (" + ",".join([str(nType) for nType in nTypeOrList]) + ")"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (billId,))
        return bSuc

    def updateRpsMultiTrans(self, queryList, paramList):
        if not len(queryList) or len(queryList) != len(paramList):
            return True

        bSuc = self._mysqlDbContainer.op_db_transaction_update_many(app.config['DATABASE'], queryList, paramList)
        return bSuc

    def getInstallationOfDevice(self, assetCodeList):
        if not len(assetCodeList):
            return {}

        strQuery = "SELECT online_addr, installLocation FROM equip_list WHERE online_addr in (" + ",".join([str(assetCode) for assetCode in assetCodeList]) + ")"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return {}

        res = {}
        for item in items:
            if not len(item):
                continue

            res.update({item[0]: int(float(item[1]))})
        return res


    def addRpsAuditSingle(self, paramList):
        strInsert = "INSERT INTO rps_audit (id, bill_id, type, create_time, creator_id, audit_type, original, target, deadline, enabled, result) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strInsert, tuple(paramList))
        return bSuc

    def getRpsEditAuditTarget(self, nBillId, nType, nAuditType, nAuditId):
        strQuery = "SELECT target FROM rps_audit WHERE bill_id=%s AND type=%s AND audit_type=%s AND id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (nBillId, nType, nAuditType, nAuditId))
        if items == None:
            return None
        try:
            res = json.loads(items[0][0])
            return res
        except:
            return None

    # nId: 单据ID    nType: 单据类型: 1-销售订单;2-出入库单;3-出库单(暂不用);4-物流单(暂不用);5-报废单   nFlag: 结果类型  -1: 未审核；0：不通过；1：通过
    def getRpsAuditRecord(self, nId, nType, nFlag=None):
        strQuery = "SELECT id, creator_id, create_time, original, target, audit_time, auditor_id, opinion, " \
                   "deadline, result, audit_type FROM rps_audit WHERE bill_id=%s AND type=%s AND enabled=1"
        param = [nId, nType]

        if nFlag != None:
            strQuery += " AND result=%s"
            param.append(nFlag)

        strQuery += " ORDER BY create_time ASC"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return None

        dataList = []
        for item in items:
            createTime = ""
            try:
                createTime = item[2].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            original = {}
            try:
                original = json.loads(item[3])
            except:
                pass

            target = {}
            try:
                target = json.loads(item[4])
            except:
                pass

            auditTime = ""
            try:
                auditTime = item[5].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            deadline = ""
            try:
                deadline = item[8].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            try:
                original = json.loads(item[3])
            except:
                pass

            auditorID = item[6]  # 从拿到的审核单里获取auditorid
            if auditorID != None:
                dUserInfo = BEOPDataAccess.getInstance().get_user_info(int(auditorID))  # 通过auditorid获取该用户的英文名、中文名等信息
                Userfullname = dUserInfo.get('userfullname')
                Username = dUserInfo.get('username')
                auditorname = ""
                if Userfullname != None:
                    auditorname = Userfullname
                else:
                    auditorname = Username
            else:
                auditorname = ''
            auditType = item[10]
            if original=={} and target== {} and auditType==2:#过滤掉废单
                continue
            dataList.append(dict(
                nAuditId=item[0],
                creatorId=item[1],
                createTime=createTime,
                original=original,
                target=target,
                auditorId=item[6],
                opinion=item[7],
                deadline=deadline,
                result=item[9],
                auditTime=auditTime,
                auditType=item[10],
                auditorname=auditorname

            ))
        #if dataList[-1]['original']=={}:
            #del dataList[-1]['original']
        #if dataList[-1]['target']=={}:
            #del dataList[-1]['target']

        return dataList

    def addRpsLBSEventLog(self, nId, strTime, nLbsId, strLbsXy, strLbsPosition, strDataPackage, strAssetCode, nQuality, strQualityDesc, nEnabled):
        strQuery = "INSERT INTO rps_locate_plate_log (id, time, plate_id, lbs_xy, lbs_position, data_package, asset_code, quality, quality_desc, enabled) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (nId, strTime, nLbsId, strLbsXy,
                                                                                      strLbsPosition, strDataPackage,
                                                                                      strAssetCode, nQuality, strQualityDesc, nEnabled))
        return bSuc


    # 获取同一类型的一批单据的所有待校核、审核项
    def getRpsCheckAndAuditRecords(self, nIdList, nBillType):
        if not len(nIdList):
            return {}

        if isinstance(nBillType, int):
            strQuery = "SELECT id, bill_id, creator_id, create_time, audit_time, auditor_id, opinion, result, deadline, audit_type, original, target FROM rps_audit WHERE enabled=1 AND type=%s AND bill_id in (" + ",".join([str(nId) for nId in nIdList]) + ")"
            param = [nBillType]
        else:
            strQuery = "SELECT id, bill_id, creator_id, create_time, audit_time, auditor_id, opinion, result, deadline, audit_type, original, target FROM rps_audit WHERE enabled=1 AND type in ("+ ",".join([str(nItem) for nItem in nBillType]) +") AND bill_id in (" + ",".join(
                [str(nId) for nId in nIdList]) + ")"
            param = []

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return None

        res = {}
        for item in items:
            nBillId = item[1]
            if nBillId not in res.keys():
                res.update({nBillId: {}})

            auditType = item[9]
            createTime = ""
            try:
                createTime = item[3].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            auditTime = ""
            try:
                auditTime = item[4].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            deadline = ""
            try:
                deadline = item[8].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            original = {}
            try:
                original = json.loads(item[10])
            except:
                pass

            target = {}
            try:
                target = json.loads(item[11])
            except:
                pass

            if not res.get(nBillId, {}):
                res[nBillId].update(dict(check=[], createAudit=[], editAudit=[], terminateAudit=[], finishAudit=[]))

            strKey = None
            if auditType == 0:
                strKey = "check"
            elif auditType == 1:
                strKey = "createAudit"
            elif auditType == 2:
                strKey = "editAudit"
            elif auditType == 3:
                strKey = "terminateAudit"
            elif auditType == 4:
                strKey = "finishAudit"

            if strKey == None:
                continue

            res[nBillId][strKey].append(dict(nAuditId=item[0],
                                                 creatorId=item[2],
                                                 createTime=createTime,
                                                 auditTime=auditTime,
                                                 auditorId=item[5],
                                                 opinion=item[6],
                                                 result=item[7],
                                                 deadline=deadline,
                                                 original=original,
                                                 target=target))

            # if not res[nBillId].get(strKey, []):
            #     res[nBillId][strKey].update(dict(nAuditId=item[0],
            #                                      creatorId=item[2],
            #                                      createTime=createTime,
            #                                      auditTime=auditTime,
            #                                      auditorId=item[5],
            #                                      opinion=item[6],
            #                                      result=item[7],
            #                                      deadline=deadline))
            # else:
            #     if datetime.strptime(createTime, "%Y-%m-%d %H:%M:%S") > datetime.strptime(res[nBillId].get(strKey).get("createTime"), "%Y-%m-%d %H:%M:%S"):
            #         res[nBillId][strKey].update(dict(nAuditId=item[0],
            #                                           creatorId=item[2],
            #                                           createTime=createTime,
            #                                           auditTime=auditTime,
            #                                           auditorId=item[5],
            #                                           opinion=item[6],
            #                                           result=item[7],
            #                                           deadline=deadline))

        return res

    def getRelatedAssetFromRFIDInPeriod(self, warehouseId, strBegin, strEnd):
        strQuery = "SELECT rrl.asset_code, rrl.time, rrl.receiver_id, rrl.quality, rrl.quality_desc, el.model_id, eatd.name AS templateName, eatd.describe AS templateDescription FROM rps_rfid_log rrl LEFT JOIN equip_list el ON rrl.asset_code=el.online_addr LEFT JOIN equip_asset_tmpl_def eatd ON el.model_id=eatd.id WHERE rrl.warehouse_id=%s AND rrl.time >= %s AND rrl.time <= %s AND rrl.enabled=1 ORDER BY rrl.time ASC"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (warehouseId, strBegin, strEnd))
        if items == None:
            return None

        dataList = []
        for item in items:
            assetCode = item[0]
            strTime = ""
            try:
                strTime = item[1].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass
            receiverId = item[2]
            quality = item[3]
            qualityDesc = item[4]
            templateId = item[5]
            templateName = item[6]
            templateDescription = item[7]

            dataList.append(dict(assetCode=assetCode,
                                 time=strTime,
                                 receiverId=receiverId,
                                 quality=quality,
                                 qualityDesc=qualityDesc,
                                 templateId=templateId,
                                 templateName=templateName,
                                 description=templateDescription
                                 ))

        return dataList

    def getRelatedAssetFromBeaconInPeriod(self, warehouseId, strBegin, strEnd):
        strQuery = "SELECT rbgl.asset_code, rbgl.time, rbgl.gateway_id, rbgl.quality, rbgl.quality_desc, el.model_id, eatd.name AS templateName, eatd.describe AS templateDescription FROM rps_bluetooth_gateway_log rbgl LEFT JOIN equip_list el ON rbgl.asset_code=el.online_addr LEFT JOIN equip_asset_tmpl_def eatd ON el.model_id=eatd.id WHERE rbgl.warehouse_id=%s AND rbgl.time >= %s AND rbgl.time <= %s AND rbgl.enabled=1 ORDER BY rbgl.time ASC"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (warehouseId, strBegin, strEnd))
        if items == None:
            return None

        dataList = []
        for item in items:
            assetCode = item[0]
            strTime = ""
            try:
                strTime = item[1].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass
            gatewayId = item[2]
            quality = item[3]
            qualityDesc = item[4]
            templateId = item[5]
            templateName = item[6]
            templateDescription = item[7]

            dataList.append(dict(assetCode=assetCode,
                                 time=strTime,
                                 gatewayId=gatewayId,
                                 quality=quality,
                                 qualityDesc=qualityDesc,
                                 templateId=templateId,
                                 templateName=templateName,
                                 description=templateDescription
                                 ))

        return dataList

    def getRelatedAssetFromLbsInPeriod(self, warehouseId, strBegin, strEnd):
        strQuery = "SELECT rlpl.asset_code, rlpl.time, rlpl.plate_id, rlpl.lbs_xy, rlpl.lbs_position, rlpl.quality, rlpl.quality_desc, el.model_id, eatd.name AS templateName, eatd.describe AS templateDescription FROM rps_locate_plate_log rlpl LEFT JOIN equip_list el ON rlpl.asset_code=el.online_addr LEFT JOIN equip_asset_tmpl_def eatd ON el.model_id=eatd.id WHERE rlpl.lbs_predict_warehouse_id=%s AND rlpl.time >= %s AND rlpl.time <= %s AND rlpl.enabled=1 ORDER BY rlpl.time ASC"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (warehouseId, strBegin, strEnd))
        if items == None:
            return None

        dataList = []
        for item in items:
            assetCode = item[0]
            strTime = ""
            try:
                strTime = item[1].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            lon = None
            lat = None
            try:
                lon = item[3].split(",")[0]
                lat = item[3].split(",")[1]
            except:
                pass

            plateId = item[2]
            lbsPosition = item[4]
            quality = item[5]
            qualityDesc = item[6]
            templateId = item[7]
            templateName = item[8]
            templateDescription = item[9]

            dataList.append(dict(assetCode=assetCode,
                                 time=strTime,
                                 plateId=plateId,
                                 quality=quality,
                                 qualityDesc=qualityDesc,
                                 templateId=templateId,
                                 templateName=templateName,
                                 description=templateDescription,
                                 lon=lon,
                                 lat=lat
                                 ))

        return dataList

    def addRetirement(self, nId, equipId, userId, warehouseId, reason, nStatus, strApplyTime, strAttachmentId, nEnabled, remark):
        strQuery = "INSERT INTO rps_retirement (id, equip_id, status, reason, applicant_id, apply_date, warehouse_id, attachment_id, enabled, remark) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (nId, equipId, nStatus, reason, userId, strApplyTime, warehouseId, strAttachmentId, nEnabled, remark))
        return bSuc

    def updateRpsRetirement(self, nId, updateList):
        strQuery = "UPDATE rps_retirement SET"

        paramList = []
        for item in updateList:
            paramList.append(" {field}='{value}'".format(field=item[0], value=item[1]))

        strQuery += ",".join(paramList)
        strQuery += " WHERE id=%s"

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (nId,))
        return bSuc

    # 获取审核单信息
    def getRpsAuditInfoByAuditId(self, nAuditIdList, nBillType=None, nAuditType=None, nResult=None):
        if not isinstance(nAuditIdList, list):
            return {}
        if not len(nAuditIdList):
            return {}
        strQuery = "SELECT id, bill_id, type, creator_id, create_time, audit_type, original, target, audit_time, auditor_id, opinion, result, deadline FROM rps_audit WHERE enabled=1 AND id in (" + ",".join([str(nId) for nId in nAuditIdList]) + ")"
        param = []

        if nBillType != None:
            strQuery += " AND type=%s"
            param.append(nBillType)

        if nAuditType != None:
            strQuery += " AND audit_type=%s"
            param.append(nAuditType)

        if nResult != None:
            strQuery += " AND result=%s"
            param.append(nResult)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return None

        res = {}
        for item in items:
            createTime = ""
            try:
                createTime = item[4].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            auditTime = ""
            try:
                auditTime = item[8].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            deadline = ""
            try:
                deadline = item[12].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass

            original = ""
            try:
                original = json.loads(item[6])
            except:
                pass

            target = ""
            try:
                target = json.loads(item[7])
            except:
                pass

            res.update({item[0]: dict(billId=item[1],
                                      billType=item[2],
                                      creatorId=item[3],
                                      createTime=createTime,
                                      auditType=item[5],
                                      original=original,
                                      target=target,
                                      auditTime=auditTime,
                                      auditorId=item[9],
                                      opinion=item[10],
                                      result=item[11],
                                      deadline=deadline)})
        return res

    def getRpsCurrentAttachmentIdList(self, nId, strField, strTable):
        strQuery = "SELECT {field} FROM {tableName} WHERE id=%s".format(tableName=strTable, field=strField)
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (nId,))
        if items == None:
            return None

        nIdList = []
        for item in items:
            for obj in item:
                if obj == None:
                    continue

                if not isinstance(obj, str):
                    continue

                try:
                    strIdList = obj.split(",")
                    for strId in strIdList:
                        nIdList.append(int(float(strId)))
                except:
                    pass

        nIdList = list(set(nIdList))
        return nIdList

    def getAllRpsAuditRecord(self, nType=None,nResult=None,nCreatorid=None,nContractid=None,nInoutwarehouseid=None):
        strQuery = "SELECT bill_id, type, creator_id, create_time,original,target,result,recved,audit_type FROM rps_audit WHERE enabled=1"
        param = []
        #
        if nType != None:
            strQuery += " AND type=%s"
            param.append(nType)
        if nResult != None:
            strQuery += " AND result=%s"
            param.append(nResult)
        if nCreatorid != None:
            strQuery += " AND creator_id=%s"
            param.append(nCreatorid)
        if nContractid != None:
            strQuery += " AND bill_id=%s"
            param.append(nContractid)
        if nInoutwarehouseid != None:
            strQuery += " AND bill_id=%s"
            param.append(nInoutwarehouseid)
        strQuery += " ORDER BY type ASC"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        #items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        dataList = []
        for item in items:
            createTime = ""
            try:
                createTime = item[3].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass
            original = {}
            try:
                original = json.loads(item[4])
            except:
                pass

            target = {}
            try:
                target = json.loads(item[5])
            except:
                pass


            nResult = -10
            try:
                nResult = int(float(item[6]))
            except:
                pass


            nRecved = -10
            try:
                nRecved = int(float(item[7]))
            except:
                pass

            nAutidType = -10
            try:
                nAutidType = int(float(item[8]))
            except:
                pass
            if original=={} and target== {} and nAutidType==2:
                continue
            if item[1]==1:
                dataList.append(dict(contractID=item[0],
                                     type=item[1],
                                     creatorId=item[2],
                                     createTime=createTime,
                                     original =original,
                                     target =target,
                                     result= nResult,
                                     recved = nRecved,
                                     auditType = nAutidType))

            if item[1] == 2:
                dataList.append(dict(inoutwarehouseID=item[0],
                                     type=item[1],
                                     creatorId=item[2],
                                     createTime=createTime,
                                     original=original,
                                     target=target,
                                     result=nResult,
                                     recved=nRecved,
                                     auditType=nAutidType))
            if item[1] == 5:
                dataList.append(dict(retirementID=item[0],
                                     type=item[1],
                                     creatorId=item[2],
                                     createTime=createTime,
                                     original=original,
                                     target=target,
                                     result=nResult,
                                     recved=nRecved,
                                     auditType=nAutidType))
        return dataList

    def getContractByCreatorID(self, nCreatorID=None,nBillID=None):
        strQuery = "SELECT id,creator_id,status,create_time FROM rps_contract WHERE "

        param = []
        if nCreatorID != None:
            strQuery += " creator_id=%s"
            param.append(nCreatorID)
        if nBillID != None:
            strQuery += " AND id=%s"
            param.append(nBillID)
        strQuery += " ORDER BY create_time ASC"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        #items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        dataList = []
        for item in items:
            createTime = ""
            try:
                createTime = item[3].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass


            dataList.append(dict(contractID=item[0],
                                 creatorId=item[1],
                                 Status=item[2],
                                 createTime=createTime
                                        ))
        return dataList

    def getInoutwarehouseByCreatorID(self, nCreatorID=None,nBillID=None):
        strQuery = "SELECT id,creator_id,type,create_time FROM rps_in_out_warehouse WHERE enabled=1"
        #
        param = []
        if nCreatorID != None:
            strQuery += " AND creator_id=%s"
            param.append(nCreatorID)
        if nBillID != None:
            strQuery += " AND id=%s"
            param.append(nBillID)
        strQuery += " ORDER BY create_time ASC"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        #items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        dataList = []
        for item in items:
            createTime = ""
            try:
                createTime = item[3].strftime("%Y-%m-%d %H:%M:%S")
            except:
                pass



            dataList.append(dict(inoutwarehouseID=item[0],
                                 creatorId=item[1],
                                 Type=item[2],
                                 createTime=createTime
                                        ))
        return dataList

    def addCommentToWorkOrder(self, nCommentId, orderId, userName, content, imgNameList):
        strInsert = "INSERT INTO work_order_comment (id, orderId, creator, createTime,"
        sqlList = []
        paramList = [nCommentId, orderId, userName, datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
        if content != None:
            sqlList.append("content")
            paramList.append(content)

        if imgNameList != None:
            sqlList.append("img")
            paramList.append(",".join([str(item) for item in imgNameList]))

        strInsert += ",".join(sqlList)
        strInsert += ") VALUES ("
        strInsert += ",".join(["%s"] * len(paramList))
        strInsert += ")"

        if len(paramList) < 5:
            return False

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strInsert, tuple(paramList))
        return bSuc

    def deleteCommentFromWorkOrder(self, commentId):
        strQuery = "DELETE FROM work_order_comment WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (commentId,))
        return bSuc

    def getWorkOrderCommentDetail(self, commentId):
        strQuery = "SELECT id, orderId, creator, createTime, modifyTime, content, img FROM work_order_comment WHERE id=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (commentId,))
        if items == None:
            return None
        if not isinstance(items, list):
            return None

        dData = {}
        for item in items:
            commentId = item[0]
            imgList = []
            if isinstance(item[6], str):
                tarList = item[6].split(",")
                for tar in tarList:
                    if not len(tar):
                        continue
                    imgList.append(tar)

            if commentId not in dData.keys():
                dData.update({commentId: {}})
            dData.get(commentId).update(dict(id=commentId,
                                             orderId=item[1],
                                             creator=item[2],
                                             createTime=item[3],
                                             modifyTime=item[4],
                                             content=item[5],
                                             imgList=imgList))
        return dData

    def editCommentOfWorkOrder(self, commentId, content):
        modifyTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        strQuery = "UPDATE work_order_comment SET content=%s, modifyTime=%s WHERE id=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (content, modifyTime, commentId))
        return bSuc

    def getCommentOfWorkOrder(self, workOrderId=None):
        strQuery = "SELECT id, orderId, creator, createTime, content, img FROM work_order_comment ORDER BY createTime"
        param = []
        if workOrderId != None:
            strQuery = "SELECT id, orderId, creator, createTime, content, img FROM work_order_comment WHERE orderId=%s ORDER BY createTime"
            param = [workOrderId]

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return {}

        dData = {}
        for item in items:
            orderId = item[1]
            strImg = item[5]

            imgNameList = []
            if isinstance(strImg, str):
                nameList = strImg.split(",")
                for strNm in nameList:
                    if len(strNm):
                        imgNameList.append(strNm)

            if orderId not in dData.keys():
                dData.update({orderId: []})
            dData.get(orderId).append(
                dict(id=item[0], orderId=item[1], creator=item[2], createTime=item[3], content=item[4], imgNameList=imgNameList)
            )

        return dData

    def getWarehouseIdOfRFIDReceiver(self, receiverId):
        strQuery = "select id ,installLocation from equip_list where installLocation != ''"
        param = []
        if receiverId != None:
            strQuery = "select id ,installLocation from equip_list where installLocation != '' and id=%s"
            param = [receiverId]

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return {}

        dData = {}
        for item in items:
            nWarehouseid = None
            try:
                nWarehouseid = int(float(item[1]))
            except:
                pass

            if nWarehouseid == None:
                continue

            dData.update({item[0]: nWarehouseid})
        return dData

    def addImgToFddWorkOrder(self, orderId, imgList):
        try:
            strQuery = "SELECT img FROM fdd_work_order WHERE orderId=%s AND status=0"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (orderId,))
            if items == None:
                return None

            originalImgList = []
            originalImg = items[0][0]
            if isinstance(originalImg, str):
                originalImgList = originalImg.split(",")

            for img in imgList:
                if not isinstance(img, str):
                    continue

                if img not in originalImgList:
                    originalImgList.append(img)

            strImgNew = ",".join(originalImgList)
            strQuery = "UPDATE fdd_work_order SET img=%s WHERE orderId=%s AND status=0"
            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (strImgNew, orderId))
            return bSuc
        except Exception as e:
            logging.error("ERROR in addImgToFddWorkOrder: %s" % e.__str__())
            return False

    def removeImgFromWorkOrder(self, orderId, imgList):
        try:
            strQuery = "SELECT img FROM fdd_work_order WHERE orderId=%s AND status=0"
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (orderId,))
            if items == None:
                return None

            originalImgList = []
            originalImg = items[0][0]
            if isinstance(originalImg, str):
                originalImgList = originalImg.split(",")

            for img in imgList:
                if not isinstance(img, str):
                    continue

                if img in originalImgList:
                    originalImgList.remove(img)

            strImgNew = ",".join(originalImgList)
            strQuery = "UPDATE fdd_work_order SET img=%s WHERE orderId=%s AND status=0"
            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (strImgNew, orderId))
            return bSuc
        except Exception as e:
            logging.error("ERROR in removeImgFromWorkOrder: %s" % e.__str__())
            return False

    def importFddWorkOrderByExcel(self, dataList):
        strQuery = "INSERT INTO fdd_work_order (fddName, opUserName, opType, orderId, modifyTime, detail, status, createTime, title, ownUser) VALUES (%s, %s,%s,%s,%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(dataList))
        return bSuc, ""

    def updateMultiKeyValueOfUnit01(self, dTar):
        strMsg = ""
        try:
            strQuery = "SELECT unitproperty01, unitproperty02,unitproperty03, unitproperty04,unitproperty05,unitproperty06,unitproperty07,unitproperty08,unitproperty09," \
                       "unitproperty10,unitproperty11,unitproperty12,unitproperty13,unitproperty14,unitproperty15 FROM unit01 WHERE unitproperty01 in (" + ",".join(["'{0}'".format(obj) for obj in dTar.keys()]) + ")"
            items = self._mysqlDbContainer.op_db_query(app.config["DATABASE"], strQuery)
            if items == None:
                return False, "数据繁忙，请稍后再试"

            dExisting = {}
            for item in items:
                dExisting.update({item[0]: item[1:len(item)]})

            lInsert = []
            lUpdate = []
            nParamLength = 0
            for key, valueList in dTar.items():
                if dExisting.get(key):  # update
                    for idx, value in enumerate(valueList):
                        if value == None:
                            if dExisting.get(key) and len(dExisting.get(key)) - 1 >= idx:
                                valueList[idx] = dExisting.get(key)[idx]

                    valueList.append(key)
                    lUpdate.append(valueList)

                else:   # insert
                    valueList.insert(0, key)
                    lInsert.append(valueList)

                if len(valueList) - 1 > nParamLength:
                    nParamLength = len(valueList)

            insertParam = []
            updateParam = []
            for lis in lInsert:
                if len(lis) - 1 < nParamLength:
                    lis.extend([""] * (nParamLength - (len(lis) - 1)))
                insertParam.append(tuple(lis))

            for lis in lUpdate:
                if len(lis) - 1 < nParamLength:
                    lis_new = lis[0:-1] + [""] * (nParamLength - (len(lis) - 1)) + [lis[-1]]
                else:
                    lis_new = lis
                updateParam.append(tuple(lis_new))

            strInsertColumnName = ",".join(["unitproperty{d:02d}".format(d=idx) for idx in range(2, 2+nParamLength)])
            strValue = ",".join(["%s"] * (nParamLength+1))
            strInsert = "INSERT INTO unit01 (unitproperty01," + strInsertColumnName + ") VALUES (" + strValue + ")"

            strUpdateColumnName = ",".join(["unitproperty{d:02d}=%s".format(d=idx) for idx in range(2, 2+nParamLength)])
            strUpdate = "UPDATE unit01 SET " + strUpdateColumnName + " WHERE unitproperty01=%s"

            queryList = []
            paramList = []
            if len(insertParam):
                queryList.append(strInsert)
                paramList.append(insertParam)

            if len(updateParam):
                queryList.append(strUpdate)
                paramList.append(updateParam)

            bSuc = self._mysqlDbContainer.op_db_transaction_update_many(app.config['DATABASE'], queryList, paramList)

            return bSuc, strMsg
        except Exception as e:
            strMsg = e.__str__()
            return False, strMsg

    def getMultiKeyValueFromUnit01(self, keyOrList):
        keyList = []
        if isinstance(keyOrList, list):
            keyList = keyOrList
        else:
            keyList.append(keyOrList)

        strQuery = "SELECT * FROM unit01 WHERE unitproperty01 in (" + ",".join(["'{0}'".format(key) for key in keyList]) + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)

        res = {}
        for item in items:
            res.update({item[0]: item[1:]})

        #系统预置key list：
        sysConfigItemDict = {
            "readinterval":"5",
            "readtypeinterval":"1",
            "readcmdinterval":"100",
            "readbacnetmode":"0",
            "readtimeout":"2000",
            "bacnetreadlimit":"20",
            "writepriority":"16",
            "writepriority":"16",
            "backupday":"30",
            "precision":"2",
            "logic_log_level":"0",
            "logic_thread_mode":"0",
            "logic_batch_log_count": "100",
            "logic_log_level":"0",
            "debugopc":"0",
            "debugmodbus":"0",
            "outputerrcode":"0",
            "outputerrpoint":"0",
            "outputerrminute":"10",
            "outputmemoryinterval":"10",
            "realtime_vpoint_calculation_disabled":"0",
            "DTUEnabled":"0",
            "DTURecCmd":"0",
            "DTUChecked":"0",
            "DTUDisableSendAll":"0",
            "DTUComPort":"0",
            "DTUMinType":"2",
            "TcpSenderEnabled":"0",
            "TcpSenderPort":"9500",
            "TcpSenderName":"XXXXXXXXXXX",
            "TCPSendFilePackageSize":"4000",
            "TCPSendFilePackageInterval":"1000",
            "sendall":"1",
            "DTUType":"0",
            "TcpSenderIP":"",
            "modbusreadonebyone":"0",
            "ModbusReadIntervalMS":"100",
            "modbusmutilcount":"99",
            "modbusidinterval":"500",
            "modbustimeout":"5000",
            "modbuspollinterval":"2",
            "co3pcmdinterval":"50",
            "co3ptimeoutl":"5000",
            "co3prollinterval":"2",
            "modbus_equipment_engine_auto_restart_seconds_when_read_fail":"300",
            "modbus_client_core_upload_interval_seconds":"2",
            "opcclientmaxpoint":"10000",
            "opcclientthread":"0",
            "enableSecurity":"0",
            "disableCheckQuality":"0",
            "OPCServerIP":"0",
            "opcmutilcount":"50",
            "opccmdsleep":"100",
            "opcreconnect":"10",
            "opccheckreconnect":"5",
            "opccmdsleepfromdevice":"50",
            "opcpollsleep":"5",
            "opcpollsleepfromdevice":"60",
            "opcupdaterate":"500",
            "opcmultiadd":"1000",
            "opcasync20interval":"60",
            "opclanguageid":"0",
            "opcasync20mode":"0",
            "opcreconnectignore":"0",
            "opcmainpollsleep":"10",
            "remotesetopc":"0",
            "remotesetmodbus":"0",
            "remotesetbacnet":"0",
            "remotesetsimens":"0",
            "remotesetmysql":"0",
            "remoteSetSqlite":"0",
            "remoteSetSQLserver":"0",
            "servertime":"",
            "watchtime":"",
            "errrconnectcount":"30",
            "dave_proto_iso_tcp243":"0",
            "siemense_communication_type":"1",
            "disablewatch":"1",
            "autostartcore":"1",
            "checkdtu":"1",
            "checkservertime":"1",
            "DTURouterIP":"0",
            "checks7udp":"1"
            }

        res_ret = {}
        try:
            dTar = {}
            for kk in keyList:
                if sysConfigItemDict.get(kk) is not None and res.get(kk) is None:
                    if kk=='modbus_equipment_engine_auto_restart_seconds_when_read_fail':
                        aa=8
                    newValueList = [sysConfigItemDict.get(kk), "", "", "", "", "", "", "", "", "", "", "", ""]
                    dTar.update({kk: newValueList})
                    res.update({kk: newValueList})
                elif res.get(kk) is None:
                    newValueList = ["", "", "", "", "", "", "", "", "", "", "", "", ""]
                    res.update({kk: newValueList})
            res_ret = copy.deepcopy(res)
            if dTar:
                bSuc, strMsg = self.updateMultiKeyValueOfUnit01(dTar)
        except:
            pass

        return res_ret


    def startStopLogic(self, logicName, onoff):
        strQuery1 = "SELECT count(unitproperty01) FROM unit01 WHERE unitproperty01=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery1, (logicName,))
        if items == None:
            return False
        if not len(items):
            return False
        if not len(items[0]):
            return False

        queryList = []
        paramList = []
        if items[0][0] == 0:
            strInsert = "INSERT INTO unit01 (unitproperty01, unitproperty02) VALUES (%s, %s)"
            queryList.append(strInsert)
            paramList.append(
                (logicName, onoff)
            )
        elif items[0][0] > 1:
            queryList.append("DELETE FROM unit01 WHERE unitproperty01=%s")
            paramList.append((logicName,))

            queryList.append("INSERT INTO unit01 (unitproperty01, unitproperty02) VALUES (%s, %s)")
            paramList.append((logicName, onoff))
        else:
            queryList.append("UPDATE unit01 SET unitproperty02=%s WHERE unitproperty01=%s")
            paramList.append((onoff, logicName))

        bSuc = self._mysqlDbContainer.op_db_transaction_update_many(app.config['DATABASE'], queryList, paramList)
        # bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery2, params)
        return bSuc

    def deleteLogic(self, logicName):
        strQuery = 'DELETE FROM unit01 WHERE unitproperty01=%s'
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (logicName,))
        return bSuc

    # 2：修改策略周期；4：新建策略；5：删除策略
    def operateUnit02ForLogic(self, opType, logicName, period=None):
        strQuery = None
        params = None

        if opType == 2:
            strQuery = "INSERT INTO unit02 (unitproperty01, unitproperty02, unitproperty03, unitproperty04, unitproperty05, unitproperty06, unitproperty07, unitproperty08) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)"
            params = ("{logicName}.dll".format(logicName=logicName), "{logicName}.dll".format(logicName=logicName), 2, period, "", "", "", "")

        # 暂时屏蔽，否则会影响其他正常的unit02表操作
        # elif opType == 4:
        #     strQuery = "INSERT INTO unit02 (unitproperty01, unitproperty02, unitproperty03, unitproperty04) VALUES (%s,%s,%s,%s)"
        #     params = (logicName, "{logicName}.dll".format(logicName=logicName), 4, "{logicName}.dll".format(logicName=logicName))
        # elif opType == 5:
        #     strQuery = "INSERT INTO unit02 (unitproperty01, unitproperty02, unitproperty03) VALUES (%s,%s,%s)"
        #     params = (logicName, "{logicName}.dll".format(logicName=logicName), 5)

        if strQuery and params:
            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, params)
            return bSuc
        return False

    # 更新dll文件
    def operateUnit02UpdateDllFile(self, dllNameWithExt):
        strQuery = "INSERT INTO unit02 (unitproperty01, unitproperty02, unitproperty03, unitproperty04, unitproperty05, unitproperty06, unitproperty07, unitproperty08) VALUES (%s,%s,%s,%s, %s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, ("", "", 7, dllNameWithExt, "", "", "", ""))
        return bSuc

    # 修改策略配参(输入参数）
    def operateUnit02EditConfig(self, logicName, editList):
        strQuery = "INSERT INTO unit02 (unitproperty01, unitproperty02, unitproperty03, unitproperty04, unitproperty05, unitproperty06, unitproperty07, unitproperty08) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)"
        params = []
        for item in editList:
            paramName = item.get("paramName", None)
            paramValue = item.get("curValue", None)
            paramType = item.get("paramType", None)
            if paramName == None or paramValue == None or paramType not in ["const", "point"]:
                continue

            params.append(
                ("{logicName}.dll".format(logicName=logicName), "{logicName}.dll".format(logicName=logicName), 0, paramName, paramValue, paramType, "", "")
            )

        bSuc = self._mysqlDbContainer.op_db_update_many(app.config["DATABASE"], strQuery, tuple(params))
        return bSuc

    """
    1： add
    0: delete
    """
    def operateUnit01ForVpointAddDel(self, vpointNameList, nAddOrDel):
        if not isinstance(vpointNameList, list):
            vpointNameList = [vpointNameList]

        strQuery = None
        params = None
        if nAddOrDel == 1:
            strQuery = "INSERT INTO unit01 (unitproperty01, unitproperty02) VALUES (%s, %s)"
            params = ("vpoint_add_del", ",".join(vpointNameList))
        elif nAddOrDel == 0:
            strQuery = "INSERT INTO unit01 (unitproperty01, unitproperty03) VALUES (%s, %s)"
            params = ("vpoint_add_del", ",".join(vpointNameList))

        if strQuery != None:
            bSuc = self._mysqlDbContainer.op_db_update(app.config["DATABASE"], strQuery, params)
            return bSuc
        return None

    def editLogicNameInUnit01(self, oldName, newName):
        strQuery1 = "SELECT unitproperty01, unitproperty02 FROM unit01 WHERE unitproperty01=%s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery1, (oldName,))
        if items == None:
            return True
        if not len(items):
            return True
        if not len(items[0]):
            return True

        strQuery = "UPDATE unit01 SET unitproperty01=%s WHERE unitproperty01=%s"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, (newName, oldName))
        return bSuc

    def getEquipAssetInfo(self, equipIdList):
        strEquipId = ",".join([str(nid) for nid in equipIdList])
        strQuery = "select ea.id, ea.equip_id, ea.en_name, ea.param_value, el.model_id, eat.cn_name from equip_asset ea left join equip_list el on ea.equip_id = el.id left join equip_asset_tmpl eat on (el.model_id = eat.tmpl_def_id and ea.en_name = eat.en_name) WHERE el.id in ( " + strEquipId + ")"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)

        if items == None:
            return {}

        dRes = {}
        for item in items:
            nEquipId = item[1]
            if nEquipId not in dRes.keys():
                dRes.update({nEquipId: dict(identity="", params={}, nameCh="")})

            if item[2] == "name":
                if not dRes[nEquipId]["nameCh"]:
                    dRes[nEquipId]["nameCh"] = item[3]

            elif item[2] == "identity":
                if not dRes[nEquipId]["identity"]:
                    dRes[nEquipId]["identity"] = item[3]

            else:
                paramName = item[2]
                paramValue = item[3]
                paramDesc = item[5]

                dRes[nEquipId]["params"].update({paramName: dict(paramValue=paramValue,
                                                                 paramDesc=paramDesc)})

        return dRes

    def getMaintainLockRecord(self):
        strQuery = "SELECT unitproperty02 FROM unit01 where unitproperty01 = 'maintain_lock_record'"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        if not isinstance(items, list):
            return None

        if not len(items):
            return {}

        if not len(items[0]):
            return {}

        try:
            dJson = json.loads(items[0][0])
            return dJson
        except:
            return {}

    def saveMaintainLockRecord(self, dInfo):
        strQuery = "SELECT unitproperty02 FROM unit01 where unitproperty01 = 'maintain_lock_record'"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return

        bExists = False
        if isinstance(items, list):
            if len(items):
                if len(items[0]):
                    bExists = True

        if bExists:
            strQuery = "UPDATE unit01 SET unitproperty02 = %s WHERE unitproperty01='maintain_lock_record'"
            param = (json.dumps(dInfo),)
        else:
            strQuery = "INSERT INTO unit01 (unitproperty01, unitproperty02) VALUES (%s, %s)"
            param = ('maintain_lock_record', json.dumps(dInfo))

        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strQuery, param)
        return bSuc

    #1 确认  2忽略  3 消除
    def dealWarning(self, conditionList, nOpType, strRemark, nUserId, nIgnoreMinutes):
        strQuery = None
        param = []
        tOpTime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if nOpType == 1:
            strQuery = "UPDATE warningrecord SET confirmed=1, confirmeduser=%s, confirmRemark=%s, confirmOpTime=%s WHERE info=%s"
            for condition in conditionList:
                param.append(
                    (nUserId, strRemark, tOpTime, condition)
                )
        elif nOpType == 2:
            tIgnoreTime = (datetime.now() + timedelta(minutes=nIgnoreMinutes)).strftime("%Y-%m-%d %H:%M:%S")
            strQuery = "UPDATE warningrecord SET ignoreUser=%s, ignoreRemark=%s, ignoreOpTime=%s, ignoreToTime=%s WHERE info=%s"
            for condition in conditionList:
                param.append(
                    (nUserId, strRemark, tOpTime, tIgnoreTime, condition)
                )
        elif nOpType == 3:
            strQuery = "UPDATE warningrecord SET closeStatus=1, closeUser=%s, closeRemark=%s, closeOpTime=%s WHERE info=%s"
            for condition in conditionList:
                param.append(
                    (nUserId, strRemark, tOpTime, condition)
                )

        if strQuery == None:
            return False

        try:
            bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strQuery, tuple(param))
            return bSuc
        except:
            return False

    def logicPushRecord(self, paramList):
        strInsert = "INSERT INTO logic_push_record_2023 (id, groupName, logicSourceName, content, reason, occurTime, important, urgent, result,actTime, positionName,  remark) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update_many(app.config['DATABASE'], strInsert, tuple(paramList))
        return bSuc

    def getLogicPushList(self, strBegin=None, strEnd=None, roomName=None):
        tBegin = datetime.strptime(strBegin, "%Y-%m-%d %H:%M:%S")
        nYear = tBegin.year

        strQuery = "SELECT id, groupName, logicSourceName, content, reason, occurTime, important, urgent, result, actTime, positionName, remark FROM logic_push_record_{year}".format(year=nYear)

        whereList = []
        paramList = []
        if strBegin != None and strEnd != None:
            whereList.append("actTime >= %s AND actTime <= %s")
            paramList.append(strBegin)
            paramList.append(strEnd)
        if roomName != None:
            whereList.append("positionName=%s")
            paramList.append(roomName)

        if len(whereList):
            strQuery += " WHERE "
            strQuery += " AND ".join(whereList)

        strQuery += " ORDER BY actTime"
        itemList = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(paramList))
        if itemList == None:
            return None

        dataList = []
        for item in itemList:
            dataList.append(dict(id=item[0],
                                 groupName=item[1],
                                 logicSourceName=item[2],
                                 content=item[3],
                                 reason=item[4],
                                 occurTime=item[5],
                                 important=item[6],
                                 urgent=item[7],
                                 result=item[8],
                                 actTime=item[9],
                                 positionName=item[10],
                                 remark=item[11]))
        return dataList

    def updateLogicPush(self, strId, nResult, nYear):
        if nResult in [1, -1]:
            strUpdate = "UPDATE logic_push_record_{year} SET result=%s, actTime=%s WHERE id =%s".format(year=nYear)
            paramList = [nResult, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), strId]
        else:
            strUpdate = "UPDATE logic_push_record_{year} SET result=%s WHERE id =%s".format(year=nYear)
            paramList = [nResult, strId]
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], strUpdate, tuple(paramList))
        return bSuc

    # 先查到所有工单记录，然后对于每个工单id挑选出modifyTime距今最近的一条记录（这样每个工单id对应一条记录）然后对这些记录进行逐个遍历，去除opUserName 为 exceptUserName 的记录
    def getFaultUpdate(self, beginTime, endTime, exceptUserName=None):
        try:
            strQuery = "SELECT fddName, opUserName, opType, opContentData, orderId, modifyTime, detail, status, title, ownUser, img, createTime FROM fdd_work_order WHERE modifyTime >= %s AND modifyTime <= %s"
            param = [beginTime, endTime]
            items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
            if items is None:
                return None, "查询失败，请稍后再试"

            dataList = []
            dAll = {}
            for item in items:
                orderId = item[4]
                createTime = ""
                if isinstance(item[11], str):
                    createTime = item[11]
                elif isinstance(item[11], datetime):
                    createTime = item[11].strftime("%Y-%m-%d %H:%M:%S")

                modifyTime = ""
                tModifyTime = None
                if isinstance(item[5], str):
                    modifyTime = item[5]
                    tModifyTime = datetime.strptime(item[5], "%Y-%m-%d %H:%M:%S")
                elif isinstance(item[5], datetime):
                    modifyTime = item[5].strftime("%Y-%m-%d %H:%M:%S")
                    tModifyTime = item[5]

                if not isinstance(tModifyTime, datetime):
                    continue

                if orderId not in dAll.keys():
                    dAll.update({orderId: dict(fddName=item[0], opUserName=item[1], opType=item[2], opContentData=item[3],
                                              modifyTime=modifyTime, detail=item[6], status=item[7], name=item[8],
                                              owner=item[9], img=item[10], orderId=item[4], createTime=createTime, tModifyTime=tModifyTime)})
                else:
                    if tModifyTime > dAll[orderId].get("tModifyTime"):
                        dAll.update({orderId: dict(fddName=item[0], opUserName=item[1], opType=item[2], opContentData=item[3],
                                              modifyTime=modifyTime, detail=item[6], status=item[7], name=item[8],
                                              owner=item[9], img=item[10], orderId=item[4], createTime=createTime, tModifyTime=tModifyTime)})

            dRes = {}
            if exceptUserName:
                for orderId, dTemp in dAll.items():
                    if dTemp.get("opUserName") == exceptUserName:
                        continue

                    dRes.update({orderId: dTemp})

            return dRes, ""
        except Exception as e:
            logging.error("ERROR in getFaultUpdate: %s" % e.__str__())
            return None, "查询失败，原因: %s" % e.__str__()

    # 查询在起止时间内的且orderid在nOrderIdList 内或 creator在relatedUserNameList的记录
    def getFaultCommentUpdate(self, beginTime, endTime, exceptUserName=None):
        strQuery = "SELECT id, createTime, modifyTime, content, orderId, creator FROM work_order_comment WHERE ((createTime >= %s AND createTime <= %s) OR (modifyTime >= %s AND modifyTime <= %s))"
        param = [beginTime, endTime, beginTime, endTime]
        if exceptUserName != None:
            strQuery += " AND creator != %s"
            param.append(exceptUserName)

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return None, "查询失败，请稍后再试"

        dataList = []
        for item in items:
            strCreateTime = ""
            if isinstance(item[1], str):
                strCreateTime = item[1]
            elif isinstance(item[1], datetime):
                strCreateTime = item[1].strftime("%Y-%m-%d %H:%M:%S")

            strModifyTime = ""
            if isinstance(item[2], str):
                strModifyTime = item[2]
            elif isinstance(item[2], datetime):
                strModifyTime = item[2].strftime("%Y-%m-%d %H:%M:%S")

            dataList.append(dict(commentId=item[0],
                                 createTime=strCreateTime,
                                 modifyTime=strModifyTime,
                                 content=item[3],
                                 orderId=item[4],
                                 creator=item[5]))

        return dataList, ""



    # 获取与用户相关的工单  1-userName自己创建的工单  2-被指派给userName的工单
    def getRelatedWorkOrder(self, userName, nAssignToOrdrerIdList=None):
        strQuery = "SELECT orderId, createTime, title FROM fdd_work_order WHERE opType = -1 AND opUserName = %s"
        param = [userName]

        if isinstance(nAssignToOrdrerIdList, list) and len(nAssignToOrdrerIdList):
            strQuery += " OR (opType = -1 AND orderId in (" + ",".join([str(item) for item in nAssignToOrdrerIdList]) + "))"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, tuple(param))
        if items == None:
            return None, "查询失败，请稍后再试"

        dOrder = {}
        for item in items:
            if item[0] not in dOrder.keys():
                dOrder.update({item[0]: {}})

            createTime = ""
            if isinstance(item[1], str):
                createTime = item[1]
            elif isinstance(item[1], datetime):
                createTime = item[1].strftime("%Y-%m-%d %H:%M:%S")

            dOrder[item[0]].update(dict(orderId=item[0], createTime=createTime, name=item[2]))

        return dOrder, ""

    def addFddWorkOrderUserVisitRecord(self, userId, orderId):
        tNow = datetime.now()
        tFrom = tNow - timedelta(days=7)
        sqlQuery = "SELECT userId, orderId, readTime FROM fdd_work_order_read_info WHERE userId=%s AND orderId=%s AND readTime >= %s AND readTime <= %s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], sqlQuery, (userId, orderId, tFrom.strftime("%Y-%m-%d %H:%M:%S"), tNow.strftime("%Y-%m-%d %H:%M:%S")))
        if items == None:
            return False, "查阅记录查询失败，请稍后再试"

        readTime = None
        if len(items):
            for item in items:
                try:
                    if isinstance(item[2], datetime):
                        rTime = item[2]
                    elif isinstance(item[2], str):
                        rTime = datetime.strptime(item[2], "%Y-%m-%d %H:%M:%S")
                    else:
                        rTime = item[2]

                    if readTime == None:
                        readTime = rTime
                    elif isinstance(readTime, datetime) and isinstance(rTime, datetime):
                        if rTime > readTime:
                            readTime = rTime
                except:
                    pass

        if isinstance(readTime, datetime):
            strReadTime = readTime.strftime("%Y-%m-%d %H:%M:%S")
            sqlUpdate = "UPDATE fdd_work_order_read_info SET readTime =%s WHERE userId=%s AND orderId=%s AND readTime=%s"
            bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sqlUpdate, (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), userId, orderId, strReadTime))
            return bSuc, ""

        sqlInsert = "INSERT INTO fdd_work_order_read_info (userId, orderId, readTime) VALUES (%s, %s,%s)"
        bSuc = self._mysqlDbContainer.op_db_update(app.config['DATABASE'], sqlInsert, (userId, orderId, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        return bSuc, ""


    def getWorkOrderUserVisitRecord(self, userId, strBeginTime, strEndTime):
        strQuery = "SELECT userId, orderId, readTime from fdd_work_order_read_info WHERE userId =%s AND readTime >= %s AND readTime <= %s"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery, (userId, strBeginTime, strEndTime))
        if items == None:
            return None, "查询失败，请稍后再试"

        dRes = {}
        for item in items:
            nId = item[1]
            tReadTime = None
            if isinstance(item[2], str):
                tReadTime = datetime.strptime(item[2], "%Y-%m-%d %H:%M:%S")
            elif isinstance(item[2], datetime):
                tReadTime = item[2]

            if nId not in dRes.keys():
                dRes.update({nId: None})

            if dRes.get(nId) == None and isinstance(tReadTime, datetime):
                dRes[nId] = tReadTime
            else:
                if isinstance(dRes.get(nId), datetime) and isinstance(tReadTime, datetime):
                    if tReadTime > dRes.get(nId):
                        dRes[nId] = tReadTime

        return dRes, ""

    def getWorkOrderCreateTimeAndTitle(self):
        strQuery = "SELECT orderId, createTime, title FROM fdd_work_order WHERE opType = -1"

        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None, "查询失败，请稍后再试"

        dOrder = {}
        for item in items:
            if item[0] not in dOrder.keys():
                dOrder.update({item[0]: {}})

            createTime = ""
            if isinstance(item[1], str):
                createTime = item[1]
            elif isinstance(item[1], datetime):
                createTime = item[1].strftime("%Y-%m-%d %H:%M:%S")

            dOrder[item[0]].update(dict(orderId=item[0], createTime=createTime, name=item[2]))

        return dOrder, ""

    def listWarningTables(self):
        strQuery = "SHOW TABLES LIKE '%warningrecord%'"
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
        if items == None:
            return None

        tableNameList = []
        for item in items:
            tableNameList.append(item[0])
        return tableNameList

    def getColumnNameOfTable(self, tableName):
        strQ = "DESC %s" % tableName
        items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQ)
        if items == None:
            return None

        columnNameList = []
        for item in items:
            if len(item):
                columnNameList.append(item[0])
        return columnNameList

    def exportHistoryWarningRecordInExcel(self, targetDir, tableNameList):
        nCount = 0
        for tableName in tableNameList:
            try:
                columnNameList = self.getColumnNameOfTable(tableName)
                if columnNameList == None:
                    return None, 0

                book = Workbook()
                sheet = book.create_sheet(tableName, 0)
                for i, columnName in enumerate(columnNameList):
                    sheet.cell(row=1, column=i+1, value=columnName)

                strQuery = "SELECT * FROM %s" % tableName
                items = self._mysqlDbContainer.op_db_query(app.config['DATABASE'], strQuery)
                if items == None:
                    return None, 0

                for i, values in enumerate(items):
                    for j, value in enumerate(values):
                        strValue = ""
                        if isinstance(value, str):
                            strValue = value
                        elif isinstance(value, datetime):
                            strValue = value.strftime("%Y-%m-%d %H:%M:%S")
                        elif isinstance(value, int) or isinstance(value, float):
                            strValue = str(value)
                        elif value == None:
                            strValue = "None"

                        sheet.cell(row=i+2, column=j+1, value=strValue)

                filePath = os.path.join(targetDir, "{name}.xlsx".format(name=tableName))
                if os.path.exists(filePath):
                    try:
                        os.remove(filePath)
                    except:
                        pass

                book.save(filePath)
                nCount += 1
            except:
                pass
            finally:
                time.sleep(0.2)

        return True, nCount


