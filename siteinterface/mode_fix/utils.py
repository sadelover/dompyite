# -*- coding: utf-8 -*-
from siteinterface.mode_fix import bp_fix
from flask import request, jsonify
from siteinterface.BEOPSqliteAccess import *
from siteinterface.BEOPDataAccess import BEOPDataAccess
import logging
from PyQt4.QtGui import QTextDocument, QPrinter, QApplication
import sys
import os
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import Workbook
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
import win32com.client


def fixIdExists(fixId):
    """
    检查fixId是否已存在
    """
    res = False
    try:
        rvConduct = BEOPDataAccess.getInstance().fixIdExistsInFix(fixId)
        if rvConduct and len(rvConduct):
            if len(rvConduct[0]):
                if rvConduct[0][0] > 0:
                    res = True

    except Exception as e:
        strLog = "ERROR in fixIdExists: %s" % e.__str__()
        logging.error(strLog)
    finally:
        return res

def isValidDate(strDate, strFormat):
    res = False
    if not isinstance(strDate, str):
        return res
    try:
        tTime = datetime.strptime(strDate, strFormat)
        res = True
    except:
        pass
    finally:
        return res

def variableToString(variable):
    res = ""
    try:
        for name in globals():
             if eval(name) == variable:
                res = name
    except:
        pass
    return res

def htmlToStr(strHtml):
    htmlList = BeautifulSoup(strHtml, "html.parser").strings
    stringList = []
    for html in htmlList:
        stringList.append(html)

    if not len(stringList):
        return ""
    return "。".join(stringList)

def saveImageToLocal(fileList):
    try:
        staticFolder = app.static_folder
        imagesDir = os.path.join(staticFolder, "images")
        if not os.path.exists(imagesDir):
            os.mkdir(imagesDir)

        fixDir = os.path.join(imagesDir, "fix")
        if not os.path.exists(fixDir):
            os.mkdir(fixDir)

        fileNameList = []
        for file in fileList:
            strFileName = "{0}_{1}".format(datetime.now().strftime("%Y%m%d%H%M%S"), file.filename)
            strFilePath = os.path.join(fixDir, strFileName)
            if os.path.exists(strFilePath):
                os.remove(strFilePath)
            file.save(strFilePath)

            nFileSize = os.path.getsize(strFilePath)
            nFileSizeM = nFileSize / float(1024 * 1024)
            if nFileSizeM > 10:
                os.remove(strFilePath)
                return dict(err=1, msg="图片大小必须小于10M", data="")

            fileNameList.append(strFileName)

        return dict(err=0, msg="上传成功", data=fileNameList)

    except Exception as e:
        strLog = "ERROR in saveImageToLocal: %s" % e.__str__()
        logging.error(strLog)
        return dict(err=1, msg="上传失败", data="")

def get_sub_tag_list(object):
    tagList = []
    try:
        nameList = []
        for content in object.contents:
            try:
                name = content.name
            except:
                name = None
            nameList.append(name)
        nameList = list(set(nameList))
        tagList = [obj for obj in nameList if obj is not None]
    except Exception as e:
        strLog = "ERROR in /fix/downloadFixContentInDocx - get_sub_tag_list: %s" % e.__str__()
        logging.error(strLog)
    finally:
        return tagList

def get_text(object):
    strText = None
    try:
        strText = object.text
    except:
        pass
    return strText

def process_fix_content(strContent):
    try:
        soup = BeautifulSoup(strContent, "html.parser")
        parList = []
        for item in soup.contents:
            subTagList = get_sub_tag_list(item)
            text = None if not get_text(item) else get_text(item)
            name = item.name
            imgSrc = None
            if "img" in subTagList:
                img = item.find_all("img")[0]
                imgSrc = img.get("src")
            parList.append({
                "text": text.strip() if text is not None and isinstance(text, str) else None,
                "name": name,
                "img": imgSrc
            })

        return parList
    except Exception as e:
        strLog = "ERROR in /fix/downloadFixContentInDocx - process_fix_content: %s" % e.__str__()
        logging.error(strLog)
        return []

def genStyle(styles, strStyleName, strFontName, nFontSize, bBold):
    try:
        titleStyle = styles.add_style(strStyleName, WD_STYLE_TYPE.CHARACTER)
        titleFont = titleStyle.font
        titleFont.name = strFontName
        titleFont.size = Pt(nFontSize)
        titleStyle._element.rPr.rFonts.set(qn('w:eastAsia'), strFontName)
        titleFont.bold = bBold
    except Exception as e:
        strLog = "ERROR in /fix/downloadFixContentInDocx - genStyle: %s" % e.__str__()
        logging.error(strLog)

def genParagraph(doc, strStyle, strContent, nBefore, nAfter):
    try:
        paragraph = doc.add_paragraph()
        paragraphFormat = paragraph.paragraph_format
        paragraphFormat.space_before = Pt(nBefore)
        paragraphFormat.space_after = Pt(nAfter)
        paragraph.add_run(strContent, style=strStyle)
    except Exception as e:
        strLog = "ERROR in /fix/downloadFixContentInDocx - genParagraph: %s" % e.__str__()
        logging.error(strLog)

def genTable(doc, nRows, nCols, formList, nFontSize):
    try:
        table = doc.add_table(rows=nRows, cols=nCols)
        table.style = 'Table Grid'
        table.style.font.size = Pt(nFontSize)
        table.style.font.name = "黑体"
        table.style.paragraph_format.space_before = Pt(1)
        table.style.paragraph_format.space_after = Pt(1)
        tableCols = table.columns
        for form in formList:
            for cell in tableCols[form[0]].cells:
                cell.width = Inches(form[1])
        return table
    except Exception as e:
        strLog = "ERROR in /fix/downloadFixContentInDocx - genTable: %s" % e.__str__()
        logging.error(strLog)

def fillTable(table, style, rowOrColNum, contentList):
    try:
        rowsOrCols = table.rows
        if style == "col":
            rowsOrCols = table.columns

        for idx, cell in enumerate(rowsOrCols[rowOrColNum].cells):
            cell.text = contentList[idx]
    except:
        pass

def get_img_path(imgLink):
    imgPath = None
    try:
        imgName = os.path.basename(imgLink)
        imgPath = os.path.join(app.static_folder, "images", "fix", imgName)
    except:
        pass
    finally:
        return imgPath

def addImg(doc, imgPath, nWidth):
    try:
        doc.add_picture(imgPath, width=Inches(nWidth))
    except Exception as e:
        strLog = "ERROR in /fix/downloadFixContentInDocx - addImg: %s" % e.__str__()
        logging.error(strLog)

def genReportDocx(defaultDocx, destPath, reportTitle, reportTime, reportUser, urgent, importance, energyEffects, solveUser, result, closeTime, parList):
        bSuc = False
        try:
            doc = Document(defaultDocx)
            styles = doc.styles

            # 生成报告标题样式
            genStyle(styles, "titleStyle", u"黑体", 24, False)
            # 生成一级标题样式
            genStyle(styles, "priTitleStyle", u"黑体", 11, True)
            # 生成二级标题样式
            genStyle(styles, "secTitleStyle", u"楷体", 10.5, False)
            # 生成正文样式
            genStyle(styles, "bodyStyle", u"宋体", 10.5, False)

            # 添加报告标题
            genParagraph(doc, "titleStyle", reportTitle, 15, 10)
            # 添加一级标题：基本信息
            genParagraph(doc, "priTitleStyle", "基本信息", 2, 2)

            # 创建表格
            table = genTable(doc, 8, 2, [(0, 1.4), (1, 4.5)], 9)
            if not table:
                return jsonify(dict(err=1, msg="下载失败", data=""))
            # 填入表头列
            tableHeaderList = ["创建时间", "报告人员", "紧急程度", "重要程度", "节能影响程度", "处理人员", "解决结果", "关闭时间"]
            fillTable(table, "col", 0, tableHeaderList)
            # 填入第二列
            basicInfoList = [reportTime, reportUser, urgent, importance, energyEffects, solveUser, result, closeTime]
            fillTable(table, "col", 1, basicInfoList)

            # 添加一级标题：详细内容
            genParagraph(doc, "priTitleStyle", "详细内容", 15, 2)

            # 添加正文
            secTitleCount = 0
            for idx, par in enumerate(parList):
                if par.get("img") != None:
                    imgPath = get_img_path(par.get("img"))
                    addImg(doc, imgPath, 4.5)
                elif par.get("name") in ["h1", "h2", "h3", "h4", "h5"]:
                    genParagraph(doc, "secTitleStyle", "({0}){1}".format(secTitleCount+1, par.get("text")), 2, 2)
                    secTitleCount += 1
                else:
                    genParagraph(doc, "bodyStyle", par.get("text"), 2, 2)

            doc.save(destPath)
            bSuc = True

        except Exception as e:
            strLog = "ERROR in processDocReport: %s" % e.__str__()
            logging.error(strLog)
        finally:
            return bSuc

def genReportPdf(srcPath, destPath):
    bSuc = False
    try:
        wdFormatPDF = 17
        word = win32com.client.Dispatch('Word.Application')
        document = word.Documents.Open(srcPath)
        document.SaveAs(destPath, FileFormat=wdFormatPDF)
        document.Close()
        word.Quit()
        bSuc = True

    except Exception as e:
        strLog = "ERROR in genPdf: %s" % e.__str__()
        logging.error(strLog)
    finally:
        return bSuc