
import csv, tablib
import os.path
from docx.oxml.shared import OxmlElement, qn
from docx import Document
from docx.shared import Inches
from siteinterface.BEOPDataAccess import BEOPDataAccess
import json
import xlwt,xlrd
from openpyxl import load_workbook
from openpyxl.chart import BarChart, LineChart, Series, Reference,PieChart,ScatterChart,marker
from openpyxl.chart import label
from datetime import datetime,timedelta
from openpyxl import load_workbook
from openpyxl import Workbook
from docx.oxml.ns import qn
import re
import numpy
import logging
from siteinterface.commonUtils import *
from siteinterface.utils import *
import PIL
from siteinterface.SystemAPI import g_systemAPIList
from siteinterface.BEOPSqliteAccess import BEOPSqliteAccess
import random
from siteinterface.RedisManager import RedisManager
import calendar


def is_digit(tar):
    try:
        n = float(tar)
        return True
    except:
        return False

def insert_acttime_for_api(strExpression, apiName, strTime):
    try:
        strExpression = strExpression.replace("'", '"')
        pattern = r'{apiName}[(].*?[)]'.format(apiName=apiName)
        targetList = re.findall(pattern, strExpression)
        if not len(targetList):
            return strExpression

        targetList = [target for target in targetList if target.find("actTime") == -1]

        targetProList = [target.replace(")", ',actTime="{time}")'.format(time=strTime)) for target in targetList]

        for idx, target in enumerate(targetList):
            strExpression = strExpression.replace(target, targetProList[idx])

        return strExpression
    except Exception as e:
        return strExpression

def alpha2num(str):
    index = 1
    length = len(str)
    sum = 0
    for i in range(length):
        sum += ((ord(str[length - i - 1]) - ord('A') + 1) * index)
        index *= 26
    return sum

def alphaNumSplit(str):
    length = len(str)

    for i in range(length):
        if str[i].isdigit():
            alphastr = str[0:i]
            numstr = str[i:]
            break
    return alpha2num(alphastr), int(numstr)


def plotInExcel(wb):
    collist = ['B', 'L']
    plt_row_num = 3
    rowNo = 0
    try:
        try:
            ws = wb.get_sheet_by_name('plot')
        except Exception:
            try:
                ws = wb.get_sheet_by_name('plotcmd')
            except Exception:
                return dict(error=0, msg='没有找到名为plotcmd的工作表。')

        rows = ws.rows
        wb.create_sheet(title='结果统计图')
        ws = wb.get_sheet_by_name('结果统计图')

        for row in rows:
            for cell in row:
                try:
                    dataIndexList = []  # [{data_min_col, data_min_row, data_max_col, data_max_row}...]
                    strEval = cell.value
                    if strEval is None or strEval=='':
                        continue

                    strEval = ''.join(strEval.split())
                    strEval = strEval.replace('，', ',')
                    strEval = strEval.replace('：', ':')
                    strEval = strEval.replace("\'", "")
                    dic = eval(strEval)
                    type = dic['type']
                    charname = dic['name']
                    ymin = dic['ymin']
                    ymax = dic['ymax']

                    x = dic['x']
                    ylist = dic['y']
                    sheetName = x.split('!')[0]
                    xarea = x.split('!')[1]
                    cats_min_col, cats_min_row = alphaNumSplit(xarea.split(':')[0])
                    cats_max_col, cats_max_row = alphaNumSplit(xarea.split(':')[1])
                    for y in ylist:
                        yarea = y.split('!')[1]
                        data_min_col, data_min_row = alphaNumSplit(yarea.split(':')[0])
                        data_max_col, data_max_row = alphaNumSplit(yarea.split(':')[1])
                        dataIndexList.append(
                            dict(data_min_col=data_min_col, data_min_row=data_min_row, data_max_col=data_max_col,
                                 data_max_row=data_max_row))
                except Exception as ee:
                    raise Exception('plotcmd格式错误:' + ee.__str__())

                ws_data = wb.get_sheet_by_name(sheetName)

                # chart1 = BarChart()
                # chart1.type = "col"
                # chart1.style = 10
                # chart1.title = "Bar Chart"
                # chart1.y_axis.title = 'Test number'
                # chart1.x_axis.title = 'Sample length (mm)'
                # data = Reference(ws, min_col=2, min_row=1, max_row=7, max_col=3)
                # cats = Reference(ws, min_col=1, min_row=2, max_row=7)
                # chart1.add_data(data, titles_from_data=True)
                # chart1.set_categories(cats)
                # chart1.shape = 4
                if type in ['line', 'bar']:
                    if type == 'line':
                        c1 = LineChart()  # 新建一张图
                    if type == 'bar':
                        c1 = BarChart()
                    c1.title = charname  # 图的标题
                    c1.style = 5  # 线条的style
                    # c1.y_axis.title = 'price'  # y坐标的标题
                    # c1.x_axis = DateAxis(crossAx=100)
                    # c1.x_axis.number_format = 'mm-e'  # 规定日期格式  这是月,年格式
                    # c1.x_axis.majorTimeUnit = "Months"  # 规定日期间隔 注意days；Months大写
                    c1.x_axis.title = "时间"  # x坐标的标题
                    if ymin:
                        c1.y_axis.scaling.min = ymin  # y坐标的区间
                    if ymax:
                        c1.y_axis.scaling.max = ymax

                    lineColorList = ["0066FF", "FFCC00", "969696", "33CC00", "6699FF", "AAB903", "3489AA", "770389", "20BBEE", "FE9008"]
                    nItemIndex = 0
                    for item in dataIndexList:
                        data = Reference(ws_data, min_col=item['data_min_col'], min_row=item['data_min_row'],
                                         max_col=item['data_max_col'], max_row=item['data_max_row'])
                        c1.add_data(data, titles_from_data=True)
                        if nItemIndex<len(lineColorList):
                            c1.series[nItemIndex].graphicalProperties.line.solidFill = lineColorList[nItemIndex]
                        nItemIndex+=1

                    # s2 = c1.series[0]
                    # s1 = c1.series[1]
                    dates = Reference(ws_data, min_col=cats_min_col, min_row=cats_min_row, max_col=cats_max_col,
                                      max_row=cats_max_row)
                    c1.set_categories(dates)

                elif type == 'pie':
                    c1 = PieChart()
                    for item in dataIndexList:
                        data = Reference(ws_data, min_col=item['data_min_col'], min_row=item['data_min_row']-1,max_col=item['data_max_col'], max_row=item['data_max_row'])
                        c1.add_data(data, titles_from_data=True)
                    labels = Reference(ws_data, min_col=cats_min_col, min_row=cats_min_row, max_col=cats_max_col,max_row=cats_max_row)
                    c1.dataLabels = label.DataLabelList()
                    c1.dataLabels.showPercent = True
                    c1.set_categories(labels)
                    c1.title = charname
                    c1.splitType = "percent"

                elif type == 'scatter':
                    c1 = ScatterChart(scatterStyle='marker')
                    # c1.style = 8
                    if ymin:
                        c1.y_axis.scaling.min = ymin  # y坐标的区间
                    if ymax:
                        c1.y_axis.scaling.max = ymax
                    c1.title = charname
                    xvalues = Reference(ws_data, min_col=cats_min_col, min_row=cats_min_row, max_col=cats_max_col,max_row=cats_max_row)
                    for item in dataIndexList:
                        data = Reference(ws_data, min_col=item['data_min_col'], min_row=item['data_min_row'],max_col=item['data_max_col'], max_row=item['data_max_row'])
                        series = Series(data, xvalues, title_from_data=True)
                        series.marker = marker.Marker('circle')
                        series.graphicalProperties.line.noFill = True
                        c1.series.append(series)

                else:
                    raise Exception('chart type wrong')

                ws.add_chart(c1, collist[rowNo] + str(plt_row_num))
                plt_row_num += (18 * (rowNo % 2))
                rowNo ^= 1
        try:
            del wb['plotcmd']
        except Exception as delErr:
            log_info_to_file("excel_report_debug_%s.log" % datetime.now().strftime("%Y-%m-%d"), "plotcmd delete err: %s" % delErr.__str__())

        try:
            del wb['plot']
        except Exception as delErr:
            log_info_to_file("excel_report_debug_%s.log" % datetime.now().strftime("%Y-%m-%d"), "plot delete err: %s" % delErr.__str__())

    except Exception as e:
        print(e.__str__())
        return dict(error=1, msg=e.__str__())
    return dict(error=0, msg=wb)

'''
将ptNameFrom积分累加，更新ptNameTo
'''
def accum_point(ptNameFrom, ptNameTo, fRatio=1.0,actTime=None):
    if app.config['MODE_HISTORY']:
        tNow = app.config['MODE_HISTORY_AT_TIME']
    elif isinstance(actTime, str):
        try:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        except:
            tNow = datetime.now()
    elif isinstance(actTime, datetime):
        tNow = actTime
    else:
        tNow = datetime.now()
    bFirstTimeOfDay = False
    if tNow.hour == 0 and tNow.minute == 0 and tNow.second == 0:
        bFirstTimeOfDay = True

    fAccumToday = None

    if bFirstTimeOfDay:
        tYesterdayStart = tNow - timedelta(days=1)
        strYestodayStart = tYesterdayStart.strftime('%Y-%m-%d %H:%M:%S')
        fValueStart = get_data_at_time(ptNameTo, strYestodayStart)

        if fValueStart is None:
            fValueStart = 0

        if isinstance(ptNameFrom, int) or isinstance(ptNameFrom, float):
            fAccumToday = ptNameFrom * 24
        else:
            fAccumToday = get_yesterday_data_sum(ptNameFrom, 3, tNow.strftime('%Y-%m-%d %H:%M:%S'))

        if fAccumToday is None:
            return None
    else:
        tTodayStart = time_get_today_start(tNow)


        strTodayStart = tTodayStart.strftime('%Y-%m-%d %H:%M:%S')
        fValueStart = get_data_at_time(ptNameTo, strTodayStart)

        if fValueStart is None or fValueStart == "":
            fValueStart = 0

        if isinstance(ptNameFrom, int) or isinstance(ptNameFrom, float):
            fAccumToday = ptNameFrom*(tNow.hour+ tNow.minute/60.0 + tNow.second/60.0/60.0)
        else:
            fAccumToday = get_today_data_sum(ptNameFrom, 3, tNow.strftime('%Y-%m-%d %H:%M:%S'))
        if fAccumToday is None:
            return None

    try:
        fNew = fValueStart+fAccumToday*fRatio
        return fNew
    except:
        return None

def get_data_max_in_day_range(strPointName, day, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    year = tNow.strftime('%Y')
    month = tNow.strftime('%m')

    datebegin = '{}-{}-{:0>2} {}:{}:{}'.format(year, month, day, '00', '00', '00')
    tBegin = datetime.strptime(datebegin, '%Y-%m-%d %H:%M:%S')
    tEnd = tBegin + timedelta(days=1)
    dateend = tEnd.strftime('%Y-%m-%d %H:%M:%S')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, datebegin,dateend, 'h1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fMax = numpy.max(ddArry)
        fMax = round(fMax, nDecimal)
    except:
        fMax = ''
    return fMax


def get_data_min_in_day_range(strPointName, day, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    year = tNow.strftime('%Y')
    month = tNow.strftime('%m')

    datebegin = '{}-{}-{:0>2} {}:{}:{}'.format(year, month, day, '00', '00', '00')
    tBegin = datetime.strptime(datebegin, '%Y-%m-%d %H:%M:%S')
    tEnd = tBegin + timedelta(days=1)
    dateend = tEnd.strftime('%Y-%m-%d %H:%M:%S')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, datebegin,dateend, 'h1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fMin = numpy.min(ddArry)
        fMin = round(fMin, nDecimal)
    except:
        fMin = ''
    return fMin



def get_data_avg_in_day_range(strPointName, day, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    year = tNow.strftime('%Y')
    month = tNow.strftime('%m')

    datebegin = '{}-{}-{:0>2} {}:{}:{}'.format(year, month, day, '00', '00', '00')
    tBegin = datetime.strptime(datebegin, '%Y-%m-%d %H:%M:%S')
    tEnd = tBegin + timedelta(days=1)
    dateend = tEnd.strftime('%Y-%m-%d %H:%M:%S')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, datebegin,dateend, 'h1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fAvg = numpy.average(ddArry)
        fAvg = round(fAvg, nDecimal)
    except:
        fAvg = ''
    return fAvg

def get_lastmonth_delta(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tBeginMonth = tNow.replace(day=1, hour=0, minute=0, second=0)
    tBeginLastMonth = tBeginMonth - timedelta(days=1)
    tBeginLastMonth = tBeginLastMonth.replace(day=1, hour=0, minute=0, second=0)

    strTimeFromAllFormat = tBeginLastMonth.strftime('%Y-%m-%d 00:00:00')
    strTimeToAllFormat = tBeginMonth.strftime('%Y-%m-%d 00:00:00')

    fValueBase = get_data_at_time(strPointName, strTimeFromAllFormat)
    if fValueBase is None:
        fValueBase = 0

    fValueCur = get_data_at_time(strPointName, strTimeToAllFormat)
    if fValueCur is None:
        return None

    try:
        return round(fValueCur-fValueBase,nDecimal)

    except:
        return None


"""
strStartTime：字符串，格式为 "nDay HH:MM"，如"26 08:00"，即开始计算的天和时间分别为 26号和08:00
"""
def get_thismonth_delta(strPointName, nDecimal=2, actTime=None, strStartTime=None):
    bCalRealTimeMode = True
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        bCalRealTimeMode = False
    else:
        tNow = datetime.now()

    tBeginMonth = tNow.replace(day=1, hour=0, minute=0, second=0)
    if strStartTime:
        tHourMinute = None
        nStartDay = None
        nStartMonth = None
        nStartYear = None
        try:
            nStartDay = int(strStartTime.split(" ")[0])
            tHourMinute = datetime.strptime(strStartTime.split(" ")[1], "%H:%M")

            if datetime.strptime("1900-01-{0} {1}:{2}:00".format(tNow.day, tNow.hour, tNow.minute), '%Y-%m-%d %H:%M:%S') >= \
                    datetime.strptime("1900-01-{0} {1}:{2}:00".format(nStartDay, tHourMinute.hour, tHourMinute.minute), '%Y-%m-%d %H:%M:%S'):
                nStartMonth = tNow.month
            else:
                nStartMonth = (time_get_first_day_of_this_month(tNow) - timedelta(days=1)).month

            if tNow.month != 1:
                nStartYear = tNow.year
            else:
                if datetime.strptime("1900-01-{0} {1}:{2}:00".format(tNow.day, tNow.hour, tNow.minute), '%Y-%m-%d %H:%M:%S') >=\
                        datetime.strptime("1900-01-{0} {1}:{2}:00".format(nStartDay, tHourMinute.hour, tHourMinute.minute), '%Y-%m-%d %H:%M:%S'):
                    nStartYear = tNow.year
                else:
                    nStartYear = (time_get_first_day_of_this_month(tNow) - timedelta(days=1)).year

        except Exception as e:
            logging.error("ERROR in get_thismonth_delta: %s" % e.__str__())

        if not tHourMinute or not nStartDay or not nStartMonth:
            return None

        tBeginMonth = tNow.replace(year=nStartYear, month=nStartMonth, day=nStartDay, hour=tHourMinute.hour, minute=tHourMinute.minute, second=0)

    strTimeFromAllFormat = tBeginMonth.strftime('%Y-%m-%d %H:%M:00')
    strTimeToAllFormat = tNow.strftime('%Y-%m-%d %H:%M:00')

    fValueBase = get_data_at_time(strPointName, strTimeFromAllFormat)
    if fValueBase is None:
        fValueBase = 0

    fValueCur = None
    if bCalRealTimeMode:
        try:
            strCurValue = BEOPDataAccess.getInstance().getInputTable([strPointName])[1].get(strPointName)
            fValueCur = float(strCurValue)
        except Exception as ee:
            logging.error('ERROR in get_thismonth_delta:%s' % (ee.__str__()))
            return None
    else:
        fValueCur = get_data_at_time(strPointName, strTimeToAllFormat)

    if fValueCur is None:
        return None

    try:
        return round(fValueCur-fValueBase,nDecimal)
    except:
        return None


def get_thismonth_delta_from_special_day(strPointName, nDayFrom=1, nDecimal=2, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
        tNow = tNow.replace(second=0)
    tBeginMonth = tNow.replace(day=nDayFrom, hour=0, minute=0, second=0)
    if tNow<=tBeginMonth:
        nNewMonth =tBeginMonth.month-1
        nNewYear = tBeginMonth.year
        if nNewMonth<=0:
            nNewMonth = 12
            nNewYear = tBeginMonth.year-1
        tBeginMonth = tBeginMonth.replace(year= nNewYear, month = nNewMonth)

    strTimeFromAllFormat = tBeginMonth.strftime('%Y-%m-%d 00:00:00')
    strTimeToAllFormat = tNow.strftime('%Y-%m-%d %H:%M:00')

    fValueBase = get_data_at_time(strPointName, strTimeFromAllFormat)
    if fValueBase is None:
        fValueBase = 0

    fValueCur = get_data_at_time(strPointName, strTimeToAllFormat)
    if fValueCur is None:
        return None

    try:
        return round(fValueCur-fValueBase,nDecimal)

    except:
        return None


    return  None

def get_lastyear_sametime_value(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tSametimeLastYear = tNow.replace(year = tNow.year-1)
    strTimeFromAllFormat = tSametimeLastYear.strftime('%Y-%m-%d %H:%M:00')

    fValueBase = get_data_at_time(strPointName, strTimeFromAllFormat)
    if fValueBase is None:
        fValueBase = 0

    try:
        return round(fValueBase, nDecimal)
    except:
        return 0


    return  0

def get_recent_days_data_sum(strPointName, nDecimal=2, nDays=0,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    tFrom = tNow - timedelta(days=nDays)

    strTimeFromAllFormat = tFrom.strftime('%Y-%m-%d 00:00:00')
    strTimeToAllFormat = tNow.strftime('%Y-%m-%d %H:%M:00')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fSum = numpy.sum(ddArry)/60.0
        fSum = round(fSum, nDecimal)
    except:
        fSum = None
    return fSum

def get_today_data_sum(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    strTimeFromAllFormat = tNow.strftime('%Y-%m-%d 00:00:00')
    strTimeToAllFormat = tNow.strftime('%Y-%m-%d %H:%M:00')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fSum = numpy.sum(ddArry)/60.0
        fSum = round(fSum, nDecimal)
    except:
        fSum = None
    return fSum


def get_today_data_average(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    strTimeFromAllFormat = tNow.strftime('%Y-%m-%d 00:00:00')
    strTimeToAllFormat = tNow.strftime('%Y-%m-%d %H:%M:00')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fSum = numpy.sum(ddArry)/len(dd)
        fSum = round(fSum, nDecimal)
    except:
        fSum = ''
    return fSum


def get_yesterday_data_average(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    tYesterdayBegin = time_get_day_start(tNow-timedelta(days=1))
    tYesterdayEnd = time_get_day_end(tNow - timedelta(days=1))

    strTimeFromAllFormat = tYesterdayBegin.strftime('%Y-%m-%d 00:00:00')
    strTimeToAllFormat = tYesterdayEnd.strftime('%Y-%m-%d 23:59:59')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fSum = numpy.sum(ddArry)/len(dd)
        fSum = round(fSum, nDecimal)
    except:
        fSum = ''
    return fSum

def get_yesterday_data_sum(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    tYesterdayBegin = time_get_day_start(tNow-timedelta(days=1))
    tYesterdayEnd = time_get_day_end(tNow - timedelta(days=1))

    strTimeFromAllFormat = tYesterdayBegin.strftime('%Y-%m-%d 00:00:00')
    strTimeToAllFormat = tYesterdayEnd.strftime('%Y-%m-%d %H:%M:%S')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fSum = numpy.sum(ddArry)/60.0
        fSum = round(fSum, nDecimal)
    except:
        fSum = None
    return fSum
"""
strTimeFrom, strTimeTo: 形如 09:00
"""
def get_today_data_avg_in_time_range(strPointName, strTimeFrom, strTimeTo, nDecimal=2, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    if not re.match("^[0-9]{2}:[0-9]{2}$", strTimeFrom) or not re.match("^[0-9]{2}:[0-9]{2}$", strTimeTo):
        return ""
    tTimeFrom = datetime.strptime(strTimeFrom, "%H:%M")
    tTimeTo = datetime.strptime(strTimeTo, "%H:%M")
    tTimeStart = tNow.replace(hour=tTimeFrom.hour, minute=tTimeFrom.minute, second=0)
    tTimeEnd = tNow.replace(hour=tTimeTo.hour, minute=tTimeTo.minute, second=0)
    strTimeStart = tTimeStart.strftime("%Y-%m-%d %H:%M:%S")
    strTimeEnd = tTimeEnd.strftime("%Y-%m-%d %H:%M:%S")
    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeStart,
                                                                   strTimeEnd, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fAvg = numpy.average(ddArry)
        fAvg = round(fAvg, nDecimal)
    except:
        fAvg = ''
    return fAvg


def get_yesterday_data_avg_in_time_range(strPointName, strTimeFrom, strTimeTo, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tYesterday = tNow - timedelta(days=1)
    strTimeFromAllFormat = '%s %s:00'%(tYesterday.strftime('%Y-%m-%d'), strTimeFrom)
    strTimeToAllFormat = '%s %s:00'%(tYesterday.strftime('%Y-%m-%d'), strTimeTo)
    tFrom = datetime.strptime(strTimeFromAllFormat,'%Y-%m-%d %H:%M:%S')
    tTo = datetime.strptime(strTimeToAllFormat,'%Y-%m-%d %H:%M:%S')
    if tFrom>tTo: #to solve 23:00 to 0:00 problem
        tTo = tTo+ timedelta(days=1)
        strTimeToAllFormat = tTo.strftime('%Y-%m-%d %H:%M:%S')
    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''
    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ''

    try:
        ddArry = numpy.array(dd)
        fAvg = numpy.average(ddArry)
        fAvg = round(fAvg, nDecimal)
    except:
        fAvg = ''
    return fAvg


def get_today_data_at_time(strPointName, strTimeFrom, strTimeTo=None, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    strKey = "HIS__get_today_data_at_time_{point}".format(point=strPointName)

    if RedisManager.is_alive():
        resHisData = RedisManager.get(strKey)
        if resHisData != None:
            strTimeKey = "{year},{month:0>2d},{day:0>2d},{time}".format(year=tNow.year, month=tNow.month,
                                                                        day=tNow.day, time=strTimeFrom)
            value = resHisData.get(strTimeKey, None)
            if is_digit(value):
                return round(float(value), 2)

    strTimeFromAllFormat = '%s %s:00'%(tNow.strftime('%Y-%m-%d'), strTimeFrom)
    strTimeToAllFormat = '%s %s:00' % (tNow.strftime('%Y-%m-%d'), strTimeFrom)
    tFrom = datetime.strptime(strTimeFromAllFormat, '%Y-%m-%d %H:%M:%S')
    tTo = datetime.strptime(strTimeToAllFormat, '%Y-%m-%d %H:%M:%S')
    if tFrom > tTo:  # to solve 23:00 to 0:00 problem
        tTo = tTo + timedelta(days=1)
        strTimeToAllFormat = tTo.strftime('%Y-%m-%d %H:%M:%S')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat,strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return 0

    resValue = round(dd[0], nDecimal)
    if datetime.now() > tFrom:
        RedisManager.set_api_point_data_at_time("get_today_data_at_time", strPointName, tNow.year, tNow.month, tNow.day, strTimeFrom, resValue)

    try:
        return resValue
    except:
        return ''


def get_yesterday_data_at_same_time(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tNow = tNow.replace(second=0)
    tYesterday = tNow - timedelta(days=1)

    strTimeFromAllFormat = tYesterday.strftime('%Y-%m-%d')
    strTimeToAllFormat = tYesterday.strftime('%Y-%m-%d')
    tFrom = datetime.strptime(strTimeFromAllFormat, '%Y-%m-%d %H:%M:%S')
    tTo = datetime.strptime(strTimeToAllFormat, '%Y-%m-%d %H:%M:%S')
    if tFrom > tTo:  # to solve 23:00 to 0:00 problem
        tTo = tTo + timedelta(days=1)
        strTimeToAllFormat = tTo.strftime('%Y-%m-%d %H:%M:%S')

    if RedisManager.is_alive() and tFrom.minute==0 and tFrom.second==0:
        strDataValue1 = RedisManager.get_history_data(strPointName, tFrom)
        if strDataValue1=='':
            return ''

        if strDataValue1  is not None:
            fDelta = float(strDataValue1)
            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat,strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''


    if len(dd)==0:
        return 0

    try:
        return round(dd[0], nDecimal)

    except:
        return ''


    return  ''


def get_last_week_data_at_same_time(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tNow = tNow.replace(second=0)
    tSameTimeLastWeek = tNow - timedelta(days=7)

    strTimeFromAllFormat = tSameTimeLastWeek.strftime('%Y-%m-%d')
    strTimeToAllFormat = tSameTimeLastWeek.strftime('%Y-%m-%d')
    tFrom = datetime.strptime(strTimeFromAllFormat, '%Y-%m-%d %H:%M:%S')
    tTo = datetime.strptime(strTimeToAllFormat, '%Y-%m-%d %H:%M:%S')
    if tFrom > tTo:  # to solve 23:00 to 0:00 problem
        tTo = tTo + timedelta(days=1)
        strTimeToAllFormat = tTo.strftime('%Y-%m-%d %H:%M:%S')

    if RedisManager.is_alive() and tFrom.minute==0 and tFrom.second==0:
        strDataValue1 = RedisManager.get_history_data(strPointName, tFrom)
        if strDataValue1=='':
            return ''

        if strDataValue1  is not None:
            fDelta = float(strDataValue1)
            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat,strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''


    if len(dd)==0:
        return 0

    try:
        return round(dd[0], nDecimal)

    except:
        return ''

def get_yesterday_data_at_time(strPointName, strTimeFrom, strTimeTo=None, nDecimal=2, reportCall=False, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tYesterday = tNow - timedelta(days=1)
    strTimeFromAllFormat = '%s %s:00'%(tYesterday.strftime('%Y-%m-%d'), strTimeFrom)
    strTimeToAllFormat = '%s %s:00' % (tYesterday.strftime('%Y-%m-%d'), strTimeFrom)
    tFrom = datetime.strptime(strTimeFromAllFormat, '%Y-%m-%d %H:%M:%S')
    tTo = datetime.strptime(strTimeToAllFormat, '%Y-%m-%d %H:%M:%S')
    if tFrom > tTo:  # to solve 23:00 to 0:00 problem
        tTo = tTo + timedelta(days=1)
        strTimeToAllFormat = tTo.strftime('%Y-%m-%d %H:%M:%S')

    strKey = "HIS__get_yesterday_data_at_time_{point}".format(point=strPointName)

    if RedisManager.is_alive():
        resHisData = RedisManager.get(strKey)
        if resHisData != None:
            strTimeKey = "{year},{month:0>2d},{day:0>2d},{time}".format(year=tYesterday.year, month=tYesterday.month, day=tYesterday.day, time=strTimeFrom)
            value = resHisData.get(strTimeKey, None)
            if is_digit(value):
                return round(float(value), 2)

    if RedisManager.is_alive() and tFrom.minute==0 and tFrom.second==0:
        strDataValue1 = RedisManager.get_history_data(strPointName, tFrom)

        if reportCall:
            strLogInfo = "data from redis: {data}(pointName:{point})".format(data=strDataValue1, point=strPointName)
            api_log_info("get_yesterday_data_at_time_{time}.log".format(time=datetime.now().strftime("%Y-%m-%d")), strLogInfo)

        # 2022-01-21 (russell与golding讨论) 由于上汽问题，redis中拿不到值后去mysql再查一次
        if strDataValue1 != None and strDataValue1 != "" and is_digit(strDataValue1):
            fDelta = float(strDataValue1)
            if datetime.now() > tFrom:
                RedisManager.set_api_point_data_at_time("get_yesterday_data_at_time", strPointName, tYesterday.year, tYesterday.month, tYesterday.day, strTimeFrom, fDelta)

            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat,strTimeToAllFormat, 'm1')

    if reportCall:
        strLogInfo = "data from mysql: {data}(pointName:{point})".format(data=allData.get("map", {}).get(strPointName, []), point=strPointName)
        api_log_info("get_yesterday_data_at_time_{time}.log".format(time=datetime.now().strftime("%Y-%m-%d")), strLogInfo)

    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return ""

    resValue = round(dd[0], nDecimal)
    if datetime.now() > tFrom:
        RedisManager.set_api_point_data_at_time("get_yesterday_data_at_time", strPointName, tYesterday.year, tYesterday.month, tYesterday.day, strTimeFrom, resValue)

    try:
        return resValue
    except:
        return ''


def get_data_at_time(strPointName, strTimeAt,  nDecimal=2,actTime=None):
    if isinstance(strTimeAt, str):
        strTimeToAllFormat = strTimeAt
        tTimeAt = datetime.strptime(strTimeAt, '%Y-%m-%d %H:%M:%S')
        tTimeFrom = tTimeAt.replace(hour=0, minute=0, second=0)
        strTimeFromAllFormat = tTimeFrom.strftime('%Y-%m-%d %H:%M:%S')
    else:
        strTimeToAllFormat = strTimeAt.strftime('%Y-%m-%d %H:%M:%S')
        tTimeAt = strTimeAt
        tTimeFrom = tTimeAt.replace(hour=0, minute=0, second=0)
        strTimeFromAllFormat = tTimeFrom.strftime('%Y-%m-%d %H:%M:%S')

    if RedisManager.is_alive() and tTimeAt.minute==0 and tTimeAt.second==0:
        strDataValue1 = RedisManager.get_history_data(strPointName, tTimeAt)

        if is_digit(strDataValue1):
            fDelta = float(strDataValue1)
            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat,strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return None


    if len(dd)==0:
        return None

    try:
        return round(dd[-1], nDecimal)

    except:
        return None


    return None

"""
作用：获取当前小时的累积量差值，如nHour输入4，则计算该日内4:00与3:00 的累积量之差 
strPointName: string, 点名
nHour: int, 小时
示例：get_this_hour_data_delta(strPointName, 8)
"""
def get_this_hour_data_delta(strPointName, nHour, nDecimal=2, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tTimeTo = tNow.replace(hour=nHour, minute=0, second=0)
    tTimeStart = tTimeTo - timedelta(hours=1)

    strTimeStart = tTimeStart.strftime("%Y-%m-%d %H:%M:%S")
    strTomeTo = tTimeTo.strftime("%Y-%m-%d %H:%M:%S")

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeStart, strTomeTo, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''
    if len(dd) == 0:
        return 0

    try:
        ddArry = numpy.array(dd)

    except:
        return ''

    return round(ddArry[-1] - ddArry[0], nDecimal)

"""
strTimeLine: str, %H:%M 时间分界线，默认为None，即时间分界线为00:00，时间分界线可以任意给定，如给定08:00，则08:00开始今日累加
"""
def get_hour_data_delta(strPointName, nDecimal=2, actTime=None, strTimeLine=None):
    if strTimeLine:
        if not isValidDate(strTimeLine, "%H:%S"):
            logging.error("ERROR in get_hour_data_delta: invalid strStartTime: %s" % strTimeLine)
            return ""

    nHour = 0
    nMinute = 0
    if strTimeLine:
        tTimeLine = datetime.strptime(strTimeLine, "%H:%S")
        nHour = tTimeLine.hour
        nMinute = tTimeLine.minute


    bCalRealtimeMode = True
    tNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        bCalRealtimeMode = False

    tFrom = tNow.replace(minute=0, second=0, microsecond=0)
    strTimeFrom = tFrom.strftime('%Y-%m-%d %H:%M:%S')
    strTimeTo = tNow.strftime('%Y-%m-%d %H:%M:%S')

    fDataValueAtActTime = None
    if bCalRealtimeMode:
        try:
            strDataValue2 = BEOPDataAccess.getInstance().getInputTable([strPointName])[1].get(strPointName)
            fDataValueAtActTime = float(strDataValue2)
        except Exception as ee:
            logging.error('ERROR in get_today_data_delta:%s'%(ee.__str__()))
            return None
    else:
        fDataValueAtActTime = get_data_at_time(strPointName, strTimeTo)

    if RedisManager.is_alive():
        strDataValue1 = RedisManager.get_history_data(strPointName, tFrom)
        if strDataValue1=='':
            return ''
        if strDataValue1  is not None and fDataValueAtActTime is not None:
            fDelta = fDataValueAtActTime - float(strDataValue1)
            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFrom, strTimeTo, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''
    if len(dd)==0:
        return 0

    try:
        ddArry = numpy.array(dd)
    except:
        return ''

    dataFrom = 0 if tFrom.hour == nHour and tFrom.minute == nMinute else ddArry[0]

    return round(ddArry[-1] - dataFrom, nDecimal)

"""
strTimeLine: str, %H:%M 时间分界线，默认为None，即时间分界线为00:00，时间分界线可以任意给定，如给定08:00，则08:00开始今日累加
"""
def get_today_data_delta(strPointName, nDecimal=2, actTime=None, strTimeLine=None):
    if strTimeLine:
        if not isValidDate(strTimeLine, "%H:%S"):
            logging.error("ERROR in get_today_data_delta: invalid strStartTime: %s" % strTimeLine)
            return ""

    nHour = 0
    nMinute = 0
    if strTimeLine:
        tTimeLine = datetime.strptime(strTimeLine, "%H:%S")
        nHour = tTimeLine.hour
        nMinute = tTimeLine.minute

    bCalRealtimeMode = True
    tNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        bCalRealtimeMode = False

    if tNow.hour == nHour and tNow.minute == nMinute:  #0:00要算昨日的一天差（零点到零点）
        tFrom = (tNow - timedelta(days=1)).replace(hour=nHour, minute=nMinute, second=0, microsecond=0)
        tTo = tNow
    else:
        tDateLine = tNow.replace(hour=nHour, minute=nMinute, second=0, microsecond=0)
        tFrom = tDateLine
        tTo = tNow
        if tNow < tDateLine:
            tFrom = (tNow - timedelta(days=1)).replace(hour=nHour, minute=nMinute, second=0, microsecond=0)

    strFrom = tFrom.strftime("%Y-%m-%d %H:%M:%S")
    strTo = tTo.strftime("%Y-%m-%d %H:%M:%S")

    if bCalRealtimeMode:
        try:
            strDataValue2 = BEOPDataAccess.getInstance().getInputTable([strPointName])[1].get(strPointName)
            fDataValueAtActTime = float(strDataValue2)
        except Exception as ee:
            logging.error('ERROR in get_today_data_delta:%s'%(ee.__str__()))
            return None
    else:
        fDataValueAtActTime = get_data_at_time(strPointName, strTo)

    if RedisManager.is_alive():
        strDataValue1 = RedisManager.get_history_data(strPointName, tFrom)
        if strDataValue1=='':
            return ''
        if strDataValue1  is not None and fDataValueAtActTime is not None:
            fDelta = fDataValueAtActTime - float(strDataValue1)
            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strFrom, strTo, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd)==0:
        return 0

    try:
        ddArry = numpy.array(dd)

    except:
        return ''

    return round(ddArry[-1] - ddArry[0], nDecimal)


def get_yesterday_data_delta(strPointName, strTimeFrom="00:00", strTimeTo="00:00", nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tYesterday = tNow - timedelta(days=1)
    strTimeFromAllFormat = '%s %s:00'%(tYesterday.strftime('%Y-%m-%d'), strTimeFrom)
    if strTimeTo=="00:00":
        strTimeToAllFormat = '%s %s:00' % (tNow.strftime('%Y-%m-%d'), strTimeTo)
    else:
        strTimeToAllFormat = '%s %s:00' % (tYesterday.strftime('%Y-%m-%d'), strTimeTo)
    tFrom = datetime.strptime(strTimeFromAllFormat, '%Y-%m-%d %H:%M:%S')
    tTo = datetime.strptime(strTimeToAllFormat, '%Y-%m-%d %H:%M:%S')
    if tFrom >= tTo:  # to solve 23:00 to 0:00 problem
        tTo = tTo + timedelta(days=1)
        strTimeToAllFormat = tTo.strftime('%Y-%m-%d %H:%M:%S')

    if RedisManager.is_alive():
        strDataValue1 = RedisManager.get_history_data(strPointName, tFrom)
        strDataValue2 = RedisManager.get_history_data(strPointName, tTo)
        if strDataValue1=='' or strDataValue2=='':
            return ''

        if strDataValue1  is not None and strDataValue2 is not None:
            fDelta = float(strDataValue2) - float(strDataValue1)
            return round(fDelta, nDecimal)


    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''


    if len(dd)==0:
        return 0

    try:
        ddArry = numpy.array(dd)

    except:
        return ''


    return  round(ddArry[-1] - ddArry[0], nDecimal)


def get_month_delta_this_year(strPointName,nMonth,nDecimal=2,actTime=None):
    tRealNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = tRealNow

    try:
        tFrom = time_get_first_day_of_one_month_this_year(nMonth, tNow)
        tTo = time_get_first_day_of_next_month_this_year(nMonth, tNow)
        if tFrom> tRealNow:
            return ''

        if tTo>tRealNow:
            tTo = tRealNow.replace(second=0)
        fv_start = get_data_at_time(strPointName, tFrom)
        fv_end = get_data_at_time(strPointName, tTo)

        if fv_start is not None and fv_end is not None:
            return round(fv_end - fv_start, nDecimal)
        else:
            return ''
    except Exception as e:
        return ''

def get_day_time_delta_this_month(strPointName, nDay, strHourMinute, nDecimal=2, actTime=None):
    tNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

    #每月1号做本计算时，this_month还是上个月，还是在计算上月的值
    if tNow.day==1:
        tNow = tNow- timedelta(days=2)

    tHourMinute = datetime.strptime(strHourMinute, "%H:%M")



    tBegin = tNow.replace(day=nDay, hour=tHourMinute.hour, minute=tHourMinute.minute, second=0, microsecond=0)
    tEnd = tBegin + timedelta(days=1)
    if tEnd>datetime.now(): #如果还没发生，不能计算出真实结果
        return ''

    strBegin = tBegin.strftime("%Y-%m-%d %H:%M:%S")
    stEnd = tEnd.strftime("%Y-%m-%d %H:%M:%S")

    if RedisManager.is_alive():
        strDataValue1 = RedisManager.get_history_data(strPointName, tBegin)
        strDataValue2 = RedisManager.get_history_data(strPointName, tEnd)
        if strDataValue1=='' or strDataValue2=='':
            return ''

        if strDataValue1  is not None and strDataValue2 is not None:
            fDelta = float(strDataValue2) - float(strDataValue1)
            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strBegin, stEnd, 'h1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return 0

    try:
        ddArry = numpy.array(dd)
    except:
        return ''

    return round(ddArry[-1] - ddArry[0], nDecimal)


def get_day_delta_this_month(strPointName,day,nDecimal=2,actTime=None, reportCall=False):

    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    year = tNow.strftime('%Y')
    month = tNow.strftime('%m')

    datebegin = '{}-{}-{:0>2} {}:{}:{}'.format(year, month, day, '00', '00', '00')
    tBegin = datetime.strptime(datebegin, '%Y-%m-%d %H:%M:%S')
    if tBegin > datetime.now():
        return 0

    tEnd = tBegin + timedelta(days=1)
    if tEnd > datetime.now():
        tEnd = datetime.now() - timedelta(minutes=1)

    dateend = tEnd.strftime('%Y-%m-%d %H:%M:%S')

    if RedisManager.is_alive():
        strDataValue1 = RedisManager.get_history_data(strPointName, tBegin)
        strDataValue2 = RedisManager.get_history_data(strPointName, tEnd)

        strError = "strDataValue1:%s, strDataValue2: %s" % (strDataValue1, strDataValue2)
        if reportCall:
            log_info_to_file('dompysite_report_get_day_delta_this_month_log_%s.log' % datetime.now().strftime('%Y-%m-%d'), strError)

        # if strDataValue1=='' or strDataValue2=='':
        #     return ''

        if strDataValue1 is not None and strDataValue2 is not None and strDataValue1 != "" and strDataValue2 != "":
            fDelta = float(strDataValue2) - float(strDataValue1)
            if reportCall:
                strError = "fDelta: %s" % fDelta
                log_info_to_file('dompysite_report_get_day_delta_this_month_log_%s.log' % datetime.now().strftime('%Y-%m-%d'), strError)

            return round(fDelta, nDecimal)

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, datebegin, dateend, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return ""

    try:
        ddArry = numpy.array(dd)
    except:
        return ''

    if reportCall:
        strError = "ddArray: %s, %s" % (ddArry[-1], ddArry[0])
        log_info_to_file('dompysite_report_get_day_delta_this_month_log_%s.log' % datetime.now().strftime('%Y-%m-%d'), strError)

    return round(ddArry[-1] - ddArry[0], nDecimal)


def get_day_delta_at_month_this_year(strPointName, day, month, nDecimal=2,actTime=None):

    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    year = tNow.strftime('%Y')

    datebegin = '{}-{}-{:0>2} {}:{}:{}'.format(year, month, day, '00', '00', '00')
    tBegin = datetime.strptime(datebegin, '%Y-%m-%d %H:%M:%S')
    tEnd = tBegin + timedelta(days=1)
    dateend = tEnd.strftime('%Y-%m-%d %H:%M:%S')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, datebegin, dateend, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return 0

    try:
        ddArry = numpy.array(dd)
    except:
        return ''

    return round(ddArry[-1] - ddArry[0], nDecimal)

def get_day_delta_last_month(strPointName, nDay, nDecimal=2, actTime=None):
    try:
        tNow = datetime.now()
        if actTime:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        firstDayThisMonth = tNow.replace(day=1)
        tLastMonth = firstDayThisMonth - timedelta(days=1)
        tTheDay = tLastMonth.replace(day=nDay)
        tTheDayBegin = time_get_day_start(tTheDay)

        # tTheDayEnd = time_get_day_end(tTheDay)
        # tTheDayEnd 修改为第二天的0点
        tTheDayEnd = tTheDayBegin + timedelta(days=1)

        strTimeFromFormat = tTheDayBegin.strftime('%Y-%m-%d 00:00:00')
        strTimeToFormat = tTheDayEnd.strftime('%Y-%m-%d 00:00:00')

        if RedisManager.is_alive():
            strDataValue1 = RedisManager.get_history_data(strPointName, tTheDayBegin)
            strDataValue2 = RedisManager.get_history_data(strPointName, tTheDayEnd)
            if strDataValue1 == '' or strDataValue2 == '':
                return None

            if strDataValue1 is not None and strDataValue2 is not None:
                fDelta = float(strDataValue2) - float(strDataValue1)
                return round(fDelta, nDecimal)

        allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromFormat, strTimeToFormat, 'm1')
        if not allData:
            return ''

        if allData.get('map') is None:
            return ''

        dd = allData['map'].get(strPointName)

        if dd is None or not isinstance(dd, list):
            return ''

        if len(dd) == 0:
            return 0

        try:
            ddArry = numpy.array(dd)
        except:
            return ''

        return round(ddArry[-1] - ddArry[0], nDecimal)
    except:
        return ""


def fillSheetMonthStatic(wb,actTime=None):
    ws = wb.get_sheet_by_name('本月统计')
    dateForSearchCol = []
    maxcol = ws.max_column
    for i in range(maxcol):
        content = ws.cell(row=2, column=i+1).value
        if isinstance(content, str):
            if 'get' in content:
                dateForSearchCol.append(dict(col=i+1, content=content.replace('get', '')))
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    datelist = []
    year = tNow.strftime('%Y')
    month = tNow.strftime('%m')
    day = tNow.strftime('%d')
    day_ = 2
    while day_ < int(day):
        strtime = '{}-{}-{:0>2} {}:{}:{}'.format(year, month, day_, '00', '00', '00')
        # time = datetime.strptime(strtime, '%Y-%m-%d %H:%M:%S')
        datelist.append(strtime)
        day_ += 1
    rowNo = 2
    for date in datelist:
        ws.cell(row=rowNo, column=1).value = '{}-{}-{:0>2}'.format(date[0:4], date[5:7], int(date[8:10])-1)
        for item in dateForSearchCol:
            ws.cell(row=rowNo, column=item['col']).value = get_yesterday_data_delta(item['content'], '0:00', '0:00', nDecimal=2, actTime=date)
        rowNo += 1
    return wb

class ReportTool:

    def __init__(self, templateName, fileName):
        self._templateName = templateName
        self._fileName = fileName

    def getPointListFromString(self, myText):
        strPointNameList = []
        if not isinstance(myText, str):
            return strPointNameList
        if len(myText) <= 1:
            return strPointNameList
        nEvalStart = myText.find('<%')
        nEvalEnd = myText.find('%>')
        while nEvalStart>=0 and nEvalEnd >= 0:
            strPointName = myText[nEvalStart + 2:nEvalEnd ]
            strPointNameList.append(strPointName)
            nEvalStart = myText.find('<%', nEvalEnd)
            nEvalEnd = myText.find('%>',nEvalStart)

        return strPointNameList

    def replaceVariableInText(self, myText):
        newText = myText.replace('<%', '')
        newText = newText.replace('%>', '')
        return newText

    def fillAllTables(self, strFileName, actTime=None):
        RedisManager.set_report_start(strFileName)
        is_lastTable = True
        map_var_cmd = {}
        realFileName = os.path.split(self._fileName)[1]
        try:
            d = Document(self._fileName)
            d.styles['Normal'].font.name = u'宋体'
            d.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        except Exception as e:
            strError = 'ERROR in Open word document:%s: %s'%(self._fileName, e.__str__())
            logging.error(strError)
            print(strError)
            return False

        tables = d.tables
        strPointNameList = []

        nTotalCells = 0
        for table in reversed(tables):
            # 将配置写入表格1
            for row in table.rows:
                for cell in row.cells:
                    nTotalCells += 1
                    content = str(cell.text)

                    if 'get_' not in content:
                        ptList = self.getPointListFromString(content)
                        strPointNameList.extend(ptList)

        RedisManager.set_report_process(self._templateName, 0, nTotalCells)

        rvdata = BEOPDataAccess.getInstance().getInputTable(strPointNameList)[0]
        pointValueMap = {}
        for item in rvdata:
            pointValueMap[item['name']] = item['value']

        nFinishedCells = 0
        for table in reversed(tables):
            # 将配置写入表格1
            for row in table.rows:
                for cell in row.cells:

                    myText = cell.text

                    if re.match('{{(.+)}}', myText):
                        var = re.match('{{(.+)}}', myText).group(1)
                        resultValue = map_var_cmd.get(var, '')
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(str(resultValue))
                        run.font.name = '宋体'

                    elif 'get_' not in myText:
                        if 'return' not in myText:
                            myText = myText.replace('‘', '\'')
                            myText = myText.replace('’', '\'')
                            myText = myText.replace('“', '\'')
                            myText = myText.replace('”', '\'')
                            myText = myText.replace('　', ' ') #中文空格
                            myText = myText.replace('\'', '"')
                            ptList = self.getPointListFromString(myText)
                            for strPointName in ptList:
                                strValue = pointValueMap.get(strPointName, '')
                                try:
                                    exec('%s = \'%s\'' % (strPointName, strValue))
                                    strEval = self.replaceVariableInText(myText)
                                    resultValue = eval(strEval)
                                    cell.paragraphs[0].clear()
                                    run = cell.paragraphs[0].add_run(str(resultValue))
                                    run.font.name = '宋体'
                                except Exception as e:
                                    cell.text = '-ERROR:%s-'%(e.__str__())
                        else:
                            func_def = '\n '.join(myText.split('\n'))
                            func_str = 'def func():\n '+func_def
                            try:
                                exec(func_str)
                                resultValue = eval('func()')
                            except Exception as e:
                                resultValue = '-ERROR:%s-'%(e.__str__())
                            for para in cell.paragraphs:
                                para.clear()
                            run = cell.paragraphs[0].add_run(str(resultValue))
                            run.font.name = '宋体'
                    else:
                        strEval = myText
                        if strEval and len(strEval) > 0 and strEval.find('get_') >= 0:
                            # print(strEval)
                            if actTime:
                                strEval = strEval[0:len(strEval) - 1] + ',actTime="' + actTime + '"' + strEval[-1]
                            resultValue = eval(strEval)
                            cell.paragraphs[0].clear()
                            run = cell.paragraphs[0].add_run(str(resultValue))
                            run.font.name = '宋体'

                    nFinishedCells += 1
                    RedisManager.set_report_process(self._templateName, nFinishedCells, nTotalCells)

            if is_lastTable:
                for row in table.rows:
                    cells = row.cells
                    try:
                        map_var_cmd[cells[0].text] = cells[1].text
                    except Exception as ee:
                        strError = 'ERROR in table.rows deal in word document:%s: %s' % (self._fileName, ee.__str__())
                        logging.error(strError)
                        print(strError)
                is_lastTable = False

        'json.loads(ConditionList).get("data")[1][3]　'
        # last_table = tables[-1]
        # last_table._element.getparent().remove(last_table._element)
        d.save(self._fileName)

        RedisManager.set_report_process(self._templateName, nTotalCells, nTotalCells)
        RedisManager.set_report_stop(self._templateName)

        return True

    def string_is_expression(self, strExpression):
        try:
            exec(strExpression)
            return True
        except:
            return False

    # 2022-02-16 russell: 处理当报表单元格中含有#时，eval函数返回结果只留数字的情况
    def calculate_report_cell_expression(self, strExp):
        checkList = ["#"]

        bIsText = False
        for item in checkList:
            if strExp.find(item) != -1:
                bIsText = True
                break

        if bIsText:
            return strExp, ""
        else:
            try:
                fValue = float(eval(strExp))
                return fValue, ""
            except Exception as err:
                return strExp, err.__str__()

    def replacePlantPrefixOfPointName(self, strReplaceFrom, strReplaceTo):
        book = None
        try:
            book = load_workbook(filename=self._fileName)
        except Exception as e:
            return False, "报表中机房前缀替换失败:{err}".format(err=e.__str__())

        if book == None:
            return False, "报表中机房前缀替换失败"

        sheetList = book.worksheets
        if not len(sheetList):
            return True, ""

        try:
            for oSheet in sheetList:
                rowList = oSheet.rows
                for row in rowList:
                    for cell in row:
                        content = cell.value
                        if not isinstance(content, str):
                            continue

                        rt = content.replace(strReplaceFrom, strReplaceTo)
                        cell.value = rt

        except Exception as e:
            return False, "报表中机房前缀替换失败: {err}".format(err=e.__str__())

        try:
            book.save(self._fileName)
        except Exception as e:
            return False, "报表中机房前缀替换失败: {err}".format(err=e.__str__())

        return True, ""


    def fillAllTablesInExcel(self, actTime, strFileName):
        RedisManager.set_report_start(strFileName)
        is_varSheet = True
        file_home = self._fileName
        realFileName = os.path.split(self._fileName)[1]

        try:
            wb = load_workbook(filename=file_home)
        except Exception as e:
            strError = 'ERROR in fillAllTablesInExcel' + e.__str__()
            print(strError)
            log_info_to_file('domJobs_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')),
                             strError)
            return dict(msg='加载模板出错，请确认是否在factory中上传了模板以及模板名是否正确。', err=1)

        map_var_cmd = {}
        wslist = wb.worksheets
        try:
            var_ws = wb.get_sheet_by_name('var')
            wslist.remove(var_ws)
            wslist.insert(0, var_ws)
        except:
            print('no vars in excel, pass')

        # 遍历所有行获取相关点名列表
        nTotalCells = 0
        strPointNameList = []
        for ws in wslist:
            rows = ws.rows

            for row in rows:
                for cell in row:
                    content = cell.value
                    if not isinstance(content, str):
                        continue

                    if "get_" in content:
                        continue

                    content = str(content)
                    ptList = self.getPointListFromString(content)
                    strPointNameList.extend(ptList)

                    nTotalCells += 1

        RedisManager.set_report_process(self._templateName, 0, nTotalCells)

        # 读相关点的实时值
        pointValueMap = {}
        if strPointNameList:
            rvdata = BEOPDataAccess.getInstance().getInputTable(strPointNameList)[0]
            for item in rvdata:
                pointValueMap[item['name']] = item['value']

        nFinishedCells = 0
        for ws in wslist:
            rows = ws.rows
            # 遍历所有的行
            for row in rows:
                for cell in row:
                    myText = cell.value
                    if not isinstance(myText, str):
                        continue

                    if re.match('{{(.+)}}', myText):
                        var = re.match('{{(.+)}}', myText).group(1)
                        resultValue = map_var_cmd.get(var, '')
                        try:
                            resultValue = float(resultValue)
                        except:
                            pass
                        cell.value = resultValue

                    # 不含get_的api单元格
                    if 'get_' not in myText:
                        if 'return' not in myText:  # 若表达式中不包含return
                            myText = myText.replace('‘', '\'')
                            myText = myText.replace('’', '\'')
                            myText = myText.replace('“', '\'')
                            myText = myText.replace('”', '\'')
                            myText = myText.replace('　', ' ')  # 中文空格
                            myText = myText.replace('\'', '"')
                            ptList = self.getPointListFromString(myText)

                            for strPointName in ptList:
                                strValue = pointValueMap.get(strPointName, '')
                                try:
                                    myText = myText.replace("<%{pointName}%>".format(pointName=strPointName), strValue)
                                except Exception as e:
                                    strErrMsg = "ERROR happend when try to replace {point} to realtime value(exp: {exp},errDetail:{errDetail})".format(point=strPointName, errDetail=e.__str__(), exp=myText)
                                    cell.value = strErrMsg
                                    log_info_to_file('dompysite_report_%s_log_%s.log' % (
                                    self._templateName, datetime.now().strftime('%Y_%m')), strErrMsg)

                            if not self.string_is_expression(myText):
                                continue

                            fResultValue, errorMessage = self.calculate_report_cell_expression(myText)
                            cell.number_format = "General"
                            if len(errorMessage):
                                cell.value = "-ERROR in eval cell content: (content:{content}, err:{err})".format(content=myText, err=errorMessage)
                                strError = '****[ERROR]****expression:{expression}; err: {err}; result: {result}'.format(expression=myText, err=errorMessage, result=fResultValue)
                                log_info_to_file('dompysite_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')), strError)
                            else:
                                cell.value = fResultValue
                                strSuccessInfo = '[SUCCESS]expression:{expression}; result: {result}'.format(expression=myText, result=fResultValue)
                                log_info_to_file('dompysite_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')), strSuccessInfo)

                        else:
                            func_def = '\n '.join(myText.split('\n'))
                            func_str = 'def func():\n ' + func_def
                            try:
                                exec(func_str)
                                resultValue = eval('func()')
                                try:
                                    resultValue = float(resultValue)
                                except Exception:
                                    pass
                            except Exception as e:
                                resultValue = '-ERROR:%s-' % (e.__str__())
                                log_info_to_file('dompysite_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')),'ERROR in exec %s: %s' % (func_str, resultValue))

                            cell.number_format = "General"
                            cell.value = resultValue
                    else: # 含有get_的api单元格
                        strEval = myText
                        if not strEval:
                            continue

                        if not len(strEval):
                            continue

                        if actTime:
                            for api in g_systemAPIList:
                                strEval = insert_acttime_for_api(strEval, api, actTime)

                        if not self.string_is_expression(strEval):
                            continue

                        fResultValue, errorMessage = self.calculate_report_cell_expression(strEval)
                        cell.number_format = "General"
                        if len(errorMessage):
                            cell.value = "-ERROR in eval cell content: (content:{content}, err:{err})".format(content=myText, err=errorMessage)
                            strError = '****[ERROR]****expression:{expression}; err: {err}; result: {result}'.format(expression=myText, err=errorMessage, result=fResultValue)
                            log_info_to_file('dompysite_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')), strError)
                        else:
                            cell.value = fResultValue
                            strSuccessInfo = '[SUCCESS]expression:{expression}; result: {result}'.format(expression=myText, result=fResultValue)
                            log_info_to_file('dompysite_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')), strSuccessInfo)

                    nFinishedCells += 1
                    RedisManager.set_report_process(self._templateName, nFinishedCells, nTotalCells)

            if is_varSheet:
                rows = ws.rows
                for row in rows:
                    if len(row) >= 2:
                        map_var_cmd[row[0].value] = row[1].value
                is_varSheet = False

        RedisManager.set_report_process(self._templateName, nTotalCells, nTotalCells)
        RedisManager.set_report_stop(self._templateName)

        # wb = fillSheetMonthStatic(wb, actTime)
        ret = plotInExcel(wb)
        if ret['error'] == 1:
            strError = 'error in plotInExcel: ' + ret['msg']
            print(strError)

            log_info_to_file('domJobs_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')),
                             strError)
        try:
            del wb['var']
        except:
            pass

        try:
            wb.save(self._fileName)  # 保存修改后的excel
        except Exception as ee:
            log_info_to_file('domJobs_report_%s_log_%s.log' % (self._templateName, datetime.now().strftime('%Y_%m')),
                             'ERROR in wb.save: %s' % (ee.__str__()))
        return dict(msg='', err=0)

    def add_paragraph_after(document, search_phrase, message):
        for paragraph in document.paragraphs:
            # if paragraph.text and regexp.search(paragraph.text) :
            if paragraph.text and paragraph.text==search_phrase:
                p = document.add_paragraph()
                p.text = message
                a, b = paragraph._p, p._p
                a.addnext(b)
                return p

    def add_image_after(document, imgPath, search_phrase):
        # regexp = re.compile(search_phrase)
        for paragraph in document.paragraphs:
            # if paragraph.text and regexp.search(paragraph.text) :
            if paragraph.text and paragraph.text==search_phrase:
                p = document.add_paragraph()
                r = p.add_run()
                # r.add_text('\n功率图表')
                r.add_picture(imgPath, width=Inches(6.0), height=Inches(8))
                a, b = paragraph._p, p._p
                a.addnext(b)
                return p

def find_quota_end_string(str, nQuataLeftIndex):
    nRightQuotaCount = 0
    nCurFrom = nQuataLeftIndex
    nTotalSum = 1
    while True:
        nFindIndex = str.find('(', nCurFrom+1)
        nFindIndexRight = str.find(')', nCurFrom + 1)
        if nFindIndex>=0 and nFindIndex<nFindIndexRight:
            nTotalSum+=1
            nCurFrom = nFindIndex
            continue
        elif nFindIndexRight>=0:
            nTotalSum-=1
            if nTotalSum==0:
                return nFindIndexRight
            else:
                nCurFrom = nFindIndexRight+1
                continue
        else:
            return -1



def eval_string_expression_strict(strExpression, strMode = '1', strActTime = '', pointData=[]):
    rvData = None

    containSystemAPIList = []

    try:
        pointNameList = find_vars_in_str(strExpression)

        strExpressionNew = strExpression

        #替换
        for api in g_systemAPIList:
            nFoundAPIStart = strExpressionNew.find(api + '(')
            if nFoundAPIStart >= 0:
                if api in containSystemAPIList:#已经有了，防止重复
                    continue
                containSystemAPIList.append(api)
                if strActTime:
                    nFoundAPIStartQuota = strExpressionNew.find('(', nFoundAPIStart)
                    nFoundAPIEnd = find_quota_end_string(strExpressionNew, nFoundAPIStartQuota)

                    strExpressionNew = strExpressionNew[:nFoundAPIEnd] + ','+ 'actTime="'+strActTime +'"'+ strExpressionNew[nFoundAPIEnd:]

        if containSystemAPIList:
            print('script is system api, no replace')
        elif strMode == '0':  # history mode
            strTimeStart = strActTime
            strTimeEnd = strActTime
            strTimeFormat = 'm1'
            result = BEOPDataAccess.getInstance().get_history_data_padded(pointNameList, strTimeStart, strTimeEnd,
                                                                          strTimeFormat)
            if result.get('error'):
                return json.dumps(dict(err=1, msg='no history data', data=-1), ensure_ascii=False)
            for k, v in result['map'].items():
                if isinstance(v, list) and v:
                    p_name = k
                    p_value = '(%s)' % (str(v[0]))
                    strVar = '<%%%s%%>' % (p_name)
                    strExpressionNew = strExpressionNew.replace(strVar, p_value)
        elif strMode=='2': #repair history
            pv = {}
            # pointData = BEOPDataAccess.getInstance().getInputTable(pointNameList)[0]
            for item in pointData:
                pv[item.get('name')] = item.get('value')
                strVar = '<%%%s%%>' % (item.get('name'))
                strExpressionNew = strExpressionNew.replace(strVar, item.get('value'))


        else:  # realtime mode
            pv = {}
            #pointData = BEOPDataAccess.getInstance().getInputTable(pointNameList)[0]
            for item in pointData:
                pv[item.get('name')] = item.get('value')
                strVar = '<%%%s%%>'%(item.get('name'))
                try:
                    strExpressionNew = strExpressionNew.replace(strVar, item.get('value'))
                except Exception as eReplace:
                    # 记录log
                    logging.error('ERROR in eval_string_expression_strict strExpressionNew.replace' + eReplace.__str__())

        try:
            if strExpressionNew.find("<%")>=0 and strExpressionNew.find('condition_last_minutes')<0:
                strError = 'ERROR in eval_string_expression_strict org(%s), point in equation lost' % (strExpression)
                print(strError)
                logging.error(strError)
                return None
            xxx = eval(strExpressionNew)
        except Exception as e:
            strError = 'ERROR in eval_string_expression_strict org(%s), eval(%s): %s'%(strExpression, strExpressionNew, e.__str__())
            print(strError)
            logging.error(strError)
            return None

        if isinstance(xxx, bool):
            rvData = 1 if xxx else 0
        else:
            rvData = xxx
        if app.config['MODE_HISTORY']:
            print('Calculation:%s ->%s' % (strExpression, rvData))
    except Exception as e:
        strError = 'ERROR in eval_string_expression_strict org(%s):%s'%(strExpression, e.__str__())
        traceback.print_exc()
        print(strError)
        logging.error(strError)
        rvData = None

    return rvData

def replace_string_expression_strict(strExpression, pointData=[]):
    try:
        strExpressionNew = strExpression

        pv = {}
        #pointData = BEOPDataAccess.getInstance().getInputTable(pointNameList)[0]
        for item in pointData:
            pv[item.get('name')] = item.get('value')
            strVar = '<%%%s%%>'%(item.get('name'))
            strExpressionNew = strExpressionNew.replace(strVar, item.get('value'))

    except Exception as e:
        print('ERROR in replace_string_expression_strict:%s'%(e.__str__()))

    return strExpressionNew

'''
condition_last_seconds("<%A%>>500", 15*60)
'''
def condition_last_minutes(strCondition, nMinutes,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    tFrom = tNow - timedelta(minutes=nMinutes)
    tFrom = tFrom.replace(second=0)
    tNow = tNow.replace(second=0)

    tAct = tFrom
    while tAct<tNow:
        strActTime = tAct.strftime('%Y-%m-%d %H:%M:00')
        rv = eval_string_expression_strict(strCondition, '0', strActTime)
        if rv is None:
            strError = 'ERROR in condition_last_minutes: eval_string_expression_strict return None, Expression: %s' % (strCondition)
            print(strError)
            logging.error(strError)
            return 0
        if rv==0:
            return 0
        tAct+= timedelta(minutes=1)

    return 1

'''
获取今天strPointName点在strOnOffPointName值>1的时段内的平均值
'''
def get_today_data_average_when_running(strPointName, strOnOffPointName, nDecimal=2, actTime=None):
    strTimeFormat = "%Y-%m-%d %H:%M:%S"
    try:
        tNow = datetime.now()
        tStart = tNow.replace(hour=0, minute=0, second=0)
        tTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")
        tTimeEnd = tNow.strftime("%Y-%m-%d %H:%M:00")

        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""

            tActTime = datetime.strptime(actTime, strTimeFormat)
            tTimeStart = tActTime.strftime("%Y-%m-%d 00:00:00")
            tTimeEnd = tActTime.strftime("%Y-%m-%d %H:%M:00")
            if is_0_oclock(actTime):
                tStart = tActTime - timedelta(days=1)
                tTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName,strOnOffPointName], tTimeStart, tTimeEnd, 'm1')

        if not allData:
            return ''

        if allData.get('map') is None:
            return ''

        dd = allData['map'].get(strPointName)
        if dd is None or not isinstance(dd, list):
            return ''

        if len(dd)==0:
            return ''

        ddStatus = allData['map'].get(strOnOffPointName)
        if ddStatus is None or not isinstance(ddStatus, list):
            return ''

        if len(ddStatus) == 0:
            return ''

        try:
            fSum = 0
            nCount = 0
            for iIndex in range(len(dd)):
                nOnOff = int(ddStatus[iIndex])
                if nOnOff>0:
                    fSum+= dd[iIndex]
                    nCount+=1

            if nCount>0:
                return round(fSum / nCount, nDecimal)
            else:
                return ''
        except:
            return ''
    except:
        return ''

"""
获取日期字符串
"""
def string_date_of_this_month(nDayIndex, strFormat='%Y/%m/%d', actTime=None):
    strTimeFormat = "%Y-%m-%d %H:%M:%S"
    try:
        tNow = datetime.now()
        tDay = tNow.replace(day=nDayIndex, hour=0, minute=0, second=0)

        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""

            tActTime = datetime.strptime(actTime, strTimeFormat)
            tDay = tActTime.replace(day=nDayIndex, hour=0, minute=0, second=0)

        try:
            return tDay.strftime(strFormat)
        except:
            return ""
    except:
        return ""

"""
获取日期字符串
"""
def string_date_of_last_month(nDayIndex, strFormat='%Y/%m/%d', actTime=None):
    strTimeFormat = "%Y-%m-%d %H:%M:%S"
    try:
        tNow = datetime.now()
        tDay = tNow.replace(day=1, hour=0, minute=0, second=0)

        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""

            tActTime = datetime.strptime(actTime, strTimeFormat)
            tDay = tActTime.replace(day=1, hour=0, minute=0, second=0)

        try:
            tDay = tDay - timedelta(days=1)
            tDay = tDay.replace(day=nDayIndex)
            return tDay.strftime(strFormat)
        except:
            return ""
    except:
        return ""

"""
获取今日从0点到当前时间strPointName点的最大值
"""
def get_today_data_max(strPointName, nDecimal=2, actTime=None):
    strTimeFormat = "%Y-%m-%d %H:%M:%S"
    try:
        tNow = datetime.now()
        tStart = tNow.replace(hour=0, minute=0, second=0)
        tTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")
        tTimeEnd = tNow.strftime("%Y-%m-%d %H:%M:00")

        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""

            tActTime = datetime.strptime(actTime, strTimeFormat)
            tTimeStart = tActTime.strftime("%Y-%m-%d 00:00:00")
            tTimeEnd = tActTime.strftime("%Y-%m-%d %H:%M:00")

            # 如果actTime是0点则往前推一天
            if is_0_oclock(actTime):
                tStart = tActTime - timedelta(days=1)
                tTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], tTimeStart, tTimeEnd, 'm1')

        if not allData:
            return ""

        if allData.get("map") is None:
            return ""

        dList = allData.get("map").get(strPointName)
        if dList is None or not isinstance(dList, list):
            return ""

        if not dList:
            return ""

        try:
            return round(max(dList), nDecimal)
        except:
            return ""
    except:
        return ""

"""
判断输入的时间是否为0点
"""
def is_0_oclock(strTime):
    res = False
    try:
        tTime = datetime.strptime(strTime, "%Y-%m-%d %H:%M:%S")
        if tTime.hour == 0 and tTime.minute == 0 and tTime.second == 0:
            res = True
    except:
        pass
    finally:
        return res

"""
获取今日从0点到当前时间strPointName点的最小值

特殊说明：如果是00:00:00，那么要特殊处理计算tStart必须是前一天的零点
"""
def get_today_data_min(strPointName, nDecimal=2, actTime=None):
    strTimeFormat = "%Y-%m-%d %H:%M:%S"
    try:
        tNow = datetime.now()
        tStart = tNow.replace(hour=0, minute=0, second=0)
        tTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")
        tTimeEnd = tNow.strftime("%Y-%m-%d %H:%M:00")

        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""

            tActTime = datetime.strptime(actTime, strTimeFormat)
            tTimeStart = tActTime.strftime("%Y-%m-%d 00:00:00")
            tTimeEnd = tActTime.strftime("%Y-%m-%d %H:%M:00")

            # 如果actTime是0点则往前推一天
            if is_0_oclock(actTime):
                tStart = tActTime - timedelta(days=1)
                tTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], tTimeStart, tTimeEnd, 'm1')

        if not allData:
            return ""

        if allData.get("map") is None:
            return ""

        dList = allData.get("map").get(strPointName)
        if dList is None or not isinstance(dList, list):
            return ""

        if not dList:
            return ""

        try:
            return round(min(dList), nDecimal)
        except:
            return ""
    except:
        return ""

"""
获取昨天strPointName点在strOnOffPointName值>0时的平均值
"""

def get_yesterday_data_average_when_running(strPointName, strOnOffPointName,  nDecimal=2,actTime=None):
    try:
        tNow = datetime.now()
        if actTime:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

        tYesterdayBegin = time_get_day_start(tNow-timedelta(days=1))
        tYesterdayEnd = time_get_day_end(tNow - timedelta(days=1))

        strTimeFromAllFormat = tYesterdayBegin.strftime('%Y-%m-%d 00:00:00')
        strTimeToAllFormat = tYesterdayEnd.strftime('%Y-%m-%d 23:59:59')

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName, strOnOffPointName],
                                                                       strTimeFromAllFormat, strTimeToAllFormat, 'm1')
        if not allData:
            return ''

        if allData.get('map') is None:
            return ''

        dd = allData['map'].get(strPointName)
        if dd is None or not isinstance(dd, list):
            return ''

        if len(dd) == 0:
            return ''

        ddStatus = allData['map'].get(strOnOffPointName)
        if ddStatus is None or not isinstance(ddStatus, list):
            return ''

        if len(ddStatus) == 0:
            return ''

        try:
            fSum = 0
            nCount = 0
            for iIndex in range(len(dd)):
                nOnOff = int(ddStatus[iIndex])
                if nOnOff > 0:
                    fSum += dd[iIndex]
                    nCount += 1

            if nCount > 0:
                return round(fSum / nCount,nDecimal)
            else:
                return ''
        except:
            return ''
    except:
        return ''

'''
todo: 
获取本月某日（nDay, 1-31） strPointName点在strOnOffPointName点值>0时段的平均值
'''
def get_day_average_this_month_when_running(strPointName, strOnOffPointName, nDay, nDecimal=2, actTime=None):
    try:
        tNow = datetime.now()
        if actTime:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

        tTheDay = tNow.replace(day=nDay)
        tTheDayBegin = time_get_day_start(tTheDay)
        tTheDayEnd = time_get_day_end(tTheDay)

        strTimeFromFormat = tTheDayBegin.strftime('%Y-%m-%d 00:00:00')
        strTimeToFormat = tTheDayEnd.strftime('%Y-%m-%d 23:59:59')

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName, strOnOffPointName],
                                                                       strTimeFromFormat, strTimeToFormat, 'm1')
        if not allData:
            return ''

        if allData.get('map') is None:
            return ''

        dd = allData['map'].get(strPointName)
        if dd is None or not isinstance(dd, list):
            return ''

        if len(dd) == 0:
            return ''

        ddStatus = allData['map'].get(strOnOffPointName)
        if ddStatus is None or not isinstance(ddStatus, list):
            return ''

        if len(ddStatus) == 0:
            return ''

        try:
            fSum = 0
            nCount = 0
            for iIndex in range(len(dd)):
                nOnOff = int(ddStatus[iIndex])
                if nOnOff > 0:
                    fSum += dd[iIndex]
                    nCount += 1

            if nCount > 0:
                return round(fSum / nCount, nDecimal)
            else:
                return ''
        except:
            return ''
    except:
        return ''

'''
todo: 
获取上月某日（nDay, 1-31） strPointName点在strOnOffPointName点值>0时段的平均值
'''
def get_day_average_last_month_when_running(strPointName, strOnOffPointName, nDay,  nDecimal=2,actTime=None):
    try:
        tNow = datetime.now()
        if actTime:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

        firstDayThisMonth = tNow.replace(day=1)
        tLastMonth = firstDayThisMonth - timedelta(days=1)
        tTheDay = tLastMonth.replace(day=nDay)
        tTheDayBegin = time_get_day_start(tTheDay)
        tTheDayEnd = time_get_day_end(tTheDay)

        strTimeFromFormat = tTheDayBegin.strftime('%Y-%m-%d 00:00:00')
        strTimeToFormat = tTheDayEnd.strftime('%Y-%m-%d 23:59:59')

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName, strOnOffPointName],
                                                                       strTimeFromFormat, strTimeToFormat, 'm1')
        if not allData:
            return ''

        if allData.get('map') is None:
            return ''

        dd = allData['map'].get(strPointName)
        if dd is None or not isinstance(dd, list):
            return ''

        if len(dd) == 0:
            return ''

        ddStatus = allData['map'].get(strOnOffPointName)
        if ddStatus is None or not isinstance(ddStatus, list):
            return ''

        if len(ddStatus) == 0:
            return ''

        try:
            fSum = 0
            nCount = 0
            for iIndex in range(len(dd)):
                nOnOff = int(ddStatus[iIndex])
                if nOnOff > 0:
                    fSum += dd[iIndex]
                    nCount += 1

            if nCount > 0:
                return round(fSum / nCount, nDecimal)
            else:
                return ''
        except:
            return ''
    except:
        return ''


"""
获取一段时间内strPointName的差值：strTimeEnd时刻的值减去strTimeStart时刻的值，时间格式：YYYY-MM-DD HH:MM:SS
"""
def get_data_delta_of_period(strPointName, strTimeStart, strTimeEnd):
    try:
        strFormat = "%Y-%m-%d %H:%M:%S"

        if not is_valid_date(strTimeStart, strFormat):
            return ""

        if not is_valid_date(strTimeEnd, strFormat):
            return ""

        tTimeFrom1 = datetime.strptime(strTimeStart, strFormat)
        tTimeTo1 = tTimeFrom1 + timedelta(minutes=1)

        tTimeFrom2 = datetime.strptime(strTimeEnd, strFormat)
        tTimeTo2 = tTimeFrom2 + timedelta(minutes=1)

        allDataStart = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], tTimeFrom1.strftime(strFormat), tTimeTo1.strftime(strFormat), 'm1')

        allDataEnd = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], tTimeFrom2.strftime(strFormat), tTimeTo2.strftime(strFormat), 'm1')

        if allDataStart.get("map") is None:
            return ""
        if allDataEnd.get("map") is None:
            return ""

        dListStart =  allDataStart.get("map").get(strPointName)
        dListEnd =  allDataEnd.get("map").get(strPointName)

        if dListStart is None or not isinstance(dListStart, list) or dListEnd is None or not isinstance(dListEnd, list):
            return ""

        if not len(dListStart) or not len(dListEnd):
            return ""

        return dListEnd[0] - dListStart[0]

    except:
        return ""

def is_valid_date(strTime, strFormat):
    try:
        time.strptime(strTime, strFormat)
        return True
    except:
        return False


def get_data_is_not_changing(strPointName,nMinutes=30,actTime=None):
    try:
        if actTime:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        else:
            tNow = datetime.now()
        tStart = tNow - timedelta(minutes= nMinutes)
        tTimeStart = tStart.strftime("%Y-%m-%d %H:%M:00")
        tTimeEnd = tNow.strftime("%Y-%m-%d %H:%M:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], tTimeStart, tTimeEnd, 'm1')

        if not allData:
            return 0

        if allData.get("map") is None:
            return 0

        dList = allData.get("map").get(strPointName)
        if dList is None or not isinstance(dList, list):
            return 0

        if not dList:
            return 0

        try:
            for iIndex in range(len(dList)-1):
                fv1 = float(dList[iIndex])
                fv2 = float(dList[iIndex+1])
                if abs(fv1-fv2)>1e-3:
                    return 0

            return 1
        except:
            return  0
    except:
        return 0

    return 0


'''
todo: 
获取本月某日（nDay, 1-31） strPointName点在strOnOffPointName点值>0时段的平均值
'''
def get_moving_average(strPointName, nMinutes,  nDecimal=2, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    tBegin = tNow - timedelta(minutes= nMinutes)

    strTimeFromFormat = tBegin.strftime('%Y-%m-%d %H:%M:00')
    strTimeToFormat = tNow.strftime('%Y-%m-%d %H:%M:00')

    allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName],
                                                                   strTimeFromFormat, strTimeToFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return ''


    try:
        fSum = 0
        nCount = 0
        for iIndex in range(len(dd)):
            fSum += dd[iIndex]
            nCount += 1

        if nCount > 0:
            return round(fSum / nCount, nDecimal)
        else:
            return ''
    except:
        return ''

    return ''


def get_moving_average_when_running(strPointName, strOnOffPointName, nMinutes, nDecimal=2, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    tBegin = tNow - timedelta(minutes= nMinutes)

    strTimeFromFormat = tBegin.strftime('%Y-%m-%d %H:%M:00')
    strTimeToFormat = tNow.strftime('%Y-%m-%d %H:%M:00')

    allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName, strOnOffPointName],
                                                                   strTimeFromFormat, strTimeToFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)
    if dd is None or not isinstance(dd, list):
        return ''

    if len(dd) == 0:
        return ''

    ddStatus = allData['map'].get(strOnOffPointName)
    if ddStatus is None or not isinstance(ddStatus, list):
        return ''

    if len(ddStatus) == 0:
        return ''

    try:
        fSum = 0
        nCount = 0
        for iIndex in range(len(dd)):
            nOnOff = int(ddStatus[iIndex])
            if nOnOff > 0:
                fSum += dd[iIndex]
                nCount += 1

        if nCount > 0:
            return round(fSum / nCount, nDecimal)
        else:
            return ''
    except:
        return ''

    return ''


def get_point_value_string(strPointName, actTime=None):
    pv = BEOPDataAccess.getInstance().getInputTable([strPointName])
    try:
        return pv[0][0]['value']
    except:
        return None



def get_thisweek_data_delta(strPointName, nDecimal=2,actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    if tNow.hour==0 and tNow.minute==0 and tNow.weekday()==0:#0:00要算昨日的一天差（零点到零点）
        tBegin = tNow - timedelta(days=7)
        strTimeFromAllFormat = '%s 00:00:00' % (tBegin.strftime('%Y-%m-%d'))
        strTimeToAllFormat = tNow.strftime('%Y-%m-%d 00:00:00')
    else:
        tBegin = tNow
        while tBegin.weekday()!=0:
            tBegin = tBegin-  timedelta(days=1)

        strTimeFromAllFormat = '%s 00:00:00'%(tBegin.strftime('%Y-%m-%d'))
        strTimeToAllFormat =  tNow.strftime('%Y-%m-%d %H:%M:%S')

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFromAllFormat, strTimeToAllFormat, 'm1')
    if not allData:
        return ''

    if allData.get('map') is None:
        return ''

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return ''


    if len(dd)==0:
        return 0

    try:
        ddArry = numpy.array(dd)

    except:
        return ''


    return  round(ddArry[-1] - ddArry[0], nDecimal)



def is_today_chinese_holiday(actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()



    calendarInfo = BEOPDataAccess.getInstance().getCalendarOfDay(tNow)
    if not calendarInfo:
        return ''

    try:
        calendarInfo = json.loads(calendarInfo, encoding='UTF-8')
        strHoliday = calendarInfo.get('result').get('data').get('holiday')
        if strHoliday:
            if strHoliday in ['元旦', '春节','清明节', '劳动节','端午节', '中秋节', '国庆节']: #防止将植树节等非假日类节日视作为节假日
                return 1
    except:
        return ''


    return 0

def point_exists(strPointName):
    res = False
    try:
        rvConduct = BEOPDataAccess.getInstance().PointExists(strPointName)
        res = rvConduct
    except:
        pass
    finally:
        return res

def run_num_count(pointNameList):
    """
    run_num_count(["OnOff01", "OnOff02"])
    计算当前点名列表（运行状态）点值之和
    """
    result = BEOPDataAccess.getInstance().getInputTable(pointNameList)
    if not result:
        return ""
    if not len(result):
        return ""
    if not len(result[0]):
        return ""

    count = 0
    for item in result[0]:
        try:
            value = int(item.get("value"))
            count += value
        except:
            pass

    return count

def get_current_active_mode_name(nType):
    if not isinstance(nType, int):
        return ""
    result = BEOPDataAccess.getInstance().getCurrentActiveModeName(nType)
    return result

"""
输入：
nTarget - 目标值
[(strPointNamePrefix01, nIndexFrom01, nIndexTo01), (strPointNamePrefix02, nIndexFrom02, nIndexTo02)...] - 点信息列表

返回：
满足率
"""
def stats_running_count(nTarget, infoList):
    if not isinstance(nTarget, int):
        return -1

    if not isinstance(infoList, list):
        return -1

    # 生成点名列表
    pointNameList = []
    for info in infoList:
        if not isinstance(info, tuple):
            return -1

        idx = info[1]
        while idx <= info[2]:
            pointNameList.append("%s%02d" % (info[0], idx))
            idx += 1

    # 取当前值
    result = BEOPDataAccess.getInstance().getInputTable(pointNameList)
    dictList = None
    try:
        dictList = result[0]
    except:
        return -1

    if not dictList or not len(dictList):
        return -1

    # 满足个数统计
    nAchieve = 0
    for dict in dictList:
        try:
            value = int(dict.get("value"))
        except:
            continue
        if value == nTarget:
            nAchieve += 1

    if not len(dictList):
        return -1

    return round(nAchieve/len(dictList), 3)

"""
点名在当前模式中：返回1（若点值为1），返回0（若点值为0）
点名不在当前模式中：返回 -1
"""
def is_point_now_on_in_schedule(strPointName):
    if not isinstance(strPointName, str):
        return ""
    dInfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db([strPointName])
    if not dInfo:
        return -1
    value = BEOPDataAccess.getInstance().pointExistsInSchedule(strPointName)
    if value is None:
        return -1
    if int(value) == 1:
        return 1
    return 0

"""
点名在当前模式中：返回1（若点值为0），返回0（若点值为1）
点名不在当前模式中：返回 -1
"""
def is_point_now_off_in_schedule(strPointName):
    if not isinstance(strPointName, str):
        return ""
    dInfo = BEOPSqliteAccess.getInstance().getPointInfoFromS3db([strPointName])
    if not dInfo:
        return -1
    value = BEOPDataAccess.getInstance().pointExistsInSchedule(strPointName)
    if value is None:
        return -1
    if int(value) == 0:
        return 1
    return 0

def random_between_start_stop(nStart, nStop, decimal):
    if not type(nStart) in [int, float] or not type(nStop) in [int, float]:
        return ""
    if not isinstance(decimal, int):
        return ""
    value = random.uniform(nStart, nStop)
    return round(value, decimal)

# [(strPointNamePrefix01, nIndexFrom01, nIndexTo01), (strPointNamePrefix02, nIndexFrom02, nIndexTo02)...] - 点信息列表
# decimal - 小数位
def get_avg_value(pointStruList, decimal):
    if not isinstance(pointStruList, list):
        return ""
    if not isinstance(decimal, int):
        return ""

    # 生成点名列表
    pointNameList = []
    for info in pointStruList:
        if not isinstance(info, tuple):
            return ""

        idx = info[1]
        while idx <= info[2]:
            pointNameList.append("%s%02d" % (info[0], idx))
            idx += 1

    result = BEOPDataAccess.getInstance().getInputTable(pointNameList)
    dictList = None
    try:
        dictList = result[0]
    except:
        return ""

    if not dictList or not len(dictList):
        return ""

    nTotal = 0
    nCount = 0
    for item in dictList:
        nTotal += float(item.get("value"))
        nCount += 1

    if nCount == 0:
        return ""

    return round(nTotal/nCount, decimal)

"""
0- YYYY-MM-DD 1- Y年m月d日  2- yyyy/mm/dd
"""
def get_date_string_today(nFormat, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    if nFormat==0:
        return tNow.strftime('%Y-%m-%d')
    elif nFormat==1:
        return tNow.strftime('%Y年%m月%d日')
    elif nFormat==2:
        return tNow.strftime('%d/%m/%Y')
    else:
        return tNow.strftime('%Y-%m-%d')

"""
0- YYYY-MM-DD 1- Y年m月d日  2- yyyy/mm/dd
"""
def get_date_string_yesterday(nFormat, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    tYesterday = tNow - timedelta(days=1)
    if nFormat==0:
        return tYesterday.strftime('%Y-%m-%d')
    elif nFormat==1:
        return tYesterday.strftime('%Y年%m月%d日')
    elif nFormat==2:
        return tYesterday.strftime('%d/%m/%Y')
    else:
        return tYesterday.strftime('%Y-%m-%d')

"""
获取运行中设备的某个参数的平均值
input: 
strEquipPrefix: 点名中设备前缀，如PriChWP
strPointAttr: 点名中点位属性，如 VSDFreq
equipNoList： 设备号列表，如[1,2,3,4]
"""
def get_avg_value_running(strEquipPrefix, strPointAttr, equipNoList, decimal=2, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()
    try:
        if not isinstance(strEquipPrefix, str):
            return ""
        if not isinstance(strPointAttr, str):
            return ""
        if not isinstance(equipNoList, list):
            return ""
        tarPointList = ["%s%s%02d" % (strEquipPrefix, strPointAttr, equipNo) for equipNo in equipNoList]
        onoffPointList = ["%sOnOff%02d" % (strEquipPrefix, equipNo) for equipNo in equipNoList]

        arr = BEOPDataAccess.getInstance().getInputTable(tarPointList + onoffPointList)
        if not len(arr):
            return ""

        dataDict = {}
        for item in arr[0]:
            dataDict.update({
                item.get("name"): float(item.get("value"))
            })

        tarTotal = 0
        count = 0
        for idx, point in enumerate(onoffPointList):
            if dataDict.get(point) > 0:
                tarTotal += dataDict.get(tarPointList[idx])
                count += 1

        if count > 0:
            return round(tarTotal/count, decimal)
        else:
            return 0

    except:
        return ""

# strPointNameTimeStart, strPointNameTimeEnd: format: %H:%S
def get_hours_delta(strPointNameTimeStart, strPointNameTimeEnd):
    try:
        rtList = BEOPDataAccess.getInstance().getInputTable([strPointNameTimeStart, strPointNameTimeEnd])
        strTimeStart = datetime.now().strftime("%Y-%m-%d ") + rtList[1].get(strPointNameTimeStart) + ":00"
        strTimeEnd = datetime.now().strftime("%Y-%m-%d ") + rtList[1].get(strPointNameTimeEnd) + ":00"
        seconds = (datetime.strptime(strTimeEnd, "%Y-%m-%d %H:%M:%S") - datetime.strptime(strTimeStart, "%Y-%m-%d %H:%M:%S")).total_seconds()
        hours = seconds / 3600
        return round(hours, 2)
    except Exception as e:
        logging.error("ERROR in get_hours_delta: %s" % e.__str__())
        return ""

def get_work_days_of_this_month(month=None):
    try:
        if not month:
            month = datetime.now().month
        tFrom = time_get_first_day_of_one_month_this_year(month, datetime.now())
        curMonth = tFrom.month
        nextMonth = time_get_first_day_of_next_month(tFrom).month

        workday = 0
        while curMonth != nextMonth:
            if tFrom.weekday() < 5:
                workday += 1
            tFrom += timedelta(days=1)
            curMonth = tFrom.month
        return workday
    except Exception as e:
        logging.error("ERROR in get_work_days_of_this_month: %s" % e.__str__())
        return ""

def get_days_of_this_month(month=None):
    try:
        if not month:
            month = datetime.now().month

        tFrom = time_get_first_day_of_one_month_this_year(month, datetime.now())
        curMonth = tFrom.month
        nextMonth = time_get_first_day_of_next_month(tFrom).month

        days = 0
        while curMonth != nextMonth:
            days += 1
            tFrom += timedelta(days=1)
            curMonth = tFrom.month
        return days

    except Exception as e:
        logging.error("ERROR in get_days_of_this_month: %s" % e.__str__())
        return ""




'''
nMode: 1-realtime
0-history

2020-12-28: russell,golding 配对编程，取消了这个函数的使用，认为该函数有内部机制不协调于现有的计算脚本。
统一替换为 eval_string_expression_strict 来进行计算
'''
def tool_eval_string_v2(strExpressionList, allPointRealtimeDataLMap, nMode, strActTime = None):
    rvDataList= []
    if strActTime is None:
        strActTime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    if allPointRealtimeDataLMap is None:
        return json.dumps(dict(err=1, msg='allPointRealtimeDataList param lost', data=-1), ensure_ascii=False)

    try:
        for strExpression in strExpressionList:
            strExpressionModify = strExpression.replace('<%', '')
            strExpressionModify = strExpressionModify.replace('%>', '')

            try:
                p = re.compile(r'[a-zA-Z]+[a-zA-Z0-9_]+', re.IGNORECASE)

                pointNameList = p.findall(strExpressionModify)
                for pitem in pointNameList:
                    if pitem in ['and', 'or', 'is' 'None','True', 'False', 'sqrt', 'pow','exp', 'log', 'log10', 'ceil', 'floor',
                                'fabs' ,'if', 'else' ]:
                        continue
                    else:
                        strVV = allPointRealtimeDataLMap.get(pitem)
                        if strVV is not None:
                            strExpressionModify = strExpressionModify.replace(pitem, strVV)

                xxx = eval(strExpressionModify)
                if isinstance(xxx, bool):
                    rvData = 1 if xxx else 0
                else:
                    rvData = xxx
            except Exception as e2:
                print('ERROR in eval expression:%s, exception:%s' %(strExpression, e2.__str__()))
                rvData = None
            if rvData is not None:
                rvDataList.append(dict(name=strExpression, value=rvData))


    except Exception as e:
        strError  = 'ERROR in tool_eval_string :%s, expression:%s'%(e.__str__(), strExpression)
        logging.error(strError)
        return []

    return rvDataList

def get_points_from_script_list(scriptList):
    try:
        pointList = []
        match = re.compile(r'[a-zA-Z]+[a-zA-Z0-9_]+', re.IGNORECASE)
        for script in scriptList:
            pointFromThisScript = []
            scriptProcessed = script.replace("<%", "")
            scriptProcessed = scriptProcessed.replace("%>", "")
            itemList = match.findall(scriptProcessed)
            for item in itemList:
                try:
                    if item in ['and', 'or', 'is' 'None','True', 'False', 'sqrt', 'pow','exp', 'log', 'log10', 'ceil', 'floor',
                                        'fabs' ,'if', 'else' ]:
                        continue
                    if item not in pointFromThisScript:
                        pointFromThisScript.append(item)
                except Exception as e:
                    strLog = "ERROR in filtering points from get_points_from_script_list: %s" % e.__str__()
                    logging.error(strLog)

            pointList.append(dict(script=script, pointList=pointFromThisScript))

        return pointList
    except Exception as e:
        strLog = "ERROR in get_points_from_script_list: %s" % e.__str__()
        logging.error(strLog)
        return []

def get_history_data_of_scripts(historyData, pointsFromScriptList):
    rstDict = {}
    try:
        for item in pointsFromScriptList:
            script = item.get("script")
            ptList = item.get("pointList")

            if not len(ptList):
                continue

            scriptPro = script.replace("<%", "")
            scriptPro = scriptPro.replace("%>", "")
            rstDataList = []
            for i, object in enumerate(historyData.get(ptList[0])):
                for j, pt in enumerate(ptList):
                    strTarget = str(historyData.get(ptList[j])[i])
                    scriptPro = scriptPro.replace(pt, strTarget)
                value = eval(scriptPro)
                if isinstance(value, bool):
                    rstData = 1 if value else 0
                else:
                    rstData = value
                rstDataList.append(rstData)

            rstDict.update({script: rstDataList})

        return rstDict
    except Exception as e:
        strLog = "ERROR in get_history_data_of_scripts: %s" % e.__str__()
        logging.error(strLog)
        return {}

"""
计算这个点在 nCheckMinutes 分钟内是否会有日程指令给1
"""
def is_point_will_on_in_schedule(strPointName, nCheckMinutes=30):
    dInfo = BEOPDataAccess.getInstance().getPointLatestTriggerTime(strPointName)
    if dInfo.get("code") > 0:
        return 0

    if not len(dInfo.get("data")):
        return 0

    timeValue = dInfo.get("data")

    valueTar = None
    if (timeValue[0] - datetime.now()).total_seconds() >= 0 and (timeValue[0] - datetime.now()).total_seconds() <= nCheckMinutes * 60:
        valueTar = timeValue[1]

    if isinstance(valueTar, str):
        if valueTar.isdigit():
            if int(float(valueTar)) == 1:
                return 1
    return 0

"""
计算这个点在 nCheckMinutes 分钟内是否会有日程指令给0
"""
def is_point_will_off_in_schedule(strPointName, nCheckMinutes=30):
    dInfo = BEOPDataAccess.getInstance().getPointLatestTriggerTime(strPointName)
    if dInfo.get("code") > 0:
        return 0

    if not len(dInfo.get("data")):
        return 0

    timeValue = dInfo.get("data")

    valueTar = None
    if (timeValue[0] - datetime.now()).total_seconds() >= 0 and (timeValue[0] - datetime.now()).total_seconds() <= nCheckMinutes * 60:
        valueTar = timeValue[1]

    if isinstance(valueTar, str):
        if valueTar.isdigit():
            if int(float(valueTar)) == 0:
                return 1
    return 0


def get_thismonth_data_max(strPointName, nDecimal=2, actTime=None):
    try:
        tNow = datetime.now()
        strTimeFormat = "%Y-%m-%d %H:%M:%S"
        tStart = tNow.replace(day=1, hour=0, minute=0, second=0)
        strTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")
        strTimeEnd = tNow.strftime("%Y-%m-%d %H:%M:00")
        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""
            tActTime = datetime.strptime(actTime, strTimeFormat)
            tTimeStart = tActTime.replace(day=1, hour=0, minute=0, second=0)
            strTimeStart = tTimeStart.strftime("%Y-%m-%d %H:%M:%S")
            strTimeEnd = tActTime.strftime("%Y-%m-%d %H:%M:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], strTimeStart, strTimeEnd, 'h1')

        if not allData:
            return ""

        if allData.get("map") is None:
            return ""

        dList = allData.get("map").get(strPointName)
        if dList is None or not isinstance(dList, list):
            return ""

        if not dList:
            return ""

        return round(max(dList), nDecimal)
    except Exception as e:
        logging.error("ERROR in get_thismonth_data_max: %s" % e.__str__())
        return ""


def get_thismonth_data_min(strPointName, nDecimal=2, actTime=None):
    try:
        tNow = datetime.now()
        strTimeFormat = "%Y-%m-%d %H:%M:%S"
        tStart = tNow.replace(day=1, hour=0, minute=0, second=0)
        strTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")
        strTimeEnd = tNow.strftime("%Y-%m-%d %H:%M:00")
        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""
            tActTime = datetime.strptime(actTime, strTimeFormat)
            tTimeStart = tActTime.replace(day=1, hour=0, minute=0, second=0)
            strTimeStart = tTimeStart.strftime("%Y-%m-%d %H:%M:%S")
            strTimeEnd = tActTime.strftime("%Y-%m-%d %H:%M:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], strTimeStart, strTimeEnd, 'h1')

        if not allData:
            return ""

        if allData.get("map") is None:
            return ""

        dList = allData.get("map").get(strPointName)
        if dList is None or not isinstance(dList, list):
            return ""

        if not dList:
            return ""

        return round(min(dList), nDecimal)
    except Exception as e:
        logging.error("ERROR in get_thismonth_data_min: %s" % e.__str__())
        return ""

def get_thisyear_data_avg(strPointName, nDecimal=2, actTime=None):
    try:
        tNow = datetime.now()
        strTimeFormat = "%Y-%m-%d %H:%M:%S"
        tStart = tNow.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
        strTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")
        strTimeEnd = tNow.strftime("%Y-%m-%d 00:00:00")
        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""
            tActTime = datetime.strptime(actTime, strTimeFormat)
            tTimeStart = tActTime.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
            strTimeStart = tTimeStart.strftime("%Y-%m-%d %H:%M:%S")
            strTimeEnd = tActTime.strftime("%Y-%m-%d 00:00:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], strTimeStart, strTimeEnd, 'd1')

        if not allData:
            return ""

        if allData.get("map") is None:
            return ""

        dList = allData.get("map").get(strPointName)
        if dList is None or not isinstance(dList, list):
            return ""

        if not dList:
            return ""

        fMean = numpy.mean(dList)
        if isinstance(fMean, float) or isinstance(fMean, int):
            return round(fMean, nDecimal)
        else:
            return ""

    except Exception as e:
        logging.error("ERROR in get_thisyear_data_avg: %s" % e.__str__())
        return ""


def get_thismonth_data_avg(strPointName, nDecimal=2, actTime=None):
    try:
        tNow = datetime.now()
        strTimeFormat = "%Y-%m-%d %H:%M:%S"
        tStart = tNow.replace(day=1, hour=0, minute=0, second=0)
        strTimeStart = tStart.strftime("%Y-%m-%d 00:00:00")
        strTimeEnd = tNow.strftime("%Y-%m-%d %H:%M:00")
        if actTime:
            if not is_valid_date(actTime, strTimeFormat):
                return ""
            tActTime = datetime.strptime(actTime, strTimeFormat)
            tTimeStart = tActTime.replace(day=1, hour=0, minute=0, second=0)
            strTimeStart = tTimeStart.strftime("%Y-%m-%d %H:%M:%S")
            strTimeEnd = tActTime.strftime("%Y-%m-%d %H:%M:00")

        allData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], strTimeStart, strTimeEnd, 'h1')

        if not allData:
            return ""

        if allData.get("map") is None:
            return ""

        dList = allData.get("map").get(strPointName)
        if dList is None or not isinstance(dList, list):
            return ""

        if not dList:
            return ""

        return round(numpy.mean(dList), nDecimal)
    except Exception as e:
        logging.error("ERROR in get_thismonth_data_avg: %s" % e.__str__())
        return ""


def get_thismonth_daily_data_delta_at_hour(strPointName, nDay=None, nHour=None, nDecimal=2, actTime=None):
    try:
        if actTime:
            if not is_valid_date(actTime, '%Y-%m-%d %H:%M:%S'):
                return ""
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        else:
            tNow = datetime.now()

        if nHour is not None:
            if not isinstance(nHour, int):
                return ""
            if nHour < 0 or nHour > 23:
                return ""

        if nDay is not None:
            if not isinstance(nDay, int):
                return ""
            if nDay < 1 or nDay > calendar.monthrange(tNow.year, tNow.month)[1]:
                return ""

        day = tNow.day
        hour = tNow.hour
        if nDay:
            day = nDay
        if nHour:
            hour = nHour

        tFrom = tNow.replace(day=day, hour=0, minute=0, second=0, microsecond=0)
        tTo = tNow.replace(day=day, hour=hour, minute=0, second=0, microsecond=0)

        strFrom = tFrom.strftime("%Y-%m-%d %H:%M:%S")
        strTo = tTo.strftime("%Y-%m-%d %H:%M:%S")

        allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strFrom, strTo, 'm1')
        if not allData:
            return ''

        if allData.get('map') is None:
            return ''

        dd = allData['map'].get(strPointName)

        if dd is None or not isinstance(dd, list):
            return ''

        if len(dd)==0:
            return 0

        try:
            ddArry = numpy.array(dd)
        except:
            return ''
        return round(ddArry[-1] - ddArry[0], nDecimal)
    except Exception as e:
        logging.error("ERROR in get_today_data_delta_at_hour: %s" % e.__str__())
        return ""


def get_this_month_data_sum(strPointName, nFromDay=2, nDecimal=2, actTime=None):
    try:
        if actTime:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        else:
            tNow = datetime.now()

        realtimeData = BEOPDataAccess.getInstance().getInputTable([strPointName])[1].get(strPointName)

        if tNow.day == nFromDay:
            return realtimeData

        plusRealtimeDataOrNot = False
        if tNow.day > nFromDay:
            strTimeFrom = tNow.replace(day=nFromDay, hour=0, minute=0, second=0).strftime("%Y-%m-%d %H:%M:%S")
            strTimeTo = tNow.replace(hour=0, minute=0, second=0).strftime("%Y-%m-%d %H:%M:%S")
            plusRealtimeDataOrNot = True

        else:
            strTimeFrom = (tNow.replace(day=1) - timedelta(days=2)).replace(day=nFromDay, hour=0, minute=0, second=0).strftime("%Y-%m-%d %H:%M:%S")
            strTimeTo = (tNow + timedelta(days=1)).replace(hour=0, minute=0, second=0).strftime("%Y-%m-%d %H:%M:%S")

        dHisData = BEOPDataAccess.getInstance().get_history_data_padded([strPointName], strTimeFrom, strTimeTo, 'd1')

        dataList = dHisData.get("map").get(strPointName)

        res = sum(dataList) + float(realtimeData) if plusRealtimeDataOrNot else sum(dataList)
        return round(res, nDecimal)

    except Exception as e:
        logging.error("ERROR in get_this_month_data_sum: %s" % e.__str__())
        return ""


def get_date_string_this_month_of_day(nDay=1, actTime=None):
    try:
        if actTime:
            tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
        else:
            tNow = datetime.now()

        if tNow.day==1:#如果是月初第一天，那么本月this month其实是上月
            tNow = tNow - timedelta(days=2)
        tDate = tNow.replace(day=nDay)
        return tDate.strftime("%Y-%m-%d")
    except Exception as e:
        logging.error("ERROR in get_date_string_this_month_of_day: %s" % e.__str__())
        return ""

"""
strTime格式： HH:MM   例 09:00
"""
def get_data_this_month_at_day_time(strPointName, nDay, strTime, reportCall=False, actTime=None):
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')
    else:
        tNow = datetime.now()

    if not is_valid_date(strTime, "%H:%M"):
        return ""

    if not isinstance(nDay, int):
        return ""

    tTime = datetime.strptime(strTime, "%H:%M")

    strKey = "HIS__get_data_this_month_at_day_time_{point}".format(point=strPointName)

    if RedisManager.is_alive():
        resHisData = RedisManager.get(strKey)
        if resHisData != None:
            strTimeKey = "{year},{month:0>2d},{day:0>2d},{time}".format(year=tNow.year, month=tNow.month, day=nDay, time=strTime)
            value = resHisData.get(strTimeKey, None)
            if is_digit(value):
                return round(float(value), 2)

    if tNow.day == 1 and tNow.replace(hour=0, minute=0) <= tNow <= tNow.replace(hour=6, minute=0):
        nYearLastMonth = (tNow.replace(day=1) - timedelta(days=1)).year
        nMonthLastMonth = (tNow.replace(day=1) - timedelta(days=1)).month
        tNow = tNow.replace(year=nYearLastMonth, month=nMonthLastMonth)

    # 目标时间
    tTimeFrom = tNow.replace(day=nDay, hour=tTime.hour, minute=tTime.minute, second=0, microsecond=0)

    if RedisManager.is_alive() and tTimeFrom.minute == 0 and tTimeFrom.second == 0:
        strDataValue1 = RedisManager.get_history_data(strPointName, tTimeFrom)

        if reportCall:
            strLogInfo = "data from redis: {data}(pointName:{point})".format(data=strDataValue1, point=strPointName)
            api_log_info("get_data_this_month_at_day_time{time}.log".format(time=datetime.now().strftime("%Y-%m-%d")),
                         strLogInfo)

        # 2022-02-7 (russell) 由于上汽问题，redis中拿不到值后去mysql再查一次
        if strDataValue1 != None and strDataValue1 != "" and is_digit(strDataValue1):
            fDelta = float(strDataValue1)

            if datetime.now() > tTimeFrom:
                RedisManager.set_api_point_data_at_time("get_data_this_month_at_day_time", strPointName, tNow.year, tNow.month, nDay, strTime, fDelta)

            return round(fDelta, 2)

    strTimeFrom = tTimeFrom.strftime("%Y-%m-%d %H:%M:%S")
    strTimeTo = strTimeFrom

    debugLogFile = "get_data_this_month_at_day_time_{date}.log".format(date=datetime.now().strftime("%Y-%m-%d"))
    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTimeFrom, strTimeTo, "m1", debugLogFile=debugLogFile)

    if reportCall:
        strLogInfo = "data from mysql: {data}(pointName:{point})".format(
            data=allData.get("map", {}).get(strPointName, []), point=strPointName)
        api_log_info("get_data_this_month_at_day_time{time}.log".format(time=datetime.now().strftime("%Y-%m-%d")),
                     strLogInfo)

    if not allData:
        return ""

    if allData.get('map') is None:
        return ""

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return None

    if not len(dd):
        return None

    targetValueMysql = round(dd[-1], 2)

    if datetime.now() > tTimeFrom:
        RedisManager.set_api_point_data_at_time("get_data_this_month_at_day_time", strPointName,tNow.year, tNow.month, nDay, strTime, targetValueMysql)

    try:
        return targetValueMysql
    except:
        return ""

def get_data_thisyear_at_month_day_time(strPointName, nMonth=1, nDay=1, strTime="00:00", actTime=None):
    tNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

    if nMonth not in [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]:
        return ""

    if not isinstance(nDay, int) or not (1 <= nDay <= 31):
        return ""

    tTime = datetime.strptime(strTime, "%H:%M")

    tBegin = tNow.replace(month=nMonth, day=nDay, hour=tTime.hour, minute=tTime.minute, second=0, microsecond=0)
    strBegin = tBegin.strftime("%Y-%m-%d %H:%M:%S")
    strEnd = strBegin

    allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strBegin, strEnd, "m1")

    if not allData:
        return ""

    if allData.get('map') is None:
        return ""

    dd = allData['map'].get(strPointName)

    if dd is None or not isinstance(dd, list):
        return None

    if not len(dd):
        return None

    try:
        return round(dd[-1], 2)
    except:
        return ""


def get_thisyear_delta_from_month_to_now(strPointName, nMonth=1, nDecimal=2, actTime=None):
    tNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

    if nMonth not in [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]:
        return ""

    tFrom = tNow.replace(month=nMonth, day=1, hour=0, minute=0, second=0, microsecond=0)
    tNextMonthStart = (tFrom.replace(day=28) + timedelta(days=5)).replace(day=1)

    fDataStart = None

    bContinue = True
    tTar = tFrom
    while bContinue:
        try:
            tCut = tTar + timedelta(days=7)
            if tCut > tNextMonthStart:
                tCut = tNextMonthStart
                bContinue = False

            strTar = tTar.strftime("%Y-%m-%d %H:%M:%S")
            strCut = tCut.strftime("%Y-%m-%d %H:%M:%S")

            allData = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTar, strCut, "m1")

            if not allData:
                continue

            if allData.get('map') is None:
                continue

            dataPoint = allData['map'].get(strPointName)

            if dataPoint is None or not isinstance(dataPoint, list):
                continue

            if not len(dataPoint):
                continue

            fDataStart = float(dataPoint[0])
            break

        except Exception as e:
            logging.error("ERROR in get_thisyear_delta_from_month_to_now(pointName: {point}; month: {month}; "
                          "err: {err}".format(point=strPointName, month=nMonth, err=e.__str__()))
        finally:
            tTar += timedelta(days=7)

    if fDataStart is None:
        logging.error("get_thisyear_delta_from_month_to_now: no data found in {month} for {point}".format(month=nMonth, point=strPointName))
        return ""

    fDataEnd = None
    try:
        strCurrValue = BEOPDataAccess.getInstance().getInputTable([strPointName])[1].get(strPointName)
        fDataEnd = float(strCurrValue)
    except Exception as e:
        logging.error("ERROR in get_thisyear_delta_from_month_to_now: fail to get current "
                      "value of {point} (err: {err})".format(point=strPointName, err=e.__str__()))

    if fDataEnd is None:
        logging.error("get_thisyear_delta_from_month_to_now: no current data found for {point}".format(point=strPointName))
        return ""

    fDelta = round(fDataEnd - fDataStart, nDecimal)

    return fDelta


def get_today_data_begin(strPointName, nDecimal=2, actTime=None):
    tNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

    tTimeFrom = tNow.replace(hour=0, minute=0, second=0, microsecond=0)
    if RedisManager.is_alive():
        reValue = RedisManager.get_history_data(strPointName, tTimeFrom)
        if reValue != None and reValue != "" and is_digit(reValue):
            return round(float(reValue), nDecimal)

    strTime = tTimeFrom.strftime("%Y-%m-%d %H:%M:%S")

    myValue = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTime, strTime, "m1")
    if not myValue:
        return ""
    if myValue.get("map", None) == None:
        return ""
    if myValue.get("map").get(strPointName, None) == None:
        return ""
    dataList = myValue.get("map").get(strPointName)
    if not isinstance(dataList, list):
        return ""
    if not len(dataList):
        return ""
    try:
        data = round(float(dataList[-1]), nDecimal)
        return data
    except:
        return ""


def get_this_year_data_delta(strPointName, nDecimal=2, actTime=None):
    tNow = datetime.now()
    if actTime:
        tNow = datetime.strptime(actTime, '%Y-%m-%d %H:%M:%S')

    try:
        tFrom = tNow.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
        strFrom = tFrom.strftime("%Y-%m-%d %H:%M:%S")
        strTo = tNow.strftime("%Y-%m-%d %H:%M:%S")

        if RedisManager.is_alive():
            dataFrom = RedisManager.get_history_data(strPointName, tFrom)
            dataTo = RedisManager.get_history_data(strPointName, tNow)
            if dataFrom is not None and dataTo is not None:
                fDelta = round(dataTo - dataFrom, nDecimal)
                return fDelta

        dataFrom = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strFrom, strFrom, "m1")
        dataTo = BEOPDataAccess.getInstance().get_history_data_padded(strPointName, strTo, strTo, "m1")

        if not dataFrom.get("map", {}).get(strPointName, []) or not dataTo.get("map", {}).get(strPointName, []):
            return ""

        dataFromList = dataFrom.get("map", {}).get(strPointName, [])
        dataToList = dataTo.get("map", {}).get(strPointName, [])

        fDelta = round(dataToList[-1] - dataFromList[-1], 2)
        return fDelta
    except:
        return ""
